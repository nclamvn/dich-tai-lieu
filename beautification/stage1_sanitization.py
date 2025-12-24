#!/usr/bin/env python3
"""
Stage 1: Document Sanitization
Loại bỏ ký tự rác, watermark và chuẩn hóa văn bản
"""

import re
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH


def remove_garbage_characters(text):
    """
    Loại bỏ các ký tự rác và chuỗi ký tự đặc biệt lặp lại
    """
    # Phase 4.4 FIX: Remove SINGLE stray exclamation marks on their own lines
    text = re.sub(r'^\s*!\s*$', '', text, flags=re.MULTILINE)

    # Phase 4.4 FIX: Remove standalone page numbers (digits only on a line)
    text = re.sub(r'^\s*\d+\s*$', '', text, flags=re.MULTILINE)

    # Phase 4.4 FIX: Remove HTML entities like &#123;
    text = re.sub(r'&#\d+;', '', text)

    # Loại bỏ chuỗi ký tự đặc biệt lặp lại từ 3 lần trở lên
    text = re.sub(r'[!@#$%^&*()_+=|~`\[\]{}:;"<>,.?/\-]{3,}', '', text)

    # Phase 4.4 FIX: Remove encoded/corrupted strings (long sequences of special chars mixed with quotes)
    text = re.sub(r'["#$%&\'()*+,\-./]{10,}', '', text)

    # Loại bỏ các ký tự điều khiển (control characters) trừ newline và tab
    text = re.sub(r'[\x00-\x08\x0B-\x0C\x0E-\x1F\x7F]', '', text)

    return text


def remove_watermarks(text):
    """
    Loại bỏ watermark và artifacts từ công cụ dịch
    """
    # Loại bỏ "AI Translator Pro | Page X"
    text = re.sub(r'AI Translator Pro \| Page \d+', '', text)
    
    # Loại bỏ "--- PAGE X ---"
    text = re.sub(r'--- PAGE \d+ ---', '', text)
    
    # Loại bỏ "Translate TheLittlePrince.pdf"
    text = re.sub(r'Translate .+?\.pdf', '', text)
    
    return text


def normalize_whitespace(text):
    """
    Chuẩn hóa khoảng trắng và dấu câu
    """
    # Thay thế nhiều khoảng trắng liên tiếp bằng một khoảng trắng
    text = re.sub(r' {2,}', ' ', text)
    
    # Thay thế nhiều hơn 2 newlines liên tiếp bằng 2 newlines (1 dòng trống)
    text = re.sub(r'\n{3,}', '\n\n', text)
    
    # Đảm bảo có khoảng trắng sau dấu câu
    text = re.sub(r'([.,;:!?])([^\s\n])', r'\1 \2', text)
    
    # Loại bỏ khoảng trắng ở đầu và cuối mỗi dòng
    lines = [line.strip() for line in text.split('\n')]
    text = '\n'.join(lines)
    
    return text


def clean_vietnamese_text(text):
    """
    Làm sạch văn bản tiếng Việt, giữ lại các ký tự hợp lệ
    """
    # Giữ lại: chữ cái Latin, số, dấu câu cơ bản, khoảng trắng, và ký tự tiếng Việt (Unicode U+00C0-U+1EF9)
    # Loại bỏ tất cả các ký tự khác
    
    # Pattern cho ký tự hợp lệ
    valid_pattern = r'[a-zA-Z0-9\s\u00C0-\u1EF9.,;:?!()\[\]{}"\'—–\-\n]'
    
    # Tạo danh sách các ký tự hợp lệ
    cleaned_chars = []
    for char in text:
        if re.match(valid_pattern, char):
            cleaned_chars.append(char)
        elif char in ['\r', '\t']:
            # Chuyển đổi \r và \t thành khoảng trắng
            cleaned_chars.append(' ')
    
    return ''.join(cleaned_chars)


def sanitize_docx(input_path, output_path):
    """
    Làm sạch toàn bộ file DOCX
    """
    print(f"Đang xử lý file: {input_path}")
    
    # Đọc document
    doc = Document(input_path)
    
    # Đếm số đoạn văn ban đầu
    original_para_count = len(doc.paragraphs)
    print(f"Số đoạn văn ban đầu: {original_para_count}")
    
    # Xử lý từng đoạn văn
    cleaned_count = 0
    removed_count = 0
    
    for para in doc.paragraphs:
        original_text = para.text
        
        # Áp dụng các bước làm sạch
        cleaned_text = remove_garbage_characters(original_text)
        cleaned_text = remove_watermarks(cleaned_text)
        cleaned_text = clean_vietnamese_text(cleaned_text)
        cleaned_text = normalize_whitespace(cleaned_text)
        
        # Cập nhật lại đoạn văn
        if cleaned_text.strip():
            para.text = cleaned_text
            cleaned_count += 1
        else:
            # Nếu đoạn văn trống sau khi làm sạch, xóa nó
            para.text = ""
            removed_count += 1
    
    print(f"Đã làm sạch: {cleaned_count} đoạn văn")
    print(f"Đã loại bỏ: {removed_count} đoạn văn rỗng")
    
    # Lưu file
    doc.save(output_path)
    print(f"Đã lưu file sạch: {output_path}")


def main():
    """
    Hàm chính để chạy sanitization
    """
    import sys
    
    if len(sys.argv) < 3:
        print("Cách sử dụng: python stage1_sanitization.py <input.docx> <output.docx>")
        sys.exit(1)
    
    input_file = sys.argv[1]
    output_file = sys.argv[2]
    
    try:
        sanitize_docx(input_file, output_file)
        print("\n✅ Hoàn thành Giai đoạn 1: Sanitization")
    except Exception as e:
        print(f"\n❌ Lỗi: {e}")
        sys.exit(1)


if __name__ == "__main__":
    main()
