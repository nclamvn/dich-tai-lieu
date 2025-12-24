#!/usr/bin/env python3
"""
Stage 3: Document Polishing
Tinh chỉnh chi tiết cuối cùng để đạt chất lượng xuất bản
"""

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


def fix_widows_and_orphans(doc):
    """
    Bật tính năng chống widows và orphans cho tất cả đoạn văn
    """
    count = 0
    for para in doc.paragraphs:
        if para.text.strip():
            # Bật widow/orphan control
            para.paragraph_format.widow_control = True
            count += 1
    
    print(f"✓ Đã bật widow/orphan control cho {count} đoạn văn")


def ensure_chapter_starts_on_new_page(doc):
    """
    Đảm bảo các chương bắt đầu ở trang mới
    """
    chapter_count = 0
    
    for para in doc.paragraphs:
        if para.style.name == 'Heading 1':
            para.paragraph_format.page_break_before = True
            chapter_count += 1
    
    print(f"✓ Đã đảm bảo {chapter_count} chương bắt đầu ở trang mới")


def add_table_of_contents(doc):
    """
    Thêm Table of Contents (Mục lục) vào đầu document
    """
    # Tìm vị trí sau trang bìa (sau Heading 1 đầu tiên hoặc Title)
    insert_position = 0
    
    for i, para in enumerate(doc.paragraphs):
        if para.style.name in ['Title', 'Heading 1']:
            insert_position = i + 1
            break
    
    # Chèn một đoạn mới cho TOC
    if insert_position > 0:
        toc_para = doc.paragraphs[insert_position].insert_paragraph_before()
        toc_para.text = "MỤC LỤC"
        toc_para.style = doc.styles['Heading 1']
        
        # Thêm một đoạn trống cho TOC content
        toc_content = doc.paragraphs[insert_position + 1].insert_paragraph_before()
        toc_content.text = "[Mục lục sẽ được tự động tạo khi mở file trong Word và chọn References > Update Table]"
        
        print("✓ Đã thêm placeholder cho Table of Contents")
    else:
        print("⚠ Không tìm thấy vị trí phù hợp để chèn TOC")


def verify_font_consistency(doc):
    """
    Kiểm tra và báo cáo về tính nhất quán của font
    """
    fonts_used = {}
    
    for para in doc.paragraphs:
        if para.text.strip():
            style_name = para.style.name
            for run in para.runs:
                font_name = run.font.name
                if font_name:
                    key = f"{style_name}:{font_name}"
                    fonts_used[key] = fonts_used.get(key, 0) + 1
    
    print("✓ Font usage report:")
    for key, count in sorted(fonts_used.items(), key=lambda x: x[1], reverse=True)[:10]:
        print(f"  - {key}: {count} lần")


def add_document_properties(doc, title="", author="", subject=""):
    """
    Thêm metadata cho document
    """
    core_properties = doc.core_properties
    
    if title:
        core_properties.title = title
    if author:
        core_properties.author = author
    if subject:
        core_properties.subject = subject
    
    print(f"✓ Đã cập nhật document properties")


def optimize_for_pdf_export(doc):
    """
    Tối ưu hóa document để xuất PDF chất lượng cao
    """
    # Đảm bảo tất cả fonts được embed
    # (Lưu ý: python-docx không hỗ trợ trực tiếp, cần làm trong Word)
    
    # Đảm bảo margins phù hợp cho in ấn
    for section in doc.sections:
        # Kiểm tra và điều chỉnh margins nếu cần
        if section.left_margin < Inches(0.75):
            section.left_margin = Inches(0.9)
        if section.right_margin < Inches(0.5):
            section.right_margin = Inches(0.6)
    
    print("✓ Đã tối ưu hóa cho PDF export")


def final_quality_check(doc):
    """
    Kiểm tra chất lượng cuối cùng
    """
    issues = []
    
    # Kiểm tra đoạn văn rỗng
    empty_paras = sum(1 for para in doc.paragraphs if not para.text.strip())
    if empty_paras > 20:
        issues.append(f"Có {empty_paras} đoạn văn rỗng (có thể cần dọn dẹp)")
    
    # Kiểm tra heading structure
    heading_count = sum(1 for para in doc.paragraphs if 'Heading' in para.style.name)
    if heading_count == 0:
        issues.append("Không có heading nào (cần thêm cấu trúc)")
    
    # Kiểm tra body text
    body_count = sum(1 for para in doc.paragraphs if para.style.name in ['Body Text', 'Normal'])
    if body_count == 0:
        issues.append("Không có nội dung body text")
    
    if issues:
        print("\n⚠ Các vấn đề cần lưu ý:")
        for issue in issues:
            print(f"  - {issue}")
    else:
        print("\n✅ Không phát hiện vấn đề nghiêm trọng")
    
    print(f"\nThống kê:")
    print(f"  - Tổng đoạn văn: {len(doc.paragraphs)}")
    print(f"  - Headings: {heading_count}")
    print(f"  - Body paragraphs: {body_count}")
    print(f"  - Empty paragraphs: {empty_paras}")


def polish_docx(input_path, output_path, title="", author=""):
    """
    Tinh chỉnh và hoàn thiện file DOCX
    """
    print(f"Đang xử lý file: {input_path}")
    
    # Đọc document
    doc = Document(input_path)
    
    # Bước 1: Fix widows and orphans
    fix_widows_and_orphans(doc)
    
    # Bước 2: Đảm bảo chương bắt đầu ở trang mới
    ensure_chapter_starts_on_new_page(doc)
    
    # Bước 3: Thêm Table of Contents
    add_table_of_contents(doc)
    
    # Bước 4: Kiểm tra font consistency
    verify_font_consistency(doc)
    
    # Bước 5: Thêm document properties
    add_document_properties(doc, title=title, author=author)
    
    # Bước 6: Tối ưu cho PDF export
    optimize_for_pdf_export(doc)
    
    # Bước 7: Quality check cuối cùng
    final_quality_check(doc)
    
    # Lưu file
    doc.save(output_path)
    print(f"\nĐã lưu file polished: {output_path}")


def main():
    """
    Hàm chính để chạy polishing
    """
    import sys
    
    if len(sys.argv) < 3:
        print("Cách sử dụng: python stage3_polishing.py <input.docx> <output.docx> [title] [author]")
        sys.exit(1)
    
    input_file = sys.argv[1]
    output_file = sys.argv[2]
    title = sys.argv[3] if len(sys.argv) > 3 else ""
    author = sys.argv[4] if len(sys.argv) > 4 else ""
    
    try:
        polish_docx(input_file, output_file, title=title, author=author)
        print("\n✅ Hoàn thành Giai đoạn 3: Polishing")
    except Exception as e:
        print(f"\n❌ Lỗi: {e}")
        import traceback
        traceback.print_exc()
        sys.exit(1)


if __name__ == "__main__":
    main()
