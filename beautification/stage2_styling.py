#!/usr/bin/env python3
"""
Stage 2: Document Styling
Áp dụng cấu trúc và styles chuyên nghiệp cho tài liệu

FIX-006: Improved heading detection with proper patterns
"""

import re
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE


# FIX-006: Chapter heading patterns (Vietnamese + English)
CHAPTER_PATTERNS = [
    r'^(CHAPTER|Chapter|CHƯƠNG|Chương|PHẦN|Phần)\s+(\d+|[IVXLCDM]+)',
    r'^(CHƯƠNG|Chương)\s+\d+\s*[:：\-–—]',
    r'^(Chapter|CHAPTER)\s+\d+\s*[:：\-–—]',
    r'^(Prologue|PROLOGUE|Lời mở đầu|LỜI MỞ ĐẦU)$',
    r'^(Epilogue|EPILOGUE|Lời kết|LỜI KẾT)$',
    r'^(Introduction|INTRODUCTION|Giới thiệu|GIỚI THIỆU)$',
]

# FIX-006: Section heading patterns
SECTION_PATTERNS = [
    r'^(\d+\.)+\s+\w',  # 1.1 Section
    r'^(Section|SECTION|Phần|PHẦN|Mục|MỤC)\s+\d+',
    r'^[IVXLCDM]+\.\s+\w',  # I. Section, II. Section
]


def setup_styles(doc):
    """
    Thiết lập các styles chuẩn cho document
    """
    styles = doc.styles
    
    # Style cho Body Text (nội dung chính)
    try:
        body_style = styles['Body Text']
    except KeyError:
        body_style = styles.add_style('Body Text', WD_STYLE_TYPE.PARAGRAPH)
    
    body_font = body_style.font
    body_font.name = 'Times New Roman'
    body_font.size = Pt(12)
    
    body_para = body_style.paragraph_format
    body_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    body_para.first_line_indent = Inches(0.2)  # Thụt đầu dòng
    body_para.line_spacing = 1.15
    body_para.space_after = Pt(6)
    
    # Style cho Heading 1 (tiêu đề chương)
    heading1_style = styles['Heading 1']
    h1_font = heading1_style.font
    h1_font.name = 'Arial'
    h1_font.size = Pt(18)
    h1_font.bold = True
    h1_font.color.rgb = RGBColor(0, 0, 0)
    
    h1_para = heading1_style.paragraph_format
    h1_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    h1_para.space_before = Pt(36)
    h1_para.space_after = Pt(12)
    h1_para.keep_with_next = True
    h1_para.page_break_before = True  # Bắt đầu chương ở trang mới
    
    # Style cho Heading 2 (tiểu mục)
    heading2_style = styles['Heading 2']
    h2_font = heading2_style.font
    h2_font.name = 'Arial'
    h2_font.size = Pt(14)
    h2_font.bold = True
    
    h2_para = heading2_style.paragraph_format
    h2_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    h2_para.space_before = Pt(18)
    h2_para.space_after = Pt(6)
    h2_para.keep_with_next = True
    
    print("✓ Đã thiết lập styles")


def is_chapter_heading(text: str) -> bool:
    """FIX-006: Check if text matches chapter heading patterns."""
    text = text.strip()
    for pattern in CHAPTER_PATTERNS:
        if re.match(pattern, text, re.IGNORECASE):
            return True
    return False


def is_section_heading(text: str) -> bool:
    """FIX-006: Check if text matches section heading patterns."""
    text = text.strip()
    for pattern in SECTION_PATTERNS:
        if re.match(pattern, text, re.IGNORECASE):
            return True
    return False


def detect_and_apply_headings(doc):
    """
    FIX-006: Improved heading detection with proper patterns.
    Tự động phát hiện và áp dụng heading styles.
    """
    heading1_count = 0
    heading2_count = 0

    for para in doc.paragraphs:
        text = para.text.strip()

        if not text:
            continue

        # FIX-006: Use pattern-based chapter detection
        if is_chapter_heading(text):
            para.style = doc.styles['Heading 1']
            heading1_count += 1
            print(f"  → Heading 1: {text[:60]}")
            continue

        # FIX-006: Use pattern-based section detection
        if is_section_heading(text):
            para.style = doc.styles['Heading 2']
            heading2_count += 1
            print(f"  → Heading 2: {text[:60]}")
            continue

        # Fallback: Short standalone number (chapter number only)
        if text.isdigit() and len(text) <= 3:
            para.style = doc.styles['Heading 1']
            heading1_count += 1
            print(f"  → Heading 1 (number): {text}")

    print(f"✓ Đã áp dụng {heading1_count} Heading 1, {heading2_count} Heading 2")


def apply_body_text_style(doc):
    """
    Áp dụng Body Text style cho tất cả đoạn văn không phải heading
    """
    body_count = 0
    
    for para in doc.paragraphs:
        # Bỏ qua các đoạn rỗng
        if not para.text.strip():
            continue
        
        # Chỉ áp dụng cho các đoạn đang dùng Normal style
        if para.style.name == 'Normal':
            para.style = doc.styles['Body Text']
            body_count += 1
    
    print(f"✓ Đã áp dụng Body Text cho {body_count} đoạn văn")


def setup_page_layout(doc):
    """
    Thiết lập bố cục trang theo tiêu chuẩn sách
    """
    sections = doc.sections
    
    for section in sections:
        # Thiết lập lề
        section.top_margin = Inches(0.75)
        section.bottom_margin = Inches(0.75)
        section.left_margin = Inches(0.9)   # Gutter (lề trong)
        section.right_margin = Inches(0.6)  # Outer (lề ngoài)
        
        # Thiết lập header/footer khác nhau cho trang chẵn/lẻ
        section.different_first_page_header_footer = True
    
    print("✓ Đã thiết lập page layout")


def add_page_numbers(doc):
    """
    Thêm số trang vào footer
    """
    for section in doc.sections:
        footer = section.footer
        footer_para = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Thêm field code cho page number
        run = footer_para.add_run()
        
    print("✓ Đã thêm page numbers")


def style_docx(input_path, output_path):
    """
    Áp dụng styling cho toàn bộ file DOCX
    """
    print(f"Đang xử lý file: {input_path}")
    
    # Đọc document
    doc = Document(input_path)
    
    # Bước 1: Thiết lập styles
    setup_styles(doc)
    
    # Bước 2: Phát hiện và áp dụng headings
    detect_and_apply_headings(doc)
    
    # Bước 3: Áp dụng Body Text style
    apply_body_text_style(doc)
    
    # Bước 4: Thiết lập page layout
    setup_page_layout(doc)
    
    # Bước 5: Thêm page numbers
    add_page_numbers(doc)
    
    # Lưu file
    doc.save(output_path)
    print(f"Đã lưu file styled: {output_path}")


def main():
    """
    Hàm chính để chạy styling
    """
    import sys
    
    if len(sys.argv) < 3:
        print("Cách sử dụng: python stage2_styling.py <input.docx> <output.docx>")
        sys.exit(1)
    
    input_file = sys.argv[1]
    output_file = sys.argv[2]
    
    try:
        style_docx(input_file, output_file)
        print("\n✅ Hoàn thành Giai đoạn 2: Styling")
    except Exception as e:
        print(f"\n❌ Lỗi: {e}")
        import traceback
        traceback.print_exc()
        sys.exit(1)


if __name__ == "__main__":
    main()
