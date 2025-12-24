"""
LLM-Native Document Renderer
AI Publisher Pro

Renders structured document content to DOCX/PDF
while preserving tables, headers, and layout.

Philosophy: Simple rendering, LLM already did the hard work.
"""

from pathlib import Path
from typing import Optional, List, Dict, Any
from dataclasses import dataclass
import json

from .document_analyzer import (
    StructuredDocument,
    DocumentPage,
    ContentBlock,
    ContentType,
    Table,
    TableCell
)


class DocumentRenderer:
    """
    Renders StructuredDocument to various output formats.

    Supports:
    - DOCX (Microsoft Word)
    - PDF (via reportlab or weasyprint)
    - Markdown
    - HTML
    """

    def __init__(self):
        pass

    def render_docx(
        self,
        document: StructuredDocument,
        output_path: str,
        title: Optional[str] = None
    ) -> str:
        """
        Render document to DOCX format.

        Args:
            document: StructuredDocument to render
            output_path: Path for output file
            title: Optional document title

        Returns:
            Path to created file
        """
        from docx import Document
        from docx.shared import Inches, Pt, Cm
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.enum.table import WD_TABLE_ALIGNMENT
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement

        doc = Document()

        # Set default font for Vietnamese/Chinese support
        style = doc.styles['Normal']
        style.font.name = 'Arial'
        style.font.size = Pt(11)

        # Add title if provided
        if title:
            heading = doc.add_heading(title, 0)
            heading.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Process each page
        for page in document.pages:
            for block in page.blocks:
                self._render_block_docx(doc, block)

            # Add page break between pages (except last)
            if page != document.pages[-1]:
                doc.add_page_break()

        # Save document
        output_path = Path(output_path)
        output_path.parent.mkdir(parents=True, exist_ok=True)
        doc.save(str(output_path))

        return str(output_path)

    def _render_block_docx(self, doc, block: ContentBlock):
        """Render a single content block to DOCX"""
        from docx.shared import Pt
        from docx.enum.text import WD_ALIGN_PARAGRAPH

        if block.type == ContentType.HEADER:
            level = min(block.level, 9) if block.level > 0 else 1
            doc.add_heading(block.content, level)

        elif block.type == ContentType.PARAGRAPH:
            para = doc.add_paragraph(block.content)
            if block.style:
                if block.style.get("bold"):
                    para.runs[0].bold = True
                if block.style.get("italic"):
                    para.runs[0].italic = True
                if block.style.get("alignment") == "center":
                    para.alignment = WD_ALIGN_PARAGRAPH.CENTER

        elif block.type == ContentType.TABLE:
            self._render_table_docx(doc, block.content)

        elif block.type == ContentType.LIST:
            for item in block.content:
                doc.add_paragraph(item, style='List Bullet')

        elif block.type == ContentType.FOOTER:
            para = doc.add_paragraph(block.content)
            para.runs[0].font.size = Pt(9)
            para.runs[0].font.italic = True

    def _render_table_docx(self, doc, table_data):
        """Render a table to DOCX"""
        from docx.shared import Pt, Cm
        from docx.enum.table import WD_TABLE_ALIGNMENT
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement

        # Handle both Table object and dict
        if isinstance(table_data, Table):
            num_rows = table_data.num_rows
            num_cols = table_data.num_cols
            cells = table_data.cells
            caption = table_data.caption
        else:
            num_rows = table_data.get("num_rows", 0)
            num_cols = table_data.get("num_cols", 0)
            cells = [TableCell(**c) if isinstance(c, dict) else c for c in table_data.get("cells", [])]
            caption = table_data.get("caption")

        if num_rows == 0 or num_cols == 0:
            return

        # Add caption if present
        if caption:
            cap_para = doc.add_paragraph(caption)
            cap_para.runs[0].bold = True
            cap_para.runs[0].font.size = Pt(10)

        # Create table
        table = doc.add_table(rows=num_rows, cols=num_cols)
        table.style = 'Table Grid'
        table.alignment = WD_TABLE_ALIGNMENT.CENTER

        # Fill cells
        for cell in cells:
            if isinstance(cell, dict):
                row, col = cell.get("row", 0), cell.get("col", 0)
                content = cell.get("content", "")
                is_header = cell.get("is_header", False)
            else:
                row, col = cell.row, cell.col
                content = cell.content
                is_header = cell.is_header

            if row < num_rows and col < num_cols:
                table_cell = table.rows[row].cells[col]
                table_cell.text = str(content)

                # Style header cells
                if is_header:
                    for paragraph in table_cell.paragraphs:
                        for run in paragraph.runs:
                            run.bold = True
                    # Add shading
                    shading = OxmlElement('w:shd')
                    shading.set(qn('w:fill'), 'E0E0E0')
                    table_cell._tc.get_or_add_tcPr().append(shading)

        # Add spacing after table
        doc.add_paragraph()

    def render_markdown(
        self,
        document: StructuredDocument,
        output_path: str,
        title: Optional[str] = None
    ) -> str:
        """
        Render document to Markdown format.
        """
        lines = []

        if title:
            lines.append(f"# {title}\n")

        for page in document.pages:
            lines.append(f"\n---\n*Page {page.page_number}*\n")

            for block in page.blocks:
                lines.append(self._render_block_markdown(block))

        content = "\n".join(lines)

        output_path = Path(output_path)
        output_path.parent.mkdir(parents=True, exist_ok=True)
        output_path.write_text(content, encoding="utf-8")

        return str(output_path)

    def _render_block_markdown(self, block: ContentBlock) -> str:
        """Render a single content block to Markdown"""
        if block.type == ContentType.HEADER:
            level = min(block.level, 6) if block.level > 0 else 1
            return f"{'#' * level} {block.content}\n"

        elif block.type == ContentType.PARAGRAPH:
            return f"{block.content}\n"

        elif block.type == ContentType.TABLE:
            return self._render_table_markdown(block.content)

        elif block.type == ContentType.LIST:
            return "\n".join(f"- {item}" for item in block.content) + "\n"

        elif block.type == ContentType.FOOTER:
            return f"*{block.content}*\n"

        return ""

    def _render_table_markdown(self, table_data) -> str:
        """Render a table to Markdown"""
        # Handle both Table object and dict
        if isinstance(table_data, Table):
            num_rows = table_data.num_rows
            num_cols = table_data.num_cols
            cells = table_data.cells
            caption = table_data.caption
        else:
            num_rows = table_data.get("num_rows", 0)
            num_cols = table_data.get("num_cols", 0)
            cells = table_data.get("cells", [])
            caption = table_data.get("caption")

        if num_rows == 0 or num_cols == 0:
            return ""

        # Create grid
        grid = [["" for _ in range(num_cols)] for _ in range(num_rows)]
        header_row = 0

        for cell in cells:
            if isinstance(cell, dict):
                row, col = cell.get("row", 0), cell.get("col", 0)
                content = cell.get("content", "")
                is_header = cell.get("is_header", False)
            else:
                row, col = cell.row, cell.col
                content = cell.content
                is_header = cell.is_header

            if row < num_rows and col < num_cols:
                grid[row][col] = str(content)
                if is_header and row == 0:
                    header_row = 1

        # Build markdown table
        lines = []

        if caption:
            lines.append(f"**{caption}**\n")

        # Header row
        lines.append("| " + " | ".join(grid[0]) + " |")
        lines.append("| " + " | ".join(["---"] * num_cols) + " |")

        # Data rows
        for row in grid[1:]:
            lines.append("| " + " | ".join(row) + " |")

        lines.append("")
        return "\n".join(lines)

    def render_html(
        self,
        document: StructuredDocument,
        output_path: str,
        title: Optional[str] = None
    ) -> str:
        """
        Render document to HTML format.
        """
        html_parts = [
            "<!DOCTYPE html>",
            "<html>",
            "<head>",
            f"<title>{title or 'Document'}</title>",
            "<meta charset='utf-8'>",
            "<style>",
            "body { font-family: Arial, sans-serif; max-width: 800px; margin: 0 auto; padding: 20px; }",
            "table { border-collapse: collapse; width: 100%; margin: 20px 0; }",
            "th, td { border: 1px solid #ddd; padding: 8px; text-align: left; }",
            "th { background-color: #f0f0f0; font-weight: bold; }",
            ".page-break { page-break-after: always; border-top: 2px dashed #ccc; margin: 30px 0; }",
            ".caption { font-weight: bold; margin-bottom: 5px; }",
            "</style>",
            "</head>",
            "<body>",
        ]

        if title:
            html_parts.append(f"<h1>{title}</h1>")

        for i, page in enumerate(document.pages):
            for block in page.blocks:
                html_parts.append(self._render_block_html(block))

            if i < len(document.pages) - 1:
                html_parts.append("<div class='page-break'></div>")

        html_parts.extend(["</body>", "</html>"])

        content = "\n".join(html_parts)

        output_path = Path(output_path)
        output_path.parent.mkdir(parents=True, exist_ok=True)
        output_path.write_text(content, encoding="utf-8")

        return str(output_path)

    def _render_block_html(self, block: ContentBlock) -> str:
        """Render a single content block to HTML"""
        if block.type == ContentType.HEADER:
            level = min(block.level, 6) if block.level > 0 else 1
            return f"<h{level}>{block.content}</h{level}>"

        elif block.type == ContentType.PARAGRAPH:
            style = ""
            if block.style:
                parts = []
                if block.style.get("bold"):
                    parts.append("font-weight: bold")
                if block.style.get("italic"):
                    parts.append("font-style: italic")
                if parts:
                    style = f" style='{'; '.join(parts)}'"
            return f"<p{style}>{block.content}</p>"

        elif block.type == ContentType.TABLE:
            return self._render_table_html(block.content)

        elif block.type == ContentType.LIST:
            items = "\n".join(f"<li>{item}</li>" for item in block.content)
            return f"<ul>\n{items}\n</ul>"

        elif block.type == ContentType.FOOTER:
            return f"<footer><em>{block.content}</em></footer>"

        return ""

    def _render_table_html(self, table_data) -> str:
        """Render a table to HTML"""
        # Handle both Table object and dict
        if isinstance(table_data, Table):
            num_rows = table_data.num_rows
            num_cols = table_data.num_cols
            cells = table_data.cells
            caption = table_data.caption
        else:
            num_rows = table_data.get("num_rows", 0)
            num_cols = table_data.get("num_cols", 0)
            cells = table_data.get("cells", [])
            caption = table_data.get("caption")

        if num_rows == 0 or num_cols == 0:
            return ""

        # Create grid
        grid = [[{"content": "", "is_header": False} for _ in range(num_cols)] for _ in range(num_rows)]

        for cell in cells:
            if isinstance(cell, dict):
                row, col = cell.get("row", 0), cell.get("col", 0)
                content = cell.get("content", "")
                is_header = cell.get("is_header", False)
            else:
                row, col = cell.row, cell.col
                content = cell.content
                is_header = cell.is_header

            if row < num_rows and col < num_cols:
                grid[row][col] = {"content": str(content), "is_header": is_header}

        # Build HTML table
        parts = ["<table>"]

        if caption:
            parts.append(f"<caption class='caption'>{caption}</caption>")

        for row_idx, row in enumerate(grid):
            parts.append("<tr>")
            for cell in row:
                tag = "th" if cell["is_header"] else "td"
                parts.append(f"<{tag}>{cell['content']}</{tag}>")
            parts.append("</tr>")

        parts.append("</table>")
        return "\n".join(parts)


# =========================================
# Convenience Functions
# =========================================

def render_to_docx(
    document: StructuredDocument,
    output_path: str,
    title: Optional[str] = None
) -> str:
    """Quick render to DOCX"""
    renderer = DocumentRenderer()
    return renderer.render_docx(document, output_path, title)


def render_to_markdown(
    document: StructuredDocument,
    output_path: str,
    title: Optional[str] = None
) -> str:
    """Quick render to Markdown"""
    renderer = DocumentRenderer()
    return renderer.render_markdown(document, output_path, title)


def render_to_html(
    document: StructuredDocument,
    output_path: str,
    title: Optional[str] = None
) -> str:
    """Quick render to HTML"""
    renderer = DocumentRenderer()
    return renderer.render_html(document, output_path, title)
