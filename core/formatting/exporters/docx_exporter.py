#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
DOCX Style Exporter - Export StyledDocument to Microsoft Word format.

Uses python-docx library for DOCX generation.
Supports:
- Heading styles (H1-H4) with proper outline levels for TOC
- Body text with typography
- Lists (bullet and numbered)
- Tables
- Code blocks
- Block quotes
- Table of Contents placeholder
"""

import os
from pathlib import Path
from typing import Optional, List

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn, nsmap
from docx.oxml import OxmlElement

from ..style_engine import (
    StyledDocument,
    StyledElement,
    StyledList,
    StyledListItem,
    StyledTable,
    StyledTableCell,
    StyledTableRow,
    StyledCodeBlock,
    StyledBlockquote,
    StyledFigure,
    StyledHorizontalRule,
)
from ..utils.constants import ELEMENT_TYPES, LIST_STYLES
from ..page_layout import PageLayoutManager
from ..toc_generator import TocGenerator, TocElement


class DocxStyleExporter:
    """
    Export StyledDocument to DOCX format.

    Usage:
        exporter = DocxStyleExporter()
        path = exporter.export(styled_doc, "output.docx")

        # With custom page layout
        layout = PageLayoutManager(page_size="A4", margins="book")
        path = exporter.export(styled_doc, "output.docx", page_layout=layout)
    """

    def __init__(self):
        """Initialize exporter."""
        self.doc: Optional[Document] = None
        self.page_layout: Optional[PageLayoutManager] = None
        self.current_chapter: str = ""  # Track current chapter for headers
        self.element_count: int = 0     # Track for page break logic

    def export(
        self,
        styled_doc: StyledDocument,
        output_path: str,
        page_layout: PageLayoutManager = None,
    ) -> str:
        """
        Export StyledDocument to DOCX file.

        Args:
            styled_doc: StyledDocument with formatting
            output_path: Path for output file
            page_layout: Optional PageLayoutManager for custom layout

        Returns:
            Absolute path to saved file
        """
        # Create new document
        self.doc = Document()
        self.page_layout = page_layout
        self.element_count = 0
        self.current_chapter = ""

        # Setup page layout (use provided layout or document defaults)
        self._setup_page_layout(styled_doc)

        # Setup header and footer
        self._setup_header_footer(styled_doc)

        # Create/modify styles
        self._create_styles(styled_doc)

        # Add document title if present
        if styled_doc.title:
            self._add_title(styled_doc.title)

        # Add TOC placeholder
        if styled_doc.include_toc and styled_doc.toc:
            self._add_toc(styled_doc.toc_max_level)

        # Add all elements
        for element in styled_doc.elements:
            self._add_element(element)
            self.element_count += 1

        # Set document metadata
        self._set_metadata(styled_doc)

        # Ensure directory exists
        output_path = Path(output_path)
        output_path.parent.mkdir(parents=True, exist_ok=True)

        # Save document
        self.doc.save(str(output_path))

        return str(output_path.absolute())

    def _setup_page_layout(self, styled_doc: StyledDocument) -> None:
        """Configure page size and margins."""
        section = self.doc.sections[0]

        if self.page_layout:
            # Use PageLayoutManager settings
            section.page_width = Inches(self.page_layout.page_size.width)
            section.page_height = Inches(self.page_layout.page_size.height)
            section.top_margin = Inches(self.page_layout.margins.top)
            section.bottom_margin = Inches(self.page_layout.margins.bottom)
            section.left_margin = Inches(self.page_layout.margins.left)
            section.right_margin = Inches(self.page_layout.margins.right)
        else:
            # Use StyledDocument defaults
            section.page_width = Inches(styled_doc.page_width_inches)
            section.page_height = Inches(styled_doc.page_height_inches)
            section.top_margin = Inches(styled_doc.margin_top_inches)
            section.bottom_margin = Inches(styled_doc.margin_bottom_inches)
            section.left_margin = Inches(styled_doc.margin_left_inches)
            section.right_margin = Inches(styled_doc.margin_right_inches)

    def _setup_header_footer(self, styled_doc: StyledDocument) -> None:
        """Setup header and footer with page numbers."""
        section = self.doc.sections[0]

        # Get configuration
        if self.page_layout:
            hf = self.page_layout.header_footer
            header_text = hf.header_text.replace("{title}", styled_doc.title or "")
            footer_alignment = hf.footer_alignment
            header_alignment = hf.header_alignment
            different_first = hf.different_first_page
        else:
            header_text = styled_doc.title or ""
            footer_alignment = "center"
            header_alignment = "center"
            different_first = False

        # Configure different first page if needed
        section.different_first_page_header_footer = different_first

        # Setup header
        if header_text:
            header = section.header
            header_para = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
            header_para.text = header_text
            header_para.alignment = self._get_alignment(header_alignment)

            # Style header text
            for run in header_para.runs:
                run.font.size = Pt(10)

        # Setup footer with page number
        footer = section.footer
        footer_para = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        footer_para.alignment = self._get_alignment(footer_alignment)

        # Add page number field
        self._add_page_number_field(footer_para)

    def _add_page_number_field(self, paragraph) -> None:
        """Add auto-updating page number field to paragraph."""
        run = paragraph.add_run()

        # Begin field
        fldChar1 = OxmlElement('w:fldChar')
        fldChar1.set(qn('w:fldCharType'), 'begin')
        run._r.append(fldChar1)

        # Field instruction
        run2 = paragraph.add_run()
        instrText = OxmlElement('w:instrText')
        instrText.set(qn('xml:space'), 'preserve')
        instrText.text = "PAGE"
        run2._r.append(instrText)

        # Separate
        run3 = paragraph.add_run()
        fldChar2 = OxmlElement('w:fldChar')
        fldChar2.set(qn('w:fldCharType'), 'separate')
        run3._r.append(fldChar2)

        # Placeholder text
        run4 = paragraph.add_run("1")
        run4.font.size = Pt(10)

        # End field
        run5 = paragraph.add_run()
        fldChar3 = OxmlElement('w:fldChar')
        fldChar3.set(qn('w:fldCharType'), 'end')
        run5._r.append(fldChar3)

    def _get_alignment(self, align_str: str):
        """Convert alignment string to WD_ALIGN_PARAGRAPH enum."""
        alignment_map = {
            "left": WD_ALIGN_PARAGRAPH.LEFT,
            "center": WD_ALIGN_PARAGRAPH.CENTER,
            "right": WD_ALIGN_PARAGRAPH.RIGHT,
            "justify": WD_ALIGN_PARAGRAPH.JUSTIFY,
        }
        return alignment_map.get(align_str, WD_ALIGN_PARAGRAPH.CENTER)

    def _create_styles(self, styled_doc: StyledDocument) -> None:
        """
        Create or modify document styles.

        Uses built-in Heading 1-4 styles for TOC compatibility.
        """
        styles = self.doc.styles

        # Get heading elements to determine what styles we need
        headings = styled_doc.get_headings()
        heading_levels = set(h.level for h in headings if h.level)

        # Modify built-in heading styles
        for level in range(1, 5):
            if level not in heading_levels:
                continue

            style_name = f"Heading {level}"
            if style_name in styles:
                style = styles[style_name]

                # Find a representative element for this level
                sample = next((h for h in headings if h.level == level), None)
                if sample:
                    # Font settings
                    style.font.name = sample.font_name
                    style.font.size = Pt(sample.font_size_pt)
                    style.font.bold = sample.bold
                    style.font.italic = sample.italic

                    # Paragraph format
                    pf = style.paragraph_format
                    pf.space_before = Pt(sample.space_before_pt)
                    pf.space_after = Pt(sample.space_after_pt)
                    pf.keep_with_next = sample.keep_with_next

                    # Page break before H1
                    if level == 1 and sample.page_break_before:
                        pf.page_break_before = True

        # Modify Normal style for body text
        normal_style = styles["Normal"]
        normal_style.font.name = styled_doc.default_font
        normal_style.font.size = Pt(12)
        normal_style.paragraph_format.space_after = Pt(6)
        normal_style.paragraph_format.line_spacing = 1.15

    def _add_title(self, title: str) -> None:
        """Add document title."""
        para = self.doc.add_paragraph()
        run = para.add_run(title)
        run.font.size = Pt(24)
        run.font.bold = True
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        para.paragraph_format.space_after = Pt(24)

    def _add_toc(self, max_level: int = 3) -> None:
        """
        Add Table of Contents placeholder.

        Note: User must update TOC in Word (Ctrl+A, F9 or right-click > Update Field)
        """
        # Add TOC heading
        toc_heading = self.doc.add_paragraph("Table of Contents")
        toc_heading.style = self.doc.styles["Heading 1"]
        # Don't include TOC heading in actual TOC
        toc_heading.paragraph_format.outline_level = None

        # Add TOC field
        paragraph = self.doc.add_paragraph()
        run = paragraph.add_run()
        fldChar1 = OxmlElement('w:fldChar')
        fldChar1.set(qn('w:fldCharType'), 'begin')

        instrText = OxmlElement('w:instrText')
        instrText.set(qn('xml:space'), 'preserve')
        instrText.text = f'TOC \\o "1-{max_level}" \\h \\z \\u'

        fldChar2 = OxmlElement('w:fldChar')
        fldChar2.set(qn('w:fldCharType'), 'separate')

        fldChar3 = OxmlElement('w:fldChar')
        fldChar3.set(qn('w:fldCharType'), 'end')

        run._r.append(fldChar1)
        run._r.append(instrText)
        run._r.append(fldChar2)
        run._r.append(fldChar3)

        # Add note about updating
        note = self.doc.add_paragraph()
        note_run = note.add_run("[Right-click and select 'Update Field' to populate TOC]")
        note_run.font.size = Pt(10)
        note_run.font.italic = True
        note_run.font.color.rgb = RGBColor(128, 128, 128)

        # Page break after TOC
        self.doc.add_page_break()

    def _add_element(self, element: StyledElement) -> None:
        """Add a single styled element to the document."""
        element_type = element.type

        if element_type == ELEMENT_TYPES["HEADING"]:
            self._add_heading(element)
        elif element_type == ELEMENT_TYPES["PARAGRAPH"]:
            self._add_paragraph(element)
        elif element_type == ELEMENT_TYPES["LIST_BULLET"]:
            self._add_bullet_list(element)
        elif element_type == ELEMENT_TYPES["LIST_NUMBERED"]:
            self._add_numbered_list(element)
        elif element_type == ELEMENT_TYPES["TABLE"]:
            self._add_table(element)
        elif element_type == ELEMENT_TYPES["CODE_BLOCK"]:
            self._add_code_block(element)
        elif element_type == ELEMENT_TYPES["QUOTE"]:
            self._add_quote(element)
        elif element_type == ELEMENT_TYPES["IMAGE"]:
            self._add_figure(element)
        elif element_type == ELEMENT_TYPES["HORIZONTAL_RULE"]:
            self._add_horizontal_rule(element)
        else:
            # Default: add as paragraph
            self._add_paragraph(element)

    def _add_heading(self, element: StyledElement) -> None:
        """Add heading with proper style and track chapters."""
        level = element.level or 1

        # Track current chapter (H1) for headers
        if level == 1:
            self.current_chapter = element.content

        # Determine if page break is needed
        should_break = False
        if self.page_layout:
            should_break = self.page_layout.should_page_break_before(
                element.type,
                level=level,
                is_first_content=(self.element_count == 0)
            )
        else:
            # Default: page break before H1 if not first
            should_break = (level == 1 and element.page_break_before and
                           len(self.doc.paragraphs) > 3)

        if should_break:
            self.doc.add_page_break()

        # Use built-in heading style for TOC compatibility
        heading = self.doc.add_heading(element.content, level=level)

        # Apply additional formatting
        self._apply_paragraph_formatting(heading, element)

    def _add_paragraph(self, element: StyledElement) -> None:
        """Add body paragraph with formatting."""
        para = self.doc.add_paragraph()
        run = para.add_run(element.content)

        # Apply formatting
        self._apply_run_formatting(run, element)
        self._apply_paragraph_formatting(para, element)

    def _add_bullet_list(self, element: StyledElement) -> None:
        """Add bullet list with proper nesting support."""
        # Check if we have structured items
        if isinstance(element, StyledList) and element.items:
            for item in element.items:
                self._add_list_item(item, "bullet")
        else:
            # Fallback: parse content as lines
            lines = element.content.split('\n')
            for line in lines:
                line = line.strip()
                if not line:
                    continue

                # Remove bullet markers
                for marker in ['- ', '* ', '+ ', '• ', '○ ', '● ', '▪ ', '▫ ']:
                    if line.startswith(marker):
                        line = line[len(marker):]
                        break

                para = self.doc.add_paragraph(line, style='List Bullet')
                self._apply_paragraph_formatting(para, element)

    def _add_numbered_list(self, element: StyledElement) -> None:
        """Add numbered list with proper nesting support."""
        # Check if we have structured items
        if isinstance(element, StyledList) and element.items:
            for item in element.items:
                self._add_list_item(item, "numbered")
        else:
            # Fallback: parse content as lines
            import re
            lines = element.content.split('\n')
            for line in lines:
                line = line.strip()
                if not line:
                    continue

                # Remove number markers (1. 2. a. b. etc.)
                line = re.sub(r'^[\d\w]+[\.\)]\s*', '', line)
                line = re.sub(r'^\([^)]+\)\s*', '', line)

                para = self.doc.add_paragraph(line, style='List Number')
                self._apply_paragraph_formatting(para, element)

    def _add_list_item(self, item: StyledListItem, list_type: str) -> None:
        """Add a single list item with proper indent level."""
        # Determine style based on level
        level = item.level

        if list_type == "bullet":
            # Use List Bullet style variants
            if level == 0:
                style = 'List Bullet'
            elif level == 1:
                style = 'List Bullet 2' if 'List Bullet 2' in self.doc.styles else 'List Bullet'
            else:
                style = 'List Bullet 3' if 'List Bullet 3' in self.doc.styles else 'List Bullet'
        else:
            # Use List Number style variants
            if level == 0:
                style = 'List Number'
            elif level == 1:
                style = 'List Number 2' if 'List Number 2' in self.doc.styles else 'List Number'
            else:
                style = 'List Number 3' if 'List Number 3' in self.doc.styles else 'List Number'

        para = self.doc.add_paragraph(item.content, style=style)

        # Apply item-specific formatting
        run = para.runs[0] if para.runs else para.add_run()
        run.font.name = item.font_name
        run.font.size = Pt(item.font_size_pt)
        run.font.bold = item.bold
        run.font.italic = item.italic

        # Apply indentation for nesting
        para.paragraph_format.left_indent = Inches(item.left_indent_inches)

    def _add_table(self, element: StyledElement) -> None:
        """Add table with proper styling for headers, cells, and alignment."""
        # Check if we have structured table data
        if isinstance(element, StyledTable):
            self._add_styled_table(element)
        else:
            # Fallback: parse content as markdown table
            self._add_markdown_table(element)

    def _add_styled_table(self, table_elem: StyledTable) -> None:
        """Add table from StyledTable with full formatting."""
        # Calculate dimensions
        num_cols = 0
        if table_elem.header_row:
            num_cols = len(table_elem.header_row.cells)
        elif table_elem.data_rows:
            num_cols = len(table_elem.data_rows[0].cells)

        if num_cols == 0:
            return

        num_rows = len(table_elem.data_rows)
        if table_elem.header_row:
            num_rows += 1

        # Create table
        table = self.doc.add_table(rows=num_rows, cols=num_cols)
        table.style = 'Table Grid'

        row_idx = 0

        # Add header row
        if table_elem.header_row:
            row = table.rows[row_idx]
            for col_idx, cell in enumerate(table_elem.header_row.cells):
                if col_idx < len(row.cells):
                    self._format_table_cell(row.cells[col_idx], cell)
            row_idx += 1

        # Add data rows
        for data_row in table_elem.data_rows:
            row = table.rows[row_idx]
            for col_idx, cell in enumerate(data_row.cells):
                if col_idx < len(row.cells):
                    self._format_table_cell(row.cells[col_idx], cell)
            row_idx += 1

        # Space after table
        self.doc.add_paragraph()

    def _format_table_cell(self, doc_cell, styled_cell: StyledTableCell) -> None:
        """Apply formatting to a table cell."""
        # Set cell text
        doc_cell.text = styled_cell.content

        # Apply formatting to paragraph and run
        for para in doc_cell.paragraphs:
            # Alignment
            alignment_map = {
                "left": WD_ALIGN_PARAGRAPH.LEFT,
                "center": WD_ALIGN_PARAGRAPH.CENTER,
                "right": WD_ALIGN_PARAGRAPH.RIGHT,
            }
            para.alignment = alignment_map.get(styled_cell.alignment, WD_ALIGN_PARAGRAPH.LEFT)

            for run in para.runs:
                run.font.name = styled_cell.font_name
                run.font.size = Pt(styled_cell.font_size_pt)
                run.font.bold = styled_cell.bold

                # Font color
                if styled_cell.font_color:
                    try:
                        r = int(styled_cell.font_color[0:2], 16)
                        g = int(styled_cell.font_color[2:4], 16)
                        b = int(styled_cell.font_color[4:6], 16)
                        run.font.color.rgb = RGBColor(r, g, b)
                    except (ValueError, IndexError):
                        pass

        # Background color (using XML for cell shading)
        if styled_cell.background_color:
            try:
                shading = OxmlElement('w:shd')
                shading.set(qn('w:fill'), styled_cell.background_color)
                doc_cell._tc.get_or_add_tcPr().append(shading)
            except Exception:
                pass  # Ignore shading errors

    def _add_markdown_table(self, element: StyledElement) -> None:
        """Fallback: Add table from markdown content."""
        lines = element.content.strip().split('\n')
        if len(lines) < 2:
            return

        # Parse table
        rows = []
        for line in lines:
            line = line.strip()
            if line.startswith('|') and line.endswith('|'):
                cells = [c.strip() for c in line[1:-1].split('|')]
                # Skip separator rows
                if not all(c.replace('-', '').replace(':', '') == '' for c in cells):
                    rows.append(cells)

        if not rows:
            return

        # Create table
        num_cols = len(rows[0])
        table = self.doc.add_table(rows=len(rows), cols=num_cols)
        table.style = 'Table Grid'

        # Fill table
        for i, row_data in enumerate(rows):
            row = table.rows[i]
            for j, cell_text in enumerate(row_data):
                if j < len(row.cells):
                    row.cells[j].text = cell_text

                    # Bold first row (header)
                    if i == 0:
                        for para in row.cells[j].paragraphs:
                            for run in para.runs:
                                run.font.bold = True

        # Space after table
        self.doc.add_paragraph()

    def _add_code_block(self, element: StyledElement) -> None:
        """Add code block with monospace font and optional background."""
        # Get code content
        if isinstance(element, StyledCodeBlock):
            content = element.code
            language = element.language
        else:
            # Fallback: parse content
            content = element.content
            lines = content.split('\n')
            if lines[0].strip().startswith('```'):
                lines = lines[1:]
            if lines and lines[-1].strip().startswith('```'):
                lines = lines[:-1]
            content = '\n'.join(lines)
            language = ""

        # Use a single-cell table for better background support
        if element.background_color:
            table = self.doc.add_table(rows=1, cols=1)
            table.autofit = True
            cell = table.rows[0].cells[0]
            cell.text = content

            # Apply cell shading
            shading = OxmlElement('w:shd')
            shading.set(qn('w:fill'), element.background_color)
            cell._tc.get_or_add_tcPr().append(shading)

            # Format text in cell
            for para in cell.paragraphs:
                para.alignment = WD_ALIGN_PARAGRAPH.LEFT
                for run in para.runs:
                    run.font.name = element.font_name
                    run.font.size = Pt(element.font_size_pt)

            # Add spacing after table
            spacer = self.doc.add_paragraph()
            spacer.paragraph_format.space_before = Pt(element.space_after_pt)
        else:
            # Simple paragraph approach
            para = self.doc.add_paragraph()
            run = para.add_run(content)
            run.font.name = element.font_name
            run.font.size = Pt(element.font_size_pt)
            para.paragraph_format.space_before = Pt(element.space_before_pt)
            para.paragraph_format.space_after = Pt(element.space_after_pt)
            para.paragraph_format.left_indent = Inches(element.left_indent_inches)

    def _add_quote(self, element: StyledElement) -> None:
        """Add block quote with indentation and optional left border."""
        # Get quote details
        if isinstance(element, StyledBlockquote):
            quote_text = element.quote_text
            attribution = element.attribution
            left_border_color = element.left_border_color
            left_border_width = element.left_border_width_pt
        else:
            quote_text = element.content
            attribution = ""
            left_border_color = "CCCCCC"
            left_border_width = 3.0

        # Add quote paragraph
        para = self.doc.add_paragraph()
        run = para.add_run(quote_text)

        # Quote formatting
        run.font.name = element.font_name
        run.font.size = Pt(element.font_size_pt)
        run.font.italic = element.italic

        # Indentation
        para.paragraph_format.left_indent = Inches(element.left_indent_inches)
        para.paragraph_format.right_indent = Inches(element.right_indent_inches)
        para.paragraph_format.space_before = Pt(element.space_before_pt)
        para.paragraph_format.space_after = Pt(element.space_after_pt)

        # Add left border
        if left_border_width > 0 and left_border_color:
            self._add_paragraph_left_border(para, left_border_color, left_border_width)

        # Add attribution if present
        if attribution:
            attr_para = self.doc.add_paragraph()
            attr_run = attr_para.add_run(f"— {attribution}")
            attr_run.font.name = element.font_name
            attr_run.font.size = Pt(element.font_size_pt - 1)
            attr_run.font.italic = True
            attr_para.paragraph_format.left_indent = Inches(element.left_indent_inches)
            attr_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    def _add_paragraph_left_border(self, paragraph, color: str, width_pt: float) -> None:
        """Add left border to a paragraph using XML manipulation."""
        try:
            pPr = paragraph._p.get_or_add_pPr()
            pBdr = OxmlElement('w:pBdr')
            left = OxmlElement('w:left')
            left.set(qn('w:val'), 'single')
            left.set(qn('w:sz'), str(int(width_pt * 8)))  # eighths of a point
            left.set(qn('w:color'), color)
            left.set(qn('w:space'), '12')  # space between border and text
            pBdr.append(left)
            pPr.append(pBdr)
        except Exception:
            pass  # Ignore border errors

    def _add_figure(self, element: StyledElement) -> None:
        """Add figure with caption."""
        if isinstance(element, StyledFigure):
            figure_number = element.figure_number
            caption = element.caption
            image_url = element.image_url
            alt_text = element.alt_text
        else:
            figure_number = 0
            caption = element.content
            image_url = ""
            alt_text = ""

        # Add image placeholder if no actual image
        if image_url:
            # Check if image file exists and add it
            if os.path.exists(image_url):
                try:
                    para = self.doc.add_paragraph()
                    para.alignment = self._get_alignment(element.alignment)
                    run = para.add_run()
                    run.add_picture(image_url, width=Inches(element.max_width_inches if hasattr(element, 'max_width_inches') else 5.0))
                except Exception:
                    # If image can't be added, add placeholder
                    para = self.doc.add_paragraph()
                    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    run = para.add_run(f"[Image: {alt_text or image_url}]")
                    run.font.italic = True
                    run.font.color.rgb = RGBColor(128, 128, 128)
            else:
                # Image file doesn't exist - add placeholder
                para = self.doc.add_paragraph()
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = para.add_run(f"[Image: {alt_text or image_url}]")
                run.font.italic = True
                run.font.color.rgb = RGBColor(128, 128, 128)
        else:
            # No image URL - this might be just a caption
            pass

        # Add caption
        if caption or figure_number > 0:
            caption_para = self.doc.add_paragraph()
            caption_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

            if figure_number > 0:
                num_run = caption_para.add_run(f"Figure {figure_number}: ")
                num_run.font.bold = True
                num_run.font.size = Pt(element.caption_font_size_pt if hasattr(element, 'caption_font_size_pt') else 10)

            if caption:
                cap_run = caption_para.add_run(caption)
                cap_run.font.italic = (element.caption_style == "italic" if hasattr(element, 'caption_style') else True)
                cap_run.font.size = Pt(element.caption_font_size_pt if hasattr(element, 'caption_font_size_pt') else 10)

            caption_para.paragraph_format.space_after = Pt(element.space_after_pt)

    def _add_horizontal_rule(self, element: StyledElement) -> None:
        """Add horizontal rule as paragraph with bottom border."""
        if isinstance(element, StyledHorizontalRule):
            line_color = element.line_color
            line_weight = element.line_weight_pt
        else:
            line_color = "AAAAAA"
            line_weight = 0.5

        # Add empty paragraph with bottom border
        para = self.doc.add_paragraph()
        para.paragraph_format.space_before = Pt(element.space_before_pt)
        para.paragraph_format.space_after = Pt(element.space_after_pt)

        # Add bottom border
        try:
            pPr = para._p.get_or_add_pPr()
            pBdr = OxmlElement('w:pBdr')
            bottom = OxmlElement('w:bottom')
            bottom.set(qn('w:val'), 'single')
            bottom.set(qn('w:sz'), str(int(line_weight * 8)))  # eighths of a point
            bottom.set(qn('w:color'), line_color)
            pBdr.append(bottom)
            pPr.append(pBdr)
        except Exception:
            # Fallback: add a text-based rule
            run = para.add_run("─" * 60)
            run.font.color.rgb = RGBColor(170, 170, 170)

    def _apply_run_formatting(self, run, element: StyledElement) -> None:
        """Apply character formatting to a run."""
        run.font.name = element.font_name
        run.font.size = Pt(element.font_size_pt)
        run.font.bold = element.bold
        run.font.italic = element.italic
        run.font.underline = element.underline

        if element.font_color:
            try:
                r = int(element.font_color[0:2], 16)
                g = int(element.font_color[2:4], 16)
                b = int(element.font_color[4:6], 16)
                run.font.color.rgb = RGBColor(r, g, b)
            except (ValueError, IndexError):
                pass

    def _apply_paragraph_formatting(self, para, element: StyledElement) -> None:
        """Apply paragraph formatting."""
        pf = para.paragraph_format

        # Spacing
        pf.space_before = Pt(element.space_before_pt)
        pf.space_after = Pt(element.space_after_pt)

        # Line spacing
        if element.line_spacing != 1.0:
            pf.line_spacing = element.line_spacing

        # Indentation
        if element.first_line_indent_inches:
            pf.first_line_indent = Inches(element.first_line_indent_inches)
        if element.left_indent_inches:
            pf.left_indent = Inches(element.left_indent_inches)

        # Alignment
        alignment_map = {
            "left": WD_ALIGN_PARAGRAPH.LEFT,
            "center": WD_ALIGN_PARAGRAPH.CENTER,
            "right": WD_ALIGN_PARAGRAPH.RIGHT,
            "justify": WD_ALIGN_PARAGRAPH.JUSTIFY,
        }
        para.alignment = alignment_map.get(element.alignment, WD_ALIGN_PARAGRAPH.LEFT)

        # Page control
        pf.keep_with_next = element.keep_with_next
        pf.keep_together = element.keep_together

    def _set_metadata(self, styled_doc: StyledDocument) -> None:
        """Set document metadata."""
        core_props = self.doc.core_properties

        if styled_doc.title:
            core_props.title = styled_doc.title
        if styled_doc.author:
            core_props.author = styled_doc.author
        if styled_doc.subject:
            core_props.subject = styled_doc.subject

        core_props.language = styled_doc.language
