"""
DOCX Adapter - Converts Document AST to DOCX Format

This module provides the bridge between the rendering-agnostic Document AST
and the python-docx library for DOCX export, with support for commercial-quality
typography.

Architecture:
    DocumentAST → render_docx_from_ast() → .docx file

Commercial Typography Features:
    - Professional ebook formatting (Georgia 11pt, 1.15 line spacing)
    - Smart paragraph indentation (first paragraph: no indent, body: 0.32" indent)
    - Proper heading hierarchy (H1=16pt, H2=14pt, H3=12pt)
    - Page breaks before chapter headings
    - Blockquotes with bilateral indentation
    - Right-aligned epigraphs for chapter openings
    - Generous spacing for scene breaks

Usage:
    from core.rendering.docx_adapter import render_docx_from_ast
    from pathlib import Path

    # Create or load DocumentAST
    ast = build_book_ast(doc_nodes)

    # Render to DOCX with commercial typography
    output_path = Path("output.docx")
    render_docx_from_ast(ast, output_path)
"""

import logging
from pathlib import Path
from typing import Optional

from docx import Document
from docx.shared import Mm, Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

from core.rendering.document_ast import (
    DocumentAST,
    Block,
    Heading,
    HeadingLevel,
    Paragraph,
    ParagraphRole,
    Blockquote,
    Epigraph,
    SceneBreak,
    Equation,
    EquationMode,
    TheoremBox,
    TheoremType,
    ProofBox,
    ReferenceEntry,
    TableBlock,
    Figure,
    ListBlock,
    Caption,
    PageBreak,
    FontStyle,
    SpacingStyle,
    ParagraphStyle,
)

from core.rendering.omml_converter import (
    latex_to_omml,
    inject_omml_as_display,
    is_pandoc_available,
)

logger = logging.getLogger(__name__)

# Monospace face applied to inline-code runs (Word substitutes a metric-compatible
# mono font when Consolas is absent).
_CODE_FONT = "Consolas"


# ============================================================================
# Main Rendering Function
# ============================================================================

def _stylesheet_for_template(name: str):
    """Map a template name to a StyleSheet (parity with the legacy renderer's
    ``template=`` argument): ``ebook`` (Georgia serif book), ``academic``
    (Cambria), ``business`` (Arial sans, left-aligned). Unknown → ebook."""
    from core.rendering.document_ast import (
        create_academic_stylesheet,
        create_book_stylesheet,
    )

    key = (name or "").strip().lower()
    if key in ("academic", "stem"):
        return create_academic_stylesheet()
    if key in ("business", "report"):
        sheet = create_book_stylesheet()
        for ps in (sheet.heading_1, sheet.heading_2, sheet.heading_3, sheet.body):
            ps.font.family = "Arial"
        sheet.body.alignment = "left"
        sheet.body.spacing.first_line_indent_pt = 0.0
        return sheet
    return create_book_stylesheet()  # ebook / book / default


def _render_title_page(doc: Document, ast: DocumentAST, title: Optional[str]) -> None:
    """Prepend a simple centered cover page (title + author) and a page break."""
    md = ast.metadata
    heading_font = ast.styles.heading_1.font.family
    for _ in range(6):  # push the title toward the vertical middle
        doc.add_paragraph()
    tp = doc.add_paragraph()
    tp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    tr = tp.add_run(title or md.title or "")
    tr.bold = True
    tr.font.size = Pt(26)
    tr.font.name = heading_font
    if md.author:
        doc.add_paragraph()
        ap = doc.add_paragraph()
        ap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        ar = ap.add_run(md.author)
        ar.font.size = Pt(15)
        ar.font.name = heading_font
    doc.add_page_break()


def _append_field(run, instr: str, placeholder: str = "") -> None:
    """Append a Word field (begin / instrText / separate / placeholder / end)
    to *run* — used for the TOC and PAGE-number fields, which Word computes when
    the document is opened/updated."""
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    itxt = OxmlElement("w:instrText")
    itxt.set(qn("xml:space"), "preserve")
    itxt.text = instr
    sep = OxmlElement("w:fldChar")
    sep.set(qn("w:fldCharType"), "separate")
    ph = OxmlElement("w:t")
    ph.text = placeholder
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    for el in (begin, itxt, sep, ph, end):
        run._r.append(el)


def _render_toc(doc: Document, ast: DocumentAST) -> None:
    """Insert a localized 'Table of contents' heading + an auto-updating Word
    TOC field (levels 1–3), then a page break."""
    lang = (ast.metadata.language or "").lower()
    label = "Mục lục" if lang.startswith("vi") else "Table of Contents"
    hp = doc.add_paragraph()
    hr = hp.add_run(label)
    hr.bold = True
    hr.font.size = Pt(16)
    hr.font.name = ast.styles.heading_1.font.family
    field_p = doc.add_paragraph()
    _append_field(
        field_p.add_run(),
        'TOC \\o "1-3" \\h \\z \\u',
        "Cập nhật mục lục trong Word: Ctrl+A rồi F9." if lang.startswith("vi")
        else "Update this field in Word: select all, then F9.",
    )
    doc.add_page_break()


def _add_running_header_footer(doc: Document, ast: DocumentAST, title: Optional[str]) -> None:
    """Running header (document title) + footer (centered PAGE number), with the
    first page (cover) left blank via different-first-page."""
    section = doc.sections[0]
    section.different_first_page_header_footer = True

    header = section.header
    hp = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
    hp.text = title or ast.metadata.title or ""
    hp.alignment = WD_ALIGN_PARAGRAPH.CENTER

    footer = section.footer
    fp = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _append_field(fp.add_run(), "PAGE", "1")


def render_docx_from_ast(
    ast: DocumentAST,
    output_path: Path,
    title: Optional[str] = None,
    template: Optional[str] = None,
    title_page: bool = False,
    toc: bool = False,
    header_footer: bool = False,
) -> None:
    """
    Render DocumentAST to DOCX format.

    Args:
        ast: The DocumentAST to render
        output_path: Path where the DOCX file will be saved
        title: Optional document title (overrides ast.metadata.title)
        template: Optional style template — "ebook" (Georgia serif book),
            "academic" (Cambria) or "business" (Arial sans, left-aligned).
            When given, its stylesheet replaces ``ast.styles`` for this render;
            when None, the AST's own styles are used unchanged.
        title_page: when True, prepend a centered title/author cover page and a
            page break before the content.
        toc: when True, insert a localized heading + an auto-updating Word table
            of contents (levels 1–3) after the cover.
        header_footer: when True, add a running header (title) and a footer with
            a centered page number (cover page left blank).

    Raises:
        ValueError: If AST is invalid
        IOError: If file cannot be written

    Example:
        >>> ast = build_book_ast(doc_nodes)
        >>> render_docx_from_ast(ast, Path("book.docx"), template="ebook")
    """
    if template:
        import dataclasses

        ast = dataclasses.replace(ast, styles=_stylesheet_for_template(template))

    logger.info(f"Rendering DocumentAST to DOCX: {output_path}")
    logger.info(f"AST summary: {ast}")

    # Create new Document
    doc = Document()

    # Set document-level properties
    _setup_document_properties(doc, ast, title)

    if title_page:
        _render_title_page(doc, ast, title)

    if toc:
        _render_toc(doc, ast)

    if header_footer:
        _add_running_header_footer(doc, ast, title)

    # Render each block
    for idx, block in enumerate(ast.blocks):
        try:
            _render_block(doc, block, ast)
        except Exception as e:
            logger.error(f"Error rendering block {idx} ({block.block_type}): {e}")
            # Continue rendering other blocks
            continue

    # Phase 4.3: Apply advanced book layout (OPTIONAL - disabled by default)
    try:
        from config.settings import settings
        if settings.enable_advanced_book_layout:
            logger.info("Applying advanced book layout features (Phase 4.3 - EXPERIMENTAL)...")
            from core.export.book_layout import apply_book_layout
            apply_book_layout(doc, ast.metadata)
            logger.info("✅ Advanced book layout applied")
    except ImportError as e:
        logger.warning(f"Could not import book layout dependencies: {e}")
    except Exception as e:
        logger.warning(f"Book layout application failed (non-critical): {e}")
        # Continue with save even if book layout fails

    # Save document
    try:
        output_path.parent.mkdir(parents=True, exist_ok=True)
        doc.save(str(output_path))
        logger.info(f"✅ DOCX saved successfully: {output_path}")
    except Exception as e:
        logger.error(f"Failed to save DOCX: {e}")
        raise IOError(f"Cannot save DOCX to {output_path}: {e}")


# ============================================================================
# Document Setup
# ============================================================================

def _setup_document_properties(
    doc: Document,
    ast: DocumentAST,
    title: Optional[str]
) -> None:
    """Set document core properties and page geometry from AST metadata."""
    # Set core properties
    props = doc.core_properties
    props.title = title or ast.metadata.title or ""
    props.author = ast.metadata.author or ""
    props.language = ast.metadata.language

    # Page geometry from metadata — parity with the legacy layout engine, which
    # applies page size + margins. python-docx defaults to US Letter otherwise,
    # so the AST's A4/margins (document_ast.DocumentMetadata) would be lost.
    md = ast.metadata
    try:
        section = doc.sections[0]
        section.page_width = Mm(md.page_width_mm)
        section.page_height = Mm(md.page_height_mm)
        section.top_margin = Mm(md.margin_top_mm)
        section.bottom_margin = Mm(md.margin_bottom_mm)
        section.left_margin = Mm(md.margin_left_mm)
        section.right_margin = Mm(md.margin_right_mm)
    except Exception as e:  # pragma: no cover - defensive
        logger.warning(f"Page setup from metadata failed: {e}")

    logger.debug(f"Document properties set: title={props.title}, author={props.author}")


# ============================================================================
# Block Rendering Dispatch
# ============================================================================

def _render_block(doc: Document, block: Block, ast: DocumentAST) -> None:
    """
    Dispatch block rendering based on block type.

    Args:
        doc: python-docx Document
        block: AST Block to render
        ast: Full DocumentAST (for style lookups)
    """
    if isinstance(block, Heading):
        _render_heading(doc, block, ast)
    elif isinstance(block, Paragraph):
        _render_paragraph(doc, block, ast)
    elif isinstance(block, Blockquote):
        _render_blockquote(doc, block, ast)
    elif isinstance(block, Epigraph):
        _render_epigraph(doc, block, ast)
    elif isinstance(block, SceneBreak):
        _render_scene_break(doc, block, ast)
    elif isinstance(block, Equation):
        _render_equation(doc, block, ast)
    elif isinstance(block, TheoremBox):
        _render_theorem_box(doc, block, ast)
    elif isinstance(block, ProofBox):
        _render_proof_box(doc, block, ast)
    elif isinstance(block, ReferenceEntry):
        _render_reference_entry(doc, block, ast)
    elif isinstance(block, TableBlock):
        _render_table(doc, block, ast)
    elif isinstance(block, Figure):
        _render_figure(doc, block, ast)
    elif isinstance(block, ListBlock):
        _render_list(doc, block, ast)
    elif isinstance(block, Caption):
        _render_caption(doc, block, ast)
    elif isinstance(block, PageBreak):
        _render_page_break(doc, block, ast)
    else:
        logger.warning(f"Unknown block type: {type(block)}")


# ============================================================================
# Block Renderers
# ============================================================================

def _render_heading(doc: Document, heading: Heading, ast: DocumentAST) -> None:
    """Render a heading block using Word's built-in heading styles."""
    # Get style from AST
    if heading.level == HeadingLevel.H1:
        style = ast.styles.heading_1
        word_style = 'Heading 1'
    elif heading.level == HeadingLevel.H2:
        style = ast.styles.heading_2
        word_style = 'Heading 2'
    else:
        style = ast.styles.heading_3
        word_style = 'Heading 3'

    # Format heading text
    if heading.number:
        text = f"{heading.number}. {heading.text}"
    else:
        text = heading.text

    # Add heading paragraph with built-in Word style
    p = doc.add_paragraph(text, style=word_style)

    # Override the default heading style with our custom formatting
    # Get first run and apply our font style
    if p.runs:
        run = p.runs[0]
        _apply_font_style(run, style.font)

    # Apply paragraph styling (spacing, alignment, etc.)
    _apply_paragraph_style(p, style)

    logger.debug(f"Rendered heading: {heading.level.name} ({word_style}) - {heading.text}")


def _render_paragraph(doc: Document, para: Paragraph, ast: DocumentAST) -> None:
    """
    Render a paragraph block with commercial typography.

    First paragraphs (after headings/breaks) have no indent,
    subsequent body paragraphs have ~0.32" first-line indent.
    """
    # Determine style based on role
    if para.role == ParagraphRole.FIRST_PARAGRAPH:
        # First paragraph after heading/break: no indent (commercial convention)
        style = ParagraphStyle(
            font=ast.styles.body.font,
            spacing=SpacingStyle(
                line_spacing=ast.styles.body.spacing.line_spacing,
                first_line_indent_pt=0.0,  # No indent for first paragraph
                space_after_pt=ast.styles.body.spacing.space_after_pt
            ),
            alignment=ast.styles.body.alignment
        )
    else:
        # Regular body paragraph: use full commercial style with indent
        style = ast.styles.body

    # Use custom style if provided
    if para.style:
        style = para.style

    # Add paragraph
    p = doc.add_paragraph()
    if para.runs:
        # Inline-formatting overlay: one Word run per span. The base font/size/
        # color come from the paragraph style; bold/italic/code are layered on
        # top so a span inherits the design and only adds its own emphasis.
        for span in para.runs:
            run = p.add_run(span.text)
            _apply_font_style(run, style.font)
            if span.bold:
                run.font.bold = True
            if span.italic:
                run.font.italic = True
            if span.code:
                run.font.name = _CODE_FONT
    else:
        run = p.add_run(para.text)
        _apply_font_style(run, style.font)

    # Apply styling
    _apply_paragraph_style(p, style)

    logger.debug(f"Rendered paragraph: role={para.role.name}, indent={style.spacing.first_line_indent_pt}pt")


def _render_blockquote(doc: Document, quote: Blockquote, ast: DocumentAST) -> None:
    """Render a blockquote."""
    style = quote.style or ast.styles.blockquote

    # Add quote text
    p = doc.add_paragraph()
    run = p.add_run(quote.text)

    # Apply styling
    _apply_font_style(run, style.font)
    _apply_paragraph_style(p, style)

    # Add attribution if present
    if quote.attribution:
        attr_p = doc.add_paragraph()
        attr_run = attr_p.add_run(f"— {quote.attribution}")
        attr_run.font.italic = True
        attr_run.font.size = Pt(10.0)
        attr_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    logger.debug(f"Rendered blockquote with attribution={quote.attribution}")


def _render_epigraph(doc: Document, epigraph: Epigraph, ast: DocumentAST) -> None:
    """Render an epigraph (chapter opening quote)."""
    style = epigraph.style or ast.styles.epigraph

    # Add epigraph text
    p = doc.add_paragraph()
    run = p.add_run(epigraph.text)

    # Apply styling
    _apply_font_style(run, style.font)
    _apply_paragraph_style(p, style)

    # Add attribution if present
    if epigraph.attribution:
        attr_p = doc.add_paragraph()
        attr_run = attr_p.add_run(f"— {epigraph.attribution}")
        attr_run.font.italic = True
        attr_run.font.size = style.font.size_pt
        attr_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    logger.debug(f"Rendered epigraph")


def _render_scene_break(doc: Document, scene_break: SceneBreak, ast: DocumentAST) -> None:
    """Render a scene break separator."""
    style = scene_break.style or ast.styles.scene_break

    # Add scene break symbol
    p = doc.add_paragraph()
    run = p.add_run(scene_break.symbol)

    # Apply styling
    _apply_font_style(run, style.font)
    _apply_paragraph_style(p, style)

    logger.debug(f"Rendered scene break: {scene_break.symbol}")


def _render_equation(doc: Document, equation: Equation, ast: DocumentAST) -> None:
    """
    Render an equation using OMML (Office Math Markup Language) if available.

    Rendering Strategy:
    1. If pandoc is available → Convert LaTeX to OMML for native Word rendering
    2. If pandoc unavailable → Fallback to LaTeX text (plain text with $ delimiters)

    Display mode equations are center-aligned and use <m:oMathPara> wrapper to
    prevent line wrapping.
    """
    p = doc.add_paragraph()

    if equation.mode == EquationMode.DISPLAY:
        # Display equation: center-aligned
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Try OMML rendering first (native Word equations)
        if is_pandoc_available():
            omml_xml = latex_to_omml(equation.latex)
            if omml_xml:
                # Successfully converted to OMML
                success = inject_omml_as_display(p, omml_xml)
                if success:
                    logger.debug(f"Rendered display equation with OMML: {equation.latex[:50]}...")

                    # Add equation number if present (as separate paragraph for proper alignment)
                    if equation.number:
                        # TODO: Implement proper equation numbering with right alignment
                        # For now, skip numbering in OMML mode to avoid layout issues
                        pass
                    return
                else:
                    logger.warning(f"OMML injection failed for display equation, using fallback")
            else:
                logger.warning(f"LaTeX to OMML conversion failed for: {equation.latex[:50]}")

        # Fallback: render as plain text with $$ delimiters
        text = f"$$ {equation.latex} $$"
        if equation.number:
            text += f"  ({equation.number})"

        run = p.add_run(text)
        run.font.name = "Cambria Math"
        run.font.size = Pt(11.0)
        logger.debug(f"Rendered display equation with LaTeX fallback: {equation.latex[:50]}")

    else:
        # Inline equation: no OMML support for inline mode yet (requires different XML structure)
        # TODO: Implement inline OMML rendering (requires <m:oMath> without <m:oMathPara>)
        text = f"$ {equation.latex} $"
        run = p.add_run(text)
        run.font.name = "Cambria Math"
        run.font.size = Pt(11.0)
        logger.debug(f"Rendered inline equation with LaTeX: {equation.latex[:30]}")


def _render_theorem_box(doc: Document, theorem: TheoremBox, ast: DocumentAST) -> None:
    """
    Render a theorem box.

    Note: This is a simplified implementation.
    Full version would add borders/shading.
    """
    style = theorem.style or ast.styles.theorem_box

    # Add theorem title
    title_p = doc.add_paragraph()
    title_text = f"{theorem.title}"
    if theorem.number:
        title_text = f"{theorem.title} {theorem.number}"

    title_run = title_p.add_run(title_text)
    title_run.bold = True
    title_run.font.size = Pt(11.0)

    # Add theorem content
    content_p = doc.add_paragraph()
    content_run = content_p.add_run(theorem.content)

    # Apply styling
    _apply_font_style(content_run, style.font)
    _apply_paragraph_style(content_p, style)

    logger.debug(f"Rendered theorem: type={theorem.theorem_type.name}, number={theorem.number}")


def _render_proof_box(doc: Document, proof: ProofBox, ast: DocumentAST) -> None:
    """Render a proof environment."""
    style = proof.style or ast.styles.proof_box

    # Add "Proof." label
    label_p = doc.add_paragraph()
    label_run = label_p.add_run("Proof. ")
    label_run.italic = True

    # Add proof content
    content_p = doc.add_paragraph()
    content_run = content_p.add_run(proof.content)

    # Apply styling
    _apply_font_style(content_run, style.font)
    _apply_paragraph_style(content_p, style)

    # Add QED symbol
    qed_p = doc.add_paragraph()
    qed_run = qed_p.add_run(proof.qed_symbol)
    qed_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    logger.debug(f"Rendered proof box")


def _render_reference_entry(doc: Document, ref: ReferenceEntry, ast: DocumentAST) -> None:
    """Render a bibliography/reference entry."""
    p = doc.add_paragraph()

    if ref.key:
        text = f"[{ref.key}] {ref.citation}"
    else:
        text = ref.citation

    run = p.add_run(text)
    run.font.size = Pt(10.0)

    # Hanging indent for references
    p.paragraph_format.left_indent = Inches(0.5)
    p.paragraph_format.first_line_indent = Inches(-0.5)

    logger.debug(f"Rendered reference: key={ref.key}")


def _render_table(doc: Document, table_block: TableBlock, ast: DocumentAST) -> None:
    """Render a table with a bordered grid; header rows are bold."""
    rows = table_block.rows or []
    if not rows:
        return
    n_cols = max((len(r) for r in rows), default=0)
    if n_cols == 0:
        return

    table = doc.add_table(rows=len(rows), cols=n_cols)
    try:
        table.style = "Table Grid"  # visible borders
    except Exception:
        pass  # style may be absent in a stripped template

    for r, row in enumerate(rows):
        for c in range(n_cols):
            cell = table.rows[r].cells[c]
            cell.text = row[c] if c < len(row) else ""
            if r < table_block.header_rows:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.bold = True

    if table_block.caption:
        cap = doc.add_paragraph()
        run = cap.add_run(table_block.caption)
        run.font.italic = True
        run.font.size = Pt(10.0)
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER

    logger.debug(f"Rendered table: {len(rows)}x{n_cols}")


def _render_figure(doc: Document, figure: Figure, ast: DocumentAST) -> None:
    """Embed the figure image, else a captioned placeholder.

    Preference order: carried ``image_bytes`` (survives a round trip) → a
    readable file at ``image_ref`` → a centered ``[Figure: …]`` placeholder when
    no pixels are available.
    """
    added = False
    ref = figure.image_ref or ""

    if figure.image_bytes:
        try:
            from io import BytesIO

            stream = BytesIO(figure.image_bytes)
            if figure.width_pt:
                doc.add_picture(stream, width=Inches(figure.width_pt / 72.0))
            else:
                doc.add_picture(stream)
            added = True
        except Exception as e:  # fall through to ref/placeholder
            logger.warning(f"Figure image add from bytes failed: {e}")

    if not added:
        try:
            if ref and not ref.startswith(("embedded", "/word/")) and Path(ref).is_file():
                if figure.width_pt:
                    doc.add_picture(ref, width=Inches(figure.width_pt / 72.0))
                else:
                    doc.add_picture(ref)
                added = True
        except Exception as e:
            logger.warning(f"Figure image add failed ({ref}): {e}")

    if not added:
        placeholder = doc.add_paragraph()
        label = figure.alt_text or figure.caption or ref or "image"
        run = placeholder.add_run(f"[Figure: {label}]")
        run.font.italic = True
        placeholder.alignment = WD_ALIGN_PARAGRAPH.CENTER

    if figure.caption:
        cap = doc.add_paragraph()
        prefix = f"Hình {figure.number}. " if figure.number else ""
        run = cap.add_run(prefix + figure.caption)
        run.font.italic = True
        run.font.size = Pt(10.0)
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER

    logger.debug(f"Rendered figure: ref={ref}, embedded={added}")


def _render_list(doc: Document, list_block: ListBlock, ast: DocumentAST) -> None:
    """Render an ordered/unordered list using Word's list styles."""
    style_name = "List Number" if list_block.ordered else "List Bullet"
    for item in list_block.items:
        try:
            doc.add_paragraph(item, style=style_name)
        except Exception:
            # Template may lack the list style — fall back to a manual marker.
            paragraph = doc.add_paragraph()
            paragraph.add_run(("• " if not list_block.ordered else "- ") + item)
    logger.debug(f"Rendered list: {len(list_block.items)} items, ordered={list_block.ordered}")


def _render_caption(doc: Document, caption: Caption, ast: DocumentAST) -> None:
    """Render a standalone caption (centered italic)."""
    paragraph = doc.add_paragraph()
    label = ""
    if caption.target and caption.number:
        label = f"{caption.target.capitalize()} {caption.number}. "
    run = paragraph.add_run(label + caption.text)
    run.font.italic = True
    run.font.size = Pt(10.0)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    logger.debug(f"Rendered caption: target={caption.target}, number={caption.number}")


def _render_page_break(doc: Document, page_break: PageBreak, ast: DocumentAST) -> None:
    """Render an explicit page break."""
    doc.add_page_break()
    logger.debug("Rendered page break")


# ============================================================================
# Style Application Helpers
# ============================================================================

def _apply_font_style(run, font_style: FontStyle) -> None:
    """Apply font styling to a run, honoring the AST's requested family.

    The chosen typography (e.g. the book stylesheet's Georgia, which supports
    Vietnamese) is applied as-is so the DOCX matches the intended design and the
    legacy engine's font handling — Georgia/Cambria are no longer silently
    rewritten to Times New Roman. A genuine assignment error falls back to a
    Vietnamese-safe face rather than dropping the styling.
    """
    try:
        run.font.name = font_style.family
    except Exception as e:  # pragma: no cover - defensive
        logger.warning(f"Font assignment failed for '{font_style.family}', using fallback: {e}")
        run.font.name = "Times New Roman"

    # Apply other font properties
    run.font.size = Pt(font_style.size_pt)
    run.font.bold = font_style.bold
    run.font.italic = font_style.italic

    # Apply color
    if font_style.color:
        try:
            # Parse hex color (e.g., "FF0000" -> RGB(255, 0, 0))
            hex_color = font_style.color
            r = int(hex_color[0:2], 16)
            g = int(hex_color[2:4], 16)
            b = int(hex_color[4:6], 16)
            run.font.color.rgb = RGBColor(r, g, b)
        except (ValueError, IndexError):
            logger.warning(f"Invalid color format: {font_style.color}")


def _apply_paragraph_style(p, para_style: ParagraphStyle) -> None:
    """Apply paragraph styling."""
    fmt = p.paragraph_format

    # Spacing
    fmt.line_spacing = para_style.spacing.line_spacing
    fmt.space_before = Pt(para_style.spacing.space_before_pt)
    fmt.space_after = Pt(para_style.spacing.space_after_pt)
    fmt.first_line_indent = Pt(para_style.spacing.first_line_indent_pt)
    fmt.left_indent = Pt(para_style.spacing.left_indent_pt)
    fmt.right_indent = Pt(para_style.spacing.right_indent_pt)

    # Alignment
    alignment_map = {
        "left": WD_ALIGN_PARAGRAPH.LEFT,
        "right": WD_ALIGN_PARAGRAPH.RIGHT,
        "center": WD_ALIGN_PARAGRAPH.CENTER,
        "justify": WD_ALIGN_PARAGRAPH.JUSTIFY,
    }
    p.alignment = alignment_map.get(para_style.alignment, WD_ALIGN_PARAGRAPH.LEFT)

    # Keep with next / keep together
    fmt.keep_with_next = para_style.keep_with_next
    fmt.keep_together = para_style.keep_together

    # Page break before
    fmt.page_break_before = para_style.page_break_before


# ============================================================================
# Convenience Functions
# ============================================================================

def render_book_docx(ast: DocumentAST, output_path: Path) -> None:
    """Render book-style DOCX — Georgia serif template with full front matter
    (title page + table of contents + running header/footer)."""
    render_docx_from_ast(
        ast, output_path, template="ebook",
        title_page=True, toc=True, header_footer=True,
    )


def render_academic_docx(ast: DocumentAST, output_path: Path) -> None:
    """Render academic-style DOCX — Cambria template with full front matter
    (title page + table of contents + running header/footer)."""
    render_docx_from_ast(
        ast, output_path, template="academic",
        title_page=True, toc=True, header_footer=True,
    )
