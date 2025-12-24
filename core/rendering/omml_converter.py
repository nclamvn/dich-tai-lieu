"""
Phase 2.0.3b - OMML Equation Converter

Converts LaTeX equations to OMML (Office Math Markup Language) for native Word math rendering.

Uses pandoc to convert LaTeX → DOCX → extract OMML XML.
"""

import subprocess
import tempfile
import os
import re
import logging
from typing import Optional
from lxml import etree
from docx import Document as DocxDocument

logger = logging.getLogger(__name__)


def strip_latex_delimiters(latex_str: str) -> str:
    """
    Remove LaTeX delimiters from equation string.

    Handles: $, $$, \\[, \\], \\(, \\)

    Args:
        latex_str: LaTeX string with or without delimiters

    Returns:
        Clean LaTeX without delimiters

    Examples:
        "$x^2$" → "x^2"
        "$$\\int_0^\\infty$$" → "\\int_0^\\infty"
        "\\[E=mc^2\\]" → "E=mc^2"
        "x^2" → "x^2" (no delimiters)
    """
    latex_str = latex_str.strip()

    # Remove display math delimiters: \[ ... \]
    if latex_str.startswith('\\[') and latex_str.endswith('\\]'):
        return latex_str[2:-2].strip()

    # Remove inline math delimiters: \( ... \)
    if latex_str.startswith('\\(') and latex_str.endswith('\\)'):
        return latex_str[2:-2].strip()

    # Remove $$ delimiters (display math)
    if latex_str.startswith('$$') and latex_str.endswith('$$'):
        return latex_str[2:-2].strip()

    # Remove $ delimiters (inline math)
    if latex_str.startswith('$') and latex_str.endswith('$'):
        return latex_str[1:-1].strip()

    # No delimiters found
    return latex_str


def latex_to_omml(latex_str: str, timeout: int = 5) -> Optional[str]:
    """
    Convert LaTeX equation to OMML XML using pandoc.

    Process:
        1. Strip LaTeX delimiters ($, $$, etc.)
        2. Create temporary .tex file
        3. Call pandoc to convert .tex → .docx
        4. Extract OMML XML from generated DOCX
        5. Return OMML XML string

    Args:
        latex_str: LaTeX equation with or without delimiters
        timeout: Pandoc subprocess timeout in seconds (default: 5)

    Returns:
        OMML XML string if successful, None if conversion fails

    Safety:
        - Try/except around all operations
        - Returns None on failure (caller should fallback to LaTeX text)
        - Timeout protection on pandoc subprocess
        - Validates XML before returning
        - Cleans up temporary files

    Examples:
        >>> latex_to_omml(r"$x^2$")
        '<m:oMath xmlns:m="http://schemas.openxmlformats.org/officeDocument/2006/math">...</m:oMath>'

        >>> latex_to_omml(r"$$\\int_0^\\infty e^{-x^2} dx$$")
        '<m:oMath ...>...</m:oMath>'

        >>> latex_to_omml(r"$\\invalid{command}$")  # Invalid LaTeX
        None
    """
    tex_path = None
    docx_path = None

    try:
        # 1. Strip delimiters
        clean_latex = strip_latex_delimiters(latex_str)

        if not clean_latex:
            logger.warning("Empty LaTeX string after stripping delimiters")
            return None

        # 2. Create temporary LaTeX file
        with tempfile.NamedTemporaryFile(
            mode='w',
            suffix='.tex',
            delete=False,
            encoding='utf-8'
        ) as tex_file:
            # Wrap in display math delimiters for pandoc
            tex_content = f"$${clean_latex}$$"
            tex_file.write(tex_content)
            tex_path = tex_file.name

        # 3. Create temporary output DOCX path
        docx_path = tex_path.replace('.tex', '.docx')

        # 4. Call pandoc subprocess
        try:
            result = subprocess.run(
                ['pandoc', '-f', 'latex', '-t', 'docx', tex_path, '-o', docx_path],
                capture_output=True,
                text=True,
                timeout=timeout
            )

            if result.returncode != 0:
                logger.debug(f"Pandoc conversion failed (exit {result.returncode}): {result.stderr[:100]}")
                return None

        except subprocess.TimeoutExpired:
            logger.warning(f"Pandoc timeout ({timeout}s) for LaTeX: {latex_str[:50]}...")
            return None

        except FileNotFoundError:
            logger.error("Pandoc not found - OMML rendering unavailable. Install pandoc or use equation_rendering_mode='latex'")
            return None

        # 5. Check if DOCX was created
        if not os.path.exists(docx_path):
            logger.warning("Pandoc did not create output DOCX")
            return None

        # 6. Extract OMML XML from generated DOCX
        try:
            doc = DocxDocument(docx_path)
        except Exception as e:
            logger.warning(f"Failed to open generated DOCX: {e}")
            return None

        if len(doc.paragraphs) == 0:
            logger.warning("Generated DOCX has no paragraphs")
            return None

        # Get first paragraph's XML element
        para_elem = doc.paragraphs[0]._element

        # Find <m:oMath> elements (OMML namespace)
        omml_ns = '{http://schemas.openxmlformats.org/officeDocument/2006/math}'
        omath_elems = para_elem.findall(f'.//{omml_ns}oMath')

        if not omath_elems:
            logger.debug(f"No OMML elements found in generated DOCX for: {latex_str[:50]}...")
            return None

        # 7. Serialize OMML XML to string
        try:
            omml_xml = etree.tostring(omath_elems[0], encoding='unicode')
        except Exception as e:
            logger.warning(f"Failed to serialize OMML XML: {e}")
            return None

        # 8. Validate XML structure (basic check)
        if '<m:oMath' not in omml_xml:
            logger.warning("Generated XML does not contain expected OMML structure")
            return None

        logger.debug(f"Successfully converted LaTeX to OMML ({len(omml_xml)} chars)")
        return omml_xml

    except Exception as e:
        logger.error(f"Unexpected error in latex_to_omml: {e}")
        return None

    finally:
        # Cleanup temporary files
        if tex_path and os.path.exists(tex_path):
            try:
                os.unlink(tex_path)
            except Exception as e:
                logger.warning(f"Failed to delete temp .tex file: {e}")

        if docx_path and os.path.exists(docx_path):
            try:
                os.unlink(docx_path)
            except Exception as e:
                logger.warning(f"Failed to delete temp .docx file: {e}")


def inject_omml_into_paragraph(para, omml_xml: str) -> bool:
    """
    Inject OMML XML into python-docx paragraph element.

    Process:
        1. Parse OMML XML string using lxml
        2. Access paragraph's low-level XML element
        3. Create a <w:r> (run) element
        4. Append OMML to the run

    Args:
        para: python-docx Paragraph object
        omml_xml: OMML XML string from latex_to_omml()

    Returns:
        True if injection succeeded, False if failed

    Safety:
        - Validates XML structure before injection
        - Try/except around all XML operations
        - Returns False on failure (caller should fallback)

    Note:
        The paragraph should already have alignment set (CENTER for equations).
        This function only injects the OMML content.
    """
    try:
        # 1. Parse OMML XML string
        try:
            omml_elem = etree.fromstring(omml_xml)
        except etree.XMLSyntaxError as e:
            logger.error(f"Invalid OMML XML: {e}")
            return False

        # 2. Get paragraph's low-level XML element
        para_elem = para._element

        # 3. Create a <w:r> (run) element to contain the OMML
        # Namespace for Word XML
        w_ns = '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}'
        run_elem = etree.SubElement(para_elem, f'{w_ns}r')

        # 4. Append OMML element to the run
        run_elem.append(omml_elem)

        logger.debug("Successfully injected OMML into paragraph")
        return True

    except AttributeError as e:
        logger.error(f"Paragraph object missing expected attributes: {e}")
        return False

    except Exception as e:
        logger.error(f"Unexpected error in inject_omml_into_paragraph: {e}")
        return False


def inject_omml_as_display(para, omml_xml: str) -> bool:
    """
    Inject OMML XML as DISPLAY MODE equation (wrapped in <m:oMathPara>).

    Display mode equations:
    - Appear on their own line (not inline with text)
    - Centered automatically by Word
    - Larger display, won't wrap across lines

    OMML Structure for Display Mode:
        <m:oMathPara>
            <m:oMathParaPr/>   <!-- Display mode properties -->
            <m:oMath>          <!-- The actual equation -->
                ...
            </m:oMath>
        </m:oMathPara>

    Args:
        para: python-docx Paragraph object
        omml_xml: OMML XML string from latex_to_omml() (just <m:oMath>...</m:oMath>)

    Returns:
        True if injection succeeded, False if failed

    Phase 2.1.1: This fixes the "equations split into multiple lines" issue
    """
    try:
        # 1. Parse OMML XML string (<m:oMath>...</m:oMath>)
        try:
            omath_elem = etree.fromstring(omml_xml)
        except etree.XMLSyntaxError as e:
            logger.error(f"Invalid OMML XML: {e}")
            return False

        # 2. Get paragraph's low-level XML element
        para_elem = para._element

        # 3. Create <m:oMathPara> wrapper for display mode
        m_ns = '{http://schemas.openxmlformats.org/officeDocument/2006/math}'
        omath_para_elem = etree.Element(f'{m_ns}oMathPara')

        # 4. Add <m:oMathParaPr> (display mode properties - optional but recommended)
        omath_para_pr = etree.SubElement(omath_para_elem, f'{m_ns}oMathParaPr')

        # 5. Append the actual <m:oMath> equation inside the para wrapper
        omath_para_elem.append(omath_elem)

        # 6. Append the entire <m:oMathPara> structure to paragraph
        para_elem.append(omath_para_elem)

        logger.debug("Successfully injected OMML as DISPLAY MODE equation")
        return True

    except AttributeError as e:
        logger.error(f"Paragraph object missing expected attributes: {e}")
        return False

    except Exception as e:
        logger.error(f"Unexpected error in inject_omml_as_display: {e}")
        return False


def check_pandoc_available() -> bool:
    """
    Check if pandoc is installed and accessible.

    Returns:
        True if pandoc is available, False otherwise

    Usage:
        if not check_pandoc_available():
            logger.warning("Pandoc not available - falling back to LaTeX text rendering")
            config.equation_rendering_mode = "latex"
    """
    try:
        result = subprocess.run(
            ['pandoc', '--version'],
            capture_output=True,
            timeout=2
        )
        return result.returncode == 0
    except (FileNotFoundError, subprocess.TimeoutExpired):
        return False


# Module-level pandoc availability check (cached)
_PANDOC_AVAILABLE = None


def is_pandoc_available() -> bool:
    """
    Cached check for pandoc availability.

    Returns:
        True if pandoc is available, False otherwise

    Note:
        Result is cached after first call to avoid repeated subprocess calls.
    """
    global _PANDOC_AVAILABLE
    if _PANDOC_AVAILABLE is None:
        _PANDOC_AVAILABLE = check_pandoc_available()
        if _PANDOC_AVAILABLE:
            logger.info("Pandoc found - OMML rendering available")
        else:
            logger.warning("Pandoc not found - OMML rendering unavailable")
    return _PANDOC_AVAILABLE
