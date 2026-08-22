"""Apply a chosen cover template to an already-produced export file.

Engine-agnostic on purpose: these helpers operate on the *finished* PDF/DOCX, so
a template cover works whether the document came from the legacy engine or the
AST pipeline. The cover is rendered at the produced file's own page size, so it
always matches.

``apply_cover_to_pdf`` merges a vector cover as page 1 (pypdf).
``apply_cover_to_docx`` inserts a full-bleed cover image as page 1 in its own
zero-margin section (python-docx), leaving the body's section/margins intact.

Both are best-effort: unknown template, missing file, or any error returns
``False`` and leaves the original file untouched (never fails the export).
"""

from __future__ import annotations

import logging
import tempfile
from pathlib import Path
from types import SimpleNamespace

from core.rendering.cover_templates import (
    has_template,
    render_cover_image,
    render_cover_pdf,
)

logger = logging.getLogger(__name__)

_PT_PER_MM = 72.0 / 25.4          # PDF points per millimetre
_EMU_PER_MM = 36000.0            # DOCX English Metric Units per millimetre


def _meta(width_mm: float, height_mm: float, title: str, author: str, language: str):
    return SimpleNamespace(
        page_width_mm=width_mm,
        page_height_mm=height_mm,
        title=(title or "Untitled").strip() or "Untitled",
        author=(author or "").strip(),
        language=(language or "vi"),
    )


def apply_cover_to_pdf(
    pdf_path, template_id: str, *, title: str = "", author: str = "", language: str = "vi"
) -> bool:
    """Prepend a template cover (matched to the PDF's page size) as page 1."""
    pdf_path = Path(pdf_path)
    if not template_id or not has_template(template_id) or not pdf_path.is_file():
        return False
    try:
        import pypdf

        reader = pypdf.PdfReader(str(pdf_path))
        if not reader.pages:
            return False
        box = reader.pages[0].mediabox
        w_mm = float(box.width) / _PT_PER_MM
        h_mm = float(box.height) / _PT_PER_MM
        with tempfile.TemporaryDirectory() as td:
            cover = Path(td) / "cover.pdf"
            render_cover_pdf(template_id, _meta(w_mm, h_mm, title, author, language), cover)
            writer = pypdf.PdfWriter()
            for pg in pypdf.PdfReader(str(cover)).pages:
                writer.add_page(pg)
            for pg in reader.pages:
                writer.add_page(pg)
            with open(pdf_path, "wb") as fh:
                writer.write(fh)
        logger.info("applied '%s' cover to PDF %s", template_id, pdf_path.name)
        return True
    except Exception as e:  # pragma: no cover - never fail the export over a cover
        logger.warning("apply_cover_to_pdf failed (%s); PDF left unchanged", e)
        return False


def _docx_insert_cover(docx_path, image_provider) -> bool:
    """Insert a full-bleed page-1 cover in its own zero-margin section.

    ``image_provider(width_mm, height_mm) -> image_path`` supplies the cover art
    at the document's page size (a template render, or a user-supplied image). The
    body keeps its own final sectPr / margins as the following section.
    """
    import docx
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.shared import Emu, Pt

    document = docx.Document(str(docx_path))
    if not document.paragraphs:
        return False
    section = document.sections[0]
    pw, ph = int(section.page_width), int(section.page_height)  # EMU
    img = image_provider(pw / _EMU_PER_MM, ph / _EMU_PER_MM)
    if not img or not Path(img).is_file():
        return False

    first = document.paragraphs[0]
    p = first.insert_paragraph_before()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pf = p.paragraph_format
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)
    p.add_run().add_picture(str(img), width=Emu(pw), height=Emu(ph))

    # OOXML page size / margins are in twips (1/1440"), NOT EMU (1/914400").
    emu_to_twips = 635  # 914400 / 1440
    sect_pr = OxmlElement("w:sectPr")
    pg_sz = OxmlElement("w:pgSz")
    pg_sz.set(qn("w:w"), str(pw // emu_to_twips))
    pg_sz.set(qn("w:h"), str(ph // emu_to_twips))
    sect_pr.append(pg_sz)
    pg_mar = OxmlElement("w:pgMar")
    for edge in ("top", "right", "bottom", "left", "header", "footer", "gutter"):
        pg_mar.set(qn(f"w:{edge}"), "0")
    sect_pr.append(pg_mar)
    p._p.get_or_add_pPr().append(sect_pr)

    document.save(str(docx_path))
    return True


def apply_cover_to_docx(
    docx_path, template_id: str, *, title: str = "", author: str = "", language: str = "vi"
) -> bool:
    """Insert a rendered template cover as a full-bleed page 1."""
    docx_path = Path(docx_path)
    if not template_id or not has_template(template_id) or not docx_path.is_file():
        return False
    try:
        with tempfile.TemporaryDirectory() as td:
            def provider(w_mm, h_mm):
                out = Path(td) / "cover.png"
                render_cover_image(
                    template_id, _meta(w_mm, h_mm, title, author, language), out, scale=2.0
                )
                return out

            ok = _docx_insert_cover(docx_path, provider)
        if ok:
            logger.info("applied '%s' cover to DOCX %s", template_id, docx_path.name)
        return ok
    except Exception as e:  # pragma: no cover - never fail the export over a cover
        logger.warning("apply_cover_to_docx failed (%s); DOCX left unchanged", e)
        return False


def apply_cover_image_to_docx(docx_path, image_path) -> bool:
    """Insert a user-supplied image as a full-bleed page-1 cover."""
    docx_path = Path(docx_path)
    if not Path(image_path).is_file() or not docx_path.is_file():
        return False
    try:
        ok = _docx_insert_cover(docx_path, lambda w_mm, h_mm: image_path)
        if ok:
            logger.info("applied custom image cover to DOCX %s", docx_path.name)
        return ok
    except Exception as e:  # pragma: no cover - never fail the export over a cover
        logger.warning("apply_cover_image_to_docx failed (%s); DOCX left unchanged", e)
        return False


def apply_cover_image_to_pdf(pdf_path, image_path) -> bool:
    """Prepend a user-supplied image as a full-bleed page-1 cover."""
    pdf_path = Path(pdf_path)
    if not Path(image_path).is_file() or not pdf_path.is_file():
        return False
    try:
        import pypdf
        from reportlab.lib.utils import ImageReader
        from reportlab.pdfgen import canvas as _canvas

        reader = pypdf.PdfReader(str(pdf_path))
        if not reader.pages:
            return False
        box = reader.pages[0].mediabox
        w, h = float(box.width), float(box.height)
        with tempfile.TemporaryDirectory() as td:
            cover = Path(td) / "cover.pdf"
            c = _canvas.Canvas(str(cover), pagesize=(w, h))
            c.drawImage(ImageReader(str(image_path)), 0, 0, width=w, height=h,
                        preserveAspectRatio=False, mask="auto")
            c.showPage()
            c.save()
            writer = pypdf.PdfWriter()
            for pg in pypdf.PdfReader(str(cover)).pages:
                writer.add_page(pg)
            for pg in reader.pages:
                writer.add_page(pg)
            with open(pdf_path, "wb") as fh:
                writer.write(fh)
        logger.info("applied custom image cover to PDF %s", pdf_path.name)
        return True
    except Exception as e:  # pragma: no cover - never fail the export over a cover
        logger.warning("apply_cover_image_to_pdf failed (%s); PDF left unchanged", e)
        return False


def apply_cover(
    path, fmt: str, *, cover_template=None, cover_image=None,
    title: str = "", author: str = "", language: str = "vi",
) -> bool:
    """Dispatch cover application for a finished PDF/DOCX. A user ``cover_image``
    wins over a ``cover_template``. EPUB covers are baked at build time (see
    ``render_epub_from_ast``), not here. Best-effort; returns True if applied."""
    if cover_image and Path(cover_image).is_file():
        if fmt == "pdf":
            return apply_cover_image_to_pdf(path, cover_image)
        if fmt == "docx":
            return apply_cover_image_to_docx(path, cover_image)
        return False
    if cover_template and has_template(cover_template):
        if fmt == "pdf":
            return apply_cover_to_pdf(path, cover_template, title=title, author=author, language=language)
        if fmt == "docx":
            return apply_cover_to_docx(path, cover_template, title=title, author=author, language=language)
    return False
