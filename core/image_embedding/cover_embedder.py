"""
Cover Image Embedder - AI Publisher Pro

Embed cover images as Page 1 in DOCX and PDF documents.
Supports full-page covers with no margins.

Usage:
    from core.image_embedding import CoverEmbedder

    embedder = CoverEmbedder()
    embedder.add_cover_to_docx(doc, cover_image_base64)
"""

import base64
import io
from pathlib import Path
from typing import Optional, Union

try:
    from docx import Document
    from docx.shared import Inches, Pt, Cm, Emu
    from docx.enum.section import WD_ORIENT
    from docx.enum.text import WD_BREAK
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    WD_BREAK = None

try:
    from PIL import Image as PILImage
    PILLOW_AVAILABLE = True
except ImportError:
    PILLOW_AVAILABLE = False

try:
    from reportlab.lib.pagesizes import letter, A4
    from reportlab.lib.units import inch
    from reportlab.platypus import Image, PageBreak
    REPORTLAB_AVAILABLE = True
except ImportError:
    REPORTLAB_AVAILABLE = False


class CoverEmbedder:
    """
    Embed cover images as Page 1 in documents.

    Features:
    - Full-page cover (no margins)
    - Maintains aspect ratio
    - Supports DOCX and PDF
    - Accepts base64 or file path
    """

    # Default page dimensions (Letter size in inches)
    DEFAULT_PAGE_WIDTH = 8.5
    DEFAULT_PAGE_HEIGHT = 11.0

    def __init__(
        self,
        page_width: float = None,
        page_height: float = None
    ):
        """
        Initialize cover embedder.

        Args:
            page_width: Page width in inches (default: 8.5)
            page_height: Page height in inches (default: 11.0)
        """
        self.page_width = page_width or self.DEFAULT_PAGE_WIDTH
        self.page_height = page_height or self.DEFAULT_PAGE_HEIGHT

    def add_cover_to_docx(
        self,
        doc: "Document",
        cover_image: Union[str, bytes],
        is_base64: bool = True
    ) -> "Document":
        """
        Add cover image as Page 1 of DOCX document.

        The cover will be inserted at the beginning with:
        - Zero margins
        - Full page width/height
        - Centered image
        - Page break after

        Args:
            doc: python-docx Document object
            cover_image: Base64 string or image bytes
            is_base64: True if cover_image is base64 encoded

        Returns:
            Modified Document object
        """
        if not DOCX_AVAILABLE:
            raise ImportError("python-docx required for DOCX cover embedding")

        # Decode base64 if needed
        if is_base64 and isinstance(cover_image, str):
            # Handle data URI format
            if cover_image.startswith('data:'):
                # Extract base64 part from data URI
                cover_image = cover_image.split(',', 1)[1]
            image_data = base64.b64decode(cover_image)
        elif isinstance(cover_image, bytes):
            image_data = cover_image
        else:
            # Assume file path
            with open(cover_image, 'rb') as f:
                image_data = f.read()

        # Get image dimensions
        image_stream = io.BytesIO(image_data)

        if PILLOW_AVAILABLE:
            pil_image = PILImage.open(image_stream)
            img_width, img_height = pil_image.size
            image_stream.seek(0)
        else:
            # Default to page dimensions
            img_width, img_height = int(self.page_width * 96), int(self.page_height * 96)

        # Calculate scaled dimensions to fit page
        width_inches, height_inches = self._calculate_cover_size(
            img_width, img_height
        )

        # Create cover section at the beginning
        self._insert_cover_section(doc, image_stream, width_inches, height_inches)

        return doc

    def _insert_cover_section(
        self,
        doc: "Document",
        image_stream: io.BytesIO,
        width_inches: float,
        height_inches: float
    ):
        """
        Insert cover image at the beginning of document.
        """
        # Get first section (or create one)
        first_section = doc.sections[0]

        # Save original margins
        original_margins = {
            'top': first_section.top_margin,
            'bottom': first_section.bottom_margin,
            'left': first_section.left_margin,
            'right': first_section.right_margin
        }

        # Insert cover paragraph at the very beginning
        # We need to insert BEFORE any existing content
        body = doc.element.body

        # Create a new paragraph for cover
        cover_para = doc.add_paragraph()
        cover_para.alignment = 1  # Center

        # Move cover paragraph to the beginning
        body.insert(0, cover_para._element)

        # Add some vertical spacing to center the image
        # Calculate vertical padding
        page_height_inches = self.page_height
        vertical_padding = (page_height_inches - height_inches) / 2

        if vertical_padding > 0:
            cover_para.paragraph_format.space_before = Inches(vertical_padding * 0.8)

        # Add image to cover paragraph
        run = cover_para.add_run()
        run.add_picture(image_stream, width=Inches(width_inches))

        # Insert page break after cover
        page_break_para = doc.add_paragraph()
        body.insert(1, page_break_para._element)
        page_break_run = page_break_para.add_run()
        page_break_run.add_break(WD_BREAK.PAGE)

    def _calculate_cover_size(
        self,
        img_width: int,
        img_height: int
    ) -> tuple:
        """
        Calculate cover image size to fit page while maintaining aspect ratio.

        Args:
            img_width: Original image width in pixels
            img_height: Original image height in pixels

        Returns:
            Tuple of (width_inches, height_inches)
        """
        # Target: fill page with some margin
        target_width = self.page_width - 1.0  # 0.5" margin each side
        target_height = self.page_height - 1.5  # 0.75" margin top/bottom

        # Calculate aspect ratios
        img_aspect = img_width / img_height
        target_aspect = target_width / target_height

        if img_aspect > target_aspect:
            # Image is wider than target - fit to width
            width_inches = target_width
            height_inches = target_width / img_aspect
        else:
            # Image is taller than target - fit to height
            height_inches = target_height
            width_inches = target_height * img_aspect

        return width_inches, height_inches

    def add_cover_to_pdf_flowables(
        self,
        cover_image: Union[str, bytes],
        is_base64: bool = True
    ) -> list:
        """
        Create PDF flowables for cover page.

        Args:
            cover_image: Base64 string or image bytes
            is_base64: True if cover_image is base64 encoded

        Returns:
            List of ReportLab flowables [Image, PageBreak]
        """
        if not REPORTLAB_AVAILABLE:
            raise ImportError("ReportLab required for PDF cover embedding")

        # Decode base64 if needed
        if is_base64 and isinstance(cover_image, str):
            if cover_image.startswith('data:'):
                cover_image = cover_image.split(',', 1)[1]
            image_data = base64.b64decode(cover_image)
        elif isinstance(cover_image, bytes):
            image_data = cover_image
        else:
            with open(cover_image, 'rb') as f:
                image_data = f.read()

        image_stream = io.BytesIO(image_data)

        # Get image dimensions
        if PILLOW_AVAILABLE:
            pil_image = PILImage.open(io.BytesIO(image_data))
            img_width, img_height = pil_image.size
        else:
            img_width, img_height = int(self.page_width * 96), int(self.page_height * 96)

        # Calculate size
        width_inches, height_inches = self._calculate_cover_size(img_width, img_height)

        # Create flowables
        cover_img = Image(
            image_stream,
            width=width_inches * inch,
            height=height_inches * inch
        )

        return [cover_img, PageBreak()]


def add_cover_to_document(
    doc_path: Union[str, Path],
    cover_image: Union[str, bytes],
    output_path: Union[str, Path] = None,
    is_base64: bool = True
) -> Path:
    """
    Convenience function to add cover to existing DOCX document.

    Args:
        doc_path: Path to existing DOCX file
        cover_image: Base64 string or image bytes
        output_path: Output path (default: overwrites original)
        is_base64: True if cover_image is base64 encoded

    Returns:
        Path to output document
    """
    if not DOCX_AVAILABLE:
        raise ImportError("python-docx required for DOCX operations")

    doc_path = Path(doc_path)
    output_path = Path(output_path) if output_path else doc_path

    # Load document
    doc = Document(str(doc_path))

    # Add cover
    embedder = CoverEmbedder()
    embedder.add_cover_to_docx(doc, cover_image, is_base64)

    # Save
    doc.save(str(output_path))

    return output_path


def decode_cover_image(cover_image_input: str) -> bytes:
    """
    Decode cover image from various input formats.

    Supports:
    - Base64 string
    - Data URI (data:image/png;base64,...)
    - File path

    Args:
        cover_image_input: Cover image in any supported format

    Returns:
        Image bytes
    """
    if cover_image_input.startswith('data:'):
        # Data URI format
        base64_data = cover_image_input.split(',', 1)[1]
        return base64.b64decode(base64_data)
    elif cover_image_input.startswith('/') or Path(cover_image_input).exists():
        # File path
        with open(cover_image_input, 'rb') as f:
            return f.read()
    else:
        # Assume base64
        return base64.b64decode(cover_image_input)
