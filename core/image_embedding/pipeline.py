"""
Image Embedding Pipeline - AI Publisher Pro

Integrates image extraction and embedding with the translation pipeline.
Provides a complete workflow for handling images in document translation.

Usage:
    pipeline = ImageEmbeddingPipeline()

    # Full workflow: extract from source PDF, embed in output DOCX
    result = pipeline.process_document(
        source_pdf="input.pdf",
        output_docx="output.docx",
        translated_text="..."
    )

    # Or step by step
    images = pipeline.extract(source_pdf)
    pipeline.embed(output_docx, images)
"""

from pathlib import Path
from typing import List, Optional, Union, Dict, Any
from dataclasses import dataclass, field
import json

from .models import ImageBlock
from .extractor import ImageExtractor, ExtractionConfig
from .docx_embedder import DocxImageEmbedder

try:
    from docx import Document
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False


@dataclass
class PipelineConfig:
    """Configuration for the image embedding pipeline"""
    # Extraction settings
    min_image_size: int = 50  # pixels
    max_image_size: int = 2000  # pixels
    skip_duplicates: bool = True

    # Embedding settings
    max_width_ratio: float = 0.8  # 80% of page width
    with_captions: bool = True
    center_images: bool = True

    # Output
    output_format: str = "png"  # png or jpeg

    # Processing
    preserve_order: bool = True  # Maintain original image order


@dataclass
class ProcessingResult:
    """Result of image pipeline processing"""
    success: bool
    images_extracted: int = 0
    images_embedded: int = 0
    source_path: Optional[str] = None
    output_path: Optional[str] = None
    image_blocks: List[ImageBlock] = field(default_factory=list)
    errors: List[str] = field(default_factory=list)
    metadata: Dict[str, Any] = field(default_factory=dict)

    def to_dict(self) -> Dict:
        return {
            "success": self.success,
            "images_extracted": self.images_extracted,
            "images_embedded": self.images_embedded,
            "source_path": self.source_path,
            "output_path": self.output_path,
            "image_count": len(self.image_blocks),
            "errors": self.errors,
            "metadata": self.metadata,
        }


class ImageEmbeddingPipeline:
    """
    Complete pipeline for extracting and embedding images.

    Workflow:
    1. Extract images from source PDF
    2. Store image data in ImageBlocks
    3. Embed images into output DOCX/PDF

    The pipeline preserves:
    - Image order (by page, then by position)
    - Image quality (configurable)
    - Captions (if detected)
    """

    def __init__(self, config: Optional[PipelineConfig] = None):
        """
        Initialize pipeline.

        Args:
            config: Pipeline configuration
        """
        self.config = config or PipelineConfig()

        # Initialize components
        extraction_config = ExtractionConfig(
            min_width=self.config.min_image_size,
            min_height=self.config.min_image_size,
            max_width=self.config.max_image_size,
            max_height=self.config.max_image_size,
            skip_duplicates=self.config.skip_duplicates,
        )
        self.extractor = ImageExtractor(extraction_config)
        self.embedder = DocxImageEmbedder(
            max_width_ratio=self.config.max_width_ratio,
            center_images=self.config.center_images,
        )

    def extract(
        self,
        source_path: Union[str, Path],
        pages: Optional[List[int]] = None
    ) -> List[ImageBlock]:
        """
        Extract images from a PDF file.

        Args:
            source_path: Path to source PDF
            pages: Specific pages to extract from (1-indexed)

        Returns:
            List of ImageBlock objects
        """
        source_path = Path(source_path)

        if not source_path.exists():
            raise FileNotFoundError(f"Source file not found: {source_path}")

        if source_path.suffix.lower() != '.pdf':
            raise ValueError(f"Source must be a PDF file: {source_path}")

        return self.extractor.extract_from_pdf(source_path, pages)

    def embed(
        self,
        document: "Document",
        images: List[ImageBlock],
        with_captions: bool = None
    ) -> int:
        """
        Embed images into a DOCX document.

        Args:
            document: python-docx Document object
            images: List of ImageBlock objects to embed
            with_captions: Include captions (None = use config)

        Returns:
            Number of images embedded
        """
        if not DOCX_AVAILABLE:
            raise ImportError("python-docx is required for DOCX embedding")

        if with_captions is None:
            with_captions = self.config.with_captions

        self.embedder.reset_counter()
        self.embedder.embed_images(document, images, with_captions=with_captions)

        return len(images)

    def process_document(
        self,
        source_pdf: Union[str, Path],
        output_docx: Union[str, Path],
        document: Optional["Document"] = None,
        pages: Optional[List[int]] = None
    ) -> ProcessingResult:
        """
        Complete workflow: extract from PDF, embed in DOCX.

        Args:
            source_pdf: Path to source PDF
            output_docx: Path for output DOCX
            document: Existing Document to add to (None = create new)
            pages: Specific pages to process

        Returns:
            ProcessingResult with statistics
        """
        result = ProcessingResult(
            success=False,
            source_path=str(source_pdf),
            output_path=str(output_docx),
        )

        try:
            # Step 1: Extract images
            images = self.extract(source_pdf, pages)
            result.images_extracted = len(images)
            result.image_blocks = images

            if not images:
                result.success = True
                result.metadata["message"] = "No images found in source PDF"
                return result

            # Step 2: Create or use document
            if document is None:
                if not DOCX_AVAILABLE:
                    raise ImportError("python-docx required")
                document = Document()

            # Step 3: Embed images
            embedded = self.embed(document, images)
            result.images_embedded = embedded

            # Step 4: Save document
            output_docx = Path(output_docx)
            document.save(output_docx)

            result.success = True
            result.metadata["message"] = f"Successfully processed {embedded} images"

        except Exception as e:
            result.errors.append(str(e))
            result.metadata["error_type"] = type(e).__name__

        return result

    def get_image_summary(
        self,
        source_pdf: Union[str, Path]
    ) -> Dict[str, Any]:
        """
        Get a summary of images in a PDF without extracting.

        Args:
            source_pdf: Path to PDF file

        Returns:
            Dictionary with image statistics
        """
        source_pdf = Path(source_pdf)

        counts = self.extractor.get_image_count(source_pdf)
        total = sum(counts.values())

        return {
            "total_images": total,
            "total_pages": len(counts),
            "images_per_page": counts,
            "pages_with_images": [p for p, c in counts.items() if c > 0],
            "empty_pages": [p for p, c in counts.items() if c == 0],
        }


# Convenience functions

def extract_and_embed(
    source_pdf: Union[str, Path],
    output_docx: Union[str, Path],
    with_captions: bool = True,
    min_size: int = 50
) -> ProcessingResult:
    """
    Quick function to extract images from PDF and embed in DOCX.

    Args:
        source_pdf: Path to source PDF
        output_docx: Path for output DOCX
        with_captions: Include figure captions
        min_size: Minimum image size in pixels

    Returns:
        ProcessingResult with statistics

    Example:
        result = extract_and_embed("input.pdf", "output.docx")
        print(f"Embedded {result.images_embedded} images")
    """
    config = PipelineConfig(
        min_image_size=min_size,
        with_captions=with_captions,
    )
    pipeline = ImageEmbeddingPipeline(config)
    return pipeline.process_document(source_pdf, output_docx)


def add_images_to_document(
    document: "Document",
    source_pdf: Union[str, Path],
    with_captions: bool = True
) -> int:
    """
    Add images from a PDF to an existing DOCX document.

    Args:
        document: python-docx Document object
        source_pdf: Path to source PDF
        with_captions: Include figure captions

    Returns:
        Number of images added

    Example:
        from docx import Document
        doc = Document()
        doc.add_heading("My Document", 0)
        doc.add_paragraph("Some text...")

        # Add images from PDF
        count = add_images_to_document(doc, "source.pdf")
        doc.save("output.docx")
    """
    pipeline = ImageEmbeddingPipeline()
    images = pipeline.extract(source_pdf)
    return pipeline.embed(document, images, with_captions=with_captions)


def save_extracted_images(
    source_pdf: Union[str, Path],
    output_dir: Union[str, Path],
    format: str = "png"
) -> List[Path]:
    """
    Extract images from PDF and save as individual files.

    Args:
        source_pdf: Path to source PDF
        output_dir: Directory to save images
        format: Output format (png or jpeg)

    Returns:
        List of saved file paths

    Example:
        paths = save_extracted_images("document.pdf", "images/")
        for path in paths:
            print(f"Saved: {path}")
    """
    from .models import ImageFormat

    output_dir = Path(output_dir)
    output_dir.mkdir(parents=True, exist_ok=True)

    config = PipelineConfig(output_format=format)
    pipeline = ImageEmbeddingPipeline(config)
    images = pipeline.extract(source_pdf)

    saved_paths = []
    for i, img in enumerate(images):
        ext = "jpg" if format == "jpeg" else format
        filename = f"image_{img.source_page:03d}_{i:03d}.{ext}"
        filepath = output_dir / filename

        with open(filepath, "wb") as f:
            f.write(img.image_data)

        saved_paths.append(filepath)

    return saved_paths
