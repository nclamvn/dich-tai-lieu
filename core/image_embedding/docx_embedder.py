"""
DOCX Image Embedder - AI Publisher Pro

Embed images into DOCX documents using python-docx.
Supports auto-sizing, captions, and inline positioning.

Usage:
    embedder = DocxImageEmbedder()

    # Embed single image
    embedder.embed_image(document, image_block)

    # Embed with caption
    embedder.embed_image_with_caption(document, image_block, "Figure 1: Example")

    # Batch embed
    embedder.embed_images(document, image_blocks)
"""

import io
from pathlib import Path
from typing import Optional, List, Union, Tuple

try:
    from docx import Document
    from docx.shared import Inches, Pt, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    PYTHON_DOCX_AVAILABLE = True
except ImportError:
    PYTHON_DOCX_AVAILABLE = False

from .models import ImageBlock


class DocxImageEmbedder:
    """
    Embed images into DOCX documents.

    Features:
    - Auto-resize to fit page width (max 80% by default)
    - Maintain aspect ratio
    - Add captions below images
    - Center alignment
    - Batch embedding
    """

    # Default page dimensions (Letter size in inches)
    DEFAULT_PAGE_WIDTH = 8.5
    DEFAULT_PAGE_MARGIN = 1.0  # Left + Right margin total
    DEFAULT_CONTENT_WIDTH = 6.5  # Usable width

    def __init__(
        self,
        max_width_ratio: float = 0.8,
        default_caption_style: str = "Caption",
        center_images: bool = True
    ):
        """
        Initialize embedder.

        Args:
            max_width_ratio: Max width as ratio of content area (0.0-1.0)
            default_caption_style: Style name for captions
            center_images: Whether to center images by default
        """
        if not PYTHON_DOCX_AVAILABLE:
            raise ImportError(
                "python-docx is required for DOCX embedding. "
                "Install with: pip install python-docx"
            )

        self.max_width_ratio = max_width_ratio
        self.default_caption_style = default_caption_style
        self.center_images = center_images
        self._figure_counter = 0

    def embed_image(
        self,
        document: "Document",
        image_block: ImageBlock,
        width_inches: Optional[float] = None,
        paragraph: Optional["Paragraph"] = None
    ) -> "InlineShape":
        """
        Embed a single image into the document.

        Args:
            document: python-docx Document object
            image_block: ImageBlock with image data
            width_inches: Override width (None = auto-calculate)
            paragraph: Existing paragraph to add to (None = create new)

        Returns:
            InlineShape object representing the embedded image
        """
        # Calculate width if not specified
        if width_inches is None:
            width_inches = self._calculate_width(image_block)

        # Create image stream
        image_stream = io.BytesIO(image_block.image_data)

        # Create or use paragraph
        if paragraph is None:
            paragraph = document.add_paragraph()

        # Center if configured
        if self.center_images:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Add the picture
        run = paragraph.add_run()
        picture = run.add_picture(image_stream, width=Inches(width_inches))

        return picture

    def embed_image_with_caption(
        self,
        document: "Document",
        image_block: ImageBlock,
        caption: Optional[str] = None,
        auto_number: bool = True,
        width_inches: Optional[float] = None
    ) -> Tuple["InlineShape", "Paragraph"]:
        """
        Embed image with a caption below it.

        Args:
            document: python-docx Document object
            image_block: ImageBlock with image data
            caption: Caption text (uses image_block.caption if None)
            auto_number: Add "Figure X:" prefix
            width_inches: Override width

        Returns:
            Tuple of (InlineShape, caption_paragraph)
        """
        # Embed the image
        picture = self.embed_image(document, image_block, width_inches)

        # Determine caption text
        caption_text = caption or image_block.caption or ""

        # Add figure number if requested
        if auto_number and caption_text:
            self._figure_counter += 1
            caption_text = f"Hình {self._figure_counter}: {caption_text}"
        elif auto_number and not caption_text:
            self._figure_counter += 1
            caption_text = f"Hình {self._figure_counter}"

        # Add caption paragraph
        caption_para = document.add_paragraph()
        caption_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Style the caption
        caption_run = caption_para.add_run(caption_text)
        caption_run.italic = True
        caption_run.font.size = Pt(10)

        # Try to apply caption style if it exists
        try:
            caption_para.style = self.default_caption_style
        except KeyError:
            pass  # Style doesn't exist, use manual formatting

        return picture, caption_para

    def embed_images(
        self,
        document: "Document",
        image_blocks: List[ImageBlock],
        with_captions: bool = True,
        add_spacing: bool = True
    ) -> List["InlineShape"]:
        """
        Embed multiple images into the document.

        Args:
            document: python-docx Document object
            image_blocks: List of ImageBlock objects
            with_captions: Include captions for each image
            add_spacing: Add paragraph spacing between images

        Returns:
            List of InlineShape objects
        """
        pictures = []

        for i, image_block in enumerate(image_blocks):
            if with_captions:
                picture, _ = self.embed_image_with_caption(document, image_block)
            else:
                picture = self.embed_image(document, image_block)

            pictures.append(picture)

            # Add spacing between images
            if add_spacing and i < len(image_blocks) - 1:
                spacing_para = document.add_paragraph()
                spacing_para.paragraph_format.space_after = Pt(12)

        return pictures

    def embed_at_position(
        self,
        document: "Document",
        image_block: ImageBlock,
        after_paragraph_index: int,
        with_caption: bool = True
    ) -> "InlineShape":
        """
        Embed image at a specific position in the document.

        Args:
            document: python-docx Document object
            image_block: ImageBlock with image data
            after_paragraph_index: Insert after this paragraph index
            with_caption: Include caption

        Returns:
            InlineShape object
        """
        # Get the paragraph to insert after
        paragraphs = document.paragraphs

        if after_paragraph_index >= len(paragraphs):
            # Append to end
            if with_caption:
                picture, _ = self.embed_image_with_caption(document, image_block)
            else:
                picture = self.embed_image(document, image_block)
            return picture

        # Insert after specified paragraph
        target_para = paragraphs[after_paragraph_index]

        # Create new paragraph for image
        new_para = self._insert_paragraph_after(target_para)

        # Center the paragraph
        if self.center_images:
            new_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Calculate width
        width_inches = self._calculate_width(image_block)

        # Create image stream and add picture
        image_stream = io.BytesIO(image_block.image_data)
        run = new_para.add_run()
        picture = run.add_picture(image_stream, width=Inches(width_inches))

        # Add caption if requested
        if with_caption:
            caption_text = image_block.caption or ""
            if caption_text:
                self._figure_counter += 1
                caption_text = f"Hình {self._figure_counter}: {caption_text}"

                caption_para = self._insert_paragraph_after(new_para)
                caption_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                caption_run = caption_para.add_run(caption_text)
                caption_run.italic = True
                caption_run.font.size = Pt(10)

        return picture

    def _calculate_width(self, image_block: ImageBlock) -> float:
        """
        Calculate optimal width for image.

        Args:
            image_block: ImageBlock with dimensions

        Returns:
            Width in inches
        """
        # Max width based on content area and ratio
        max_width = self.DEFAULT_CONTENT_WIDTH * self.max_width_ratio

        # Use image's specified max width if smaller
        if image_block.max_width_inches < max_width:
            max_width = image_block.max_width_inches

        # Calculate width based on original dimensions
        if image_block.width_px > 0:
            # Assume 96 DPI for screen images
            original_width_inches = image_block.width_px / 96.0

            # Use smaller of original or max
            width = min(original_width_inches, max_width)
        else:
            width = max_width

        return width

    def _insert_paragraph_after(self, paragraph) -> "Paragraph":
        """
        Insert a new paragraph after the given one.

        Args:
            paragraph: Paragraph to insert after

        Returns:
            New paragraph object
        """
        new_p = OxmlElement("w:p")
        paragraph._p.addnext(new_p)

        # Get the new paragraph object
        new_para = paragraph._element.getnext()

        # Create a Paragraph wrapper
        from docx.text.paragraph import Paragraph
        return Paragraph(new_para, paragraph._parent)

    def reset_counter(self):
        """Reset the figure counter (call at start of new document)"""
        self._figure_counter = 0

    def set_figure_counter(self, value: int):
        """Set the figure counter to a specific value"""
        self._figure_counter = value


def create_document_with_images(
    image_blocks: List[ImageBlock],
    output_path: Union[str, Path],
    title: Optional[str] = None,
    with_captions: bool = True
) -> Path:
    """
    Create a new DOCX document with embedded images.

    Args:
        image_blocks: List of ImageBlock objects
        output_path: Output file path
        title: Optional document title
        with_captions: Include captions

    Returns:
        Path to created document

    Example:
        images = extract_images("source.pdf")
        create_document_with_images(images, "output.docx", title="Extracted Images")
    """
    doc = Document()

    # Add title if provided
    if title:
        title_para = doc.add_heading(title, level=0)
        doc.add_paragraph()  # Spacing

    # Embed images
    embedder = DocxImageEmbedder()
    embedder.embed_images(doc, image_blocks, with_captions=with_captions)

    # Save
    output_path = Path(output_path)
    doc.save(output_path)

    return output_path
