#!/usr/bin/env python3
# -*- coding: utf-8 -*-

"""
Batch Processor - Process translation jobs from queue

Executes translation jobs with priority scheduling, progress tracking,
and automatic error recovery.

Phase 1.5 Refactor:
- Extracted sub-modules to core/batch/ for maintainability
- Job lifecycle management: core.batch.job_handler
- Chunk processing: core.batch.chunk_processor
- Result aggregation: core.batch.result_aggregator
- Progress tracking: core.batch.progress_tracker
"""

import asyncio
import time
import traceback
from pathlib import Path
from typing import Optional, List, Any, Dict, Tuple
from collections.abc import Callable
import httpx
from config.logging_config import get_logger

logger = get_logger(__name__)

# Phase 1.5: Import batch sub-modules
from .batch import (
    JobHandler, JobState, JobResult,
    ChunkProcessor, ChunkResult,
    ResultAggregator, AggregatedResult,
    ProgressTracker, ProgressCallback,
    create_logging_callback,
    # V2 Orchestrator
    BatchOrchestrator, OrchestratorConfig, OrchestratorResult,
)

# Document readers
try:
    from pypdf import PdfReader
    HAS_PDF = True
except ImportError:
    HAS_PDF = False

try:
    from docx import Document
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False

from .job_queue import JobQueue, TranslationJob, JobStatus, JobPriority
from .chunker import SmartChunker
from .cache import TranslationCache  # Legacy cache (keep for compatibility)
from .cache.chunk_cache import ChunkCache  # Phase 5.1: New chunk-level cache
from .cache import CheckpointManager, serialize_translation_result, deserialize_translation_result  # Phase 5.2: Checkpoints
from .validator import QualityValidator
from .glossary import GlossaryManager
from .translator import TranslatorEngine
from .merger import SmartMerger
from .translation_memory import TranslationMemory
# Import from export.py file (not export/ directory)
import sys
from pathlib import Path as _Path
_export_file = _Path(__file__).parent / 'export.py'
import importlib.util
_spec = importlib.util.spec_from_file_location("core.export_file", _export_file)
_export_module = importlib.util.module_from_spec(_spec)
_spec.loader.exec_module(_export_module)
UniversalExporter = _export_module.UniversalExporter

# Phase 1.6 - Academic Presentation Layer
try:
    from .postprocess.academic_vietnamese import AcademicVietnamesePolisher
    from .export.docx_presentation import export_academic_docx
    HAS_ACADEMIC_LAYER = True
except ImportError:
    HAS_ACADEMIC_LAYER = False

# Phase 2.0.2 - Semantic Structure & Academic DOCX Builder
try:
    from .structure.semantic_extractor import extract_semantic_structure
    from .export.docx_academic_builder import build_academic_docx, AcademicLayoutConfig
    HAS_PHASE202 = True
except ImportError:
    HAS_PHASE202 = False

# Phase 2.0.7 - Professional Post-Formatting (Typography & Polish)
try:
    from .post_formatting.professional_formatter import (
        ProfessionalFormatter,
        PostFormattingConfig
    )
    HAS_PHASE207 = True
except ImportError:
    HAS_PHASE207 = False

# Phase 2.1.0 - arXiv LaTeX Source Extraction
try:
    from .arxiv_integration import create_arxiv_integration, enrich_docnodes_with_latex
    HAS_PHASE210 = True
    logger.info(f"✅ Phase 2.1.0 modules imported successfully")
except ImportError as e:
    HAS_PHASE210 = False
    logger.info(f"⚠️  Phase 2.1.0 import failed: {e}")

# Phase 3.2 - Book DOCX Builder (Commercial Translation)
try:
    from .export.docx_book_builder import build_book_docx, BookLayoutConfig
    HAS_PHASE31 = True
except ImportError:
    HAS_PHASE31 = False

# Phase 3.3 - Book Paragraph Merging Engine
try:
    from .post_formatting.paragraph_merger import merge_paragraphs_for_book, ParagraphMergeConfig
    HAS_PHASE33 = True
except ImportError:
    HAS_PHASE33 = False

# Phase ADN - Content DNA Extraction (Agent #1 Output)
try:
    from .adn import ADNExtractor, ContentADN
    HAS_ADN = True
    logger.info("✅ ADN module imported successfully")
except ImportError as e:
    HAS_ADN = False
    logger.info(f"⚠️  ADN module not available: {e}")


# OCR Language Mapping for PaddleOCR
OCR_LANG_MAP = {
    'en': 'en',
    'vi': 'vi',
    'ja': 'japan',
    'jp': 'japan',
    'japan': 'japan',
    'japanese': 'japan',
    'zh': 'ch',
    'chinese': 'ch',
    'ko': 'korean',
    'korean': 'korean',
    'fr': 'french',
    'de': 'german',
    'es': 'es',
    'pt': 'pt',
    'ru': 'ru',
    'ar': 'ar',
    'auto': 'en',  # Default to English for auto
}


def get_ocr_lang(source_lang: str) -> str:
    """Map source language to PaddleOCR language code."""
    return OCR_LANG_MAP.get(source_lang.lower() if source_lang else 'en', 'en')


def read_document(file_path: Path) -> str:
    """
    Read document content from various formats

    Args:
        file_path: Path to input document

    Returns:
        Extracted text content

    Raises:
        ValueError: If file format is unsupported
    """
    suffix = file_path.suffix.lower()

    # PDF files
    if suffix == '.pdf':
        if not HAS_PDF:
            raise ValueError("pypdf not installed. Install with: pip install pypdf")

        reader = PdfReader(file_path)
        text_parts = []
        for page in reader.pages:
            text_parts.append(page.extract_text())
        return '\n\n'.join(text_parts)

    # DOCX files
    elif suffix == '.docx':
        if not HAS_DOCX:
            raise ValueError("python-docx not installed. Install with: pip install python-docx")

        doc = Document(file_path)
        text_parts = []
        for para in doc.paragraphs:
            text_parts.append(para.text)
        return '\n\n'.join(text_parts)

    # Text files (TXT, SRT, etc.)
    elif suffix in ['.txt', '.srt', '.md', '.rst']:
        # Try UTF-8 first, fallback to other encodings
        encodings = ['utf-8', 'utf-8-sig', 'latin-1', 'cp1252']
        for encoding in encodings:
            try:
                return file_path.read_text(encoding=encoding)
            except UnicodeDecodeError:
                continue
        raise ValueError(f"Could not decode {file_path.name} with any known encoding")

    else:
        raise ValueError(f"Unsupported file format: {suffix}")


def _merge_equation_blocks(paragraphs: List[str]) -> List[str]:
    """
    Merge consecutive equation fragments into complete equation blocks.

    Phase 2.0.4: Fixes OMML fragmentation where multi-line equations are split
    into separate paragraphs, causing each fragment to get its own OMML tag.

    Uses same heuristics as semantic_extractor._detect_equation_block():
    - Display math delimiters: $$, \\[...\\]
    - Short text (<200 chars) with >20% math symbols

    Args:
        paragraphs: List of paragraph strings

    Returns:
        List with consecutive equation fragments merged

    Example:
        Input:  ["∑", "j=1", "n", "f(jd)", "Regular text"]
        Output: ["∑\\nj=1\\nn\\nf(jd)", "Regular text"]
    """
    def _is_equation_fragment(text: str) -> bool:
        """Detect if text is likely an equation fragment"""
        # Check for display math delimiters
        if '$$' in text:
            return True
        if '\\[' in text and '\\]' in text:
            return True

        # Check if very short and math-heavy
        if len(text) < 200:
            # Count math symbols (same set as semantic_extractor)
            math_chars = sum(1 for c in text if c in r'\{}[]^_=+-*/<>≤≥≠∈∉⊂⊃∩∪∀∃∞∑∏∫')
            if len(text) > 0 and math_chars / len(text) > 0.2:  # >20% math symbols
                return True

        return False

    merged = []
    equation_buffer = []

    for para in paragraphs:
        if _is_equation_fragment(para):
            # Accumulate equation fragments
            equation_buffer.append(para)
        else:
            # Hit a non-equation paragraph
            if equation_buffer:
                # Merge buffered equation fragments with newlines
                merged.append('\n'.join(equation_buffer))
                equation_buffer = []
            merged.append(para)

    # Don't forget final buffered equations
    if equation_buffer:
        merged.append('\n'.join(equation_buffer))

    return merged


class BatchProcessor:
    """Process translation jobs from queue with priority scheduling"""

    def __init__(
        self,
        queue: Optional[JobQueue] = None,
        max_concurrent_jobs: int = 1,
        auto_start: bool = False,
        websocket_manager: Optional[Any] = None
    ):
        """
        Initialize batch processor

        Args:
            queue: Job queue (creates new if None)
            max_concurrent_jobs: Maximum number of concurrent jobs
            auto_start: Auto-start processing on init
        """
        self.queue = queue or JobQueue()
        self.max_concurrent_jobs = max_concurrent_jobs
        self.is_running = False
        self.current_jobs: List[str] = []
        self.background_tasks: List[asyncio.Task] = []  # Track all background tasks
        self.websocket_manager = websocket_manager  # For realtime progress broadcast

        # Phase 5.2: Initialize checkpoint manager
        from config.settings import settings
        if settings.checkpoint_enabled:
            self.checkpoint_manager = CheckpointManager(
                db_path=settings.checkpoint_dir / "checkpoints.db"
            )
            logger.info(f" Checkpoints enabled (DB: {settings.checkpoint_dir / 'checkpoints.db'})")
        else:
            self.checkpoint_manager = None
            logger.warning(" Checkpoints disabled")

        # V2 Orchestrator (opt-in via enable_v2())
        self._orchestrator: Optional[BatchOrchestrator] = None
        self._use_v2 = False

        if auto_start:
            asyncio.create_task(self.start())

    async def start(self, continuous: bool = True):
        """
        Start processing jobs from queue

        Args:
            continuous: If True, keep processing until stopped
        """
        self.is_running = True
        logger.info("🚀 Batch Processor started")

        while self.is_running:
            # Check if we can process more jobs
            if len(self.current_jobs) >= self.max_concurrent_jobs:
                await asyncio.sleep(1)
                continue

            # Get next job from queue
            job = self.queue.get_next_job()

            if job:
                logger.info(f"\n📋 Processing job: {job.job_name} (ID: {job.job_id})")
                logger.info(f"  Priority: {job.priority} | Status: {job.status}")

                # Process job in background with exception handling
                task = asyncio.create_task(self._process_job(job))
                task.add_done_callback(self._handle_task_exception)

                # Track this task so we can cancel it later
                self.background_tasks.append(task)

            else:
                # No jobs available
                if not continuous:
                    logger.info(" No more jobs in queue. Stopping.")
                    break

                # Wait before checking again
                await asyncio.sleep(2)

        self.is_running = False
        logger.info(" Batch Processor stopped")

    def stop(self):
        """Stop processing jobs and cancel all background tasks"""
        self.is_running = False

        # Cancel all running background tasks
        logger.info(f" Cancelling {len(self.background_tasks)} background tasks...")
        for task in self.background_tasks:
            if not task.done():
                task.cancel()
                logger.info(f" Cancelled task")

        # Clear the list
        self.background_tasks.clear()
        logger.info(f" All background tasks cancelled")

    def _handle_task_exception(self, task: asyncio.Task):
        """Handle exceptions from background tasks"""
        try:
            # Get exception if task failed
            exception = task.exception()
            if exception:
                logger.error(f" CRITICAL: Background task failed with exception:")
                logger.info(f"  {type(exception).__name__}: {str(exception)}")
                import traceback
                logger.info(f"  Traceback: {traceback.format_exception(type(exception), exception, exception.__traceback__)}")
        except asyncio.CancelledError:
            logger.info(f" Background task was cancelled")
        except Exception as e:
            logger.error(f" Error retrieving task exception: {str(e)}")
        finally:
            # Remove completed/failed task from list to prevent memory leak
            if task in self.background_tasks:
                self.background_tasks.remove(task)

    # =========================================================================
    # V2 Orchestrator Integration
    # =========================================================================

    def enable_v2(self, enable: bool = True):
        """
        Enable V2 processing using BatchOrchestrator.

        V2 uses modular batch/ sub-modules instead of monolithic _process_job_impl.
        This is an incremental migration strategy - V1 remains as fallback.

        Args:
            enable: Whether to enable V2 processing
        """
        self._use_v2 = enable
        logger.info(f"V2 orchestrator {'enabled' if enable else 'disabled'}")

    def is_v2_enabled(self) -> bool:
        """Check if V2 processing is enabled."""
        return self._use_v2

    # =========================================================================
    # Phase 1.5: Helper methods extracted from _process_job_impl
    # =========================================================================

    async def _load_input_content(
        self,
        job: TranslationJob,
        input_path: Path
    ) -> Tuple[str, bool, Dict[str, Any]]:
        """
        Load input content from file, with optional OCR.

        Args:
            job: Translation job
            input_path: Path to input file

        Returns:
            Tuple of (text_content, ocr_used, ocr_stats)
        """
        input_type = job.metadata.get('input_type', 'native_pdf')
        enable_ocr = job.metadata.get('enable_ocr', False)
        ocr_mode = job.metadata.get('ocr_mode', 'auto')
        mathpix_app_id = job.metadata.get('mathpix_app_id')
        mathpix_app_key = job.metadata.get('mathpix_app_key')

        input_text = None
        ocr_used = False
        ocr_stats = {}

        # Smart detection for auto mode
        if input_type == 'native_pdf' or (enable_ocr and input_path.suffix.lower() == '.pdf'):
            try:
                from .ocr import SmartDetector, PDFType
                detector = SmartDetector()
                detection_result = detector.detect_pdf_type(input_path)

                if input_type == 'native_pdf' and detection_result.pdf_type in [PDFType.SCANNED, PDFType.MIXED]:
                    logger.info(f"Smart Detection: PDF is {detection_result.pdf_type.value}")
                    # AUTO-ENABLE OCR when SmartDetector recommends it
                    if detection_result.ocr_needed:
                        logger.info(f"Auto-enabling OCR based on SmartDetector recommendation")
                        enable_ocr = True
                        input_type = 'scanned_pdf'  # Update type so OCR block runs
                        if ocr_mode == 'auto':
                            ocr_mode = detection_result.recommendation.value if detection_result.recommendation else 'paddle'
                        logger.info(f"OCR mode set to: {ocr_mode}")
            except ImportError:
                logger.warning(f"Smart detection not available")
            except Exception as e:
                logger.warning(f"Smart detection failed: {e}")

        # Perform OCR if needed
        if enable_ocr and input_type in ['scanned_pdf', 'handwritten_pdf']:
            # Determine OCR language from source_lang
            ocr_lang = get_ocr_lang(job.source_lang)
            logger.info(f"OCR mode: {ocr_mode} for {input_type}, language: {ocr_lang}")
            try:
                from .ocr import OcrPipeline, PaddleOcrClient, HybridOcrClient, MathPixOcrClient

                ocr_client = None
                if ocr_mode == 'hybrid':
                    try:
                        ocr_client = HybridOcrClient(
                            paddle_lang=ocr_lang,
                            mathpix_app_id=mathpix_app_id,
                            mathpix_app_key=mathpix_app_key
                        )
                    except Exception:
                        ocr_client = PaddleOcrClient(lang=ocr_lang)
                elif ocr_mode == 'mathpix':
                    try:
                        ocr_client = MathPixOcrClient(
                            app_id=mathpix_app_id,
                            app_key=mathpix_app_key
                        )
                    except Exception:
                        ocr_client = PaddleOcrClient(lang=ocr_lang)
                else:
                    ocr_client = PaddleOcrClient(lang=ocr_lang)

                if ocr_client:
                    pipeline = OcrPipeline(ocr_client, dpi=150)  # Reduced for faster processing
                    ocr_processing_mode = 'handwriting' if input_type == 'handwritten_pdf' else 'document'
                    ocr_pages = pipeline.process_pdf(input_path, mode=ocr_processing_mode)
                    input_text = pipeline.merge_pages_to_text(ocr_pages)
                    ocr_used = True
                    total_confidence = sum(p.confidence for p in ocr_pages) / len(ocr_pages) if ocr_pages else 0
                    ocr_stats = {
                        'total_pages': len(ocr_pages),
                        'avg_confidence': total_confidence,
                        'mode': ocr_mode,
                        'engine': ocr_client.__class__.__name__
                    }
                    logger.info(f"OCR completed: {len(ocr_pages)} pages, confidence: {total_confidence:.2%}")
            except ImportError as e:
                logger.error(f"OCR dependencies not installed: {e}")
                input_text = read_document(input_path)
            except Exception as e:
                logger.error(f"OCR failed: {e}")
                input_text = read_document(input_path)

        # Read document normally if OCR not used
        if input_text is None:
            input_text = read_document(input_path)

        logger.info(f"Loaded input: {len(input_text)} characters ({input_path.suffix})")
        return input_text, ocr_used, ocr_stats

    def _initialize_translator(
        self,
        job: TranslationJob,
        glossary_mgr: Optional[GlossaryManager],
        cache: TranslationCache,
        chunk_cache: Optional[Any],
        validator: QualityValidator,
        tm: TranslationMemory,
        pipeline_mode: str,
        domain: str,
    ) -> TranslatorEngine:
        """
        Initialize translator engine with all dependencies.

        Returns:
            Configured TranslatorEngine
        """
        from config.settings import settings
        api_key = settings.openai_api_key if job.provider == "openai" else settings.anthropic_api_key

        if not api_key or len(api_key.strip()) < 10:
            raise ValueError(f"Invalid or missing API key for provider '{job.provider}'")

        logger.info(f"API key validated for provider: {job.provider}")

        return TranslatorEngine(
            provider=job.provider,
            model=job.model,
            api_key=api_key,
            source_lang=job.source_lang,
            target_lang=job.target_lang,
            glossary_mgr=glossary_mgr,
            cache=cache,
            validator=validator,
            tm=tm,
            tm_fuzzy_threshold=0.85,
            chunk_cache=chunk_cache,
            mode=pipeline_mode,
            domain=domain
        )

    def _preprocess_stem(
        self,
        input_text: str,
        job: TranslationJob,
    ) -> Tuple[str, Optional[Any], List[Any], List[Any]]:
        """
        Preprocess STEM content (formulas, code).

        Args:
            input_text: Original input text
            job: Translation job

        Returns:
            Tuple of (text_to_chunk, stem_preprocessed, formula_matches, code_matches)
        """
        from .stem import FormulaDetector, CodeDetector, PlaceholderManager

        formula_detector = FormulaDetector()
        code_detector = CodeDetector()
        placeholder_manager = PlaceholderManager()

        enable_chemical = job.metadata.get('enable_chemical_formulas', True)
        formula_matches = formula_detector.detect_formulas(input_text, include_chemical=enable_chemical)
        code_matches = code_detector.detect_code(input_text)

        if enable_chemical:
            chemical_count = len([f for f in formula_matches if f.formula_type.value == 'chemical'])
            if chemical_count > 0:
                logger.info(f"Detected {chemical_count} chemical formulas")

        logger.info(f"Detected: {len(formula_matches)} formulas, {len(code_matches)} code blocks")

        stem_preprocessed = placeholder_manager.preprocess(
            text=input_text,
            formula_matches=formula_matches,
            code_matches=code_matches
        )

        job.metadata['stem_preprocessed'] = stem_preprocessed.to_dict()
        logger.info(f"Created {len(stem_preprocessed.mapping)} placeholders")

        return stem_preprocessed.text, stem_preprocessed, formula_matches, code_matches

    def _calculate_job_stats(
        self,
        results: List[Any],
        chunks: List[Any],
    ) -> Tuple[float, float]:
        """
        Calculate job statistics.

        Returns:
            Tuple of (avg_quality, estimated_cost)
        """
        quality_scores = [r.quality_score for r in results if r.quality_score > 0]
        avg_quality = sum(quality_scores) / len(quality_scores) if quality_scores else 0.0

        total_chars = sum(len(chunk.text) for chunk in chunks)
        est_cost = (total_chars / 1000) * 0.01  # ~$0.01 per 1000 chars

        return avg_quality, est_cost

    def _extract_adn(
        self,
        segments: List[str],
        source_lang: str,
        target_lang: str,
        document_type: str = "unknown",
        glossary: Optional[Dict[str, str]] = None,
    ) -> Optional['ContentADN']:
        """
        Extract Content ADN (DNA) from document segments.

        Phase ADN: Agent #1 output for Agent #2 consumption.

        Args:
            segments: List of text segments (paragraphs/chunks)
            source_lang: Source language code
            target_lang: Target language code
            document_type: Type of document (book, article, report, etc.)
            glossary: Optional glossary dict {original: translation}

        Returns:
            ContentADN object or None if extraction fails/not available
        """
        if not HAS_ADN:
            logger.warning("ADN extraction skipped: module not available")
            return None

        try:
            extractor = ADNExtractor(
                source_lang=source_lang,
                target_lang=target_lang,
                glossary=glossary or {},
            )
            adn = extractor.extract(segments, document_type)
            logger.info(f"📊 ADN extracted: {len(adn.proper_nouns)} nouns, {len(adn.patterns)} patterns, {len(adn.characters)} characters")
            return adn
        except Exception as e:
            logger.warning(f"ADN extraction failed: {e}")
            return None

    # =========================================================================
    # End of Phase 1.5 helper methods
    # =========================================================================

    async def _process_job(self, job: TranslationJob):
        """
        Process a single translation job with timeout

        Args:
            job: Job to process
        """
        self.current_jobs.append(job.job_id)

        # Set overall timeout for job (2 hours)
        job_timeout = 7200  # 2 hours in seconds

        try:
            # Wrap entire processing in timeout
            # Route to V2 or V1 based on flag
            if self._use_v2:
                await asyncio.wait_for(
                    self._process_job_v2(job),
                    timeout=job_timeout
                )
            else:
                await asyncio.wait_for(
                    self._process_job_impl(job),
                    timeout=job_timeout
                )
        except asyncio.TimeoutError:
            error_msg = f"Job exceeded maximum time limit of {job_timeout/3600:.1f} hours"
            logger.info(f" {error_msg}")
            job.mark_failed(error_msg)
            self.queue.update_job(job)
        except Exception as e:
            error_msg = str(e)

            # Check if this is a cancellation
            if job.cancellation_requested or "cancelled by user" in error_msg.lower():
                logger.info(f" Job cancelled: {job.job_name}")
                job.mark_cancelled()
            else:
                # Regular error handling
                logger.error(f" Job failed: {job.job_name}")
                logger.info(f"  Error: {str(e)}")

                # Check if can retry
                if job.can_retry():
                    job.increment_retry()
                    logger.info(f" Will retry (attempt {job.retry_count}/{job.max_retries})")
                else:
                    full_error = f"{str(e)}\n{traceback.format_exc()}"
                    job.mark_failed(full_error)

            self.queue.update_job(job)

        finally:
            # Remove from current jobs
            if job.job_id in self.current_jobs:
                self.current_jobs.remove(job.job_id)

    async def _process_job_v2(self, job: TranslationJob):
        """
        V2 implementation using BatchOrchestrator.

        This is the new modular implementation that uses batch/ sub-modules.
        Falls back to V1 (_process_job_impl) if orchestrator fails to initialize.

        Args:
            job: Job to process
        """
        logger.info(f"Processing job with V2 orchestrator: {job.job_id}")

        # Mark job as running
        job.mark_started()
        self.queue.update_job(job)

        # Initialize orchestrator lazily if needed
        if self._orchestrator is None:
            try:
                # Get translator and http client from existing setup
                from config.settings import settings
                import httpx

                translator = TranslatorEngine(settings)
                http_client = httpx.AsyncClient(timeout=120.0)

                self._orchestrator = BatchOrchestrator(
                    translator=translator,
                    http_client=http_client,
                    config=OrchestratorConfig(
                        max_workers=settings.parallel_workers,
                        timeout_seconds=settings.job_timeout,
                        enable_validation=settings.enable_validation,
                        enable_cache=settings.cache_enabled,
                    ),
                )
                logger.info("V2 orchestrator initialized")
            except Exception as e:
                logger.error(f"Failed to initialize V2 orchestrator: {e}")
                logger.warning("Falling back to V1 implementation")
                return await self._process_job_impl(job)

        # Determine input
        input_path = Path(job.input_file) if job.input_file else None
        input_text = job.metadata.get('input_text')

        # Determine output path
        output_path = None
        if job.output_file:
            output_path = Path(job.output_file)

        # Create progress callback for WebSocket broadcasting
        async def progress_callback(percentage: float, message: str, data: dict):
            job.progress = int(percentage * 100)
            job.current_phase = message
            self.queue.update_job(job)

            # Broadcast via WebSocket if available
            if self.websocket_manager:
                try:
                    await self.websocket_manager.broadcast({
                        "event": "job_progress",
                        "job_id": job.job_id,
                        "progress": job.progress,
                        "message": message,
                        **data
                    })
                except Exception:
                    pass

        # Process with orchestrator
        result = await self._orchestrator.process(
            input_path=input_path,
            input_text=input_text,
            source_lang=job.source_lang,
            target_lang=job.target_lang,
            domain=job.metadata.get('domain', 'general'),
            output_path=output_path,
            output_format=job.output_format or 'txt',
            progress_callback=progress_callback,
            options=job.metadata,
        )

        # Update job with result
        if result.success:
            job.output_file = str(result.output_path) if result.output_path else None
            job.mark_completed()
            job.metadata['v2_result'] = {
                'chunk_count': result.chunk_count,
                'total_chars': result.total_chars,
                'duration_seconds': result.duration_seconds,
                'quality_score': result.quality_score,
            }
            logger.info(f"V2 job completed: {job.job_id}")
        else:
            raise Exception(result.error or "V2 processing failed")

        self.queue.update_job(job)

    async def _process_job_impl(self, job: TranslationJob):
        """
        Internal implementation of job processing (separated for timeout handling)

        Args:
            job: Job to process
        """
        # Mark job as running
        job.mark_started()
        self.queue.update_job(job)

        # Load input file
        input_path = Path(job.input_file)
        output_path = Path(job.output_file)  # Initialize output_path early for checkpoint saving
        if not input_path.exists():
                raise FileNotFoundError(f"Input file not found: {job.input_file}")

        # Phase 3: Hybrid OCR Integration
        input_type = job.metadata.get('input_type', 'native_pdf')
        enable_ocr = job.metadata.get('enable_ocr', False)
        ocr_mode = job.metadata.get('ocr_mode', 'auto')  # auto, paddle, hybrid, none
        mathpix_app_id = job.metadata.get('mathpix_app_id')  # Per-job override
        mathpix_app_key = job.metadata.get('mathpix_app_key')

        input_text = None
        ocr_used = False

        # Smart detection for auto mode
        if input_type == 'native_pdf' or (enable_ocr and input_path.suffix.lower() == '.pdf'):
            try:
                from .ocr import SmartDetector, PDFType
                detector = SmartDetector()
                detection_result = detector.detect_pdf_type(input_path)

                # Override input_type based on detection
                if input_type == 'native_pdf' and detection_result.pdf_type in [PDFType.SCANNED, PDFType.MIXED]:
                    logger.info(f"🔍 Smart Detection: PDF appears to be {detection_result.pdf_type.value}")
                    logger.info(f"  Recommendation: {detection_result.recommendation.value}")

                    # AUTO-ENABLE OCR when SmartDetector recommends it
                    if detection_result.ocr_needed:
                        logger.info(f"  Auto-enabling OCR based on SmartDetector recommendation")
                        enable_ocr = True
                        input_type = 'scanned_pdf'  # Update type so OCR block runs
                        # Set OCR mode based on recommendation
                        if ocr_mode == 'auto':
                            ocr_mode = detection_result.recommendation.value if detection_result.recommendation else 'paddle'
                        logger.info(f"  OCR mode set to: {ocr_mode}")

            except ImportError:
                logger.warning(f"  Smart detection not available (missing dependencies)")
            except Exception as e:
                logger.warning(f"  Smart detection failed: {str(e)}")

        # Perform OCR if needed
        if enable_ocr and input_type in ['scanned_pdf', 'handwritten_pdf']:
            logger.info(f"  OCR mode: {ocr_mode} for {input_type}")

            try:
                # Import OCR components
                from .ocr import OcrPipeline, PaddleOcrClient, HybridOcrClient, MathPixOcrClient

                # Create OCR client based on mode
                ocr_client = None

                # Determine OCR language from source_lang
                ocr_lang = get_ocr_lang(job.source_lang)
                logger.info(f"  OCR language: {ocr_lang} (from source: {job.source_lang})")

                if ocr_mode == 'hybrid':
                    # Hybrid mode (PaddleOCR + MathPix)
                    try:
                        ocr_client = HybridOcrClient(
                            paddle_lang=ocr_lang,
                            mathpix_app_id=mathpix_app_id,
                            mathpix_app_key=mathpix_app_key
                        )
                        logger.info(f"  🔬 Hybrid OCR: PaddleOCR + MathPix")
                    except Exception as e:
                        logger.warning(f"  Hybrid OCR unavailable: {str(e)}")
                        logger.info(f"  Falling back to PaddleOCR only")
                        ocr_client = PaddleOcrClient(lang=ocr_lang)

                elif ocr_mode == 'mathpix':
                    # MathPix only (for formula-heavy documents)
                    try:
                        ocr_client = MathPixOcrClient(
                            app_id=mathpix_app_id,
                            app_key=mathpix_app_key
                        )
                        logger.info(f"  ⚗️  MathPix OCR (formula-specialized)")
                    except Exception as e:
                        logger.warning(f"  MathPix unavailable: {str(e)}")
                        logger.info(f"  Falling back to PaddleOCR")
                        ocr_client = PaddleOcrClient(lang=ocr_lang)

                else:
                    # Default: PaddleOCR only
                    ocr_client = PaddleOcrClient(lang=ocr_lang)
                    logger.info(f"  📄 PaddleOCR (local, free)")

                if ocr_client:
                    # Create OCR pipeline
                    pipeline = OcrPipeline(ocr_client, dpi=150)  # Reduced for faster processing

                    # Determine OCR mode
                    ocr_processing_mode = 'handwriting' if input_type == 'handwritten_pdf' else 'document'

                    # Process PDF
                    logger.info(f"  Processing pages...")
                    ocr_pages = pipeline.process_pdf(
                        input_path,
                        mode=ocr_processing_mode
                    )

                    # Merge results
                    input_text = pipeline.merge_pages_to_text(ocr_pages)
                    ocr_used = True

                    # Store OCR stats in metadata
                    total_confidence = sum(p.confidence for p in ocr_pages) / len(ocr_pages) if ocr_pages else 0
                    job.metadata['ocr_stats'] = {
                        'total_pages': len(ocr_pages),
                        'avg_confidence': total_confidence,
                        'mode': ocr_mode,
                        'engine': ocr_client.__class__.__name__
                    }

                    logger.info(f" OCR completed: {len(ocr_pages)} pages, avg confidence: {total_confidence:.2%}")

            except ImportError as e:
                logger.error(f" OCR dependencies not installed: {str(e)}")
                logger.info(f"  Install with: pip install paddleocr paddlepaddle opencv-python-headless")
                logger.info(f"  Falling back to text extraction...")
                input_text = read_document(input_path)

            except Exception as e:
                logger.error(f" OCR failed: {str(e)}")
                logger.info(f"  Falling back to text extraction...")
                input_text = read_document(input_path)

        # Read document normally if OCR not used
        if input_text is None:
            input_text = read_document(input_path)
        logger.info(f" Loaded input: {len(input_text)} characters ({input_path.suffix})")

        # Check if STEM mode is enabled
        is_stem_mode = (job.domain and job.domain.lower() == 'stem')
        if is_stem_mode:
            logger.info(f" STEM mode enabled - formulas and code will be preserved")

        # Initialize components
        chunker = SmartChunker(
                max_chars=job.chunk_size,
                context_window=500,
                stem_mode=is_stem_mode  # Enable STEM-aware chunking
        )

        # Phase 5.1: Use new ChunkCache for better caching with hash keys
        from config.settings import settings
        if settings.chunk_cache_enabled:
            chunk_cache = ChunkCache(
                db_path=settings.cache_dir / "chunks.db"
            )
            logger.info(f" Chunk cache enabled (DB: {settings.cache_dir / 'chunks.db'})")
        else:
            chunk_cache = None
            logger.warning(" Chunk cache disabled")

        # Legacy cache (fallback for compatibility)
        cache = TranslationCache(
                cache_dir=Path("data/cache"),
                enabled=True
        )

        glossary_mgr = None
        if job.glossary:
                glossary_mgr = GlossaryManager(
                    glossary_dir=Path("glossary"),
                    glossary_name=job.glossary
                )

        validator = QualityValidator()

        tm = TranslationMemory(
                db_path=Path("data/translation_memory/tm.db")
        )

        # Get API key
        from config.settings import settings
        api_key = settings.openai_api_key if job.provider == "openai" else settings.anthropic_api_key

        # Validate API key before starting
        if not api_key or len(api_key.strip()) < 10:
                error_msg = f"Invalid or missing API key for provider '{job.provider}'. Please configure your API key in config/settings.py"
                logger.error(f" {error_msg}")
                raise ValueError(error_msg)

        logger.info(f" API key validated for provider: {job.provider}")

        # Determine pipeline mode for cache key
        pipeline_mode = job.metadata.get('layout_mode', 'simple') if job.metadata else 'simple'
        domain = job.domain

        # Create base translator engine
        base_translator = TranslatorEngine(
                provider=job.provider,
                model=job.model,
                api_key=api_key,
                source_lang=job.source_lang,
                target_lang=job.target_lang,
                glossary_mgr=glossary_mgr,
                cache=cache,
                validator=validator,
                tm=tm,
                tm_fuzzy_threshold=0.85,
                # Phase 5.1: Pass chunk cache and mode parameters
                chunk_cache=chunk_cache,
                mode=pipeline_mode,
                domain=domain
        )

        # Wrap with STEM translator if STEM mode is enabled
        stem_preprocessed = None
        stem_formula_matches = []
        stem_code_matches = []

        if is_stem_mode:
            from .stem import FormulaDetector, CodeDetector, PlaceholderManager

            # Detect formulas and code
            formula_detector = FormulaDetector()
            code_detector = CodeDetector()
            placeholder_manager = PlaceholderManager()

            # Phase 3: Enable chemical formula detection if requested
            enable_chemical = job.metadata.get('enable_chemical_formulas', True)
            stem_formula_matches = formula_detector.detect_formulas(
                input_text,
                include_chemical=enable_chemical
            )
            stem_code_matches = code_detector.detect_code(input_text)

            if enable_chemical:
                chemical_count = len([f for f in stem_formula_matches if f.formula_type.value == 'chemical'])
                if chemical_count > 0:
                    logger.info(f"  ⚗️  Detected {chemical_count} chemical formulas")

            logger.info(f"  Detected: {len(stem_formula_matches)} formulas, {len(stem_code_matches)} code blocks")

            # Preprocess: replace with placeholders
            stem_preprocessed = placeholder_manager.preprocess(
                text=input_text,
                formula_matches=stem_formula_matches,
                code_matches=stem_code_matches
            )

            # Use preprocessed text for chunking and translation
            text_to_chunk = stem_preprocessed.text
            job.metadata['stem_preprocessed'] = stem_preprocessed.to_dict()
            logger.info(f"  Created {len(stem_preprocessed.mapping)} placeholders")
        else:
            text_to_chunk = input_text

        translator = base_translator

        # Create chunks
        chunks = chunker.create_chunks(text_to_chunk)
        job.total_chunks = len(chunks)
        self.queue.update_job(job)
        logger.info(f" Created {len(chunks)} chunks")

        # Phase ADN: Extract Content DNA from source segments
        content_adn = None
        if HAS_ADN and job.metadata.get('enable_adn_extraction', True):
            # Get glossary dict for ADN extraction
            glossary_dict = {}
            if glossary_mgr:
                glossary_dict = glossary_mgr.get_terms()

            # Determine document type from job metadata
            doc_type = job.metadata.get('document_type', 'unknown')
            if job.domain == 'stem':
                doc_type = 'article'
            elif job.metadata.get('layout_mode') == 'book':
                doc_type = 'book'

            # Extract ADN from original input text (pre-STEM processing)
            source_segments = [chunk.text for chunk in chunks]
            content_adn = self._extract_adn(
                segments=source_segments,
                source_lang=job.source_lang,
                target_lang=job.target_lang,
                document_type=doc_type,
                glossary=glossary_dict,
            )

            if content_adn:
                # Store ADN summary in job metadata
                job.metadata['adn_summary'] = {
                    'proper_nouns': len(content_adn.proper_nouns),
                    'characters': len(content_adn.characters),
                    'terms': len(content_adn.terms),
                    'patterns': len(content_adn.patterns),
                }
                self.queue.update_job(job)

        # Phase 5.2: Check for existing checkpoint and resume if possible
        completed_results = {}  # Map of chunk_id -> TranslationResult
        chunks_to_process = chunks.copy()

        if self.checkpoint_manager and self.checkpoint_manager.has_checkpoint(job.job_id):
            checkpoint = self.checkpoint_manager.load_checkpoint(job.job_id)
            if checkpoint:
                logger.info(f" Resuming from checkpoint: {len(checkpoint.completed_chunk_ids)}/{checkpoint.total_chunks} chunks completed")

                # Restore completed results
                # FIX-003: Convert chunk_id from STRING to INT (JSON keys are always strings)
                for chunk_id_str, result_data in checkpoint.results_data.items():
                    chunk_id = int(chunk_id_str)  # Convert string key to int
                    completed_results[chunk_id] = deserialize_translation_result(result_data)

                # Filter out completed chunks
                chunks_to_process = [c for c in chunks if c.id not in checkpoint.completed_chunk_ids]

                logger.info(f" Restored {len(completed_results)} cached results")
                logger.info(f"  → Processing remaining {len(chunks_to_process)} chunks")
        else:
            # Save initial checkpoint before starting translation
            if self.checkpoint_manager:
                self.checkpoint_manager.save_checkpoint(
                    job_id=job.job_id,
                    input_file=str(input_path),
                    output_file=str(output_path),
                    total_chunks=len(chunks),
                    completed_chunk_ids=[],
                    results_data={},
                    job_metadata=job.metadata
                )
                logger.info(f" Initial checkpoint saved")

        # Phase 5.4: Check if streaming mode should be used
        from config.settings import settings
        use_streaming = (
            settings.streaming_enabled and
            len(chunks_to_process) >= settings.streaming_batch_size and
            job.output_format in ['docx', 'pdf', 'txt']  # Streaming for all supported formats
        )

        if use_streaming:
            # Use streaming pipeline for memory efficiency
            logger.info(f" Streaming mode enabled: {len(chunks_to_process)} chunks → batched processing")

            # Initialize all_completed_results for streaming mode (same as parallel mode)
            all_completed_results = completed_results.copy()

            from core.streaming import StreamingBatchProcessor

            streaming_processor = StreamingBatchProcessor(
                batch_size=settings.streaming_batch_size,
                enable_streaming=settings.streaming_broadcast_chunks,
                enable_partial_export=settings.streaming_partial_export,
                websocket_manager=self.websocket_manager
            )

            # Define progress callback for streaming mode
            async def streaming_progress_callback(completed_chunks: int, total_chunks: int, progress: float):
                """Update job progress in database during streaming"""
                job.update_progress(completed_chunks, total_chunks)
                self.queue.update_job(job)
                logger.info(f"Progress: {completed_chunks}/{total_chunks} ({progress*100:.1f}%)")

            # Process in streaming batches
            async with httpx.AsyncClient(timeout=httpx.Timeout(300.0)) as client:
                all_results_list, batch_stats = await streaming_processor.process_streaming(
                    job=job,
                    chunks=chunks_to_process,
                    translator=translator,
                    http_client=client,
                    output_path=output_path,
                    progress_callback=streaming_progress_callback
                )

                # Convert list to results array and merge with completed results
                new_results = []
                for chunk in chunks_to_process:
                    # Find result for this chunk
                    chunk_result = next((r for r in all_results_list if r.chunk_id == chunk.id), None)
                    if chunk_result:
                        new_results.append(chunk_result)
                        all_completed_results[chunk.id] = chunk_result

                # Merge with restored results for final output
                results = []
                for chunk in chunks:
                    if chunk.id in all_completed_results:
                        results.append(all_completed_results[chunk.id])

                logger.info(f" Streaming complete: {batch_stats['batches_processed']} batches")
                logger.info(f"  Memory saved: {batch_stats['memory_saved_bytes'] / 1024 / 1024:.1f} MB")

        else:
            # Use standard parallel processing (existing code path)
            logger.info(f" Starting parallel translation with concurrency: {job.concurrency or 10}")

            # Phase 5.2: Track all completed results for checkpoint updates
            all_completed_results = completed_results.copy()  # Start with restored results

            # Define progress callback for real-time updates
            def progress_callback(completed: int, total: int, quality_score: float = 0.0):
                    """Callback to update job progress in real-time"""
                    # FIX: Use len(all_completed_results) directly - it already includes new results
                    # Previously was double-counting: len(all_completed_results) + completed
                    actual_completed = len(all_completed_results)
                    actual_total = len(chunks)

                    job.update_progress(actual_completed, actual_total)
                    job.tm_hits = translator.tm_exact_matches + translator.tm_fuzzy_matches
                    job.cache_hits = cache.hits if cache else 0
                    self.queue.update_job(job)

                    # Broadcast WebSocket event for realtime updates
                    if self.websocket_manager:
                        asyncio.create_task(self.websocket_manager.broadcast({
                            "event": "job_updated",
                            "job_id": job.job_id,
                            "job_name": job.job_name,
                            "status": job.status,
                            "progress": job.progress,
                            "completed_chunks": actual_completed,
                            "total_chunks": actual_total,
                            "quality_score": quality_score
                        }))

                    # Show progress periodically
                    if completed % 5 == 0 or completed == total:
                        logger.info(f"Progress: {actual_completed}/{actual_total} ({job.progress*100:.1f}%) - Quality: {quality_score:.2f}")

            # Use parallel translation with proper concurrency
            async with httpx.AsyncClient(timeout=httpx.Timeout(300.0)) as client:
                    # Create a modified translate_chunk that works with the existing http_client
                    async def translate_with_client(client_param, chunk):
                        return await translator.translate_chunk(client_param, chunk)

                    # Use ParallelProcessor with progress callback
                    from .parallel import ParallelProcessor
                    processor = ParallelProcessor(
                        max_concurrency=job.concurrency or 10,
                        max_retries=5,
                        timeout=120.0,
                        show_progress=False  # We use our own progress callback
                    )

                    # Track progress
                    completed_count = 0

                    # Wrap translate_chunk to include progress callback and cancellation check
                    async def translate_with_progress(client_param, chunk):
                        nonlocal completed_count
                        # Check for cancellation before processing (fetch fresh job state from queue)
                        fresh_job = self.queue.get_job(job.job_id)
                        if fresh_job and fresh_job.cancellation_requested:
                            raise Exception("Job cancelled by user")
                        result = await translator.translate_chunk(client_param, chunk)
                        completed_count += 1

                        # Phase 5.2: Add result to tracking dict
                        all_completed_results[chunk.id] = result

                        # Phase 5.2: Save checkpoint every N chunks
                        from config.settings import settings
                        if self.checkpoint_manager and completed_count % settings.checkpoint_interval == 0:
                            # Serialize all completed results
                            serialized_results = {
                                cid: serialize_translation_result(res)
                                for cid, res in all_completed_results.items()
                            }
                            self.checkpoint_manager.save_checkpoint(
                                job_id=job.job_id,
                                input_file=str(input_path),
                                output_file=str(output_path),
                                total_chunks=len(chunks),
                                completed_chunk_ids=list(all_completed_results.keys()),
                                results_data=serialized_results,
                                job_metadata=job.metadata
                            )
                            logger.info(f"  💾 Checkpoint saved ({len(all_completed_results)}/{len(chunks)} chunks)")

                        progress_callback(completed_count, len(chunks_to_process), result.quality_score)
                        return result

                    # Process remaining chunks in parallel
                    new_results, stats = await processor.process_all(
                        chunks_to_process,
                        translate_with_progress,
                        http_client=client
                    )

                    # Phase 5.2: Merge new results with restored results
                    # Results must be in original chunk order
                    results = []
                    for chunk in chunks:
                        if chunk.id in all_completed_results:
                            results.append(all_completed_results[chunk.id])
                        else:
                            # This shouldn't happen, but handle gracefully
                            logger.warning(f"  Warning: Missing result for chunk {chunk.id}")
                            # Create empty result as fallback
                            from .validator import TranslationResult
                            results.append(TranslationResult(
                                chunk_id=chunk.id,
                                source=chunk.text,
                                translated="[MISSING]",
                                quality_score=0.0
                            ))

                    # Update failed chunks count
                    job.failed_chunks = stats.failed
                    self.queue.update_job(job)

        # Merge results
        merger = SmartMerger()
        merged_text = merger.merge_translations(results)

        # Restore STEM content if in STEM mode
        if is_stem_mode and stem_preprocessed:
            from .stem import PlaceholderManager
            placeholder_manager = PlaceholderManager()

            # Restore formulas and code from placeholders
            final_text = placeholder_manager.restore(
                translated_text=merged_text,
                mapping=stem_preprocessed.mapping
            )

            # Verify restoration
            verification = placeholder_manager.verify_restoration(
                original_text=input_text,
                restored_text=final_text,
                formula_matches=stem_formula_matches,
                code_matches=stem_code_matches
            )

            logger.info(f" STEM preservation: {verification['preservation_rate']:.1%}")
            if verification['formulas_lost'] > 0 or verification['code_lost'] > 0:
                logger.warning(f"  Lost: {verification['formulas_lost']} formulas, {verification['code_lost']} code blocks")

            # Store verification in job metadata
            job.metadata['stem_verification'] = verification
        else:
            final_text = merged_text

        # Phase 1.6: Academic Vietnamese Polishing (opt-in)
        if job.metadata.get('academic_mode', False) and HAS_ACADEMIC_LAYER:
            logger.info(f"\n📚 Applying academic Vietnamese polishing...")
            polisher = AcademicVietnamesePolisher(glossary_mgr=glossary_mgr if glossary_mgr else None)
            final_text, stats = polisher.polish_with_stats(final_text)
            logger.info(f"  Terms normalized: {stats.terms_normalized}")
            logger.info(f"  Phrases improved: {stats.phrases_improved}")
            logger.info(f"  Total changes: {stats.total_changes}")
            job.metadata['academic_polish_stats'] = {
                'terms_normalized': stats.terms_normalized,
                'phrases_improved': stats.phrases_improved,
                'total_changes': stats.total_changes
            }

        # Phase 3.5a: Translation Quality Polish (Rule-Based Only)
        # DEFAULT: OFF (must be explicitly enabled via metadata)
        # SCOPE: Rule-based polish only - no LLM, no API calls
        # SAFETY: Only for book/general domains, NEVER STEM
        quality_mode = job.metadata.get('translation_quality_mode', 'off')
        if quality_mode != 'off' and not is_stem_mode:
            try:
                logger.info(f"\n🔧 Phase 3.5a: Translation quality polish (mode={quality_mode})...")
                from core.quality import TranslationQualityEngine, TranslationQualityConfig

                # Create config based on domain and mode
                domain = 'book' if job.layout_mode == 'book' else 'general'
                config = TranslationQualityConfig(
                    mode=quality_mode,
                    domain=domain,
                    enable_rule_based_pass=True,
                    enable_llm_rewrite=False  # Phase 3.5b not yet implemented
                )

                # Create engine and polish text
                engine = TranslationQualityEngine(config)

                # Analyze first to get issue counts
                report = engine.analyze(final_text)

                # Apply polish
                final_text = engine.polish(final_text)

                logger.info(f"  ✅ Quality polish complete")
                logger.info(f"  Issues found: {report.issues_found}")
                logger.info(f"  Issues fixed: {report.issues_fixed}")

                # Store stats in metadata
                job.metadata['translation_quality_stats'] = {
                    'mode': quality_mode,
                    'domain': domain,
                    'issues_found': report.issues_found,
                    'issues_fixed': report.issues_fixed,
                    'issue_breakdown': report.issue_breakdown
                }

            except Exception as e:
                # Non-fatal: Log error and continue
                logger.warning(f"  Phase 3.5a failed (non-fatal): {e}")
                job.metadata['translation_quality_error'] = str(e)

        # Phase 3: Quality checking
        if job.metadata.get('enable_quality_check', False):
            from .quality import build_quality_report

            logger.info(f"\n✅ Running quality checker...")
            quality_report = build_quality_report(
                source_text=text_to_chunk,  # Text with placeholders
                translated_text=merged_text,  # Translated text with placeholders
                original_source=input_text if is_stem_mode else None  # Original for STEM checks
            )

            logger.info(f"  Quality check: {'✓ PASS' if quality_report.overall_pass else '✗ FAIL'}")
            logger.info(f"  Length ratio: {quality_report.length_ratio:.2f}")

            if not quality_report.placeholder_consistency_ok:
                logger.warning(f"  Placeholder issues detected")
            if not quality_report.stem_preservation_ok:
                logger.warning(f"  STEM preservation issues detected")

            if quality_report.warnings:
                logger.info(f"  Warnings ({len(quality_report.warnings)}):")
                for warning in quality_report.warnings[:3]:  # Show first 3
                    logger.info(f"    - {warning}")

            # Store quality report in metadata
            job.metadata['quality_report'] = {
                'overall_pass': quality_report.overall_pass,
                'length_ratio': quality_report.length_ratio,
                'warnings_count': len(quality_report.warnings),
                'warnings': quality_report.warnings[:5]  # Store first 5 warnings
            }

        # Calculate stats
        quality_scores = [r.quality_score for r in results if r.quality_score > 0]
        avg_quality = sum(quality_scores) / len(quality_scores) if quality_scores else 0.0

        # Estimate cost (rough estimate: ~$0.01 per 1000 chars for gpt-4o-mini)
        total_chars = sum(len(chunk.text) for chunk in chunks)
        est_cost = (total_chars / 1000) * 0.01

        # Save output (output_path already initialized at line 396)
        output_path.parent.mkdir(exist_ok=True, parents=True)

        # Phase 3: Check output mode for layout preservation
        output_mode = job.metadata.get('output_mode', 'docx_reflow')

        # Export based on format
        if job.output_format == "txt":
                output_path.write_text(final_text, encoding='utf-8')
        else:
                # Phase 3: Advanced layout-aware export
                if output_mode == 'pdf_preserve' and input_path.suffix.lower() == '.pdf':
                    logger.info(f" Output mode: {output_mode} (layout-preserving PDF)")
                    logger.warning(f"  Note: Layout preservation requires LayoutExtractor + PDFReconstructor")
                    logger.info(f"  For now, generating standard PDF...")

                    # TODO Phase 3: Integrate layout extraction and PDF reconstruction
                    # from .stem import LayoutExtractor, PDFReconstructor
                    # extractor = LayoutExtractor()
                    # layout = extractor.extract(input_path)
                    # reconstructor = PDFReconstructor()
                    # reconstructor.rebuild_preserve_layout(layout, translated_blocks, output_path)

                elif output_mode == 'docx_reflow':
                    logger.info(f" Output mode: {output_mode} (clean, editable DOCX)")

                # PHASE 1.7.2: Initialize exporter to None (prevents UnboundLocalError)
                exporter = None
                success = False

                # Phase 2.0.2: Semantic Academic DOCX export (layout_mode == 'academic')
                # Phase 2.0.8: Force academic mode for STEM domain
                layout_mode = job.metadata.get('layout_mode', 'simple')
                if job.domain == 'stem' and job.output_format == 'docx':
                    layout_mode = 'academic'  # Override for STEM
                    logger.debug(f" Phase 2.0.8: STEM domain detected → forcing academic pipeline")

                # Phase 4.4: Debug logging - which export path will be taken
                logger.info(f"🔧 Phase 4.4 DEBUG: layout_mode={layout_mode}, format={job.output_format}, domain={job.domain}")

                if (layout_mode == 'academic' and
                    HAS_PHASE202 and
                    job.output_format == 'docx'):
                    logger.info(f" Using Phase 2.0.2 semantic academic DOCX builder...")
                    try:
                        # Split text into paragraphs (single newline = line break)
                        # Note: Phase 1.9 Polisher merges paragraphs, so we split by '\n' not '\n\n'
                        paragraphs = [p.strip() for p in final_text.split('\n') if p.strip()]

                        # Phase 2.0.4: Merge fragmented equation blocks before semantic extraction
                        # This fixes OMML fragmentation issue where equations are split into single symbols
                        original_count = len(paragraphs)
                        paragraphs = _merge_equation_blocks(paragraphs)
                        if len(paragraphs) < original_count:
                            logger.info(f"• Merged equation fragments: {original_count} → {len(paragraphs)} paragraphs")

                        # Extract semantic structure
                        logger.info(f"• Extracting semantic structure from {len(paragraphs)} paragraphs...")
                        nodes = extract_semantic_structure(paragraphs)
                        logger.info(f"• Detected {len(nodes)} semantic nodes")

                        # Phase 2.1.0: Enrich equation nodes with LaTeX source from arXiv .tar.gz
                        import sys
                        logger.debug(f"FORENSIC: HAS_PHASE210 = {HAS_PHASE210}")
                        logger.debug(f"FORENSIC: type(HAS_PHASE210) = {type(HAS_PHASE210)}")
                        logger.debug(f"FORENSIC: bool(HAS_PHASE210) = {bool(HAS_PHASE210)}")
                        logger.debug(f"FORENSIC: HAS_PHASE210 is True = {HAS_PHASE210 is True}")
                        logger.debug(f"FORENSIC: HAS_PHASE210 == True = {HAS_PHASE210 == True}")
                        if HAS_PHASE210:
                            try:
                                logger.debug(f"FORENSIC: Creating arxiv_integration for: {input_path}")
                                arxiv_integration = create_arxiv_integration(str(input_path))
                                logger.debug(f"FORENSIC: arxiv_integration created: {arxiv_integration is not None}")
                                if arxiv_integration:
                                    is_enabled = arxiv_integration.is_enabled()
                                    logger.debug(f"FORENSIC: is_enabled() = {is_enabled}")
                                if arxiv_integration and arxiv_integration.is_enabled():
                                    enriched_count = enrich_docnodes_with_latex(
                                        doc_nodes=nodes,
                                        arxiv_integration=arxiv_integration,
                                        min_confidence=0.5
                                    )
                                    if enriched_count > 0:
                                        logger.info(f"• Phase 2.1.0: Enriched {enriched_count} equations with LaTeX source")
                                    else:
                                        logger.info(f"• Phase 2.1.0: No equations enriched (confidence threshold not met)")
                                else:
                                    logger.warning(f"  Phase 2.1.0: ArXiv integration not enabled")
                            except Exception as e:
                                import traceback
                                logger.warning(f"  Phase 2.1.0: LaTeX enrichment failed: {e}")
                                logger.info(f"  Traceback: {traceback.format_exc()}")
                                # Non-critical: continue without LaTeX enrichment

                        # Phase 2.0.4: Create config with equation rendering mode
                        equation_mode = job.metadata.get('equation_rendering_mode', 'latex_text')
                        config = AcademicLayoutConfig(equation_rendering_mode=equation_mode)

                        # Log configuration
                        logger.info(f"• Equation rendering: {equation_mode}")

                        # 🔧 FORENSIC: Track execution
                        import sys
                        logger.debug(f"FORENSIC: About to call build_academic_docx()")
                        logger.debug(f"   Module: {build_academic_docx.__module__}")
                        logger.debug(f"   File: {build_academic_docx.__code__.co_filename}")
                        logger.debug(f"   Output: {output_path}")

                        # Build academic DOCX with config (Phase 2.0.8: Capture return value)
                        output_file_path = build_academic_docx(nodes, str(output_path), config)

                        logger.debug(f"FORENSIC: build_academic_docx() COMPLETED")

                        # Phase 2.0.8: Verify output was created
                        from pathlib import Path as PathLib
                        if not output_file_path or not PathLib(output_file_path).exists():
                            raise RuntimeError(f"Phase 2.0.8: Academic builder failed - {output_path} not created")

                        logger.debug(f" Phase 2.0.8: Academic DOCX verified: {output_file_path}")

                        # Phase 4.4: Apply document beautification
                        try:
                            from beautification import beautify_docx
                            from config.settings import Settings

                            settings = Settings()

                            # Extract metadata for beautification
                            title = job.metadata.get('title', '') if hasattr(job, 'metadata') and job.metadata else ''
                            author = job.metadata.get('author', '') if hasattr(job, 'metadata') and job.metadata else ''

                            output_file_path = beautify_docx(
                                output_file_path,
                                title=title,
                                author=author,
                                enable=settings.enable_beautification
                            )
                            if settings.enable_beautification:
                                logger.info(f"✅ Beautification applied: {output_file_path}")
                        except Exception as e:
                            logger.warning(f"Beautification skipped: {e}")

                        # Phase 2.0.7: Apply professional post-formatting
                        if HAS_PHASE207:
                            try:
                                logger.info(f" Applying Phase 2.0.7 professional formatting...")

                                # 🔧 FORENSIC: Track Phase 2.0.7 execution
                                logger.debug(f"FORENSIC: Phase 2.0.7 STARTING")
                                logger.debug(f"   ProfessionalFormatter module: {ProfessionalFormatter.__module__}")
                                logger.debug(f"   File: {ProfessionalFormatter.__init__.__code__.co_filename}")

                                post_config = PostFormattingConfig(
                                    enable_heading_detection=True,
                                    remove_technical_footer=True,
                                    clean_page_headers=True,
                                    fix_unicode_artifacts=True,
                                    paragraph_style='modern'  # Block with spacing
                                )
                                formatter = ProfessionalFormatter(config=post_config)

                                # Phase 2.0.8: Use force_apply and verify success
                                formatting_success = formatter.process_document(
                                    str(output_path),
                                    force_apply=True
                                )

                                if not formatting_success:
                                    raise RuntimeError(f"Phase 2.0.8: Professional formatting failed for {job.job_id}")

                                logger.debug(f"FORENSIC: Phase 2.0.7 COMPLETED SUCCESSFULLY")
                                logger.info(f" Professional typography applied (16/14/12/11pt hierarchy)")
                            except Exception as e:
                                logger.debug(f"FORENSIC: Phase 2.0.7 FAILED: {e}")
                                # Phase 2.0.8: Re-raise for academic pipeline (was non-critical before)
                                raise RuntimeError(f"Phase 2.0.7 formatting is mandatory for STEM documents") from e

                        success = True
                        logger.info(f" Semantic academic formatting applied")
                    except Exception as e:
                        logger.warning(f"  Phase 2.0.2 academic export failed: {e}")
                        logger.info(f"  Falling back to standard export...")
                        success = False

                # Phase 3.2: Book DOCX export (layout_mode == 'book')
                elif (layout_mode == 'book' and
                      HAS_PHASE31 and
                      job.output_format == 'docx'):
                    logger.info(f" Phase 3.2: Using book DOCX builder for layout_mode='book'")
                    try:
                        # Split text into paragraphs
                        paragraphs = [p.strip() for p in final_text.split('\n') if p.strip()]

                        # Phase 3.3: Paragraph Merging Engine
                        # Merge paragraphs that were artificially split mid-sentence during PDF conversion
                        if HAS_PHASE33:
                            logger.info(f"  🔗 Phase 3.3: Merging split paragraphs...")
                            logger.info(f"• Input: {len(paragraphs)} raw paragraphs")
                            paragraphs = merge_paragraphs_for_book(paragraphs)
                            logger.info(f"• Output: {len(paragraphs)} merged paragraphs")
                        else:
                            logger.warning(f"  Phase 3.3 not available - skipping paragraph merging")

                        # Extract semantic structure with book-specific detection
                        logger.info(f"• Extracting semantic structure from {len(paragraphs)} paragraphs...")
                        nodes = extract_semantic_structure(paragraphs)
                        logger.info(f"• Detected {len(nodes)} semantic nodes")

                        # Count book-specific elements for logging
                        chapters = sum(1 for n in nodes if n.node_type.value == 'chapter')
                        blockquotes = sum(1 for n in nodes if n.node_type.value == 'blockquote')
                        scene_breaks = sum(1 for n in nodes if n.node_type.value == 'scene_break')
                        logger.info(f"• Book elements: {chapters} chapters, {blockquotes} blockquotes, {scene_breaks} scene breaks")

                        # Create default book layout config
                        config = BookLayoutConfig()

                        # Phase AST Integration: Route to AST pipeline if flag is enabled
                        from config.settings import settings
                        if settings.use_ast_pipeline:
                            logger.info(f"  🎨 Phase AST: Using NEW AST pipeline for book rendering...")
                            try:
                                from core.rendering.ast_builder import build_book_ast
                                from core.rendering.docx_adapter import render_docx_from_ast

                                # Build Document AST from semantic nodes
                                ast = build_book_ast(nodes, language=job.target_lang or "vi")

                                # Render AST to DOCX
                                render_docx_from_ast(ast, output_path)
                                output_file_path = str(output_path)

                                logger.info(f"  ✅ AST pipeline: Document rendered successfully")
                            except Exception as e:
                                logger.error(f" AST pipeline failed: {e}")
                                logger.info(f"  → Falling back to legacy book builder...")
                                # Fallback to legacy
                                output_file_path = build_book_docx(nodes, str(output_path), config)
                        else:
                            # Legacy pipeline (default)
                            output_file_path = build_book_docx(nodes, str(output_path), config)

                        # Verify output was created
                        from pathlib import Path as PathLib
                        if not output_file_path or not PathLib(output_file_path).exists():
                            raise RuntimeError(f"Phase 3.2: Book builder failed - {output_path} not created")

                        # Phase 4.4: Apply document beautification
                        try:
                            from beautification import beautify_docx
                            from config.settings import Settings

                            settings = Settings()

                            # Extract metadata for beautification
                            title = job.metadata.get('title', '') if hasattr(job, 'metadata') and job.metadata else ''
                            author = job.metadata.get('author', '') if hasattr(job, 'metadata') and job.metadata else ''

                            output_file_path = beautify_docx(
                                output_file_path,
                                title=title,
                                author=author,
                                enable=settings.enable_beautification
                            )
                            if settings.enable_beautification:
                                logger.info(f"✅ Beautification applied: {output_file_path}")
                        except Exception as e:
                            logger.warning(f"Beautification skipped: {e}")

                        # Phase 3.4: Apply commercial ebook polish
                        try:
                            logger.info(f"  ✨ Phase 3.4: Applying commercial ebook polish...")
                            from core.post_formatting.book_polisher import BookPolisher, BookPolishConfig
                            from docx import Document

                            # Load the document
                            document = Document(output_file_path)

                            # Create polisher with default config
                            polisher = BookPolisher(BookPolishConfig())

                            # Apply all 12 polish rules
                            document = polisher.polish(document)

                            # Save polished document
                            document.save(output_file_path)

                            logger.info(f"  ✅ Phase 3.4: Commercial polish complete (typography, spacing, scene breaks)")
                        except Exception as e:
                            logger.warning(f"  Phase 3.4: Polish failed (non-fatal): {e}")
                            # Don't fail the entire job if polish fails - document is still usable

                        success = True
                        logger.info(f" Professional book formatting applied")
                    except Exception as e:
                        logger.warning(f"  Phase 3.2 book export failed: {e}")
                        import traceback
                        logger.info(f"  Traceback: {traceback.format_exc()}")
                        # Phase 3.2: Raise exception rather than fallback
                        raise RuntimeError(f"Phase 3.2: Book pipeline failed for {job.job_id}") from e

                # Phase 1.6: Legacy Academic DOCX export (backward compatible)
                elif (job.metadata.get('academic_mode', False) and
                      HAS_ACADEMIC_LAYER and
                      job.output_format == 'docx'):
                    logger.info(f" Using Phase 1.6 academic DOCX presentation layer...")
                    metadata = {
                        'title': job.job_name,
                        'author': 'Academic Translation',
                        'subject': f"STEM Translation - {job.domain or 'general'}"
                    }
                    try:
                        export_academic_docx(
                            text=final_text,
                            output_path=str(output_path),
                            metadata=metadata
                        )
                        success = True
                        logger.info(f" Legacy academic formatting applied")
                    except Exception as e:
                        logger.warning(f"  Academic export failed: {e}")
                        logger.info(f"  Falling back to standard export...")
                        success = False

                # Use UniversalExporter for standard export (fallback or non-academic mode)
                if not success:
                    # Phase 2.0.8: Prevent fallback for academic pipeline
                    if layout_mode == 'academic':
                        raise RuntimeError(
                            f"Phase 2.0.8: Academic pipeline failed for {job.job_id}. "
                            f"Cannot fallback to UniversalExporter for STEM documents."
                        )

                    exporter = UniversalExporter()
                    metadata = {
                        'title': job.job_name,
                        'source_lang': job.source_lang,
                        'target_lang': job.target_lang,
                        'domain': job.domain or "general"
                    }

                    # Generate primary format (DOCX)
                    success = exporter.export(
                        text=final_text,
                        output_path=str(output_path),
                        format=job.output_format,
                        metadata=metadata
                    )

                    # Phase 4.4: Apply document beautification (professional mode path)
                    if success and job.output_format == 'docx':
                        try:
                            from beautification import beautify_docx
                            from config.settings import Settings

                            settings = Settings()
                            title = job.metadata.get('title', '') if hasattr(job, 'metadata') and job.metadata else ''
                            author = job.metadata.get('author', '') if hasattr(job, 'metadata') and job.metadata else ''

                            output_path_str = beautify_docx(
                                str(output_path),
                                title=title,
                                author=author,
                                enable=settings.enable_beautification
                            )
                            if settings.enable_beautification:
                                logger.info(f"✅ Phase 4.4: Beautification applied to professional mode output: {output_path_str}")
                        except Exception as e:
                            logger.warning(f"Beautification skipped: {e}")

                if not success:
                    # Fallback to txt
                    output_path.write_text(final_text, encoding='utf-8')

                # ALSO generate PDF if DOCX was requested and PDF is supported
                # PHASE 1.7.2: Only if exporter was initialized (not None)
                if exporter is not None and job.output_format == 'docx' and 'pdf' in exporter.supported_formats:
                    pdf_path = output_path.parent / f"{output_path.stem}.pdf"
                    try:
                        success_pdf = exporter.export(
                            text=final_text,
                            output_path=str(pdf_path),
                            format='pdf',
                            metadata=metadata
                        )
                        if success_pdf:
                            logger.info(f" Also saved PDF: {pdf_path}")
                        else:
                            logger.warning(f"  PDF generation failed (exporter returned False)")
                    except Exception as e:
                        import traceback
                        logger.warning(f"  PDF generation error: {e}")
                        logger.info(f"Full traceback:\n{traceback.format_exc()}")

        logger.info(f" Saved primary format: {output_path}")

        # Phase ADN: Save ADN JSON file
        if content_adn:
            try:
                adn_path = output_path.parent / f"{output_path.stem}_adn.json"
                adn_json = content_adn.to_json(indent=2)
                adn_path.write_text(adn_json, encoding='utf-8')
                logger.info(f"📊 Saved ADN file: {adn_path}")

                # Store ADN in job metadata for API response
                job.metadata['adn'] = content_adn.to_dict()
                job.metadata['adn_file'] = str(adn_path)
            except Exception as e:
                logger.warning(f"Failed to save ADN file: {e}")

        # Save cache
        if cache:
                cache.save()

        # Phase 5.2: Delete checkpoint on successful completion
        if self.checkpoint_manager and self.checkpoint_manager.has_checkpoint(job.job_id):
            self.checkpoint_manager.delete_checkpoint(job.job_id)
            logger.info(f"  Checkpoint deleted (job completed successfully)")

        # Mark job as completed
        job.mark_completed(avg_quality=avg_quality, total_cost=est_cost)
        self.queue.update_job(job)

        # Broadcast completion event via WebSocket
        if self.websocket_manager:
            asyncio.create_task(self.websocket_manager.broadcast({
                "event": "job_completed",
                "job_id": job.job_id,
                "job_name": job.job_name,
                "status": "completed",
                "progress": 1.0,
                "quality_score": avg_quality,
                "total_cost_usd": est_cost
            }))

        logger.info(f" Job completed: {job.job_name}")
        logger.info(f"  Quality: {avg_quality:.2f} | Cost: ${est_cost:.4f}")
        logger.info(f"  TM hits: {job.tm_hits} | Cache hits: {job.cache_hits}")

    async def process_single_job(self, job_id: str) -> bool:
        """
        Process a single job by ID

        Args:
            job_id: Job ID to process

        Returns:
            True if successful, False otherwise
        """
        job = self.queue.get_job(job_id)
        if not job:
            logger.error(f" Job not found: {job_id}")
            return False

        if job.status not in [JobStatus.PENDING, JobStatus.QUEUED, JobStatus.RETRYING]:
            logger.error(f" Job cannot be processed (status: {job.status})")
            return False

        await self._process_job(job)
        return job.status == JobStatus.COMPLETED

    def get_status(self) -> dict:
        """Get processor status"""
        stats = self.queue.get_queue_stats()

        return {
            "is_running": self.is_running,
            "current_jobs": len(self.current_jobs),
            "max_concurrent_jobs": self.max_concurrent_jobs,
            "queue_stats": stats
        }


class JobScheduler:
    """Schedule jobs to run at specific times or intervals"""

    def __init__(self, queue: Optional[JobQueue] = None):
        self.queue = queue or JobQueue()
        self.is_running = False

    async def start(self):
        """Start scheduler"""
        self.is_running = True
        logger.info("⏰ Job Scheduler started")

        while self.is_running:
            # Check for scheduled jobs
            current_time = time.time()

            # Get all pending jobs with scheduled_at <= now
            jobs = self.queue.list_jobs(status=JobStatus.PENDING, limit=1000)

            for job in jobs:
                if job.scheduled_at and job.scheduled_at <= current_time:
                    # Mark as queued (ready to process)
                    job.status = JobStatus.QUEUED
                    self.queue.update_job(job)
                    logger.info(f" Scheduled job queued: {job.job_name}")

            # Wait before next check
            await asyncio.sleep(10)

    def stop(self):
        """Stop scheduler"""
        self.is_running = False


async def run_batch_processor(
    queue: Optional[JobQueue] = None,
    max_concurrent_jobs: int = 1,
    enable_scheduler: bool = True
):
    """
    Run batch processor with optional scheduler

    Args:
        queue: Job queue
        max_concurrent_jobs: Max concurrent jobs
        enable_scheduler: Enable job scheduler
    """
    processor = BatchProcessor(queue, max_concurrent_jobs)

    tasks = [processor.start(continuous=True)]

    if enable_scheduler:
        scheduler = JobScheduler(queue)
        tasks.append(scheduler.start())

    try:
        await asyncio.gather(*tasks)
    except KeyboardInterrupt:
        logger.info("\n🛑 Shutting down...")
        processor.stop()
        if enable_scheduler:
            scheduler.stop()
