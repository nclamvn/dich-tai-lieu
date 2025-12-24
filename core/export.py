#!/usr/bin/env python3
# -*- coding: utf-8 -*-

"""
AI Translator Pro - Advanced Export Module
==========================================
Professional document export system supporting multiple formats:
- DOCX with advanced styling and formatting
- PDF with ReportLab (professional layout)
- HTML with embedded styles
- Markdown for documentation
- RTF for compatibility

Author: AI Translator Pro Team
Version: 2.0.0
"""

import os
import re
import json
import base64
import html
from pathlib import Path
from datetime import datetime
from typing import Dict, List, Optional, Tuple, Any
from dataclasses import dataclass, field
from io import BytesIO

# Import MathReconstructor for Unicode normalization
try:
    from .math_reconstructor import MathReconstructor
    MATH_RECONSTRUCTOR_AVAILABLE = True
except ImportError:
    MATH_RECONSTRUCTOR_AVAILABLE = False

# ========== DOCX Export (Primary) ==========
try:
    from docx import Document
    from docx.shared import Pt, Inches, RGBColor, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
    from docx.enum.style import WD_STYLE_TYPE
    from docx.oxml import parse_xml
    from docx.oxml.ns import nsdecls
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    print("⚠️ python-docx not installed. DOCX export limited.")

# ========== PDF Export (ReportLab) ==========
try:
    from reportlab.lib import colors
    from reportlab.lib.pagesizes import A4, letter
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import inch, cm
    from reportlab.lib.enums import TA_JUSTIFY, TA_CENTER
    from reportlab.platypus import (
        SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle,
        PageBreak, Image, KeepTogether, Flowable
    )
    from reportlab.pdfbase import pdfmetrics
    from reportlab.pdfbase.ttfonts import TTFont
    from reportlab.pdfgen import canvas
    REPORTLAB_AVAILABLE = True
except ImportError:
    REPORTLAB_AVAILABLE = False
    print("⚠️ ReportLab not installed. PDF export limited.")

# ========== Alternative PDF (docx2pdf) ==========
try:
    from docx2pdf import convert as docx_to_pdf
    DOCX2PDF_AVAILABLE = True
except ImportError:
    DOCX2PDF_AVAILABLE = False

# ========== XML Sanitization Utility ==========
def sanitize_for_xml(text: str) -> str:
    """
    Remove control characters and NULL bytes for XML compatibility.

    Fixes critical DOCX corruption issue where control characters in translated text
    cause XML serialization errors: "All strings must be XML compatible: Unicode or ASCII,
    no NULL bytes or control characters"

    Args:
        text: Input text that may contain control characters

    Returns:
        Sanitized text safe for XML/DOCX serialization

    Note:
        - Removes NULL bytes (\x00)
        - Removes control characters (\x01-\x1F) except tab (\t), newline (\n), carriage return (\r)
        - Removes DEL character (\x7F)
        - Preserves all valid Unicode and ASCII characters
    """
    if not text:
        return text

    # Remove NULL bytes
    text = text.replace('\x00', '')

    # Remove control characters except tab, newline, carriage return
    # Control characters: \x01-\x08, \x0B-\x0C, \x0E-\x1F, \x7F
    # Keep: \x09 (tab), \x0A (newline), \x0D (carriage return)
    text = re.sub(r'[\x01-\x08\x0B-\x0C\x0E-\x1F\x7F]', '', text)

    return text


# ========== Configuration ==========
@dataclass
class ExportConfig:
    """Export configuration settings"""
    # Document metadata
    title: str = "Translation Document"
    author: str = "AI Translator Pro"
    subject: str = "Automated Translation"
    keywords: List[str] = field(default_factory=lambda: ["translation", "AI"])

    # Formatting
    font_family: str = "Times-Roman"  # ReportLab standard font name
    font_size: int = 12
    line_spacing: float = 1.5
    paragraph_spacing: int = 6
    margin_top: float = 1.0  # inches
    margin_bottom: float = 1.0
    margin_left: float = 1.25
    margin_right: float = 1.25

    # Features
    add_header: bool = True
    add_footer: bool = True
    add_page_numbers: bool = True
    add_toc: bool = False  # Table of contents
    add_watermark: bool = False
    watermark_text: str = "DRAFT"

    # Quality settings
    preserve_formatting: bool = True
    detect_headers: bool = False  # Disabled by default to avoid XML parsing errors with special characters
    smart_quotes: bool = True
    hyphenation: bool = False

    # Export specific
    pdf_compression: bool = True
    pdf_encrypt: bool = False
    pdf_password: str = ""
    embed_fonts: bool = True

    # Phase 1.6 - Academic presentation layer settings
    academic_mode: bool = False  # Enable academic Vietnamese polishing and formatting
    academic_polish_vietnamese: bool = False  # Apply Vietnamese academic prose enhancement
    academic_docx_formatting: bool = False  # Apply academic DOCX presentation
    clean_boilerplate: bool = False  # Remove "AI Translator Pro" footer and metadata
    format_equations: bool = False  # Apply academic equation spacing and formatting
    detect_academic_structure: bool = False  # Detect and format academic structure (headings, references)

# ========== Style Manager ==========
class StyleManager:
    """Manages document styles across formats"""
    
    def __init__(self, config: ExportConfig):
        self.config = config
        self.styles = self._init_styles()
    
    def _init_styles(self) -> Dict:
        """Initialize style definitions"""
        base_font = self.config.font_family
        base_size = self.config.font_size
        
        return {
            'normal': {
                'font': base_font,
                'size': base_size,
                'color': (0, 0, 0),
                'alignment': 'justify',
                'spacing_after': self.config.paragraph_spacing
            },
            'heading1': {
                'font': base_font,
                'size': int(base_size * 1.8),
                'color': (0, 0, 139),  # Dark blue
                'bold': True,
                'alignment': 'left',
                'spacing_before': 12,
                'spacing_after': 6
            },
            'heading2': {
                'font': base_font,
                'size': int(base_size * 1.4),
                'color': (25, 25, 112),  # Midnight blue
                'bold': True,
                'alignment': 'left',
                'spacing_before': 10,
                'spacing_after': 4
            },
            'heading3': {
                'font': base_font,
                'size': int(base_size * 1.2),
                'color': (70, 130, 180),  # Steel blue
                'bold': True,
                'italic': True,
                'alignment': 'left',
                'spacing_before': 8,
                'spacing_after': 3
            },
            'quote': {
                'font': base_font,
                'size': base_size,
                'color': (105, 105, 105),  # Dim gray
                'italic': True,
                'alignment': 'center',
                'indent_left': 36,
                'indent_right': 36,
                'spacing_before': 6,
                'spacing_after': 6
            },
            'code': {
                'font': 'Courier New',
                'size': int(base_size * 0.9),
                'color': (0, 100, 0),  # Dark green
                'background': (245, 245, 245),  # Light gray
                'alignment': 'left'
            },
            'footnote': {
                'font': base_font,
                'size': int(base_size * 0.8),
                'color': (128, 128, 128),  # Gray
                'alignment': 'justify'
            }
        }
    
    def apply_docx_style(self, paragraph, style_name: str):
        """Apply style to DOCX paragraph"""
        if not DOCX_AVAILABLE:
            return
        
        style = self.styles.get(style_name, self.styles['normal'])
        
        # Font settings
        for run in paragraph.runs:
            run.font.name = style['font']
            run.font.size = Pt(style['size'])
            if 'bold' in style:
                run.font.bold = style['bold']
            if 'italic' in style:
                run.font.italic = style['italic']
            if 'color' in style:
                run.font.color.rgb = RGBColor(*style['color'])
        
        # Paragraph settings
        if style['alignment'] == 'justify':
            paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        elif style['alignment'] == 'center':
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif style['alignment'] == 'right':
            paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        else:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        
        if 'spacing_before' in style:
            paragraph.paragraph_format.space_before = Pt(style['spacing_before'])
        if 'spacing_after' in style:
            paragraph.paragraph_format.space_after = Pt(style['spacing_after'])
        if 'indent_left' in style:
            paragraph.paragraph_format.left_indent = Pt(style['indent_left'])
        if 'indent_right' in style:
            paragraph.paragraph_format.right_indent = Pt(style['indent_right'])
    
    def get_reportlab_style(self, style_name: str) -> 'ParagraphStyle':
        """Get ReportLab paragraph style"""
        if not REPORTLAB_AVAILABLE:
            return None
        
        style = self.styles.get(style_name, self.styles['normal'])
        
        alignment_map = {
            'justify': TA_JUSTIFY,
            'center': TA_CENTER,
            'left': 0,
            'right': 2
        }
        
        return ParagraphStyle(
            name=style_name,
            fontName=style['font'].replace(' ', '-'),
            fontSize=style['size'],
            leading=style['size'] * 1.2,
            alignment=alignment_map.get(style['alignment'], 0),
            spaceAfter=style.get('spacing_after', 0),
            spaceBefore=style.get('spacing_before', 0),
            textColor=colors.Color(
                style['color'][0]/255, 
                style['color'][1]/255, 
                style['color'][2]/255
            ) if 'color' in style else colors.black,
            leftIndent=style.get('indent_left', 0),
            rightIndent=style.get('indent_right', 0)
        )

# ========== Document Analyzer ==========
class DocumentAnalyzer:
    """Analyzes document structure for intelligent formatting"""
    
    @staticmethod
    def detect_structure(text: str) -> List[Dict]:
        """Detect document structure (headings, paragraphs, lists, etc.)"""
        lines = text.split('\n')
        structured_content = []
        
        for i, line in enumerate(lines):
            line = line.strip()
            if not line:
                continue
            
            element = {'text': line, 'type': 'paragraph'}
            
            # Detect headings
            if re.match(r'^#{1,3}\s+', line):
                level = len(re.match(r'^(#{1,3})', line).group(1))
                element['type'] = f'heading{level}'
                element['text'] = re.sub(r'^#{1,3}\s+', '', line)
            
            # Detect chapter/section markers
            elif re.match(r'^(Chapter|Section|Part)\s+\d+', line, re.IGNORECASE):
                element['type'] = 'heading1'
            
            # Detect numbered headings
            elif re.match(r'^\d+\.(\d+\.)*\s+[A-Z]', line):
                depth = line.count('.')
                element['type'] = f'heading{min(depth + 1, 3)}'
            
            # Detect lists
            elif re.match(r'^[\*\-\+]\s+', line):
                element['type'] = 'bullet'
                element['text'] = re.sub(r'^[\*\-\+]\s+', '', line)
            elif re.match(r'^\d+\.\s+', line):
                element['type'] = 'numbered'
                element['text'] = re.sub(r'^\d+\.\s+', '', line)
            
            # Detect quotes
            elif line.startswith('"') and line.endswith('"'):
                element['type'] = 'quote'
            
            # Detect code blocks
            elif line.startswith('```'):
                element['type'] = 'code_marker'
            
            # Detect tables (simple)
            elif '|' in line and line.count('|') > 2:
                element['type'] = 'table_row'
            
            # Length-based heading detection
            elif len(line) < 50 and line[0].isupper() and not line.endswith('.'):
                # Short lines starting with capital might be headings
                next_line = lines[i + 1] if i + 1 < len(lines) else ""
                if not next_line or len(next_line) > 100:
                    element['type'] = 'heading3'
            
            structured_content.append(element)
        
        return structured_content

# ========== DOCX Exporter ==========
class DocxExporter:
    """Advanced DOCX export with rich formatting"""
    
    def __init__(self, config: ExportConfig):
        if not DOCX_AVAILABLE:
            raise ImportError("python-docx required for DOCX export")

        self.config = config
        self.style_manager = StyleManager(config)
        self.doc = Document()

        # Initialize Unicode normalizer for STEM content
        if MATH_RECONSTRUCTOR_AVAILABLE:
            self.math_reconstructor = MathReconstructor()
        else:
            self.math_reconstructor = None

        self._setup_document()
    
    def _setup_document(self):
        """Setup document properties and styles"""
        # Set document properties
        core_properties = self.doc.core_properties
        core_properties.title = self.config.title
        core_properties.author = self.config.author
        core_properties.subject = self.config.subject
        core_properties.keywords = ', '.join(self.config.keywords)
        core_properties.created = datetime.now()
        
        # Set margins
        sections = self.doc.sections
        for section in sections:
            section.top_margin = Inches(self.config.margin_top)
            section.bottom_margin = Inches(self.config.margin_bottom)
            section.left_margin = Inches(self.config.margin_left)
            section.right_margin = Inches(self.config.margin_right)
        
        # Add custom styles
        self._create_custom_styles()
    
    def _create_custom_styles(self):
        """Create custom paragraph and character styles"""
        styles = self.doc.styles
        
        # Normal style
        normal_style = styles['Normal']
        normal_style.font.name = self.config.font_family
        normal_style.font.size = Pt(self.config.font_size)
        normal_style.paragraph_format.line_spacing = self.config.line_spacing
        
        # Create heading styles
        for i in range(1, 4):
            heading_name = f'CustomHeading{i}'
            if heading_name not in styles:
                heading_style = styles.add_style(heading_name, WD_STYLE_TYPE.PARAGRAPH)
                style_config = self.style_manager.styles[f'heading{i}']
                heading_style.font.name = style_config['font']
                heading_style.font.size = Pt(style_config['size'])
                heading_style.font.bold = style_config.get('bold', False)
                heading_style.font.color.rgb = RGBColor(*style_config['color'])
    
    def add_header_footer(self):
        """Add header and footer to document"""
        if not self.config.add_header and not self.config.add_footer:
            return
        
        section = self.doc.sections[0]
        
        # Header
        if self.config.add_header:
            header = section.header
            header_para = header.paragraphs[0]
            header_para.text = self.config.title
            header_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            header_para.style.font.size = Pt(10)
            header_para.style.font.color.rgb = RGBColor(128, 128, 128)
        
        # Footer
        if self.config.add_footer:
            footer = section.footer
            footer_para = footer.paragraphs[0]
            if self.config.add_page_numbers:
                footer_para.text = f"{self.config.author} | Page "
                # Add page number field (complex - requires XML manipulation)
                self._add_page_number_field(footer_para)
            else:
                footer_para.text = f"{self.config.author} | {datetime.now().strftime('%Y-%m-%d')}"
            footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            footer_para.style.font.size = Pt(9)
    
    def _add_page_number_field(self, paragraph):
        """Add page number field to paragraph"""
        run = paragraph.add_run()
        fldChar1 = parse_xml(r'<w:fldChar w:fldCharType="begin" xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"/>')
        instrText = parse_xml(r'<w:instrText xml:space="preserve" xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"> PAGE </w:instrText>')
        fldChar2 = parse_xml(r'<w:fldChar w:fldCharType="end" xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"/>')
        
        run._r.append(fldChar1)
        run._r.append(instrText)
        run._r.append(fldChar2)
    
    def add_watermark(self):
        """Add watermark to document"""
        if not self.config.add_watermark:
            return
        
        # Watermark implementation is complex in python-docx
        # This is a simplified version
        section = self.doc.sections[0]
        # Would require XML manipulation for full watermark support
        pass
    
    def add_toc(self):
        """Add table of contents"""
        if not self.config.add_toc:
            return
        
        self.doc.add_heading('Table of Contents', 1)
        # TOC generation would require tracking all headings
        # and their page numbers - complex in python-docx
        self.doc.add_paragraph('(Table of Contents will be updated in Word)')
        self.doc.add_page_break()
    
    def add_content(self, text: str, structured: bool = True):
        """Add content to document with intelligent formatting"""
        # CRITICAL: Normalize Unicode BEFORE adding to document (fixes Erdős, math symbols, etc.)
        if self.math_reconstructor:
            text = self.math_reconstructor.normalize_unicode(text)

        # CRITICAL: Sanitize text for XML compatibility (fixes DOCX corruption from control characters)
        text = sanitize_for_xml(text)

        if structured and self.config.detect_headers:
            # Use structure detection
            elements = DocumentAnalyzer.detect_structure(text)
            self._add_structured_content(elements)
        else:
            # Add as simple paragraphs
            paragraphs = text.split('\n\n')
            for para_text in paragraphs:
                if para_text.strip():
                    p = self.doc.add_paragraph(para_text.strip())
                    self.style_manager.apply_docx_style(p, 'normal')
    
    def _add_structured_content(self, elements: List[Dict]):
        """Add structured content with appropriate formatting"""
        in_code_block = False
        code_content = []
        list_items = []
        
        for element in elements:
            elem_type = element['type']
            elem_text = element['text']
            
            # Handle code blocks
            if elem_type == 'code_marker':
                if in_code_block:
                    # End code block
                    code_text = '\n'.join(code_content)
                    p = self.doc.add_paragraph(code_text)
                    self.style_manager.apply_docx_style(p, 'code')
                    code_content = []
                    in_code_block = False
                else:
                    in_code_block = True
                continue
            
            if in_code_block:
                code_content.append(elem_text)
                continue
            
            # Handle different element types
            if elem_type.startswith('heading'):
                level = int(elem_type[-1])
                heading = self.doc.add_heading(elem_text, level)
                
            elif elem_type == 'bullet':
                p = self.doc.add_paragraph(elem_text, style='List Bullet')
                
            elif elem_type == 'numbered':
                p = self.doc.add_paragraph(elem_text, style='List Number')
                
            elif elem_type == 'quote':
                p = self.doc.add_paragraph(elem_text)
                self.style_manager.apply_docx_style(p, 'quote')
                
            elif elem_type == 'table_row':
                # Simple table handling
                cells = [cell.strip() for cell in elem_text.split('|')]
                # Would need to accumulate rows and create table
                p = self.doc.add_paragraph(elem_text)
                
            else:  # paragraph
                p = self.doc.add_paragraph(elem_text)
                self.style_manager.apply_docx_style(p, 'normal')
    
    def save(self, filename: str):
        """Save document to file"""
        # Add header/footer
        self.add_header_footer()
        
        # Add watermark if configured
        self.add_watermark()
        
        # Save
        self.doc.save(filename)

# ========== PDF Exporter (ReportLab) ==========
class PdfExporter:
    """Professional PDF export using ReportLab"""
    
    def __init__(self, config: ExportConfig):
        if not REPORTLAB_AVAILABLE:
            raise ImportError("ReportLab required for PDF export")
        
        self.config = config
        self.style_manager = StyleManager(config)
        self.story = []
        self.styles = getSampleStyleSheet()
        self._register_fonts()
        self._setup_custom_styles()
    
    def _register_fonts(self):
        """Register custom fonts with Vietnamese Unicode support (Phase 4.4 FIX)"""
        import platform

        # Phase 4.4 FIX: Support both macOS and Linux font paths
        # macOS fonts support Vietnamese Unicode (U+00C0-U+1EF9)
        macos_fonts = {
            'Times-Roman': '/System/Library/Fonts/Supplemental/Times New Roman.ttf',
            'Times-Bold': '/System/Library/Fonts/Supplemental/Times New Roman Bold.ttf',
            'Helvetica': '/System/Library/Fonts/Supplemental/Arial.ttf',
            'Courier': '/System/Library/Fonts/Supplemental/Courier New.ttf'
        }

        # Linux Liberation fonts (fallback)
        linux_fonts = {
            'Times-Roman': '/usr/share/fonts/liberation/LiberationSerif-Regular.ttf',
            'Times-Bold': '/usr/share/fonts/liberation/LiberationSerif-Bold.ttf',
            'Helvetica': '/usr/share/fonts/liberation/LiberationSans-Regular.ttf',
            'Courier': '/usr/share/fonts/liberation/LiberationMono-Regular.ttf'
        }

        # Choose font set based on platform
        if platform.system() == 'Darwin':  # macOS
            font_paths = macos_fonts
            fallback_paths = linux_fonts
        else:  # Linux
            font_paths = linux_fonts
            fallback_paths = macos_fonts

        # Try primary font paths first, then fallback
        for font_name in font_paths.keys():
            registered = False

            # Try primary path
            if Path(font_paths[font_name]).exists():
                try:
                    pdfmetrics.registerFont(TTFont(font_name, font_paths[font_name]))
                    registered = True
                except Exception:
                    pass

            # Try fallback path if primary failed
            if not registered and font_name in fallback_paths:
                if Path(fallback_paths[font_name]).exists():
                    try:
                        pdfmetrics.registerFont(TTFont(font_name, fallback_paths[font_name]))
                    except Exception:
                        pass
    
    def _setup_custom_styles(self):
        """Setup custom paragraph styles"""
        # Add custom styles from StyleManager
        for style_name in ['normal', 'heading1', 'heading2', 'heading3', 'quote', 'code']:
            custom_style = self.style_manager.get_reportlab_style(style_name)
            if custom_style:
                self.styles.add(custom_style)
    
    def _create_header_footer(self, canvas, doc):
        """Create header and footer for each page"""
        canvas.saveState()
        
        # Header
        if self.config.add_header:
            canvas.setFont('Times-Roman', 9)
            canvas.setFillGray(0.5)
            canvas.drawCentredString(
                doc.pagesize[0] / 2,
                doc.pagesize[1] - 0.5 * inch,
                self.config.title
            )
            canvas.line(
                doc.leftMargin,
                doc.pagesize[1] - 0.6 * inch,
                doc.pagesize[0] - doc.rightMargin,
                doc.pagesize[1] - 0.6 * inch
            )
        
        # Footer
        if self.config.add_footer:
            canvas.setFont('Times-Roman', 9)
            canvas.setFillGray(0.5)
            
            footer_text = f"{self.config.author}"
            if self.config.add_page_numbers:
                footer_text += f" | Page {doc.page}"
            footer_text += f" | {datetime.now().strftime('%Y-%m-%d')}"
            
            canvas.drawCentredString(
                doc.pagesize[0] / 2,
                0.5 * inch,
                footer_text
            )
            canvas.line(
                doc.leftMargin,
                0.6 * inch,
                doc.pagesize[0] - doc.rightMargin,
                0.6 * inch
            )
        
        # Watermark
        if self.config.add_watermark:
            canvas.setFont('Helvetica', 60)
            canvas.setFillGray(0.9)
            canvas.saveState()
            canvas.translate(doc.pagesize[0] / 2, doc.pagesize[1] / 2)
            canvas.rotate(45)
            canvas.drawCentredString(0, 0, self.config.watermark_text)
            canvas.restoreState()
        
        canvas.restoreState()
    
    def add_content(self, text: str, structured: bool = True):
        """Add content to PDF story"""
        if structured and self.config.detect_headers:
            elements = DocumentAnalyzer.detect_structure(text)
            self._add_structured_content(elements)
        else:
            # Add as simple paragraphs
            paragraphs = text.split('\n\n')
            for para_text in paragraphs:
                if para_text.strip():
                    # Escape HTML/XML special characters to prevent parsing errors
                    escaped_text = html.escape(para_text.strip())
                    p = Paragraph(escaped_text, self.styles['normal'])
                    self.story.append(p)
                    self.story.append(Spacer(1, 0.2 * inch))
    
    def _add_structured_content(self, elements: List[Dict]):
        """Add structured content to PDF"""
        in_code_block = False
        code_lines = []
        table_rows = []
        in_table = False
        
        for element in elements:
            elem_type = element['type']
            elem_text = element['text']
            
            # Handle code blocks
            if elem_type == 'code_marker':
                if in_code_block:
                    # End code block
                    code_text = '<pre>' + '\n'.join(code_lines) + '</pre>'
                    p = Paragraph(code_text, self.styles.get('code', self.styles['Code']))
                    self.story.append(p)
                    self.story.append(Spacer(1, 0.1 * inch))
                    code_lines = []
                    in_code_block = False
                else:
                    in_code_block = True
                continue
            
            if in_code_block:
                code_lines.append(elem_text)
                continue
            
            # Handle tables
            if elem_type == 'table_row':
                if not in_table:
                    in_table = True
                    table_rows = []
                cells = [cell.strip() for cell in elem_text.split('|') if cell.strip()]
                table_rows.append(cells)
                
                # Check if next element is not a table row
                # (would need look-ahead logic)
                continue
            elif in_table:
                # End table
                if table_rows:
                    t = Table(table_rows)
                    t.setStyle(TableStyle([
                        ('BACKGROUND', (0, 0), (-1, 0), colors.grey),
                        ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
                        ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
                        ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                        ('FONTSIZE', (0, 0), (-1, 0), 12),
                        ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
                        ('BACKGROUND', (0, 1), (-1, -1), colors.beige),
                        ('GRID', (0, 0), (-1, -1), 1, colors.black),
                    ]))
                    self.story.append(t)
                    self.story.append(Spacer(1, 0.2 * inch))
                in_table = False
                table_rows = []
            
            # Handle different element types
            if elem_type.startswith('heading'):
                style_name = elem_type
                p = Paragraph(elem_text, self.styles.get(style_name, self.styles['Heading1']))
                self.story.append(p)
                self.story.append(Spacer(1, 0.1 * inch))
                
            elif elem_type == 'bullet':
                p = Paragraph(f"• {elem_text}", self.styles['normal'])
                self.story.append(p)
                
            elif elem_type == 'numbered':
                p = Paragraph(elem_text, self.styles['normal'])
                self.story.append(p)
                
            elif elem_type == 'quote':
                p = Paragraph(f"<i>{elem_text}</i>", self.styles.get('quote', self.styles['Italic']))
                self.story.append(p)
                self.story.append(Spacer(1, 0.1 * inch))
                
            elif elem_type == 'paragraph':
                p = Paragraph(elem_text, self.styles['normal'])
                self.story.append(p)
                self.story.append(Spacer(1, 0.1 * inch))
    
    def save(self, filename: str):
        """Build and save PDF document"""
        doc = SimpleDocTemplate(
            filename,
            pagesize=A4,
            topMargin=self.config.margin_top * inch,
            bottomMargin=self.config.margin_bottom * inch,
            leftMargin=self.config.margin_left * inch,
            rightMargin=self.config.margin_right * inch,
            title=self.config.title,
            author=self.config.author,
            subject=self.config.subject,
            keywords=self.config.keywords
        )
        
        # Build document
        if self.config.add_header or self.config.add_footer or self.config.add_watermark:
            doc.build(
                self.story,
                onFirstPage=self._create_header_footer,
                onLaterPages=self._create_header_footer
            )
        else:
            doc.build(self.story)

# ========== Universal Exporter ==========
class UniversalExporter:
    """Main exporter class supporting multiple formats"""
    
    def __init__(self, config: Optional[ExportConfig] = None):
        self.config = config or ExportConfig()
        self.supported_formats = self._detect_supported_formats()
    
    def _detect_supported_formats(self) -> List[str]:
        """Detect available export formats"""
        formats = ['txt', 'md', 'html']  # Always available
        
        if DOCX_AVAILABLE:
            formats.append('docx')
        if REPORTLAB_AVAILABLE:
            formats.append('pdf')
        if DOCX_AVAILABLE and DOCX2PDF_AVAILABLE:
            formats.append('pdf_from_docx')
        
        return formats
    
    def export(self, text: str, output_path: str, format: str = 'auto', 
               metadata: Optional[Dict] = None) -> bool:
        """
        Export text to specified format
        
        Args:
            text: Content to export
            output_path: Output file path
            format: Export format (auto, docx, pdf, html, md, txt)
            metadata: Additional metadata for the document
        
        Returns:
            Success status
        """
        # Auto-detect format from extension
        if format == 'auto':
            ext = Path(output_path).suffix.lower()
            format = ext[1:] if ext else 'docx'
        
        # Update config with metadata
        if metadata:
            for key, value in metadata.items():
                if hasattr(self.config, key):
                    setattr(self.config, key, value)
        
        # Export based on format
        try:
            if format == 'docx':
                return self._export_docx(text, output_path)
            elif format == 'pdf':
                return self._export_pdf(text, output_path)
            elif format == 'html':
                return self._export_html(text, output_path)
            elif format == 'md':
                return self._export_markdown(text, output_path)
            elif format == 'txt':
                return self._export_text(text, output_path)
            else:
                print(f"Unsupported format: {format}")
                return False
        except Exception as e:
            print(f"Export error: {e}")
            return False
    
    def _export_docx(self, text: str, output_path: str) -> bool:
        """Export to DOCX format"""
        if not DOCX_AVAILABLE:
            print("DOCX export not available. Install python-docx.")
            return self._fallback_export(text, output_path, 'docx')
        
        exporter = DocxExporter(self.config)
        exporter.add_content(text, structured=True)
        exporter.save(output_path)
        return True
    
    def _export_pdf(self, text: str, output_path: str) -> bool:
        """Export to PDF format"""
        if REPORTLAB_AVAILABLE:
            # Use ReportLab for direct PDF generation
            exporter = PdfExporter(self.config)
            exporter.add_content(text, structured=True)
            exporter.save(output_path)
            return True
        
        elif DOCX_AVAILABLE and DOCX2PDF_AVAILABLE:
            # Fallback: Create DOCX then convert to PDF
            temp_docx = output_path.replace('.pdf', '_temp.docx')
            if self._export_docx(text, temp_docx):
                try:
                    docx_to_pdf(temp_docx, output_path)
                    Path(temp_docx).unlink()  # Remove temp file
                    return True
                except Exception as e:
                    print(f"DOCX to PDF conversion failed: {e}")
                    return False
        
        else:
            print("PDF export not available. Install reportlab or docx2pdf.")
            return self._fallback_export(text, output_path, 'pdf')
    
    def _export_html(self, text: str, output_path: str) -> bool:
        """Export to HTML format"""
        html_template = """
<!DOCTYPE html>
<html lang="en">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>{title}</title>
    <meta name="author" content="{author}">
    <meta name="description" content="{subject}">
    <meta name="keywords" content="{keywords}">
    <style>
        body {{
            font-family: '{font_family}', serif;
            font-size: {font_size}pt;
            line-height: {line_spacing};
            max-width: 800px;
            margin: 0 auto;
            padding: 40px 20px;
            color: #333;
        }}
        h1 {{ color: #00008B; margin-top: 30px; }}
        h2 {{ color: #191970; margin-top: 25px; }}
        h3 {{ color: #4682B4; margin-top: 20px; font-style: italic; }}
        p {{ text-align: justify; margin: 10px 0; }}
        blockquote {{
            font-style: italic;
            color: #696969;
            margin: 20px 40px;
            padding: 10px 20px;
            border-left: 3px solid #ddd;
        }}
        pre, code {{
            font-family: 'Courier New', monospace;
            background-color: #f5f5f5;
            padding: 2px 4px;
            border-radius: 3px;
        }}
        pre {{ padding: 10px; overflow-x: auto; }}
        .watermark {{
            position: fixed;
            top: 50%;
            left: 50%;
            transform: translate(-50%, -50%) rotate(-45deg);
            font-size: 120px;
            color: rgba(0,0,0,0.05);
            z-index: -1;
            user-select: none;
        }}
    </style>
</head>
<body>
    {watermark}
    <article>
        {content}
    </article>
    <footer style="margin-top: 50px; padding-top: 20px; border-top: 1px solid #ddd; text-align: center; color: #888; font-size: 0.9em;">
        <p>{author} | Generated on {date}</p>
    </footer>
</body>
</html>
        """
        
        # Convert text to HTML
        content_html = self._text_to_html(text)
        
        # Add watermark if configured
        watermark_html = ""
        if self.config.add_watermark:
            watermark_html = f'<div class="watermark">{self.config.watermark_text}</div>'
        
        # Format HTML
        html = html_template.format(
            title=self.config.title,
            author=self.config.author,
            subject=self.config.subject,
            keywords=', '.join(self.config.keywords),
            font_family=self.config.font_family,
            font_size=self.config.font_size,
            line_spacing=self.config.line_spacing,
            watermark=watermark_html,
            content=content_html,
            date=datetime.now().strftime('%Y-%m-%d %H:%M')
        )
        
        # Save
        Path(output_path).write_text(html, encoding='utf-8')
        return True
    
    def _text_to_html(self, text: str) -> str:
        """Convert plain text to HTML with structure detection"""
        if self.config.detect_headers:
            elements = DocumentAnalyzer.detect_structure(text)
            html_parts = []
            
            in_code_block = False
            code_lines = []
            
            for element in elements:
                elem_type = element['type']
                elem_text = element['text']
                
                # Escape HTML
                elem_text = elem_text.replace('&', '&amp;')
                elem_text = elem_text.replace('<', '&lt;')
                elem_text = elem_text.replace('>', '&gt;')
                
                if elem_type == 'code_marker':
                    if in_code_block:
                        html_parts.append(f'<pre><code>{"<br>".join(code_lines)}</code></pre>')
                        code_lines = []
                        in_code_block = False
                    else:
                        in_code_block = True
                elif in_code_block:
                    code_lines.append(elem_text)
                elif elem_type == 'heading1':
                    html_parts.append(f'<h1>{elem_text}</h1>')
                elif elem_type == 'heading2':
                    html_parts.append(f'<h2>{elem_text}</h2>')
                elif elem_type == 'heading3':
                    html_parts.append(f'<h3>{elem_text}</h3>')
                elif elem_type == 'quote':
                    html_parts.append(f'<blockquote>{elem_text}</blockquote>')
                elif elem_type == 'bullet':
                    html_parts.append(f'<li>{elem_text}</li>')
                elif elem_type == 'numbered':
                    html_parts.append(f'<li>{elem_text}</li>')
                else:
                    html_parts.append(f'<p>{elem_text}</p>')
            
            return '\n'.join(html_parts)
        else:
            # Simple paragraph conversion
            paragraphs = text.split('\n\n')
            return '\n'.join(f'<p>{p.strip()}</p>' for p in paragraphs if p.strip())
    
    def _export_markdown(self, text: str, output_path: str) -> bool:
        """Export to Markdown format"""
        md_header = f"""---
title: {self.config.title}
author: {self.config.author}
date: {datetime.now().strftime('%Y-%m-%d')}
subject: {self.config.subject}
keywords: [{', '.join(self.config.keywords)}]
---

# {self.config.title}

"""
        
        # Add content with basic formatting
        md_content = md_header + text
        
        Path(output_path).write_text(md_content, encoding='utf-8')
        return True
    
    def _export_text(self, text: str, output_path: str) -> bool:
        """Export to plain text format"""
        header = f"""{'=' * 70}
{self.config.title.center(70)}
{self.config.author.center(70)}
{datetime.now().strftime('%Y-%m-%d').center(70)}
{'=' * 70}

"""
        
        content = header + text
        
        Path(output_path).write_text(content, encoding='utf-8')
        return True
    
    def _fallback_export(self, text: str, output_path: str, format: str) -> bool:
        """Fallback export when libraries are not available"""
        print(f"Using fallback export for {format}")
        
        # For unsupported formats, export as text with notice
        fallback_path = output_path.replace(f'.{format}', '.txt')
        self._export_text(text, fallback_path)
        
        print(f"Exported as text to: {fallback_path}")
        print(f"To enable {format} export, install required libraries.")
        return True
    
    def export_batch(self, texts: List[str], output_dir: str, 
                    format: str = 'docx', prefix: str = 'document') -> List[str]:
        """
        Export multiple documents
        
        Args:
            texts: List of text contents
            output_dir: Output directory
            format: Export format
            prefix: Filename prefix
        
        Returns:
            List of output file paths
        """
        output_dir = Path(output_dir)
        output_dir.mkdir(exist_ok=True)
        
        output_files = []
        for i, text in enumerate(texts, 1):
            output_path = output_dir / f"{prefix}_{i:03d}.{format}"
            if self.export(text, str(output_path), format):
                output_files.append(str(output_path))
        
        return output_files

# ========== Utility Functions ==========

def create_comparison_document(
    original_text: str,
    translated_text: str,
    output_path: str,
    format: str = 'docx',
    side_by_side: bool = False
) -> bool:
    """
    Create a comparison document showing original and translated text
    
    Args:
        original_text: Original text
        translated_text: Translated text
        output_path: Output file path
        format: Output format (docx or pdf)
        side_by_side: If True, show texts side by side (tables)
    
    Returns:
        Success status
    """
    config = ExportConfig(
        title="Translation Comparison",
        detect_headers=False
    )
    
    if side_by_side and format == 'docx' and DOCX_AVAILABLE:
        # Create side-by-side comparison in DOCX
        doc = Document()
        
        # Add title
        doc.add_heading('Translation Comparison', 1)
        doc.add_paragraph(f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M')}")
        doc.add_page_break()
        
        # Create comparison table
        orig_paras = original_text.split('\n\n')
        trans_paras = translated_text.split('\n\n')
        
        table = doc.add_table(rows=1, cols=2)
        table.style = 'Table Grid'
        
        # Header row
        header_cells = table.rows[0].cells
        header_cells[0].text = 'Original'
        header_cells[1].text = 'Translated'
        
        # Add content rows
        max_paras = max(len(orig_paras), len(trans_paras))
        for i in range(max_paras):
            row_cells = table.add_row().cells
            row_cells[0].text = orig_paras[i] if i < len(orig_paras) else ''
            row_cells[1].text = trans_paras[i] if i < len(trans_paras) else ''
        
        doc.save(output_path)
        return True
    
    else:
        # Sequential comparison
        combined_text = f"""
{'=' * 70}
ORIGINAL TEXT
{'=' * 70}

{original_text}

{'=' * 70}
TRANSLATED TEXT
{'=' * 70}

{translated_text}

{'=' * 70}
END OF COMPARISON
{'=' * 70}
"""
        
        exporter = UniversalExporter(config)
        return exporter.export(combined_text, output_path, format)

# ========== Main Demo ==========

def demo():
    """Demo function showing export capabilities"""
    print("=" * 70)
    print("AI Translator Pro - Export Module Demo".center(70))
    print("=" * 70)
    
    # Sample text
    sample_text = """
# Chapter 1: Introduction

This is a demonstration of the advanced export capabilities of AI Translator Pro.

## Features

The system supports multiple export formats:

* **DOCX**: Full formatting with styles
* **PDF**: Professional layout using ReportLab
* **HTML**: Web-ready with embedded styles
* **Markdown**: For documentation
* **Plain Text**: Universal compatibility

## Code Example

```python
def translate(text):
    # Advanced translation logic
    return translated_text
```

### Quality Metrics

The translation quality is measured using multiple metrics:

1. Length ratio analysis
2. Completeness checking
3. Terminology consistency
4. Grammar validation

> "The best translation is invisible - it reads as if it were originally written in the target language."

## Conclusion

This export module provides professional-quality document generation for all your translation needs.
"""
    
    # Create output directory
    output_dir = Path("export_demo")
    output_dir.mkdir(exist_ok=True)
    
    # Initialize exporter with custom config
    config = ExportConfig(
        title="AI Translator Pro Demo",
        author="Translation Team",
        subject="Export Capabilities Demo",
        keywords=["translation", "export", "demo"],
        add_watermark=True,
        watermark_text="DEMO",
        detect_headers=True
    )
    
    exporter = UniversalExporter(config)
    
    print(f"\nAvailable formats: {', '.join(exporter.supported_formats)}")
    print(f"Output directory: {output_dir}")
    print("-" * 70)
    
    # Export to different formats
    formats = ['docx', 'pdf', 'html', 'md', 'txt']
    
    for format in formats:
        output_file = output_dir / f"demo.{format}"
        print(f"Exporting to {format}...", end=" ")
        
        if exporter.export(sample_text, str(output_file), format):
            print(f"✓ Success: {output_file}")
        else:
            print(f"✗ Failed")
    
    # Create comparison document
    print("\nCreating comparison document...", end=" ")
    comparison_file = output_dir / "comparison.docx"
    if create_comparison_document(
        sample_text,
        sample_text,  # Would be translated text in real use
        str(comparison_file),
        'docx',
        side_by_side=True
    ):
        print(f"✓ Success: {comparison_file}")
    else:
        print("✗ Failed")
    
    print("\n" + "=" * 70)
    print("Demo completed! Check the 'export_demo' directory for output files.")

if __name__ == "__main__":
    demo()
