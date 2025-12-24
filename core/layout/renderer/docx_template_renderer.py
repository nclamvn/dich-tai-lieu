#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Template-Based DOCX Renderer

Uses pre-built templates for maximum performance:
1. Load template with pre-defined styles
2. Clear placeholder content
3. Insert actual content using existing styles
4. Save document

Expected: ~2x faster than creating from scratch.

Version: 1.0.0
"""

from typing import List, Dict, Optional, Any
from pathlib import Path
import logging

try:
    from docx import Document
    from docx.shared import Pt, Inches, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.style import WD_STYLE_TYPE
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False
    Document = None

from core.contracts import (
    LayoutIntentPackage,
    Block,
    BlockType,
)
from .base_renderer import BaseRenderer

logger = logging.getLogger(__name__)

# Templates directory
TEMPLATES_DIR = Path(__file__).parent.parent / "templates"


class TemplateCache:
    """
    Cache for loaded template documents.

    Templates are loaded once and their bytes cached in memory.
    Each render gets a fresh Document from cached bytes (no disk I/O).
    """
    _path_cache: Dict[str, Path] = {}
    _bytes_cache: Dict[str, bytes] = {}

    @classmethod
    def get_template_path(cls, template_type: str) -> Optional[Path]:
        """Get path to template file"""
        if template_type not in cls._path_cache:
            template_path = TEMPLATES_DIR / f"base_{template_type}.docx"
            if template_path.exists():
                cls._path_cache[template_type] = template_path
            else:
                # Fallback to book template
                fallback = TEMPLATES_DIR / "base_book.docx"
                if fallback.exists():
                    cls._path_cache[template_type] = fallback
                else:
                    # Try default template
                    default = TEMPLATES_DIR / "base_default.docx"
                    if default.exists():
                        cls._path_cache[template_type] = default
                    else:
                        cls._path_cache[template_type] = None

        return cls._path_cache.get(template_type)

    @classmethod
    def get_template_bytes(cls, template_type: str) -> Optional[bytes]:
        """Get cached template bytes (load once, reuse forever)"""
        if template_type not in cls._bytes_cache:
            path = cls.get_template_path(template_type)
            if path and path.exists():
                cls._bytes_cache[template_type] = path.read_bytes()
            else:
                cls._bytes_cache[template_type] = None

        return cls._bytes_cache.get(template_type)

    @classmethod
    def clear(cls):
        """Clear cache"""
        cls._path_cache.clear()
        cls._bytes_cache.clear()


class TemplateDocxRenderer(BaseRenderer):
    """
    Template-based DOCX Renderer.

    Performance optimizations:
    1. Pre-built templates with all styles
    2. No style creation at runtime
    3. Direct content insertion
    4. Minimal python-docx API calls

    Usage:
        renderer = TemplateDocxRenderer(template="book")
        renderer.render(lip, blocks, "output.docx")
    """

    # Map BlockType to template style name
    STYLE_MAP = {
        BlockType.TITLE: "APS_Title",
        BlockType.SUBTITLE: "APS_Subtitle",
        BlockType.CHAPTER: "APS_Chapter",
        BlockType.SECTION: "APS_Section",
        BlockType.HEADING_1: "APS_Heading1",
        BlockType.HEADING_2: "APS_Heading2",
        BlockType.HEADING_3: "APS_Heading3",
        BlockType.PARAGRAPH: "APS_Paragraph",
        BlockType.QUOTE: "APS_Quote",
        BlockType.CODE: "APS_Code",
        BlockType.LIST: "APS_List",
        BlockType.FOOTNOTE: "APS_Footnote",
        BlockType.TOC_ENTRY: "APS_TOC1",
    }

    def __init__(
        self,
        template: str = "book",
        page_size: str = "A4",
    ):
        super().__init__(template, page_size)
        self._template_type = template
        # Pre-load template bytes into cache
        self._template_bytes = TemplateCache.get_template_bytes(template)

        if self._template_bytes is None:
            logger.warning(f"Template '{template}' not found, will create from scratch")

    def render(
        self,
        lip: LayoutIntentPackage,
        flowed_blocks: List,
        output_path: str,
        section_manager=None,
    ) -> Path:
        """
        Render using template.
        """
        output_path = Path(output_path)

        if not HAS_DOCX:
            return self._simulate_render(lip, flowed_blocks, output_path)

        # Load template from cached bytes (no disk I/O) or create new document
        if self._template_bytes:
            from io import BytesIO
            doc = Document(BytesIO(self._template_bytes))
            # Clear placeholder content
            self._clear_placeholder(doc)
        else:
            # Fallback: create from scratch
            doc = Document()
            self._setup_fallback(doc)
            logger.warning("Using fallback (no template)")

        # Render all blocks
        self._render_blocks(doc, flowed_blocks)

        # Save document
        doc.save(str(output_path))

        logger.debug(f"Template DOCX saved: {output_path}")
        return output_path

    def _simulate_render(
        self,
        lip: LayoutIntentPackage,
        flowed_blocks: List,
        output_path: Path,
    ) -> Path:
        """Simulate render when python-docx not available"""
        lines = [
            f"# {lip.title}",
            f"## {lip.subtitle}" if lip.subtitle else "",
            "",
            f"Template: {self.template}",
            f"Total Blocks: {len(flowed_blocks)}",
            "",
        ]

        for fb in flowed_blocks:
            block = fb.block
            lines.append(f"[{block.type.value}] {block.content[:100]}...")

        output_path.write_text("\n".join(lines), encoding="utf-8")
        return output_path

    def _setup_fallback(self, doc: Document):
        """Setup document when template not available"""
        section = doc.sections[0]
        section.page_width = Cm(21)
        section.page_height = Cm(29.7)
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    def _clear_placeholder(self, doc: Document):
        """Remove placeholder paragraph from template"""
        for para in doc.paragraphs:
            if "{{APS_CONTENT_START}}" in para.text:
                # Remove this paragraph
                p = para._element
                p.getparent().remove(p)
                break

    def _render_blocks(self, doc: Document, flowed_blocks: List):
        """
        Render all blocks using template styles.

        Optimization: Use style name directly (already exists in template).
        """
        current_page = 1

        for fb in flowed_blocks:
            block = fb.block

            # Handle page break
            if hasattr(fb, 'page_break_before') and fb.page_break_before:
                if hasattr(fb, 'page_number') and fb.page_number > current_page:
                    doc.add_page_break()
                    current_page = fb.page_number

            # Get style name for this block type
            style_name = self.STYLE_MAP.get(block.type, "APS_Paragraph")

            # Add paragraph with pre-existing style
            try:
                doc.add_paragraph(block.content, style=style_name)
            except KeyError:
                # Style doesn't exist, use default
                doc.add_paragraph(block.content)

    @classmethod
    def supports_format(cls, format_name: str) -> bool:
        return format_name.lower() == "docx"

    @classmethod
    def get_supported_formats(cls) -> List[str]:
        return ["docx"]

    @classmethod
    def templates_available(cls) -> bool:
        """Check if templates are available"""
        return TEMPLATES_DIR.exists() and any(TEMPLATES_DIR.glob("base_*.docx"))


def ensure_templates_exist():
    """
    Ensure templates exist, create if not.

    Call this on application startup.
    """
    if not TEMPLATES_DIR.exists() or not any(TEMPLATES_DIR.glob("base_*.docx")):
        logger.info("Creating DOCX templates...")

        # Import and run template creator
        try:
            import sys
            scripts_dir = Path(__file__).parent.parent.parent.parent / "scripts"
            sys.path.insert(0, str(scripts_dir.parent))
            from scripts.create_docx_templates import create_all_templates
            create_all_templates()
        except ImportError as e:
            logger.warning(f"Could not import template creator: {e}")
            # Create minimal template inline
            _create_minimal_template()


def _create_minimal_template():
    """Create a minimal template if script not available"""
    if not HAS_DOCX:
        logger.warning("python-docx not available, skipping template creation")
        return

    TEMPLATES_DIR.mkdir(parents=True, exist_ok=True)

    doc = Document()

    # Set page size
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

    # Create basic styles
    styles = doc.styles

    basic_styles = [
        ("APS_Title", 24, True, False),
        ("APS_Subtitle", 14, False, True),
        ("APS_Chapter", 18, True, False),
        ("APS_Section", 14, True, False),
        ("APS_Heading1", 16, True, False),
        ("APS_Heading2", 14, True, False),
        ("APS_Heading3", 12, True, False),
        ("APS_Paragraph", 11, False, False),
        ("APS_Quote", 11, False, True),
        ("APS_Code", 10, False, False),
        ("APS_List", 11, False, False),
        ("APS_Footnote", 9, False, False),
        ("APS_TOC1", 12, True, False),
        ("APS_TOC2", 11, False, False),
    ]

    for name, size, bold, italic in basic_styles:
        try:
            style = styles.add_style(name, WD_STYLE_TYPE.PARAGRAPH)
            style.font.size = Pt(size)
            style.font.bold = bold
            style.font.italic = italic
        except Exception:
            pass

    doc.add_paragraph("{{APS_CONTENT_START}}")

    # Save as both book and default
    for template_name in ["book", "default"]:
        doc.save(str(TEMPLATES_DIR / f"base_{template_name}.docx"))

    logger.info(f"Created minimal templates in: {TEMPLATES_DIR}")
