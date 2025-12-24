#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
DOCX Renderer

Renders flowed blocks to DOCX format.
Integrates with existing core/formatting exporters.

Version: 1.0.0
"""

from typing import List, Dict, Optional, Any, TYPE_CHECKING
from pathlib import Path
import logging

try:
    from docx import Document
    from docx.shared import Pt, Inches, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.style import WD_STYLE_TYPE
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False
    Document = None

from core.contracts import (
    LayoutIntentPackage,
    Block,
    BlockType,
    SectionType,
)

if TYPE_CHECKING:
    from ..executor.block_flow import FlowedBlock
    from ..sections.manager import SectionManager

logger = logging.getLogger(__name__)


class DocxRenderer:
    """
    Renders LayoutIntentPackage to DOCX.

    Usage:
        renderer = DocxRenderer(template="book")
        renderer.render(lip, flowed_blocks, "output.docx")
    """

    # Block type to Word style mapping
    STYLE_MAP = {
        BlockType.TITLE: "Title",
        BlockType.SUBTITLE: "Subtitle",
        BlockType.CHAPTER: "Heading 1",
        BlockType.SECTION: "Heading 2",
        BlockType.HEADING_1: "Heading 1",
        BlockType.HEADING_2: "Heading 2",
        BlockType.HEADING_3: "Heading 3",
        BlockType.PARAGRAPH: "Normal",
        BlockType.QUOTE: "Quote",
        BlockType.LIST: "List Paragraph",
        BlockType.CODE: "Normal",  # Will apply monospace
        BlockType.FOOTNOTE: "Normal",
        BlockType.TOC_ENTRY: "TOC 1",
    }

    def __init__(
        self,
        template: str = "default",
        page_size: str = "A4",
    ):
        """
        Initialize DOCX renderer.

        Args:
            template: Template name (affects styling)
            page_size: Page size
        """
        if not HAS_DOCX:
            logger.warning("python-docx not installed. DOCX rendering will be simulated.")

        self.template = template
        self.page_size = page_size

    def render(
        self,
        lip: LayoutIntentPackage,
        flowed_blocks: List,  # List[FlowedBlock]
        output_path: str,
        section_manager: Optional[Any] = None,
    ) -> Path:
        """
        Render to DOCX file.

        Args:
            lip: LayoutIntentPackage
            flowed_blocks: Flowed blocks from executor
            output_path: Output file path
            section_manager: Optional SectionManager for headers/footers

        Returns:
            Path to created file
        """
        logger.info(f"Rendering DOCX: {len(flowed_blocks)} blocks")

        output_path = Path(output_path)

        if not HAS_DOCX:
            # Simulate output for testing
            return self._simulate_render(lip, flowed_blocks, output_path)

        # Create document
        doc = Document()

        # Set page size
        self._set_page_size(doc)

        # Add title if present
        if lip.title:
            self._add_title(doc, lip.title, lip.subtitle)

        # Track current page for page breaks
        current_page = 1

        # Render blocks
        for fb in flowed_blocks:
            block = fb.block

            # Handle page breaks
            if fb.page_break_before and fb.page_number > current_page:
                doc.add_page_break()
                current_page = fb.page_number

            # Render block
            self._render_block(doc, block)

        # Save
        doc.save(output_path)

        logger.info(f"DOCX saved: {output_path}")

        return output_path

    def _simulate_render(
        self,
        lip: LayoutIntentPackage,
        flowed_blocks: List,
        output_path: Path,
    ) -> Path:
        """Simulate render when python-docx not available"""
        # Create a simple text representation
        lines = [
            f"# {lip.title}",
            f"## {lip.subtitle}" if lip.subtitle else "",
            "",
            f"Template: {self.template}",
            f"Page Size: {self.page_size}",
            f"Total Blocks: {len(flowed_blocks)}",
            "",
            "=" * 50,
            "",
        ]

        for fb in flowed_blocks:
            block = fb.block
            lines.append(f"[Page {fb.page_number}] [{block.type.value}] {block.content[:100]}...")

        # Write as text file with .docx extension (for testing)
        output_path.write_text("\n".join(lines), encoding="utf-8")

        logger.info(f"Simulated DOCX saved: {output_path}")

        return output_path

    def _set_page_size(self, doc):
        """Set document page size"""
        if not HAS_DOCX:
            return

        section = doc.sections[0]

        if self.page_size == "A4":
            section.page_width = Cm(21)
            section.page_height = Cm(29.7)
        elif self.page_size == "letter":
            section.page_width = Inches(8.5)
            section.page_height = Inches(11)
        elif self.page_size == "A5":
            section.page_width = Cm(14.8)
            section.page_height = Cm(21)

        # Margins
        section.top_margin = Cm(2.54)
        section.bottom_margin = Cm(2.54)
        section.left_margin = Cm(2.54)
        section.right_margin = Cm(2.54)

    def _add_title(self, doc, title: str, subtitle: str = ""):
        """Add document title"""
        if not HAS_DOCX:
            return

        # Title
        p = doc.add_paragraph(title)
        p.style = 'Title'
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Subtitle
        if subtitle:
            p = doc.add_paragraph(subtitle)
            p.style = 'Subtitle'
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Page break after title page
        doc.add_page_break()

    def _render_block(self, doc, block: Block):
        """Render a single block"""
        if not HAS_DOCX:
            return

        style_name = self.STYLE_MAP.get(block.type, "Normal")

        # Add paragraph
        p = doc.add_paragraph(block.content)

        # Apply style
        try:
            p.style = style_name
        except KeyError:
            p.style = "Normal"

        # Apply spacing
        if block.spacing:
            p.paragraph_format.space_before = Pt(block.spacing.before)
            p.paragraph_format.space_after = Pt(block.spacing.after)
            p.paragraph_format.line_spacing = block.spacing.line_spacing

        # Special handling for code blocks
        if block.type == BlockType.CODE:
            for run in p.runs:
                run.font.name = "Courier New"
                run.font.size = Pt(10)

        # Special handling for quotes
        if block.type == BlockType.QUOTE:
            p.paragraph_format.left_indent = Cm(1)
            p.paragraph_format.right_indent = Cm(1)

    def render_toc(
        self,
        doc,
        lip: LayoutIntentPackage,
    ):
        """Render table of contents"""
        if not HAS_DOCX or not lip.toc_blocks:
            return

        # TOC heading
        doc.add_paragraph("Table of Contents", style='Heading 1')

        # TOC entries
        for block in lip.toc_blocks:
            if block.type == BlockType.TOC_ENTRY:
                p = doc.add_paragraph(block.content)

                # Indent by level
                indent = block.level * 0.5
                p.paragraph_format.left_indent = Cm(indent)

                # Add chapter number if present
                if block.number:
                    p.insert_paragraph_before(f"{block.number}. ")
