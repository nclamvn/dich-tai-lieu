#!/usr/bin/env python3
# -*- coding: utf-8 -*-

"""
Incremental DOCX Builder - Phase 5.3 (Updated for 5.4 + FIX-005)

Build DOCX documents incrementally to reduce memory usage.
Instead of building entire DOCX in RAM, write batches to temp files.

FIX-005: Smart formatting - detect headings, paragraphs, apply styles.
"""

from pathlib import Path
from typing import List
import re
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

from config.logging_config import get_logger
logger = get_logger(__name__)

from core.validator import TranslationResult
from .base_builder import BaseIncrementalBuilder


# FIX-005: Heading detection patterns
CHAPTER_PATTERNS = [
    r'^(CHAPTER|Chapter|CHƯƠNG|Chương|PHẦN|Phần)\s+(\d+|[IVXLCDM]+)',  # CHAPTER 1, Chương 1, etc.
    r'^(CHƯƠNG|Chương)\s+\d+\s*[:：\-–—]',  # Chương 1: Title
    r'^(Chapter|CHAPTER)\s+\d+\s*[:：\-–—]',  # Chapter 1: Title
]

SECTION_PATTERNS = [
    r'^(\d+\.)+\s+\w',  # 1.1 Section, 1.2.3 Subsection
    r'^(Section|SECTION|Phần|PHẦN|Mục|MỤC)\s+\d+',  # Section 1
    r'^[IVXLCDM]+\.\s+\w',  # I. Section, II. Section
]


class IncrementalDocxBuilder(BaseIncrementalBuilder):
    """
    Build DOCX incrementally in batches

    Instead of:
      - Load ALL results in memory → Build DOCX → Save
      - Peak memory: ~1GB for large doc

    Do:
      - Build batch 1 DOCX → Save to temp
      - Build batch 2 DOCX → Save to temp
      - ...
      - Merge all temp DOCX → Final output
      - Peak memory: ~100MB (10x reduction)

    Usage:
        builder = IncrementalDocxBuilder(output_path)

        # Add batches
        for batch in batches:
            temp_file = await builder.add_batch(batch, batch_idx)

        # Merge all batches
        final_file = await builder.merge_all()
    """

    def get_format(self) -> str:
        """Get format identifier"""
        return 'docx'

    # =========================================================================
    # FIX-005: Smart formatting methods
    # =========================================================================

    def _is_chapter_heading(self, text: str) -> bool:
        """Check if text is a chapter heading."""
        text = text.strip()
        for pattern in CHAPTER_PATTERNS:
            if re.match(pattern, text, re.IGNORECASE):
                return True
        return False

    def _is_section_heading(self, text: str) -> bool:
        """Check if text is a section heading."""
        text = text.strip()
        for pattern in SECTION_PATTERNS:
            if re.match(pattern, text, re.IGNORECASE):
                return True
        return False

    def _is_short_line(self, text: str, max_chars: int = 80) -> bool:
        """Check if text is a short line (potential heading/title)."""
        return len(text.strip()) < max_chars and '\n' not in text.strip()

    def _setup_base_styles(self, doc: Document) -> None:
        """Setup base document styles for professional output."""
        styles = doc.styles

        # Normal style - body text
        normal = styles['Normal']
        normal.font.name = 'Times New Roman'
        normal.font.size = Pt(11)
        normal.paragraph_format.line_spacing = 1.3
        normal.paragraph_format.space_after = Pt(6)

        # Heading 1 - Chapter headings
        if 'Heading 1' in styles:
            h1 = styles['Heading 1']
            h1.font.name = 'Times New Roman'
            h1.font.size = Pt(18)
            h1.font.bold = True
            h1.paragraph_format.space_before = Pt(24)
            h1.paragraph_format.space_after = Pt(12)

        # Heading 2 - Section headings
        if 'Heading 2' in styles:
            h2 = styles['Heading 2']
            h2.font.name = 'Times New Roman'
            h2.font.size = Pt(14)
            h2.font.bold = True
            h2.paragraph_format.space_before = Pt(18)
            h2.paragraph_format.space_after = Pt(6)

    def _add_formatted_text(self, doc: Document, text: str) -> None:
        """
        Add text to document with smart formatting.

        FIX-005: Detects structure and applies appropriate styles.
        """
        # Split by double newlines to get paragraphs
        paragraphs = re.split(r'\n\s*\n', text)

        for para_text in paragraphs:
            para_text = para_text.strip()
            if not para_text:
                continue

            # Check for chapter heading
            if self._is_chapter_heading(para_text):
                doc.add_heading(para_text, level=1)
                continue

            # Check for section heading
            if self._is_section_heading(para_text):
                doc.add_heading(para_text, level=2)
                continue

            # Check for short uppercase line (potential title)
            if self._is_short_line(para_text) and para_text.isupper():
                para = doc.add_paragraph(para_text)
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in para.runs:
                    run.bold = True
                continue

            # Handle multi-line paragraphs (preserve single line breaks)
            lines = para_text.split('\n')
            if len(lines) > 1:
                # Multiple lines - add as separate paragraphs or preserve breaks
                for line in lines:
                    line = line.strip()
                    if line:
                        doc.add_paragraph(line)
            else:
                # Single paragraph
                doc.add_paragraph(para_text)

    def _verify_docx(self, file_path: Path) -> None:
        """
        Verify DOCX file is valid and readable

        Args:
            file_path: Path to DOCX file to verify

        Raises:
            RuntimeError: If DOCX is corrupted or unreadable
        """
        try:
            # Try to open the DOCX
            doc = Document(file_path)

            # Verify it has content
            if len(doc.paragraphs) == 0:
                raise RuntimeError(f"DOCX has no paragraphs: {file_path}")

            # Verify paragraphs have text (at least some)
            has_content = any(para.text.strip() for para in doc.paragraphs)
            if not has_content:
                raise RuntimeError(f"DOCX has no text content: {file_path}")

        except Exception as e:
            if "corrupted" in str(e).lower() or "invalid" in str(e).lower():
                raise RuntimeError(f"DOCX file is corrupted: {file_path}: {e}")
            raise RuntimeError(f"Failed to verify DOCX: {file_path}: {e}")

    async def add_batch(
        self,
        batch_results: List[TranslationResult],
        batch_idx: int
    ) -> Path:
        """
        Add a batch and save to temp DOCX file with error handling

        Args:
            batch_results: Translation results for this batch
            batch_idx: Batch index (0-based)

        Returns:
            Path to temp batch DOCX file

        Raises:
            RuntimeError: If DOCX creation fails
        """
        batch_file = self.temp_dir / f"batch_{batch_idx:04d}.docx"

        try:
            # Create DOCX for this batch with smart formatting (FIX-005)
            doc = Document()

            # Setup base styles for professional output
            self._setup_base_styles(doc)

            for result in batch_results:
                try:
                    # FIX-005: Add translated text with smart formatting
                    self._add_formatted_text(doc, result.translated)
                except Exception as e:
                    logger.warning(f"Failed to add chunk {result.chunk_id}: {e}")
                    doc.add_paragraph(f"[Error: chunk {result.chunk_id} failed]")

            # Save batch file
            doc.save(str(batch_file))

            # Verify file created
            if not batch_file.exists():
                raise RuntimeError(f"DOCX batch not created: {batch_file}")

            if batch_file.stat().st_size == 0:
                raise RuntimeError(f"DOCX batch is empty: {batch_file}")

            # Deep verification - check DOCX is valid
            self._verify_docx(batch_file)

            self.batch_files.append(batch_file)
            return batch_file

        except PermissionError as e:
            raise RuntimeError(f"Permission denied: {batch_file}: {e}")
        except OSError as e:
            if "No space left" in str(e):
                raise RuntimeError(f"Disk full: {batch_file}: {e}")
            raise RuntimeError(f"OS error: {e}")
        except Exception as e:
            # Cleanup partial file
            if batch_file.exists():
                try:
                    batch_file.unlink()
                except Exception:
                    pass
            raise RuntimeError(f"Failed to create DOCX batch {batch_idx}: {e}")

    async def merge_all(self) -> Path:
        """
        Merge all batch DOCX files into final output with error handling

        Returns:
            Path to final merged DOCX

        Raises:
            RuntimeError: If merge fails
        """
        if not self.batch_files:
            raise ValueError("No batch files to merge")

        logger.info(f"Merging {len(self.batch_files)} DOCX batches...")

        try:
            # Create final document with professional styles (FIX-005)
            final_doc = Document()
            self._setup_base_styles(final_doc)

            # Merge each batch
            for batch_idx, batch_file in enumerate(self.batch_files):
                if not batch_file.exists():
                    raise RuntimeError(f"Batch file missing: {batch_file}")

                try:
                    # Load batch document
                    batch_doc = Document(batch_file)

                    # Copy all paragraphs from batch to final doc
                    for para in batch_doc.paragraphs:
                        # Create new paragraph with same text and style
                        new_para = final_doc.add_paragraph(
                            para.text,
                            style=para.style.name if para.style else None
                        )

                        # Copy paragraph formatting
                        new_para.paragraph_format.alignment = para.paragraph_format.alignment
                        new_para.paragraph_format.left_indent = para.paragraph_format.left_indent
                        new_para.paragraph_format.right_indent = para.paragraph_format.right_indent
                        new_para.paragraph_format.space_before = para.paragraph_format.space_before
                        new_para.paragraph_format.space_after = para.paragraph_format.space_after

                    logger.debug(f"Batch {batch_idx + 1}/{len(self.batch_files)} merged")

                except Exception as e:
                    raise RuntimeError(f"Failed to merge batch {batch_idx}: {e}")

            # Save final document
            final_doc.save(str(self.output_path))

            # Verify output exists
            if not self.output_path.exists():
                raise RuntimeError("Final DOCX not created")

            file_size = self.output_path.stat().st_size
            if file_size == 0:
                raise RuntimeError("Final DOCX is empty")

            # Deep verification - check final DOCX is valid and readable
            self._verify_docx(self.output_path)

            logger.info(f"Final DOCX saved: {self.output_path} ({file_size / 1024 / 1024:.1f} MB)")

            # Cleanup temp files
            await self.cleanup()

            return self.output_path

        except Exception as e:
            # Ensure cleanup even on error
            try:
                await self.cleanup()
            except Exception:
                pass
            raise
