"""
PDF Reconstructor Module (Phase 3 - Production Quality)

Two-mode PDF/DOCX rebuilding from translated layout:
1. Preserve Layout Mode: Maintains original PDF layout and positioning
2. Reflow DOCX Mode: Creates structured, single-column DOCX

Supports block-type aware formatting (titles, headings, captions, etc.)
"""

import fitz  # PyMuPDF
from pathlib import Path
from typing import List, Dict, Optional, Tuple
from .layout_extractor import DocumentLayout, PageLayout, TextBlock, BlockType

from config.logging_config import get_logger
logger = get_logger(__name__)

try:
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches
    from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False


class PDFReconstructor:
    """
    Rebuild PDFs/DOCX from translated layout

    Phase 3 Production Quality Features:
    - Preserve Layout Mode: PDF with original positioning
    - Reflow DOCX Mode: Structured single-column document
    - Block-type aware formatting (titles, headings, captions)
    - Multi-column layout support
    - Table and caption handling

    Modes:
    - rebuild_preserve_layout(): Maintains PDF layout
    - rebuild_reflow_docx(): Creates structured DOCX
    - rebuild_pdf(): Legacy method (backward compatible)
    """

    def __init__(
        self,
        default_font: str = "helv",  # Helvetica
        default_font_size: float = 12.0,
        preserve_formatting: bool = True
    ):
        """
        Initialize PDF reconstructor

        Args:
            default_font: Default font name (PyMuPDF font names)
            default_font_size: Default font size in points
            preserve_formatting: Attempt to preserve text formatting
        """
        self.default_font = default_font
        self.default_font_size = default_font_size
        self.preserve_formatting = preserve_formatting

    def rebuild_preserve_layout(
        self,
        original_layout: DocumentLayout,
        translated_blocks: Dict[int, List[str]],
        output_path: Path
    ) -> bool:
        """
        Rebuild PDF preserving original layout

        This mode maintains the original PDF layout including:
        - Page dimensions
        - Text positioning (approximate)
        - Multi-column layouts
        - Headers and footers

        Args:
            original_layout: Layout extracted from original PDF
            translated_blocks: Dict mapping page_num -> list of translated texts
            output_path: Path for output PDF

        Returns:
            True if successful

        Note:
            Text may overflow bboxes if translation is significantly longer.
            Font matching is approximate.
        """
        try:
            doc = fitz.open()

            for page_layout in original_layout.pages:
                page_num = page_layout.page_num

                if page_num not in translated_blocks:
                    logger.warning(f"No translation for page {page_num + 1}")
                    continue

                translated_texts = translated_blocks[page_num]

                # Create page with same dimensions
                page = doc.new_page(
                    width=page_layout.width,
                    height=page_layout.height
                )

                # Sort blocks by reading order
                sorted_blocks = sorted(page_layout.blocks, key=lambda b: b.reading_order)

                # Place each block
                for i, block in enumerate(sorted_blocks):
                    if i >= len(translated_texts):
                        break

                    translated_text = translated_texts[i]

                    # Skip headers/footers if requested (optional)
                    if block.block_type in [BlockType.HEADER, BlockType.FOOTER]:
                        # Could choose to skip or include
                        pass

                    # Determine font size (use original or scale if needed)
                    font_size = block.font_size if block.font_size > 0 else self.default_font_size

                    # Scale font if text is too long (basic auto-fit)
                    available_width = block.width
                    estimated_width = len(translated_text) * font_size * 0.6  # Rough estimate

                    if estimated_width > available_width and available_width > 0:
                        scale_factor = available_width / estimated_width
                        font_size *= max(scale_factor, 0.7)  # Don't go below 70%

                    # Place text at original position
                    try:
                        # Insert at baseline (y0 + font_size for baseline approximation)
                        baseline_y = block.y0 + font_size

                        # Use textbox for better wrapping
                        page.insert_textbox(
                            rect=block.bbox,
                            buffer=translated_text,
                            fontname=self.default_font,
                            fontsize=font_size,
                            color=(0, 0, 0),
                            align=0  # Left align
                        )
                    except Exception as e:
                        logger.warning(f"Failed to place block {i}: {e}")

            doc.save(output_path)
            doc.close()

            return True

        except Exception as e:
            logger.error(f"Error rebuilding PDF with layout preservation: {e}", exc_info=True)
            return False

    def rebuild_reflow_docx(
        self,
        original_layout: DocumentLayout,
        translated_blocks: Dict[int, List[str]],
        output_path: Path
    ) -> bool:
        """
        Rebuild as structured DOCX (single-column, reflowed)

        This mode creates a clean, structured document with:
        - Titles as Heading 1
        - Headings as Heading 2
        - Normal paragraphs
        - Captions as italic text
        - Tables preserved (simple representation)

        Args:
            original_layout: Layout extracted from original PDF
            translated_blocks: Dict mapping page_num -> list of translated texts
            output_path: Path for output DOCX

        Returns:
            True if successful

        Note:
            Requires python-docx package.
            Layout structure is not preserved (single column flow).
        """
        if not DOCX_AVAILABLE:
            logger.error("python-docx not available. Install with: pip install python-docx")
            return False

        try:
            doc = Document()

            # Configure default styles
            style = doc.styles['Normal']
            style.font.name = 'Arial'
            style.font.size = Pt(11)

            for page_layout in original_layout.pages:
                page_num = page_layout.page_num

                if page_num not in translated_blocks:
                    continue

                translated_texts = translated_blocks[page_num]

                # Sort blocks by reading order
                sorted_blocks = sorted(page_layout.blocks, key=lambda b: b.reading_order)

                for i, block in enumerate(sorted_blocks):
                    if i >= len(translated_texts):
                        break

                    translated_text = translated_texts[i]
                    block_type = block.block_type

                    # Skip headers/footers in reflowed document
                    if block_type in [BlockType.HEADER, BlockType.FOOTER]:
                        continue

                    # Add based on block type
                    if block_type == BlockType.TITLE:
                        # Title as Heading 1
                        p = doc.add_heading(translated_text, level=1)

                    elif block_type == BlockType.HEADING:
                        # Heading as Heading 2
                        p = doc.add_heading(translated_text, level=2)

                    elif block_type == BlockType.CAPTION:
                        # Caption as italic paragraph
                        p = doc.add_paragraph()
                        run = p.add_run(translated_text)
                        run.italic = True
                        run.font.size = Pt(10)
                        p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

                    elif block_type == BlockType.TABLE:
                        # Simple table representation (could be enhanced)
                        p = doc.add_paragraph(translated_text)
                        p.style = 'List Bullet'  # Or create custom table style

                    else:
                        # Normal paragraph
                        p = doc.add_paragraph(translated_text)

                # Add page break between pages (optional)
                if page_num < original_layout.total_pages - 1:
                    doc.add_page_break()

            doc.save(output_path)

            return True

        except Exception as e:
            logger.error(f"Error rebuilding DOCX: {e}", exc_info=True)
            return False

    def rebuild_pdf(
        self,
        original_layout: DocumentLayout,
        translated_blocks: Dict[int, List[str]],  # page_num -> list of translated texts
        output_path: Path
    ) -> bool:
        """
        Rebuild PDF from translated content

        Args:
            original_layout: Layout extracted from original PDF
            translated_blocks: Dictionary mapping page numbers to translated texts
            output_path: Path for output PDF

        Returns:
            True if successful, False otherwise

        Note:
            This is a basic implementation. Text is placed at original
            coordinates but may overflow or have spacing issues.
        """
        try:
            # Create new PDF document
            doc = fitz.open()

            # Reconstruct each page
            for page_layout in original_layout.pages:
                page_num = page_layout.page_num

                # Get translated texts for this page
                if page_num not in translated_blocks:
                    logger.warning(f"No translation for page {page_num + 1}")
                    continue

                translated_texts = translated_blocks[page_num]

                # Create new page with same dimensions
                page = doc.new_page(
                    width=page_layout.width,
                    height=page_layout.height
                )

                # Place each translated block
                for i, block in enumerate(page_layout.blocks):
                    if i >= len(translated_texts):
                        logger.warning(f"Not enough translations for page {page_num + 1}")
                        break

                    translated_text = translated_texts[i]

                    # Place text at original position
                    self._place_text_at_position(
                        page=page,
                        text=translated_text,
                        bbox=block.bbox,
                        font_size=block.font_size if hasattr(block, 'font_size') else self.default_font_size
                    )

            # Save PDF
            doc.save(output_path)
            doc.close()

            return True

        except Exception as e:
            logger.error(f"Error rebuilding PDF: {e}", exc_info=True)
            return False

    def _place_text_at_position(
        self,
        page: fitz.Page,
        text: str,
        bbox: Tuple[float, float, float, float],
        font_size: float = None
    ):
        """
        Place text at a specific position on the page

        Args:
            page: PyMuPDF page object
            text: Text to place
            bbox: Bounding box (x0, y0, x1, y1)
            font_size: Font size in points

        Note:
            This is a simplified implementation. Text may overflow
            the bounding box or have alignment issues.
        """
        x0, y0, x1, y1 = bbox
        font_size = font_size or self.default_font_size

        # Calculate available space
        available_width = x1 - x0
        available_height = y1 - y0

        # Insert text (basic implementation)
        # PyMuPDF insert_text places text at baseline, not top-left
        # So we adjust y0 to approximate the baseline
        baseline_y = y0 + font_size

        try:
            # Simple text insertion
            # Note: This may overflow the box - future work: text wrapping
            page.insert_text(
                point=(x0, baseline_y),
                text=text,
                fontname=self.default_font,
                fontsize=font_size,
                color=(0, 0, 0)  # Black text
            )
        except Exception as e:
            logger.warning(f"Failed to place text: {e}")

    def rebuild_pdf_simple(
        self,
        page_size: Tuple[float, float],
        translated_text: str,
        output_path: Path
    ) -> bool:
        """
        Create a simple PDF from translated text (no layout preservation)

        This is a fallback method when original layout is not available.

        Args:
            page_size: (width, height) in points
            translated_text: Full translated text
            output_path: Output PDF path

        Returns:
            True if successful
        """
        try:
            doc = fitz.open()
            width, height = page_size

            # Split text into pages (simple line-based splitting)
            lines_per_page = int(height / (self.default_font_size * 1.5))  # Rough estimate
            lines = translated_text.split('\n')

            current_page_lines = []
            for line in lines:
                current_page_lines.append(line)

                if len(current_page_lines) >= lines_per_page:
                    # Create page
                    page = doc.new_page(width=width, height=height)
                    page_text = '\n'.join(current_page_lines)

                    # Place text with margins
                    margin = 50
                    page.insert_textbox(
                        rect=(margin, margin, width - margin, height - margin),
                        buffer=page_text,
                        fontname=self.default_font,
                        fontsize=self.default_font_size,
                        color=(0, 0, 0)
                    )

                    current_page_lines = []

            # Add remaining lines
            if current_page_lines:
                page = doc.new_page(width=width, height=height)
                page_text = '\n'.join(current_page_lines)
                margin = 50
                page.insert_textbox(
                    rect=(margin, margin, width - margin, height - margin),
                    buffer=page_text,
                    fontname=self.default_font,
                    fontsize=self.default_font_size,
                    color=(0, 0, 0)
                )

            doc.save(output_path)
            doc.close()

            return True

        except Exception as e:
            logger.error(f"Error creating simple PDF: {e}")
            return False

    def estimate_text_dimensions(
        self,
        text: str,
        font_size: float
    ) -> Tuple[float, float]:
        """
        Estimate text dimensions (width, height)

        This is a rough approximation. Actual rendering may differ.

        Args:
            text: Text to measure
            font_size: Font size in points

        Returns:
            (width, height) in points
        """
        # Rough estimation: average character width is ~60% of font size
        avg_char_width = font_size * 0.6
        text_width = len(text) * avg_char_width

        # Height is approximately font size * number of lines
        lines = text.split('\n')
        text_height = len(lines) * font_size * 1.2  # 1.2 = line spacing factor

        return (text_width, text_height)

    def calculate_scale_factor(
        self,
        original_text: str,
        translated_text: str,
        bbox: Tuple[float, float, float, float],
        font_size: float
    ) -> float:
        """
        Calculate font size scale factor to fit translated text

        If translated text is longer/shorter than original,
        adjust font size to approximately fit the original bounding box.

        Args:
            original_text: Original text
            translated_text: Translated text
            bbox: Bounding box
            font_size: Original font size

        Returns:
            Adjusted font size

        Note:
            This is a simple heuristic. Real layout preservation
            requires more sophisticated text flow algorithms.
        """
        x0, y0, x1, y1 = bbox
        available_width = x1 - x0

        # Estimate original text width
        orig_width, _ = self.estimate_text_dimensions(original_text, font_size)

        # Estimate translated text width
        trans_width, _ = self.estimate_text_dimensions(translated_text, font_size)

        # If translated text is longer, scale down font
        if trans_width > available_width:
            scale = available_width / trans_width
            adjusted_font_size = font_size * scale
            # Don't go too small
            return max(adjusted_font_size, font_size * 0.7)

        return font_size


# Example usage
if __name__ == "__main__":
    import sys

    if len(sys.argv) < 3:
        print("Usage: python pdf_reconstructor.py <input.pdf> <output.pdf>")
        print("\nThis demo creates a simple reconstructed PDF.")
        sys.exit(1)

    input_path = Path(sys.argv[1])
    output_path = Path(sys.argv[2])

    if not input_path.exists():
        print(f"Error: Input file not found: {input_path}")
        sys.exit(1)

    print(f"Reconstructing PDF...")
    print(f"  Input: {input_path}")
    print(f"  Output: {output_path}")
    print("=" * 80)

    # Import layout extractor
    from layout_extractor import LayoutExtractor

    # Extract layout
    print("1. Extracting layout...")
    extractor = LayoutExtractor()
    layout = extractor.extract_layout(input_path)
    print(f"   ✓ Extracted {layout.total_pages} pages")

    # Simulate translation (just copy original text for demo)
    print("2. Simulating translation...")
    translated_blocks = {}
    for page in layout.pages:
        page_translations = [block.text + " [TRANSLATED]" for block in page.blocks]
        translated_blocks[page.page_num] = page_translations
    print(f"   ✓ Prepared {len(translated_blocks)} pages")

    # Rebuild PDF
    print("3. Rebuilding PDF...")
    reconstructor = PDFReconstructor()
    success = reconstructor.rebuild_pdf(
        original_layout=layout,
        translated_blocks=translated_blocks,
        output_path=output_path
    )

    if success:
        print(f"   ✓ PDF saved: {output_path}")
        print("\n" + "=" * 80)
        print("✓ Reconstruction complete!")
    else:
        print("   ✗ Failed to rebuild PDF")
        sys.exit(1)
