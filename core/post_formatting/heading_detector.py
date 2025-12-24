"""
Phase 2.0.6 - Heading Detection and Styling

Detects section headings in paragraphs and applies Word Heading styles (Heading 1/2/3).
This enables proper document navigation, TOC generation, and professional appearance.

Key Features:
- Pattern-based heading detection (numbered sections, named sections)
- Bilingual support (English + Vietnamese)
- Context-aware (doesn't convert theorems to headings)
- Preserves existing content, only changes style
"""

import re
from typing import Optional, Tuple
from docx import Document
from docx.text.paragraph import Paragraph
import logging

logger = logging.getLogger(__name__)


class HeadingDetector:
    """
    Detects section headings in document text.

    Strategies:
    1. Numbered sections (1., 1.1, 1.1.1)
    2. Named sections (Introduction, References, etc.)
    3. Vietnamese equivalents (Giới thiệu, Tài liệu tham khảo)

    Anti-patterns (NOT headings):
    - Theorems, Lemmas, Definitions (e.g., "Theorem 1.1")
    - Proofs
    - Equations
    """

    # Numbered section patterns
    NUMBERED_SECTION_PATTERNS = [
        # Level 1: "1. Introduction" or "1 Introduction"
        (r'^(\d+)\.\s+([A-ZĐ][\w\s]+)', 1),
        (r'^(\d+)\s+([A-ZĐ][\w\s]+)', 1),

        # Level 2: "1.1 Background" or "1.1. Background"
        (r'^(\d+)\.(\d+)\.?\s+([A-ZĐ][\w\s]+)', 2),

        # Level 3: "1.1.1 Details" or "1.1.1. Details"
        (r'^(\d+)\.(\d+)\.(\d+)\.?\s+', 3),
    ]

    # Named section patterns (major sections)
    MAJOR_SECTIONS_EN = [
        'Abstract', 'Introduction', 'Background', 'Methodology', 'Methods',
        'Results', 'Discussion', 'Conclusion', 'Acknowledgments',
        'References', 'Bibliography', 'Appendix'
    ]

    MAJOR_SECTIONS_VI = [
        'Tóm tắt', 'Giới thiệu', 'Phương pháp', 'Kết quả', 'Thảo luận',
        'Kết luận', 'Lời cảm ơn', 'Tài liệu tham khảo', 'Phụ lục'
    ]

    # Theorem-like keywords (NOT headings)
    THEOREM_KEYWORDS = [
        'Theorem', 'Lemma', 'Corollary', 'Proposition', 'Definition',
        'Example', 'Remark', 'Proof', 'Problem', 'Exercise',
        # Vietnamese
        'Định lý', 'Bổ đề', 'Hệ quả', 'Mệnh đề', 'Định nghĩa',
        'Ví dụ', 'Nhận xét', 'Chứng minh', 'Bài tập'
    ]

    def __init__(self):
        """Initialize heading detector."""
        pass

    def detect_heading_level(self, text: str) -> Optional[int]:
        """
        Detect if text is a heading and return its level (1, 2, or 3).

        Args:
            text: Paragraph text to analyze

        Returns:
            int: Heading level (1, 2, or 3) if detected, None otherwise

        Examples:
            "1. Introduction" → 1
            "1.1 Background" → 2
            "1.1.1 Details" → 3
            "Introduction" → 1
            "Theorem 1.1" → None (not a heading)
        """
        text = text.strip()
        if not text:
            return None

        # Check if it's a theorem-like structure (NOT a heading)
        if self._is_theorem_like(text):
            return None

        # Check numbered sections
        for pattern, level in self.NUMBERED_SECTION_PATTERNS:
            if re.match(pattern, text):
                return level

        # Check named major sections (English)
        for section_name in self.MAJOR_SECTIONS_EN:
            if self._is_major_section(text, section_name):
                return 1

        # Check named major sections (Vietnamese)
        for section_name in self.MAJOR_SECTIONS_VI:
            if self._is_major_section(text, section_name):
                return 1

        return None

    def _is_theorem_like(self, text: str) -> bool:
        """
        Check if text starts with theorem-like keywords.

        Args:
            text: Text to check

        Returns:
            bool: True if theorem-like, False otherwise
        """
        for keyword in self.THEOREM_KEYWORDS:
            # Match "Theorem 1.1" but not "Theorems are..."
            pattern = rf'^{re.escape(keyword)}\s+\d+'
            if re.match(pattern, text):
                return True
        return False

    def _is_major_section(self, text: str, section_name: str) -> bool:
        """
        Check if text is a major section heading.

        Args:
            text: Text to check
            section_name: Section name to match

        Returns:
            bool: True if matches, False otherwise

        Matching rules:
        - Exact match: "Introduction"
        - Case-insensitive: "introduction", "INTRODUCTION"
        - With trailing text: "Introduction to Graph Theory" (for subsections)
        """
        text_lower = text.lower()
        section_lower = section_name.lower()

        # Exact match or starts with section name
        if text_lower == section_lower or text_lower.startswith(section_lower + ' '):
            return True

        return False


class HeadingStyleApplicator:
    """
    Applies Word Heading styles to detected headings.

    Modifies paragraph style to Heading 1/2/3 based on detected level.
    Preserves text content, only changes formatting.
    """

    def __init__(self, detector: Optional[HeadingDetector] = None):
        """
        Initialize heading style applicator.

        Args:
            detector: HeadingDetector instance (creates new if None)
        """
        self.detector = detector or HeadingDetector()

    def apply_heading_styles(self, doc: Document) -> int:
        """
        Apply Heading styles to all detected headings in document.

        Args:
            doc: python-docx Document object

        Returns:
            int: Number of headings detected and styled

        Safety:
        - Read-only first pass to detect headings
        - Only applies styles to detected headings
        - Preserves text content
        - Graceful error handling (continues on failure)
        """
        headings_count = 0

        for para in doc.paragraphs:
            # Skip empty paragraphs
            if not para.text.strip():
                continue

            # Detect heading level
            level = self.detector.detect_heading_level(para.text)

            if level:
                try:
                    # Apply heading style
                    para.style = f'Heading {level}'
                    headings_count += 1
                    logger.debug(f"Applied Heading {level}: {para.text[:50]}...")

                except Exception as e:
                    logger.warning(f"Failed to apply Heading {level} to '{para.text[:50]}': {e}")
                    # Continue processing other paragraphs

        logger.info(f"Applied Heading styles to {headings_count} paragraphs")
        return headings_count

    def preview_headings(self, doc: Document) -> list[Tuple[str, int]]:
        """
        Preview detected headings without modifying document.

        Args:
            doc: python-docx Document object

        Returns:
            list: List of (text, level) tuples for detected headings

        Usage:
            headings = applicator.preview_headings(doc)
            for text, level in headings:
                print(f"Heading {level}: {text}")
        """
        headings = []

        for para in doc.paragraphs:
            level = self.detector.detect_heading_level(para.text)
            if level:
                headings.append((para.text, level))

        return headings


# Convenience function for quick usage
def apply_heading_styles_to_document(doc: Document) -> int:
    """
    Quick function to apply heading styles to document.

    Args:
        doc: python-docx Document object

    Returns:
        int: Number of headings styled

    Example:
        from docx import Document
        from core.post_formatting.heading_detector import apply_heading_styles_to_document

        doc = Document("input.docx")
        count = apply_heading_styles_to_document(doc)
        doc.save("output.docx")
        print(f"Styled {count} headings")
    """
    applicator = HeadingStyleApplicator()
    return applicator.apply_heading_styles(doc)
