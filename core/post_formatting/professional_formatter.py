"""
Phase 2.0.7 - Professional Document Formatter

Main integration point for post-formatting Phase 2.0.5 output.
Transforms "acceptable for internal use" to "exemplary Word document" quality.

Phase 2.0.7 Wave 1 Features:
1. Typography System - Proper font size hierarchy (18/16/14/12/11pt)
2. Enhanced Heading Formatting - Visual hierarchy with sizes
3. Remove technical footers
4. Clean page headers
5. Fix Unicode artifacts and spacing
6. Professional paragraph formatting

Future Wave 2:
- Title/Abstract detection and formatting
- Paragraph merging (complex heuristics)
- Advanced layout optimization
"""

import re
import logging
from dataclasses import dataclass
from typing import Literal, Optional
from pathlib import Path
from docx import Document
from docx.shared import Pt, Inches
from core.post_formatting.heading_detector import HeadingStyleApplicator
from core.post_formatting.typography_system import TypographyManager, AcademicTypography

logger = logging.getLogger(__name__)


@dataclass
class PostFormattingConfig:
    """
    Configuration for Phase 2.0.6 post-formatting.

    Attributes:
        enable_heading_detection: Apply Heading 1/2/3 styles to sections
        remove_technical_footer: Remove "AI Translator Pro | Page N" footers
        clean_page_headers: Remove "Translate *.pdf" from headers
        fix_unicode_artifacts: Fix broken characters (e.g., "Erd ős" → "Erdős")
        paragraph_style: Paragraph formatting style
            - 'traditional': First-line indent (0.3")
            - 'modern': Block with spacing (12pt after)
            - 'none': No changes (Phase 2.0.5 defaults)
    """
    enable_heading_detection: bool = True
    remove_technical_footer: bool = True
    clean_page_headers: bool = True
    fix_unicode_artifacts: bool = True
    paragraph_style: Literal['traditional', 'modern', 'none'] = 'modern'


class ProfessionalFormatter:
    """
    Main formatter for Phase 2.0.7 post-processing.

    Usage:
        formatter = ProfessionalFormatter(config=PostFormattingConfig())
        formatter.process_document("input.docx", "output.docx")
    """

    def __init__(self, config: Optional[PostFormattingConfig] = None):
        """
        Initialize formatter.

        Args:
            config: Post-formatting configuration (uses defaults if None)
        """
        self.config = config or PostFormattingConfig()
        self.heading_applicator = HeadingStyleApplicator()
        self.typography_manager = TypographyManager()  # Phase 2.0.7

    def process_document(
        self,
        input_path: str,
        output_path: Optional[str] = None,
        force_apply: bool = False
    ) -> bool:
        """
        Process DOCX document with all enabled post-formatting steps.

        Args:
            input_path: Path to input DOCX (Phase 2.0.5 output)
            output_path: Path to save output (overwrites input if None)
            force_apply: If True, force formatting even if conditions not met (Phase 2.0.8)

        Returns:
            bool: True if processing succeeded, False otherwise

        Steps:
            1. Load document
            2. Apply heading styles
            3. Clean headers/footers
            4. Fix Unicode artifacts
            5. Apply paragraph formatting
            6. Save document
        """
        # Default: overwrite input (in-place processing)
        if output_path is None:
            output_path = input_path

        logger.info(f"Starting Phase 2.0.7 post-formatting: {input_path}")

        try:
            # Load document
            doc = Document(input_path)

            # Step 1: Apply heading styles
            if self.config.enable_heading_detection:
                headings_count = self.heading_applicator.apply_heading_styles(doc)
                logger.info(f"Applied Heading styles to {headings_count} paragraphs")

                # Phase 2.0.7: Enhance headings with professional typography
                typography_count = 0
                for para in doc.paragraphs:
                    if para.style.name.startswith('Heading'):
                        # Extract level from style name "Heading 1" → 1
                        level = int(para.style.name.split()[-1])
                        self.typography_manager.enhance_heading(para, level)
                        typography_count += 1
                logger.info(f"Enhanced typography for {typography_count} headings")

            # Step 2: Remove technical footers
            if self.config.remove_technical_footer:
                self._remove_technical_footers(doc)

            # Step 3: Clean page headers
            if self.config.clean_page_headers:
                self._clean_page_headers(doc)

            # Step 4: Fix Unicode artifacts
            if self.config.fix_unicode_artifacts:
                self._fix_unicode_artifacts(doc)

            # Step 5: Apply paragraph formatting
            if self.config.paragraph_style != 'none':
                self._apply_paragraph_formatting(doc)

            # Phase 2.0.7: Apply professional typography to body paragraphs
            body_typo_count = self.typography_manager.enhance_body_paragraphs(doc)
            logger.info(f"Phase 2.0.7: Enhanced typography for {body_typo_count} body paragraphs")

            # Save
            doc.save(output_path)
            logger.info(f"Phase 2.0.7 post-formatting complete: {output_path}")

            # Phase 2.0.8: Return success status
            return True

        except Exception as e:
            logger.error(f"Failed to process document {input_path}: {e}")
            # Phase 2.0.8: Re-raise if force_apply, otherwise return False
            if force_apply:
                raise
            return False

    def _remove_technical_footers(self, doc: Document) -> int:
        """
        Remove "AI Translator Pro | Page N" footers.

        Args:
            doc: python-docx Document

        Returns:
            int: Number of footers cleaned
        """
        cleaned_count = 0

        for section in doc.sections:
            footer = section.footer
            for para in footer.paragraphs:
                # Check if paragraph contains technical footer text
                if self._is_technical_footer(para.text):
                    # Clear paragraph (remove all runs)
                    para.clear()
                    cleaned_count += 1
                    logger.debug(f"Removed technical footer: {para.text[:50]}")

        logger.info(f"Removed {cleaned_count} technical footers")
        return cleaned_count

    def _is_technical_footer(self, text: str) -> bool:
        """
        Check if text is a technical footer.

        Args:
            text: Footer paragraph text

        Returns:
            bool: True if technical footer, False otherwise

        Patterns:
            - "AI Translator Pro"
            - "Page 1", "Page 2", etc.
            - Combination: "AI Translator Pro | Page 1"
        """
        patterns = [
            r'AI Translator Pro',
            r'Page\s+\d+',
            r'\|\s*Page\s+\d+',
        ]

        for pattern in patterns:
            if re.search(pattern, text):
                return True

        return False

    def _clean_page_headers(self, doc: Document) -> int:
        """
        Clean page headers (remove "Translate *.pdf" artifacts).

        Args:
            doc: python-docx Document

        Returns:
            int: Number of headers cleaned
        """
        cleaned_count = 0

        for section in doc.sections:
            header = section.header
            for para in header.paragraphs:
                original_text = para.text

                # Remove "Translate filename.pdf" pattern
                cleaned_text = re.sub(r'Translate\s+\w+\.pdf\s*', '', original_text)

                # Remove empty lines and excessive whitespace
                cleaned_text = re.sub(r'\s+', ' ', cleaned_text).strip()

                # Only update if text changed
                if cleaned_text != original_text:
                    para.text = cleaned_text
                    cleaned_count += 1
                    logger.debug(f"Cleaned header: '{original_text}' → '{cleaned_text}'")

        logger.info(f"Cleaned {cleaned_count} page headers")
        return cleaned_count

    def _fix_unicode_artifacts(self, doc: Document) -> int:
        """
        Fix Unicode artifacts and spacing issues in all paragraphs.

        Args:
            doc: python-docx Document

        Returns:
            int: Number of paragraphs fixed

        Fixes:
            - "Erd ő s" → "Erdős"
            - Multiple spaces → Single space
            - Space before punctuation
            - Space in brackets
        """
        fixed_count = 0

        for para in doc.paragraphs:
            original_text = para.text

            # Apply Unicode fixes
            fixed_text = self._normalize_text(original_text)

            # Only update if text changed
            if fixed_text != original_text:
                para.text = fixed_text
                fixed_count += 1
                logger.debug(f"Fixed Unicode: '{original_text[:50]}' → '{fixed_text[:50]}'")

        logger.info(f"Fixed Unicode artifacts in {fixed_count} paragraphs")
        return fixed_count

    def _normalize_text(self, text: str) -> str:
        """
        Normalize text with Unicode and spacing fixes.

        Args:
            text: Original text

        Returns:
            str: Normalized text

        Transformations:
            - Fix split diacritics: "Erd ő s" → "Erdős"
            - Multiple spaces → single space
            - Remove space before punctuation: "text ." → "text."
            - Fix brackets: "( text )" → "(text)"
        """
        # Multiple spaces → single space
        text = re.sub(r'\s{2,}', ' ', text)

        # Remove space before punctuation
        text = re.sub(r'\s+([.,;:!?])', r'\1', text)

        # Fix space after opening bracket
        text = re.sub(r'\(\s+', '(', text)

        # Fix space before closing bracket
        text = re.sub(r'\s+\)', ')', text)

        # Fix split diacritics (common in PDF OCR)
        # Pattern: letter + space + combining character + space + letter
        # Example: "Erd ő s" has "ő" as combining char
        text = re.sub(r'(\w)\s+([\u0300-\u036f])\s+(\w)', r'\1\2\3', text)

        return text.strip()

    def _apply_paragraph_formatting(self, doc: Document) -> int:
        """
        Apply professional paragraph formatting.

        Args:
            doc: python-docx Document

        Returns:
            int: Number of paragraphs formatted

        Formatting:
            - Traditional: First-line indent (0.3")
            - Modern: Block with spacing (12pt after)

        Skips:
            - Headings (Heading 1/2/3)
            - Theorems (has borders from Phase 2.0.5)
            - Proofs (indented from Phase 2.0.5)
            - Equations (centered)
        """
        formatted_count = 0

        for para in doc.paragraphs:
            # Skip if empty
            if not para.text.strip():
                continue

            # Skip headings
            if para.style.name.startswith('Heading'):
                continue

            # Skip if already has special formatting (theorem boxes, etc.)
            # These have borders or background from Phase 2.0.5
            if self._has_special_formatting(para):
                continue

            # Skip centered paragraphs (equations)
            if para.alignment == 1:  # WD_ALIGN_PARAGRAPH.CENTER = 1
                continue

            # Apply paragraph style
            if self.config.paragraph_style == 'traditional':
                # First-line indent
                para.paragraph_format.first_line_indent = Inches(0.3)
                para.paragraph_format.space_after = Pt(0)
                formatted_count += 1

            elif self.config.paragraph_style == 'modern':
                # Block with spacing
                para.paragraph_format.first_line_indent = Inches(0)
                para.paragraph_format.space_after = Pt(12)
                formatted_count += 1

        logger.info(f"Applied {self.config.paragraph_style} formatting to {formatted_count} paragraphs")
        return formatted_count

    def _has_special_formatting(self, para) -> bool:
        """
        Check if paragraph has special formatting (borders, shading, etc.).

        Args:
            para: python-docx Paragraph

        Returns:
            bool: True if has special formatting, False otherwise

        Detection:
            - Has borders (theorem boxes)
            - Has background shading
            - Is indented (proofs)
        """
        try:
            pPr = para._element.get_or_add_pPr()

            # Check for borders
            if pPr.find('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}pBdr') is not None:
                return True

            # Check for shading
            if pPr.find('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}shd') is not None:
                return True

            # Check for significant indentation (proof blocks)
            if para.paragraph_format.left_indent and para.paragraph_format.left_indent > Inches(0.1):
                return True

        except Exception:
            pass

        return False


# Convenience function for quick usage
def apply_professional_formatting(
    input_path: str,
    output_path: Optional[str] = None,
    config: Optional[PostFormattingConfig] = None
) -> None:
    """
    Quick function to apply Phase 2.0.7 post-formatting.

    Args:
        input_path: Path to input DOCX (Phase 2.0.5 output)
        output_path: Path to save output (overwrites input if None)
        config: Post-formatting configuration (uses defaults if None)

    Example:
        from core.post_formatting.professional_formatter import apply_professional_formatting

        # Process with defaults (overwrite input)
        apply_professional_formatting("translated.docx")

        # Process with custom config
        from core.post_formatting.professional_formatter import PostFormattingConfig
        config = PostFormattingConfig(paragraph_style='traditional')
        apply_professional_formatting("input.docx", "output.docx", config)
    """
    formatter = ProfessionalFormatter(config=config)
    formatter.process_document(input_path, output_path)
