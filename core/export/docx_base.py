#!/usr/bin/env python3
# -*- coding: utf-8 -*-

"""
DOCX Base Exporter - Phase 1.4 Consolidation
=============================================

Abstract base class for all DOCX exporters.
Provides common functionality:
- Document creation and setup
- Page layout configuration
- Metadata handling
- Style application
- Save operations
"""

from abc import ABC, abstractmethod
from dataclasses import dataclass, field
from pathlib import Path
from typing import Optional, Dict, Any, List

from docx import Document
from docx.shared import Pt, Inches

from .docx_styles import (
    PAGE_MARGINS,
    PAGE_SIZES,
    DEFAULT_OPTIONS,
    AcademicStyles,
    StyleApplicator,
)
from config.logging_config import get_logger

logger = get_logger(__name__)


@dataclass
class ExportConfig:
    """
    Base configuration for DOCX export.

    All exporter configurations should extend this class.
    """
    # Page layout
    page_size: str = 'a4'
    margin_style: str = 'academic'

    # Typography
    font_name: str = 'Times New Roman'
    font_size: int = 11
    line_spacing: float = 1.15

    # Paragraph spacing (in points)
    paragraph_before: int = 6
    paragraph_after: int = 6

    # Document metadata
    title: str = ''
    author: str = ''
    subject: str = ''

    # Features
    include_toc: bool = False
    include_page_numbers: bool = True

    def to_dict(self) -> Dict[str, Any]:
        """Convert config to dictionary."""
        return {
            'page_size': self.page_size,
            'margin_style': self.margin_style,
            'font_name': self.font_name,
            'font_size': self.font_size,
            'line_spacing': self.line_spacing,
            'paragraph_before': self.paragraph_before,
            'paragraph_after': self.paragraph_after,
            'title': self.title,
            'author': self.author,
            'subject': self.subject,
            'include_toc': self.include_toc,
            'include_page_numbers': self.include_page_numbers,
        }

    @classmethod
    def from_dict(cls, data: Dict[str, Any]) -> 'ExportConfig':
        """Create config from dictionary."""
        return cls(**{k: v for k, v in data.items() if k in cls.__dataclass_fields__})


class DocxExporterBase(ABC):
    """
    Abstract base class for DOCX exporters.

    Provides common functionality for document creation, styling, and export.
    Subclasses must implement the `_add_content` method.
    """

    def __init__(self, config: Optional[ExportConfig] = None):
        """
        Initialize exporter with configuration.

        Args:
            config: Export configuration (uses defaults if None)
        """
        self.config = config or ExportConfig()
        self.doc: Optional[Document] = None

    def export(
        self,
        content: Any,
        output_path: str,
        metadata: Optional[Dict[str, str]] = None
    ) -> str:
        """
        Export content to DOCX file.

        Args:
            content: Content to export (type depends on subclass)
            output_path: Path to output DOCX file
            metadata: Optional document metadata

        Returns:
            str: Absolute path to created file
        """
        # Create new document
        self.doc = Document()

        # Apply page layout
        self._setup_page_layout()

        # Apply base styles
        self._setup_base_styles()

        # Set metadata
        self._setup_metadata(metadata or {})

        # Add content (implemented by subclass)
        self._add_content(content)

        # Save document
        self.doc.save(output_path)

        # Return absolute path
        import os
        return os.path.abspath(output_path)

    @abstractmethod
    def _add_content(self, content: Any) -> None:
        """
        Add content to document.

        Must be implemented by subclasses.

        Args:
            content: Content to add (type depends on implementation)
        """
        pass

    def _setup_page_layout(self) -> None:
        """Configure page size and margins."""
        if not self.doc:
            return

        # Get page size
        page_size = PAGE_SIZES.get(
            self.config.page_size,
            PAGE_SIZES['a4']
        )

        # Get margins
        margins = PAGE_MARGINS.get(
            self.config.margin_style,
            PAGE_MARGINS['default']
        )

        # Apply to all sections
        for section in self.doc.sections:
            # Page size
            section.page_width = page_size['width']
            section.page_height = page_size['height']

            # Margins
            section.top_margin = margins['top']
            section.bottom_margin = margins['bottom']
            section.left_margin = margins['left']
            section.right_margin = margins['right']

    def _setup_base_styles(self) -> None:
        """Configure base document styles (Normal, Headings)."""
        if not self.doc:
            return

        styles = self.doc.styles

        # Normal style
        normal_style = styles['Normal']
        normal_style.font.name = self.config.font_name
        normal_style.font.size = Pt(self.config.font_size)
        normal_style.paragraph_format.line_spacing = self.config.line_spacing
        normal_style.paragraph_format.space_before = Pt(self.config.paragraph_before)
        normal_style.paragraph_format.space_after = Pt(self.config.paragraph_after)

        # Heading styles
        for level in range(1, 4):
            heading_name = f'Heading {level}'
            if heading_name in styles:
                heading = styles[heading_name]
                heading.font.name = self.config.font_name
                heading.font.size = Pt(14 - level + 1)  # 14pt, 13pt, 12pt
                heading.font.bold = True
                heading.paragraph_format.space_before = Pt(12)
                heading.paragraph_format.space_after = Pt(6)

    def _setup_metadata(self, metadata: Dict[str, str]) -> None:
        """Set document metadata (title, author, subject)."""
        if not self.doc:
            return

        core_props = self.doc.core_properties

        # Use metadata dict or fall back to config
        core_props.title = metadata.get('title', self.config.title) or 'Document'
        core_props.author = metadata.get('author', self.config.author) or 'Academic Translator'
        core_props.subject = metadata.get('subject', self.config.subject) or ''

    def add_paragraph(self, text: str, style: str = 'Normal') -> None:
        """
        Add a paragraph to the document.

        Args:
            text: Paragraph text
            style: Style name (default: 'Normal')
        """
        if self.doc:
            self.doc.add_paragraph(text, style=style)

    def add_heading(self, text: str, level: int = 1) -> None:
        """
        Add a heading to the document.

        Args:
            text: Heading text
            level: Heading level (1-9)
        """
        if self.doc:
            self.doc.add_heading(text, level=level)

    def add_page_break(self) -> None:
        """Add a page break."""
        if self.doc:
            self.doc.add_page_break()
