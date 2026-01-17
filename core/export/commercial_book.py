#!/usr/bin/env python3
# -*- coding: utf-8 -*-

"""
Commercial Book Exporter
========================

Generates professional ebook-quality documents matching major publishers:
- Penguin Random House, Simon & Schuster, Doubleday style
- Proper title page with elegant typography
- Chapter headings with "CHAPTER 1" style and letter-spacing
- Drop caps for chapter openings
- First-line indent for paragraphs
- Scene breaks with ornaments

Reference: Dan Brown's commercial ebook layout

Author: AI Publisher Pro
Version: 1.0.0
"""

import re
from pathlib import Path
from typing import Optional, Dict, List, Tuple, Any
from dataclasses import dataclass

# DOCX support
try:
    from docx import Document
    from docx.shared import Pt, Inches, Cm, RGBColor, Twips
    from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
    from docx.enum.section import WD_ORIENT
    from docx.enum.style import WD_STYLE_TYPE
    from docx.oxml.ns import qn
    from docx.oxml import parse_xml
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

# Title formatter for smart title extraction
try:
    from core.utils.title_formatter import (
        extract_title_from_filename,
        format_document_title,
        clean_header_text
    )
    TITLE_FORMATTER_AVAILABLE = True
except ImportError:
    TITLE_FORMATTER_AVAILABLE = False

from config.logging_config import get_logger

logger = get_logger(__name__)


@dataclass
class BookConfig:
    """Configuration for commercial book export."""
    # Book metadata
    title: str = "Untitled"
    author: str = ""
    subtitle: str = ""
    publisher: str = ""
    copyright_year: str = ""
    isbn: str = ""

    # Page layout - Trade paperback (5.5" x 8.5") - professional standard
    page_width: float = 5.5  # inches
    page_height: float = 8.5  # inches
    margin_top: float = 0.9  # inches - tighter for more content
    margin_bottom: float = 0.9
    margin_inner: float = 1.0  # binding side - slightly less for better balance
    margin_outer: float = 0.75

    # Typography - Professional book standards
    font_body: str = "Times New Roman"
    font_heading: str = "Times New Roman"
    font_size_body: int = 11  # Standard for trade paperback
    font_size_chapter: int = 14
    line_spacing: float = 1.15  # Tighter for professional look (was 1.5)
    paragraph_indent: float = 0.25  # inches - standard book indent
    paragraph_spacing: float = 0  # No extra space between paragraphs

    # Features
    include_title_page: bool = True
    include_copyright_page: bool = False
    include_toc: bool = False
    include_running_header: bool = True  # Add running headers
    chapter_style: str = "commercial"  # commercial, literary, minimal
    drop_cap_enabled: bool = True
    scene_break_style: str = "line"  # line, dots, blank (minimalist default)

    # Advanced typography
    widow_orphan_control: bool = True  # Prevent single lines at page breaks
    hyphenation: bool = False  # Disable for cleaner look
    justify_text: bool = True  # Full justification


class CommercialBookExporter:
    """
    Exports documents in commercial ebook format.

    Creates professional book layout matching major publishers:
    - Title page with elegant typography
    - CHAPTER 1 style headings
    - Drop caps
    - First-line indent
    - Scene ornaments
    """

    # Chapter patterns for detection
    CHAPTER_PATTERNS = [
        # English
        (r'^Chapter\s+(\d+|[IVXLC]+)\s*[:\.]?\s*(.*)', 'en'),
        (r'^CHAPTER\s+(\d+|[IVXLC]+)\s*[:\.]?\s*(.*)', 'en'),
        # Vietnamese
        (r'^Chương\s+(\d+|[IVXLC]+)\s*[:\.]?\s*(.*)', 'vi'),
        (r'^CHƯƠNG\s+(\d+|[IVXLC]+)\s*[:\.]?\s*(.*)', 'vi'),
        (r'^Phần\s+(\d+|[IVXLC]+)\s*[:\.]?\s*(.*)', 'vi'),
        (r'^Hồi\s+(\d+|[IVXLC]+)\s*[:\.]?\s*(.*)', 'vi'),
        # Markdown headers
        (r'^#\s+(.+)', 'md'),
    ]

    SCENE_BREAK_PATTERNS = [
        r'^\s*\*\s*\*\s*\*\s*$',
        r'^\s*\*{3,}\s*$',
        r'^\s*-{3,}\s*$',
        r'^\s*•\s*•\s*•\s*$',
        r'^\s*\.\s*\.\s*\.\s*$',
    ]

    SCENE_BREAK_ORNAMENTS = {
        'dots': '•   •   •',
        'ornament': '❧',
        'line': '───────',
        'stars': '✦  ✦  ✦',
        'blank': '',
    }

    # Vietnamese literature title patterns for first-line detection
    # Pattern: "Title Author Content..." where Title and Author are known
    VIETNAMESE_TITLE_AUTHOR_PATTERNS = [
        # "Chí Phèo Nam Cao ..." - title followed by author at start
        (r'^(Chí Phèo)\s+(Nam Cao)\s+(.+)', 'Chí Phèo', 'Nam Cao'),
        (r'^(Truyện Kiều)\s+(Nguyễn Du)\s+(.+)', 'Truyện Kiều', 'Nguyễn Du'),
        (r'^(Tắt Đèn)\s+(Ngô Tất Tố)\s+(.+)', 'Tắt Đèn', 'Ngô Tất Tố'),
        (r'^(Số Đỏ)\s+(Vũ Trọng Phụng)\s+(.+)', 'Số Đỏ', 'Vũ Trọng Phụng'),
        (r'^(Lão Hạc)\s+(Nam Cao)\s+(.+)', 'Lão Hạc', 'Nam Cao'),
        (r'^(Đời Thừa)\s+(Nam Cao)\s+(.+)', 'Đời Thừa', 'Nam Cao'),
        (r'^(Sống Mòn)\s+(Nam Cao)\s+(.+)', 'Sống Mòn', 'Nam Cao'),
        (r'^(Nỗi Buồn Chiến Tranh)\s+(Bảo Ninh)\s+(.+)', 'Nỗi Buồn Chiến Tranh', 'Bảo Ninh'),
    ]

    def __init__(self, config: Optional[BookConfig] = None):
        """Initialize exporter with configuration."""
        if not DOCX_AVAILABLE:
            raise ImportError("python-docx required for commercial book export")

        self.config = config or BookConfig()
        self.doc = None
        self.chapter_count = 0

    def export(
        self,
        content: str,
        output_path: str,
        source_filename: Optional[str] = None
    ) -> str:
        """
        Export content as commercial book.

        Args:
            content: Book content (plain text or markdown)
            output_path: Output DOCX path
            source_filename: Original filename for smart title extraction

        Returns:
            Path to generated DOCX file
        """
        # Step 1: Try to detect title/author from content FIRST
        # This handles cases like "Chí Phèo Nam Cao He was walking..."
        detected_title, detected_author, cleaned_content = self._extract_title_from_first_line(content)

        if detected_title:
            self.config.title = detected_title
        if detected_author:
            self.config.author = detected_author

        # Step 2: Fallback to filename-based extraction if still no title
        if source_filename and TITLE_FORMATTER_AVAILABLE:
            if not self.config.title or self.config.title == "Untitled":
                title, author = extract_title_from_filename(source_filename)
                self.config.title = title
                if not self.config.author and author:
                    self.config.author = author

        # Create document
        self.doc = Document()
        self._setup_page_layout()
        self._create_styles()

        # Add front matter (now with detected title/author)
        if self.config.include_title_page:
            self._add_title_page()

        if self.config.include_copyright_page:
            self._add_copyright_page()

        # Parse and add content (using cleaned content without title/author header)
        self._add_book_content_internal(cleaned_content)

        # Save document
        self.doc.save(output_path)
        logger.info(f"Commercial book exported: {output_path}")

        return output_path

    def _setup_page_layout(self):
        """Configure page size and margins for book format."""
        section = self.doc.sections[0]

        # Page size - Trade paperback
        section.page_width = Inches(self.config.page_width)
        section.page_height = Inches(self.config.page_height)

        # Margins - Professional book layout
        section.top_margin = Inches(self.config.margin_top)
        section.bottom_margin = Inches(self.config.margin_bottom)
        section.left_margin = Inches(self.config.margin_inner)  # inner/binding
        section.right_margin = Inches(self.config.margin_outer)  # outer

        # Gutter for binding
        section.gutter = Inches(0.15)

        # Header/Footer distance
        section.header_distance = Inches(0.4)
        section.footer_distance = Inches(0.4)

        # Add running header if enabled
        if self.config.include_running_header:
            self._add_running_header(section)

    def _add_running_header(self, section):
        """Add running header with author name on left pages, title on right."""
        header = section.header
        header.is_linked_to_previous = False

        # Create header paragraph
        header_para = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
        header_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Add title (simplified - both pages same for now)
        header_text = self.config.title.upper() if self.config.title else ""
        if self.config.author:
            header_text = f"{self.config.author}  •  {header_text}"

        run = header_para.add_run(header_text)
        run.font.name = self.config.font_body
        run.font.size = Pt(8)
        run.font.small_caps = True
        run.font.color.rgb = RGBColor(100, 100, 100)

    def _create_styles(self):
        """Create custom styles for commercial book layout."""
        styles = self.doc.styles

        # Body text style - Professional book typography
        body_style = styles['Normal']
        body_style.font.name = self.config.font_body
        body_style.font.size = Pt(self.config.font_size_body)
        body_style.paragraph_format.line_spacing = self.config.line_spacing
        body_style.paragraph_format.first_line_indent = Inches(self.config.paragraph_indent)
        body_style.paragraph_format.space_after = Pt(self.config.paragraph_spacing)
        body_style.paragraph_format.space_before = Pt(0)
        body_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY if self.config.justify_text else WD_ALIGN_PARAGRAPH.LEFT

        # Widow/orphan control
        if self.config.widow_orphan_control:
            body_style.paragraph_format.widow_control = True

        # Title page - Book title
        if 'BookTitle' not in [s.name for s in styles]:
            title_style = styles.add_style('BookTitle', WD_STYLE_TYPE.PARAGRAPH)
            title_style.font.name = self.config.font_heading
            title_style.font.size = Pt(28)
            title_style.font.bold = False
            title_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
            title_style.paragraph_format.space_after = Pt(12)
            title_style.paragraph_format.first_line_indent = Pt(0)

        # Title page - Author name
        if 'BookAuthor' not in [s.name for s in styles]:
            author_style = styles.add_style('BookAuthor', WD_STYLE_TYPE.PARAGRAPH)
            author_style.font.name = self.config.font_heading
            author_style.font.size = Pt(16)
            author_style.font.bold = False
            author_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
            author_style.paragraph_format.space_before = Pt(24)
            author_style.paragraph_format.first_line_indent = Pt(0)

        # Chapter heading - "CHAPTER"
        if 'ChapterLabel' not in [s.name for s in styles]:
            ch_label = styles.add_style('ChapterLabel', WD_STYLE_TYPE.PARAGRAPH)
            ch_label.font.name = self.config.font_heading
            ch_label.font.size = Pt(11)
            ch_label.font.bold = False
            ch_label.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
            ch_label.paragraph_format.space_before = Pt(72)  # ~1 inch from top
            ch_label.paragraph_format.space_after = Pt(6)
            ch_label.paragraph_format.first_line_indent = Pt(0)

        # Chapter heading - Number
        if 'ChapterNumber' not in [s.name for s in styles]:
            ch_num = styles.add_style('ChapterNumber', WD_STYLE_TYPE.PARAGRAPH)
            ch_num.font.name = self.config.font_heading
            ch_num.font.size = Pt(32)
            ch_num.font.bold = False
            ch_num.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
            ch_num.paragraph_format.space_before = Pt(6)
            ch_num.paragraph_format.space_after = Pt(36)
            ch_num.paragraph_format.first_line_indent = Pt(0)

        # Chapter title (if any)
        if 'ChapterTitle' not in [s.name for s in styles]:
            ch_title = styles.add_style('ChapterTitle', WD_STYLE_TYPE.PARAGRAPH)
            ch_title.font.name = self.config.font_heading
            ch_title.font.size = Pt(14)
            ch_title.font.italic = True
            ch_title.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
            ch_title.paragraph_format.space_after = Pt(36)
            ch_title.paragraph_format.first_line_indent = Pt(0)

        # First paragraph (no indent, for after headings)
        if 'BodyFirst' not in [s.name for s in styles]:
            first_style = styles.add_style('BodyFirst', WD_STYLE_TYPE.PARAGRAPH)
            first_style.font.name = self.config.font_body
            first_style.font.size = Pt(self.config.font_size_body)
            first_style.paragraph_format.line_spacing = self.config.line_spacing
            first_style.paragraph_format.first_line_indent = Pt(0)
            first_style.paragraph_format.space_after = Pt(0)
            first_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

        # Scene break
        if 'SceneBreak' not in [s.name for s in styles]:
            scene_style = styles.add_style('SceneBreak', WD_STYLE_TYPE.PARAGRAPH)
            scene_style.font.name = self.config.font_body
            scene_style.font.size = Pt(12)
            scene_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
            scene_style.paragraph_format.space_before = Pt(24)
            scene_style.paragraph_format.space_after = Pt(24)
            scene_style.paragraph_format.first_line_indent = Pt(0)

    def _add_title_page(self):
        """Add professional title page - minimalist design."""
        # Push title to ~35% from top for elegant positioning
        for _ in range(10):
            spacer = self.doc.add_paragraph()
            spacer.paragraph_format.space_after = Pt(0)
            spacer.paragraph_format.space_before = Pt(0)

        # Title - elegant, uppercase with letter-spacing
        title_para = self.doc.add_paragraph()
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_para.paragraph_format.space_after = Pt(18)
        title_para.paragraph_format.space_before = Pt(0)

        title_run = title_para.add_run(self.config.title.upper())
        title_run.font.name = self.config.font_heading
        title_run.font.size = Pt(26)
        title_run.font.bold = False
        self._add_letter_spacing(title_run, 4)

        # Subtitle (if any) - italic, smaller
        if self.config.subtitle:
            subtitle = self.doc.add_paragraph()
            subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
            subtitle.paragraph_format.space_after = Pt(24)
            sub_run = subtitle.add_run(self.config.subtitle)
            sub_run.font.name = self.config.font_heading
            sub_run.font.size = Pt(13)
            sub_run.font.italic = True
            sub_run.font.color.rgb = RGBColor(80, 80, 80)

        # Thin decorative line
        line_para = self.doc.add_paragraph()
        line_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        line_para.paragraph_format.space_before = Pt(12)
        line_para.paragraph_format.space_after = Pt(12)
        line_run = line_para.add_run("─────")
        line_run.font.size = Pt(10)
        line_run.font.color.rgb = RGBColor(180, 180, 180)

        # Author - clean, professional
        if self.config.author:
            author_para = self.doc.add_paragraph()
            author_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            author_para.paragraph_format.space_before = Pt(12)
            author_run = author_para.add_run(self.config.author)
            author_run.font.name = self.config.font_heading
            author_run.font.size = Pt(14)
            author_run.font.bold = False

        # Publisher at bottom (if any)
        if self.config.publisher:
            for _ in range(12):
                self.doc.add_paragraph()
            publisher = self.doc.add_paragraph()
            publisher.alignment = WD_ALIGN_PARAGRAPH.CENTER
            pub_run = publisher.add_run(self.config.publisher.upper())
            pub_run.font.name = self.config.font_heading
            pub_run.font.size = Pt(9)
            self._add_letter_spacing(pub_run, 2)

        # Page break after title page
        self.doc.add_page_break()

    def _add_copyright_page(self):
        """Add copyright page."""
        # Push content to bottom
        for _ in range(20):
            self.doc.add_paragraph()

        # Copyright text
        lines = [
            self.config.title,
            f"Copyright © {self.config.copyright_year or '2024'} {self.config.author}" if self.config.author else "",
            "",
            "All rights reserved.",
            "",
            "No part of this publication may be reproduced, distributed, or transmitted",
            "in any form or by any means without the prior written permission of the publisher.",
            "",
        ]

        if self.config.isbn:
            lines.append(f"ISBN: {self.config.isbn}")

        if self.config.publisher:
            lines.extend(["", f"Published by {self.config.publisher}"])

        for line in lines:
            p = self.doc.add_paragraph(line)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.first_line_indent = Pt(0)
            for run in p.runs:
                run.font.size = Pt(9)

        self.doc.add_page_break()

    def _extract_title_from_first_line(self, content: str) -> Tuple[Optional[str], Optional[str], str]:
        """
        Detect and extract title/author from first line if present.

        Pattern: "Title Author Content..." where title and author are known Vietnamese works.
        Example: "Chí Phèo Nam Cao He was walking while cursing."
                 → title="Chí Phèo", author="Nam Cao", content="He was walking..."

        Returns:
            Tuple of (title, author, remaining_content)
        """
        # Check Vietnamese title patterns
        for pattern, title, author in self.VIETNAMESE_TITLE_AUTHOR_PATTERNS:
            match = re.match(pattern, content, re.IGNORECASE | re.DOTALL)
            if match:
                remaining = match.group(3).strip() if len(match.groups()) >= 3 else content
                logger.info(f"Detected title from content: '{title}' by '{author}'")
                return title, author, remaining

        # Generic pattern: First line looks like "Title Author" (both capitalized, short)
        # E.g., "The Great Gatsby F. Scott Fitzgerald"
        first_line_match = re.match(r'^([A-ZÀ-Ỹ][^\n]{2,30}?)\s+([A-ZÀ-Ỹ][a-zA-ZÀ-ỹ\s\.]{3,25}?)\s+([A-Z])', content)
        if first_line_match:
            potential_title = first_line_match.group(1).strip()
            potential_author = first_line_match.group(2).strip()
            # Check if this looks like a real title/author combo (not sentence start)
            if len(potential_title.split()) <= 5 and len(potential_author.split()) <= 4:
                # Further validation: author should be name-like (2-4 words, capitalized)
                author_words = potential_author.split()
                if 2 <= len(author_words) <= 4 and all(w[0].isupper() for w in author_words if w):
                    remaining = content[first_line_match.end(2):].strip()
                    logger.info(f"Detected generic title: '{potential_title}' by '{potential_author}'")
                    return potential_title, potential_author, remaining

        return None, None, content

    def _add_book_content(self, content: str):
        """Parse and add book content with proper formatting (legacy compatibility)."""
        detected_title, detected_author, cleaned = self._extract_title_from_first_line(content)
        if detected_title and (not self.config.title or self.config.title == "Untitled"):
            self.config.title = detected_title
        if detected_author and not self.config.author:
            self.config.author = detected_author
        self._add_book_content_internal(cleaned)

    def _add_book_content_internal(self, content: str):
        """Parse and add book content with proper formatting (internal)."""
        lines = content.split('\n')
        current_para_lines = []
        is_first_para_in_chapter = True

        i = 0
        while i < len(lines):
            line = lines[i]
            line_stripped = line.strip()

            # Skip empty lines
            if not line_stripped:
                # Flush current paragraph
                if current_para_lines:
                    self._add_paragraph(current_para_lines, is_first_para_in_chapter)
                    current_para_lines = []
                    is_first_para_in_chapter = False
                i += 1
                continue

            # Check for chapter heading
            chapter_match = self._is_chapter_heading(line_stripped)
            if chapter_match:
                # Flush current paragraph
                if current_para_lines:
                    self._add_paragraph(current_para_lines, is_first_para_in_chapter)
                    current_para_lines = []

                # Add chapter heading
                self._add_chapter_heading(chapter_match)
                is_first_para_in_chapter = True
                i += 1
                continue

            # Check for scene break
            if self._is_scene_break(line_stripped):
                # Flush current paragraph
                if current_para_lines:
                    self._add_paragraph(current_para_lines, is_first_para_in_chapter)
                    current_para_lines = []

                # Add scene break
                self._add_scene_break()
                is_first_para_in_chapter = True
                i += 1
                continue

            # Regular text - accumulate into paragraph
            current_para_lines.append(line_stripped)
            i += 1

        # Flush remaining paragraph
        if current_para_lines:
            self._add_paragraph(current_para_lines, is_first_para_in_chapter)

    def _is_chapter_heading(self, line: str) -> Optional[Dict]:
        """Check if line is a chapter heading and extract info."""
        for pattern, lang in self.CHAPTER_PATTERNS:
            match = re.match(pattern, line, re.IGNORECASE)
            if match:
                if lang == 'md':
                    # Markdown header - treat as chapter
                    return {
                        'number': None,
                        'title': match.group(1).strip(),
                        'lang': lang
                    }
                else:
                    return {
                        'number': match.group(1),
                        'title': match.group(2).strip() if len(match.groups()) > 1 else '',
                        'lang': lang
                    }
        return None

    def _is_scene_break(self, line: str) -> bool:
        """Check if line is a scene break."""
        for pattern in self.SCENE_BREAK_PATTERNS:
            if re.match(pattern, line):
                return True
        return False

    def _add_chapter_heading(self, chapter_info: Dict):
        """Add chapter heading in commercial style."""
        self.chapter_count += 1

        # Page break before chapter (except first)
        if self.chapter_count > 1:
            self.doc.add_page_break()

        # Add some vertical space at top of chapter
        for _ in range(3):
            spacer = self.doc.add_paragraph()
            spacer.paragraph_format.space_after = Pt(0)

        # "CHAPTER" label with elegant letter-spacing
        if chapter_info['number']:
            chapter_label = self.doc.add_paragraph()
            chapter_label.alignment = WD_ALIGN_PARAGRAPH.CENTER
            chapter_label.paragraph_format.space_after = Pt(8)

            label_run = chapter_label.add_run('CHAPTER')
            label_run.font.name = self.config.font_heading
            label_run.font.size = Pt(10)
            label_run.font.color.rgb = RGBColor(100, 100, 100)
            self._add_letter_spacing(label_run, 5)

            # Chapter number - elegant, large
            number_para = self.doc.add_paragraph()
            number_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            number_para.paragraph_format.space_after = Pt(20)

            num_run = number_para.add_run(str(chapter_info['number']))
            num_run.font.name = self.config.font_heading
            num_run.font.size = Pt(28)
            num_run.font.bold = False

        # Chapter title (if any) - italic, refined
        if chapter_info['title']:
            title_para = self.doc.add_paragraph()
            title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            title_para.paragraph_format.space_after = Pt(36)

            title_run = title_para.add_run(chapter_info['title'])
            title_run.font.name = self.config.font_heading
            title_run.font.size = Pt(13)
            title_run.font.italic = True
            title_run.font.color.rgb = RGBColor(60, 60, 60)

    def _add_scene_break(self):
        """Add scene break - minimalist design."""
        # Get ornament or use subtle line
        ornament = self.SCENE_BREAK_ORNAMENTS.get(
            self.config.scene_break_style,
            self.SCENE_BREAK_ORNAMENTS['line']
        )

        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(18)
        p.paragraph_format.space_after = Pt(18)
        p.paragraph_format.first_line_indent = Pt(0)

        if ornament:
            run = p.add_run(ornament)
            run.font.size = Pt(10)
            run.font.color.rgb = RGBColor(150, 150, 150)
        else:
            # Just extra space for blank style
            p.paragraph_format.space_before = Pt(24)
            p.paragraph_format.space_after = Pt(24)

    def _add_paragraph(self, lines: List[str], is_first: bool = False):
        """Add paragraph with proper formatting."""
        text = ' '.join(lines)

        if not text.strip():
            return

        if is_first and self.config.drop_cap_enabled:
            # Add paragraph with drop cap effect
            self._add_drop_cap_paragraph(text)
        elif is_first:
            # No indent for first paragraph after heading
            p = self.doc.add_paragraph()
            p.paragraph_format.first_line_indent = Pt(0)
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0)
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY if self.config.justify_text else WD_ALIGN_PARAGRAPH.LEFT

            run = p.add_run(text)
            run.font.name = self.config.font_body
            run.font.size = Pt(self.config.font_size_body)
        else:
            # Normal paragraph with indent
            p = self.doc.add_paragraph()
            p.paragraph_format.first_line_indent = Inches(self.config.paragraph_indent)
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = self.config.line_spacing
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY if self.config.justify_text else WD_ALIGN_PARAGRAPH.LEFT

            # Widow/orphan control
            if self.config.widow_orphan_control:
                p.paragraph_format.widow_control = True

            run = p.add_run(text)
            run.font.name = self.config.font_body
            run.font.size = Pt(self.config.font_size_body)

    def _add_drop_cap_paragraph(self, text: str):
        """Add paragraph with elegant drop cap effect."""
        if not text:
            return

        # Extract first letter and first few words for small caps effect
        first_letter = text[0].upper()
        rest_of_text = text[1:] if len(text) > 1 else ''

        # Create paragraph
        p = self.doc.add_paragraph()
        p.paragraph_format.first_line_indent = Pt(0)
        p.paragraph_format.space_before = Pt(6)
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY if self.config.justify_text else WD_ALIGN_PARAGRAPH.LEFT

        # Add drop cap letter - elegant, not too large
        drop_run = p.add_run(first_letter)
        drop_run.font.name = self.config.font_heading
        drop_run.font.size = Pt(42)  # Refined size
        drop_run.font.bold = False
        drop_run.font.color.rgb = RGBColor(30, 30, 30)

        # Add rest of text
        if rest_of_text:
            # First few words in small caps for elegant transition
            words = rest_of_text.split(' ', 3)
            if len(words) >= 3:
                first_words = ' '.join(words[:3])
                remaining = ' '.join(words[3:])

                # Small caps for first words
                first_run = p.add_run(first_words.upper() + ' ')
                first_run.font.name = self.config.font_body
                first_run.font.size = Pt(self.config.font_size_body - 1)
                first_run.font.small_caps = True

                # Regular text
                if remaining:
                    text_run = p.add_run(remaining)
                    text_run.font.name = self.config.font_body
                    text_run.font.size = Pt(self.config.font_size_body)
            else:
                text_run = p.add_run(rest_of_text)
                text_run.font.name = self.config.font_body
                text_run.font.size = Pt(self.config.font_size_body)

    def _add_letter_spacing(self, run, spacing_pts: int):
        """Add letter spacing to a run (using character spacing)."""
        # This uses the w:spacing element for character spacing
        try:
            spacing_twips = spacing_pts * 20  # Convert points to twips
            r = run._r
            rPr = r.get_or_add_rPr()
            spacing_elem = parse_xml(
                f'<w:spacing {nsdecls("w")} w:val="{spacing_twips}"/>'
            )
            rPr.append(spacing_elem)
        except Exception:
            # If XML manipulation fails, skip spacing
            pass


def export_commercial_book(
    content: str,
    output_path: str,
    title: str = "Untitled",
    author: str = "",
    source_filename: Optional[str] = None,
    **kwargs
) -> str:
    """
    Convenience function to export content as commercial book.

    Args:
        content: Book content
        output_path: Output DOCX path
        title: Book title
        author: Author name
        source_filename: Original filename for smart title extraction
        **kwargs: Additional BookConfig options

    Returns:
        Path to generated file
    """
    config = BookConfig(
        title=title,
        author=author,
        **kwargs
    )

    exporter = CommercialBookExporter(config)
    return exporter.export(content, output_path, source_filename)


# CLI interface
if __name__ == "__main__":
    import argparse

    parser = argparse.ArgumentParser(description="Commercial Book Exporter")
    parser.add_argument("input", help="Input text file")
    parser.add_argument("-o", "--output", required=True, help="Output DOCX path")
    parser.add_argument("-t", "--title", default="Untitled", help="Book title")
    parser.add_argument("-a", "--author", default="", help="Author name")
    parser.add_argument("--no-title-page", action="store_true", help="Skip title page")
    parser.add_argument("--no-drop-cap", action="store_true", help="Disable drop caps")

    args = parser.parse_args()

    # Read input
    content = Path(args.input).read_text(encoding='utf-8')

    # Export
    config = BookConfig(
        title=args.title,
        author=args.author,
        include_title_page=not args.no_title_page,
        drop_cap_enabled=not args.no_drop_cap
    )

    exporter = CommercialBookExporter(config)
    output = exporter.export(content, args.output, args.input)

    print(f"Exported: {output}")
