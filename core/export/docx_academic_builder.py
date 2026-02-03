"""
Phase 2.0.2 - Academic DOCX Builder
====================================

DEPRECATED: This module is maintained for backward compatibility.
New code should use `core.export.docx_academic` instead.

Full implementation for Phase 2.0.2:
- Create DOCX with semantic structure
- Proper styling (headings, bold/italic for theorems/proofs)
- Centered equations (LaTeX text, no OMML conversion yet)
- Global formatting (fonts, spacing)

Phase 2.0.3b: OMML equation rendering
Phase 2.0.5: Professional aesthetic improvements
"""

from dataclasses import dataclass
from typing import List, Optional, Literal, Dict, Any
from pathlib import Path
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from core.structure.semantic_model import DocNode, DocNodeType, DocNodeList
from core.rendering import omml_converter
from core.export.docx_styles import StyleManager, THEME_ACADEMIC
from core.export.docx_page_layout import _setup_page_layout
from core.export.docx_front_matter import FrontMatterGenerator
from core.export.config import AcademicLayoutConfig
from config.logging_config import get_logger

logger = get_logger(__name__)




def build_academic_docx(
    nodes: DocNodeList,
    output_path: str,
    config: Optional[AcademicLayoutConfig] = None,
    metadata: Optional[Dict[str, str]] = None,
) -> str:
    """
    Build academic DOCX from semantic node list.

    Phase 2.0.2 implementation:
    - CHAPTER/SECTION/SUBSECTION → heading styles (Heading 1/2/3)
    - THEOREM/LEMMA/etc → bold label + normal text
    - PROOF → italic header + normal text
    - EQUATION_BLOCK → centered paragraph with LaTeX text (no OMML rendering)
    - PARAGRAPH → normal paragraph
    - REFERENCES_SECTION → heading
    - REFERENCE_ENTRY → hanging indent paragraph

    Args:
        nodes: List of DocNode from semantic extractor
        output_path: Output DOCX file path
        config: Optional layout configuration (uses defaults if None)

    Returns:
        str: Absolute path to created DOCX file (Phase 2.0.8)
    """
    # Check equation nodes for LaTeX source
    logger.debug("build_academic_docx() START")
    logger.debug(f"Total nodes: {len(nodes)}")
    equation_nodes = [n for n in nodes if n.is_equation()]
    logger.debug(f"Equation nodes: {len(equation_nodes)}")

    # Check if any equations have latex_source metadata
    enriched_count = sum(1 for n in equation_nodes if n.metadata.get('latex_source'))
    logger.debug(f"Equations with latex_source: {enriched_count}/{len(equation_nodes)}")

    if equation_nodes and enriched_count == 0:
        logger.warning("No equations have latex_source metadata!")
        logger.warning("arXiv enrichment may not have been called!")

    if config is None:
        config = AcademicLayoutConfig()

    doc = Document()

    # Initialize StyleManager
    style_manager = StyleManager(theme_name=config.theme)
    logger.info(f"Initialized StyleManager with theme: {config.theme}")

    # Phase 9.2: Setup Page Architecture (Margins, Headers, Footers)
    _setup_page_layout(doc, config, style_manager)

    # Phase 9.4: Front Matter
    fm_generator = FrontMatterGenerator(doc, style_manager)

    # Phase 2026-01: Add cover page first (if provided in metadata)
    cover_image = metadata.get('cover_image') if metadata else None
    if cover_image:
        fm_generator.generate_cover_page(cover_image)

    if metadata:
        fm_generator.generate_title_page(metadata)

        # Optionally generate TOC (configurable)
        # For now, always include if metadata provided for premium feel
        fm_generator.generate_toc()

    # Process each nodes

    for node in nodes:
        if node.is_heading():
            # Add heading with appropriate level
            level = node.level or 1
            para = doc.add_heading(node.text, level=level)
            style_manager.apply_heading_style(para, level)

        elif node.is_theorem_like():
            # Add theorem block
            _format_theorem_block(doc, node, config, style_manager)

        elif node.is_proof():
            # Add proof block
            _format_proof_block(doc, node, config, style_manager)


        elif node.is_equation():
            # Add equation block (centered, LaTeX or OMML)
            _format_equation_block(doc, node, config, style_manager)

        elif node.node_type == DocNodeType.REFERENCES_SECTION:
            # Add references heading
            doc.add_heading(node.text or "References", level=1)

        elif node.node_type == DocNodeType.REFERENCE_ENTRY:
            # Add reference entry with hanging indent
            para = doc.add_paragraph(node.text)
            para.paragraph_format.left_indent = Inches(0.5)
            para.paragraph_format.first_line_indent = Inches(-0.5)

        elif node.node_type.value == 'table': # DocNodeType.TABLE
            # Phase 8: Render Smart Table
            html_content = node.metadata.get('html_content')
            if html_content:
                logger.info("  Rendering Smart Table...")
                try:
                    _render_html_table_naive(doc, html_content)
                except Exception as e:
                    logger.error(f"Failed to render table: {e}")
                    doc.add_paragraph(f"[Table: {node.text}]")
            else:
                 doc.add_paragraph(f"[Table: {node.text}] - Content Missing")

        else:  # PARAGRAPH or UNKNOWN
            # Add normal paragraph
            para = doc.add_paragraph(node.text)
            style_manager.apply_body_style(para)

    # Save
    doc.save(output_path)

    # Phase 2.0.8: Return absolute path for verification
    import os
    return os.path.abspath(output_path)


def _render_html_table_naive(doc: Document, html: str):
    """
    Naive HTML Table renderer for python-docx.
    Parses simple <table><tr><td> structure.
    """
    from html.parser import HTMLParser

    class TableParser(HTMLParser):
        def __init__(self):
            super().__init__()
            self.rows = []
            self.current_row = []
            self.in_td = False
            self.current_data = []

        def handle_starttag(self, tag, attrs):
            if tag == 'tr':
                self.current_row = []
            elif tag == 'td' or tag == 'th':
                self.in_td = True
                self.current_data = []

        def handle_endtag(self, tag):
            if tag == 'tr':
                self.rows.append(self.current_row)
            elif tag == 'td' or tag == 'th':
                self.in_td = False
                self.current_row.append("".join(self.current_data).strip())

        def handle_data(self, data):
            if self.in_td:
                self.current_data.append(data)

    parser = TableParser()
    parser.feed(html)
    
    if not parser.rows:
        return

    # Create table in docx
    table = doc.add_table(rows=len(parser.rows), cols=len(parser.rows[0]))
    table.style = 'Table Grid'
    
    for r_idx, row_data in enumerate(parser.rows):
        row = table.rows[r_idx]
        for c_idx, cell_text in enumerate(row_data):
            if c_idx < len(row.cells):
                row.cells[c_idx].text = cell_text








def _format_theorem_block(doc: Document, node: DocNode, config: AcademicLayoutConfig, style_manager: StyleManager):
    """Format theorem with dynamic styling."""
    para = doc.add_paragraph()

    # Add bold label if present
    if node.title:
        label_run = para.add_run(f"{node.title}. ")
        label_run.bold = True

    # Strip label from text overlap
    text = node.text
    if node.title and text.startswith(node.title):
        text = text[len(node.title):].lstrip('. ')

    para.add_run(text)

    # Apply premium box styling
    if config.enable_theorem_boxes:
        try:
            # Detect box type
            box_type = 'theorem'
            if node.node_type == DocNodeType.LEMMA: box_type = 'lemma'
            elif node.node_type == DocNodeType.COROLLARY: box_type = 'corollary'
            elif node.node_type == DocNodeType.DEFINITION: box_type = 'definition'
            
            style_manager.apply_theorem_box(para, box_type)
        except Exception as e:
            logger.warning(f"Failed to apply theorem styling: {e}")
    else:
        # Fallback to simple body styling
        style_manager.apply_body_style(para)


def _format_proof_block(doc: Document, node: DocNode, config: AcademicLayoutConfig, style_manager: StyleManager):
    """Format proof with dynamic styling."""
    para = doc.add_paragraph()
    
    # Header resolution logic...
    proof_header = node.title if node.title else ("Chứng minh" if "Chứng minh" in node.text[:20] else "Proof")

    header_run = para.add_run(f"{proof_header}. ")
    header_run.italic = True

    # Strip text logic...
    text = node.text
    if text.startswith(proof_header):
         text = text[len(proof_header):].lstrip('. :')
    
    # QED logic
    has_qed = any(marker in text for marker in ['□', '■', '∎', 'Q.E.D', 'Đpcm'])
    if config.enable_proof_indent and not has_qed:
        text = text.rstrip() + "  □"

    para.add_run(text)

    # Styling
    style_manager.apply_body_style(para)
    para.paragraph_format.left_indent = Pt(20) # Slight indent for proofs



def _format_equation_block(doc: Document, node: DocNode, config: AcademicLayoutConfig, style_manager: Optional[StyleManager] = None):
    """
    Format an equation block.

    Phase 2.0.2: Center-aligned LaTeX text
    Phase 2.0.3b: Convert LaTeX → OMML if config.equation_rendering_mode == "omml"
    Phase 2.0.5: Enhanced spacing and layout

    Args:
        doc: python-docx Document object
        node: DocNode with EQUATION_BLOCK type
        config: Layout configuration with equation_rendering_mode

    Rendering Logic:
        - If equation_rendering_mode == "omml" AND pandoc available:
            * Try to convert LaTeX to OMML
            * If conversion succeeds: inject OMML into paragraph
            * If conversion fails: fallback to LaTeX text
        - If equation_rendering_mode == "latex_text" OR pandoc unavailable:
            * Render as centered LaTeX text (backward compatible)

    Safety:
        - Never crashes, never loses equations
        - Math correctness > layout > formatting
        - Fallback always preserves original LaTeX text
    """
    # Create centered paragraph
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Phase 2.0.5: Apply enhanced equation spacing if enabled
    if style_manager:
        style_manager.apply_body_style(para)
    
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Check if OMML rendering is requested
    if config.equation_rendering_mode == "omml":
        # Check if pandoc is available
        if not omml_converter.is_pandoc_available():
            logger.warning(
                "OMML rendering requested but pandoc not available - "
                "falling back to LaTeX text. Install pandoc for OMML support."
            )
            # Fallback: render as LaTeX text
            para.add_run(node.text)
            return

        # Try to convert LaTeX to OMML
        try:
            # Phase 2.2.0: LaTeX equation splitter for compound blocks
            # If latex_source exists (Phase 2.1.0) but latex_equation_primary doesn't exist,
            # try splitting the compound block
            latex_source_raw = node.metadata.get('latex_source')
            latex_equation_primary = node.metadata.get('latex_equation_primary')

            if latex_source_raw and not latex_equation_primary:
                # Phase 2.2.0: Try to split compound LaTeX block
                try:
                    from core.latex.eq_splitter import split_latex_equations

                    split_result = split_latex_equations(latex_source_raw)

                    if split_result.is_confident and split_result.equation_segments:
                        logger.debug(f"Phase 2.2.0: Successfully split LaTeX into {len(split_result.equation_segments)} equation(s)")

                        # Try to convert each segment
                        # For Phase 2.2.0, we use the first segment (most important equation)
                        # Future: could render multiple equations, but keep it simple for now
                        for eq in split_result.equation_segments[:1]:  # Only first equation for Phase 2.2.0
                            omml_xml = omml_converter.latex_to_omml(eq)

                            if omml_xml:
                                success = omml_converter.inject_omml_as_display(para, omml_xml)

                                if success:
                                    logger.debug(f"Phase 2.2.0: OMML rendering successful for split equation")
                                    return
                                else:
                                    logger.debug(f"Phase 2.2.0: OMML injection failed, trying fallback")
                            else:
                                logger.debug(f"Phase 2.2.0: OMML conversion failed for split equation")

                        # If we reach here, splitting succeeded but OMML conversion failed
                        # Fall through to try the regular path below
                        logger.debug(f"Phase 2.2.0: Split succeeded but OMML failed, using fallback")
                    else:
                        # Not confident or no segments, use regular path
                        logger.debug(f"Phase 2.2.0: Not confident about splitting (reason: {split_result.reason}), using fallback")

                except Exception as e:
                    logger.debug(f"Phase 2.2.0: Splitter error: {e}, using fallback")
                    # Continue to regular path

            # Regular path: Phase 2.1.2 / 2.1.0 / fallback
            # This gives us clean, isolated equations from compound LaTeX blocks
            latex_content = (
                node.metadata.get('latex_equation_primary') or  # Phase 2.1.2: Clean single equation
                node.metadata.get('latex_source') or             # Phase 2.1.0: Fallback to compound block
                node.text or                                     # Last resort: PDF text
                ""
            )

            if node.metadata.get('latex_equation_primary'):
                logger.debug(f"✅ Phase 2.1.2: Using primary equation (length: {len(latex_content)})")
            elif node.metadata.get('latex_source'):
                logger.debug(f"⚠️  Phase 2.1.0: Using compound LaTeX source (length: {len(latex_content)})")

            # Convert to OMML
            omml_xml = omml_converter.latex_to_omml(latex_content)

            if omml_xml:
                # Phase 2.1.1: Use DISPLAY MODE (fixes equations split into multiple lines)
                # Successfully converted - inject OMML as display mode equation
                success = omml_converter.inject_omml_as_display(para, omml_xml)

                if success:
                    logger.debug(f"OMML display mode rendering successful for equation: {latex_content[:50]}...")
                    return
                else:
                    logger.warning(f"OMML display injection failed for: {latex_content[:50]}... - falling back to LaTeX text")
            else:
                logger.debug(f"OMML conversion returned None for: {latex_content[:50]}... - falling back to LaTeX text")

        except Exception as e:
            logger.error(f"Unexpected error in OMML rendering: {e} - falling back to LaTeX text")

        # Fallback: render as LaTeX text
        para.add_run(node.text)

    else:
        # Default: render as LaTeX text (backward compatible)
        para.add_run(node.text)
