"""
Phase 3.2 - Book DOCX Builder with Typography Engine

Professional book layout for commercial translation.
Handles chapters, blockquotes, epigraphs, scene breaks, dialogue, and more.

Phase 3.2: Enhanced typography system bringing output to 85-90% commercial quality.

Parallel to docx_academic_builder.py - designed for general/commercial books
rather than STEM documents.
"""

import os
from dataclasses import dataclass
from typing import Optional, List
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

from core.structure.semantic_model import DocNode, DocNodeType, DocNodeList

import logging

# Initialize logger
logger = logging.getLogger(__name__)


@dataclass
class BookLayoutConfig:
    """
    Configuration for professional book typography.

    Phase 3.2: Enhanced typography system with professional controls.
    Brings output to 85-90% of commercial ebook quality.

    Font Hierarchy:
        title_font: Font for title page and major titles
        heading_font: Font for chapter and section headings
        body_font: Font for body text, blockquotes, etc.

    Font Sizes (in points):
        title_size: Title page (20pt)
        h1_size: Chapter headings (18pt)
        h2_size: Section headings (14pt)
        h3_size: Subsection headings (12pt)
        body_font_size: Body text (11pt)
        epigraph_size: Epigraphs (10pt - slightly smaller)

    Line Spacing:
        line_spacing: Line spacing multiplier for body (1.3 = comfortable)
        heading_line_spacing: Line spacing for headings (1.15 = tighter)

    Paragraph Spacing (in points):
        para_spacing_before: Space before each body paragraph
        para_spacing_after: Space after each body paragraph
        heading_spacing_before: Space before headings
        heading_spacing_after: Space after headings

    Indentation (in twips, 1440 twips = 1 inch):
        first_line_indent: First-line indent for body paragraphs (360 = 0.25 inch)
        dialogue_indent: Left indent for dialogue (360 = 0.25 inch)
        blockquote_left_indent: Left indent for blockquotes (720 = 0.5 inch)
        blockquote_right_indent: Right indent for blockquotes (720 = 0.5 inch)
        epigraph_left_indent: Left indent for epigraphs (1440 = 1 inch)

    Scene Breaks:
        scene_break_spacing_before: Space before scene break (18pt)
        scene_break_spacing_after: Space after scene break (18pt)

    Page Breaks:
        page_break_before_chapter: Start each chapter on new page (True)
    """
    # Font families - Times New Roman for Vietnamese compatibility
    title_font: str = "Times New Roman"
    heading_font: str = "Times New Roman"
    body_font: str = "Times New Roman"

    # Font sizes (in points)
    title_size: int = 20
    h1_size: int = 18  # Chapter headings
    h2_size: int = 14  # Section headings
    h3_size: int = 12  # Subsection headings
    body_font_size: int = 11
    epigraph_size: int = 10

    # Line spacing
    line_spacing: float = 1.3  # Body text
    heading_line_spacing: float = 1.15  # Headings (tighter)

    # Paragraph spacing (in points)
    para_spacing_before: int = 0
    para_spacing_after: int = 6
    heading_spacing_before: int = 18
    heading_spacing_after: int = 12

    # Indentation (in twips, 1440 twips = 1 inch)
    first_line_indent: int = 360  # 0.25 inch for body paragraphs
    dialogue_indent: int = 360  # 0.25 inch for dialogue
    blockquote_left_indent: int = 720  # 0.5 inch
    blockquote_right_indent: int = 720  # 0.5 inch
    epigraph_left_indent: int = 1440  # 1 inch

    # Scene breaks
    scene_break_spacing_before: int = 18  # Points
    scene_break_spacing_after: int = 18  # Points

    # Page breaks
    page_break_before_chapter: bool = True


def build_book_docx(
    nodes: DocNodeList,
    output_path: str,
    config: Optional[BookLayoutConfig] = None
) -> str:
    """
    Build a professionally formatted book DOCX from semantic nodes.

    Phase 3.1: Simple but professional layout with proper typography.
    Handles chapters, sections, blockquotes, epigraphs, scene breaks.

    Args:
        nodes: List of DocNode objects representing document structure
        output_path: Path where DOCX should be saved
        config: Optional BookLayoutConfig (uses defaults if None)

    Returns:
        Absolute path to the created DOCX file

    Raises:
        ValueError: If nodes is empty or output_path is invalid
        IOError: If DOCX cannot be written to output_path

    Example:
        >>> from core.structure.semantic_model import DocNode, DocNodeType
        >>> nodes = [
        ...     DocNode(DocNodeType.CHAPTER, "The Beginning", title="Chapter 1"),
        ...     DocNode(DocNodeType.PARAGRAPH, "It was a dark and stormy night..."),
        ...     DocNode(DocNodeType.SCENE_BREAK, "* * *"),
        ... ]
        >>> path = build_book_docx(nodes, "output.docx")
        >>> assert os.path.exists(path)
    """
    if not nodes:
        raise ValueError("Cannot build DOCX from empty node list")

    if not output_path:
        raise ValueError("output_path cannot be empty")

    # Use default config if none provided
    if config is None:
        config = BookLayoutConfig()

    # Create new document
    doc = Document()

    # Process each node
    for node in nodes:
        _add_node_to_document(doc, node, config)

    # Save document
    doc.save(output_path)

    # Return absolute path
    return os.path.abspath(output_path)


def _add_node_to_document(doc: Document, node: DocNode, config: BookLayoutConfig) -> None:
    """
    Add a single DocNode to the document with appropriate formatting.

    Args:
        doc: python-docx Document object
        node: Semantic node to add
        config: Layout configuration
    """
    if node.is_heading():
        _add_heading(doc, node, config)
    elif node.is_blockquote():
        _add_blockquote(doc, node, config)
    elif node.is_epigraph():
        _add_epigraph(doc, node, config)
    elif node.is_scene_break():
        _add_scene_break(doc, node, config)
    elif node.node_type == DocNodeType.PARAGRAPH:
        _add_paragraph(doc, node, config)
    elif node.node_type == DocNodeType.DIALOGUE:
        # Phase 3.2: Dedicated dialogue formatting
        _add_dialogue(doc, node, config)
    elif node.node_type == DocNodeType.FRONT_MATTER:
        # Phase 3.2: Dedicated front matter formatting
        _add_front_matter(doc, node, config)
    elif node.node_type == DocNodeType.BACK_MATTER:
        # Phase 3.2: Dedicated back matter formatting
        _add_back_matter(doc, node, config)
    else:
        # Fallback: treat as paragraph
        _add_paragraph(doc, node, config)


def _add_heading(doc: Document, node: DocNode, config: BookLayoutConfig) -> None:
    """
    Add a heading (chapter/section/subsection) with professional formatting.

    Phase 3.2: Enhanced with page breaks, spacing, and typography.

    Args:
        doc: python-docx Document object
        node: Heading node (CHAPTER, SECTION, or SUBSECTION)
        config: Layout configuration
    """
    # Determine heading level, size, and whether to add page break
    is_chapter = node.node_type == DocNodeType.CHAPTER

    if is_chapter:
        level = 0  # Heading 1
        font_size = config.h1_size
        font_name = config.heading_font
        # Page break before chapter (if enabled)
        if config.page_break_before_chapter and len(doc.paragraphs) > 0:
            doc.add_page_break()
    elif node.node_type == DocNodeType.SECTION:
        level = 1  # Heading 2
        font_size = config.h2_size
        font_name = config.heading_font
    elif node.node_type == DocNodeType.SUBSECTION:
        level = 2  # Heading 3
        font_size = config.h3_size
        font_name = config.heading_font
    else:
        level = 0
        font_size = config.h1_size
        font_name = config.heading_font

    # Construct heading text
    heading_text = node.text
    if node.title:
        # If there's a title (e.g., "Chapter 1"), prepend it
        heading_text = f"{node.title}: {node.text}" if node.text else node.title

    # Add heading
    heading = doc.add_heading(heading_text, level=level)

    # Apply font formatting
    for run in heading.runs:
        run.font.name = font_name
        run.font.size = Pt(font_size)
        if is_chapter:
            run.font.bold = True  # Bold for chapter headings

    # Apply spacing
    heading.paragraph_format.space_before = Pt(config.heading_spacing_before)
    heading.paragraph_format.space_after = Pt(config.heading_spacing_after)
    heading.paragraph_format.line_spacing = config.heading_line_spacing

    # Center chapter headings for professional appearance
    if is_chapter:
        heading.alignment = WD_ALIGN_PARAGRAPH.CENTER


def _add_paragraph(doc: Document, node: DocNode, config: BookLayoutConfig) -> None:
    """
    Add a normal body paragraph with professional formatting.

    Phase 3.2: Enhanced with justification, first-line indent, and spacing.

    Args:
        doc: python-docx Document object
        node: Paragraph node
        config: Layout configuration
    """
    para = doc.add_paragraph(node.text)

    # Apply font formatting
    for run in para.runs:
        run.font.name = config.body_font
        run.font.size = Pt(config.body_font_size)

    # Apply line spacing
    para.paragraph_format.line_spacing = config.line_spacing

    # Apply paragraph spacing
    para.paragraph_format.space_before = Pt(config.para_spacing_before)
    para.paragraph_format.space_after = Pt(config.para_spacing_after)

    # Apply first-line indent for professional book appearance
    para.paragraph_format.first_line_indent = Inches(config.first_line_indent / 1440)

    # Justify text (standard for books)
    para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY


def _add_blockquote(doc: Document, node: DocNode, config: BookLayoutConfig) -> None:
    """
    Add a blockquote with professional formatting.

    Phase 3.2: Enhanced with bilateral indent, spacing, and italic formatting.

    Args:
        doc: python-docx Document object
        node: Blockquote node
        config: Layout configuration
    """
    para = doc.add_paragraph(node.text)

    # Apply font formatting
    for run in para.runs:
        run.font.name = config.body_font
        run.font.size = Pt(config.body_font_size)
        run.font.italic = True  # Italic for blockquotes

    # Apply bilateral indentation (both left and right)
    para.paragraph_format.left_indent = Inches(config.blockquote_left_indent / 1440)
    para.paragraph_format.right_indent = Inches(config.blockquote_right_indent / 1440)

    # Apply line spacing
    para.paragraph_format.line_spacing = config.line_spacing

    # Apply paragraph spacing
    para.paragraph_format.space_before = Pt(config.para_spacing_before)
    para.paragraph_format.space_after = Pt(config.para_spacing_after)


def _add_epigraph(doc: Document, node: DocNode, config: BookLayoutConfig) -> None:
    """
    Add an epigraph with professional positioning.

    Phase 3.2: Enhanced with refined typography and spacing.

    Args:
        doc: python-docx Document object
        node: Epigraph node
        config: Layout configuration
    """
    para = doc.add_paragraph(node.text)

    # Apply font formatting (smaller + italic)
    for run in para.runs:
        run.font.name = config.body_font
        run.font.size = Pt(config.epigraph_size)
        run.font.italic = True

    # Apply large left indentation
    para.paragraph_format.left_indent = Inches(config.epigraph_left_indent / 1440)

    # Right-align for professional epigraph appearance
    para.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    # Apply line spacing
    para.paragraph_format.line_spacing = config.line_spacing

    # Add spacing
    para.paragraph_format.space_before = Pt(config.para_spacing_before)
    para.paragraph_format.space_after = Pt(config.para_spacing_after)


def _add_scene_break(doc: Document, node: DocNode, config: BookLayoutConfig) -> None:
    """
    Add a scene break with config-driven spacing.

    Phase 3.2: Enhanced with professional spacing controls.

    Args:
        doc: python-docx Document object
        node: Scene break node
        config: Layout configuration
    """
    separator_text = node.text if node.text else "* * *"
    para = doc.add_paragraph(separator_text)

    # Center alignment
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Apply font formatting
    for run in para.runs:
        run.font.name = config.body_font
        run.font.size = Pt(config.body_font_size)

    # Apply config-driven spacing
    para.paragraph_format.space_before = Pt(config.scene_break_spacing_before)
    para.paragraph_format.space_after = Pt(config.scene_break_spacing_after)


def _add_dialogue(doc: Document, node: DocNode, config: BookLayoutConfig) -> None:
    """
    Add a dialogue paragraph with optional special formatting.

    Phase 3.2: Enhanced with dedicated dialogue indent.

    Args:
        doc: python-docx Document object
        node: Dialogue node
        config: Layout configuration
    """
    para = doc.add_paragraph(node.text)

    # Apply font formatting
    for run in para.runs:
        run.font.name = config.body_font
        run.font.size = Pt(config.body_font_size)

    # Apply line spacing
    para.paragraph_format.line_spacing = config.line_spacing

    # Apply paragraph spacing
    para.paragraph_format.space_before = Pt(config.para_spacing_before)
    para.paragraph_format.space_after = Pt(config.para_spacing_after)

    # Apply dialogue-specific indent (optional - can be same as body or different)
    para.paragraph_format.left_indent = Inches(config.dialogue_indent / 1440)

    # Justify text (standard for books)
    para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY


def _add_front_matter(doc: Document, node: DocNode, config: BookLayoutConfig) -> None:
    """
    Add front matter content (title page, copyright, dedication, TOC).

    Phase 3.2: Simple formatting suitable for front matter.

    Args:
        doc: python-docx Document object
        node: Front matter node
        config: Layout configuration
    """
    para = doc.add_paragraph(node.text)

    # Apply font formatting
    for run in para.runs:
        run.font.name = config.body_font
        run.font.size = Pt(config.body_font_size)

    # Center-align for typical front matter (title pages, dedications)
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Apply spacing
    para.paragraph_format.space_before = Pt(config.para_spacing_before)
    para.paragraph_format.space_after = Pt(config.para_spacing_after)
    para.paragraph_format.line_spacing = config.line_spacing


def _add_back_matter(doc: Document, node: DocNode, config: BookLayoutConfig) -> None:
    """
    Add back matter content (appendices, acknowledgments, author bio).

    Phase 3.2: Standard paragraph formatting suitable for back matter.

    Args:
        doc: python-docx Document object
        node: Back matter node
        config: Layout configuration
    """
    para = doc.add_paragraph(node.text)

    # Apply font formatting
    for run in para.runs:
        run.font.name = config.body_font
        run.font.size = Pt(config.body_font_size)

    # Apply line spacing
    para.paragraph_format.line_spacing = config.line_spacing

    # Apply paragraph spacing
    para.paragraph_format.space_before = Pt(config.para_spacing_before)
    para.paragraph_format.space_after = Pt(config.para_spacing_after)

    # Left-align for typical back matter
    para.alignment = WD_ALIGN_PARAGRAPH.LEFT
