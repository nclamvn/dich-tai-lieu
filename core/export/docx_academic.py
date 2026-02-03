#!/usr/bin/env python3
# -*- coding: utf-8 -*-

"""
Academic DOCX Exporter - Phase 1.4 Consolidation
=================================================

Academic document exporter using semantic node structure.
Extends base class for common functionality.

Features:
- Semantic node processing (DocNodeList)
- Theorem/Lemma/Proof formatting
- OMML equation rendering
- Professional styling
"""

from dataclasses import dataclass
from typing import Optional, Dict, Any, Literal

from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

from .docx_base import DocxExporterBase, ExportConfig
from .docx_styles import StyleManager, THEME_ACADEMIC
from .docx_page_layout import _setup_page_layout
from .docx_front_matter import FrontMatterGenerator
from core.structure.semantic_model import DocNode, DocNodeType, DocNodeList
from core.rendering import omml_converter
from config.logging_config import get_logger

logger = get_logger(__name__)


@dataclass
class AcademicExportConfig(ExportConfig):
    """
    Configuration for academic DOCX export.

    Extends base config with academic-specific options.
    """
    # Equation rendering
    equation_rendering_mode: Literal["latex_text", "omml"] = "latex_text"

    # Professional styling (Phase 2.0.5)
    theme: str = "academic"
    enable_theorem_boxes: bool = True
    enable_proof_indent: bool = True
    enable_advanced_equation_layout: bool = True

    # Override defaults for academic style
    font_name: Optional[str] = None # Defer to theme
    font_size: Optional[int] = None # Defer to theme


class AcademicDocxExporter(DocxExporterBase):
    """
    Academic DOCX exporter for semantic documents.

    Processes DocNodeList with proper formatting for:
    - Headings (Chapter/Section/Subsection)
    - Theorems, Lemmas, Corollaries, etc.
    - Proofs with QED symbols
    - Equations (LaTeX text or OMML)
    - References with hanging indent
    """

    def __init__(self, config: Optional[AcademicExportConfig] = None):
        """
        Initialize academic exporter.

        Args:
            config: Academic export configuration
        """
        super().__init__(config or AcademicExportConfig())
        self.academic_config = self.config  # Type hint
        
        # Initialize StyleManager
        self.style_manager = StyleManager(theme_name=self.academic_config.theme)
    
    def _setup_page_layout(self) -> None:
        """Override page layout to use Page Architecture module."""
        # We need a temporary config object that matches what _setup_page_layout expects
        # Or we can pass self.academic_config if it has the fields
        # Ideally, _setup_page_layout should accept a generic config with theme/margins
        if self.doc:
            # We can reuse the helper logic directly
            from .config import AcademicLayoutConfig
            # Adapt config
            layout_config = AcademicLayoutConfig(theme=self.academic_config.theme)
            _setup_page_layout(self.doc, layout_config, self.style_manager)
            
    def _setup_base_styles(self) -> None:
        """Override base styles to use StyleManager."""
        pass # StyleManager applies styles dynamically per element, or we can pre-set doc defaults here

    def _add_front_matter(self, metadata: Dict[str, str]) -> None:
        """Add Front Matter.

        Args:
            metadata: Document metadata (title, author, etc.)
        """
        fm_generator = FrontMatterGenerator(self.doc, self.style_manager)

        # Phase 2026-01: Add cover page first (if provided in metadata)
        cover_image = metadata.get('cover_image') if metadata else None
        if cover_image:
            fm_generator.generate_cover_page(cover_image)

        # Add title page and TOC
        if metadata:
            fm_generator.generate_title_page(metadata)
            fm_generator.generate_toc()


    def _add_content(self, content: DocNodeList) -> None:
        """
        Add semantic nodes to document.

        Args:
            content: DocNodeList from semantic extractor
        """
        logger.debug(f"AcademicDocxExporter: Processing {len(content)} nodes")

        for node in content:
            if node.is_heading():
                self._add_heading_node(node)
            elif node.is_theorem_like():
                self._add_theorem_node(node)
            elif node.is_proof():
                self._add_proof_node(node)
            elif node.is_equation():
                self._add_equation_node(node)
            elif node.node_type == DocNodeType.REFERENCES_SECTION:
                self._add_references_heading(node)
            elif node.node_type == DocNodeType.REFERENCE_ENTRY:
                self._add_reference_entry(node)
            else:
                # Default: add as paragraph
                para = self.doc.add_paragraph(node.text)
                self.style_manager.apply_body_style(para)

    def _add_heading_node(self, node: DocNode) -> None:
        """Add heading node."""
        level = node.level or 1
        para = self.doc.add_heading(node.text, level=level)
        self.style_manager.apply_heading_style(para, level)

    def _add_theorem_node(self, node: DocNode) -> None:
        """Add theorem-like block with styling."""
        para = self.doc.add_paragraph()

        # Add bold label
        if node.title:
            label_run = para.add_run(f"{node.title}. ")
            label_run.bold = True

        # Add content (strip duplicate label if present)
        text = node.text
        if node.title and text.startswith(node.title):
            text = text[len(node.title):].lstrip('. ')

        para.add_run(text)


        # Apply theorem box styling if enabled
        if self.academic_config.enable_theorem_boxes:
            box_type = self._get_box_type(node.node_type)
            try:
                self.style_manager.apply_theorem_box(para, box_type)
            except Exception as e:
                logger.warning(f"Failed to apply theorem box: {e}")
        else:
             self.style_manager.apply_body_style(para)

    def _add_proof_node(self, node: DocNode) -> None:
        """Add proof block with styling."""
        para = self.doc.add_paragraph()

        # Determine header
        if node.title:
            header = node.title
        elif "Chứng minh" in node.text[:20]:
            header = "Chứng minh"
        else:
            header = "Proof"

        # Add italic header
        header_run = para.add_run(f"{header}. ")
        header_run.italic = True

        # Add content (strip duplicate header)
        text = node.text
        if text.startswith(header):
            text = text[len(header):].lstrip('. :')
        elif text.startswith("Proof.") or text.startswith("Proof:"):
            text = text[6:].lstrip('. :')

        # Add QED if enabled and not present
        has_qed = any(m in text for m in ['□', '■', '∎', 'Q.E.D', 'Đpcm'])
        if self.academic_config.enable_proof_indent and not has_qed:
            text = text.rstrip() + "  □"

        para.add_run(text)

        # Apply proof styling if enabled
        # Use simple indentation via paragraph format
        self.style_manager.apply_body_style(para)
        if self.academic_config.enable_proof_indent:
            para.paragraph_format.left_indent = Pt(20)

    def _add_equation_node(self, node: DocNode) -> None:
        """Add equation block (LaTeX text or OMML)."""
        para = self.doc.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Apply equation styling if enabled
        # Center align is handled above
        self.style_manager.apply_body_style(para)
        # Assuming body style might reset alignment, so re-center
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Try OMML rendering if enabled
        if self.academic_config.equation_rendering_mode == "omml":
            if self._try_omml_rendering(para, node):
                return

        # Fallback: LaTeX text
        para.add_run(node.text)

    def _try_omml_rendering(self, para, node: DocNode) -> bool:
        """
        Try to render equation as OMML.

        Returns True if successful, False to fallback to text.
        """
        if not omml_converter.is_pandoc_available():
            logger.debug("Pandoc not available for OMML")
            return False

        try:
            # Get LaTeX source
            latex_content = (
                node.metadata.get('latex_equation_primary') or
                node.metadata.get('latex_source') or
                node.text or
                ""
            )

            # Convert to OMML
            omml_xml = omml_converter.latex_to_omml(latex_content)

            if omml_xml:
                success = omml_converter.inject_omml_as_display(para, omml_xml)
                if success:
                    logger.debug(f"OMML rendering successful")
                    return True

        except Exception as e:
            logger.debug(f"OMML rendering failed: {e}")

        return False

    def _add_references_heading(self, node: DocNode) -> None:
        """Add references section heading."""
        self.add_heading(node.text or "References", level=1)

    def _add_reference_entry(self, node: DocNode) -> None:
        """Add reference entry with hanging indent."""
        para = self.doc.add_paragraph(node.text)
        self.style_manager.apply_body_style(para)
        para.paragraph_format.left_indent = Inches(0.5)
        para.paragraph_format.first_line_indent = Inches(-0.5)

    def _get_box_type(self, node_type: DocNodeType) -> str:
        """Map node type to theorem box type."""
        mapping = {
            DocNodeType.THEOREM: 'theorem',
            DocNodeType.LEMMA: 'lemma',
            DocNodeType.COROLLARY: 'corollary',
            DocNodeType.PROPOSITION: 'proposition',
            DocNodeType.DEFINITION: 'definition',
            DocNodeType.EXAMPLE: 'example',
        }
        return mapping.get(node_type, 'theorem')


def export_academic_docx(
    nodes: DocNodeList,
    output_path: str,
    config: Optional[AcademicExportConfig] = None,
    metadata: Optional[Dict[str, str]] = None
) -> str:
    """
    Convenience function for academic DOCX export.

    Args:
        nodes: DocNodeList from semantic extractor
        output_path: Output file path
        config: Academic export configuration
        metadata: Document metadata

    Returns:
        str: Absolute path to created file
    """
    exporter = AcademicDocxExporter(config)
    return exporter.export(nodes, output_path, metadata)


# Backward compatibility aliases
AcademicLayoutConfig = AcademicExportConfig
build_academic_docx = export_academic_docx
