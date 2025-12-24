#!/usr/bin/env python3
# -*- coding: utf-8 -*-

"""
Basic DOCX Exporter - Phase 1.4 Consolidation
==============================================

Simple text-to-DOCX exporter for basic documents.
Uses the base class for common functionality.
"""

from typing import Optional, Dict, List, Union
from dataclasses import dataclass

from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

from .docx_base import DocxExporterBase, ExportConfig
from config.logging_config import get_logger

logger = get_logger(__name__)


@dataclass
class BasicExportConfig(ExportConfig):
    """Configuration for basic DOCX export."""
    # Additional options specific to basic export
    preserve_line_breaks: bool = True
    auto_detect_headings: bool = False


class BasicDocxExporter(DocxExporterBase):
    """
    Basic DOCX exporter for simple text documents.

    Features:
    - Text to DOCX conversion
    - Automatic paragraph splitting
    - Optional heading detection
    - Simple formatting
    """

    def __init__(self, config: Optional[BasicExportConfig] = None):
        """
        Initialize basic exporter.

        Args:
            config: Export configuration (uses defaults if None)
        """
        super().__init__(config or BasicExportConfig())

    def _add_content(self, content: Union[str, List[str]]) -> None:
        """
        Add text content to document.

        Args:
            content: Text string or list of paragraphs
        """
        if isinstance(content, str):
            self._add_text_content(content)
        elif isinstance(content, list):
            self._add_list_content(content)
        else:
            logger.warning(f"Unsupported content type: {type(content)}")

    def _add_text_content(self, text: str) -> None:
        """
        Add text string to document.

        Splits by double newlines to create paragraphs.

        Args:
            text: Text content
        """
        # Split into paragraphs
        paragraphs = text.split('\n\n')

        for para_text in paragraphs:
            para_text = para_text.strip()
            if not para_text:
                continue

            # Check for heading
            if self.config.auto_detect_headings:
                heading_level = self._detect_heading(para_text)
                if heading_level:
                    self.add_heading(para_text, level=heading_level)
                    continue

            # Handle line breaks within paragraph
            if self.config.preserve_line_breaks:
                # Replace single newlines with soft line breaks
                para_text = para_text.replace('\n', '\n')

            self.add_paragraph(para_text)

    def _add_list_content(self, paragraphs: List[str]) -> None:
        """
        Add list of paragraphs to document.

        Args:
            paragraphs: List of paragraph strings
        """
        for para_text in paragraphs:
            if para_text.strip():
                self.add_paragraph(para_text.strip())

    def _detect_heading(self, text: str) -> Optional[int]:
        """
        Detect if text is a heading.

        Args:
            text: Text to analyze

        Returns:
            Heading level (1-3) or None if not a heading
        """
        text = text.strip()

        # All caps and short -> Heading 1
        if text.isupper() and len(text) < 60:
            return 1

        # Numbered sections
        import re
        if re.match(r'^[IVXLC]+\.?\s+', text):  # Roman numerals
            return 1
        if re.match(r'^\d+\.\s+', text):  # "1. Title"
            return 2
        if re.match(r'^\d+\.\d+\s+', text):  # "1.1 Subtitle"
            return 3

        return None


def export_basic_docx(
    text: str,
    output_path: str,
    config: Optional[BasicExportConfig] = None,
    metadata: Optional[Dict[str, str]] = None
) -> str:
    """
    Convenience function for basic DOCX export.

    Args:
        text: Text content to export
        output_path: Output file path
        config: Export configuration
        metadata: Document metadata

    Returns:
        str: Absolute path to created file
    """
    exporter = BasicDocxExporter(config)
    return exporter.export(text, output_path, metadata)
