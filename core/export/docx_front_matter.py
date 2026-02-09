
from typing import Dict, Any, Optional, Union
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, Inches, Cm
from core.export.docx_styles import StyleManager, ThemeConfig
from core.i18n import get_string

# Cover image support (Phase 2026-01)
try:
    from core.image_embedding import CoverEmbedder
    COVER_EMBEDDER_AVAILABLE = True
except ImportError:
    COVER_EMBEDDER_AVAILABLE = False

class FrontMatterGenerator:
    """
    Generates professional front matter:
    1. Title Page (Theme-aware)
    2. Table of Contents (Automatic field code)
    """

    def __init__(self, doc: Document, style_manager: StyleManager):
        self.doc = doc
        self.style_manager = style_manager
        self.theme = style_manager.theme

    def generate_cover_page(
        self,
        cover_image: Union[str, bytes],
        is_base64: bool = True
    ) -> bool:
        """
        Generate Cover Page (Page 1) with full-page image.

        Args:
            cover_image: Base64 string, data URI, or image bytes
            is_base64: True if cover_image is base64/data URI format

        Returns:
            True if cover was added successfully
        """
        if not COVER_EMBEDDER_AVAILABLE:
            print("Warning: CoverEmbedder not available, skipping cover page")
            return False

        if not cover_image:
            return False

        try:
            embedder = CoverEmbedder()
            embedder.add_cover_to_docx(self.doc, cover_image, is_base64)
            return True
        except Exception as e:
            print(f"Warning: Failed to add cover page: {e}")
            return False

    def generate_title_page(self, metadata: Dict[str, str]):
        """
        Generate a full Title Page.
        """
        # 1. Spacer (Push down)
        self._add_spacer(3)

        # 2. Title
        title = metadata.get('title', 'Untitled Document')
        self._add_title(title)

        # 3. Subtitle / Subject
        subject = metadata.get('subject')
        if subject:
            self._add_subtitle(subject)

        # 4. Large Spacer (Push to bottom)
        self._add_spacer(8)

        # 5. Author Block (skip generic "Unknown")
        author = metadata.get('author')
        if author and author != "Unknown":
            self._add_author(author)

        # 6. Page Break
        self.doc.add_page_break()

    def generate_toc(self, lang: str = "en"):
        """
        Generate Table of Contents (TOC).
        Note: Needs docx update/refresh to show page numbers.
        """
        # Heading (localized)
        toc_title = get_string("table_of_contents", lang)
        para = self.doc.add_heading(toc_title, level=1)
        self.style_manager.apply_heading_style(para, 1)
        
        # TOC Field Code
        p = self.doc.add_paragraph()
        run = p.add_run()
        
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn

        fldChar = OxmlElement('w:fldChar')
        fldChar.set(qn('w:fldCharType'), 'begin')
        run._element.append(fldChar)

        instrText = OxmlElement('w:instrText')
        instrText.set(qn('xml:space'), 'preserve')
        instrText.text = 'TOC \\o "1-3" \\h \\z \\u'
        run._element.append(instrText)

        fldChar2 = OxmlElement('w:fldChar')
        fldChar2.set(qn('w:fldCharType'), 'separate')
        run._element.append(fldChar2)

        fldChar3 = OxmlElement('w:fldChar')
        fldChar3.set(qn('w:fldCharType'), 'end')
        run._element.append(fldChar3)

        self.doc.add_page_break()


    def _add_title(self, text: str):
        para = self.doc.add_paragraph(text)
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Override style for Title
        run = para.runs[0]
        run.bold = True
        run.font.name = self.theme.fonts.heading
        run.font.size = Pt(28) # Huge
        run.font.color.rgb = self.theme.colors['heading']
        
        para.paragraph_format.space_after = Pt(24)

    def _add_subtitle(self, text: str):
        para = self.doc.add_paragraph(text)
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        run = para.runs[0]
        run.italic = True
        run.font.name = self.theme.fonts.heading
        run.font.size = Pt(16)
        run.font.color.rgb = self.theme.colors['body']
        
    def _add_author(self, text: str):
        para = self.doc.add_paragraph(f"Author: {text}")
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        run = para.runs[0]
        run.font.name = self.theme.fonts.body
        run.font.size = Pt(12)

    def _add_spacer(self, lines: int):
        for _ in range(lines):
            self.doc.add_paragraph()
