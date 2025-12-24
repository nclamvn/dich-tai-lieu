#!/usr/bin/env python3
# -*- coding: utf-8 -*-

"""
Academic DOCX Presentation Engine - Phase 1.4 Consolidation
============================================================

Text-to-DOCX exporter for translated documents with formula preservation.
Extends base class for common functionality.

Features:
- Formula sanitization (duplicate removal, delimiter fixing)
- Auto-heading detection
- Display equation formatting
- Boilerplate removal

HOTFIX 1.6.1 changes preserved:
- Parse by paragraphs instead of lines (fixes formula duplication)
- Proper display equation isolation
- Heading styles correctly applied
- Formula deduplication logic
"""

import re
from typing import Optional, Dict
from dataclasses import dataclass

from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

from .docx_base import DocxExporterBase, ExportConfig
from config.logging_config import get_logger

logger = get_logger(__name__)


@dataclass
class PresentationExportConfig(ExportConfig):
    """Configuration for presentation DOCX export."""
    # Override defaults for presentation style
    font_name: str = 'Times New Roman'
    font_size: int = 12
    line_spacing: float = 1.5

    # Presentation-specific options
    sanitize_formulas: bool = True
    auto_detect_headings: bool = True
    clean_boilerplate: bool = True


class PresentationDocxExporter(DocxExporterBase):
    """
    Text-to-DOCX exporter with formula preservation for STEM documents.

    Extends DocxExporterBase with:
    - Formula sanitization (duplicate removal, delimiter fixing)
    - Auto-heading detection from text patterns
    - Display equation formatting
    - Boilerplate removal
    """

    def __init__(self, config: Optional[PresentationExportConfig] = None):
        """
        Initialize presentation exporter.

        Args:
            config: PresentationExportConfig instance (uses defaults if None)
        """
        super().__init__(config or PresentationExportConfig())
        self.presentation_config = self.config  # Type hint

    def _add_content(self, content: str) -> None:
        """
        Add text content to document with formula handling.

        Args:
            content: Text content (with formulas)
        """
        text = content

        # Sanitize formulas if enabled
        if self.presentation_config.sanitize_formulas:
            text = self._sanitize_formulas(text)

        # Add content with paragraph-based parsing
        self._add_academic_content_fixed(text)

        # Clean boilerplate if enabled
        if self.presentation_config.clean_boilerplate:
            self._clean_footer_and_metadata()

    def export_academic(
        self,
        text: str,
        output_path: str,
        metadata: Optional[Dict] = None
    ) -> str:
        """
        Legacy method for backward compatibility.

        Args:
            text: Vietnamese translated text (with formulas preserved)
            output_path: Path to output DOCX file
            metadata: Optional metadata dict

        Returns:
            str: Absolute path to created file
        """
        return self.export(text, output_path, metadata)

    def _sanitize_formulas(self, text: str) -> str:
        """
        HOTFIX C1: Fix duplicate formulas and invalid dollar delimiters.

        Fixes:
        1. Remove pattern `$FORMULA$$FORMULA$` → keep only one
        2. Remove invalid `$$FORMULA$` (triple dollar) → normalize to `$$FORMULA$$`
        3. Deduplicate identical adjacent formulas
        """
        # Check for malformed delimiters
        if '$$x$' in text or '$$n=' in text:
            logger.debug("Found malformed delimiters in input")
            for i, line in enumerate(text.split('\n')):
                if '$$x$' in line or '$$n=' in line:
                    logger.debug(f"  Line {i}: {line[:100]}")

        result = text

        # Fix 1: Remove exact duplicates like `$\lambda...$`$\lambda...$`
        # Pattern: $CONTENT$$CONTENT$ where CONTENT is identical
        result = re.sub(
            r'\$([^\$]+?)\$\$\1\$',
            r'$$\1$$',  # Convert to display math
            result
        )

        # Fix 2: Normalize invalid `$$...$` based on content length
        # SHORT content (≤15 chars): $$x$ → $x$ (inline math)
        # LONG content (>15 chars): $$FORMULA$ → $$FORMULA$$ (display math)
        def normalize_malformed_delimiter(match):
            content = match.group(1)
            if len(content) <= 15:
                # Short content: convert to inline math
                return f'${content}$'
            else:
                # Long content: complete display math
                return f'$${content}$$'

        result = re.sub(
            r'\$\$([^\$]+?)\$(?!\$)',  # $$CONTENT$ not followed by $
            normalize_malformed_delimiter,
            result
        )

        # Fix 3: Normalize `$...$$$` based on content length (same logic as Fix 2)
        def normalize_triple_close(match):
            content = match.group(1)
            if len(content) <= 15:
                return f'${content}$'
            else:
                return f'$${content}$$'

        result = re.sub(
            r'\$([^\$]+?)\$\$\$',
            normalize_triple_close,
            result
        )

        # Fix 4: Normalize reversed delimiters `$VAR$$` to `$VAR$` (inline with display close)
        result = re.sub(
            r'(?<!\$)\$([^\$]{1,15}?)\$\$',
            r'$\1$',
            result
        )

        # Fix 5: Deduplicate identical adjacent display formulas
        # Match $$FORMULA$$ followed by same $$FORMULA$$
        def dedup_display(match):
            # Keep only first occurrence
            return match.group(1)

        result = re.sub(
            r'(\$\$[^\$]+?\$\$)\s*\1',
            dedup_display,
            result
        )

        return result

    def _add_academic_content_fixed(self, text: str) -> None:
        """
        HOTFIX C2 & C3: Add content with proper paragraph parsing and heading detection.

        Major changes from original:
        1. Parse by PARAGRAPHS (split by \\n\\n) not by lines
        2. Detect display equations as complete blocks
        3. Properly apply heading styles
        4. Each display equation gets its own paragraph with spacing
        """
        # Split by double newlines to get paragraphs
        paragraphs = re.split(r'\n\n+', text)

        is_first = True

        for para_text in paragraphs:
            para_text = para_text.strip()

            if not para_text:
                continue

            # PHASE 1.6.3: Sanitize paragraph text before adding to document
            para_text = self._sanitize_formulas(para_text)

            # Check if this is a display equation
            if self._is_display_equation(para_text):
                # Display equation: centered, with spacing
                p = self.doc.add_paragraph(para_text)
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p.paragraph_format.space_before = Pt(12)
                p.paragraph_format.space_after = Pt(12)
                # Monospace font for math
                for run in p.runs:
                    run.font.name = 'Courier New'
                    run.font.size = Pt(11)
                is_first = False
                continue

            # Detect heading type
            heading_level = self._detect_heading_level(para_text, is_first=is_first)

            if heading_level is not None:
                # Add as heading
                if heading_level == 0:
                    # Title
                    h = self.doc.add_heading(para_text, level=0)
                    h.alignment = WD_ALIGN_PARAGRAPH.CENTER
                else:
                    self.doc.add_heading(para_text, level=heading_level)
                is_first = False
            else:
                # Normal paragraph
                self.doc.add_paragraph(para_text)
                is_first = False

    def _is_display_equation(self, text: str) -> bool:
        """Check if text is a display equation"""
        text = text.strip()
        return (
            text.startswith('$$') or
            text.startswith(r'\[') or
            text.startswith(r'\begin{equation') or
            text.startswith(r'\begin{align')
        )

    def _detect_heading_level(self, text: str, is_first: bool = False) -> Optional[int]:
        """
        HOTFIX C3: Detect heading level for a paragraph.

        Returns:
            0 for Title
            1 for Heading 1
            2 for Heading 2
            3 for Heading 3
            None for normal paragraph
        """
        text = text.strip()

        # First paragraph is Title
        if is_first:
            return 0

        # Pattern "1. " or "2. " → Heading 2
        if re.match(r'^\d+\.\s+', text):
            return 2

        # Pattern "1.1 " or "2.3 " → Heading 3
        if re.match(r'^\d+\.\d+\s+', text):
            return 3

        # ALL CAPS and short → Heading 1
        if text.isupper() and len(text) < 60:
            return 1

        # Not a heading
        return None

    def _clean_footer_and_metadata(self) -> None:
        """
        HOTFIX C5: Remove boilerplate footer and metadata.

        Removes:
        - "AI Translator Pro | Page N" footer
        - Any watermarks
        - Demo/draft markers
        """
        if not self.doc:
            return

        for section in self.doc.sections:
            # Clear footer
            footer = section.footer
            for para in footer.paragraphs:
                para.text = ""

            # Clear header if it contains boilerplate
            header = section.header
            for para in header.paragraphs:
                text = para.text.lower()
                if 'ai translator' in text or 'draft' in text or 'demo' in text:
                    para.text = ""


def export_presentation_docx(
    text: str,
    output_path: str,
    config: Optional[PresentationExportConfig] = None,
    metadata: Optional[Dict] = None
) -> str:
    """
    Convenience function for presentation DOCX export.

    Args:
        text: Vietnamese text with formulas preserved
        output_path: Output DOCX file path
        config: PresentationExportConfig instance (optional)
        metadata: Document metadata dict (optional)

    Returns:
        str: Absolute path to created file
    """
    exporter = PresentationDocxExporter(config=config)
    return exporter.export(text, output_path, metadata)


# Backward compatibility aliases
AcademicDocxExporter = PresentationDocxExporter


def export_academic_docx(
    text: str,
    output_path: str,
    config=None,
    metadata: Optional[Dict] = None
) -> str:
    """
    DEPRECATED: Use export_presentation_docx instead.

    Kept for backward compatibility.
    """
    return export_presentation_docx(text, output_path, config, metadata)
