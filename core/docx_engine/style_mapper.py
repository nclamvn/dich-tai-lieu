"""
Style Mapper - Maps semantic content blocks to DOCX styles.
Handles the translation between our ContentBlock types and actual DOCX formatting.
"""

from typing import Dict, Optional, List, TYPE_CHECKING
from dataclasses import dataclass
import logging

from docx import Document
from docx.text.paragraph import Paragraph
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING, WD_BREAK
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

from .models import (
    ContentBlock, BlockType, TextRun, InlineStyle,
    ListItem, TableData, TableCell, ListType
)
from .templates.base import DocxTemplate, ParagraphSpec, FontSpec

if TYPE_CHECKING:
    from docx.table import Table

logger = logging.getLogger(__name__)


@dataclass
class RenderContext:
    """Context for rendering decisions"""
    is_first_para_in_chapter: bool = False
    is_first_para_after_heading: bool = False
    previous_block_type: Optional[BlockType] = None
    chapter_number: int = 0
    chapter_title: str = ""
    list_level: int = 0


class StyleMapper:
    """
    Maps ContentBlocks to styled DOCX elements.

    Usage:
        mapper = StyleMapper(document, template)
        for block in chapter.content:
            mapper.render_block(block, context)
    """

    def __init__(self, document: Document, template: DocxTemplate):
        self.doc = document
        self.template = template
        self.styles = template.get_styles()

        # Ensure styles are registered in document
        self._register_styles()

    def _register_styles(self):
        """Register template styles in DOCX document"""
        # DOCX has built-in styles, we'll use paragraph formatting directly
        # rather than creating new styles (simpler and more portable)
        pass

    def render_block(self, block: ContentBlock, context: RenderContext) -> Optional[Paragraph]:
        """
        Render a ContentBlock to the document.

        Args:
            block: The content block to render
            context: Rendering context

        Returns:
            The created paragraph (or None for tables)
        """
        if block.type == BlockType.HEADING:
            return self._render_heading(block, context)
        elif block.type == BlockType.PARAGRAPH:
            return self._render_paragraph(block, context)
        elif block.type == BlockType.LIST:
            return self._render_list(block, context)
        elif block.type == BlockType.TABLE:
            self._render_table(block, context)
            return None
        elif block.type == BlockType.QUOTE:
            return self._render_quote(block, context)
        elif block.type == BlockType.CODE:
            return self._render_code(block, context)
        elif block.type == BlockType.FIGURE:
            return self._render_figure(block, context)
        elif block.type == BlockType.PAGE_BREAK:
            return self._render_page_break()
        else:
            logger.warning(f"Unknown block type: {block.type}")
            return self._render_paragraph(block, context)

    def _render_heading(self, block: ContentBlock, context: RenderContext) -> Paragraph:
        """Render a heading"""
        level = min(block.level, 3)  # We support H1-H3
        style_name = f"heading_{level}"
        spec = self.styles.get(style_name, self.styles['heading_1'])

        para = self.doc.add_paragraph()
        self._apply_paragraph_spec(para, spec)

        # Add text
        text = block.content if isinstance(block.content, str) else str(block.content)
        run = para.add_run(text)
        self._apply_font_spec(run, spec.font)

        return para

    def _render_paragraph(self, block: ContentBlock, context: RenderContext) -> Paragraph:
        """Render a paragraph with inline formatting"""
        # Choose style based on context
        if context.is_first_para_after_heading or context.is_first_para_in_chapter:
            spec = self.styles.get('body_first', self.styles['body'])
        else:
            spec = self.styles['body']

        para = self.doc.add_paragraph()
        self._apply_paragraph_spec(para, spec)

        # Handle content
        content = block.content

        if isinstance(content, str):
            run = para.add_run(content)
            self._apply_font_spec(run, spec.font)
        elif isinstance(content, list):
            # List of TextRuns
            for text_run in content:
                if isinstance(text_run, TextRun):
                    run = para.add_run(text_run.text)
                    self._apply_font_spec(run, spec.font)
                    self._apply_inline_style(run, text_run.style)
                else:
                    run = para.add_run(str(text_run))
                    self._apply_font_spec(run, spec.font)

        return para

    def _render_list(self, block: ContentBlock, context: RenderContext) -> Optional[Paragraph]:
        """Render a list (bullet or numbered)"""
        items = block.content
        if not isinstance(items, list):
            return self._render_paragraph(block, context)

        list_type = block.style_hints.get('list_type', 'bullet')
        style_name = 'list_numbered' if list_type == 'numbered' else 'list_bullet'
        spec = self.styles.get(style_name, self.styles['body'])

        last_para = None
        for i, item in enumerate(items):
            para = self.doc.add_paragraph()
            self._apply_paragraph_spec(para, spec)

            # Add bullet/number prefix
            if list_type == 'numbered':
                prefix = f"{i + 1}. "
            else:
                prefix = "• "

            # Add prefix
            run = para.add_run(prefix)
            self._apply_font_spec(run, spec.font)

            # Add content
            if isinstance(item, ListItem):
                for text_run in item.content:
                    run = para.add_run(text_run.text)
                    self._apply_font_spec(run, spec.font)
                    self._apply_inline_style(run, text_run.style)
            else:
                run = para.add_run(str(item))
                self._apply_font_spec(run, spec.font)

            # Add left indent for list items
            para.paragraph_format.left_indent = Cm(0.75)

            last_para = para

        return last_para

    def _render_table(self, block: ContentBlock, context: RenderContext):
        """Render a table"""
        table_data = block.content
        if not isinstance(table_data, TableData) or not table_data.rows:
            return

        rows = table_data.rows
        num_cols = max(len(row) for row in rows)

        table = self.doc.add_table(rows=len(rows), cols=num_cols)
        table.style = 'Table Grid'

        for i, row in enumerate(rows):
            for j, cell in enumerate(row):
                if j < num_cols:
                    table_cell = table.cell(i, j)

                    # Clear default paragraph
                    table_cell.paragraphs[0].clear()

                    # Add cell content
                    for text_run in cell.content:
                        run = table_cell.paragraphs[0].add_run(text_run.text)

                        # Header row styling
                        if cell.is_header or i == 0:
                            run.bold = True

                        self._apply_inline_style(run, text_run.style)

    def _render_quote(self, block: ContentBlock, context: RenderContext) -> Paragraph:
        """Render a block quote"""
        spec = self.styles.get('quote', self.styles['body'])

        para = self.doc.add_paragraph()
        self._apply_paragraph_spec(para, spec)

        # Add left indent for quote
        para.paragraph_format.left_indent = Cm(1)
        para.paragraph_format.right_indent = Cm(1)

        # Add content
        content = block.content
        if isinstance(content, str):
            run = para.add_run(content)
            self._apply_font_spec(run, spec.font)
        elif isinstance(content, list):
            for text_run in content:
                if isinstance(text_run, TextRun):
                    run = para.add_run(text_run.text)
                    self._apply_font_spec(run, spec.font)
                    self._apply_inline_style(run, text_run.style)
                else:
                    run = para.add_run(str(text_run))
                    self._apply_font_spec(run, spec.font)

        return para

    def _render_code(self, block: ContentBlock, context: RenderContext) -> Paragraph:
        """Render a code block"""
        spec = self.styles.get('code', self.styles['body'])

        para = self.doc.add_paragraph()
        self._apply_paragraph_spec(para, spec)

        # Add light gray background (via shading)
        self._add_shading(para, "F0F0F0")

        # Add content preserving whitespace
        content = block.content if isinstance(block.content, str) else str(block.content)
        run = para.add_run(content)
        self._apply_font_spec(run, spec.font)

        return para

    def _render_figure(self, block: ContentBlock, context: RenderContext) -> Paragraph:
        """Render a figure (placeholder - actual image handling in renderer)"""
        spec = self.styles.get('caption', self.styles['body'])

        para = self.doc.add_paragraph()
        self._apply_paragraph_spec(para, spec)

        caption = block.caption or f"[Figure: {block.content}]"
        run = para.add_run(caption)
        self._apply_font_spec(run, spec.font)

        return para

    def _render_page_break(self) -> Paragraph:
        """Insert a page break"""
        para = self.doc.add_paragraph()
        run = para.add_run()
        run.add_break(WD_BREAK.PAGE)
        return para

    def _apply_paragraph_spec(self, para: Paragraph, spec: ParagraphSpec):
        """Apply ParagraphSpec to a paragraph"""
        pf = para.paragraph_format

        # Alignment
        pf.alignment = spec.alignment

        # Spacing
        pf.space_before = spec.space_before
        pf.space_after = spec.space_after

        # Line spacing
        if spec.line_spacing:
            pf.line_spacing = spec.line_spacing

        # First line indent
        if spec.first_line_indent is not None:
            pf.first_line_indent = spec.first_line_indent

        # Keep with next
        pf.keep_with_next = spec.keep_with_next

        # Page break before
        pf.page_break_before = spec.page_break_before

        # Widow control
        pf.widow_control = spec.widow_control

    def _apply_font_spec(self, run, spec: FontSpec):
        """Apply FontSpec to a run"""
        run.font.name = spec.name
        run.font.size = spec.size
        run.bold = spec.bold
        run.italic = spec.italic

        if spec.color:
            run.font.color.rgb = spec.color

        # Set East Asian font for CJK support
        r = run._element
        rPr = r.get_or_add_rPr()
        rFonts = rPr.get_or_add_rFonts()
        rFonts.set(qn('w:eastAsia'), spec.name)

    def _apply_inline_style(self, run, style: InlineStyle):
        """Apply InlineStyle to a run"""
        if style.bold:
            run.bold = True
        if style.italic:
            run.italic = True
        if style.underline:
            run.underline = True
        if style.strikethrough:
            run.font.strike = True
        if style.superscript:
            run.font.superscript = True
        if style.subscript:
            run.font.subscript = True
        if style.code:
            run.font.name = "Consolas"
            run.font.size = Pt(10)

    def _add_shading(self, para: Paragraph, color: str):
        """Add background shading to paragraph"""
        shd = OxmlElement('w:shd')
        shd.set(qn('w:fill'), color)
        para._p.get_or_add_pPr().append(shd)
