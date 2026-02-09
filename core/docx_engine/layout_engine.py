"""
Layout Engine - Handles document-level layout: TOC, headers/footers, pagination.
"""

from typing import Optional, List
from pathlib import Path
import logging

from docx import Document
from docx.shared import Pt, Cm, Inches
from docx.enum.section import WD_ORIENT, WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

from .models import NormalizedDocument, DocumentMeta, TableOfContents, TocItem
from .templates.base import DocxTemplate, PageSetup, HeaderFooterSpec, TocSpec, ParagraphSpec
from core.i18n import format_chapter_title, get_string

logger = logging.getLogger(__name__)


class LayoutEngine:
    """
    Handles document-level layout concerns:
    - Page setup (size, margins, orientation)
    - Headers and footers
    - Table of contents generation
    - Section breaks
    - Page numbering
    """

    def __init__(self, document: Document, template: DocxTemplate):
        self.doc = document
        self.template = template

    def setup_document(self, meta: DocumentMeta):
        """Configure document page setup"""
        page_setup = self.template.get_page_setup()

        for section in self.doc.sections:
            # Page size
            section.page_width = page_setup.width
            section.page_height = page_setup.height

            # Margins
            section.top_margin = page_setup.top_margin
            section.bottom_margin = page_setup.bottom_margin
            section.left_margin = page_setup.left_margin
            section.right_margin = page_setup.right_margin

            # Gutter for binding
            section.gutter = page_setup.gutter

            # Mirror margins for book
            if page_setup.mirror_margins:
                section.different_first_page_header_footer = True

    def add_header_footer(self, meta: DocumentMeta):
        """Add headers and footers to document"""
        spec = self.template.get_header_footer()

        if not spec.show_header and not spec.show_footer:
            return

        for section in self.doc.sections:
            # Different first page
            section.different_first_page_header_footer = spec.different_first_page

            # Header
            if spec.show_header:
                self._setup_header(section, meta, spec)

            # Footer
            if spec.show_footer:
                self._setup_footer(section, meta, spec)

    def _setup_header(self, section, meta: DocumentMeta, spec: HeaderFooterSpec):
        """Setup header for a section"""
        header = section.header

        # Clear existing
        if header.paragraphs:
            header.paragraphs[0].clear()

        # Build header content
        para = header.paragraphs[0] if header.paragraphs else header.add_paragraph()

        # Determine content
        content_parts = []
        if spec.header_left:
            content_parts.append(('left', self._format_placeholder(spec.header_left, meta)))
        if spec.header_center:
            content_parts.append(('center', self._format_placeholder(spec.header_center, meta)))
        if spec.header_right:
            content_parts.append(('right', self._format_placeholder(spec.header_right, meta)))

        # Simple implementation: just center if center, otherwise left
        if content_parts:
            for position, text in content_parts:
                if position == 'center':
                    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = para.add_run(text)
                run.font.name = spec.font.name
                run.font.size = spec.font.size

    def _setup_footer(self, section, meta: DocumentMeta, spec: HeaderFooterSpec):
        """Setup footer for a section"""
        footer = section.footer

        # Clear existing
        if footer.paragraphs:
            footer.paragraphs[0].clear()

        para = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()

        # Page number handling
        content_parts = []
        if spec.footer_left:
            content_parts.append(('left', spec.footer_left))
        if spec.footer_center:
            content_parts.append(('center', spec.footer_center))
        if spec.footer_right:
            content_parts.append(('right', spec.footer_right))

        for position, text in content_parts:
            if '{page}' in text:
                # Add page number field
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER if position == 'center' else WD_ALIGN_PARAGRAPH.LEFT

                # Text before {page}
                before = text.split('{page}')[0]
                after = text.split('{page}')[1] if '{page}' in text else ''

                if before:
                    run = para.add_run(before)
                    run.font.name = spec.font.name
                    run.font.size = spec.font.size

                # Add PAGE field
                self._add_page_number_field(para, spec)

                if after:
                    run = para.add_run(after)
                    run.font.name = spec.font.name
                    run.font.size = spec.font.size
            else:
                formatted = self._format_placeholder(text, meta)
                run = para.add_run(formatted)
                run.font.name = spec.font.name
                run.font.size = spec.font.size

    def _add_page_number_field(self, para, spec: HeaderFooterSpec):
        """Add a PAGE field for page numbering"""
        run = para.add_run()

        # Create field code
        fldChar1 = OxmlElement('w:fldChar')
        fldChar1.set(qn('w:fldCharType'), 'begin')

        instrText = OxmlElement('w:instrText')
        instrText.text = 'PAGE'

        fldChar2 = OxmlElement('w:fldChar')
        fldChar2.set(qn('w:fldCharType'), 'end')

        run._r.append(fldChar1)
        run._r.append(instrText)
        run._r.append(fldChar2)

        run.font.name = spec.font.name
        run.font.size = spec.font.size

    def _format_placeholder(self, text: str, meta: DocumentMeta) -> str:
        """Replace placeholders in header/footer text"""
        replacements = {
            '{title}': meta.title or '',
            '{author}': meta.author or '',
            '{date}': meta.date or '',
            '{running_title}': meta.running_title or meta.title or '',
        }

        result = text
        for placeholder, value in replacements.items():
            result = result.replace(placeholder, value)

        return result

    def generate_toc(self, toc: TableOfContents, insert_at_beginning: bool = True, lang: str = "en"):
        """
        Generate table of contents.

        Note: DOCX TOC requires field codes for auto-updating.
        This creates a simple static TOC.
        For auto-updating TOC, use Word's built-in TOC feature.
        """
        toc_spec = self.template.get_toc_spec()

        # TOC title (localized)
        toc_title = get_string("table_of_contents", lang)
        title_para = self.doc.add_paragraph()
        if toc_spec.title_style:
            self._apply_para_spec(title_para, toc_spec.title_style)
        run = title_para.add_run(toc_title)
        if toc_spec.title_style:
            self._apply_font_from_spec(run, toc_spec.title_style)

        # TOC entries
        for item in toc.items:
            para = self.doc.add_paragraph()

            # Get style for level
            if item.level == 1 and toc_spec.level1_style:
                self._apply_para_spec(para, toc_spec.level1_style)
                # Add left indent based on level
            elif item.level == 2 and toc_spec.level2_style:
                self._apply_para_spec(para, toc_spec.level2_style)
                para.paragraph_format.left_indent = Cm(0.5)
            elif item.level >= 3 and toc_spec.level3_style:
                self._apply_para_spec(para, toc_spec.level3_style)
                para.paragraph_format.left_indent = Cm(1.0)

            # Chapter number + title
            if item.chapter_number:
                text = format_chapter_title(item.chapter_number, item.title, lang)
            else:
                text = item.title

            run = para.add_run(text)

            # Apply font based on level
            if item.level == 1 and toc_spec.level1_style:
                self._apply_font_from_spec(run, toc_spec.level1_style)
            elif item.level == 2 and toc_spec.level2_style:
                self._apply_font_from_spec(run, toc_spec.level2_style)
            elif item.level >= 3 and toc_spec.level3_style:
                self._apply_font_from_spec(run, toc_spec.level3_style)

            # Page number (placeholder - real page numbers need field codes)
            if toc_spec.show_page_numbers and item.page_number:
                if toc_spec.dot_leader:
                    # Add tab with dot leader
                    para.add_run('\t')
                para.add_run(str(item.page_number))

        # Page break after TOC
        self.doc.add_page_break()

    def add_section_break(self, break_type: str = 'page'):
        """
        Add a section break.

        Args:
            break_type: 'page', 'odd_page', 'continuous'
        """
        break_map = {
            'page': WD_SECTION.NEW_PAGE,
            'odd_page': WD_SECTION.ODD_PAGE,
            'continuous': WD_SECTION.CONTINUOUS,
        }

        self.doc.add_section(break_map.get(break_type, WD_SECTION.NEW_PAGE))

    def _apply_para_spec(self, para, spec: ParagraphSpec):
        """Apply paragraph spec"""
        pf = para.paragraph_format
        pf.alignment = spec.alignment
        pf.space_before = spec.space_before
        pf.space_after = spec.space_after
        if spec.line_spacing:
            pf.line_spacing = spec.line_spacing

    def _apply_font_from_spec(self, run, spec: ParagraphSpec):
        """Apply font from paragraph spec"""
        if spec.font:
            run.font.name = spec.font.name
            run.font.size = spec.font.size
            run.bold = spec.font.bold
            run.italic = spec.font.italic
            if spec.font.color:
                run.font.color.rgb = spec.font.color
