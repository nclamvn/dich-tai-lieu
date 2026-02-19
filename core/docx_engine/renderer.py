"""
Main DOCX Renderer - Orchestrates all components to produce final document.
"""

from typing import Optional, Union
from pathlib import Path
import logging

from docx import Document
from docx.shared import Pt, Cm

from .models import NormalizedDocument, Chapter, ContentBlock, BlockType
from .normalizer import DocumentNormalizer
from .style_mapper import StyleMapper, RenderContext
from .layout_engine import LayoutEngine
from .templates.base import DocxTemplate, create_template
from core.i18n import get_string, format_chapter_title

logger = logging.getLogger(__name__)


class DocxRenderer:
    """
    Main DOCX renderer that orchestrates all components.

    Usage:
        # From Agent 2 output folder
        renderer = DocxRenderer(template='ebook')
        output_path = renderer.render(
            source_folder='book_output/',
            output_path='final_book.docx'
        )

        # From NormalizedDocument
        renderer = DocxRenderer(template='academic')
        output_path = renderer.render_document(doc, 'paper.docx')
    """

    def __init__(
        self,
        template: Union[str, DocxTemplate] = 'ebook',
        normalizer: Optional[DocumentNormalizer] = None
    ):
        """
        Initialize renderer.

        Args:
            template: Template name ('ebook', 'academic', 'business') or DocxTemplate instance
            normalizer: Optional custom normalizer
        """
        if isinstance(template, str):
            self.template = create_template(template)
        else:
            self.template = template

        self.normalizer = normalizer or DocumentNormalizer()

    def render(
        self,
        source_folder: str,
        output_path: str,
        include_toc: bool = True,
        include_glossary: bool = True
    ) -> Path:
        """
        Render Agent 2 output folder to DOCX.

        Args:
            source_folder: Path to Agent 2 output folder
            output_path: Output DOCX file path
            include_toc: Whether to include table of contents
            include_glossary: Whether to include glossary

        Returns:
            Path to created DOCX file
        """
        logger.info(f"Rendering {source_folder} with {self.template.name} template")

        # Normalize input
        doc = self.normalizer.from_agent2_output(source_folder)

        return self.render_document(
            doc,
            output_path,
            include_toc=include_toc,
            include_glossary=include_glossary
        )

    def render_document(
        self,
        doc: NormalizedDocument,
        output_path: str,
        include_toc: bool = True,
        include_glossary: bool = True
    ) -> Path:
        """
        Render NormalizedDocument to DOCX.

        Args:
            doc: NormalizedDocument to render
            output_path: Output DOCX file path
            include_toc: Whether to include table of contents
            include_glossary: Whether to include glossary

        Returns:
            Path to created DOCX file
        """
        output = Path(output_path)

        # Create DOCX document
        docx = Document()

        # Initialize components
        style_mapper = StyleMapper(docx, self.template)
        layout_engine = LayoutEngine(docx, self.template)

        # Setup document layout
        layout_engine.setup_document(doc.meta)
        layout_engine.add_header_footer(doc.meta)

        # Render title page
        lang = doc.meta.language or "en"
        self._render_title_page(docx, doc, style_mapper)

        # Render TOC
        if include_toc and doc.toc.items:
            layout_engine.generate_toc(doc.toc, lang=lang)

        # Render front matter
        if doc.front_matter.items:
            self._render_front_matter(docx, doc, style_mapper)

        # Render chapters
        for i, chapter in enumerate(doc.chapters):
            self._render_chapter(docx, chapter, style_mapper, layout_engine, i == 0, lang=lang)

        # Render glossary
        if include_glossary and doc.glossary and doc.glossary.items:
            self._render_glossary(docx, doc, style_mapper)

        # Render bibliography
        if doc.bibliography and doc.bibliography.items:
            self._render_bibliography(docx, doc, style_mapper)

        # Save document
        output.parent.mkdir(parents=True, exist_ok=True)
        docx.save(str(output))

        logger.info(f"Successfully rendered to {output}")
        return output

    def render_markdown(
        self,
        markdown_content: str,
        output_path: str,
        title: str = "Untitled",
        author: str = "Unknown"
    ) -> Path:
        """
        Render markdown content directly to DOCX.

        Args:
            markdown_content: Markdown text to render
            output_path: Output DOCX file path
            title: Document title
            author: Document author

        Returns:
            Path to created DOCX file
        """
        from .models import DocumentMeta

        meta = DocumentMeta(title=title, author=author)
        doc = self.normalizer.from_markdown(markdown_content, meta)

        return self.render_document(doc, output_path, include_toc=False)

    def _render_title_page(
        self,
        docx: Document,
        doc: NormalizedDocument,
        style_mapper: StyleMapper
    ):
        """Render the title page"""
        styles = self.template.get_styles()

        # Title
        if doc.meta.title:
            title_spec = styles.get('title')
            if title_spec:
                para = docx.add_paragraph()
                style_mapper._apply_paragraph_spec(para, title_spec)
                run = para.add_run(doc.meta.title)
                style_mapper._apply_font_spec(run, title_spec.font)

        # Subtitle
        if doc.meta.subtitle:
            subtitle_spec = styles.get('subtitle')
            if subtitle_spec:
                para = docx.add_paragraph()
                style_mapper._apply_paragraph_spec(para, subtitle_spec)
                run = para.add_run(doc.meta.subtitle)
                style_mapper._apply_font_spec(run, subtitle_spec.font)

        # Author
        if doc.meta.author:
            author_spec = styles.get('author', styles['body'])
            para = docx.add_paragraph()
            style_mapper._apply_paragraph_spec(para, author_spec)
            run = para.add_run(doc.meta.author)
            style_mapper._apply_font_spec(run, author_spec.font)

        # Translator
        if doc.meta.translator:
            lang = doc.meta.language or "en"
            translator_label = get_string("translator", lang)
            para = docx.add_paragraph()
            run = para.add_run(f"{translator_label}: {doc.meta.translator}")

        # Page break after title
        docx.add_page_break()

    def _render_front_matter(
        self,
        docx: Document,
        doc: NormalizedDocument,
        style_mapper: StyleMapper
    ):
        """Render front matter sections"""
        for item in doc.front_matter.items:
            # Section title
            if item.title:
                context = RenderContext()
                heading = ContentBlock(type=BlockType.HEADING, level=1, content=item.title)
                style_mapper.render_block(heading, context)

            # Content
            context = RenderContext(is_first_para_in_chapter=True)
            for i, block in enumerate(item.content):
                style_mapper.render_block(block, context)
                context.is_first_para_in_chapter = False
                context.is_first_para_after_heading = (block.type == BlockType.HEADING)
                context.previous_block_type = block.type

        docx.add_page_break()

    def _render_chapter(
        self,
        docx: Document,
        chapter: Chapter,
        style_mapper: StyleMapper,
        layout_engine: LayoutEngine,
        is_first_chapter: bool,
        lang: str = "en",
    ):
        """Render a single chapter"""
        styles = self.template.get_styles()

        # Chapter break (if not first chapter)
        break_type = self.template.get_chapter_break_type()
        if not is_first_chapter and break_type != 'none':
            if break_type == 'page':
                docx.add_page_break()
            elif break_type == 'odd_page':
                layout_engine.add_section_break('odd_page')

        # Chapter title
        chapter_title = format_chapter_title(chapter.number, chapter.title, lang) if chapter.number else chapter.title
        heading_spec = styles.get('heading_1')

        if heading_spec:
            para = docx.add_paragraph()
            style_mapper._apply_paragraph_spec(para, heading_spec)
            run = para.add_run(chapter_title)
            style_mapper._apply_font_spec(run, heading_spec.font)

        # Epigraph
        if chapter.epigraph:
            epigraph_spec = styles.get('epigraph', styles.get('quote', styles['body']))
            para = docx.add_paragraph()
            style_mapper._apply_paragraph_spec(para, epigraph_spec)
            run = para.add_run(chapter.epigraph)
            style_mapper._apply_font_spec(run, epigraph_spec.font)

        # Chapter content
        context = RenderContext(
            is_first_para_in_chapter=True,
            chapter_number=chapter.number,
            chapter_title=chapter.title
        )

        for i, block in enumerate(chapter.content):
            # Skip if first block is H1 (already rendered chapter title)
            if i == 0 and block.type == BlockType.HEADING and block.level == 1:
                continue

            style_mapper.render_block(block, context)

            # Update context
            context.is_first_para_in_chapter = False
            context.is_first_para_after_heading = (block.type == BlockType.HEADING)
            context.previous_block_type = block.type

    def _render_glossary(
        self,
        docx: Document,
        doc: NormalizedDocument,
        style_mapper: StyleMapper
    ):
        """Render glossary section"""
        styles = self.template.get_styles()

        # Page break
        docx.add_page_break()

        # Title
        heading_spec = styles.get('heading_1')
        if heading_spec:
            para = docx.add_paragraph()
            style_mapper._apply_paragraph_spec(para, heading_spec)
            run = para.add_run(doc.glossary.title)
            style_mapper._apply_font_spec(run, heading_spec.font)

        # Glossary items
        body_spec = styles.get('body')
        for item in doc.glossary.items:
            para = docx.add_paragraph()
            style_mapper._apply_paragraph_spec(para, body_spec)

            # Term (bold)
            run = para.add_run(item.term)
            style_mapper._apply_font_spec(run, body_spec.font)
            run.bold = True

            # Separator
            para.add_run(': ')

            # Definition
            run = para.add_run(item.definition)
            style_mapper._apply_font_spec(run, body_spec.font)

            # Source term
            if item.source_term:
                run = para.add_run(f' ({item.source_term})')
                style_mapper._apply_font_spec(run, body_spec.font)
                run.italic = True

    def _render_bibliography(
        self,
        docx: Document,
        doc: NormalizedDocument,
        style_mapper: StyleMapper
    ):
        """Render bibliography section"""
        styles = self.template.get_styles()

        # Page break
        docx.add_page_break()

        # Title
        heading_spec = styles.get('heading_1')
        if heading_spec:
            para = docx.add_paragraph()
            style_mapper._apply_paragraph_spec(para, heading_spec)
            run = para.add_run(doc.bibliography.title)
            style_mapper._apply_font_spec(run, heading_spec.font)

        # Bibliography items
        bib_spec = styles.get('bibliography', styles['body'])
        for item in doc.bibliography.items:
            para = docx.add_paragraph()
            style_mapper._apply_paragraph_spec(para, bib_spec)

            # Hanging indent
            para.paragraph_format.first_line_indent = Cm(-1.27)
            para.paragraph_format.left_indent = Cm(1.27)

            run = para.add_run(item.formatted)
            style_mapper._apply_font_spec(run, bib_spec.font)
