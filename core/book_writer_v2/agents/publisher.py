"""
Publisher Agent

Generates final output files with optional illustration support.
"""

import os
import re
from pathlib import Path
from typing import Dict, List, Optional

from .base import BaseAgent, AgentContext
from ..models import BookBlueprint, BookProject
from ..config import OutputFormat
from ..illustration_models import (
    IllustrationPlan,
    ImagePlacement,
    LayoutMode,
    ImageSize,
    LayoutConfig,
)


class PublisherAgent(BaseAgent[BookProject, Dict[str, str]]):
    """
    Agent 9: Publisher

    Generates final output files:
    - DOCX with professional formatting
    - PDF with proper layout
    - Markdown source
    - HTML (optional)
    - EPUB (optional)

    If the project has an illustration_plan, images are inserted.
    For illustrated DOCX/EPUB/PDF, delegates to LayoutEngine.
    """

    @property
    def name(self) -> str:
        return "Publisher"

    @property
    def description(self) -> str:
        return "Generate final output files in requested formats"

    async def execute(
        self,
        input_data: BookProject,
        context: AgentContext
    ) -> Dict[str, str]:
        project = input_data
        blueprint = project.blueprint

        context.report_progress("Generating output files...", 0)

        output_dir = Path(self.config.output_dir) / project.id
        output_dir.mkdir(parents=True, exist_ok=True)

        # Illustration plan (None for text-only books)
        illus_plan = project.illustration_plan

        # Image directory for illustrated books
        image_dir = f"data/uploads/books/{project.id}/images"

        # Detect genre from illustration plan
        genre = None
        if illus_plan and hasattr(illus_plan, "genre"):
            genre = illus_plan.genre

        import asyncio

        output_files = {}
        formats = self.config.output_formats

        # Delegate illustrated DOCX/EPUB/PDF to LayoutEngine
        layout_formats = []
        direct_formats = []
        for fmt in formats:
            if illus_plan and illus_plan.placements and fmt in (
                OutputFormat.DOCX, OutputFormat.EPUB, OutputFormat.PDF
            ):
                layout_formats.append(fmt.value)
            else:
                direct_formats.append(fmt)

        # Render illustrated formats via LayoutEngine
        if layout_formats:
            try:
                from ..renderers.layout_engine import LayoutEngine
                from ..illustration_models import BookGenre
                engine = LayoutEngine()
                rendered = engine.render(
                    blueprint=blueprint,
                    plan=illus_plan,
                    image_dir=image_dir,
                    formats=layout_formats,
                    output_dir=output_dir,
                    genre=genre or BookGenre.NON_FICTION,
                )
                output_files.update(rendered)
            except Exception as e:
                self.logger.warning(f"LayoutEngine failed, falling back: {e}")
                # Fall back to direct generation
                for fmt_str in layout_formats:
                    direct_formats.append(OutputFormat(fmt_str))

        # Direct generation for text-only or non-layout formats
        async def generate_format(fmt):
            try:
                if fmt == OutputFormat.MARKDOWN:
                    path = await self._generate_markdown(blueprint, output_dir, illus_plan)
                    return "markdown", str(path)
                elif fmt == OutputFormat.DOCX:
                    path = await self._generate_docx(blueprint, output_dir, illus_plan, image_dir)
                    return "docx", str(path)
                elif fmt == OutputFormat.PDF:
                    path = await self._generate_pdf(blueprint, output_dir, illus_plan)
                    return "pdf", str(path)
                elif fmt == OutputFormat.HTML:
                    path = await self._generate_html(blueprint, output_dir, illus_plan)
                    return "html", str(path)
                elif fmt == OutputFormat.EPUB:
                    path = await self._generate_epub(blueprint, output_dir, illus_plan, genre)
                    return "epub", str(path)
            except Exception as e:
                self.logger.error(f"Failed to generate {fmt.value}: {e}")
            return None, None

        if direct_formats:
            results = await asyncio.gather(*[generate_format(fmt) for fmt in direct_formats])
            for key, path in results:
                if key and path:
                    output_files[key] = path

        context.report_progress("Publishing complete", 100)

        return output_files

    # ── Markdown ──────────────────────────────────────────────

    async def _generate_markdown(
        self,
        blueprint: BookBlueprint,
        output_dir: Path,
        plan: Optional[IllustrationPlan] = None,
    ) -> Path:
        """Generate Markdown file with optional illustrations."""

        content_parts = []

        # Title page
        content_parts.append(f"# {blueprint.title}\n")
        if blueprint.subtitle:
            content_parts.append(f"## {blueprint.subtitle}\n")
        content_parts.append(f"\n*By {blueprint.author}*\n\n")
        content_parts.append("---\n\n")

        # Front matter
        if blueprint.front_matter.dedication:
            content_parts.append("## Dedication\n\n")
            content_parts.append(blueprint.front_matter.dedication + "\n\n")

        if blueprint.front_matter.preface:
            content_parts.append("## Preface\n\n")
            content_parts.append(blueprint.front_matter.preface + "\n\n")

        # Table of Contents
        content_parts.append("## Table of Contents\n\n")
        for part in blueprint.parts:
            content_parts.append(f"### {part.title}\n")
            for chapter in part.chapters:
                content_parts.append(f"- Chapter {chapter.number}: {chapter.title}\n")
                for section in chapter.sections:
                    content_parts.append(f"  - {section.title}\n")
        content_parts.append("\n---\n\n")

        # Main content
        chapter_idx = 0
        for part in blueprint.parts:
            content_parts.append(f"# Part {part.number}: {part.title}\n\n")

            if part.introduction:
                content_parts.append(part.introduction + "\n\n")

            for chapter in part.chapters:
                content_parts.append(f"## Chapter {chapter.number}: {chapter.title}\n\n")

                if chapter.introduction:
                    content_parts.append(chapter.introduction + "\n\n")

                for sec_idx, section in enumerate(chapter.sections):
                    content_parts.append(f"### {section.title}\n\n")
                    content_parts.append(section.content + "\n\n")

                    # Insert illustrations after section content
                    if plan:
                        sec_placements = plan.get_placements_for_section(
                            chapter_idx, sec_idx
                        )
                        for p in sec_placements:
                            content_parts.append(
                                self._markdown_illustration(p)
                            )

                if chapter.summary:
                    content_parts.append("#### Summary\n\n")
                    content_parts.append(chapter.summary + "\n\n")

                if chapter.key_takeaways:
                    content_parts.append("#### Key Takeaways\n\n")
                    for takeaway in chapter.key_takeaways:
                        content_parts.append(f"- {takeaway}\n")
                    content_parts.append("\n")

                chapter_idx += 1

        # Back matter
        if blueprint.back_matter.conclusion:
            content_parts.append("# Conclusion\n\n")
            content_parts.append(blueprint.back_matter.conclusion + "\n\n")

        if blueprint.back_matter.glossary:
            content_parts.append("# Glossary\n\n")
            for term, definition in sorted(blueprint.back_matter.glossary.items()):
                content_parts.append(f"**{term}**: {definition}\n\n")

        # Write file
        filename = self._sanitize_filename(blueprint.title) + ".md"
        filepath = output_dir / filename

        with open(filepath, "w", encoding="utf-8") as f:
            f.write("".join(content_parts))

        return filepath

    def _markdown_illustration(self, placement: ImagePlacement) -> str:
        """Generate markdown for a single illustration."""
        caption = placement.caption or ""
        # Use image_id as reference (actual path resolved at render time)
        img_ref = placement.image_id
        return f"![{caption}]({img_ref})\n\n"

    # ── DOCX ──────────────────────────────────────────────────

    async def _generate_docx(
        self,
        blueprint: BookBlueprint,
        output_dir: Path,
        plan: Optional[IllustrationPlan] = None,
        image_dir: str = "",
    ) -> Path:
        """Generate DOCX file with optional illustrations."""

        # Use DocxIllustratedRenderer for illustrated books
        if plan and plan.placements:
            try:
                from ..renderers.docx_renderer import DocxIllustratedRenderer
                genre = plan.genre if hasattr(plan, "genre") else None
                config = LayoutConfig.for_genre(genre) if genre else None
                renderer = DocxIllustratedRenderer(config, genre or plan.genre)
                filename = self._sanitize_filename(blueprint.title) + ".docx"
                output_path = str(output_dir / filename)
                renderer.render(blueprint, plan, image_dir, output_path)
                return Path(output_path)
            except Exception as e:
                self.logger.warning(
                    f"DocxIllustratedRenderer failed, falling back: {e}"
                )

        try:
            from docx import Document
            from docx.shared import Pt, Inches
            from docx.enum.text import WD_ALIGN_PARAGRAPH
        except ImportError:
            self.logger.error("python-docx not installed, falling back to markdown")
            return await self._generate_markdown(blueprint, output_dir, plan)

        doc = Document()

        # Title page
        title_para = doc.add_paragraph()
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_run = title_para.add_run(blueprint.title)
        title_run.bold = True
        title_run.font.size = Pt(24)

        if blueprint.subtitle:
            subtitle_para = doc.add_paragraph()
            subtitle_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            subtitle_run = subtitle_para.add_run(blueprint.subtitle)
            subtitle_run.font.size = Pt(16)

        author_para = doc.add_paragraph()
        author_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        author_para.add_run(f"By {blueprint.author}")

        doc.add_page_break()

        # Front matter
        if blueprint.front_matter.preface:
            doc.add_heading("Preface", level=1)
            doc.add_paragraph(blueprint.front_matter.preface)
            doc.add_page_break()

        # Main content
        chapter_idx = 0
        for part in blueprint.parts:
            doc.add_heading(f"Part {part.number}: {part.title}", level=1)

            if part.introduction:
                doc.add_paragraph(part.introduction)

            for chapter in part.chapters:
                doc.add_heading(f"Chapter {chapter.number}: {chapter.title}", level=2)

                if chapter.introduction:
                    doc.add_paragraph(chapter.introduction)

                for sec_idx, section in enumerate(chapter.sections):
                    doc.add_heading(section.title, level=3)

                    # Float-top illustrations (before section text)
                    if plan:
                        float_tops = [
                            p for p in plan.get_placements_for_section(chapter_idx, sec_idx)
                            if p.layout_mode == LayoutMode.FLOAT_TOP
                        ]
                        for p in float_tops:
                            self._insert_docx_illustration(doc, p)

                    paragraphs = section.content.split("\n\n")
                    for para_text in paragraphs:
                        if para_text.strip():
                            doc.add_paragraph(para_text.strip())

                    # Non-float illustrations (after section text)
                    if plan:
                        others = [
                            p for p in plan.get_placements_for_section(chapter_idx, sec_idx)
                            if p.layout_mode != LayoutMode.FLOAT_TOP
                        ]
                        for p in others:
                            self._insert_docx_illustration(doc, p)

                if chapter.summary:
                    doc.add_heading("Summary", level=4)
                    doc.add_paragraph(chapter.summary)

                if chapter.key_takeaways:
                    doc.add_heading("Key Takeaways", level=4)
                    for takeaway in chapter.key_takeaways:
                        doc.add_paragraph(takeaway, style='List Bullet')

                chapter_idx += 1

        # Conclusion
        if blueprint.back_matter.conclusion:
            doc.add_heading("Conclusion", level=1)
            doc.add_paragraph(blueprint.back_matter.conclusion)

        # Save
        filename = self._sanitize_filename(blueprint.title) + ".docx"
        filepath = output_dir / filename
        doc.save(filepath)

        return filepath

    def _insert_docx_illustration(self, doc, placement: ImagePlacement):
        """Insert an illustration into a DOCX document."""
        try:
            from docx.shared import Inches
            from docx.enum.text import WD_ALIGN_PARAGRAPH
        except ImportError:
            return

        # Resolve image file path
        image_path = self._resolve_image_path(placement.image_id)
        if not image_path or not os.path.exists(image_path):
            self.logger.warning(f"Image not found: {placement.image_id}")
            return

        size_map = {
            ImageSize.SMALL: 1.5,
            ImageSize.MEDIUM: 3.0,
            ImageSize.LARGE: 4.5,
            ImageSize.FULL: 6.0,
        }
        width_inches = size_map.get(placement.size, 3.0)

        if placement.layout_mode == LayoutMode.FULL_PAGE:
            doc.add_page_break()
            para = doc.add_paragraph()
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            para.add_run().add_picture(image_path, width=Inches(6.0))
            self._add_docx_caption(doc, placement.caption)
            doc.add_page_break()

        elif placement.layout_mode == LayoutMode.GALLERY:
            # Gallery: add images in a table
            self._insert_docx_gallery_image(doc, placement, image_path)

        else:
            # INLINE, FLOAT_TOP, MARGIN
            para = doc.add_paragraph()
            if placement.layout_mode == LayoutMode.MARGIN:
                width_inches = 1.5
            para.add_run().add_picture(image_path, width=Inches(width_inches))
            if placement.caption:
                self._add_docx_caption(doc, placement.caption)

    def _insert_docx_gallery_image(self, doc, placement: ImagePlacement, image_path: str):
        """Insert a gallery image."""
        try:
            from docx.shared import Inches
        except ImportError:
            return
        para = doc.add_paragraph()
        para.add_run().add_picture(image_path, width=Inches(2.5))
        if placement.caption:
            self._add_docx_caption(doc, placement.caption)

    def _add_docx_caption(self, doc, caption: str):
        """Add an italic caption paragraph."""
        if not caption:
            return
        try:
            from docx.shared import Pt
            from docx.enum.text import WD_ALIGN_PARAGRAPH
        except ImportError:
            return
        para = doc.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = para.add_run(caption)
        run.italic = True
        run.font.size = Pt(9)

    def _resolve_image_path(self, image_id: str) -> Optional[str]:
        """Resolve image_id to a file path. If image_id is already a path, use it."""
        if os.path.exists(image_id):
            return image_id
        # Try common upload locations
        for base in ["data/uploads/books", "data/uploads"]:
            for root, _, files in os.walk(base):
                for f in files:
                    if image_id in f:
                        return os.path.join(root, f)
        return None

    # ── PDF ──────────────────────────────────────────────────

    async def _generate_pdf(
        self,
        blueprint: BookBlueprint,
        output_dir: Path,
        plan: Optional[IllustrationPlan] = None,
    ) -> Path:
        """Generate PDF file. Tries DOCX→PDF conversion, falls back to markdown."""
        try:
            docx_path = await self._generate_docx(blueprint, output_dir, plan)
            from docx2pdf import convert
            pdf_path = output_dir / (self._sanitize_filename(blueprint.title) + ".pdf")
            convert(str(docx_path), str(pdf_path))
            return pdf_path
        except Exception:
            md_path = await self._generate_markdown(blueprint, output_dir, plan)
            self.logger.warning("PDF generation via docx2pdf failed. Created markdown fallback.")
            return md_path

    # ── HTML ──────────────────────────────────────────────────

    async def _generate_html(
        self,
        blueprint: BookBlueprint,
        output_dir: Path,
        plan: Optional[IllustrationPlan] = None,
    ) -> Path:
        """Generate HTML file with optional illustrations."""

        html_parts = [
            "<!DOCTYPE html>",
            "<html lang='en'>",
            "<head>",
            "<meta charset='UTF-8'>",
            f"<title>{blueprint.title}</title>",
            "<style>",
            "body { max-width: 800px; margin: 0 auto; padding: 20px; font-family: Georgia, serif; }",
            "h1 { text-align: center; }",
            "h2 { border-bottom: 1px solid #ccc; padding-bottom: 10px; }",
            ".chapter { margin-top: 40px; }",
            ".section { margin-top: 20px; }",
            "figure { margin: 20px 0; text-align: center; }",
            "figure img { max-width: 100%; height: auto; }",
            "figcaption { font-style: italic; color: #666; margin-top: 8px; font-size: 0.9em; }",
            ".full-page { page-break-before: always; page-break-after: always; }",
            ".gallery { display: flex; flex-wrap: wrap; gap: 12px; justify-content: center; }",
            ".gallery figure { flex: 0 0 45%; }",
            "</style>",
            "</head>",
            "<body>",
            f"<h1>{blueprint.title}</h1>",
        ]

        if blueprint.subtitle:
            html_parts.append(
                f"<h2 style='text-align:center;border:none;'>{blueprint.subtitle}</h2>"
            )

        html_parts.append(f"<p style='text-align:center;'>By {blueprint.author}</p>")
        html_parts.append("<hr>")

        # Content
        chapter_idx = 0
        for part in blueprint.parts:
            html_parts.append(f"<h1>Part {part.number}: {part.title}</h1>")

            for chapter in part.chapters:
                html_parts.append("<div class='chapter'>")
                html_parts.append(f"<h2>Chapter {chapter.number}: {chapter.title}</h2>")

                for sec_idx, section in enumerate(chapter.sections):
                    html_parts.append("<div class='section'>")
                    html_parts.append(f"<h3>{section.title}</h3>")

                    # Float-top illustrations
                    if plan:
                        float_tops = [
                            p for p in plan.get_placements_for_section(chapter_idx, sec_idx)
                            if p.layout_mode == LayoutMode.FLOAT_TOP
                        ]
                        for p in float_tops:
                            html_parts.append(self._html_figure(p))

                    paragraphs = section.content.split("\n\n")
                    for para in paragraphs:
                        if para.strip():
                            html_parts.append(f"<p>{para.strip()}</p>")

                    # Other illustrations
                    if plan:
                        others = [
                            p for p in plan.get_placements_for_section(chapter_idx, sec_idx)
                            if p.layout_mode != LayoutMode.FLOAT_TOP
                        ]
                        for p in others:
                            html_parts.append(self._html_figure(p))

                    html_parts.append("</div>")

                html_parts.append("</div>")
                chapter_idx += 1

        html_parts.extend(["</body>", "</html>"])

        filename = self._sanitize_filename(blueprint.title) + ".html"
        filepath = output_dir / filename

        with open(filepath, "w", encoding="utf-8") as f:
            f.write("\n".join(html_parts))

        return filepath

    def _html_figure(self, placement: ImagePlacement) -> str:
        """Generate an HTML <figure> for an illustration."""
        size_map = {
            ImageSize.SMALL: "25%",
            ImageSize.MEDIUM: "50%",
            ImageSize.LARGE: "75%",
            ImageSize.FULL: "100%",
        }
        width = size_map.get(placement.size, "50%")
        css_class = ""
        if placement.layout_mode == LayoutMode.FULL_PAGE:
            css_class = " class='full-page'"
            width = "100%"

        caption_html = ""
        if placement.caption:
            caption_html = f"<figcaption>{placement.caption}</figcaption>"

        img_src = placement.image_id
        resolved = self._resolve_image_path(img_src)
        if resolved:
            img_src = resolved

        return (
            f"<figure{css_class}>"
            f"<img src='{img_src}' alt='{placement.caption}' style='max-width:{width};'>"
            f"{caption_html}"
            f"</figure>"
        )

    # ── EPUB ──────────────────────────────────────────────────

    async def _generate_epub(
        self,
        blueprint: BookBlueprint,
        output_dir: Path,
        plan: Optional[IllustrationPlan] = None,
        genre=None,
    ) -> Path:
        """Generate EPUB file. Delegates to epub_renderer if available."""
        try:
            from ..renderers.epub_renderer import EpubRenderer
            from ..illustration_models import BookGenre
            epub_genre = genre or BookGenre.NON_FICTION
            renderer = EpubRenderer(epub_genre)
            return renderer.render(blueprint, output_dir, plan)
        except ImportError:
            self.logger.warning("ebooklib not installed, falling back to HTML")
            return await self._generate_html(blueprint, output_dir, plan)

    # ── Helpers ──────────────────────────────────────────────

    def _sanitize_filename(self, name: str) -> str:
        """Sanitize filename"""
        name = re.sub(r'[<>:"/\\|?*]', '', name)
        name = name.replace(' ', '_')
        return name[:100]
