"""
DOCX Illustrated Renderer (Sprint K — TIP K4)

Renders illustrated books to .docx with professional layout.
Supports all 5 layout modes, genre-aware page setup,
gallery tables, figure numbering, borders, credits, and alignment.
"""

import logging
import os
from pathlib import Path
from typing import Dict, List, Optional

from ..illustration_models import (
    BookGenre,
    GalleryGroup,
    IllustrationPlan,
    ImagePlacement,
    ImageSize,
    LayoutConfig,
    LayoutMode,
)
from ..models import BookBlueprint

logger = logging.getLogger("BookWriter.DocxRenderer")


class DocxIllustratedRenderer:
    """
    Renders illustrated books to DOCX format.

    Usage:
        config = LayoutConfig.for_genre(BookGenre.CHILDREN)
        renderer = DocxIllustratedRenderer(config)
        path = renderer.render(blueprint, plan, image_dir, output_path)
    """

    # Page dimensions in inches by genre
    PAGE_PRESETS: Dict[BookGenre, dict] = {
        BookGenre.CHILDREN: {"width": 8.5, "height": 8.5, "margin": 0.5},
        BookGenre.PHOTOGRAPHY: {"width": 10.0, "height": 12.0, "margin": 0.5},
        BookGenre.TECHNICAL: {"width": 7.0, "height": 10.0, "margin": 1.0},
        BookGenre.COOKBOOK: {"width": 8.0, "height": 10.0, "margin": 0.75},
        BookGenre.TRAVEL: {"width": 8.5, "height": 11.0, "margin": 0.75},
        BookGenre.ACADEMIC: {"width": 7.0, "height": 10.0, "margin": 1.0},
    }
    DEFAULT_PAGE = {"width": 6.0, "height": 9.0, "margin": 0.75}

    def __init__(self, config: Optional[LayoutConfig] = None, genre: BookGenre = BookGenre.NON_FICTION):
        self.config = config or LayoutConfig.for_genre(genre)
        self.genre = genre
        self.figure_counter: Dict[int, int] = {}  # chapter_index → count

    def render(
        self,
        blueprint: BookBlueprint,
        plan: Optional[IllustrationPlan],
        image_dir: str,
        output_path: str,
    ) -> str:
        """
        Render illustrated book to DOCX.

        Returns the output file path.
        """
        from docx import Document
        from docx.shared import Pt, Inches
        from docx.enum.text import WD_ALIGN_PARAGRAPH

        doc = Document()
        self.figure_counter = {}

        # Setup page size and margins
        self._setup_page(doc)

        # Title page
        self._add_title_page(doc, blueprint)

        # Front matter
        if blueprint.front_matter.preface:
            doc.add_heading("Preface", level=1)
            doc.add_paragraph(blueprint.front_matter.preface)
            doc.add_page_break()

        if blueprint.front_matter.dedication:
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(blueprint.front_matter.dedication)
            run.italic = True
            run.font.size = Pt(14)
            doc.add_page_break()

        # Main content
        chapter_idx = 0
        for part in blueprint.parts:
            doc.add_heading(f"Part {part.number}: {part.title}", level=1)
            if part.introduction:
                doc.add_paragraph(part.introduction)

            for chapter in part.chapters:
                doc.add_heading(
                    f"Chapter {chapter.number}: {chapter.title}", level=2
                )

                if chapter.introduction:
                    doc.add_paragraph(chapter.introduction)

                for sec_idx, section in enumerate(chapter.sections):
                    doc.add_heading(section.title, level=3)

                    # Get placements for this section
                    sec_placements = []
                    if plan:
                        sec_placements = plan.get_placements_for_section(
                            chapter_idx, sec_idx
                        )

                    # Float-top images (before section text)
                    float_tops = [
                        p for p in sec_placements
                        if p.layout_mode == LayoutMode.FLOAT_TOP
                    ]
                    for placement in float_tops:
                        self._insert_illustration(
                            doc, placement, image_dir, chapter_idx
                        )

                    # Section text
                    paragraphs = section.content.split("\n\n")
                    for para_text in paragraphs:
                        if para_text.strip():
                            doc.add_paragraph(para_text.strip())

                    # Non-float images (after section text)
                    others = [
                        p for p in sec_placements
                        if p.layout_mode != LayoutMode.FLOAT_TOP
                    ]
                    for placement in others:
                        self._insert_illustration(
                            doc, placement, image_dir, chapter_idx
                        )

                    # Gallery groups for this section
                    if plan:
                        for gallery in plan.galleries:
                            if gallery.chapter_index == chapter_idx:
                                self._insert_gallery(
                                    doc, gallery, plan, image_dir, chapter_idx
                                )

                if chapter.summary:
                    doc.add_heading("Summary", level=4)
                    doc.add_paragraph(chapter.summary)

                if chapter.key_takeaways:
                    doc.add_heading("Key Takeaways", level=4)
                    for takeaway in chapter.key_takeaways:
                        doc.add_paragraph(takeaway, style="List Bullet")

                chapter_idx += 1

        # Back matter
        if blueprint.back_matter.conclusion:
            doc.add_heading("Conclusion", level=1)
            doc.add_paragraph(blueprint.back_matter.conclusion)

        if blueprint.back_matter.glossary:
            doc.add_heading("Glossary", level=1)
            for term, definition in sorted(blueprint.back_matter.glossary.items()):
                p = doc.add_paragraph()
                run = p.add_run(f"{term}: ")
                run.bold = True
                p.add_run(definition)

        # Save
        Path(output_path).parent.mkdir(parents=True, exist_ok=True)
        doc.save(output_path)
        logger.info(f"DOCX rendered: {output_path}")
        return output_path

    # ── Page Setup ────────────────────────────────────────

    def _setup_page(self, doc):
        """Set page size and margins based on genre."""
        from docx.shared import Inches

        preset = self.PAGE_PRESETS.get(self.genre, self.DEFAULT_PAGE)
        section = doc.sections[0]
        section.page_width = Inches(preset["width"])
        section.page_height = Inches(preset["height"])
        section.left_margin = Inches(preset["margin"])
        section.right_margin = Inches(preset["margin"])
        section.top_margin = Inches(preset["margin"])
        section.bottom_margin = Inches(preset["margin"])

    def _get_content_width(self) -> float:
        """Content width in inches (page width minus margins)."""
        preset = self.PAGE_PRESETS.get(self.genre, self.DEFAULT_PAGE)
        return preset["width"] - 2 * preset["margin"]

    # ── Title Page ────────────────────────────────────────

    def _add_title_page(self, doc, blueprint: BookBlueprint):
        from docx.shared import Pt
        from docx.enum.text import WD_ALIGN_PARAGRAPH

        title_para = doc.add_paragraph()
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = title_para.add_run(blueprint.title)
        run.bold = True
        run.font.size = Pt(28)

        if blueprint.subtitle:
            sub_para = doc.add_paragraph()
            sub_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            sub_run = sub_para.add_run(blueprint.subtitle)
            sub_run.font.size = Pt(16)

        author_para = doc.add_paragraph()
        author_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        author_para.add_run(f"By {blueprint.author}")

        doc.add_page_break()

    # ── Illustration Insertion ────────────────────────────

    def _insert_illustration(
        self, doc, placement: ImagePlacement, image_dir: str, chapter_idx: int
    ):
        """Insert a single illustration using the appropriate layout mode."""
        image_path = self._resolve_image_path(placement.image_id, image_dir)
        if not image_path:
            logger.warning(f"Image not found: {placement.image_id}")
            return

        mode = placement.layout_mode
        if mode == LayoutMode.FULL_PAGE:
            self._insert_full_page(doc, placement, image_path, chapter_idx)
        elif mode == LayoutMode.FLOAT_TOP:
            self._insert_float_top(doc, placement, image_path, chapter_idx)
        elif mode == LayoutMode.MARGIN:
            self._insert_margin(doc, placement, image_path, chapter_idx)
        elif mode == LayoutMode.GALLERY:
            # Gallery items are handled by _insert_gallery; single fallback
            self._insert_inline(doc, placement, image_path, chapter_idx)
        else:
            self._insert_inline(doc, placement, image_path, chapter_idx)

    def _insert_full_page(
        self, doc, placement: ImagePlacement, image_path: str, chapter_idx: int
    ):
        """Full-page image: page break, centered image, caption, page break."""
        from docx.shared import Inches
        from docx.enum.text import WD_ALIGN_PARAGRAPH

        if placement.page_break_before:
            doc.add_page_break()
        else:
            doc.add_page_break()

        para = doc.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = para.add_run()
        run.add_picture(image_path, width=Inches(self._get_content_width()))

        self._add_caption(doc, placement, chapter_idx)
        self._add_credit(doc, placement)

        doc.add_page_break()

    def _insert_inline(
        self, doc, placement: ImagePlacement, image_path: str, chapter_idx: int
    ):
        """Inline image with size and alignment."""
        from docx.shared import Inches
        from docx.enum.text import WD_ALIGN_PARAGRAPH

        width = self._calculate_width(placement.size)

        para = doc.add_paragraph()
        align_map = {
            "left": WD_ALIGN_PARAGRAPH.LEFT,
            "center": WD_ALIGN_PARAGRAPH.CENTER,
            "right": WD_ALIGN_PARAGRAPH.RIGHT,
        }
        para.alignment = align_map.get(placement.alignment, WD_ALIGN_PARAGRAPH.CENTER)

        run = para.add_run()
        run.add_picture(image_path, width=Inches(width))

        self._add_caption(doc, placement, chapter_idx)
        self._add_credit(doc, placement)

    def _insert_float_top(
        self, doc, placement: ImagePlacement, image_path: str, chapter_idx: int
    ):
        """Float-top: large image at start of section."""
        from docx.shared import Inches
        from docx.enum.text import WD_ALIGN_PARAGRAPH

        width = self._calculate_width(ImageSize.LARGE)

        para = doc.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = para.add_run()
        run.add_picture(image_path, width=Inches(width))

        self._add_caption(doc, placement, chapter_idx)
        self._add_credit(doc, placement)

    def _insert_margin(
        self, doc, placement: ImagePlacement, image_path: str, chapter_idx: int
    ):
        """Small margin image (python-docx can't do true float, so small centered)."""
        from docx.shared import Inches
        from docx.enum.text import WD_ALIGN_PARAGRAPH

        para = doc.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        run = para.add_run()
        run.add_picture(image_path, width=Inches(self._calculate_width(ImageSize.SMALL)))

        self._add_caption(doc, placement, chapter_idx)

    def _insert_gallery(
        self, doc, gallery: GalleryGroup, plan: IllustrationPlan,
        image_dir: str, chapter_idx: int
    ):
        """Gallery: multi-column table with images and captions."""
        from docx.shared import Inches, Pt
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.enum.table import WD_TABLE_ALIGNMENT

        image_ids = gallery.image_ids
        if not image_ids:
            return

        cols = min(self.config.gallery_columns, len(image_ids))
        rows = (len(image_ids) + cols - 1) // cols

        # Gallery title
        if gallery.title:
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(gallery.title)
            run.bold = True

        content_width = self._get_content_width()
        cell_width = (content_width - 0.5) / cols  # gap between cells

        table = doc.add_table(rows=rows, cols=cols)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER

        for i, img_id in enumerate(image_ids):
            row_idx = i // cols
            col_idx = i % cols
            cell = table.cell(row_idx, col_idx)

            image_path = self._resolve_image_path(img_id, image_dir)
            if not image_path:
                continue

            # Insert image into cell
            para = cell.paragraphs[0]
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = para.add_run()
            run.add_picture(image_path, width=Inches(cell_width * 0.9))

            # Find placement for caption
            placement = next(
                (p for p in plan.placements if p.image_id == img_id), None
            )
            if placement and placement.caption:
                cap_para = cell.add_paragraph()
                cap_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                cap_run = cap_para.add_run(placement.caption)
                cap_run.italic = True
                cap_run.font.size = Pt(8)

        # Gallery caption
        if gallery.caption:
            cap_p = doc.add_paragraph()
            cap_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            cap_run = cap_p.add_run(gallery.caption)
            cap_run.italic = True

    # ── Caption & Credit ──────────────────────────────────

    def _add_caption(self, doc, placement: ImagePlacement, chapter_idx: int):
        """Add a styled caption paragraph with optional figure numbering."""
        caption = placement.caption
        if not caption:
            return

        from docx.shared import Pt
        from docx.enum.text import WD_ALIGN_PARAGRAPH

        # Figure numbering for technical/academic genres
        if self.genre in (BookGenre.TECHNICAL, BookGenre.ACADEMIC):
            fig_num = self.figure_counter.get(chapter_idx, 0) + 1
            self.figure_counter[chapter_idx] = fig_num
            # Only add prefix if not already present
            if not caption.startswith("Figure "):
                caption = f"Figure {chapter_idx + 1}.{fig_num}: {caption}"

        para = doc.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER

        run = para.add_run(caption)
        run.font.size = Pt(9)

        # Caption style from config
        if self.config.caption_style == "italic":
            run.italic = True
        elif self.config.caption_style == "bold":
            run.bold = True
        # "plain" → no styling

    def _add_credit(self, doc, placement: ImagePlacement):
        """Add photo credit/attribution if present."""
        if not placement.credit:
            return

        from docx.shared import Pt, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH

        para = doc.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = para.add_run(placement.credit)
        run.font.size = Pt(7)
        run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
        run.italic = True

    # ── Width Calculation ─────────────────────────────────

    def _calculate_width(self, size: ImageSize) -> float:
        """Calculate image display width in inches."""
        content_width = self._get_content_width()
        pct = self.config.size_pct.get(size.value, 0.5)
        width = content_width * pct
        return max(1.5, min(width, content_width))

    # ── Image Resolution ──────────────────────────────────

    def _resolve_image_path(self, image_id: str, image_dir: str) -> Optional[str]:
        """Resolve image_id to a file path."""
        # Direct path
        if os.path.isfile(image_id):
            return image_id

        # Check in image_dir
        if image_dir and os.path.isdir(image_dir):
            for fname in os.listdir(image_dir):
                if image_id in fname:
                    full = os.path.join(image_dir, fname)
                    if os.path.isfile(full):
                        return full

        # Check common upload locations
        for base in ["data/uploads/books", "data/uploads"]:
            if os.path.isdir(base):
                for root, _, files in os.walk(base):
                    for f in files:
                        if image_id in f:
                            return os.path.join(root, f)

        return None
