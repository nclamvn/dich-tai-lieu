"""
Batch Worker
AI Publisher Pro - Batch Queue System

Worker functions for processing batch jobs.
"""

import asyncio
import os
from pathlib import Path
from typing import Optional, List, Callable, Dict, Any
from datetime import datetime

from .batch_job import BatchJob, JobStatus


async def process_document_job(
    job: BatchJob,
    translation_service=None,
    max_concurrent: int = 10,
    on_progress: Optional[Callable[[BatchJob], None]] = None
):
    """
    Process a document translation job.

    Steps:
    1. Convert PDF to images (if needed)
    2. Extract text with OCR
    3. Analyze content
    4. Translate using tiered models
    5. Generate output document

    Args:
        job: BatchJob to process
        translation_service: Translation service instance
        max_concurrent: Max concurrent page processing
        on_progress: Progress callback
    """

    # Update status
    job.update_status(JobStatus.PREPARING)

    # =========================================
    # Step 1: Prepare input
    # =========================================

    input_path = Path(job.input_path)

    if job.input_type == "pdf":
        # Convert PDF to images
        image_paths = await pdf_to_images(input_path)
    elif job.input_type == "folder":
        # Get images from folder
        image_paths = get_images_from_folder(input_path)
    else:
        # Single image or list
        image_paths = [str(input_path)]

    job.progress.total_pages = len(image_paths)

    if on_progress:
        on_progress(job)

    # Check if too many pages
    if len(image_paths) > 1000:
        raise ValueError(f"Document too large: {len(image_paths)} pages (max 1000)")

    # =========================================
    # Step 2: Extract text with OCR
    # =========================================

    job.update_status(JobStatus.PROCESSING)

    extracted_texts = []

    # Process in batches for OCR
    batch_size = max_concurrent * 2
    for i in range(0, len(image_paths), batch_size):
        batch = image_paths[i:i + batch_size]

        # OCR extraction (parallel within batch)
        texts = await extract_text_batch(batch)
        extracted_texts.extend(texts)

        # Update progress
        job.progress.current_page = min(i + batch_size, len(image_paths))
        if on_progress:
            on_progress(job)

    # =========================================
    # Step 3: Translate
    # =========================================

    translations = []

    # Use translation service if available
    if translation_service:
        result = await translation_service.translate_document(
            texts=extracted_texts,
            source_lang=job.source_lang,
            target_lang=job.target_lang,
            on_progress=lambda p, done, total: update_translation_progress(
                job, done, total, on_progress
            )
        )
        translations = [r.translated for r in result.pages]

        # Update cost tracking
        job.progress.tokens_used = result.total_tokens
        job.progress.cost_incurred = result.total_cost
    else:
        # Fallback: use simple translation
        translations = await translate_batch(
            texts=extracted_texts,
            source_lang=job.source_lang,
            target_lang=job.target_lang,
            mode=job.translation_mode,
            max_concurrent=max_concurrent,
            on_progress=lambda done: update_translation_progress(
                job, done, len(extracted_texts), on_progress
            )
        )

    job.translations = translations

    # =========================================
    # Step 4: Generate output
    # =========================================

    output_path = await generate_output(
        translations=translations,
        output_dir=job.output_dir,
        output_format=job.output_format,
        job_name=job.name
    )

    job.output_path = output_path
    job.progress.completed_pages = len(translations)
    job.progress.completed_at = datetime.now()

    if on_progress:
        on_progress(job)

    return job


def update_translation_progress(
    job: BatchJob,
    completed: int,
    total: int,
    callback: Optional[Callable]
):
    """Update job progress during translation"""
    job.progress.completed_pages = completed
    job.update_progress(completed_pages=completed)
    if callback:
        callback(job)


# =========================================
# PDF Processing
# =========================================

async def pdf_to_images(
    pdf_path: Path,
    dpi: int = 150,
    output_dir: Optional[Path] = None
) -> List[str]:
    """Convert PDF to images"""
    try:
        from pdf2image import convert_from_path
    except ImportError:
        raise ImportError("pdf2image not installed. Run: pip install pdf2image")

    if output_dir is None:
        output_dir = pdf_path.parent / f"{pdf_path.stem}_pages"

    output_dir.mkdir(parents=True, exist_ok=True)

    # Convert in thread pool to not block
    loop = asyncio.get_event_loop()
    images = await loop.run_in_executor(
        None,
        lambda: convert_from_path(str(pdf_path), dpi=dpi)
    )

    paths = []
    for i, img in enumerate(images):
        path = output_dir / f"page_{i+1:04d}.png"
        img.save(path, "PNG")
        paths.append(str(path))

    return paths


def get_images_from_folder(folder_path: Path) -> List[str]:
    """Get image files from folder"""
    extensions = {".png", ".jpg", ".jpeg", ".tiff", ".bmp", ".gif"}

    images = []
    for ext in extensions:
        images.extend(folder_path.glob(f"*{ext}"))
        images.extend(folder_path.glob(f"*{ext.upper()}"))

    # Sort by name
    images.sort(key=lambda p: p.name)

    return [str(p) for p in images]


# =========================================
# OCR Extraction
# =========================================

async def extract_text_batch(
    image_paths: List[str],
    engine: str = "paddle"
) -> List[str]:
    """Extract text from multiple images"""

    # Try PaddleOCR first
    try:
        from paddleocr import PaddleOCR
        ocr = PaddleOCR(use_angle_cls=True, lang='ch', show_log=False)

        texts = []
        for path in image_paths:
            result = ocr.ocr(path, cls=True)
            if result and result[0]:
                lines = [line[1][0] for line in result[0]]
                texts.append("\n".join(lines))
            else:
                texts.append("")

        return texts

    except ImportError:
        pass

    # Fallback to EasyOCR
    try:
        import easyocr
        reader = easyocr.Reader(['ch_sim', 'en'], gpu=False)

        texts = []
        for path in image_paths:
            result = reader.readtext(path)
            lines = [r[1] for r in result]
            texts.append("\n".join(lines))

        return texts

    except ImportError:
        pass

    # Fallback to Tesseract
    try:
        import pytesseract
        import cv2

        texts = []
        for path in image_paths:
            img = cv2.imread(path)
            text = pytesseract.image_to_string(img, lang='chi_sim+eng')
            texts.append(text)

        return texts

    except ImportError:
        raise ImportError(
            "No OCR engine available. Install one of: "
            "paddleocr, easyocr, or pytesseract"
        )


# =========================================
# Translation
# =========================================

async def translate_batch(
    texts: List[str],
    source_lang: str,
    target_lang: str,
    mode: str = "balanced",
    max_concurrent: int = 10,
    on_progress: Optional[Callable[[int], None]] = None
) -> List[str]:
    """Translate batch of texts using cheap API"""

    # Select model based on mode
    models = {
        "economy": ("gemini", "gemini-1.5-flash"),
        "balanced": ("openai", "gpt-4o-mini"),
        "quality": ("openai", "gpt-4o")
    }

    provider, model = models.get(mode, models["balanced"])

    translations = []
    completed = 0

    # Create semaphore for concurrency control
    semaphore = asyncio.Semaphore(max_concurrent)

    async def translate_one(text: str, index: int) -> tuple:
        async with semaphore:
            translated = await call_translation_api(
                text=text,
                source_lang=source_lang,
                target_lang=target_lang,
                provider=provider,
                model=model
            )
            return index, translated

    # Create all tasks
    tasks = [
        translate_one(text, i)
        for i, text in enumerate(texts)
    ]

    # Process with progress updates
    results = [None] * len(texts)

    for coro in asyncio.as_completed(tasks):
        index, translated = await coro
        results[index] = translated
        completed += 1

        if on_progress:
            on_progress(completed)

    return results


async def call_translation_api(
    text: str,
    source_lang: str,
    target_lang: str,
    provider: str,
    model: str
) -> str:
    """Call translation API"""

    if not text.strip():
        return ""

    system_prompt = f"""You are an expert translator from {source_lang} to {target_lang}.
Translate accurately, preserving formatting and meaning.
Output ONLY the translation, no explanations."""

    if provider == "openai":
        from openai import AsyncOpenAI
        client = AsyncOpenAI()

        response = await client.chat.completions.create(
            model=model,
            messages=[
                {"role": "system", "content": system_prompt},
                {"role": "user", "content": text}
            ],
            temperature=0.3
        )
        return response.choices[0].message.content

    elif provider == "gemini":
        import google.generativeai as genai

        genai.configure(api_key=os.getenv("GOOGLE_API_KEY"))
        model_instance = genai.GenerativeModel(model)

        prompt = f"{system_prompt}\n\nText:\n{text}"
        response = await model_instance.generate_content_async(prompt)
        return response.text

    elif provider == "deepseek":
        from openai import AsyncOpenAI
        client = AsyncOpenAI(
            api_key=os.getenv("DEEPSEEK_API_KEY"),
            base_url="https://api.deepseek.com"
        )

        response = await client.chat.completions.create(
            model=model,
            messages=[
                {"role": "system", "content": system_prompt},
                {"role": "user", "content": text}
            ],
            temperature=0.3
        )
        return response.choices[0].message.content

    else:
        # Default to OpenAI compatible
        return text


# =========================================
# Output Generation
# =========================================

async def generate_output(
    translations: List[str],
    output_dir: str,
    output_format: str,
    job_name: str
) -> str:
    """Generate output document"""

    output_path = Path(output_dir)
    output_path.mkdir(parents=True, exist_ok=True)

    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    filename = f"{job_name}_{timestamp}"

    if output_format == "txt":
        file_path = output_path / f"{filename}.txt"
        with open(file_path, "w", encoding="utf-8") as f:
            for i, text in enumerate(translations):
                f.write(f"=== Page {i+1} ===\n\n")
                f.write(text)
                f.write("\n\n")
        return str(file_path)

    elif output_format == "md":
        file_path = output_path / f"{filename}.md"
        with open(file_path, "w", encoding="utf-8") as f:
            f.write(f"# {job_name}\n\n")
            for i, text in enumerate(translations):
                f.write(f"## Page {i+1}\n\n")
                f.write(text)
                f.write("\n\n---\n\n")
        return str(file_path)

    elif output_format == "docx":
        try:
            from docx import Document
            from docx.shared import Inches, Pt
            from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

            doc = Document()

            # Title
            title = doc.add_heading(job_name, 0)
            title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

            # Content
            for i, text in enumerate(translations):
                doc.add_heading(f"Page {i+1}", level=1)

                for paragraph in text.split("\n\n"):
                    if paragraph.strip():
                        doc.add_paragraph(paragraph.strip())

                doc.add_page_break()

            file_path = output_path / f"{filename}.docx"
            doc.save(str(file_path))
            return str(file_path)

        except ImportError:
            # Fallback to txt
            return await generate_output(
                translations, output_dir, "txt", job_name
            )

    elif output_format == "pdf":
        # Generate PDF using reportlab or similar
        try:
            from reportlab.lib.pagesizes import A4
            from reportlab.pdfgen import canvas
            from reportlab.lib.units import inch
            from reportlab.pdfbase import pdfmetrics
            from reportlab.pdfbase.ttfonts import TTFont

            file_path = output_path / f"{filename}.pdf"
            c = canvas.Canvas(str(file_path), pagesize=A4)
            width, height = A4

            # Try to register a Chinese font
            try:
                pdfmetrics.registerFont(TTFont('SimSun', 'simsun.ttc'))
                font_name = 'SimSun'
            except:
                font_name = 'Helvetica'

            for i, text in enumerate(translations):
                c.setFont(font_name, 12)

                # Draw page number
                c.drawString(72, height - 50, f"Page {i+1}")

                # Draw text
                y = height - 80
                for line in text.split("\n"):
                    if y < 72:
                        c.showPage()
                        y = height - 50
                    c.drawString(72, y, line[:80])  # Truncate long lines
                    y -= 14

                c.showPage()

            c.save()
            return str(file_path)

        except ImportError:
            # Fallback to txt
            return await generate_output(
                translations, output_dir, "txt", job_name
            )

    else:
        # Default to txt
        return await generate_output(
            translations, output_dir, "txt", job_name
        )


# =========================================
# Utility Functions
# =========================================

def estimate_job_time(pages: int, mode: str = "balanced") -> float:
    """Estimate job processing time in seconds"""
    # Based on empirical testing
    time_per_page = {
        "economy": 1.5,   # Gemini is fast
        "balanced": 2.0,  # GPT-4o-mini is fast
        "quality": 4.0    # GPT-4o is slower
    }

    base_time = time_per_page.get(mode, 2.0)

    # OCR time (~0.5s per page)
    ocr_time = pages * 0.5

    # Translation time (with parallelism)
    concurrent = 10
    translation_time = (pages / concurrent) * base_time

    return ocr_time + translation_time


def estimate_job_cost(pages: int, mode: str = "balanced") -> float:
    """Estimate job cost in USD"""
    cost_per_page = {
        "economy": 0.001,
        "balanced": 0.004,
        "quality": 0.05
    }

    return pages * cost_per_page.get(mode, 0.004)
