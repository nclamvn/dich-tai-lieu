# ═══════════════════════════════════════════════════════════════════
# FILE: core/book_writer/pipeline.py
# PURPOSE: 7-agent pipeline with sliding context window optimization
#          Plan → Write (×N with summary chain) → Enrich → Edit → Compile
# REPLACES: Existing pipeline.py (full rewrite)
# ═══════════════════════════════════════════════════════════════════

from __future__ import annotations
import json
import logging
import time
import os
from typing import Callable, Optional, Any

from .prompts import (
    get_analyst_prompt,
    get_architect_prompt,
    get_outliner_prompt,
    get_writer_prompt,
    build_writer_context,
    get_enricher_prompt,
    get_editor_prompt,
    get_summarizer_prompt,
    get_model_config,
    calculate_context_budget,
)

logger = logging.getLogger("book_writer.pipeline")


# ─────────────────────────────────────────────────────────────────
# AI CALL ABSTRACTION
# ─────────────────────────────────────────────────────────────────

async def call_ai(
    system_prompt: str,
    user_message: str,
    model: str = "gpt-4o",
    temperature: float = 0.5,
    max_tokens: int = 4096,
    ai_service=None,
) -> dict:
    """
    Call AI provider. Returns {"text": str, "tokens_in": int, "tokens_out": int}

    Uses UnifiedLLMClient with auto-fallback: OpenAI → Anthropic → DeepSeek → Gemini.
    If ai_service is provided (from main.py), uses that first.
    """
    if ai_service and hasattr(ai_service, 'generate'):
        # Use app's existing AI service
        result = await ai_service.generate(
            system_prompt=system_prompt,
            user_message=user_message,
            model=model,
            temperature=temperature,
            max_tokens=max_tokens,
        )
        return result

    # Use UnifiedLLMClient (OpenAI primary → Anthropic fallback)
    try:
        from ai_providers.unified_client import get_unified_client
        client = get_unified_client()
        messages = [
            {"role": "system", "content": system_prompt},
            {"role": "user", "content": user_message},
        ]
        response = await client.chat(
            messages=messages,
            max_tokens=max_tokens,
            temperature=temperature,
        )
        usage = response.usage
        return {
            "text": response.content,
            "tokens_in": usage.input_tokens if usage else 0,
            "tokens_out": usage.output_tokens if usage else 0,
        }
    except Exception as e:
        logger.error(f"AI call failed: {e}")
        raise


def parse_json_response(text: str) -> dict | list:
    """Extract JSON from AI response, handling markdown code blocks."""
    text = text.strip()
    # Remove markdown code fences
    if text.startswith("```"):
        lines = text.split("\n")
        text = "\n".join(lines[1:])
        if text.endswith("```"):
            text = text[:-3]
        text = text.strip()
    return json.loads(text)


# ─────────────────────────────────────────────────────────────────
# PIPELINE CLASS
# ─────────────────────────────────────────────────────────────────

class BookWriterPipeline:
    """
    7-agent pipeline for writing a complete book.

    Flow:
    1. Analyst: Analyze input → AnalysisReport
    2. Architect: Design structure → BookBlueprint
    3. Outliner: Detail each chapter → ChapterOutlines[]
    [User checkpoint: review & approve]
    4. Writer: Write each chapter (sliding context) → ChapterResults[]
    5. Enricher: Enrich each chapter → EnrichedChapters[]
    6. Editor: Edit for consistency → EditedChapters[]
    7. Publisher: Compile → DOCX/EPUB/PDF files

    Context Window Strategy:
    - Summary chain: After each chapter, create 200-word summary
    - Sliding window: Inject prev chapter + all summaries
    - Token budget: Calculated per call based on model context size
    """

    def __init__(
        self,
        ai_service=None,
        on_progress: Optional[Callable] = None,
        data_dir: str = "data/books",
    ):
        self.ai_service = ai_service
        self.on_progress = on_progress or (lambda *a, **k: None)
        self.data_dir = data_dir
        os.makedirs(data_dir, exist_ok=True)

    # ─── AGENT 1: ANALYST ───────────────────────────────────────

    async def analyze(
        self,
        input_text: str,
        language: str = "vi",
        user_preferences: dict | None = None,
    ) -> dict:
        """Run Analyst agent on user input."""
        logger.info("Agent 1: Analyzing input...")

        config = get_model_config("analyst")
        system = get_analyst_prompt(language)

        user_msg = f"<user_input>\n{input_text}\n</user_input>"
        if user_preferences:
            user_msg += f"\n\n<user_preferences>\n{json.dumps(user_preferences, ensure_ascii=False)}\n</user_preferences>"

        result = await call_ai(
            system_prompt=system,
            user_message=user_msg,
            ai_service=self.ai_service,
            **config,
        )

        analysis = parse_json_response(result["text"])
        logger.info(f"Analysis complete: {analysis.get('input_mode')}, {analysis.get('genre')}")
        return {
            "analysis": analysis,
            "tokens_in": result["tokens_in"],
            "tokens_out": result["tokens_out"],
        }

    # ─── AGENT 2: ARCHITECT ─────────────────────────────────────

    async def architect(
        self,
        analysis: dict,
        input_text: str,
        target_pages: int = 200,
        language: str = "vi",
    ) -> dict:
        """Run Architect agent to design book structure."""
        logger.info("Agent 2: Designing blueprint...")

        genre = analysis.get("genre", "non_fiction")
        config = get_model_config("architect")
        system = get_architect_prompt(genre, language)

        # Estimate words from pages
        words_per_page = 300  # Average
        target_words = target_pages * words_per_page

        user_msg = f"""<analysis_report>
{json.dumps(analysis, ensure_ascii=False, indent=2)}
</analysis_report>

<target>
Target: ~{target_words} từ (~{target_pages} trang)
</target>

<user_input_preview>
{input_text[:5000]}
</user_input_preview>"""

        result = await call_ai(
            system_prompt=system,
            user_message=user_msg,
            ai_service=self.ai_service,
            **config,
        )

        blueprint = parse_json_response(result["text"])
        logger.info(f"Blueprint: {blueprint.get('total_chapters')} chapters, {blueprint.get('total_words')} words")
        return {
            "blueprint": blueprint,
            "tokens_in": result["tokens_in"],
            "tokens_out": result["tokens_out"],
        }

    # ─── AGENT 3: OUTLINER ──────────────────────────────────────

    async def outline(
        self,
        blueprint: dict,
        input_text: str,
        analysis: dict,
        language: str = "vi",
    ) -> dict:
        """Run Outliner agent — may need multiple calls for many chapters."""
        logger.info("Agent 3: Creating detailed outlines...")

        config = get_model_config("outliner")
        system = get_outliner_prompt(language)

        total_chapters = blueprint.get("total_chapters", 15)
        all_outlines = []
        total_in = 0
        total_out = 0

        # Process in batches of 5 chapters to fit output in context
        batch_size = 5
        for batch_start in range(0, total_chapters, batch_size):
            batch_end = min(batch_start + batch_size, total_chapters)
            batch_chapters = blueprint["chapters"][batch_start:batch_end]

            user_msg = f"""<blueprint>
{json.dumps(blueprint, ensure_ascii=False, indent=2)}
</blueprint>

<chapters_to_outline>
Tạo outline CHI TIẾT cho chapters {batch_start+1} đến {batch_end}:
{json.dumps(batch_chapters, ensure_ascii=False, indent=2)}
</chapters_to_outline>

<source_content>
{input_text[:20000]}
</source_content>

<input_mode>
{analysis.get("input_mode", "seeds")}
</input_mode>"""

            result = await call_ai(
                system_prompt=system,
                user_message=user_msg,
                ai_service=self.ai_service,
                **config,
            )

            batch_outlines = parse_json_response(result["text"])
            if isinstance(batch_outlines, dict):
                batch_outlines = [batch_outlines]
            all_outlines.extend(batch_outlines)
            total_in += result["tokens_in"]
            total_out += result["tokens_out"]

            self.on_progress(
                agent="outliner",
                message=f"Outlined chapters {batch_start+1}-{batch_end}",
            )

        logger.info(f"Outlines complete: {len(all_outlines)} chapters")
        return {
            "outlines": all_outlines,
            "tokens_in": total_in,
            "tokens_out": total_out,
        }

    # ─── AGENT 4: WRITER (CORE — Sliding Context) ──────────────

    async def write_chapter(
        self,
        chapter_number: int,
        blueprint: dict,
        outline: dict,
        analysis: dict,
        chapter_summaries: list[dict],
        previous_chapter_text: str,
        source_material: str | None,
        user_model: str | None = None,
    ) -> dict:
        """
        Write a single chapter using sliding context window.

        Context injection strategy:
        1. System: Writer prompt + writing principles
        2. User message layers:
           - Blueprint (compressed) → structure awareness
           - Style guide → voice consistency
           - Summary chain (ALL prev chapters) → memory
           - Previous chapter FULL TEXT → flow continuity
           - Current outline → what to write
           - Source material → user's draft (if any)
           - Character/term sheet → consistency
        """
        config = get_model_config("writer", user_model)
        voice = analysis.get("voice_profile", "")
        genre = analysis.get("genre", "non_fiction")
        language = analysis.get("detected_language", "vi")

        system = get_writer_prompt(voice, genre, language)

        # Build style guide
        style_guide = f"Voice: {voice}\nTone: {analysis.get('tone', '')}\nGenre: {genre}"

        # Compress blueprint for context (only structure, not full details)
        bp_compressed = {
            "title": blueprint.get("title"),
            "total_chapters": blueprint.get("total_chapters"),
            "chapters": [
                {"n": c["chapter_number"], "title": c["title"], "purpose": c.get("purpose", "")}
                for c in blueprint.get("chapters", [])
            ]
        }

        # Character/term sheet
        ref_sheet = ""
        if blueprint.get("characters"):
            ref_sheet = json.dumps(blueprint["characters"], ensure_ascii=False, indent=1)
        elif blueprint.get("terms"):
            ref_sheet = json.dumps(blueprint["terms"], ensure_ascii=False, indent=1)

        user_msg = build_writer_context(
            blueprint_json=json.dumps(bp_compressed, ensure_ascii=False),
            style_guide=style_guide,
            chapter_summaries=chapter_summaries,
            previous_chapter_text=previous_chapter_text[-15000:] if previous_chapter_text else "",
            current_outline_json=json.dumps(outline, ensure_ascii=False, indent=2),
            source_material=source_material[:30000] if source_material else None,
            character_or_term_sheet=ref_sheet,
            chapter_number=chapter_number,
            chapter_title=outline.get("title", f"Chapter {chapter_number}"),
            word_target=outline.get("word_target", 5000),
        )

        result = await call_ai(
            system_prompt=system,
            user_message=user_msg,
            ai_service=self.ai_service,
            **config,
        )

        chapter_text = result["text"].strip()
        word_count = len(chapter_text.split())

        logger.info(
            f"Chapter {chapter_number} written: {word_count} words "
            f"(target: {outline.get('word_target', 5000)})"
        )

        return {
            "chapter_number": chapter_number,
            "title": outline.get("title", f"Chapter {chapter_number}"),
            "status": "written",
            "content": chapter_text,
            "word_count": word_count,
            "tokens_in": result["tokens_in"],
            "tokens_out": result["tokens_out"],
        }

    async def summarize_chapter(self, chapter_text: str, chapter_number: int, title: str) -> str:
        """Create compact summary for context chain."""
        config = get_model_config("summarizer")
        system = get_summarizer_prompt()

        user_msg = f"Chapter {chapter_number}: {title}\n\n{chapter_text[:12000]}"

        result = await call_ai(
            system_prompt=system,
            user_message=user_msg,
            ai_service=self.ai_service,
            **config,
        )
        return result["text"].strip()

    async def write_all_chapters(
        self,
        blueprint: dict,
        outlines: list[dict],
        analysis: dict,
        input_text: str,
        user_model: str | None = None,
        skip_chapters: set[int] | None = None,
    ) -> list[dict]:
        """
        Write all chapters sequentially with sliding context.

        The summary chain grows with each chapter:
        Ch1: [] → write → summary_ch1
        Ch2: [summary_ch1] + ch1_full → write → summary_ch2
        Ch3: [summary_ch1, summary_ch2] + ch2_full → write → summary_ch3
        ...
        """
        chapters = []
        summaries: list[dict] = []
        previous_text = ""
        skip = skip_chapters or set()

        total = len(outlines)
        total_tokens_in = 0
        total_tokens_out = 0

        for i, outline in enumerate(outlines):
            ch_num = outline.get("chapter_number", i + 1)

            if ch_num in skip:
                logger.info(f"Skipping chapter {ch_num} (user edited)")
                continue

            self.on_progress(
                agent="writer",
                chapter=ch_num,
                total=total,
                message=f"Writing chapter {ch_num}/{total}: {outline.get('title', '')}",
            )

            # Extract source material for this chapter from input
            source = self._extract_chapter_source(input_text, outline, analysis)

            result = await self.write_chapter(
                chapter_number=ch_num,
                blueprint=blueprint,
                outline=outline,
                analysis=analysis,
                chapter_summaries=summaries,
                previous_chapter_text=previous_text,
                source_material=source,
                user_model=user_model,
            )

            # Create summary for chain
            summary = await self.summarize_chapter(
                result["content"], ch_num, result["title"]
            )
            result["summary"] = summary

            chapters.append(result)
            total_tokens_in += result["tokens_in"]
            total_tokens_out += result["tokens_out"]

            summaries.append({
                "number": ch_num,
                "title": result["title"],
                "summary": summary,
            })

            # Slide window
            previous_text = result["content"]

            self.on_progress(
                agent="writer",
                chapter=ch_num,
                total=total,
                message=f"Chapter {ch_num} complete: {result['word_count']} words",
                chapters_done=len(chapters),
            )

        return chapters

    def _extract_chapter_source(
        self, input_text: str, outline: dict, analysis: dict
    ) -> str | None:
        """Extract relevant source material for a specific chapter."""
        mode = analysis.get("input_mode", "seeds")

        if mode == "seeds":
            return None  # No draft content

        # For messy_draft/enrich: look for sections marked is_from_user
        user_sections = []
        for section in outline.get("sections", []):
            if section.get("is_from_user") and section.get("source_material"):
                user_sections.append(section["source_material"])

        if user_sections:
            return "\n\n---\n\n".join(user_sections)

        # Fallback: try to find matching content by chapter title keywords
        title = outline.get("title", "")
        if title and len(input_text) > 500:
            # Simple keyword matching
            keywords = [w.lower() for w in title.split() if len(w) > 3]
            paragraphs = input_text.split("\n\n")
            matched = []
            for para in paragraphs:
                para_lower = para.lower()
                if any(kw in para_lower for kw in keywords):
                    matched.append(para)
            if matched:
                return "\n\n".join(matched[:5])  # Max 5 paragraphs

        return None

    # ─── AGENT 5: ENRICHER ─────────────────────────────────────

    async def enrich_chapter(
        self,
        chapter_text: str,
        chapter_number: int,
        outline: dict,
        analysis: dict,
        user_model: str | None = None,
    ) -> dict:
        """Enrich a single chapter with examples, data, details."""
        config = get_model_config("enricher", user_model)
        genre = analysis.get("genre", "non_fiction")
        language = analysis.get("detected_language", "vi")

        system = get_enricher_prompt(genre, language)

        user_msg = f"""<chapter_outline>
{json.dumps(outline, ensure_ascii=False, indent=2)}
</chapter_outline>

<chapter_to_enrich>
{chapter_text}
</chapter_to_enrich>"""

        result = await call_ai(
            system_prompt=system,
            user_message=user_msg,
            ai_service=self.ai_service,
            **config,
        )

        enriched = result["text"].strip()
        original_words = len(chapter_text.split())
        enriched_words = len(enriched.split())

        logger.info(
            f"Chapter {chapter_number} enriched: {original_words} → {enriched_words} words "
            f"(+{enriched_words - original_words})"
        )

        return {
            "enriched_content": enriched,
            "word_count": enriched_words,
            "additions": enriched_words - original_words,
            "tokens_in": result["tokens_in"],
            "tokens_out": result["tokens_out"],
        }

    async def enrich_all_chapters(
        self,
        chapters: list[dict],
        outlines: list[dict],
        analysis: dict,
        user_model: str | None = None,
    ) -> list[dict]:
        """Enrich all chapters."""
        enriched = []
        for i, chapter in enumerate(chapters):
            ch_num = chapter["chapter_number"]
            outline = outlines[i] if i < len(outlines) else {}

            self.on_progress(
                agent="enricher",
                chapter=ch_num,
                total=len(chapters),
                message=f"Enriching chapter {ch_num}",
            )

            result = await self.enrich_chapter(
                chapter["content"], ch_num, outline, analysis, user_model
            )

            enriched.append({
                **chapter,
                "status": "enriched",
                "enriched_content": result["enriched_content"],
                "word_count": result["word_count"],
            })

        return enriched

    # ─── AGENT 6: EDITOR ───────────────────────────────────────

    async def edit_chapter(
        self,
        chapter_text: str,
        chapter_number: int,
        previous_ending: str,
        reference_sheet: str,
        style_guide: str,
    ) -> dict:
        """Edit a single chapter for consistency and quality."""
        config = get_model_config("editor")
        system = get_editor_prompt()

        user_msg = f"""<previous_chapter_ending>
{previous_ending[-2000:] if previous_ending else "(First chapter)"}
</previous_chapter_ending>

<style_guide>
{style_guide}
</style_guide>

<reference_sheet>
{reference_sheet}
</reference_sheet>

<chapter_to_edit>
{chapter_text}
</chapter_to_edit>"""

        result = await call_ai(
            system_prompt=system,
            user_message=user_msg,
            ai_service=self.ai_service,
            **config,
        )

        text = result["text"]

        # Parse edit notes
        edited_text = text
        edit_notes = []
        if "===EDIT_NOTES===" in text:
            parts = text.split("===EDIT_NOTES===")
            edited_text = parts[0].strip()
            if len(parts) > 1:
                try:
                    edit_notes = json.loads(parts[1].strip())
                except json.JSONDecodeError:
                    edit_notes = [{"type": "general", "description": parts[1].strip()}]

        return {
            "edited_content": edited_text,
            "edit_notes": edit_notes,
            "tokens_in": result["tokens_in"],
            "tokens_out": result["tokens_out"],
        }

    async def edit_all_chapters(
        self,
        chapters: list[dict],
        blueprint: dict,
        analysis: dict,
    ) -> list[dict]:
        """Edit all chapters for consistency."""
        edited = []
        prev_ending = ""

        ref_sheet = ""
        if blueprint.get("characters"):
            ref_sheet = json.dumps(blueprint["characters"], ensure_ascii=False)
        elif blueprint.get("terms"):
            ref_sheet = json.dumps(blueprint["terms"], ensure_ascii=False)

        style = f"Voice: {analysis.get('voice_profile', '')}\nTone: {analysis.get('tone', '')}"

        for chapter in chapters:
            ch_num = chapter["chapter_number"]

            self.on_progress(
                agent="editor",
                chapter=ch_num,
                total=len(chapters),
                message=f"Editing chapter {ch_num}",
            )

            # Use best available content
            content = (
                chapter.get("enriched_content") or
                chapter.get("content", "")
            )

            result = await self.edit_chapter(
                content, ch_num, prev_ending, ref_sheet, style
            )

            edited.append({
                **chapter,
                "status": "edited",
                "edited_content": result["edited_content"],
                "edit_notes": result["edit_notes"],
                "final_content": result["edited_content"],
                "word_count": len(result["edited_content"].split()),
            })

            prev_ending = result["edited_content"]

        return edited

    # ─── AGENT 7: PUBLISHER ─────────────────────────────────────

    async def compile_book(
        self,
        book_id: str,
        title: str,
        subtitle: str | None,
        chapters: list[dict],
        output_formats: list[str],
    ) -> list[dict]:
        """Compile chapters into output files (DOCX, EPUB, TXT, MD)."""
        logger.info(f"Compiling book: {title} ({len(chapters)} chapters)")

        output_dir = os.path.join(self.data_dir, book_id, "outputs")
        os.makedirs(output_dir, exist_ok=True)

        output_files = []

        # Always create Markdown (base format)
        md_path = self._compile_markdown(output_dir, title, subtitle, chapters)
        output_files.append({"format": "markdown", "path": md_path, "filename": f"{title}.md"})

        for fmt in output_formats:
            try:
                if fmt == "docx":
                    path = await self._compile_docx(output_dir, title, subtitle, chapters)
                    output_files.append({"format": "docx", "path": path, "filename": f"{title}.docx"})
                elif fmt == "txt":
                    path = self._compile_txt(output_dir, title, subtitle, chapters)
                    output_files.append({"format": "txt", "path": path, "filename": f"{title}.txt"})
                elif fmt == "epub":
                    path = await self._compile_epub(output_dir, title, subtitle, chapters)
                    output_files.append({"format": "epub", "path": path, "filename": f"{title}.epub"})
            except Exception as e:
                logger.error(f"Failed to compile {fmt}: {e}")

        return output_files

    def _compile_markdown(
        self, output_dir: str, title: str, subtitle: str | None, chapters: list[dict]
    ) -> str:
        """Compile to Markdown."""
        lines = [f"# {title}"]
        if subtitle:
            lines.append(f"\n*{subtitle}*")
        lines.append("\n---\n")

        # TOC
        lines.append("## Mục lục\n")
        for ch in chapters:
            num = ch["chapter_number"]
            t = ch.get("title", f"Chapter {num}")
            lines.append(f"{num}. [{t}](#chapter-{num})")
        lines.append("\n---\n")

        # Chapters
        for ch in chapters:
            content = ch.get("final_content") or ch.get("edited_content") or ch.get("enriched_content") or ch.get("content", "")
            lines.append(f"\n<a id='chapter-{ch['chapter_number']}'></a>")
            lines.append(content)
            lines.append("\n---\n")

        path = os.path.join(output_dir, f"{title}.md")
        with open(path, "w", encoding="utf-8") as f:
            f.write("\n".join(lines))
        return path

    def _compile_txt(
        self, output_dir: str, title: str, subtitle: str | None, chapters: list[dict]
    ) -> str:
        """Compile to plain text."""
        lines = [title.upper(), "=" * len(title)]
        if subtitle:
            lines.append(subtitle)
        lines.append("\n")

        for ch in chapters:
            content = ch.get("final_content") or ch.get("content", "")
            lines.append(content)
            lines.append("\n" + "=" * 40 + "\n")

        path = os.path.join(output_dir, f"{title}.txt")
        with open(path, "w", encoding="utf-8") as f:
            f.write("\n".join(lines))
        return path

    async def _compile_docx(
        self, output_dir: str, title: str, subtitle: str | None, chapters: list[dict]
    ) -> str:
        """
        Compile to DOCX using python-docx.

        Professional formatting: TOC-ready headings, page numbers,
        proper typography, section breaks.
        """
        try:
            from docx import Document
            from docx.shared import Pt, Inches, RGBColor
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            from docx.enum.section import WD_ORIENT
        except ImportError:
            logger.warning("python-docx not available, skipping DOCX compilation")
            raise

        doc = Document()

        # Default style
        style = doc.styles['Normal']
        font = style.font
        font.name = 'Georgia'
        font.size = Pt(11)

        # Title page
        title_para = doc.add_paragraph()
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_para.space_before = Pt(120)
        run = title_para.add_run(title)
        run.font.size = Pt(28)
        run.font.bold = True

        if subtitle:
            sub_para = doc.add_paragraph()
            sub_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = sub_para.add_run(subtitle)
            run.font.size = Pt(16)
            run.font.italic = True
            run.font.color.rgb = RGBColor(100, 100, 100)

        doc.add_page_break()

        # TOC placeholder
        toc_para = doc.add_paragraph()
        toc_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = toc_para.add_run("MỤC LỤC")
        run.font.size = Pt(18)
        run.font.bold = True
        doc.add_paragraph()  # Spacing

        for ch in chapters:
            num = ch["chapter_number"]
            ch_title = ch.get("title", f"Chương {num}")
            toc_entry = doc.add_paragraph(f"{num}. {ch_title}")
            toc_entry.style = doc.styles['Normal']

        doc.add_page_break()

        # Chapters
        for ch in chapters:
            content = (
                ch.get("final_content") or
                ch.get("edited_content") or
                ch.get("enriched_content") or
                ch.get("content", "")
            )

            # Chapter heading
            num = ch["chapter_number"]
            ch_title = ch.get("title", f"Chương {num}")
            heading = doc.add_heading(f"Chương {num}: {ch_title}", level=1)

            # Parse content
            lines = content.split("\n")
            for line in lines:
                line = line.strip()
                if not line:
                    continue
                if line.startswith("# "):
                    continue  # Skip markdown h1 (already added as heading)
                elif line.startswith("## "):
                    doc.add_heading(line[3:], level=2)
                elif line.startswith("### "):
                    doc.add_heading(line[4:], level=3)
                elif line in ("***", "---", "* * *"):
                    # Section break
                    sep = doc.add_paragraph()
                    sep.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    run = sep.add_run("* * *")
                    run.font.color.rgb = RGBColor(150, 150, 150)
                else:
                    para = doc.add_paragraph(line)
                    para.paragraph_format.first_line_indent = Inches(0.3)
                    para.paragraph_format.space_after = Pt(6)

            doc.add_page_break()

        path = os.path.join(output_dir, f"{title}.docx")
        doc.save(path)
        logger.info(f"DOCX compiled: {path}")
        return path

    async def _compile_epub(
        self, output_dir: str, title: str, subtitle: str | None, chapters: list[dict]
    ) -> str:
        """Compile to EPUB using ebooklib."""
        try:
            from ebooklib import epub
        except ImportError:
            logger.warning("ebooklib not available, skipping EPUB")
            raise

        book = epub.EpubBook()
        book.set_identifier(f"book-{title.lower().replace(' ', '-')}")
        book.set_title(title)
        book.set_language("vi")

        # CSS
        style = epub.EpubItem(
            uid="style",
            file_name="style/default.css",
            media_type="text/css",
            content=b"""
body { font-family: Georgia, serif; line-height: 1.6; margin: 2em; }
h1 { font-size: 1.8em; margin-top: 2em; }
h2 { font-size: 1.4em; margin-top: 1.5em; }
p { text-indent: 1.5em; margin: 0.5em 0; }
.separator { text-align: center; margin: 1.5em 0; color: #999; }
"""
        )
        book.add_item(style)

        epub_chapters = []
        spine = ['nav']
        toc = []

        for ch in chapters:
            num = ch["chapter_number"]
            ch_title = ch.get("title", f"Chương {num}")
            content = (
                ch.get("final_content") or ch.get("content", "")
            )

            # Convert to HTML
            html_content = self._markdown_to_simple_html(content)

            epub_ch = epub.EpubHtml(
                title=f"Chương {num}: {ch_title}",
                file_name=f"chapter_{num:02d}.xhtml",
                lang="vi",
            )
            epub_ch.content = f"<h1>Chương {num}: {ch_title}</h1>{html_content}"
            epub_ch.add_item(style)

            book.add_item(epub_ch)
            epub_chapters.append(epub_ch)
            spine.append(epub_ch)
            toc.append(epub_ch)

        book.toc = toc
        book.spine = spine
        book.add_item(epub.EpubNcx())
        book.add_item(epub.EpubNav())

        path = os.path.join(output_dir, f"{title}.epub")
        epub.write_epub(path, book)
        logger.info(f"EPUB compiled: {path}")
        return path

    def _markdown_to_simple_html(self, text: str) -> str:
        """Simple markdown-to-HTML for EPUB content."""
        lines = text.split("\n")
        html_parts = []

        for line in lines:
            line = line.strip()
            if not line:
                continue
            if line.startswith("# "):
                continue  # Skip H1 (already in chapter header)
            elif line.startswith("## "):
                html_parts.append(f"<h2>{line[3:]}</h2>")
            elif line.startswith("### "):
                html_parts.append(f"<h3>{line[4:]}</h3>")
            elif line in ("***", "---", "* * *"):
                html_parts.append('<p class="separator">* * *</p>')
            else:
                # Handle basic markdown bold/italic
                import re
                line = re.sub(r'\*\*(.+?)\*\*', r'<strong>\1</strong>', line)
                line = re.sub(r'\*(.+?)\*', r'<em>\1</em>', line)
                html_parts.append(f"<p>{line}</p>")

        return "\n".join(html_parts)

    # ─── FULL PIPELINE ORCHESTRATOR ─────────────────────────────

    async def run_full_pipeline(
        self,
        book_id: str,
        input_text: str,
        language: str = "vi",
        target_pages: int = 200,
        user_model: str | None = None,
        output_formats: list[str] = None,
        user_preferences: dict | None = None,
        save_callback: Callable | None = None,
    ) -> dict:
        """
        Run the complete 7-agent pipeline.

        Returns full pipeline result with all artifacts.
        save_callback is called after each major step to persist state.
        """
        if output_formats is None:
            output_formats = ["docx"]

        start_time = time.time()
        total_tokens_in = 0
        total_tokens_out = 0

        def save(data: dict):
            if save_callback:
                save_callback(data)

        try:
            # ─── PHASE 1: ANALYZE ───
            self.on_progress(agent="analyst", message="Analyzing input...")
            analysis_result = await self.analyze(input_text, language, user_preferences)
            analysis = analysis_result["analysis"]
            total_tokens_in += analysis_result["tokens_in"]
            total_tokens_out += analysis_result["tokens_out"]
            save({"status": "analysis_ready", "analysis": analysis})

            # ─── PHASE 2: ARCHITECT ───
            self.on_progress(agent="architect", message="Designing book structure...")
            arch_result = await self.architect(analysis, input_text, target_pages, language)
            blueprint = arch_result["blueprint"]
            total_tokens_in += arch_result["tokens_in"]
            total_tokens_out += arch_result["tokens_out"]

            # ─── PHASE 3: OUTLINE ───
            self.on_progress(agent="outliner", message="Creating detailed outlines...")
            outline_result = await self.outline(blueprint, input_text, analysis, language)
            outlines = outline_result["outlines"]
            total_tokens_in += outline_result["tokens_in"]
            total_tokens_out += outline_result["tokens_out"]
            save({
                "status": "outline_ready",
                "blueprint": blueprint,
                "outlines": outlines,
            })

            # ─── CHECKPOINT: Wait for approval ───
            # (Service layer handles this — pipeline returns here
            #  and resumes from write_from_outline)

            return {
                "status": "outline_ready",
                "analysis": analysis,
                "blueprint": blueprint,
                "outlines": outlines,
                "tokens_in": total_tokens_in,
                "tokens_out": total_tokens_out,
                "elapsed": time.time() - start_time,
            }

        except Exception as e:
            logger.error(f"Pipeline failed: {e}", exc_info=True)
            save({"status": "failed", "error": str(e)})
            raise

    async def write_from_outline(
        self,
        book_id: str,
        blueprint: dict,
        outlines: list[dict],
        analysis: dict,
        input_text: str,
        user_model: str | None = None,
        output_formats: list[str] = None,
        save_callback: Callable | None = None,
        existing_chapters: list[dict] | None = None,
    ) -> dict:
        """
        Continue pipeline from approved outline → write → enrich → edit → compile.
        """
        if output_formats is None:
            output_formats = ["docx"]

        start_time = time.time()

        def save(data: dict):
            if save_callback:
                save_callback(data)

        try:
            # Skip chapters that user already edited
            skip = set()
            if existing_chapters:
                for ch in existing_chapters:
                    if ch.get("status") == "user_edited":
                        skip.add(ch["chapter_number"])

            # ─── PHASE 4: WRITE ───
            save({"status": "writing"})
            chapters = await self.write_all_chapters(
                blueprint, outlines, analysis, input_text, user_model, skip
            )

            # Merge with existing user-edited chapters
            if existing_chapters:
                ch_map = {ch["chapter_number"]: ch for ch in chapters}
                for existing in existing_chapters:
                    if existing["chapter_number"] in skip:
                        ch_map[existing["chapter_number"]] = existing
                chapters = [ch_map[o.get("chapter_number", i+1)] for i, o in enumerate(outlines) if o.get("chapter_number", i+1) in ch_map]

            save({"status": "enriching", "chapters": chapters})

            # ─── PHASE 5: ENRICH ───
            self.on_progress(agent="enricher", message="Enriching chapters...")
            chapters = await self.enrich_all_chapters(
                chapters, outlines, analysis, user_model
            )
            save({"status": "editing", "chapters": chapters})

            # ─── PHASE 6: EDIT ───
            self.on_progress(agent="editor", message="Editing for consistency...")
            chapters = await self.edit_all_chapters(chapters, blueprint, analysis)
            save({"status": "compiling", "chapters": chapters})

            # ─── PHASE 7: COMPILE ───
            self.on_progress(agent="publisher", message="Compiling output files...")
            title = blueprint.get("title", "Untitled")
            subtitle = blueprint.get("subtitle")
            output_files = await self.compile_book(
                book_id, title, subtitle, chapters, output_formats
            )

            # Final stats
            total_words = sum(ch.get("word_count", 0) for ch in chapters)

            save({
                "status": "complete",
                "chapters": chapters,
                "output_files": output_files,
                "total_words": total_words,
            })

            self.on_progress(
                agent="publisher",
                message=f"Complete! {len(chapters)} chapters, {total_words} words",
            )

            return {
                "status": "complete",
                "chapters": chapters,
                "output_files": output_files,
                "total_words": total_words,
                "elapsed": time.time() - start_time,
            }

        except Exception as e:
            logger.error(f"Write pipeline failed: {e}", exc_info=True)
            save({"status": "failed", "error": str(e)})
            raise
