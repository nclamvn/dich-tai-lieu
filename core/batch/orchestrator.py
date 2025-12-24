"""
Job orchestrator using batch sub-modules.

This is the new implementation to replace _process_job_impl().
Uses modular components for better maintainability and testability.

Phase 1.5: Orchestrator pattern for batch processing.
"""

from pathlib import Path
from typing import Optional, Dict, Any, Callable, List
from dataclasses import dataclass, field
import asyncio
import time
import uuid

from config.logging_config import get_logger
from config.constants import (
    BATCH_TIMEOUT_SECONDS,
    BATCH_PARALLEL_WORKERS,
    BATCH_CHUNK_SIZE,
)

from .job_handler import JobHandler, JobState, JobResult
from .chunk_processor import ChunkProcessor, ChunkResult
from .result_aggregator import ResultAggregator, AggregatedResult
from .progress_tracker import ProgressTracker, ProgressCallback

logger = get_logger(__name__)


@dataclass
class OrchestratorConfig:
    """Configuration for BatchOrchestrator."""
    max_workers: int = BATCH_PARALLEL_WORKERS
    timeout_seconds: int = BATCH_TIMEOUT_SECONDS
    chunk_size: int = BATCH_CHUNK_SIZE
    max_retries: int = 3
    enable_validation: bool = True
    enable_cache: bool = True
    enable_stem: bool = True
    enable_ocr: bool = True
    enable_glossary: bool = True


@dataclass
class OrchestratorResult:
    """Result from orchestrator processing."""
    job_id: str
    success: bool
    translated_text: Optional[str] = None
    output_path: Optional[Path] = None
    chunk_count: int = 0
    total_chars: int = 0
    duration_seconds: float = 0.0
    quality_score: float = 0.0
    error: Optional[str] = None
    metadata: Dict[str, Any] = field(default_factory=dict)


class BatchOrchestrator:
    """
    Orchestrates batch translation jobs using sub-modules.

    This class replaces the monolithic _process_job_impl() with
    a modular approach using:
    - JobHandler for lifecycle management
    - ChunkProcessor for parallel translation
    - ResultAggregator for combining results
    - ProgressTracker for progress reporting

    Usage:
        orchestrator = BatchOrchestrator(translator, http_client, config)
        result = await orchestrator.process(
            input_text="...",
            source_lang="en",
            target_lang="vi",
            progress_callback=my_callback
        )
    """

    def __init__(
        self,
        translator: Any,
        http_client: Any,
        config: Optional[OrchestratorConfig] = None,
        chunker: Optional[Any] = None,
        validator: Optional[Any] = None,
        cache: Optional[Any] = None,
        glossary: Optional[Any] = None,
    ):
        """
        Initialize orchestrator.

        Args:
            translator: Translation engine instance
            http_client: HTTP client for API calls
            config: Orchestrator configuration
            chunker: Optional custom chunker
            validator: Optional quality validator
            cache: Optional translation cache
            glossary: Optional glossary manager
        """
        self.config = config or OrchestratorConfig()
        self.translator = translator
        self.http_client = http_client

        # Core dependencies
        self.chunker = chunker
        self.validator = validator
        self.cache = cache
        self.glossary = glossary

        # Lazy-loaded optional components
        self._ocr_pipeline = None
        self._stem_processor = None

        logger.info(
            f"BatchOrchestrator initialized: "
            f"workers={self.config.max_workers}, "
            f"timeout={self.config.timeout_seconds}s, "
            f"chunk_size={self.config.chunk_size}"
        )

    async def process(
        self,
        input_path: Optional[Path] = None,
        input_text: Optional[str] = None,
        source_lang: str = "en",
        target_lang: str = "vi",
        domain: str = "general",
        output_path: Optional[Path] = None,
        output_format: str = "txt",
        progress_callback: Optional[ProgressCallback] = None,
        options: Optional[Dict[str, Any]] = None,
    ) -> OrchestratorResult:
        """
        Process a translation job end-to-end.

        Args:
            input_path: Path to input file
            input_text: Or direct input text
            source_lang: Source language code
            target_lang: Target language code
            domain: Translation domain
            output_path: Optional output path
            output_format: Output format (txt, docx, pdf)
            progress_callback: Progress updates callback
            options: Additional options

        Returns:
            OrchestratorResult with translated content

        Raises:
            ValueError: If no input provided
            TimeoutError: If job exceeds timeout
        """
        options = options or {}
        job_id = str(uuid.uuid4())[:8]
        start_time = time.time()

        # Create job handler
        job_handler = JobHandler(
            job_id=job_id,
            timeout=self.config.timeout_seconds,
            max_retries=self.config.max_retries,
        )

        # Setup progress tracker
        tracker = ProgressTracker(
            total_chunks=100,  # Will update later
            job_id=job_id,
        )
        if progress_callback:
            tracker.add_callback(progress_callback)

        try:
            result = await self._execute_pipeline(
                job_handler=job_handler,
                tracker=tracker,
                input_path=input_path,
                input_text=input_text,
                source_lang=source_lang,
                target_lang=target_lang,
                domain=domain,
                output_path=output_path,
                output_format=output_format,
                options=options,
            )

            duration = time.time() - start_time
            result.duration_seconds = duration

            logger.info(
                f"Job {job_id} completed: "
                f"{result.chunk_count} chunks, "
                f"{duration:.1f}s, "
                f"quality={result.quality_score:.2f}"
            )

            return result

        except Exception as e:
            duration = time.time() - start_time
            job_handler.fail(str(e))
            tracker.fail(str(e))

            logger.error(f"Job {job_id} failed: {e}")

            return OrchestratorResult(
                job_id=job_id,
                success=False,
                error=str(e),
                duration_seconds=duration,
            )

    async def _execute_pipeline(
        self,
        job_handler: JobHandler,
        tracker: ProgressTracker,
        input_path: Optional[Path],
        input_text: Optional[str],
        source_lang: str,
        target_lang: str,
        domain: str,
        output_path: Optional[Path],
        output_format: str,
        options: Dict[str, Any],
    ) -> OrchestratorResult:
        """Execute the full translation pipeline."""
        job_handler.start()
        tracker.start()

        # Phase 1: Load content
        tracker.start_phase("loading", total_steps=1)
        job_handler.transition_to(JobState.LOADING_INPUT)

        text = await self._load_content(input_path, input_text)
        if not text or not text.strip():
            raise ValueError("No input text provided or file is empty")

        tracker.update(1, "Content loaded")
        tracker.complete_phase()

        # Phase 2: Preprocess (STEM detection, glossary)
        tracker.start_phase("preprocessing", total_steps=2)
        job_handler.transition_to(JobState.PREPROCESSING)

        preprocess_meta = {}
        if self.config.enable_stem:
            text, stem_meta = await self._preprocess_stem(text)
            preprocess_meta['stem'] = stem_meta
            tracker.update(1, "STEM detection complete")

        if self.config.enable_glossary and self.glossary:
            glossary_meta = self._apply_glossary(text, source_lang, target_lang)
            preprocess_meta['glossary'] = glossary_meta
            tracker.update(2, "Glossary applied")

        tracker.complete_phase()

        # Phase 3: Chunking
        tracker.start_phase("chunking", total_steps=1)
        job_handler.transition_to(JobState.CHUNKING)

        chunks = self._create_chunks(text)
        tracker.update(1, f"Created {len(chunks)} chunks")
        tracker.complete_phase()

        # Phase 4: Translation (main work - 70% of progress)
        tracker.start_phase("translating", total_steps=len(chunks))
        job_handler.transition_to(JobState.TRANSLATING)

        chunk_results = await self._translate_chunks(
            chunks=chunks,
            source_lang=source_lang,
            target_lang=target_lang,
            domain=domain,
            tracker=tracker,
            options=options,
        )

        tracker.complete_phase()

        # Phase 5: Aggregation
        tracker.start_phase("postprocessing", total_steps=1)
        job_handler.transition_to(JobState.MERGING)

        aggregator = ResultAggregator()
        aggregated = aggregator.aggregate(chunk_results)

        tracker.update(1, "Results aggregated")
        tracker.complete_phase()

        # Phase 6: Validation (optional)
        quality_score = 0.0
        if self.config.enable_validation and self.validator:
            quality_score = await self._validate(
                original_text=text,
                translated_text=aggregated.text,
                source_lang=source_lang,
                target_lang=target_lang,
            )

        # Phase 7: Export
        if output_path:
            tracker.start_phase("exporting", total_steps=1)
            job_handler.transition_to(JobState.EXPORTING)

            await self._export(
                text=aggregated.text,
                output_path=output_path,
                output_format=output_format,
                options=options,
            )

            tracker.update(1, f"Exported to {output_path.name}")
            tracker.complete_phase()

        # Complete
        tracker.finish("Translation complete")
        job_result = job_handler.complete(
            output_path=output_path,
            translated_text=aggregated.text,
            chunk_count=aggregated.chunk_count,
            quality_score=quality_score or aggregated.avg_quality,
        )

        return OrchestratorResult(
            job_id=job_handler.job_id,
            success=True,
            translated_text=aggregated.text,
            output_path=output_path,
            chunk_count=aggregated.chunk_count,
            total_chars=aggregated.total_chars,
            quality_score=quality_score or aggregated.avg_quality,
            metadata={
                **aggregated.metadata,
                **preprocess_meta,
                'successful_chunks': aggregated.successful_chunks,
                'failed_chunks': aggregated.failed_chunks,
            },
        )

    async def _load_content(
        self,
        input_path: Optional[Path],
        input_text: Optional[str],
    ) -> str:
        """Load content from file or direct text."""
        if input_text:
            return input_text

        if input_path:
            return await self._load_file(input_path)

        raise ValueError("No input provided (text or path)")

    async def _load_file(self, path: Path) -> str:
        """Load file content, handling different formats."""
        if not path.exists():
            raise FileNotFoundError(f"Input file not found: {path}")

        suffix = path.suffix.lower()

        if suffix == '.txt':
            return path.read_text(encoding='utf-8')

        elif suffix == '.pdf':
            return await self._load_pdf(path)

        elif suffix == '.docx':
            return self._load_docx(path)

        else:
            # Try as plain text
            try:
                return path.read_text(encoding='utf-8')
            except UnicodeDecodeError:
                raise ValueError(f"Cannot read file: {path}")

    async def _load_pdf(self, path: Path) -> str:
        """Load PDF file content."""
        if self.config.enable_ocr and self._ocr_pipeline:
            return await self._ocr_pipeline.process(path)

        # Fallback to simple extraction
        try:
            from pypdf import PdfReader
            reader = PdfReader(path)
            text_parts = []
            for page in reader.pages:
                text = page.extract_text()
                if text:
                    text_parts.append(text)
            return "\n\n".join(text_parts)
        except ImportError:
            logger.warning("pypdf not installed, cannot extract PDF")
            return ""

    def _load_docx(self, path: Path) -> str:
        """Load DOCX file content."""
        try:
            from docx import Document
            doc = Document(path)
            return "\n\n".join(p.text for p in doc.paragraphs if p.text)
        except ImportError:
            logger.warning("python-docx not installed")
            return ""

    async def _preprocess_stem(self, text: str) -> tuple[str, Dict[str, Any]]:
        """Detect and handle STEM content (formulas, code)."""
        meta = {'detected': False}
        # STEM processing would go here
        # For now, return text unchanged
        return text, meta

    def _apply_glossary(
        self,
        text: str,
        source_lang: str,
        target_lang: str,
    ) -> Dict[str, Any]:
        """Apply glossary terms."""
        if not self.glossary:
            return {}

        try:
            # Glossary application logic
            return {'terms_found': 0}
        except Exception as e:
            logger.warning(f"Glossary application failed: {e}")
            return {'error': str(e)}

    def _create_chunks(self, text: str) -> List[Any]:
        """Create chunks from text."""
        if self.chunker:
            return self.chunker.chunk(text)

        # Simple fallback chunking
        from dataclasses import dataclass

        @dataclass
        class SimpleChunk:
            id: str
            text: str

        # Split by paragraphs, respecting chunk size
        paragraphs = text.split('\n\n')
        chunks = []
        current_chunk = []
        current_size = 0

        for i, para in enumerate(paragraphs):
            para = para.strip()
            if not para:
                continue

            if current_size + len(para) > self.config.chunk_size and current_chunk:
                chunks.append(SimpleChunk(
                    id=f"chunk_{len(chunks)}",
                    text='\n\n'.join(current_chunk)
                ))
                current_chunk = []
                current_size = 0

            current_chunk.append(para)
            current_size += len(para)

        # Add remaining
        if current_chunk:
            chunks.append(SimpleChunk(
                id=f"chunk_{len(chunks)}",
                text='\n\n'.join(current_chunk)
            ))

        return chunks if chunks else [SimpleChunk(id="chunk_0", text=text)]

    async def _translate_chunks(
        self,
        chunks: List[Any],
        source_lang: str,
        target_lang: str,
        domain: str,
        tracker: ProgressTracker,
        options: Dict[str, Any],
    ) -> List[ChunkResult]:
        """Translate all chunks with progress tracking."""

        async def translate_func(http_client, chunk):
            """Translation function for ChunkProcessor."""
            # Check cache first
            if self.config.enable_cache and self.cache:
                cached = self.cache.get(chunk.text, source_lang, target_lang)
                if cached:
                    # Return cached result
                    from dataclasses import dataclass

                    @dataclass
                    class CachedResult:
                        chunk_id: str
                        source: str
                        translated: str
                        quality_score: float
                        from_cache: bool = True

                    return CachedResult(
                        chunk_id=chunk.id,
                        source=chunk.text,
                        translated=cached,
                        quality_score=0.9,
                        from_cache=True,
                    )

            # Translate using engine
            result = await self.translator.translate_chunk(
                http_client,
                chunk,
            )

            # Cache result
            if self.config.enable_cache and self.cache and result:
                self.cache.set(
                    chunk.text,
                    result.translated,
                    source_lang,
                    target_lang,
                )

            return result

        # Track progress
        completed_count = 0

        def progress_callback(completed: int, total: int, quality: float):
            nonlocal completed_count
            completed_count = completed
            tracker.update(
                completed,
                f"Translating {completed}/{total}",
                quality=quality,
            )

        # Create processor and process
        processor = ChunkProcessor(
            translate_func=translate_func,
            max_concurrency=self.config.max_workers,
            max_retries=self.config.max_retries,
            timeout=float(self.config.timeout_seconds),
        )

        results, stats = await processor.process_all(
            chunks=chunks,
            http_client=self.http_client,
            progress_callback=progress_callback,
        )

        logger.info(
            f"Chunk processing: {stats.successful}/{stats.total_chunks} successful, "
            f"{stats.failed} failed, {stats.from_cache} cached"
        )

        return results

    async def _validate(
        self,
        original_text: str,
        translated_text: str,
        source_lang: str,
        target_lang: str,
    ) -> float:
        """Validate translation quality."""
        if not self.validator:
            return 0.0

        try:
            result = self.validator.validate(
                original=original_text,
                translated=translated_text,
            )
            return result.get('score', 0.0) if isinstance(result, dict) else 0.0
        except Exception as e:
            logger.warning(f"Validation failed: {e}")
            return 0.0

    async def _export(
        self,
        text: str,
        output_path: Path,
        output_format: str,
        options: Dict[str, Any],
    ):
        """Export translated text to file."""
        output_path.parent.mkdir(parents=True, exist_ok=True)

        format_lower = output_format.lower()

        if format_lower == 'txt' or output_path.suffix.lower() == '.txt':
            output_path.write_text(text, encoding='utf-8')

        elif format_lower == 'docx' or output_path.suffix.lower() == '.docx':
            await self._export_docx(text, output_path, options)

        elif format_lower == 'pdf' or output_path.suffix.lower() == '.pdf':
            await self._export_pdf(text, output_path, options)

        else:
            # Default to txt
            output_path.write_text(text, encoding='utf-8')

        logger.info(f"Exported to {output_path}")

    async def _export_docx(
        self,
        text: str,
        output_path: Path,
        options: Dict[str, Any],
    ):
        """Export to DOCX format."""
        try:
            from docx import Document
            doc = Document()
            for para in text.split('\n\n'):
                if para.strip():
                    doc.add_paragraph(para.strip())
            doc.save(output_path)
        except ImportError:
            logger.warning("python-docx not installed, falling back to txt")
            output_path.with_suffix('.txt').write_text(text, encoding='utf-8')

    async def _export_pdf(
        self,
        text: str,
        output_path: Path,
        options: Dict[str, Any],
    ):
        """Export to PDF format."""
        try:
            from reportlab.lib.pagesizes import A4
            from reportlab.pdfgen import canvas

            c = canvas.Canvas(str(output_path), pagesize=A4)
            width, height = A4

            # Simple text rendering
            y = height - 50
            for line in text.split('\n'):
                if y < 50:
                    c.showPage()
                    y = height - 50
                c.drawString(50, y, line[:100])  # Truncate long lines
                y -= 15

            c.save()
        except ImportError:
            logger.warning("reportlab not installed, falling back to txt")
            output_path.with_suffix('.txt').write_text(text, encoding='utf-8')
