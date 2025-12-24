"""
Layout Preserver - Enhanced DOCX Output with Layout Preservation

Creates properly formatted DOCX documents with:
- Chapter/section styling
- Footnote support
- Drop caps
- Custom typography
- Template-based output
"""

import logging
import re
from pathlib import Path
from typing import Optional, List, Dict, Any
from dataclasses import dataclass
from enum import Enum
from html.parser import HTMLParser

logger = logging.getLogger(__name__)

# Check for python-docx
try:
    from docx import Document
    from docx.shared import Pt, Inches, Cm, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
    from docx.enum.style import WD_STYLE_TYPE
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    HAS_PYTHON_DOCX = True
except ImportError:
    HAS_PYTHON_DOCX = False
    logger.warning("python-docx not installed, advanced DOCX features unavailable")


class DocumentType(Enum):
    """Document types for styling"""
    NOVEL = "novel"
    ACADEMIC = "academic"
    BUSINESS = "business"
    TECHNICAL = "technical"
    LEGAL = "legal"


@dataclass
class DocumentStyle:
    """Document styling configuration"""
    doc_type: DocumentType

    # Page setup
    page_width: float = 8.5       # inches
    page_height: float = 11.0     # inches
    margin_top: float = 1.0
    margin_bottom: float = 1.0
    margin_left: float = 1.25
    margin_right: float = 1.25

    # Typography
    body_font: str = "Times New Roman"
    body_size: int = 12
    heading_font: str = "Times New Roman"
    line_spacing: float = 1.5

    # Novel-specific
    chapter_page_break: bool = True
    drop_caps: bool = False
    scene_break_symbol: str = "* * *"

    # Academic-specific
    include_toc: bool = False
    numbered_headings: bool = True

    @classmethod
    def for_novel(cls) -> "DocumentStyle":
        """Style preset for novels"""
        return cls(
            doc_type=DocumentType.NOVEL,
            body_font="Georgia",
            body_size=11,
            heading_font="Georgia",
            line_spacing=1.5,
            chapter_page_break=True,
            drop_caps=True,
            margin_left=1.0,
            margin_right=1.0,
        )

    @classmethod
    def for_academic(cls) -> "DocumentStyle":
        """Style preset for academic papers"""
        return cls(
            doc_type=DocumentType.ACADEMIC,
            body_font="Times New Roman",
            body_size=12,
            heading_font="Times New Roman",
            line_spacing=2.0,  # Double-spaced
            numbered_headings=True,
            include_toc=True,
        )

    @classmethod
    def for_business(cls) -> "DocumentStyle":
        """Style preset for business documents"""
        return cls(
            doc_type=DocumentType.BUSINESS,
            body_font="Calibri",
            body_size=11,
            heading_font="Calibri",
            line_spacing=1.15,
        )


class LayoutPreserver:
    """
    Creates formatted DOCX documents with layout preservation

    Uses python-docx for native DOCX creation instead of Pandoc conversion.
    """

    def __init__(self, style: Optional[DocumentStyle] = None):
        """
        Initialize Layout Preserver

        Args:
            style: Document style configuration
        """
        if not HAS_PYTHON_DOCX:
            raise RuntimeError("python-docx required: pip install python-docx")

        self.style = style or DocumentStyle(doc_type=DocumentType.NOVEL)

    def create_document(
        self,
        content: str,
        output_path: Path,
        title: Optional[str] = None,
        author: Optional[str] = None,
        metadata: Optional[Dict[str, Any]] = None,
    ) -> Path:
        """
        Create formatted DOCX from Markdown content

        Args:
            content: Markdown content with structure
            output_path: Output file path
            title: Document title
            author: Document author
            metadata: Additional metadata

        Returns:
            Path to created DOCX
        """
        doc = Document()

        # Setup document
        self._setup_page(doc)
        self._setup_styles(doc)

        # Add metadata
        if title:
            doc.core_properties.title = title
        if author:
            doc.core_properties.author = author

        # Parse and add content
        self._parse_and_add_content(doc, content)

        # Save
        output_path = Path(output_path)
        doc.save(str(output_path))

        logger.info(f"Created DOCX: {output_path}")
        return output_path

    def _setup_page(self, doc: Document):
        """Configure page layout"""
        section = doc.sections[0]

        section.page_width = Inches(self.style.page_width)
        section.page_height = Inches(self.style.page_height)
        section.top_margin = Inches(self.style.margin_top)
        section.bottom_margin = Inches(self.style.margin_bottom)
        section.left_margin = Inches(self.style.margin_left)
        section.right_margin = Inches(self.style.margin_right)

    def _setup_styles(self, doc: Document):
        """Configure document styles"""
        styles = doc.styles

        # Normal/Body style
        if 'Normal' in styles:
            normal = styles['Normal']
            normal.font.name = self.style.body_font
            normal.font.size = Pt(self.style.body_size)
            normal.paragraph_format.line_spacing = self.style.line_spacing

        # Heading styles
        for i in range(1, 4):
            style_name = f'Heading {i}'
            if style_name in styles:
                heading = styles[style_name]
                heading.font.name = self.style.heading_font
                heading.font.bold = True
                heading.font.size = Pt(self.style.body_size + (4 - i) * 2)

        # Create custom styles
        self._create_custom_styles(doc)

    def _create_custom_styles(self, doc: Document):
        """Create custom paragraph styles"""
        styles = doc.styles
        existing_styles = [s.name for s in styles]

        # Chapter Title style
        if 'ChapterTitle' not in existing_styles:
            try:
                chapter_style = styles.add_style('ChapterTitle', WD_STYLE_TYPE.PARAGRAPH)
                chapter_style.font.name = self.style.heading_font
                chapter_style.font.size = Pt(24)
                chapter_style.font.bold = True
                chapter_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
                chapter_style.paragraph_format.space_before = Pt(72)
                chapter_style.paragraph_format.space_after = Pt(24)
            except Exception as e:
                logger.debug(f"Could not create ChapterTitle style: {e}")

        # Scene Break style
        if 'SceneBreak' not in existing_styles:
            try:
                scene_style = styles.add_style('SceneBreak', WD_STYLE_TYPE.PARAGRAPH)
                scene_style.font.name = self.style.body_font
                scene_style.font.size = Pt(self.style.body_size)
                scene_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
                scene_style.paragraph_format.space_before = Pt(12)
                scene_style.paragraph_format.space_after = Pt(12)
            except Exception as e:
                logger.debug(f"Could not create SceneBreak style: {e}")

        # Block Quote style
        if 'BlockQuote' not in existing_styles:
            try:
                quote_style = styles.add_style('BlockQuote', WD_STYLE_TYPE.PARAGRAPH)
                quote_style.font.name = self.style.body_font
                quote_style.font.size = Pt(self.style.body_size)
                quote_style.font.italic = True
                quote_style.paragraph_format.left_indent = Inches(0.5)
                quote_style.paragraph_format.right_indent = Inches(0.5)
            except Exception as e:
                logger.debug(f"Could not create BlockQuote style: {e}")

    def _parse_and_add_content(self, doc: Document, content: str):
        """Parse Markdown and add to document"""
        lines = content.split('\n')

        i = 0
        in_code_block = False
        code_content = []
        is_first_paragraph_of_chapter = False

        while i < len(lines):
            line = lines[i]

            # Code blocks
            if line.startswith('```'):
                if in_code_block:
                    # End code block
                    self._add_code_block(doc, '\n'.join(code_content))
                    code_content = []
                    in_code_block = False
                else:
                    in_code_block = True
                i += 1
                continue

            if in_code_block:
                code_content.append(line)
                i += 1
                continue

            # Empty line
            if not line.strip():
                i += 1
                continue

            # Chapter heading (# or ## CHAPTER)
            chapter_match = re.match(r'^#{1,2}\s+(.+)', line)
            is_chapter = chapter_match or re.match(r'^(CHAPTER|Chapter|CHƯƠNG|Chương)\s+', line)

            if is_chapter:
                if self.style.chapter_page_break and len(doc.paragraphs) > 0:
                    doc.add_page_break()

                title_text = chapter_match.group(1) if chapter_match else line
                try:
                    p = doc.add_paragraph(title_text, style='ChapterTitle')
                except KeyError:
                    p = doc.add_heading(title_text, level=1)

                is_first_paragraph_of_chapter = True
                i += 1
                continue

            # Section heading (### or ####)
            heading_match = re.match(r'^(#{3,6})\s+(.+)', line)
            if heading_match:
                level = len(heading_match.group(1)) - 2  # ### = Heading 1
                level = min(level, 3)  # Cap at Heading 3
                doc.add_heading(heading_match.group(2), level=level)
                i += 1
                continue

            # Scene break
            if line.strip() in ['***', '* * *', '---', '* * * *', '###']:
                try:
                    p = doc.add_paragraph(self.style.scene_break_symbol, style='SceneBreak')
                except KeyError:
                    p = doc.add_paragraph(self.style.scene_break_symbol)
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                i += 1
                continue

            # Block quote
            if line.startswith('>'):
                quote_lines = [line[1:].strip()]
                i += 1
                while i < len(lines) and lines[i].startswith('>'):
                    quote_lines.append(lines[i][1:].strip())
                    i += 1

                try:
                    p = doc.add_paragraph(' '.join(quote_lines), style='BlockQuote')
                except KeyError:
                    p = doc.add_paragraph(' '.join(quote_lines))
                    p.paragraph_format.left_indent = Inches(0.5)
                continue

            # HTML Table
            if line.strip().startswith('<table'):
                table_html = [line]
                i += 1
                while i < len(lines) and '</table>' not in lines[i-1]:
                    table_html.append(lines[i])
                    i += 1

                self._add_html_table(doc, '\n'.join(table_html))
                continue

            # Markdown Table
            if '|' in line and i + 1 < len(lines) and re.match(r'^\|[\s\-:|]+\|', lines[i + 1]):
                table_lines = [line]
                i += 1
                while i < len(lines) and '|' in lines[i]:
                    table_lines.append(lines[i])
                    i += 1

                self._add_markdown_table(doc, table_lines)
                continue

            # Regular paragraph
            p = doc.add_paragraph()

            # Drop cap for first paragraph of chapter
            if is_first_paragraph_of_chapter and self.style.drop_caps and line.strip():
                self._add_drop_cap(p, line)
                is_first_paragraph_of_chapter = False
            else:
                self._add_formatted_text(p, line)

            i += 1

    def _add_formatted_text(self, paragraph, text: str):
        """Add text with inline formatting"""
        # Process bold, italic, code
        parts = self._parse_inline_formatting(text)

        for part_text, part_format in parts:
            run = paragraph.add_run(part_text)
            if 'bold' in part_format:
                run.bold = True
            if 'italic' in part_format:
                run.italic = True
            if 'code' in part_format:
                run.font.name = 'Courier New'
                run.font.size = Pt(self.style.body_size - 1)

    def _parse_inline_formatting(self, text: str) -> List[tuple]:
        """Parse inline Markdown formatting"""
        result = []

        # Simple regex-based parsing
        # Pattern: **bold**, *italic*, `code`
        pattern = r'(\*\*[^*]+\*\*|\*[^*]+\*|`[^`]+`|[^*`]+)'

        for match in re.finditer(pattern, text):
            segment = match.group(0)

            if segment.startswith('**') and segment.endswith('**'):
                result.append((segment[2:-2], ['bold']))
            elif segment.startswith('*') and segment.endswith('*'):
                result.append((segment[1:-1], ['italic']))
            elif segment.startswith('`') and segment.endswith('`'):
                result.append((segment[1:-1], ['code']))
            else:
                result.append((segment, []))

        return result if result else [(text, [])]

    def _add_drop_cap(self, paragraph, text: str):
        """Add drop cap to paragraph"""
        if not text:
            return

        # First letter as drop cap
        first_letter = text[0]
        rest = text[1:]

        # Add drop cap run
        run = paragraph.add_run(first_letter)
        run.font.size = Pt(self.style.body_size * 3)
        run.font.bold = True

        # Add rest of text
        if rest:
            self._add_formatted_text(paragraph, rest)

    def _add_code_block(self, doc: Document, code: str):
        """Add code block"""
        p = doc.add_paragraph()
        run = p.add_run(code)
        run.font.name = 'Courier New'
        run.font.size = Pt(10)

        # Add shading
        p.paragraph_format.left_indent = Inches(0.25)
        p.paragraph_format.right_indent = Inches(0.25)

    def _add_markdown_table(self, doc: Document, lines: List[str]):
        """Add Markdown table to document"""
        # Parse Markdown table
        rows = []
        for line in lines:
            if re.match(r'^\|[\s\-:|]+\|', line):
                continue  # Skip separator

            cells = [c.strip() for c in line.split('|')[1:-1]]
            if cells:
                rows.append(cells)

        if not rows:
            return

        # Create table
        num_cols = len(rows[0])
        table = doc.add_table(rows=len(rows), cols=num_cols)
        table.style = 'Table Grid'

        for i, row_data in enumerate(rows):
            row = table.rows[i]
            for j, cell_text in enumerate(row_data):
                if j < len(row.cells):
                    row.cells[j].text = cell_text

                    # Header row bold
                    if i == 0:
                        for paragraph in row.cells[j].paragraphs:
                            for run in paragraph.runs:
                                run.bold = True

    def _add_html_table(self, doc: Document, html: str):
        """Add HTML table to document"""

        class TableParser(HTMLParser):
            def __init__(self):
                super().__init__()
                self.rows = []
                self.current_row = []
                self.current_cell = ""
                self.in_cell = False
                self.is_header = False

            def handle_starttag(self, tag, attrs):
                if tag == 'tr':
                    self.current_row = []
                elif tag in ('td', 'th'):
                    self.in_cell = True
                    self.is_header = (tag == 'th')
                    self.current_cell = ""

            def handle_endtag(self, tag):
                if tag == 'tr':
                    if self.current_row:
                        self.rows.append(self.current_row)
                elif tag in ('td', 'th'):
                    self.current_row.append({
                        'text': self.current_cell.strip(),
                        'is_header': self.is_header
                    })
                    self.in_cell = False

            def handle_data(self, data):
                if self.in_cell:
                    self.current_cell += data

        parser = TableParser()
        try:
            parser.feed(html)
        except Exception as e:
            logger.warning(f"Failed to parse HTML table: {e}")
            return

        if not parser.rows:
            return

        # Create table
        num_cols = max(len(row) for row in parser.rows)
        table = doc.add_table(rows=len(parser.rows), cols=num_cols)
        table.style = 'Table Grid'

        for i, row_data in enumerate(parser.rows):
            row = table.rows[i]
            for j, cell_data in enumerate(row_data):
                if j < len(row.cells):
                    row.cells[j].text = cell_data['text']

                    if cell_data['is_header']:
                        for paragraph in row.cells[j].paragraphs:
                            for run in paragraph.runs:
                                run.bold = True


# Convenience function
def create_formatted_docx(
    content: str,
    output_path: str,
    doc_type: str = "novel",
    title: Optional[str] = None,
    author: Optional[str] = None,
) -> Path:
    """
    Create formatted DOCX document

    Args:
        content: Markdown content
        output_path: Output file path
        doc_type: "novel", "academic", or "business"
        title: Document title
        author: Document author

    Returns:
        Path to created DOCX
    """
    if doc_type == "novel":
        style = DocumentStyle.for_novel()
    elif doc_type == "academic":
        style = DocumentStyle.for_academic()
    elif doc_type == "business":
        style = DocumentStyle.for_business()
    else:
        style = DocumentStyle(doc_type=DocumentType.NOVEL)

    preserver = LayoutPreserver(style)
    return preserver.create_document(
        content,
        Path(output_path),
        title=title,
        author=author,
    )
