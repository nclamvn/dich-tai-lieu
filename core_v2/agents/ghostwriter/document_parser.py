#!/usr/bin/env python3
# -*- coding: utf-8 -*-

"""
Phase 5A: Document Intelligence - Document Parser
Handles DOCX, TXT, PDF, MD file parsing for draft uploads
"""

from pathlib import Path
from typing import List, Dict, Any, Optional
import re
from dataclasses import dataclass


@dataclass
class ParsedChapter:
    """Represents a parsed chapter from uploaded document"""
    chapter_number: int
    title: str
    content: str
    word_count: int
    metadata: Dict[str, Any]


@dataclass
class ParsedDocument:
    """Complete parsed document with chapters"""
    filename: str
    total_chapters: int
    total_words: int
    chapters: List[ParsedChapter]
    metadata: Dict[str, Any]


class DocumentParser:
    """
    Parse uploaded documents and extract chapters

    Supports:
    - DOCX (Microsoft Word)
    - TXT (Plain text)
    - PDF (via text extraction)
    - MD (Markdown)
    """

    # Chapter detection patterns
    CHAPTER_PATTERNS = [
        r'^Chapter\s+(\d+)[:\s]*(.*?)$',           # Chapter 1: Title
        r'^Ch\.?\s*(\d+)[:\s]*(.*?)$',             # Ch. 1: Title
        r'^CHAPTER\s+(\d+)[:\s]*(.*?)$',           # CHAPTER 1: TITLE
        r'^Chương\s+(\d+)[:\s]*(.*?)$',            # Chương 1: Tiêu đề
        r'^第(\d+)章[:\s]*(.*?)$',                  # 第1章: 标题
        r'^(\d+)\.\s+(.+)$',                       # 1. Title
        r'^Part\s+(\d+)[:\s]*(.*?)$',              # Part 1: Title
    ]

    def __init__(self):
        self.chapter_regex = [re.compile(pattern, re.IGNORECASE) for pattern in self.CHAPTER_PATTERNS]

    def parse_file(self, file_path: Path) -> ParsedDocument:
        """
        Parse uploaded file and extract chapters

        Args:
            file_path: Path to uploaded file

        Returns:
            ParsedDocument with chapters
        """
        suffix = file_path.suffix.lower()

        if suffix == '.docx':
            return self._parse_docx(file_path)
        elif suffix == '.txt':
            return self._parse_txt(file_path)
        elif suffix == '.pdf':
            return self._parse_pdf(file_path)
        elif suffix == '.md':
            return self._parse_markdown(file_path)
        else:
            raise ValueError(f"Unsupported file format: {suffix}")

    def _parse_docx(self, file_path: Path) -> ParsedDocument:
        """Parse DOCX file"""
        try:
            from docx import Document as DocxDocument
        except ImportError:
            raise ImportError("python-docx not installed. Install with: pip install python-docx")

        doc = DocxDocument(str(file_path))

        # Extract all paragraphs
        lines = []
        for para in doc.paragraphs:
            text = para.text.strip()
            if text:
                lines.append(text)

        return self._split_chapters(lines, file_path.name)

    def _parse_txt(self, file_path: Path) -> ParsedDocument:
        """Parse plain text file"""
        with open(file_path, 'r', encoding='utf-8') as f:
            content = f.read()

        lines = [line.strip() for line in content.split('\n') if line.strip()]
        return self._split_chapters(lines, file_path.name)

    def _parse_pdf(self, file_path: Path) -> ParsedDocument:
        """Parse PDF file using PyMuPDF"""
        try:
            import fitz  # PyMuPDF
        except ImportError:
            raise ImportError("PyMuPDF not installed. Install with: pip install PyMuPDF")

        doc = fitz.open(str(file_path))

        lines = []
        for page in doc:
            text = page.get_text()
            page_lines = [line.strip() for line in text.split('\n') if line.strip()]
            lines.extend(page_lines)

        doc.close()
        return self._split_chapters(lines, file_path.name)

    def _parse_markdown(self, file_path: Path) -> ParsedDocument:
        """Parse Markdown file"""
        with open(file_path, 'r', encoding='utf-8') as f:
            content = f.read()

        # For markdown, also detect # headings as chapter markers
        lines = []
        for line in content.split('\n'):
            line = line.strip()
            if line:
                # Convert markdown headings to chapter format
                if line.startswith('# '):
                    line = line[2:].strip()
                lines.append(line)

        return self._split_chapters(lines, file_path.name)

    def _split_chapters(self, lines: List[str], filename: str) -> ParsedDocument:
        """
        Split lines into chapters based on chapter markers

        Args:
            lines: List of text lines
            filename: Original filename

        Returns:
            ParsedDocument with detected chapters
        """
        chapters = []
        current_chapter_num = 0
        current_chapter_title = ""
        current_chapter_content = []

        for line in lines:
            # Check if line is a chapter marker
            is_chapter_marker = False
            chapter_num = None
            chapter_title = ""

            for regex in self.chapter_regex:
                match = regex.match(line)
                if match:
                    is_chapter_marker = True
                    chapter_num = int(match.group(1))
                    if len(match.groups()) > 1:
                        chapter_title = match.group(2).strip()
                    else:
                        chapter_title = f"Chapter {chapter_num}"
                    break

            if is_chapter_marker and chapter_num:
                # Save previous chapter
                if current_chapter_num > 0 and current_chapter_content:
                    content_text = '\n\n'.join(current_chapter_content)
                    word_count = len(content_text.split())

                    chapters.append(ParsedChapter(
                        chapter_number=current_chapter_num,
                        title=current_chapter_title,
                        content=content_text,
                        word_count=word_count,
                        metadata={}
                    ))

                # Start new chapter
                current_chapter_num = chapter_num
                current_chapter_title = chapter_title
                current_chapter_content = []
            else:
                # Add line to current chapter
                if current_chapter_num > 0:
                    current_chapter_content.append(line)
                else:
                    # If no chapter detected yet, treat as Chapter 1
                    if not current_chapter_num:
                        current_chapter_num = 1
                        current_chapter_title = "Chapter 1"
                    current_chapter_content.append(line)

        # Save last chapter
        if current_chapter_num > 0 and current_chapter_content:
            content_text = '\n\n'.join(current_chapter_content)
            word_count = len(content_text.split())

            chapters.append(ParsedChapter(
                chapter_number=current_chapter_num,
                title=current_chapter_title,
                content=content_text,
                word_count=word_count,
                metadata={}
            ))

        # If no chapters detected, treat entire content as single chapter
        if not chapters:
            content_text = '\n\n'.join(lines)
            word_count = len(content_text.split())

            chapters.append(ParsedChapter(
                chapter_number=1,
                title="Chapter 1",
                content=content_text,
                word_count=word_count,
                metadata={}
            ))

        total_words = sum(ch.word_count for ch in chapters)

        return ParsedDocument(
            filename=filename,
            total_chapters=len(chapters),
            total_words=total_words,
            chapters=chapters,
            metadata={
                'format': Path(filename).suffix[1:] if Path(filename).suffix else 'txt',
                'has_explicit_chapters': len(chapters) > 1
            }
        )
