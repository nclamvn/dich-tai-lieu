"""
Book Export Integration (Phase 4.4)

Integrates Author Mode projects with Book Pipeline (Phase 3.x)
for professional DOCX/PDF export with:
- Character glossary appendix
- Timeline appendix
- Professional formatting
- Typography
"""

from typing import Optional
from pathlib import Path
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH


class BookExporter:
    """
    Export Author Mode projects as professional books

    Combines chapters with memory-based appendices for publication-ready output
    """

    def __init__(self, project):
        """
        Initialize book exporter

        Args:
            project: AuthorProject to export
        """
        self.project = project

    def export_to_docx(
        self,
        output_path: Path,
        include_glossary: bool = True,
        include_timeline: bool = True,
        include_plot_summary: bool = False
    ) -> Path:
        """
        Export project as professional DOCX

        Args:
            output_path: Output file path
            include_glossary: Include character glossary
            include_timeline: Include timeline appendix
            include_plot_summary: Include plot summary

        Returns:
            Path to generated file
        """
        doc = Document()

        # Set document margins
        sections = doc.sections
        for section in sections:
            section.top_margin = Inches(1)
            section.bottom_margin = Inches(1)
            section.left_margin = Inches(1.25)
            section.right_margin = Inches(1.25)

        # Title page
        self._add_title_page(doc)

        # Table of contents placeholder
        self._add_toc_placeholder(doc)

        # Chapters
        self._add_chapters(doc)

        # Appendices
        if self.project.memory:
            if include_glossary and self.project.memory.characters:
                self._add_character_glossary(doc)

            if include_timeline and self.project.memory.timeline:
                self._add_timeline_appendix(doc)

            if include_plot_summary and self.project.memory.plot_points:
                self._add_plot_summary(doc)

        # Save document
        doc.save(str(output_path))
        return output_path

    def _add_title_page(self, doc):
        """Add title page"""
        # Title
        title = doc.add_paragraph()
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_run = title.add_run(self.project.title)
        title_run.bold = True
        title_run.font.size = Pt(24)

        # Spacing
        doc.add_paragraph()

        # Author
        author = doc.add_paragraph()
        author.alignment = WD_ALIGN_PARAGRAPH.CENTER
        author_run = author.add_run(f"by {self.project.author_id}")
        author_run.font.size = Pt(14)

        # Spacing
        doc.add_paragraph()
        doc.add_paragraph()

        # Description
        if self.project.description:
            desc = doc.add_paragraph(self.project.description)
            desc.alignment = WD_ALIGN_PARAGRAPH.CENTER
            desc.runs[0].font.size = Pt(11)
            desc.runs[0].italic = True

        # Page break
        doc.add_page_break()

    def _add_toc_placeholder(self, doc):
        """Add table of contents placeholder"""
        toc_heading = doc.add_heading("Table of Contents", level=1)

        toc_note = doc.add_paragraph(
            "Note: In Microsoft Word, right-click here and select "
            "'Update Field' to generate the table of contents."
        )
        toc_note.runs[0].italic = True
        toc_note.runs[0].font.size = Pt(10)

        doc.add_page_break()

    def _add_chapters(self, doc):
        """Add all chapters"""
        for chapter_num in sorted(self.project.chapters.keys()):
            content = self.project.chapters[chapter_num]

            # Chapter heading
            chapter_title = f"Chapter {chapter_num}"
            if self.project.outline and chapter_num <= len(self.project.outline):
                chapter_title += f": {self.project.outline[chapter_num - 1]}"

            doc.add_heading(chapter_title, level=1)

            # Chapter content
            paragraphs = content.split('\n\n')
            for para in paragraphs:
                if para.strip():
                    p = doc.add_paragraph(para.strip())
                    p.paragraph_format.line_spacing = 1.5
                    p.runs[0].font.size = Pt(12)

            # Page break after chapter
            doc.add_page_break()

    def _add_character_glossary(self, doc):
        """Add character glossary appendix"""
        doc.add_heading("Appendix A: Character Glossary", level=1)

        characters = sorted(
            self.project.memory.list_characters(),
            key=lambda c: c.first_appearance_chapter or 999
        )

        for char in characters:
            # Character name (bold)
            name_para = doc.add_paragraph()
            name_run = name_para.add_run(char.name)
            name_run.bold = True
            name_run.font.size = Pt(12)

            # Description
            if char.description:
                doc.add_paragraph(f"  {char.description}")

            # Details
            details = []
            if char.role:
                details.append(f"Role: {char.role}")
            if char.traits:
                details.append(f"Traits: {', '.join(char.traits)}")
            if char.first_appearance_chapter:
                details.append(f"First appears: Chapter {char.first_appearance_chapter}")

            if details:
                detail_para = doc.add_paragraph(f"  {' | '.join(details)}")
                detail_para.runs[0].font.size = Pt(10)
                detail_para.runs[0].italic = True

            doc.add_paragraph()  # Spacing

        doc.add_page_break()

    def _add_timeline_appendix(self, doc):
        """Add timeline appendix"""
        doc.add_heading("Appendix B: Timeline of Events", level=1)

        events = self.project.memory.get_events()

        current_chapter = None
        for event in events:
            if event.chapter != current_chapter:
                current_chapter = event.chapter
                # Chapter heading
                chapter_heading = doc.add_paragraph()
                chapter_run = chapter_heading.add_run(f"Chapter {current_chapter}")
                chapter_run.bold = True
                chapter_run.font.size = Pt(11)

            # Event
            event_para = doc.add_paragraph(f"  • {event.description}")
            event_para.runs[0].font.size = Pt(11)

            # Details
            if event.participants or event.location:
                details = []
                if event.participants:
                    details.append(f"Participants: {', '.join(event.participants)}")
                if event.location:
                    details.append(f"Location: {event.location}")

                detail_para = doc.add_paragraph(f"    {' | '.join(details)}")
                detail_para.runs[0].font.size = Pt(9)
                detail_para.runs[0].italic = True

        doc.add_page_break()

    def _add_plot_summary(self, doc):
        """Add plot summary appendix"""
        doc.add_heading("Appendix C: Plot Summary", level=1)

        # Active plot threads
        active = [p for p in self.project.memory.plot_points.values() if p.status == "active"]
        if active:
            doc.add_heading("Active Plot Threads", level=2)

            for plot in active:
                plot_para = doc.add_paragraph()
                plot_run = plot_para.add_run(f"{plot.type.title()}: ")
                plot_run.bold = True
                plot_para.add_run(plot.description)

                intro_para = doc.add_paragraph(f"  Introduced: Chapter {plot.first_introduced_chapter}")
                intro_para.runs[0].font.size = Pt(10)
                intro_para.runs[0].italic = True

                doc.add_paragraph()  # Spacing

        # Resolved plot threads
        resolved = [p for p in self.project.memory.plot_points.values() if p.status == "resolved"]
        if resolved:
            doc.add_heading("Resolved Plot Threads", level=2)

            for plot in resolved:
                plot_para = doc.add_paragraph()
                plot_run = plot_para.add_run(f"{plot.type.title()}: ")
                plot_run.bold = True
                plot_para.add_run(plot.description)

                span = f"  Chapters {plot.first_introduced_chapter}"
                if plot.resolution_chapter:
                    span += f"-{plot.resolution_chapter}"

                span_para = doc.add_paragraph(span)
                span_para.runs[0].font.size = Pt(10)
                span_para.runs[0].italic = True

                doc.add_paragraph()  # Spacing

    def export_simple_text(self, output_path: Path) -> Path:
        """
        Export project as plain text file

        Args:
            output_path: Output file path

        Returns:
            Path to generated file
        """
        lines = []

        # Title
        lines.append(self.project.title.upper())
        lines.append("=" * len(self.project.title))
        lines.append(f"by {self.project.author_id}\n")

        if self.project.description:
            lines.append(self.project.description)
            lines.append("")

        lines.append("-" * 70)
        lines.append("")

        # Chapters
        for chapter_num in sorted(self.project.chapters.keys()):
            content = self.project.chapters[chapter_num]

            chapter_title = f"CHAPTER {chapter_num}"
            if self.project.outline and chapter_num <= len(self.project.outline):
                chapter_title += f": {self.project.outline[chapter_num - 1].upper()}"

            lines.append(chapter_title)
            lines.append("-" * len(chapter_title))
            lines.append("")
            lines.append(content)
            lines.append("")
            lines.append("-" * 70)
            lines.append("")

        # Write to file
        with open(output_path, 'w', encoding='utf-8') as f:
            f.write('\n'.join(lines))

        return output_path
