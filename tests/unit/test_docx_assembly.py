"""DOCX document assembly: TOC + running header/footer + page number
(Option A, stage 2b.3)."""

from docx import Document

from core.rendering.document_ast import (
    DocumentAST,
    DocumentMetadata,
    Heading,
    HeadingLevel,
    Paragraph,
    StyleSheet,
)
from core.rendering.docx_adapter import render_book_docx, render_docx_from_ast


def _ast(title="Sách Của Tôi", author="Tác Giả", lang="vi") -> DocumentAST:
    md = DocumentMetadata(title=title, author=author, language=lang)
    ast = DocumentAST(metadata=md, styles=StyleSheet())
    ast.add_block(Heading(level=HeadingLevel.H1, text="Chương 1"))
    ast.add_block(Paragraph(text="Nội dung."))
    ast.add_block(Heading(level=HeadingLevel.H2, text="Mục 1.1"))
    return ast


def test_toc_field_and_localized_heading(tmp_path):
    out = tmp_path / "t.docx"
    render_docx_from_ast(_ast(), out, toc=True)
    d = Document(str(out))
    assert "Mục lục" in [p.text for p in d.paragraphs]  # vi label
    assert "TOC \\o" in d.element.xml  # auto-updating TOC field


def test_toc_english_label(tmp_path):
    out = tmp_path / "te.docx"
    render_docx_from_ast(_ast(lang="en"), out, toc=True)
    assert "Table of Contents" in [p.text for p in Document(str(out)).paragraphs]


def test_header_and_page_number_footer(tmp_path):
    out = tmp_path / "h.docx"
    render_docx_from_ast(_ast(title="Tiêu Đề"), out, header_footer=True)
    sec = Document(str(out)).sections[0]
    assert sec.different_first_page_header_footer is True  # cover left blank
    assert "Tiêu Đề" in sec.header.paragraphs[0].text
    assert "PAGE" in sec.footer.paragraphs[0]._p.xml  # page-number field


def test_no_assembly_by_default(tmp_path):
    out = tmp_path / "n.docx"
    render_docx_from_ast(_ast(), out)
    d = Document(str(out))
    assert "TOC \\o" not in d.element.xml
    hdr = d.sections[0].header
    assert (hdr.paragraphs[0].text if hdr.paragraphs else "") == ""


def test_render_book_docx_full_front_matter(tmp_path):
    out = tmp_path / "b.docx"
    render_book_docx(_ast(title="Cuốn Sách", author="AB"), out)
    d = Document(str(out))
    texts = [p.text for p in d.paragraphs]
    assert "Cuốn Sách" in texts  # title page
    assert "Mục lục" in texts  # TOC heading
    assert "TOC \\o" in d.element.xml  # TOC field
    sec = d.sections[0]
    assert "Cuốn Sách" in sec.header.paragraphs[0].text  # running header
    assert "PAGE" in sec.footer.paragraphs[0]._p.xml  # page number
