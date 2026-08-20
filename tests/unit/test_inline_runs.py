"""Inline text runs — bold / italic / code inside paragraphs (Option A, 2b).

The AST gains an optional ``Paragraph.runs`` overlay (``InlineRun`` spans) that is
fully backward-compatible: ``runs=None`` renders exactly as before. These tests
pin the parser, the serialization round trip, both extraction paths (Markdown +
DOCX), and that all three renderers emit real emphasis.
"""

import zipfile

import pypdf

from core.rendering.ast_serialization import (
    ast_from_dict,
    ast_from_json,
    ast_to_dict,
    ast_to_json,
)
from core.rendering.docx_adapter import _CODE_FONT, render_docx_from_ast
from core.rendering.document_ast import (
    DocumentAST,
    DocumentMetadata,
    InlineRun,
    Paragraph,
    StyleSheet,
)
from core.rendering.document_extractor import (
    _docx_runs,
    extract_docx,
    extract_text,
    parse_inline,
)
from core.rendering.epub_adapter import render_epub_from_ast
from core.rendering.pdf_adapter import render_pdf_from_ast


def _flags(runs):
    return [(r.text, r.bold, r.italic, r.code) for r in runs]


# --------------------------------------------------------------------------- #
# parse_inline
# --------------------------------------------------------------------------- #
def test_plain_text_returns_none():
    assert parse_inline("just plain prose, nothing special") is None


def test_bold_italic_code_asterisk_and_backtick():
    assert _flags(parse_inline("a **b** c")) == [("a ", False, False, False),
                                                 ("b", True, False, False),
                                                 (" c", False, False, False)]
    assert _flags(parse_inline("a *b* c")) == [("a ", False, False, False),
                                               ("b", False, True, False),
                                               (" c", False, False, False)]
    assert _flags(parse_inline("a `b` c")) == [("a ", False, False, False),
                                               ("b", False, False, True),
                                               (" c", False, False, False)]


def test_bold_italic_combined_triple():
    assert _flags(parse_inline("x ***y*** z")) == [("x ", False, False, False),
                                                   ("y", True, True, False),
                                                   (" z", False, False, False)]


def test_underscore_forms():
    assert _flags(parse_inline("__b__ and _i_")) == [("b", True, False, False),
                                                     (" and ", False, False, False),
                                                     ("i", False, True, False)]


def test_underscores_inside_identifiers_are_left_alone():
    # snake_case / file paths must not be mistaken for italics
    assert parse_inline("call func_name and open a_b_c.txt") is None
    assert parse_inline("MAX_RETRIES = 3") is None


def test_code_content_is_literal_no_nested_emphasis():
    runs = parse_inline("run `a*b*c` now")
    assert _flags(runs) == [("run ", False, False, False),
                            ("a*b*c", False, False, True),
                            (" now", False, False, False)]


def test_plaintext_invariant_markers_removed():
    for text in ["a **b** c", "x ***y*** z", "mix **b** and `c` and *i*"]:
        runs = parse_inline(text)
        plain = "".join(r.text for r in runs)
        assert "*" not in plain and "`" not in plain and "__" not in plain


# --------------------------------------------------------------------------- #
# Markdown extraction
# --------------------------------------------------------------------------- #
def _md_ast(tmp_path, text):
    p = tmp_path / "in.md"
    p.write_text(text, encoding="utf-8")
    return extract_text(p)


def test_markdown_paragraph_gets_runs_and_plain_stays_none(tmp_path):
    ast = _md_ast(tmp_path, "Plain lead.\n\nHas **bold**, *italic* and `code`.\n\nPlain tail.")
    paras = [b for b in ast.blocks if isinstance(b, Paragraph)]
    assert paras[0].runs is None
    assert paras[2].runs is None
    runs = paras[1].runs
    assert runs is not None
    assert "".join(r.text for r in runs) == paras[1].text  # plaintext view intact
    assert any(r.bold for r in runs)
    assert any(r.italic for r in runs)
    assert any(r.code for r in runs)


# --------------------------------------------------------------------------- #
# Serialization round trip
# --------------------------------------------------------------------------- #
def _ast_with_runs():
    ast = DocumentAST(metadata=DocumentMetadata(title="T"), styles=StyleSheet())
    ast.add_block(Paragraph(text="plain only"))
    ast.add_block(
        Paragraph(
            text="a bold code",
            runs=[
                InlineRun(text="a "),
                InlineRun(text="bold", bold=True),
                InlineRun(text=" "),
                InlineRun(text="code", code=True),
            ],
        )
    )
    return ast


def test_runs_survive_dict_and_json_round_trip():
    ast = _ast_with_runs()
    for restore, dump in ((ast_from_dict, ast_to_dict), (ast_from_json, ast_to_json)):
        back = restore(dump(ast))
        paras = [b for b in back.blocks if isinstance(b, Paragraph)]
        assert paras[0].runs is None
        assert paras[1].runs == ast.blocks[1].runs


# --------------------------------------------------------------------------- #
# Renderers
# --------------------------------------------------------------------------- #
def test_docx_emits_bold_italic_code_runs(tmp_path):
    ast = DocumentAST(metadata=DocumentMetadata(title="T"), styles=StyleSheet())
    ast.add_block(
        Paragraph(
            text="A bold italic code end",
            runs=[
                InlineRun(text="A "),
                InlineRun(text="bold", bold=True),
                InlineRun(text=" "),
                InlineRun(text="italic", italic=True),
                InlineRun(text=" "),
                InlineRun(text="code", code=True),
                InlineRun(text=" end"),
            ],
        )
    )
    out = tmp_path / "o.docx"
    render_docx_from_ast(ast, out)

    import docx as _docx

    got = {"bold": False, "italic": False, "code": False}
    for para in _docx.Document(str(out)).paragraphs:
        for run in para.runs:
            if run.text == "bold" and run.font.bold:
                got["bold"] = True
            if run.text == "italic" and run.font.italic:
                got["italic"] = True
            if run.text == "code" and (run.font.name or "").lower() == _CODE_FONT.lower():
                got["code"] = True
    assert all(got.values()), got


def test_epub_emits_semantic_tags(tmp_path):
    ast = _ast_with_runs()
    out = tmp_path / "o.epub"
    render_epub_from_ast(ast, out)
    with zipfile.ZipFile(out) as z:
        body = "".join(
            z.read(n).decode("utf-8", "ignore")
            for n in z.namelist()
            if n.endswith((".xhtml", ".html"))
        )
    assert "<strong>bold</strong>" in body
    assert "<code>code</code>" in body


def test_pdf_embeds_emphasis_faces_and_courier(tmp_path):
    ast = DocumentAST(metadata=DocumentMetadata(title="T"), styles=StyleSheet())
    ast.add_block(
        Paragraph(
            text="a bold italic code",
            runs=[
                InlineRun(text="a "),
                InlineRun(text="bold", bold=True),
                InlineRun(text=" "),
                InlineRun(text="italic", italic=True),
                InlineRun(text=" "),
                InlineRun(text="code", code=True),
            ],
        )
    )
    out = tmp_path / "o.pdf"
    render_pdf_from_ast(ast, out)

    base = set()
    for page in pypdf.PdfReader(str(out)).pages:
        for f in (page.get("/Resources") or {}).get("/Font", {}).values():
            obj = f.get_object()
            if obj.get("/BaseFont"):
                base.add(str(obj["/BaseFont"]))
    joined = " ".join(sorted(base))
    assert "Bold" in joined  # <b> resolved to a real bold face
    assert "Italic" in joined  # <i> resolved to a real italic face
    assert "Courier" in joined  # inline code


# --------------------------------------------------------------------------- #
# DOCX extraction (round trip through a real .docx)
# --------------------------------------------------------------------------- #
def test_docx_extraction_reads_runs_back(tmp_path):
    import docx as _docx

    doc = _docx.Document()
    para = doc.add_paragraph()
    para.add_run("normal ")
    r = para.add_run("strong")
    r.bold = True
    para.add_run(" ")
    r = para.add_run("slanted")
    r.italic = True
    para.add_run(" ")
    r = para.add_run("mono")
    r.font.name = _CODE_FONT
    src = tmp_path / "src.docx"
    doc.save(str(src))

    ast = extract_docx(src)
    paras = [b for b in ast.blocks if isinstance(b, Paragraph)]
    runs = paras[0].runs
    assert runs is not None
    assert "".join(x.text for x in runs) == paras[0].text  # invariant vs stripped text
    assert any(x.text == "strong" and x.bold for x in runs)
    assert any(x.text == "slanted" and x.italic for x in runs)
    assert any(x.text == "mono" and x.code for x in runs)


def test_docx_unformatted_paragraph_stays_plain(tmp_path):
    import docx as _docx

    doc = _docx.Document()
    p = doc.add_paragraph("wholly unformatted paragraph text")
    # _docx_runs at the unit level: no emphasis anywhere -> None (stays plain).
    assert _docx_runs(p) is None

    src = tmp_path / "plain.docx"
    doc.save(str(src))
    ast = extract_docx(src)
    para = next(b for b in ast.blocks if isinstance(b, Paragraph))
    assert para.runs is None
