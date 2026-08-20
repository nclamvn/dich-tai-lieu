"""Inline emphasis inside list items, blockquotes and table cells (Option A polish).

Paragraphs already carried bold/italic/code via ``Paragraph.runs``; this extends
the same formatting to the other text-bearing blocks through the shared
``core.rendering.inline.parse_inline``, at render time, with no AST model change.
Plain text (no markers) stays plain — fully backward-compatible.
"""

import zipfile

import pypdf

from core.rendering.docx_adapter import _CODE_FONT, render_docx_from_ast
from core.rendering.document_ast import (
    Blockquote,
    DocumentAST,
    DocumentMetadata,
    ListBlock,
    Paragraph,
    StyleSheet,
    TableBlock,
)
from core.rendering.epub_adapter import render_epub_from_ast
from core.rendering.pdf_adapter import render_pdf_from_ast


def _ast(*blocks):
    ast = DocumentAST(metadata=DocumentMetadata(title="T", language="vi"), styles=StyleSheet())
    for b in blocks:
        ast.add_block(b)
    return ast


def _formatted_ast():
    return _ast(
        Blockquote(text="Trích **đậm** và `code`."),
        ListBlock(items=["mục **đậm**", "mục *nghiêng*", "mục `code`"], ordered=False),
        TableBlock(rows=[["Cột **A**", "Cột B"], ["ô `x`", "ô *y*"]], header_rows=1),
    )


def _all_docx_runs(doc):
    runs = []
    for p in doc.paragraphs:
        runs += [(r.text, bool(r.font.bold), bool(r.font.italic), r.font.name or "") for r in p.runs]
    for t in doc.tables:
        for row in t.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    runs += [
                        (r.text, bool(r.font.bold), bool(r.font.italic), r.font.name or "")
                        for r in p.runs
                    ]
    return runs


# --------------------------------------------------------------------------- #
# DOCX
# --------------------------------------------------------------------------- #
def test_docx_inline_in_list_quote_table(tmp_path):
    out = tmp_path / "o.docx"
    render_docx_from_ast(_formatted_ast(), out)

    import docx as _docx

    runs = _all_docx_runs(_docx.Document(str(out)))
    # bold appears (from blockquote/list/table markers)
    assert any(t == "đậm" and b for t, b, i, f in runs)
    # italic appears
    assert any(t == "nghiêng" and i for t, b, i, f in runs)
    # code font applied to code spans (list + table cell)
    assert any(t == "code" and f.lower() == _CODE_FONT.lower() for t, b, i, f in runs)
    assert any(t == "x" and f.lower() == _CODE_FONT.lower() for t, b, i, f in runs)
    # header cell "Cột **A**" -> "A" run is bold (emphasis) within a bolded header
    assert any(t == "A" and b for t, b, i, f in runs)


def test_docx_plain_blocks_have_no_emphasis(tmp_path):
    # List + table have non-emphasized base styles, so any emphasis here would
    # come from the inline machinery. (Blockquotes are intentionally italic by
    # style, so they are excluded from this "no extra emphasis" check.)
    out = tmp_path / "plain.docx"
    render_docx_from_ast(
        _ast(
            ListBlock(items=["một", "hai"], ordered=True),
            TableBlock(rows=[["a", "b"]], header_rows=0),
        ),
        out,
    )
    import docx as _docx

    runs = _all_docx_runs(_docx.Document(str(out)))
    # nothing was marked up -> no italic run and no code-font run
    assert not any(i for t, b, i, f in runs)
    assert not any(f.lower() == _CODE_FONT.lower() for t, b, i, f in runs)


# --------------------------------------------------------------------------- #
# EPUB
# --------------------------------------------------------------------------- #
def _epub_body(path):
    with zipfile.ZipFile(path) as z:
        return "".join(
            z.read(n).decode("utf-8", "ignore")
            for n in z.namelist()
            if n.endswith((".xhtml", ".html"))
        )


def test_epub_inline_in_list_quote_table(tmp_path):
    out = tmp_path / "o.epub"
    render_epub_from_ast(_formatted_ast(), out)
    body = _epub_body(out)
    # blockquote
    assert "<blockquote>" in body
    assert "<strong>đậm</strong>" in body
    assert "<code>code</code>" in body
    # list item emphasis
    assert "<li>mục <strong>đậm</strong></li>" in body
    assert "<li>mục <em>nghiêng</em></li>" in body
    # table header cell keeps its emphasis span
    assert "Cột <strong>A</strong>" in body


# --------------------------------------------------------------------------- #
# PDF
# --------------------------------------------------------------------------- #
def test_pdf_inline_in_blocks_embeds_emphasis_faces(tmp_path):
    out = tmp_path / "o.pdf"
    render_pdf_from_ast(_formatted_ast(), out)  # no template -> default serif

    base = set()
    for page in pypdf.PdfReader(str(out)).pages:
        for f in (page.get("/Resources") or {}).get("/Font", {}).values():
            obj = f.get_object()
            if obj.get("/BaseFont"):
                base.add(str(obj["/BaseFont"]))
    joined = " ".join(sorted(base))
    assert "Bold" in joined
    assert "Italic" in joined
    assert "Courier" in joined


def test_pdf_plain_blocks_have_no_emphasis_faces(tmp_path):
    out = tmp_path / "plain.pdf"
    render_pdf_from_ast(
        _ast(
            Paragraph(text="đoạn thường"),
            ListBlock(items=["một", "hai"], ordered=False),
        ),
        out,
    )
    base = set()
    for page in pypdf.PdfReader(str(out)).pages:
        for f in (page.get("/Resources") or {}).get("/Font", {}).values():
            obj = f.get_object()
            if obj.get("/BaseFont"):
                base.add(str(obj["/BaseFont"]))
    joined = " ".join(sorted(base))
    # nothing marked up -> no bold/italic face, no Courier
    assert "Bold" not in joined
    assert "Italic" not in joined
    assert "Courier" not in joined
