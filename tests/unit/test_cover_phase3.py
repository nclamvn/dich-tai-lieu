"""Phase 3 covers: user-supplied cover images (PDF/DOCX/EPUB) and EPUB covers.

No API key, deterministic. A rendered template image stands in for an uploaded
cover so the tests need no external asset.
"""

import asyncio

import docx
import pypdf
import pytest
from ebooklib import epub as ebook

from core.rendering.cover_apply import (
    apply_cover,
    apply_cover_image_to_docx,
    apply_cover_image_to_pdf,
)
from core.rendering.cover_templates import render_cover_image
from core.rendering.docx_adapter import render_docx_from_ast
from core.rendering.document_ast import (
    DocumentAST,
    DocumentMetadata,
    Heading,
    HeadingLevel,
    Paragraph,
    StyleSheet,
)
from core.rendering.epub_adapter import render_epub_from_ast
from core.rendering.pdf_adapter import render_pdf_from_ast


def _meta():
    return DocumentMetadata(title="BÌNH MINH XANH", author="Trần Văn Bút", language="vi")


def _ast():
    ast = DocumentAST(metadata=_meta(), styles=StyleSheet())
    ast.add_block(Heading(level=HeadingLevel.H1, text="Chương một"))
    ast.add_block(Paragraph(text="Nội dung chương một."))
    ast.add_block(Heading(level=HeadingLevel.H1, text="Chương hai"))
    ast.add_block(Paragraph(text="Nội dung chương hai."))
    return ast


def _has_cover(epub_path) -> bool:
    book = ebook.read_epub(str(epub_path))
    return any("cover" in (it.get_name() or "").lower() for it in book.get_items())


@pytest.fixture
def user_image(tmp_path):
    """A stand-in for a user-uploaded cover image."""
    p = tmp_path / "user_cover.png"
    render_cover_image("gradient", _meta(), p)
    return p


# ---- custom image covers: PDF / DOCX -------------------------------------- #
def test_image_cover_pdf_prepends_page(tmp_path, user_image):
    p = tmp_path / "book.pdf"
    render_pdf_from_ast(_ast(), p)
    n0 = len(pypdf.PdfReader(str(p)).pages)
    assert apply_cover_image_to_pdf(p, user_image) is True
    assert len(pypdf.PdfReader(str(p)).pages) == n0 + 1


def test_image_cover_pdf_missing_image_is_noop(tmp_path):
    p = tmp_path / "book.pdf"
    render_pdf_from_ast(_ast(), p)
    assert apply_cover_image_to_pdf(p, tmp_path / "absent.png") is False


def test_image_cover_docx_adds_section(tmp_path, user_image):
    d = tmp_path / "book.docx"
    render_docx_from_ast(_ast(), d)
    s0 = len(docx.Document(str(d)).sections)
    assert apply_cover_image_to_docx(d, user_image) is True
    dd = docx.Document(str(d))
    assert len(dd.sections) == s0 + 1
    assert "blip" in dd.paragraphs[0]._p.xml


def test_dispatcher_image_wins_over_template(tmp_path, user_image):
    p = tmp_path / "book.pdf"
    render_pdf_from_ast(_ast(), p)
    n0 = len(pypdf.PdfReader(str(p)).pages)
    # Both given; the user image takes precedence and the cover is applied.
    assert apply_cover(p, "pdf", cover_template="noir", cover_image=str(user_image), title="X") is True
    assert len(pypdf.PdfReader(str(p)).pages) == n0 + 1


# ---- EPUB covers (baked at build time) ------------------------------------ #
def test_epub_cover_image_is_baked_in(tmp_path, user_image):
    e = tmp_path / "book.epub"
    render_epub_from_ast(_ast(), e, "BÌNH MINH XANH", cover_image=user_image)
    assert e.exists() and _has_cover(e)


def test_epub_has_no_cover_by_default(tmp_path):
    e = tmp_path / "book.epub"
    render_epub_from_ast(_ast(), e, "BÌNH MINH XANH")
    assert e.exists() and not _has_cover(e)


def test_converter_epub_professional_template_and_image(tmp_path, user_image):
    from core_v2.output_converter import OutputConverter

    conv = OutputConverter(temp_dir=tmp_path / "oc")
    md = "# Chương một\n\nNội dung.\n\n# Chương hai\n\nHai."

    async def build(name, **kw):
        out = tmp_path / f"{name}.epub"
        await conv.convert_markdown_to_epub_professional(
            md, out, title="BÌNH MINH XANH", author="Trần Văn Bút", language="vi", **kw
        )
        return out

    e_tpl = asyncio.run(build("tpl", cover_template="emblem"))
    e_img = asyncio.run(build("img", cover_image=str(user_image)))
    assert _has_cover(e_tpl)
    assert _has_cover(e_img)


# ---- professional PDF/DOCX accept a custom cover image -------------------- #
@pytest.mark.parametrize("fmt", ["pdf", "docx"])
def test_converter_professional_accepts_cover_image(tmp_path, user_image, fmt):
    from core_v2.output_converter import OutputConverter

    conv = OutputConverter(temp_dir=tmp_path / "oc2")
    md = "# Chương một\n\nNội dung.\n\n# Hai\n\nX."

    async def run(cover_image):
        out = tmp_path / f"{'cover' if cover_image else 'plain'}.{fmt}"
        fn = (
            conv.convert_markdown_to_pdf_professional
            if fmt == "pdf"
            else conv.convert_markdown_to_docx_professional
        )
        await fn(md, out, title="BÌNH MINH XANH", author="Trần Văn Bút",
                 language="vi", cover_image=cover_image)
        return out

    plain = asyncio.run(run(None))
    covered = asyncio.run(run(str(user_image)))
    if fmt == "pdf":
        assert len(pypdf.PdfReader(str(covered)).pages) == len(pypdf.PdfReader(str(plain)).pages) + 1
    else:
        assert len(docx.Document(str(covered)).sections) == len(docx.Document(str(plain)).sections) + 1
