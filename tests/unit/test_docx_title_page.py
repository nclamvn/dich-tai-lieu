"""DOCX title page / front matter (Option A, stage 2b.2)."""

from docx import Document

from core.rendering.document_ast import (
    DocumentAST,
    DocumentMetadata,
    Heading,
    HeadingLevel,
    Paragraph,
    StyleSheet,
)
from core.rendering.docx_adapter import render_book_docx, render_docx_from_ast


def _ast(title="Cuốn Sách", author="Tác Giả") -> DocumentAST:
    ast = DocumentAST(metadata=DocumentMetadata(title=title, author=author), styles=StyleSheet())
    ast.add_block(Heading(level=HeadingLevel.H1, text="Chương 1"))
    ast.add_block(Paragraph(text="Nội dung chương."))
    return ast


def _texts(path):
    return [p.text for p in Document(str(path)).paragraphs]


def test_title_page_title_and_author_before_content(tmp_path):
    out = tmp_path / "tp.docx"
    render_docx_from_ast(_ast(), out, title_page=True)
    texts = _texts(out)
    assert "Cuốn Sách" in texts and "Tác Giả" in texts
    assert texts.index("Cuốn Sách") < texts.index("Chương 1")  # cover precedes content


def test_title_page_inserts_page_break(tmp_path):
    out = tmp_path / "tp.docx"
    render_docx_from_ast(_ast(), out, title_page=True)
    assert 'w:type="page"' in Document(str(out)).element.xml


def test_no_title_page_by_default(tmp_path):
    out = tmp_path / "n.docx"
    render_docx_from_ast(_ast(), out)  # title_page defaults False
    texts = _texts(out)
    assert "Tác Giả" not in texts  # author lives in core properties, not the body
    assert texts[0] == "Chương 1"  # content starts immediately, no cover


def test_render_book_docx_has_title_page(tmp_path):
    out = tmp_path / "b.docx"
    render_book_docx(_ast(title="Sách", author="AB"), out)
    texts = _texts(out)
    assert "Sách" in texts and "AB" in texts


def test_title_override_wins_on_cover(tmp_path):
    out = tmp_path / "o.docx"
    render_docx_from_ast(_ast(title="Meta Title"), out, title="Override Title", title_page=True)
    texts = _texts(out)
    assert "Override Title" in texts and "Meta Title" not in texts
