"""Carry-bytes: figure image bytes flow extraction → AST → every renderer.

Closes the figure round trip — a figure with carried ``image_bytes`` embeds the
real picture in DOCX/PDF/EPUB (not a placeholder), the bytes survive the AST
dict/JSON round trip, and a figure *without* bytes still renders a placeholder.
"""

import io

from core.rendering.ast_serialization import (
    ast_from_dict,
    ast_from_json,
    ast_to_dict,
    ast_to_json,
)
from core.rendering.document_ast import (
    DocumentAST,
    DocumentMetadata,
    Figure,
    Heading,
    HeadingLevel,
    Paragraph,
    StyleSheet,
)
from core.rendering.document_extractor import extract_to_ast
from core.rendering.docx_adapter import render_docx_from_ast
from core.rendering.epub_adapter import render_epub_from_ast
from core.rendering.pdf_adapter import render_pdf_from_ast


def _png_bytes(color=(220, 30, 30)) -> bytes:
    from PIL import Image as PILImage

    buf = io.BytesIO()
    PILImage.new("RGB", (16, 12), color).save(buf, format="PNG")
    return buf.getvalue()


def _ast_with_figure(png: bytes) -> DocumentAST:
    ast = DocumentAST(metadata=DocumentMetadata(title="Doc"), styles=StyleSheet())
    ast.add_block(Heading(level=HeadingLevel.H1, text="Chương 1"))
    ast.add_block(Paragraph(text="Trước hình."))
    ast.add_block(
        Figure(
            image_ref="/word/media/image1.png",  # unreadable ref: bytes must win
            image_bytes=png,
            content_type="image/png",
            caption="Chú thích hình",
        )
    )
    return ast


def _make_docx_with_image(tmp_path, png: bytes):
    from docx import Document as DocxDocument

    doc = DocxDocument()
    doc.add_paragraph("Đoạn trước hình.")
    img = tmp_path / "pic.png"
    img.write_bytes(png)
    doc.add_picture(str(img))
    src = tmp_path / "in.docx"
    doc.save(str(src))
    return src


# --------------------------------------------------------------------------- #
# Serialization round trip
# --------------------------------------------------------------------------- #
def test_bytes_survive_dict_and_json_round_trip():
    png = _png_bytes()
    ast = _ast_with_figure(png)

    assert ast_from_dict(ast_to_dict(ast)) == ast
    assert ast_from_json(ast_to_json(ast)) == ast

    rebuilt = ast_from_json(ast_to_json(ast))
    fig = next(b for b in rebuilt.blocks if isinstance(b, Figure))
    assert fig.image_bytes == png
    assert fig.content_type == "image/png"


def test_json_encodes_bytes_as_base64_tag():
    # Without the base64 tag, json.dumps(bytes) raises TypeError.
    text = ast_to_json(_ast_with_figure(_png_bytes()))
    assert "__bytes__" in text


# --------------------------------------------------------------------------- #
# Renderers embed the real image
# --------------------------------------------------------------------------- #
def test_docx_embeds_real_image(tmp_path):
    out = tmp_path / "out.docx"
    render_docx_from_ast(_ast_with_figure(_png_bytes()), out)

    from docx import Document as DocxDocument

    d = DocxDocument(str(out))
    assert len(d.inline_shapes) >= 1  # embedded picture, not a text placeholder


def test_pdf_embeds_real_image(tmp_path):
    out = tmp_path / "out.pdf"
    render_pdf_from_ast(_ast_with_figure(_png_bytes()), out)

    raw = out.read_bytes()
    assert b"/Image" in raw  # an image XObject exists (absent in the placeholder path)


def test_epub_embeds_real_image(tmp_path):
    import ebooklib
    from ebooklib import epub

    png = _png_bytes()
    out = tmp_path / "out.epub"
    render_epub_from_ast(_ast_with_figure(png), out)

    book = epub.read_epub(str(out))
    imgs = [i for i in book.get_items() if i.get_type() == ebooklib.ITEM_IMAGE]
    assert imgs and any(i.get_content() == png for i in imgs)

    docs = [i for i in book.get_items() if i.get_type() == ebooklib.ITEM_DOCUMENT]
    xhtml = " ".join(d.get_content().decode("utf-8") for d in docs)
    assert "<img" in xhtml and "images/fig_1" in xhtml


def test_figure_without_bytes_still_placeholder(tmp_path):
    import ebooklib
    from ebooklib import epub

    ast = DocumentAST(metadata=DocumentMetadata(title="D"), styles=StyleSheet())
    ast.add_block(Figure(image_ref="/word/media/x.png", caption="Chú thích"))
    out = tmp_path / "placeholder.epub"
    render_epub_from_ast(ast, out)

    book = epub.read_epub(str(out))
    docs = [i for i in book.get_items() if i.get_type() == ebooklib.ITEM_DOCUMENT]
    xhtml = " ".join(d.get_content().decode("utf-8") for d in docs)
    assert "[Figure:" in xhtml
    assert not [i for i in book.get_items() if i.get_type() == ebooklib.ITEM_IMAGE]


# --------------------------------------------------------------------------- #
# Extraction carries bytes, and the whole pipeline preserves the image
# --------------------------------------------------------------------------- #
def test_extraction_carries_bytes_from_docx(tmp_path):
    png = _png_bytes()
    src = _make_docx_with_image(tmp_path, png)

    ast = extract_to_ast(src)
    figs = [b for b in ast.blocks if isinstance(b, Figure)]
    assert figs, "a figure should be extracted from the DOCX"
    fig = figs[0]
    assert fig.image_bytes, "image bytes should be carried through extraction"
    assert fig.content_type and fig.content_type.startswith("image/")
    assert fig.image_bytes == png  # exact bytes preserved


def test_end_to_end_docx_to_epub_keeps_image(tmp_path):
    import ebooklib
    from ebooklib import epub

    png = _png_bytes(color=(20, 120, 220))
    src = _make_docx_with_image(tmp_path, png)

    ast = extract_to_ast(src)
    out = tmp_path / "out.epub"
    render_epub_from_ast(ast, out)

    book = epub.read_epub(str(out))
    imgs = [i for i in book.get_items() if i.get_type() == ebooklib.ITEM_IMAGE]
    assert imgs and any(i.get_content() == png for i in imgs)
