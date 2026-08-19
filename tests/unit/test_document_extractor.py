"""L0 phase 2: structure-preserving extract_to_ast, tested on synthetic docs."""

import pytest

from core.rendering.ast_serialization import ast_from_dict, ast_to_dict
from core.rendering.document_ast import Figure, Heading, ListBlock, TableBlock
from core.rendering.document_extractor import extract_to_ast


def _tiny_png(path):
    from PIL import Image

    Image.new("RGB", (4, 4), (200, 100, 50)).save(str(path))


def _make_docx(path, image_path):
    from docx import Document
    from docx.shared import Inches

    doc = Document()
    doc.add_heading("Chương 1", level=1)
    doc.add_paragraph("Đoạn văn thường.")
    doc.add_paragraph("Mục một", style="List Bullet")
    doc.add_paragraph("Mục hai", style="List Bullet")
    table = doc.add_table(rows=2, cols=2)
    table.rows[0].cells[0].text = "A"
    table.rows[0].cells[1].text = "B"
    table.rows[1].cells[0].text = "1"
    table.rows[1].cells[1].text = "2"
    doc.add_picture(str(image_path), width=Inches(1))
    doc.save(str(path))


@pytest.fixture
def sample_docx(tmp_path):
    image = tmp_path / "px.png"
    _tiny_png(image)
    docx = tmp_path / "sample.docx"
    _make_docx(docx, image)
    return docx


def test_docx_block_types_and_counts(sample_docx):
    stats = extract_to_ast(sample_docx).get_statistics()
    assert stats["headings"] >= 1
    assert stats["tables"] == 1
    assert stats["lists"] >= 1  # consecutive bullets grouped into one list
    assert stats["figures"] == 1


def test_docx_content_fidelity(sample_docx):
    ast = extract_to_ast(sample_docx)

    heading = next(b for b in ast.blocks if isinstance(b, Heading))
    assert "Chương 1" in heading.text

    table = next(b for b in ast.blocks if isinstance(b, TableBlock))
    assert table.rows[0] == ["A", "B"]
    assert table.rows[1] == ["1", "2"]

    lst = next(b for b in ast.blocks if isinstance(b, ListBlock))
    assert lst.items == ["Mục một", "Mục hai"]
    assert lst.ordered is False


def test_docx_preserves_reading_order(sample_docx):
    types = [type(b).__name__ for b in extract_to_ast(sample_docx).blocks]
    assert types.index("Heading") < types.index("TableBlock") < types.index("Figure")


def test_extracted_ast_round_trips(sample_docx):
    ast = extract_to_ast(sample_docx)
    assert ast_from_dict(ast_to_dict(ast)) == ast


def test_markdown_extraction(tmp_path):
    md = tmp_path / "doc.md"
    md.write_text("# Tiêu đề\n\nMột đoạn.\n\n- a\n- b\n", encoding="utf-8")
    stats = extract_to_ast(md).get_statistics()
    assert stats["headings"] == 1
    assert stats["paragraphs"] == 1
    assert stats["lists"] == 1


def test_unsupported_format_raises(tmp_path):
    pdf = tmp_path / "x.pdf"
    pdf.write_bytes(b"%PDF-1.4\n")
    with pytest.raises(NotImplementedError):
        extract_to_ast(pdf)
