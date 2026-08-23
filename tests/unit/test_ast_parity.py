"""AST-adapter parity harness (Option A, stage 2a).

Objective gates for the renderer convergence: the AST adapters must preserve all
content (round-trip fidelity), honor the AST's page geometry and fonts (no silent
Times-New-Roman override, no hardcoded A4), and lose nothing the legacy DOCX
engine captures for the same source.
"""

import re

import pytest
from docx import Document
from docx.shared import Mm

from core.rendering.document_ast import (
    Blockquote,
    DocumentAST,
    DocumentMetadata,
    Figure,
    Heading,
    HeadingLevel,
    ListBlock,
    Paragraph,
    StyleSheet,
    TableBlock,
)
from core.rendering.document_extractor import extract_to_ast
from core.rendering.docx_adapter import render_docx_from_ast
from core.rendering.pdf_adapter import render_pdf_from_ast


def _png(tmp_path):
    from PIL import Image as PILImage

    p = tmp_path / "fig.png"
    PILImage.new("RGB", (16, 12), (40, 120, 200)).save(str(p), format="PNG")
    return p


def _rich_ast(png_path) -> DocumentAST:
    ast = DocumentAST(metadata=DocumentMetadata(title="Tài liệu", author="Tác giả"), styles=StyleSheet())
    ast.add_block(Heading(level=HeadingLevel.H1, text="Chương 1"))
    ast.add_block(Paragraph(text="Đoạn văn tiếng Việt đủ dấu."))
    ast.add_block(Heading(level=HeadingLevel.H2, text="Mục 1.1"))
    ast.add_block(ListBlock(items=["mục a", "mục b"], ordered=False))
    ast.add_block(TableBlock(rows=[["Cột 1", "Cột 2"], ["ô a", "ô b"]], header_rows=1))
    ast.add_block(Blockquote(text="Một câu trích dẫn."))
    ast.add_block(Figure(image_ref=str(png_path), caption="Hình 1"))
    return ast


def test_docx_roundtrip_keeps_all_content(tmp_path):
    out = tmp_path / "o.docx"
    render_docx_from_ast(_rich_ast(_png(tmp_path)), out)

    d = Document(str(out))
    text = "\n".join(p.text for p in d.paragraphs)
    for token in ("Chương 1", "Đoạn văn tiếng Việt đủ dấu.", "Mục 1.1", "mục a", "mục b", "Một câu trích dẫn."):
        assert token in text, f"DOCX lost: {token}"
    assert len(d.tables) == 1
    cells = [c.text for row in d.tables[0].rows for c in row.cells]
    assert "Cột 1" in cells and "ô a" in cells and "ô b" in cells
    assert len(d.inline_shapes) >= 1  # figure embedded, not just a placeholder


def test_docx_honors_font_not_forced_times(tmp_path):
    ast = DocumentAST(metadata=DocumentMetadata(title="T"), styles=StyleSheet())
    ast.add_block(Paragraph(text="Xin chào thế giới"))
    out = tmp_path / "f.docx"
    render_docx_from_ast(ast, out)

    d = Document(str(out))
    fonts = {r.font.name for p in d.paragraphs for r in p.runs if r.font.name}
    assert "Times New Roman" not in fonts  # regression: no longer force-swapped
    if fonts:  # if the adapter set an explicit body font, it must be the AST's Georgia
        assert "Georgia" in fonts


def test_docx_page_setup_from_metadata(tmp_path):
    ast = DocumentAST(metadata=DocumentMetadata(title="T"), styles=StyleSheet())  # A4 210x297, 25mm margins
    ast.add_block(Paragraph(text="x"))
    out = tmp_path / "p.docx"
    render_docx_from_ast(ast, out)

    sec = Document(str(out)).sections[0]
    assert abs(sec.page_width - Mm(210)) < Mm(1)
    assert abs(sec.page_height - Mm(297)) < Mm(1)
    assert abs(sec.left_margin - Mm(25)) < Mm(1)
    assert abs(sec.top_margin - Mm(25)) < Mm(1)


def test_pdf_page_size_from_metadata(tmp_path):
    md = DocumentMetadata(title="T", page_width_mm=150.0, page_height_mm=200.0)
    ast = DocumentAST(metadata=md, styles=StyleSheet())
    ast.add_block(Paragraph(text="một đoạn"))
    out = tmp_path / "p.pdf"
    render_pdf_from_ast(ast, out)

    raw = out.read_bytes().decode("latin-1", "ignore")
    m = re.search(r"/MediaBox\s*\[\s*0\s+0\s+([\d.]+)\s+([\d.]+)", raw)
    assert m, "no MediaBox in PDF"
    width_pt = float(m.group(1))
    assert 420 < width_pt < 430  # 150mm ≈ 425pt, i.e. NOT A4 (595pt)


def test_new_docx_loses_nothing_vs_legacy(tmp_path):
    """The AST DOCX path must contain every heading/paragraph/list token the
    legacy DocxRenderer captures for the same markdown (content completeness)."""
    md = "# Title\n\nA paragraph.\n\n| A | B |\n| - | - |\n| 1 | 2 |\n\n- one\n- two\n"
    (tmp_path / "in.md").write_text(md, encoding="utf-8")

    new_out = tmp_path / "new.docx"
    render_docx_from_ast(extract_to_ast(tmp_path / "in.md"), new_out)
    nd = Document(str(new_out))
    new_text = "\n".join(p.text for p in nd.paragraphs)
    new_cells = [c.text for t in nd.tables for row in t.rows for c in row.cells]

    assert "Title" in new_text and "A paragraph." in new_text
    assert "one" in new_text and "two" in new_text
    assert "A" in new_cells and "1" in new_cells and "2" in new_cells
    # (The legacy DocxRenderer comparison leg retired with the engines in
    # stage 5 — the absolute content assertions above are the guard.)
