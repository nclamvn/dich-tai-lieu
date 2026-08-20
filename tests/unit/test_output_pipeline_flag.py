"""Output pipeline flag — wire the AST stack behind OUTPUT_PIPELINE (Option A, 3).

Default-safe: with the flag unset/``engine`` the live DOCX/PDF export keeps using
the legacy engines untouched; with ``OUTPUT_PIPELINE=ast`` it routes through the
DocumentAST + core/rendering adapters, and *always* falls back to the legacy
engine if the AST path errors — so flipping the flag can never lose an output.
"""

import asyncio
from pathlib import Path

import pypdf

from core_v2 import output_converter as oc
from core_v2.output_converter import (
    OutputConverter,
    _ast_pipeline_enabled,
    _ast_template,
    _output_pipeline,
)

MD = "# Chương 1\n\nĐoạn có **đậm**, *nghiêng* và `code` tiếng Việt đủ dấu.\n"


# --------------------------------------------------------------------------- #
# Flag resolution
# --------------------------------------------------------------------------- #
def test_flag_default_is_engine(monkeypatch):
    monkeypatch.delenv("OUTPUT_PIPELINE", raising=False)
    assert _output_pipeline() == "engine"
    assert _ast_pipeline_enabled() is False


def test_flag_env_ast_enables_case_insensitive(monkeypatch):
    monkeypatch.setenv("OUTPUT_PIPELINE", "AST")
    assert _output_pipeline() == "ast"
    assert _ast_pipeline_enabled() is True


def test_flag_env_engine_disables(monkeypatch):
    monkeypatch.setenv("OUTPUT_PIPELINE", "engine")
    assert _ast_pipeline_enabled() is False


def test_ast_template_maps_unknown_and_auto_to_ebook():
    assert _ast_template("auto") == "ebook"
    assert _ast_template("") == "ebook"
    assert _ast_template("weird") == "ebook"
    assert _ast_template("business") == "business"
    assert _ast_template("ACADEMIC") == "academic"


# --------------------------------------------------------------------------- #
# The real AST render methods produce valid, faithful output
# --------------------------------------------------------------------------- #
def test_ast_docx_method_produces_valid_docx(tmp_path):
    conv = OutputConverter(temp_dir=tmp_path / "t")
    out = tmp_path / "o.docx"
    result = asyncio.run(
        conv._markdown_to_docx_via_ast(MD, out, "ebook", "Tiêu đề", "Tác giả", "vi")
    )
    assert Path(result).exists()

    import docx as _docx

    doc = _docx.Document(str(out))
    all_text = "\n".join(p.text for p in doc.paragraphs)
    assert "Chương 1" in all_text
    # inline bold survived through the AST into a real Word run
    assert any(
        r.text == "đậm" and r.font.bold for p in doc.paragraphs for r in p.runs
    )


def test_ast_pdf_method_produces_valid_pdf(tmp_path):
    conv = OutputConverter(temp_dir=tmp_path / "t")
    out = tmp_path / "o.pdf"
    result = asyncio.run(
        conv._markdown_to_pdf_via_ast(MD, out, "business", "T", "A", "vi")
    )
    assert Path(result).exists()
    text = "".join((pg.extract_text() or "") for pg in pypdf.PdfReader(str(out)).pages)
    assert "Chương 1" in text


# --------------------------------------------------------------------------- #
# Routing: the flag picks the branch (heavy renderers stubbed for speed)
# --------------------------------------------------------------------------- #
def test_routing_on_uses_ast_docx(monkeypatch, tmp_path):
    monkeypatch.setenv("OUTPUT_PIPELINE", "ast")
    conv = OutputConverter(temp_dir=tmp_path / "t")
    seen = {}

    async def fake_ast(md, out, template, title, author, language):
        seen["args"] = (template, title, author, language)
        return Path(out)

    monkeypatch.setattr(conv, "_markdown_to_docx_via_ast", fake_ast)
    out = tmp_path / "o.docx"
    result = asyncio.run(
        conv.convert_markdown_to_docx_professional(MD, out, template="ebook", title="X")
    )
    assert seen["args"] == ("ebook", "X", "Unknown", "vi")
    assert Path(result) == out


def test_routing_on_uses_ast_pdf(monkeypatch, tmp_path):
    monkeypatch.setenv("OUTPUT_PIPELINE", "ast")
    conv = OutputConverter(temp_dir=tmp_path / "t")
    seen = {}

    async def fake_ast(md, out, template, title, author, language):
        seen["template"] = template
        return Path(out)

    monkeypatch.setattr(conv, "_markdown_to_pdf_via_ast", fake_ast)
    out = tmp_path / "o.pdf"
    asyncio.run(
        conv.convert_markdown_to_pdf_professional(MD, out, template="academic")
    )
    assert seen["template"] == "academic"


def test_routing_off_skips_ast_and_uses_legacy(monkeypatch, tmp_path):
    monkeypatch.setenv("OUTPUT_PIPELINE", "engine")
    conv = OutputConverter(temp_dir=tmp_path / "t")

    async def must_not_run(*a, **k):
        raise AssertionError("AST path must not run when the flag is off")

    monkeypatch.setattr(conv, "_markdown_to_docx_via_ast", must_not_run)

    sentinel = tmp_path / "legacy.docx"

    async def fake_run_blocking(fn, *a, **k):  # stub the heavy legacy render
        return sentinel

    monkeypatch.setattr(oc, "run_blocking", fake_run_blocking)
    out = tmp_path / "o.docx"
    result = asyncio.run(conv.convert_markdown_to_docx_professional(MD, out))
    assert result == sentinel  # legacy branch ran


# --------------------------------------------------------------------------- #
# Safety net: flag on, AST errors -> automatic fallback to the legacy engine
# --------------------------------------------------------------------------- #
def test_flag_on_falls_back_to_legacy_on_ast_error(monkeypatch, tmp_path):
    monkeypatch.setenv("OUTPUT_PIPELINE", "ast")
    conv = OutputConverter(temp_dir=tmp_path / "t")

    async def broken(*a, **k):
        raise RuntimeError("simulated AST failure")

    monkeypatch.setattr(conv, "_markdown_to_docx_via_ast", broken)

    sentinel = tmp_path / "legacy.docx"

    async def fake_run_blocking(fn, *a, **k):
        return sentinel

    monkeypatch.setattr(oc, "run_blocking", fake_run_blocking)
    out = tmp_path / "o.docx"
    result = asyncio.run(conv.convert_markdown_to_docx_professional(MD, out))
    assert result == sentinel  # fell back to legacy despite the flag being on
