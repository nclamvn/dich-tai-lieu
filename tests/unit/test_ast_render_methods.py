"""The AST renderer stack — the ONLY live DOCX/PDF renderer (Option A complete).

Stage 5 retired the legacy docx_engine/pdf_engine and the OUTPUT_PIPELINE flag;
the professional converters now route straight through the DocumentAST +
core/rendering adapters. These tests pin the template mapping, prove the render
methods produce valid faithful output, and verify the public converters delegate
to the AST methods with the right arguments.

(Successor of test_output_pipeline_flag.py, whose flag-resolution tests died
with the flag.)
"""

import asyncio
from pathlib import Path

import pypdf

from core_v2.output_converter import OutputConverter, _ast_template

MD = "# Chương 1\n\nĐoạn có **đậm**, *nghiêng* và `code` tiếng Việt đủ dấu.\n"


# --------------------------------------------------------------------------- #
# Template mapping
# --------------------------------------------------------------------------- #
def test_ast_template_maps_unknown_and_auto_to_ebook():
    assert _ast_template("auto") == "ebook"
    assert _ast_template("") == "ebook"
    assert _ast_template("weird") == "ebook"
    assert _ast_template("business") == "business"
    assert _ast_template("ACADEMIC") == "academic"


# --------------------------------------------------------------------------- #
# The real AST render methods produce valid, faithful output
# --------------------------------------------------------------------------- #
def test_ast_docx_method_produces_valid_docx(tmp_path):
    conv = OutputConverter(temp_dir=tmp_path / "t")
    out = tmp_path / "o.docx"
    result = asyncio.run(
        conv._markdown_to_docx_via_ast(MD, out, "ebook", "Tiêu đề", "Tác giả", "vi")
    )
    assert Path(result).exists()

    import docx as _docx

    doc = _docx.Document(str(out))
    all_text = "\n".join(p.text for p in doc.paragraphs)
    assert "Chương 1" in all_text
    # inline bold survived through the AST into a real Word run
    assert any(
        r.text == "đậm" and r.font.bold for p in doc.paragraphs for r in p.runs
    )


def test_ast_pdf_method_produces_valid_pdf(tmp_path):
    conv = OutputConverter(temp_dir=tmp_path / "t")
    out = tmp_path / "o.pdf"
    result = asyncio.run(
        conv._markdown_to_pdf_via_ast(MD, out, "business", "T", "A", "vi")
    )
    assert Path(result).exists()
    text = "".join((pg.extract_text() or "") for pg in pypdf.PdfReader(str(out)).pages)
    assert "Chương 1" in text


# --------------------------------------------------------------------------- #
# The public professional converters delegate to the AST methods
# --------------------------------------------------------------------------- #
def test_professional_docx_delegates_to_ast(monkeypatch, tmp_path):
    conv = OutputConverter(temp_dir=tmp_path / "t")
    seen = {}

    async def fake_ast(md, out, template, title, author, language):
        seen["args"] = (template, title, author, language)
        return Path(out)

    monkeypatch.setattr(conv, "_markdown_to_docx_via_ast", fake_ast)
    out = tmp_path / "o.docx"
    result = asyncio.run(
        conv.convert_markdown_to_docx_professional(MD, out, template="ebook", title="X")
    )
    assert seen["args"] == ("ebook", "X", "Unknown", "vi")
    assert Path(result) == out


def test_professional_pdf_delegates_to_ast(monkeypatch, tmp_path):
    conv = OutputConverter(temp_dir=tmp_path / "t")
    seen = {}

    async def fake_ast(md, out, template, title, author, language):
        seen["template"] = template
        return Path(out)

    monkeypatch.setattr(conv, "_markdown_to_pdf_via_ast", fake_ast)
    out = tmp_path / "o.pdf"
    asyncio.run(
        conv.convert_markdown_to_pdf_professional(MD, out, template="academic")
    )
    assert seen["template"] == "academic"


def test_ast_failure_propagates_for_orchestrator_fallback(monkeypatch, tmp_path):
    """With the legacy engine gone, an AST failure must RAISE (the orchestrator
    catches it and falls back to pandoc) — never silently produce nothing."""
    import pytest

    conv = OutputConverter(temp_dir=tmp_path / "t")

    async def boom(*a, **k):
        raise RuntimeError("render exploded")

    monkeypatch.setattr(conv, "_markdown_to_docx_via_ast", boom)
    with pytest.raises(RuntimeError, match="render exploded"):
        asyncio.run(
            conv.convert_markdown_to_docx_professional(MD, tmp_path / "o.docx")
        )
