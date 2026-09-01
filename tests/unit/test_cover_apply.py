"""Engine-agnostic cover application (Phase 2): stamp a chosen template cover
onto a finished PDF / DOCX, and through the professional converters.

All no-API-key and deterministic. A cover must add exactly one page (PDF) or one
zero-margin cover section (DOCX); an unknown id or missing file is a safe no-op.
"""

import asyncio

import docx
import pypdf
import pytest
from PIL import Image

from core.rendering.cover_apply import apply_cover_to_docx, apply_cover_to_pdf
from core.rendering.cover_templates import render_cover_image
from core.rendering.docx_adapter import render_docx_from_ast
from core.rendering.document_ast import (
    DocumentAST,
    DocumentMetadata,
    Heading,
    HeadingLevel,
    Paragraph,
    StyleSheet,
)
from core.rendering.pdf_adapter import render_pdf_from_ast


def _meta():
    return DocumentMetadata(title="BÌNH MINH XANH", author="Trần Văn Bút", language="vi")


def _ast():
    ast = DocumentAST(metadata=_meta(), styles=StyleSheet())
    ast.add_block(Heading(level=HeadingLevel.H1, text="Chương một"))
    ast.add_block(Paragraph(text="Nội dung chương một, tiếng Việt đủ dấu."))
    ast.add_block(Heading(level=HeadingLevel.H1, text="Chương hai"))
    ast.add_block(Paragraph(text="Nội dung chương hai."))
    return ast


def test_render_cover_image_is_portrait_png(tmp_path):
    out = tmp_path / "cover.png"
    render_cover_image("noir", _meta(), out)
    assert out.stat().st_size > 1000
    im = Image.open(out)
    assert im.width > 200 and im.height > im.width  # A4-ish portrait


def test_apply_cover_to_pdf_prepends_one_page(tmp_path):
    p = tmp_path / "book.pdf"
    render_pdf_from_ast(_ast(), p)
    n0 = len(pypdf.PdfReader(str(p)).pages)
    assert apply_cover_to_pdf(p, "gradient", title="BÌNH MINH XANH", author="Trần Văn Bút") is True
    assert len(pypdf.PdfReader(str(p)).pages) == n0 + 1


def test_apply_cover_to_pdf_unknown_is_noop(tmp_path):
    p = tmp_path / "book.pdf"
    render_pdf_from_ast(_ast(), p)
    n0 = len(pypdf.PdfReader(str(p)).pages)
    assert apply_cover_to_pdf(p, "no-such-template") is False
    assert len(pypdf.PdfReader(str(p)).pages) == n0


def test_apply_cover_to_pdf_missing_file_is_noop(tmp_path):
    assert apply_cover_to_pdf(tmp_path / "absent.pdf", "noir") is False


def test_apply_cover_to_docx_adds_zero_margin_cover_section(tmp_path):
    d = tmp_path / "book.docx"
    render_docx_from_ast(_ast(), d)
    s0 = len(docx.Document(str(d)).sections)
    assert apply_cover_to_docx(d, "emblem", title="BÌNH MINH XANH", author="Trần Văn Bút") is True
    dd = docx.Document(str(d))
    assert len(dd.sections) == s0 + 1  # cover section added ahead of the body
    assert "blip" in dd.paragraphs[0]._p.xml  # first paragraph carries the image
    # the cover section's margins are zero (true full-bleed)
    assert int(dd.sections[0].left_margin) == 0
    assert int(dd.sections[0].top_margin) == 0


def test_apply_cover_to_docx_unknown_is_noop(tmp_path):
    d = tmp_path / "book.docx"
    render_docx_from_ast(_ast(), d)
    s0 = len(docx.Document(str(d)).sections)
    assert apply_cover_to_docx(d, "no-such-template") is False
    assert len(docx.Document(str(d)).sections) == s0


@pytest.mark.parametrize("fmt", ["pdf", "docx"])
def test_converter_applies_cover_when_requested(tmp_path, fmt):
    from core_v2.output_converter import OutputConverter

    conv = OutputConverter(temp_dir=tmp_path / "oc")
    md = "# Chương một\n\nNội dung chương một.\n\n# Chương hai\n\nNội dung hai."

    async def _run(cover):
        out = tmp_path / f"{cover or 'plain'}.{fmt}"
        fn = (
            conv.convert_markdown_to_pdf_professional
            if fmt == "pdf"
            else conv.convert_markdown_to_docx_professional
        )
        await fn(md, out, title="BÌNH MINH XANH", author="Trần Văn Bút",
                 language="vi", cover_template=cover)
        return out

    plain = asyncio.run(_run(None))
    covered = asyncio.run(_run("noir"))

    if fmt == "pdf":
        assert len(pypdf.PdfReader(str(covered)).pages) == len(pypdf.PdfReader(str(plain)).pages) + 1
    else:
        assert len(docx.Document(str(covered)).sections) == len(docx.Document(str(plain)).sections) + 1
