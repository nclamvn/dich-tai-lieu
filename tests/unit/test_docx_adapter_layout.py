"""Sóng 1: docx_adapter renders the new layout blocks + AST->DOCX->AST round trip."""

from docx import Document as DocxDocument

from core.rendering.docx_adapter import render_docx_from_ast
from core.rendering.document_ast import (
    Caption,
    DocumentAST,
    DocumentMetadata,
    Heading,
    HeadingLevel,
    ListBlock,
    PageBreak,
    Paragraph,
    TableBlock,
    create_book_stylesheet,
)
from core.rendering.document_extractor import extract_to_ast


def _ast_with_layout_blocks() -> DocumentAST:
    doc = DocumentAST(
        metadata=DocumentMetadata(title="Tài liệu"),
        styles=create_book_stylesheet(),
    )
    doc.add_block(Heading(level=HeadingLevel.H1, text="Chương 1"))
    doc.add_block(Paragraph(text="Đoạn mở đầu."))
    doc.add_block(ListBlock(items=["mục a", "mục b"], ordered=False))
    doc.add_block(TableBlock(rows=[["H1", "H2"], ["1", "2"]], header_rows=1, caption="Bảng 1"))
    doc.add_block(PageBreak())
    doc.add_block(Caption(text="Chú thích", target="figure", number="1"))
    return doc


def test_render_produces_valid_docx_with_table(tmp_path):
    out = tmp_path / "out.docx"
    render_docx_from_ast(_ast_with_layout_blocks(), out)
    assert out.exists()

    rendered = DocxDocument(str(out))
    assert len(rendered.tables) == 1
    assert rendered.tables[0].rows[0].cells[0].text == "H1"
    assert any("Chương 1" in p.text for p in rendered.paragraphs)


def test_render_table_cell_content(tmp_path):
    out = tmp_path / "out.docx"
    render_docx_from_ast(_ast_with_layout_blocks(), out)
    table = DocxDocument(str(out)).tables[0]
    assert [c.text for c in table.rows[1].cells] == ["1", "2"]


def test_render_list_items_present(tmp_path):
    out = tmp_path / "out.docx"
    render_docx_from_ast(_ast_with_layout_blocks(), out)
    texts = [p.text for p in DocxDocument(str(out)).paragraphs]
    assert "mục a" in texts and "mục b" in texts


def test_round_trip_preserves_text_structure(tmp_path):
    ast1 = _ast_with_layout_blocks()
    out = tmp_path / "rt.docx"
    render_docx_from_ast(ast1, out)

    ast2 = extract_to_ast(out)
    s1, s2 = ast1.get_statistics(), ast2.get_statistics()
    assert s2["headings"] == s1["headings"]
    assert s2["tables"] == s1["tables"]
    assert s2["lists"] == s1["lists"]


def test_new_blocks_are_not_dropped(tmp_path):
    """Before this change the layout blocks were logged as 'Unknown' and skipped."""
    out = tmp_path / "out.docx"
    render_docx_from_ast(_ast_with_layout_blocks(), out)
    rendered = DocxDocument(str(out))
    # A table exists (TableBlock rendered) and the caption text made it in.
    assert len(rendered.tables) == 1
    assert any("Chú thích" in p.text for p in rendered.paragraphs)
