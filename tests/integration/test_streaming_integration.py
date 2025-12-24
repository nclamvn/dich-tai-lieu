#!/usr/bin/env python3
# -*- coding: utf-8 -*-

"""
Integration Tests for Streaming Pipeline - Phase 5.4

Tests full end-to-end streaming flow with all formats
"""

import pytest
import asyncio
import httpx
from pathlib import Path
from unittest.mock import Mock, AsyncMock
from docx import Document
from pypdf import PdfReader

from core.streaming import StreamingBatchProcessor
from core.validator import TranslationResult
from core.chunker import TranslationChunk
from core.job_queue import TranslationJob, JobStatus


class TestStreamingIntegration:
    """Integration tests for full streaming pipeline"""

    @pytest.fixture
    def mock_translator(self):
        """Mock translator that returns fake translations"""
        translator = Mock()

        async def translate_chunk(client, chunk):
            """Simulate translation"""
            await asyncio.sleep(0.01)  # Simulate API call
            return TranslationResult(
                chunk_id=chunk.id,
                source=chunk.text,
                translated=f"Translated: {chunk.text}",
                quality_score=0.9
            )

        translator.translate_chunk = AsyncMock(side_effect=translate_chunk)
        return translator

    @pytest.fixture
    def large_job_250_chunks(self):
        """Create job with 250 chunks (3 batches of 100)"""
        chunks = [
            TranslationChunk(
                id=i,
                text=f"Source text {i} for translation test",
                context_before="",
                context_after=""
            )
            for i in range(250)
        ]

        job = TranslationJob(
            job_id="test_large_job",
            job_name="Integration Test Large Job",
            input_file="/tmp/test_input.txt",
            output_file="/tmp/test_output.docx",
            status=JobStatus.RUNNING,
            output_format="docx"
        )

        return job, chunks

    @pytest.mark.asyncio
    async def test_full_docx_streaming_pipeline(self, mock_translator, large_job_250_chunks, tmp_path):
        """Test complete DOCX streaming: 250 chunks → 3 batches → final merge"""
        job, chunks = large_job_250_chunks
        output_path = tmp_path / "output.docx"

        # Create processor
        processor = StreamingBatchProcessor(
            batch_size=100,
            enable_streaming=True,
            enable_partial_export=True,
            websocket_manager=None
        )

        # Process streaming
        async with httpx.AsyncClient() as client:
            results, stats = await processor.process_streaming(
                job=job,
                chunks=chunks,
                translator=mock_translator,
                http_client=client,
                output_path=output_path
            )

        # Verify results
        assert len(results) == 250, "Should have 250 results"
        assert stats['batches_processed'] == 3, "Should process 3 batches"
        assert stats['chunks_processed'] == 250
        assert len(stats['partial_exports']) == 3, "Should have 3 partial exports"

        # Verify output file
        assert output_path.exists(), "Output file should exist"
        assert output_path.stat().st_size > 0, "Output file should not be empty"

        # Verify DOCX is readable
        doc = Document(output_path)
        assert len(doc.paragraphs) == 250, f"Should have 250 paragraphs, got {len(doc.paragraphs)}"
        assert "Translated: Source text 0" in doc.paragraphs[0].text
        assert "Translated: Source text 249" in doc.paragraphs[249].text

        # Verify temp files cleaned up
        temp_dir = output_path.parent / ".temp_docx_batches"
        if temp_dir.exists():
            remaining = list(temp_dir.glob("*.docx"))
            assert len(remaining) == 0, f"Temp files should be cleaned, found: {remaining}"

    @pytest.mark.asyncio
    async def test_full_pdf_streaming_pipeline(self, mock_translator, large_job_250_chunks, tmp_path):
        """Test complete PDF streaming pipeline"""
        job, chunks = large_job_250_chunks
        job.output_format = "pdf"
        output_path = tmp_path / "output.pdf"

        processor = StreamingBatchProcessor(batch_size=100)

        async with httpx.AsyncClient() as client:
            results, stats = await processor.process_streaming(
                job=job,
                chunks=chunks,
                translator=mock_translator,
                http_client=client,
                output_path=output_path
            )

        # Verify results
        assert len(results) == 250
        assert stats['batches_processed'] == 3

        # Verify PDF
        assert output_path.exists()
        reader = PdfReader(output_path)
        assert len(reader.pages) > 0, "PDF should have pages"

        # Extract text from PDF and verify
        text = ""
        for page in reader.pages:
            text += page.extract_text()
        assert "Translated: Source text 0" in text
        assert "Translated: Source text 249" in text

        # Verify temp cleanup
        temp_dir = output_path.parent / ".temp_pdf_batches"
        if temp_dir.exists():
            assert len(list(temp_dir.glob("*.pdf"))) == 0

    @pytest.mark.asyncio
    async def test_full_txt_streaming_pipeline(self, mock_translator, large_job_250_chunks, tmp_path):
        """Test complete TXT streaming pipeline"""
        job, chunks = large_job_250_chunks
        job.output_format = "txt"
        output_path = tmp_path / "output.txt"

        processor = StreamingBatchProcessor(batch_size=100)

        async with httpx.AsyncClient() as client:
            results, stats = await processor.process_streaming(
                job=job,
                chunks=chunks,
                translator=mock_translator,
                http_client=client,
                output_path=output_path
            )

        # Verify results
        assert len(results) == 250
        assert stats['batches_processed'] == 3

        # Verify TXT
        assert output_path.exists()
        content = output_path.read_text(encoding='utf-8')
        assert "Translated: Source text 0" in content
        assert "Translated: Source text 249" in content
        assert len(content) > 0

        # Verify temp cleanup
        temp_dir = output_path.parent / ".temp_txt_batches"
        if temp_dir.exists():
            assert len(list(temp_dir.glob("*.txt"))) == 0

    @pytest.mark.asyncio
    async def test_error_recovery_cleanup(self, mock_translator, tmp_path):
        """Test cleanup when builder raises exception during processing"""
        from core.streaming import IncrementalDocxBuilder

        output_path = tmp_path / "output.docx"

        # Create a builder and force an error scenario
        builder = IncrementalDocxBuilder(output_path)

        results = [
            TranslationResult(
                chunk_id=i,
                source=f"Source {i}",
                translated=f"Text {i}",
                quality_score=0.9
            )
            for i in range(10)
        ]

        # Add first batch successfully
        await builder.add_batch(results[:5], 0)
        batch_file = builder.batch_files[0]
        assert batch_file.exists(), "Batch file should exist"

        # Simulate error during second batch by using invalid path
        builder.temp_dir = tmp_path / "nonexistent" / "path"

        # This should raise an error
        with pytest.raises(Exception):
            await builder.add_batch(results[5:], 1)

        # Cleanup should remove temp files
        await builder.cleanup()
        assert not batch_file.exists(), "Batch file should be cleaned up after error"

    @pytest.mark.asyncio
    async def test_small_job_no_streaming(self, mock_translator, tmp_path):
        """Test that small jobs (< batch_size) work correctly"""
        # Create small job with only 50 chunks
        chunks = [
            TranslationChunk(
                id=i,
                text=f"Source text {i}",
                context_before="",
                context_after=""
            )
            for i in range(50)
        ]

        job = TranslationJob(
            job_id="test_small_job",
            job_name="Integration Test Small Job",
            input_file="/tmp/test_small_input.txt",
            output_file="/tmp/test_small_output.docx",
            status=JobStatus.RUNNING,
            output_format="docx"
        )

        output_path = tmp_path / "small_output.docx"
        processor = StreamingBatchProcessor(batch_size=100)

        async with httpx.AsyncClient() as client:
            results, stats = await processor.process_streaming(
                job=job,
                chunks=chunks,
                translator=mock_translator,
                http_client=client,
                output_path=output_path
            )

        # Should process as single batch
        assert len(results) == 50
        assert stats['batches_processed'] == 1
        assert output_path.exists()

    @pytest.mark.asyncio
    async def test_context_manager_with_streaming(self, mock_translator, large_job_250_chunks, tmp_path):
        """Test using builder with context manager in streaming"""
        from core.streaming import IncrementalDocxBuilder

        job, chunks = large_job_250_chunks
        output_path = tmp_path / "ctx_output.docx"

        # Create results manually
        results = [
            TranslationResult(
                chunk_id=i,
                source=f"Source {i}",
                translated=f"Translated {i}",
                quality_score=0.9
            )
            for i in range(250)
        ]

        # Use context manager
        async with IncrementalDocxBuilder(output_path) as builder:
            # Add batches
            for batch_idx in range(3):
                start = batch_idx * 100
                end = min(start + 100, len(results))
                batch_results = results[start:end]
                await builder.add_batch(batch_results, batch_idx)

            # Merge
            final_output = await builder.merge_all()

        # Verify
        assert final_output.exists()
        doc = Document(final_output)
        assert len(doc.paragraphs) == 250

        # Verify cleanup
        temp_dir = output_path.parent / ".temp_docx_batches"
        if temp_dir.exists():
            assert len(list(temp_dir.glob("*.docx"))) == 0


if __name__ == "__main__":
    pytest.main([__file__, "-v", "--tb=short"])
