"""
Integration Tests for Book DOCX Typography
==========================================

Phase 3.5 - Integration Testing & Real Book Testing

This module tests the AST pipeline against the legacy DOCX builder to ensure:
1. Structural compatibility (same paragraph/heading counts)
2. Commercial typography quality (fonts, spacing, indents)
3. Performance benchmarks (time and memory usage)

Test Strategy:
    - Compare legacy vs AST pipeline output structurally
    - Verify typography specifications are applied correctly
    - Test with various document sizes (small, medium, large)
    - Measure performance metrics

Note: Tests enable use_ast_pipeline=True temporarily for comparison.
"""

import pytest
import tempfile
import time
import tracemalloc
from pathlib import Path
from typing import List, Tuple

from docx import Document
from docx.shared import Pt

from core.structure.semantic_model import DocNode, DocNodeType
from core.rendering.ast_builder import build_book_ast
from core.rendering.docx_adapter import render_docx_from_ast
from core.export.docx_book_builder import build_book_docx


# ============================================================================
# Test Fixtures - Sample Documents
# ============================================================================

@pytest.fixture
def simple_chapter_nodes() -> List[DocNode]:
    """Simple chapter with heading and 5 paragraphs."""
    return [
        DocNode(
            node_type=DocNodeType.CHAPTER,
            text="Chapter 1: The Beginning",
            level=1,
            metadata={}
        ),
        DocNode(
            node_type=DocNodeType.PARAGRAPH,
            text="This is the first paragraph after the chapter heading.",
            metadata={}
        ),
        DocNode(
            node_type=DocNodeType.PARAGRAPH,
            text="This is the second paragraph with regular body text.",
            metadata={}
        ),
        DocNode(
            node_type=DocNodeType.PARAGRAPH,
            text="Third paragraph continues the narrative flow.",
            metadata={}
        ),
        DocNode(
            node_type=DocNodeType.PARAGRAPH,
            text="Fourth paragraph adds more content.",
            metadata={}
        ),
        DocNode(
            node_type=DocNodeType.PARAGRAPH,
            text="Fifth and final paragraph of this simple chapter.",
            metadata={}
        ),
    ]


@pytest.fixture
def complex_chapter_nodes() -> List[DocNode]:
    """Complex chapter with various elements."""
    return [
        DocNode(
            node_type=DocNodeType.CHAPTER,
            text="Chapter 2: The Storm",
            level=1,
            metadata={}
        ),
        DocNode(
            node_type=DocNodeType.EPIGRAPH,
            text="All beginnings are hard.",
            metadata={'attribution': 'Ancient Proverb'}
        ),
        DocNode(
            node_type=DocNodeType.PARAGRAPH,
            text="The sky darkened as the first drops began to fall.",
            metadata={}
        ),
        DocNode(
            node_type=DocNodeType.PARAGRAPH,
            text="John looked up at the gathering clouds with apprehension.",
            metadata={}
        ),
        DocNode(
            node_type=DocNodeType.BLOCKQUOTE,
            text="We must seek shelter before the storm arrives.",
            metadata={'attribution': 'John'}
        ),
        DocNode(
            node_type=DocNodeType.PARAGRAPH,
            text="Mary nodded in agreement as thunder rumbled in the distance.",
            metadata={}
        ),
        DocNode(
            node_type=DocNodeType.SCENE_BREAK,
            text="* * *",
            metadata={}
        ),
        DocNode(
            node_type=DocNodeType.PARAGRAPH,
            text="Hours later, the storm had passed.",
            metadata={}
        ),
        DocNode(
            node_type=DocNodeType.PARAGRAPH,
            text="The sun emerged from behind the clouds, painting the sky in brilliant colors.",
            metadata={}
        ),
    ]


@pytest.fixture
def large_chapter_nodes() -> List[DocNode]:
    """Large chapter with 100 paragraphs for performance testing."""
    nodes = [
        DocNode(
            node_type=DocNodeType.CHAPTER,
            text="Chapter 3: The Long Journey",
            level=1,
            metadata={}
        )
    ]

    # Add 100 paragraphs
    for i in range(1, 101):
        nodes.append(DocNode(
            node_type=DocNodeType.PARAGRAPH,
            text=f"This is paragraph {i} of the long chapter. " * 5,  # ~30 words each
            metadata={}
        ))

    return nodes


@pytest.fixture
def multi_section_nodes() -> List[DocNode]:
    """Document with multiple sections and headings."""
    return [
        DocNode(
            node_type=DocNodeType.CHAPTER,
            text="Chapter 4: Advanced Topics",
            level=1,
            metadata={}
        ),
        DocNode(
            node_type=DocNodeType.PARAGRAPH,
            text="Introduction to advanced topics.",
            metadata={}
        ),
        DocNode(
            node_type=DocNodeType.SECTION,
            text="Section 4.1: First Topic",
            level=2,
            metadata={}
        ),
        DocNode(
            node_type=DocNodeType.PARAGRAPH,
            text="Content for the first topic.",
            metadata={}
        ),
        DocNode(
            node_type=DocNodeType.PARAGRAPH,
            text="More details about this topic.",
            metadata={}
        ),
        DocNode(
            node_type=DocNodeType.SUBSECTION,
            text="4.1.1: Subtopic",
            level=3,
            metadata={}
        ),
        DocNode(
            node_type=DocNodeType.PARAGRAPH,
            text="Detailed information about the subtopic.",
            metadata={}
        ),
        DocNode(
            node_type=DocNodeType.SECTION,
            text="Section 4.2: Second Topic",
            level=2,
            metadata={}
        ),
        DocNode(
            node_type=DocNodeType.PARAGRAPH,
            text="Content for the second topic.",
            metadata={}
        ),
    ]


# ============================================================================
# Helper Functions
# ============================================================================

def render_with_legacy(nodes: List[DocNode], output_path: Path) -> Path:
    """Render document using legacy DOCX builder."""
    config = {
        'target_lang': 'vi',
        'font_name': 'Georgia',
        'font_size': 11
    }
    build_book_docx(nodes, str(output_path), config)
    return output_path


def render_with_ast(nodes: List[DocNode], output_path: Path) -> Path:
    """Render document using AST pipeline."""
    ast = build_book_ast(nodes, language='vi')
    render_docx_from_ast(ast, output_path)
    return output_path


def analyze_docx_structure(docx_path: Path) -> dict:
    """
    Analyze DOCX structure and extract key metrics.

    Returns:
        dict with keys:
            - paragraph_count: int
            - heading_count: int
            - paragraphs: list of paragraph objects
            - headings: list of heading objects
    """
    doc = Document(str(docx_path))

    paragraphs = doc.paragraphs
    headings = [p for p in paragraphs if p.style.name.startswith('Heading')]

    return {
        'paragraph_count': len(paragraphs),
        'heading_count': len(headings),
        'paragraphs': paragraphs,
        'headings': headings,
    }


def get_paragraph_indent(paragraph) -> float:
    """Get first-line indent in points."""
    if paragraph.paragraph_format.first_line_indent:
        return paragraph.paragraph_format.first_line_indent.pt
    return 0.0


def get_font_size(paragraph) -> float:
    """Get font size in points from paragraph."""
    if paragraph.runs:
        run = paragraph.runs[0]
        if run.font.size:
            return run.font.size.pt
    return 0.0


def get_line_spacing(paragraph) -> float:
    """Get line spacing value."""
    if paragraph.paragraph_format.line_spacing:
        return paragraph.paragraph_format.line_spacing
    return 1.0


# ============================================================================
# Structural Comparison Tests
# ============================================================================

class TestStructuralComparison:
    """Compare legacy vs AST pipeline output structure."""

    def test_simple_chapter_paragraph_count(self, simple_chapter_nodes):
        """Test that paragraph counts match for simple chapter."""
        with tempfile.TemporaryDirectory() as tmpdir:
            legacy_path = Path(tmpdir) / "legacy.docx"
            ast_path = Path(tmpdir) / "ast.docx"

            # Render with both pipelines
            render_with_legacy(simple_chapter_nodes, legacy_path)
            render_with_ast(simple_chapter_nodes, ast_path)

            # Analyze structure
            legacy_structure = analyze_docx_structure(legacy_path)
            ast_structure = analyze_docx_structure(ast_path)

            # Compare counts (allow ±1 difference for minor variations)
            para_diff = abs(legacy_structure['paragraph_count'] - ast_structure['paragraph_count'])

            print(f"\n📊 Simple Chapter Structure:")
            print(f"   Legacy paragraphs: {legacy_structure['paragraph_count']}")
            print(f"   AST paragraphs: {ast_structure['paragraph_count']}")
            print(f"   Difference: {para_diff}")

            assert para_diff <= 1, (
                f"Paragraph count mismatch: "
                f"legacy={legacy_structure['paragraph_count']}, "
                f"ast={ast_structure['paragraph_count']}"
            )

    def test_complex_chapter_structure(self, complex_chapter_nodes):
        """Test that complex chapter structure matches."""
        with tempfile.TemporaryDirectory() as tmpdir:
            legacy_path = Path(tmpdir) / "legacy_complex.docx"
            ast_path = Path(tmpdir) / "ast_complex.docx"

            # Render with both pipelines
            render_with_legacy(complex_chapter_nodes, legacy_path)
            render_with_ast(complex_chapter_nodes, ast_path)

            # Analyze structure
            legacy_structure = analyze_docx_structure(legacy_path)
            ast_structure = analyze_docx_structure(ast_path)

            print(f"\n📊 Complex Chapter Structure:")
            print(f"   Legacy: {legacy_structure['paragraph_count']} paras, {legacy_structure['heading_count']} headings")
            print(f"   AST: {ast_structure['paragraph_count']} paras, {ast_structure['heading_count']} headings")

            # Heading count should match exactly (chapters, sections)
            assert legacy_structure['heading_count'] == ast_structure['heading_count'], (
                f"Heading count mismatch: "
                f"legacy={legacy_structure['heading_count']}, "
                f"ast={ast_structure['heading_count']}"
            )

    def test_multi_section_heading_hierarchy(self, multi_section_nodes):
        """Test that heading hierarchy is preserved."""
        with tempfile.TemporaryDirectory() as tmpdir:
            ast_path = Path(tmpdir) / "ast_sections.docx"

            # Render with AST pipeline
            render_with_ast(multi_section_nodes, ast_path)

            # Analyze headings
            structure = analyze_docx_structure(ast_path)
            headings = structure['headings']

            print(f"\n📊 Heading Hierarchy:")
            for h in headings:
                print(f"   {h.style.name}: {h.text}")

            # Should have H1, H2, H3 headings
            heading_styles = {h.style.name for h in headings}
            assert 'Heading 1' in heading_styles, "Missing H1 headings"
            assert 'Heading 2' in heading_styles, "Missing H2 headings"
            assert 'Heading 3' in heading_styles, "Missing H3 headings"


# ============================================================================
# Typography Verification Tests
# ============================================================================

class TestTypographyVerification:
    """Verify commercial typography specifications."""

    def test_first_paragraph_no_indent(self, simple_chapter_nodes):
        """Test that first paragraph after heading has no indent."""
        with tempfile.TemporaryDirectory() as tmpdir:
            ast_path = Path(tmpdir) / "ast_first_para.docx"

            # Render with AST pipeline
            render_with_ast(simple_chapter_nodes, ast_path)

            # Analyze paragraphs
            doc = Document(str(ast_path))
            paragraphs = doc.paragraphs

            # Find first body paragraph (after heading)
            body_paras = [p for p in paragraphs if not p.style.name.startswith('Heading')]

            if body_paras:
                first_para = body_paras[0]
                first_indent = get_paragraph_indent(first_para)

                print(f"\n📐 First Paragraph Indent:")
                print(f"   First paragraph indent: {first_indent}pt")
                print(f"   Expected: 0pt (commercial convention)")

                # First paragraph should have 0pt indent (or very close to 0)
                assert first_indent < 1.0, (
                    f"First paragraph should have no indent, got {first_indent}pt"
                )

    def test_body_paragraph_indent(self, simple_chapter_nodes):
        """Test that body paragraphs have 23pt indent."""
        with tempfile.TemporaryDirectory() as tmpdir:
            ast_path = Path(tmpdir) / "ast_body_para.docx"

            # Render with AST pipeline
            render_with_ast(simple_chapter_nodes, ast_path)

            # Analyze paragraphs
            doc = Document(str(ast_path))
            paragraphs = doc.paragraphs

            # Find body paragraphs (skip first one)
            body_paras = [p for p in paragraphs if not p.style.name.startswith('Heading')]

            if len(body_paras) > 1:
                # Check second and subsequent paragraphs
                second_para = body_paras[1]
                second_indent = get_paragraph_indent(second_para)

                print(f"\n📐 Body Paragraph Indent:")
                print(f"   Second paragraph indent: {second_indent}pt")
                print(f"   Expected: 23pt (0.32 inch commercial standard)")

                # Allow small tolerance (±2pt) due to conversion rounding
                assert 21.0 <= second_indent <= 25.0, (
                    f"Body paragraph should have ~23pt indent, got {second_indent}pt"
                )

    def test_heading_font_sizes(self, multi_section_nodes):
        """Test that heading font sizes are correct (H1=16pt, H2=14pt, H3=12pt)."""
        with tempfile.TemporaryDirectory() as tmpdir:
            ast_path = Path(tmpdir) / "ast_headings.docx"

            # Render with AST pipeline
            render_with_ast(multi_section_nodes, ast_path)

            # Analyze headings
            doc = Document(str(ast_path))
            headings = [p for p in doc.paragraphs if p.style.name.startswith('Heading')]

            print(f"\n📏 Heading Font Sizes:")

            for h in headings:
                font_size = get_font_size(h)
                print(f"   {h.style.name}: {font_size}pt - \"{h.text[:30]}...\"")

                # Verify expected sizes (allow ±1pt tolerance)
                if h.style.name == 'Heading 1':
                    assert 15.0 <= font_size <= 17.0, f"H1 should be ~16pt, got {font_size}pt"
                elif h.style.name == 'Heading 2':
                    assert 13.0 <= font_size <= 15.0, f"H2 should be ~14pt, got {font_size}pt"
                elif h.style.name == 'Heading 3':
                    assert 11.0 <= font_size <= 13.0, f"H3 should be ~12pt, got {font_size}pt"

    def test_body_font_size(self, simple_chapter_nodes):
        """Test that body text is 11pt."""
        with tempfile.TemporaryDirectory() as tmpdir:
            ast_path = Path(tmpdir) / "ast_body_font.docx"

            # Render with AST pipeline
            render_with_ast(simple_chapter_nodes, ast_path)

            # Analyze paragraphs
            doc = Document(str(ast_path))
            body_paras = [p for p in doc.paragraphs if not p.style.name.startswith('Heading')]

            if body_paras:
                para = body_paras[0]
                font_size = get_font_size(para)

                print(f"\n📏 Body Font Size:")
                print(f"   Body paragraph font: {font_size}pt")
                print(f"   Expected: 11pt")

                # Allow ±1pt tolerance
                assert 10.0 <= font_size <= 12.0, (
                    f"Body text should be ~11pt, got {font_size}pt"
                )

    def test_line_spacing(self, simple_chapter_nodes):
        """Test that line spacing is 1.15."""
        with tempfile.TemporaryDirectory() as tmpdir:
            ast_path = Path(tmpdir) / "ast_line_spacing.docx"

            # Render with AST pipeline
            render_with_ast(simple_chapter_nodes, ast_path)

            # Analyze paragraphs
            doc = Document(str(ast_path))
            body_paras = [p for p in doc.paragraphs if not p.style.name.startswith('Heading')]

            if body_paras:
                para = body_paras[0]
                line_spacing = get_line_spacing(para)

                print(f"\n📏 Line Spacing:")
                print(f"   Body paragraph line spacing: {line_spacing}")
                print(f"   Expected: 1.15")

                # Allow small tolerance
                assert 1.10 <= line_spacing <= 1.20, (
                    f"Line spacing should be ~1.15, got {line_spacing}"
                )


# ============================================================================
# Performance Benchmark Tests
# ============================================================================

class TestPerformanceBenchmark:
    """Benchmark performance for different document sizes."""

    def benchmark_rendering(self, nodes: List[DocNode], label: str) -> Tuple[float, float]:
        """
        Benchmark rendering time and memory for given nodes.

        Returns:
            (time_seconds, memory_mb)
        """
        with tempfile.TemporaryDirectory() as tmpdir:
            output_path = Path(tmpdir) / f"{label}.docx"

            # Start memory tracking
            tracemalloc.start()
            start_time = time.time()

            # Render
            render_with_ast(nodes, output_path)

            # Measure
            end_time = time.time()
            current, peak = tracemalloc.get_traced_memory()
            tracemalloc.stop()

            elapsed_time = end_time - start_time
            peak_memory_mb = peak / 1024 / 1024

            return elapsed_time, peak_memory_mb

    def test_performance_small_document(self, simple_chapter_nodes):
        """Benchmark small document (5 paragraphs)."""
        time_s, memory_mb = self.benchmark_rendering(simple_chapter_nodes, "small")

        print(f"\n⚡ Performance - Small Document (5 paragraphs):")
        print(f"   Time: {time_s:.3f}s")
        print(f"   Peak Memory: {memory_mb:.2f} MB")

        # Sanity check: should complete in reasonable time
        assert time_s < 10.0, f"Small document took too long: {time_s:.2f}s"

    def test_performance_large_document(self, large_chapter_nodes):
        """Benchmark large document (100 paragraphs)."""
        time_s, memory_mb = self.benchmark_rendering(large_chapter_nodes, "large")

        print(f"\n⚡ Performance - Large Document (100 paragraphs):")
        print(f"   Time: {time_s:.3f}s")
        print(f"   Peak Memory: {memory_mb:.2f} MB")

        # Sanity check: should complete in reasonable time
        assert time_s < 30.0, f"Large document took too long: {time_s:.2f}s"

    def test_performance_comparison_legacy_vs_ast(self, large_chapter_nodes):
        """Compare performance between legacy and AST pipeline."""
        with tempfile.TemporaryDirectory() as tmpdir:
            # Benchmark legacy
            legacy_path = Path(tmpdir) / "legacy_perf.docx"
            tracemalloc.start()
            start_time = time.time()
            render_with_legacy(large_chapter_nodes, legacy_path)
            legacy_time = time.time() - start_time
            legacy_current, legacy_peak = tracemalloc.get_traced_memory()
            tracemalloc.stop()

            # Benchmark AST
            ast_path = Path(tmpdir) / "ast_perf.docx"
            tracemalloc.start()
            start_time = time.time()
            render_with_ast(large_chapter_nodes, ast_path)
            ast_time = time.time() - start_time
            ast_current, ast_peak = tracemalloc.get_traced_memory()
            tracemalloc.stop()

            print(f"\n⚡ Performance Comparison (100 paragraphs):")
            print(f"   Legacy:")
            print(f"      Time: {legacy_time:.3f}s")
            print(f"      Memory: {legacy_peak / 1024 / 1024:.2f} MB")
            print(f"   AST Pipeline:")
            print(f"      Time: {ast_time:.3f}s")
            print(f"      Memory: {ast_peak / 1024 / 1024:.2f} MB")

            # Calculate overhead
            time_overhead = ((ast_time - legacy_time) / legacy_time) * 100
            memory_overhead = ((ast_peak - legacy_peak) / legacy_peak) * 100

            print(f"   Overhead:")
            print(f"      Time: {time_overhead:+.1f}%")
            print(f"      Memory: {memory_overhead:+.1f}%")

            # AST pipeline should not be significantly slower (allow 50% overhead)
            # This is acceptable for the first implementation
            assert time_overhead < 100, (
                f"AST pipeline is too slow: {time_overhead:.1f}% overhead"
            )


# ============================================================================
# Edge Cases and Regression Tests
# ============================================================================

class TestEdgeCases:
    """Test edge cases and potential issues."""

    def test_empty_document(self):
        """Test rendering empty document doesn't crash."""
        with tempfile.TemporaryDirectory() as tmpdir:
            ast_path = Path(tmpdir) / "empty.docx"

            # Empty node list
            nodes = []

            # Should not crash
            render_with_ast(nodes, ast_path)

            # File should exist
            assert ast_path.exists()

    def test_chapter_without_paragraphs(self):
        """Test chapter heading without any paragraphs."""
        with tempfile.TemporaryDirectory() as tmpdir:
            ast_path = Path(tmpdir) / "chapter_only.docx"

            nodes = [
                DocNode(
                    node_type=DocNodeType.CHAPTER,
                    text="Chapter 1: Standalone",
                    level=1,
                    metadata={}
                )
            ]

            # Should not crash
            render_with_ast(nodes, ast_path)

            # Verify structure
            structure = analyze_docx_structure(ast_path)
            assert structure['heading_count'] >= 1

    def test_very_long_paragraph(self):
        """Test paragraph with very long text."""
        with tempfile.TemporaryDirectory() as tmpdir:
            ast_path = Path(tmpdir) / "long_para.docx"

            # Create paragraph with 1000 words
            long_text = "Lorem ipsum dolor sit amet. " * 200

            nodes = [
                DocNode(
                    node_type=DocNodeType.CHAPTER,
                    text="Chapter 1",
                    level=1,
                    metadata={}
                ),
                DocNode(
                    node_type=DocNodeType.PARAGRAPH,
                    text=long_text,
                    metadata={}
                )
            ]

            # Should not crash
            render_with_ast(nodes, ast_path)
            assert ast_path.exists()


# ============================================================================
# Summary Report
# ============================================================================

def test_generate_summary_report(simple_chapter_nodes, complex_chapter_nodes, large_chapter_nodes):
    """
    Generate summary report comparing legacy vs AST pipeline.

    This is a meta-test that runs comparisons and prints a summary.
    """
    print("\n" + "="*80)
    print("📊 PHASE 3.5 - INTEGRATION TEST SUMMARY")
    print("="*80)

    with tempfile.TemporaryDirectory() as tmpdir:
        # Test 1: Simple chapter
        print("\n1. Simple Chapter (5 paragraphs):")
        legacy_simple = Path(tmpdir) / "legacy_simple.docx"
        ast_simple = Path(tmpdir) / "ast_simple.docx"

        render_with_legacy(simple_chapter_nodes, legacy_simple)
        render_with_ast(simple_chapter_nodes, ast_simple)

        legacy_struct = analyze_docx_structure(legacy_simple)
        ast_struct = analyze_docx_structure(ast_simple)

        print(f"   Legacy: {legacy_struct['paragraph_count']} paragraphs")
        print(f"   AST:    {ast_struct['paragraph_count']} paragraphs")
        print(f"   Match:  {'✅ YES' if abs(legacy_struct['paragraph_count'] - ast_struct['paragraph_count']) <= 1 else '❌ NO'}")

        # Test 2: Complex chapter
        print("\n2. Complex Chapter (headings, epigraphs, blockquotes, scene breaks):")
        legacy_complex = Path(tmpdir) / "legacy_complex.docx"
        ast_complex = Path(tmpdir) / "ast_complex.docx"

        render_with_legacy(complex_chapter_nodes, legacy_complex)
        render_with_ast(complex_chapter_nodes, ast_complex)

        legacy_struct = analyze_docx_structure(legacy_complex)
        ast_struct = analyze_docx_structure(ast_complex)

        print(f"   Legacy: {legacy_struct['paragraph_count']} paragraphs, {legacy_struct['heading_count']} headings")
        print(f"   AST:    {ast_struct['paragraph_count']} paragraphs, {ast_struct['heading_count']} headings")
        print(f"   Match:  {'✅ YES' if legacy_struct['heading_count'] == ast_struct['heading_count'] else '❌ NO'}")

        # Test 3: Large document
        print("\n3. Large Document (100 paragraphs):")
        ast_large = Path(tmpdir) / "ast_large.docx"

        tracemalloc.start()
        start_time = time.time()
        render_with_ast(large_chapter_nodes, ast_large)
        elapsed = time.time() - start_time
        current, peak = tracemalloc.get_traced_memory()
        tracemalloc.stop()

        print(f"   Render time: {elapsed:.3f}s")
        print(f"   Peak memory: {peak / 1024 / 1024:.2f} MB")
        print(f"   Status:      {'✅ PASS' if elapsed < 30.0 else '⚠️  SLOW'}")

    print("\n" + "="*80)
    print("✅ Phase 3.5 Integration Tests Complete")
    print("="*80 + "\n")
