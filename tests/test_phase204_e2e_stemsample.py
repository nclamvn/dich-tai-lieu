#!/usr/bin/env python3
# -*- coding: utf-8 -*-

"""
Phase 2.0.4 E2E Black-box Test - Stemsample.pdf with OMML

This test validates the COMPLETE pipeline end-to-end:
- CLI translate_pdf.py with --equation-rendering omml
- Actual DOCX file output
- OMML equations in XML (not plain text LaTeX)
- Adequate content length (not truncated translation)

Philosophy: "nhìn file mới tin được" - only believe when you see the file.
"""

import subprocess
import os
from zipfile import ZipFile
from pathlib import Path

import pytest
from docx import Document


PROJECT_ROOT = Path(__file__).parent.parent
PDF_PATH = PROJECT_ROOT / "Stemsample.pdf"


def test_stemsample_e2e_academic_omml():
    """
    E2E test: Stemsample.pdf → Academic DOCX with OMML equations

    Validates:
    1. Output file is created
    2. Adequate content (not truncated translation)
    3. OMML math markup present in document.xml
    """

    # Expected output path (based on CLI behavior)
    # The actual filename may vary - let's find it dynamically
    output_dir = PROJECT_ROOT / "output"

    # 1. Clean up any previous output files for Stemsample
    if output_dir.exists():
        for f in output_dir.glob("Stemsample*Vietnamese*.docx"):
            f.unlink()

    # 2. Run CLI end-to-end with OMML mode
    cmd = [
        "python3", "translate_pdf.py", str(PDF_PATH),
        "--domain", "stem",
        "--layout-mode", "academic",
        "--equation-rendering", "omml"
    ]

    print(f"\n{'='*70}")
    print(f"Running E2E command:")
    print(f"  {' '.join(cmd)}")
    print(f"{'='*70}\n")

    result = subprocess.run(
        cmd,
        cwd=PROJECT_ROOT,
        capture_output=True,
        text=True,
        timeout=300  # 5 minutes max
    )

    print("STDOUT:")
    print(result.stdout)

    if result.returncode != 0:
        print("STDERR:")
        print(result.stderr)
        pytest.fail(f"CLI command failed with exit code {result.returncode}")

    # 3. Find the output file
    output_files = list(output_dir.glob("Stemsample*Vietnamese*.docx"))

    assert len(output_files) > 0, \
        f"No output DOCX file found in {output_dir}"

    # Use the most recent file
    out_path = max(output_files, key=lambda p: p.stat().st_mtime)

    print(f"\n✅ Found output: {out_path.name}")
    print(f"   File size: {out_path.stat().st_size:,} bytes\n")

    # 4. Check file has adequate content using python-docx
    doc = Document(str(out_path))
    paragraphs = [p.text for p in doc.paragraphs]
    all_text = "\n".join(paragraphs)

    print(f"📊 Content metrics:")
    print(f"   Paragraphs: {len(paragraphs)}")
    print(f"   Total text length: {len(all_text):,} chars")
    print(f"   Tables: {len(doc.tables)}")

    # Heuristic validation: Stemsample should have substantial content
    assert len(paragraphs) > 50, \
        f"Too few paragraphs ({len(paragraphs)}) - output likely broken or incomplete"

    assert len(all_text) > 5000, \
        f"Too little text ({len(all_text)} chars) - translation likely incomplete"

    print(f"   ✅ Content check PASSED\n")

    # 5. CRITICAL: Check for OMML markup in document.xml
    print(f"🔬 Checking for OMML equations in XML...")

    with ZipFile(str(out_path)) as z:
        doc_xml = z.read("word/document.xml").decode("utf-8")

    # OMML math elements:
    # - <m:oMath> - inline equation
    # - <m:oMathPara> - display equation
    has_omml_inline = "<m:oMath" in doc_xml
    has_omml_display = "<m:oMathPara" in doc_xml

    # Count occurrences
    omml_count = doc_xml.count("<m:oMath")
    omml_para_count = doc_xml.count("<m:oMathPara")

    print(f"   Found OMML elements:")
    print(f"     <m:oMath> (inline):   {omml_count}")
    print(f"     <m:oMathPara> (display): {omml_para_count}")

    assert has_omml_inline or has_omml_display, \
        "❌ CRITICAL FAILURE: No OMML math found in document.xml!\n" \
        "   Equations are NOT rendered as native Word math.\n" \
        "   They are likely plain text or LaTeX.\n" \
        "   Phase 2.0.4 OMML pipeline is NOT working!"

    # Stemsample has many equations - expect at least 10 OMML elements
    total_omml = omml_count + omml_para_count
    assert total_omml >= 10, \
        f"Only {total_omml} OMML elements found - expected at least 10 for Stemsample.pdf"

    print(f"   ✅ OMML validation PASSED")
    print(f"\n{'='*70}")
    print(f"🎉 Phase 2.0.4 E2E Test: PASS")
    print(f"{'='*70}\n")

    # Return output path for manual inspection if needed
    return str(out_path)


if __name__ == "__main__":
    # Allow running test directly
    output = test_stemsample_e2e_academic_omml()
    print(f"\n✅ Test completed successfully!")
    print(f"📄 Output file: {output}")
    print(f"\n💡 Open this file in Microsoft Word to verify:")
    print(f"   1. Click on equations - should be editable with Equation Editor (OMML)")
    print(f"   2. Content should be ~27 pages with full translation")
    print(f"   3. Academic structure: theorems, proofs, etc.")
