"""
Phase 2.0.3b - OMML Integration Tests

Comprehensive test suite for OMML equation rendering in Academic DOCX Builder.

Test Coverage:
1. Config defaults (latex_text)
2. Config override (omml)
3. Pandoc availability fallback
4. Conversion success/failure handling
5. OMML injection success
6. No double-display (OMML + text)
7. Various LaTeX patterns (inline, display, complex)
8. Edge cases (empty, whitespace, invalid)
9. Fallback safety (original text preserved)
10. Simple mode unchanged
11. Academic mode with latex_text unchanged
12. Academic mode with omml → OMML rendering
13. Multiple equations in document
14. Vietnamese text compatibility
"""

import pytest
import tempfile
from pathlib import Path
from docx import Document
from lxml import etree
from unittest.mock import patch, MagicMock

from core.export.docx_academic_builder import (
    build_academic_docx,
    AcademicLayoutConfig,
    _format_equation_block
)
from core.structure.semantic_model import DocNode, DocNodeType, DocNodeList
from core.rendering import omml_converter


# ============================================================================
# TEST 1: Config Defaults
# ============================================================================

def test_academic_config_defaults():
    """Test that AcademicLayoutConfig has correct default for equation_rendering_mode."""
    config = AcademicLayoutConfig()

    # Default must be "latex_text" for backward compatibility
    assert config.equation_rendering_mode == "latex_text"

    # Phase 2.0.5: Updated defaults for better academic appearance
    assert config.font_name == "Cambria"  # Changed from "Times New Roman" for better math rendering
    assert config.font_size == 11  # Changed from 12 to 11pt (academic standard)
    assert config.line_spacing == 1.15  # Changed from 1.5 to 1.15 (modern academic standard)


# ============================================================================
# TEST 2: Config Override
# ============================================================================

def test_academic_config_omml_override():
    """Test that equation_rendering_mode can be set to 'omml'."""
    config = AcademicLayoutConfig(equation_rendering_mode="omml")

    assert config.equation_rendering_mode == "omml"


# ============================================================================
# TEST 3: Pandoc Availability Check
# ============================================================================

def test_omml_fallback_when_pandoc_unavailable():
    """Test that OMML rendering falls back to LaTeX text when pandoc unavailable."""
    # Mock pandoc unavailable
    with patch('core.rendering.omml_converter.is_pandoc_available', return_value=False):
        config = AcademicLayoutConfig(equation_rendering_mode="omml")

        nodes = [
            DocNode(
                node_type=DocNodeType.EQUATION_BLOCK,
                text="$x^2 + y^2 = z^2$"
            )
        ]

        with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as f:
            output_path = f.name

        try:
            build_academic_docx(nodes, output_path, config)

            # Verify DOCX created
            assert Path(output_path).exists()

            # Open and check content
            doc = Document(output_path)
            assert len(doc.paragraphs) > 0

            # Should have LaTeX text (fallback)
            para_text = doc.paragraphs[0].text
            assert "$x^2 + y^2 = z^2$" in para_text

            # Should NOT have OMML elements
            para_elem = doc.paragraphs[0]._element
            omml_ns = '{http://schemas.openxmlformats.org/officeDocument/2006/math}'
            omath_elems = para_elem.findall(f'.//{omml_ns}oMath')
            assert len(omath_elems) == 0

        finally:
            Path(output_path).unlink(missing_ok=True)


# ============================================================================
# TEST 4: OMML Conversion Success
# ============================================================================

def test_omml_rendering_success():
    """Test successful OMML rendering when pandoc available."""
    # Only run if pandoc is actually available
    if not omml_converter.is_pandoc_available():
        pytest.skip("Pandoc not available - skipping OMML rendering test")

    config = AcademicLayoutConfig(equation_rendering_mode="omml")

    nodes = [
        DocNode(
            node_type=DocNodeType.EQUATION_BLOCK,
            text="$x^2$",
            
        )
    ]

    with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as f:
        output_path = f.name

    try:
        build_academic_docx(nodes, output_path, config)

        # Verify DOCX created
        assert Path(output_path).exists()

        # Open and check for OMML elements
        doc = Document(output_path)
        assert len(doc.paragraphs) > 0

        para_elem = doc.paragraphs[0]._element
        omml_ns = '{http://schemas.openxmlformats.org/officeDocument/2006/math}'
        omath_elems = para_elem.findall(f'.//{omml_ns}oMath')

        # Should have OMML element
        assert len(omath_elems) > 0

        # Should NOT have duplicate text content
        para_text = doc.paragraphs[0].text
        # OMML elements don't have visible text, so paragraph should be empty or whitespace
        assert para_text.strip() == "" or "$x^2$" not in para_text

    finally:
        Path(output_path).unlink(missing_ok=True)


# ============================================================================
# TEST 5: Conversion Failure Fallback
# ============================================================================

def test_omml_conversion_failure_fallback():
    """Test that conversion failure falls back to LaTeX text gracefully."""
    # Mock latex_to_omml to return None (conversion failed)
    with patch('core.rendering.omml_converter.is_pandoc_available', return_value=True):
        with patch('core.rendering.omml_converter.latex_to_omml', return_value=None):
            config = AcademicLayoutConfig(equation_rendering_mode="omml")

            nodes = [
                DocNode(
                    node_type=DocNodeType.EQUATION_BLOCK,
                    text="$\\invalid{command}$",
                    
                )
            ]

            with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as f:
                output_path = f.name

            try:
                # Should not crash
                build_academic_docx(nodes, output_path, config)

                # Verify DOCX created with fallback text
                doc = Document(output_path)
                para_text = doc.paragraphs[0].text
                assert "$\\invalid{command}$" in para_text

            finally:
                Path(output_path).unlink(missing_ok=True)


# ============================================================================
# TEST 6: OMML Injection Failure Fallback
# ============================================================================

def test_omml_injection_failure_fallback():
    """Test that injection failure falls back to LaTeX text."""
    # Mock successful conversion but failed injection
    fake_omml = '<m:oMath xmlns:m="http://schemas.openxmlformats.org/officeDocument/2006/math"></m:oMath>'

    with patch('core.rendering.omml_converter.is_pandoc_available', return_value=True):
        with patch('core.rendering.omml_converter.latex_to_omml', return_value=fake_omml):
            # Phase 2.1.1: Updated mock to use inject_omml_as_display (not inject_omml_into_paragraph)
            with patch('core.rendering.omml_converter.inject_omml_as_display', return_value=False):
                config = AcademicLayoutConfig(equation_rendering_mode="omml")

                nodes = [
                    DocNode(
                        node_type=DocNodeType.EQUATION_BLOCK,
                        text="$x^2$",

                    )
                ]

                with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as f:
                    output_path = f.name

                try:
                    build_academic_docx(nodes, output_path, config)

                    # Should fallback to LaTeX text
                    doc = Document(output_path)
                    para_text = doc.paragraphs[0].text
                    assert "$x^2$" in para_text

                finally:
                    Path(output_path).unlink(missing_ok=True)


# ============================================================================
# TEST 7: No Double-Display (OMML + Text)
# ============================================================================

def test_no_double_display_omml_and_text():
    """Test that successful OMML rendering doesn't also display LaTeX text."""
    if not omml_converter.is_pandoc_available():
        pytest.skip("Pandoc not available")

    config = AcademicLayoutConfig(equation_rendering_mode="omml")

    nodes = [
        DocNode(
            node_type=DocNodeType.EQUATION_BLOCK,
            text="$\\frac{a}{b}$",
            
        )
    ]

    with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as f:
        output_path = f.name

    try:
        build_academic_docx(nodes, output_path, config)
        doc = Document(output_path)

        para_elem = doc.paragraphs[0]._element
        omml_ns = '{http://schemas.openxmlformats.org/officeDocument/2006/math}'
        omath_elems = para_elem.findall(f'.//{omml_ns}oMath')

        if len(omath_elems) > 0:
            # OMML rendering succeeded - text should be empty
            para_text = doc.paragraphs[0].text
            assert "$\\frac{a}{b}$" not in para_text

    finally:
        Path(output_path).unlink(missing_ok=True)


# ============================================================================
# TEST 8: Various LaTeX Patterns
# ============================================================================

def test_various_latex_patterns():
    """Test OMML rendering with different LaTeX patterns."""
    if not omml_converter.is_pandoc_available():
        pytest.skip("Pandoc not available")

    config = AcademicLayoutConfig(equation_rendering_mode="omml")

    # Test inline, display, complex equations
    test_equations = [
        "$x^2$",  # Simple inline
        "$$\\int_0^\\infty e^{-x^2} dx$$",  # Display integral
        "$\\sum_{i=1}^n i = \\frac{n(n+1)}{2}$",  # Summation
        "$E = mc^2$",  # Famous equation
    ]

    for eq in test_equations:
        nodes = [
            DocNode(
                node_type=DocNodeType.EQUATION_BLOCK,
                text=eq
            )
        ]

        with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as f:
            output_path = f.name

        try:
            # Should not crash for any pattern
            build_academic_docx(nodes, output_path, config)
            assert Path(output_path).exists()

        finally:
            Path(output_path).unlink(missing_ok=True)


# ============================================================================
# TEST 9: Edge Cases (Empty, Whitespace)
# ============================================================================

def test_edge_cases_empty_whitespace():
    """Test handling of edge cases: empty equations, whitespace."""
    config = AcademicLayoutConfig(equation_rendering_mode="omml")

    edge_cases = [
        "",  # Empty
        "   ",  # Whitespace only
        "$   $",  # Empty delimiters
    ]

    for eq in edge_cases:
        nodes = [
            DocNode(
                node_type=DocNodeType.EQUATION_BLOCK,
                text=eq
            )
        ]

        with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as f:
            output_path = f.name

        try:
            # Should not crash
            build_academic_docx(nodes, output_path, config)
            assert Path(output_path).exists()

        finally:
            Path(output_path).unlink(missing_ok=True)


# ============================================================================
# TEST 10: Backward Compatibility - Default Mode Unchanged
# ============================================================================

def test_backward_compatibility_default_mode():
    """Test that default config renders LaTeX text (backward compatible)."""
    # Use default config (equation_rendering_mode = "latex_text")
    config = AcademicLayoutConfig()

    nodes = [
        DocNode(
            node_type=DocNodeType.EQUATION_BLOCK,
            text="$x^2 + y^2 = z^2$",
            
        )
    ]

    with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as f:
        output_path = f.name

    try:
        build_academic_docx(nodes, output_path, config)

        # Verify LaTeX text is rendered
        doc = Document(output_path)
        para_text = doc.paragraphs[0].text
        assert "$x^2 + y^2 = z^2$" in para_text

        # Verify NO OMML elements
        para_elem = doc.paragraphs[0]._element
        omml_ns = '{http://schemas.openxmlformats.org/officeDocument/2006/math}'
        omath_elems = para_elem.findall(f'.//{omml_ns}oMath')
        assert len(omath_elems) == 0

    finally:
        Path(output_path).unlink(missing_ok=True)


# ============================================================================
# TEST 11: Academic Mode with LaTeX Text
# ============================================================================

def test_academic_mode_with_latex_text():
    """Test academic document with latex_text rendering (existing behavior)."""
    config = AcademicLayoutConfig(equation_rendering_mode="latex_text")

    nodes = [
        DocNode(
            node_type=DocNodeType.CHAPTER,
            text="Chapter 1: Introduction",
            level=1
        ),
        DocNode(
            node_type=DocNodeType.PARAGRAPH,
            text="Consider the equation:"
        ),
        DocNode(
            node_type=DocNodeType.EQUATION_BLOCK,
            text="$x^2 + y^2 = r^2$",
            
        ),
        DocNode(
            node_type=DocNodeType.PARAGRAPH,
            text="This represents a circle."
        )
    ]

    with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as f:
        output_path = f.name

    try:
        build_academic_docx(nodes, output_path, config)

        doc = Document(output_path)

        # Verify structure: heading, para, equation (centered text), para
        assert len(doc.paragraphs) >= 3

        # Find equation paragraph (should have centered alignment)
        equation_para = None
        for para in doc.paragraphs:
            if "$x^2 + y^2 = r^2$" in para.text:
                equation_para = para
                break

        assert equation_para is not None
        assert equation_para.alignment == 1  # WD_ALIGN_PARAGRAPH.CENTER = 1

    finally:
        Path(output_path).unlink(missing_ok=True)


# ============================================================================
# TEST 12: Academic Mode with OMML
# ============================================================================

def test_academic_mode_with_omml():
    """Test academic document with OMML rendering."""
    if not omml_converter.is_pandoc_available():
        pytest.skip("Pandoc not available")

    config = AcademicLayoutConfig(equation_rendering_mode="omml")

    nodes = [
        DocNode(
            node_type=DocNodeType.CHAPTER,
            text="Chapter 1: Algebra",
            level=1
        ),
        DocNode(
            node_type=DocNodeType.THEOREM,
            text="Pythagorean Theorem",
            title="Theorem 1.1"
        ),
        DocNode(
            node_type=DocNodeType.EQUATION_BLOCK,
            text="$a^2 + b^2 = c^2$",
            
        ),
        DocNode(
            node_type=DocNodeType.PROOF,
            text="Proof by construction...",
            title="Proof"
        )
    ]

    with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as f:
        output_path = f.name

    try:
        build_academic_docx(nodes, output_path, config)

        doc = Document(output_path)
        assert len(doc.paragraphs) >= 4

        # Find equation paragraph
        found_omml = False
        for para in doc.paragraphs:
            para_elem = para._element
            omml_ns = '{http://schemas.openxmlformats.org/officeDocument/2006/math}'
            omath_elems = para_elem.findall(f'.//{omml_ns}oMath')
            if len(omath_elems) > 0:
                found_omml = True
                break

        # Should have OMML rendering
        assert found_omml

    finally:
        Path(output_path).unlink(missing_ok=True)


# ============================================================================
# TEST 13: Multiple Equations in Document
# ============================================================================

def test_multiple_equations_in_document():
    """Test document with multiple equations - all should render correctly."""
    if not omml_converter.is_pandoc_available():
        pytest.skip("Pandoc not available")

    config = AcademicLayoutConfig(equation_rendering_mode="omml")

    nodes = [
        DocNode(
            node_type=DocNodeType.PARAGRAPH,
            text="First equation:"
        ),
        DocNode(
            node_type=DocNodeType.EQUATION_BLOCK,
            text="$x = 1$",
            
        ),
        DocNode(
            node_type=DocNodeType.PARAGRAPH,
            text="Second equation:"
        ),
        DocNode(
            node_type=DocNodeType.EQUATION_BLOCK,
            text="$y = 2$",
            
        ),
        DocNode(
            node_type=DocNodeType.PARAGRAPH,
            text="Third equation:"
        ),
        DocNode(
            node_type=DocNodeType.EQUATION_BLOCK,
            text="$z = x + y$",
            
        )
    ]

    with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as f:
        output_path = f.name

    try:
        build_academic_docx(nodes, output_path, config)

        doc = Document(output_path)

        # Count OMML elements
        omml_count = 0
        omml_ns = '{http://schemas.openxmlformats.org/officeDocument/2006/math}'

        for para in doc.paragraphs:
            para_elem = para._element
            omath_elems = para_elem.findall(f'.//{omml_ns}oMath')
            omml_count += len(omath_elems)

        # Should have 3 OMML elements (one per equation)
        assert omml_count == 3

    finally:
        Path(output_path).unlink(missing_ok=True)


# ============================================================================
# TEST 14: Vietnamese Text Compatibility
# ============================================================================

def test_vietnamese_text_compatibility():
    """Test that OMML rendering works with Vietnamese text in document."""
    if not omml_converter.is_pandoc_available():
        pytest.skip("Pandoc not available")

    config = AcademicLayoutConfig(equation_rendering_mode="omml")

    nodes = [
        DocNode(
            node_type=DocNodeType.PARAGRAPH,
            text="Xét phương trình bậc hai:"
        ),
        DocNode(
            node_type=DocNodeType.EQUATION_BLOCK,
            text="$ax^2 + bx + c = 0$",
            
        ),
        DocNode(
            node_type=DocNodeType.PARAGRAPH,
            text="Nghiệm của phương trình được tính bởi công thức:"
        ),
        DocNode(
            node_type=DocNodeType.EQUATION_BLOCK,
            text="$x = \\frac{-b \\pm \\sqrt{b^2 - 4ac}}{2a}$",
            
        )
    ]

    with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as f:
        output_path = f.name

    try:
        build_academic_docx(nodes, output_path, config)

        doc = Document(output_path)

        # Verify Vietnamese text preserved
        found_vietnamese = False
        for para in doc.paragraphs:
            if "Xét phương trình" in para.text or "Nghiệm của phương trình" in para.text:
                found_vietnamese = True

        assert found_vietnamese

        # Verify OMML elements present
        omml_count = 0
        omml_ns = '{http://schemas.openxmlformats.org/officeDocument/2006/math}'

        for para in doc.paragraphs:
            para_elem = para._element
            omath_elems = para_elem.findall(f'.//{omml_ns}oMath')
            omml_count += len(omath_elems)

        assert omml_count == 2  # Two equations

    finally:
        Path(output_path).unlink(missing_ok=True)


# ============================================================================
# TEST 15: Exception Handling in OMML Pipeline
# ============================================================================

def test_exception_handling_in_omml_pipeline():
    """Test that unexpected exceptions in OMML pipeline don't crash the build."""
    # Mock latex_to_omml to raise an exception
    with patch('core.rendering.omml_converter.is_pandoc_available', return_value=True):
        with patch('core.rendering.omml_converter.latex_to_omml', side_effect=RuntimeError("Simulated error")):
            config = AcademicLayoutConfig(equation_rendering_mode="omml")

            nodes = [
                DocNode(
                    node_type=DocNodeType.EQUATION_BLOCK,
                    text="$x^2$",
                    
                )
            ]

            with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as f:
                output_path = f.name

            try:
                # Should not crash - should fallback to LaTeX text
                build_academic_docx(nodes, output_path, config)

                doc = Document(output_path)
                para_text = doc.paragraphs[0].text
                # Should have fallback text
                assert "$x^2$" in para_text

            finally:
                Path(output_path).unlink(missing_ok=True)


# ============================================================================
# TEST 16: Original LaTeX Text Never Mutated
# ============================================================================

def test_original_latex_never_mutated():
    """Test that original node.text is never modified during OMML rendering."""
    if not omml_converter.is_pandoc_available():
        pytest.skip("Pandoc not available")

    config = AcademicLayoutConfig(equation_rendering_mode="omml")

    original_text = "$x^2 + y^2 = z^2$"
    node = DocNode(
        node_type=DocNodeType.EQUATION_BLOCK,
        text=original_text
    )

    nodes = [node]

    with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as f:
        output_path = f.name

    try:
        build_academic_docx(nodes, output_path, config)

        # Verify node.text unchanged
        assert node.text == original_text
        # Original text preserved in node.text

    finally:
        Path(output_path).unlink(missing_ok=True)
