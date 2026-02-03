"""
Output File Validation Tests for AI Publisher Pro

Validates the structure and content of generated documents:
- DOCX structure (cover, title page, TOC, content)
- Image presence and captions
- PDF page structure

Run: pytest tests/e2e/test_output_validation.py -v
"""

import pytest
from pathlib import Path


# ============================================================
# TC-OUT-01: DOCX Structure Validation
# ============================================================

def test_docx_has_cover_page(output_docx_with_cover):
    """DOCX should have cover image as first element"""
    from docx import Document

    doc = Document(str(output_docx_with_cover))

    # Check for images in first few paragraphs
    has_image = False
    for para in doc.paragraphs[:5]:
        # Check for inline shapes (images)
        drawings = para._element.xpath('.//w:drawing')
        if drawings:
            has_image = True
            break

    assert has_image, "Cover page should contain an image"


# ============================================================
# TC-OUT-02: DOCX Has Title Page
# ============================================================

def test_docx_has_title_page(output_docx_with_cover):
    """DOCX should have title page after cover"""
    from docx import Document

    doc = Document(str(output_docx_with_cover))

    # Find title (usually large, bold, centered)
    title_found = False
    for para in doc.paragraphs[:15]:  # Check first 15 paragraphs
        style_name = para.style.name if para.style else ""
        if style_name in ['Title', 'Heading 1']:
            title_found = True
            break
        # Also check for bold text that could be a title
        if para.runs and len(para.runs) > 0:
            if para.runs[0].bold and len(para.text) > 5:
                title_found = True
                break

    assert title_found, "Title page should exist after cover"


# ============================================================
# TC-OUT-03: DOCX Has Embedded Images
# ============================================================

def test_docx_has_embedded_images(output_docx_with_images):
    """DOCX should contain embedded images from PDF"""
    from docx import Document

    doc = Document(str(output_docx_with_images))

    # Count images
    image_count = 0
    for para in doc.paragraphs:
        drawings = para._element.xpath('.//w:drawing')
        image_count += len(drawings)

    assert image_count > 0, "DOCX should contain embedded images"


# ============================================================
# TC-OUT-04: DOCX Image Captions
# ============================================================

def test_docx_image_captions(output_docx_with_images):
    """Images should have captions (Hình 1, Hình 2, etc.)"""
    from docx import Document

    doc = Document(str(output_docx_with_images))

    caption_found = False
    for para in doc.paragraphs:
        text = para.text.strip()
        # Check for Vietnamese captions
        if text.startswith("Hình ") and (":" in text or text[-1].isdigit()):
            caption_found = True
            break

    assert caption_found, "Images should have Vietnamese captions (Hình X)"


# ============================================================
# TC-OUT-05: PDF Has Cover Page
# ============================================================

def test_pdf_has_cover_page(output_pdf_with_cover):
    """PDF should have cover image on first page"""
    try:
        import fitz  # PyMuPDF
    except ImportError:
        pytest.skip("PyMuPDF not available")

    doc = fitz.open(str(output_pdf_with_cover))

    first_page = doc[0]
    images = first_page.get_images()

    assert len(images) > 0, "First page should have cover image"

    doc.close()


# ============================================================
# TC-OUT-06: PDF Page Count
# ============================================================

def test_pdf_page_count(output_pdf_with_cover):
    """PDF should have at least 3 pages (cover, title, content)"""
    try:
        import fitz  # PyMuPDF
    except ImportError:
        pytest.skip("PyMuPDF not available")

    doc = fitz.open(str(output_pdf_with_cover))

    assert doc.page_count >= 3, "PDF should have cover, title, and content pages"

    doc.close()


# ============================================================
# TC-OUT-07: Document Order Validation
# ============================================================

def test_document_order(output_docx_with_cover):
    """
    Document order should be:
    1. Cover Image (full page)
    2. Title Page
    3. TOC (optional)
    4. Content
    """
    from docx import Document

    doc = Document(str(output_docx_with_cover))

    elements = []
    for i, para in enumerate(doc.paragraphs[:25]):
        # Check for images
        has_image = len(para._element.xpath('.//w:drawing')) > 0

        if has_image and i < 5:
            elements.append("COVER")
        elif para.style and para.style.name in ['Title', 'Heading 1']:
            elements.append("TITLE")
        elif "Table of Contents" in para.text or "Mục lục" in para.text:
            elements.append("TOC")
        elif para.text.strip() and "Chapter" in para.text:
            elements.append("CONTENT")
            break

    # Verify COVER comes before TITLE
    if "COVER" in elements and "TITLE" in elements:
        cover_idx = elements.index("COVER")
        title_idx = elements.index("TITLE")
        assert cover_idx < title_idx, f"Cover should come before title. Order: {elements}"


# ============================================================
# TC-OUT-08: DOCX File Size Reasonable
# ============================================================

def test_docx_file_size(output_docx_with_cover):
    """DOCX with cover should have reasonable file size"""
    file_size = output_docx_with_cover.stat().st_size

    # Should be at least 10KB (has images)
    assert file_size > 10000, "DOCX with cover should be > 10KB"

    # Should be less than 50MB (not bloated)
    assert file_size < 50 * 1024 * 1024, "DOCX should be < 50MB"


# ============================================================
# TC-OUT-09: DOCX Has Valid Structure
# ============================================================

def test_docx_valid_structure(output_docx_with_cover):
    """DOCX should have valid internal structure"""
    from docx import Document

    # Should not raise exception
    doc = Document(str(output_docx_with_cover))

    # Should have paragraphs
    assert len(doc.paragraphs) > 0, "Document should have paragraphs"

    # Should have sections
    assert len(doc.sections) > 0, "Document should have sections"


# ============================================================
# TC-OUT-10: Image Quality Check
# ============================================================

def test_image_quality(output_docx_with_images):
    """Embedded images should maintain reasonable quality"""
    from docx import Document
    import zipfile
    import io
    from PIL import Image

    # DOCX is a ZIP file
    with zipfile.ZipFile(str(output_docx_with_images)) as zf:
        # Find images in the docx
        image_files = [f for f in zf.namelist() if f.startswith('word/media/')]

        for img_file in image_files:
            img_data = zf.read(img_file)
            img = Image.open(io.BytesIO(img_data))

            # Check image has reasonable dimensions
            assert img.width >= 50, f"Image {img_file} width too small"
            assert img.height >= 50, f"Image {img_file} height too small"


# ============================================================
# TC-OUT-11: No Corrupted Images
# ============================================================

def test_no_corrupted_images(output_docx_with_images):
    """Embedded images should not be corrupted"""
    import zipfile
    import io
    from PIL import Image

    with zipfile.ZipFile(str(output_docx_with_images)) as zf:
        image_files = [f for f in zf.namelist() if f.startswith('word/media/')]

        for img_file in image_files:
            img_data = zf.read(img_file)

            # Should be able to open without error
            try:
                img = Image.open(io.BytesIO(img_data))
                img.verify()  # Verify image integrity
            except Exception as e:
                pytest.fail(f"Image {img_file} is corrupted: {e}")
