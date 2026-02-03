"""
E2E Test Fixtures and Configuration

Provides shared fixtures for E2E tests.
"""

import pytest
import base64
import tempfile
from pathlib import Path
from PIL import Image
import io

# Base paths
PROJECT_ROOT = Path(__file__).parent.parent.parent
FIXTURES_DIR = PROJECT_ROOT / "tests" / "fixtures"
OUTPUTS_DIR = PROJECT_ROOT / "tests" / "outputs"

# Server configuration
BASE_URL = "http://localhost:3000"


# ============================================================
# Helper Functions
# ============================================================

def create_simple_pdf(path: Path, with_images: bool = False) -> Path:
    """Create a simple test PDF."""
    try:
        from reportlab.lib.pagesizes import letter
        from reportlab.pdfgen import canvas
        from reportlab.lib.units import inch

        c = canvas.Canvas(str(path), pagesize=letter)

        # Page 1: Title
        c.setFont("Helvetica-Bold", 24)
        c.drawString(1 * inch, 10 * inch, "Test Document")
        c.setFont("Helvetica", 12)
        c.drawString(1 * inch, 9.5 * inch, "This is a sample PDF for testing.")

        # Page 1: Content
        c.drawString(1 * inch, 8.5 * inch, "Lorem ipsum dolor sit amet, consectetur adipiscing elit.")
        c.drawString(1 * inch, 8.2 * inch, "Sed do eiusmod tempor incididunt ut labore et dolore magna aliqua.")

        if with_images:
            # Create a simple colored rectangle as "image"
            c.setFillColorRGB(0.2, 0.4, 0.8)
            c.rect(1 * inch, 5 * inch, 3 * inch, 2 * inch, fill=1)
            c.setFillColorRGB(0, 0, 0)
            c.drawString(1 * inch, 4.7 * inch, "Figure 1: Sample Image")

        c.showPage()

        # Page 2: More content
        c.setFont("Helvetica-Bold", 18)
        c.drawString(1 * inch, 10 * inch, "Chapter 2")
        c.setFont("Helvetica", 12)
        c.drawString(1 * inch, 9.5 * inch, "More content for testing translation pipeline.")
        c.drawString(1 * inch, 9.2 * inch, "This page tests multi-page document handling.")

        if with_images:
            c.setFillColorRGB(0.8, 0.3, 0.2)
            c.rect(1 * inch, 6 * inch, 2.5 * inch, 1.5 * inch, fill=1)
            c.setFillColorRGB(0, 0, 0)
            c.drawString(1 * inch, 5.7 * inch, "Figure 2: Another Image")

        c.save()
        return path

    except ImportError:
        pytest.skip("ReportLab not available for PDF creation")


def create_test_image(width: int, height: int, color: tuple = (30, 60, 114)) -> bytes:
    """Create a test image with specified dimensions."""
    img = Image.new('RGB', (width, height), color=color)

    # Add some visual elements
    from PIL import ImageDraw
    draw = ImageDraw.Draw(img)

    # Title area (white rectangle)
    draw.rectangle([width//10, height//4, width*9//10, height//3], fill=(255, 255, 255))

    # Author area (gold rectangle)
    draw.rectangle([width//4, height//2 - 30, width*3//4, height//2 + 30], fill=(255, 200, 100))

    # Bottom area (light blue)
    draw.rectangle([width//5, height*4//5, width*4//5, height*4//5 + 50], fill=(100, 200, 255))

    buffer = io.BytesIO()
    img.save(buffer, format='PNG')
    return buffer.getvalue()


# ============================================================
# Session-Scoped Fixtures (Create once)
# ============================================================

@pytest.fixture(scope="session")
def ensure_fixtures_dir():
    """Ensure fixtures directory exists."""
    FIXTURES_DIR.mkdir(parents=True, exist_ok=True)
    OUTPUTS_DIR.mkdir(parents=True, exist_ok=True)
    return FIXTURES_DIR


@pytest.fixture(scope="session")
def sample_text_only_pdf(ensure_fixtures_dir) -> Path:
    """Create/get a text-only PDF fixture."""
    path = FIXTURES_DIR / "sample_text_only.pdf"
    if not path.exists():
        create_simple_pdf(path, with_images=False)
    return path


@pytest.fixture(scope="session")
def sample_with_images_pdf(ensure_fixtures_dir) -> Path:
    """Create/get a PDF with images fixture."""
    path = FIXTURES_DIR / "sample_with_images.pdf"
    if not path.exists():
        create_simple_pdf(path, with_images=True)
    return path


@pytest.fixture(scope="session")
def sample_cover_png(ensure_fixtures_dir) -> Path:
    """Create/get a cover image (PNG) fixture."""
    path = FIXTURES_DIR / "sample_cover.png"
    if not path.exists():
        img_data = create_test_image(1200, 1800)
        path.write_bytes(img_data)
    return path


@pytest.fixture(scope="session")
def sample_cover_jpg(ensure_fixtures_dir) -> Path:
    """Create/get a cover image (JPG) fixture."""
    path = FIXTURES_DIR / "sample_cover.jpg"
    if not path.exists():
        img = Image.new('RGB', (1200, 1800), color=(30, 60, 114))
        img.save(path, format='JPEG', quality=90)
    return path


# ============================================================
# Function-Scoped Fixtures
# ============================================================

@pytest.fixture
def sample_pdf_text_only_b64(sample_text_only_pdf) -> str:
    """PDF without images as base64."""
    return base64.b64encode(sample_text_only_pdf.read_bytes()).decode()


@pytest.fixture
def sample_pdf_with_images_b64(sample_with_images_pdf) -> str:
    """PDF with images as base64."""
    return base64.b64encode(sample_with_images_pdf.read_bytes()).decode()


@pytest.fixture
def sample_cover_image_b64(sample_cover_png) -> str:
    """Cover image as data URI."""
    data = base64.b64encode(sample_cover_png.read_bytes()).decode()
    return f"data:image/png;base64,{data}"


@pytest.fixture
def temp_output_dir(tmp_path) -> Path:
    """Temporary output directory for test files."""
    return tmp_path


# ============================================================
# Output File Fixtures (for validation tests)
# ============================================================

@pytest.fixture
def output_docx_with_cover(sample_cover_png, temp_output_dir) -> Path:
    """Create a DOCX with cover for validation tests."""
    from docx import Document
    from core.image_embedding import CoverEmbedder
    from core.export.docx_styles import StyleManager
    from core.export.docx_front_matter import FrontMatterGenerator

    doc = Document()
    style_manager = StyleManager(theme_name='professional')

    # Add cover
    cover_b64 = base64.b64encode(sample_cover_png.read_bytes()).decode()
    fm_gen = FrontMatterGenerator(doc, style_manager)
    fm_gen.generate_cover_page(cover_b64)

    # Add title page
    fm_gen.generate_title_page({
        'title': 'Test Document',
        'author': 'Test Author',
        'subject': 'E2E Testing'
    })

    # Add content
    doc.add_heading('Chapter 1: Introduction', level=1)
    doc.add_paragraph('This is test content for E2E validation.')

    output_path = temp_output_dir / "test_with_cover.docx"
    doc.save(str(output_path))
    return output_path


@pytest.fixture
def output_docx_with_images(sample_with_images_pdf, temp_output_dir) -> Path:
    """Create a DOCX with embedded images for validation tests."""
    from docx import Document
    from core.image_embedding import ImageExtractor, DocxImageEmbedder

    doc = Document()
    doc.add_heading('Document with Images', level=1)
    doc.add_paragraph('This document contains embedded images.')

    # Try to extract images from PDF
    try:
        extractor = ImageExtractor()
        images = extractor.extract_from_pdf(str(sample_with_images_pdf))

        if images:
            embedder = DocxImageEmbedder()
            embedder.embed_images(doc, images, with_captions=True)
    except Exception:
        # Create a simple image if extraction fails
        img_data = create_test_image(400, 300)
        from core.image_embedding import ImageBlock, ImageFormat
        img_block = ImageBlock(
            image_data=img_data,
            format=ImageFormat.PNG,
            width_px=400,
            height_px=300
        )
        embedder = DocxImageEmbedder()
        embedder.embed_image(doc, img_block)

    output_path = temp_output_dir / "test_with_images.docx"
    doc.save(str(output_path))
    return output_path


@pytest.fixture
def output_pdf_with_cover(sample_cover_png, temp_output_dir) -> Path:
    """Create a PDF with cover for validation tests."""
    try:
        from reportlab.lib.pagesizes import letter
        from reportlab.pdfgen import canvas
        from reportlab.lib.units import inch

        output_path = temp_output_dir / "test_with_cover.pdf"
        c = canvas.Canvas(str(output_path), pagesize=letter)

        # Page 1: Cover
        c.drawImage(str(sample_cover_png), 0.5*inch, 1*inch,
                   width=7.5*inch, height=9*inch, preserveAspectRatio=True)
        c.showPage()

        # Page 2: Title
        c.setFont("Helvetica-Bold", 24)
        c.drawString(1*inch, 10*inch, "Test Document")
        c.showPage()

        # Page 3: Content
        c.setFont("Helvetica", 12)
        c.drawString(1*inch, 10*inch, "Chapter 1: Introduction")
        c.drawString(1*inch, 9.5*inch, "This is test content.")
        c.save()

        return output_path

    except ImportError:
        pytest.skip("ReportLab not available")


# ============================================================
# Server Fixtures
# ============================================================

@pytest.fixture
def server_url():
    """Return the server base URL."""
    return BASE_URL


@pytest.fixture
def check_server_running(server_url):
    """Check if server is running, skip if not."""
    import httpx
    try:
        response = httpx.get(f"{server_url}/health", timeout=5.0)
        if response.status_code != 200:
            pytest.skip("Server not running or unhealthy")
    except Exception:
        pytest.skip("Server not running")
