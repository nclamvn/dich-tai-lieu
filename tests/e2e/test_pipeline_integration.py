"""
Pipeline Integration Tests for AI Publisher Pro

Tests the image embedding pipeline components:
- Image extraction from PDF
- DOCX image embedding
- PDF image embedding
- Cover image embedding

Run: pytest tests/e2e/test_pipeline_integration.py -v
"""

import pytest
from pathlib import Path
import base64


# ============================================================
# TC-PIPE-01: Image Extraction
# ============================================================

def test_image_extraction(sample_with_images_pdf):
    """ImageExtractor should extract images from PDF"""
    from core.image_embedding import ImageExtractor

    if not sample_with_images_pdf.exists():
        pytest.skip("Sample PDF not found")

    extractor = ImageExtractor()
    images = extractor.extract_from_pdf(str(sample_with_images_pdf))

    assert isinstance(images, list)

    # If images were extracted, verify structure
    for img in images:
        assert hasattr(img, 'image_data')
        assert hasattr(img, 'format')
        assert hasattr(img, 'width_px')
        assert hasattr(img, 'height_px')
        assert len(img.image_data) > 0


# ============================================================
# TC-PIPE-02: Image Extraction with Config
# ============================================================

def test_image_extraction_with_config(sample_with_images_pdf):
    """ImageExtractor should respect configuration"""
    from core.image_embedding import ImageExtractor, ExtractionConfig

    if not sample_with_images_pdf.exists():
        pytest.skip("Sample PDF not found")

    config = ExtractionConfig(
        min_width=100,
        min_height=100,
        skip_duplicates=True
    )

    extractor = ImageExtractor(config)
    images = extractor.extract_from_pdf(str(sample_with_images_pdf))

    assert isinstance(images, list)
    # All extracted images should meet minimum size
    for img in images:
        assert img.width_px >= 100 or img.height_px >= 100


# ============================================================
# TC-PIPE-03: DOCX Image Embedding
# ============================================================

def test_docx_image_embedding(sample_with_images_pdf, temp_output_dir):
    """DocxImageEmbedder should embed images into DOCX"""
    from docx import Document
    from core.image_embedding import ImageExtractor, DocxImageEmbedder

    if not sample_with_images_pdf.exists():
        pytest.skip("Sample PDF not found")

    # Extract images
    extractor = ImageExtractor()
    images = extractor.extract_from_pdf(str(sample_with_images_pdf))

    if not images:
        # Create a test image if none extracted
        from PIL import Image
        import io
        img = Image.new('RGB', (200, 150), color=(100, 150, 200))
        buffer = io.BytesIO()
        img.save(buffer, format='PNG')

        from core.image_embedding import ImageBlock, ImageFormat
        images = [ImageBlock(
            image_data=buffer.getvalue(),
            format=ImageFormat.PNG,
            width_px=200,
            height_px=150
        )]

    # Create DOCX and embed
    doc = Document()
    doc.add_paragraph("Test document with images")

    embedder = DocxImageEmbedder()
    embedder.embed_images(doc, images, with_captions=True)

    # Save and verify
    output_path = temp_output_dir / "test_embedded.docx"
    doc.save(str(output_path))

    assert output_path.exists()
    assert output_path.stat().st_size > 5000  # Should be > 5KB with images


# ============================================================
# TC-PIPE-04: DOCX Image Embedding with Captions
# ============================================================

def test_docx_image_captions(temp_output_dir):
    """DocxImageEmbedder should add numbered captions"""
    from docx import Document
    from core.image_embedding import DocxImageEmbedder, ImageBlock, ImageFormat
    from PIL import Image
    import io

    # Create test images
    images = []
    for i in range(3):
        img = Image.new('RGB', (200, 150), color=(100 + i*50, 100, 200))
        buffer = io.BytesIO()
        img.save(buffer, format='PNG')
        images.append(ImageBlock(
            image_data=buffer.getvalue(),
            format=ImageFormat.PNG,
            width_px=200,
            height_px=150,
            caption=f"Test image {i+1}"
        ))

    doc = Document()
    doc.add_heading("Document with Captioned Images", level=1)

    embedder = DocxImageEmbedder()
    embedder.embed_images(doc, images, with_captions=True)

    output_path = temp_output_dir / "test_captions.docx"
    doc.save(str(output_path))

    # Verify captions exist
    doc_check = Document(str(output_path))
    caption_found = False
    for para in doc_check.paragraphs:
        if "Hình" in para.text:
            caption_found = True
            break

    assert caption_found, "Captions should be present"


# ============================================================
# TC-PIPE-05: PDF Image Embedding
# ============================================================

def test_pdf_image_embedding(sample_with_images_pdf, temp_output_dir):
    """PdfImageEmbedder should create PDF with images"""
    from core.image_embedding import ImageExtractor, create_pdf_with_images

    if not sample_with_images_pdf.exists():
        pytest.skip("Sample PDF not found")

    # Extract images
    extractor = ImageExtractor()
    images = extractor.extract_from_pdf(str(sample_with_images_pdf))

    if not images:
        # Create test image
        from PIL import Image
        import io
        img = Image.new('RGB', (200, 150), color=(100, 150, 200))
        buffer = io.BytesIO()
        img.save(buffer, format='PNG')

        from core.image_embedding import ImageBlock, ImageFormat
        images = [ImageBlock(
            image_data=buffer.getvalue(),
            format=ImageFormat.PNG,
            width_px=200,
            height_px=150
        )]

    # Create PDF with images
    output_path = temp_output_dir / "test_embedded.pdf"
    create_pdf_with_images(images, str(output_path), title="Test PDF")

    assert output_path.exists()
    assert output_path.stat().st_size > 5000


# ============================================================
# TC-PIPE-06: Cover Embedding - DOCX
# ============================================================

def test_cover_embedding_docx(sample_cover_png, temp_output_dir):
    """CoverEmbedder should add cover to DOCX"""
    from docx import Document
    from core.image_embedding import CoverEmbedder

    if not sample_cover_png.exists():
        pytest.skip("Sample cover not found")

    doc = Document()

    # Read cover as bytes
    cover_data = sample_cover_png.read_bytes()
    cover_b64 = base64.b64encode(cover_data).decode()

    embedder = CoverEmbedder()
    embedder.add_cover_to_docx(doc, cover_b64)

    doc.add_paragraph("Content after cover")

    output_path = temp_output_dir / "test_cover.docx"
    doc.save(str(output_path))

    assert output_path.exists()
    assert output_path.stat().st_size > 10000  # Should include image


# ============================================================
# TC-PIPE-07: Cover from Data URI
# ============================================================

def test_cover_from_data_uri(sample_cover_png, temp_output_dir):
    """CoverEmbedder should handle data URI input"""
    from docx import Document
    from core.image_embedding import CoverEmbedder

    if not sample_cover_png.exists():
        pytest.skip("Sample cover not found")

    # Convert to data URI
    cover_data = sample_cover_png.read_bytes()
    cover_b64 = base64.b64encode(cover_data).decode()
    data_uri = f"data:image/png;base64,{cover_b64}"

    doc = Document()

    embedder = CoverEmbedder()
    embedder.add_cover_to_docx(doc, data_uri)

    output_path = temp_output_dir / "test_cover_uri.docx"
    doc.save(str(output_path))

    assert output_path.exists()


# ============================================================
# TC-PIPE-08: FrontMatterGenerator with Cover
# ============================================================

def test_front_matter_with_cover(sample_cover_png, temp_output_dir):
    """FrontMatterGenerator should add cover before title page"""
    from docx import Document
    from core.export.docx_styles import StyleManager
    from core.export.docx_front_matter import FrontMatterGenerator

    if not sample_cover_png.exists():
        pytest.skip("Sample cover not found")

    doc = Document()
    style_manager = StyleManager(theme_name='professional')

    cover_b64 = base64.b64encode(sample_cover_png.read_bytes()).decode()

    fm_gen = FrontMatterGenerator(doc, style_manager)

    # Add cover
    result = fm_gen.generate_cover_page(cover_b64)
    assert result is True, "Cover page should be added successfully"

    # Add title page
    fm_gen.generate_title_page({
        'title': 'Test Document',
        'author': 'Test Author'
    })

    # Add content
    doc.add_heading('Chapter 1', level=1)
    doc.add_paragraph('Test content')

    output_path = temp_output_dir / "test_front_matter.docx"
    doc.save(str(output_path))

    assert output_path.exists()
    assert output_path.stat().st_size > 10000


# ============================================================
# TC-PIPE-09: Image Block Serialization
# ============================================================

def test_image_block_serialization():
    """ImageBlock should serialize/deserialize correctly"""
    from core.image_embedding import ImageBlock, ImageFormat
    from PIL import Image
    import io

    # Create test image
    img = Image.new('RGB', (100, 100), color=(255, 0, 0))
    buffer = io.BytesIO()
    img.save(buffer, format='PNG')

    block = ImageBlock(
        image_data=buffer.getvalue(),
        format=ImageFormat.PNG,
        width_px=100,
        height_px=100,
        caption="Test image"
    )

    # Test properties
    assert block.size_bytes > 0
    assert block.aspect_ratio == 1.0
    assert block.mime_type == "image/png"
    assert block.file_extension == "png"

    # Test serialization
    data = block.to_dict(include_data=True)
    assert "image_data_base64" in data
    assert data["format"] == "png"
    assert data["width_px"] == 100

    # Test base64 encoding
    b64 = block.to_base64()
    assert len(b64) > 0

    # Test data URI
    uri = block.to_data_uri()
    assert uri.startswith("data:image/png;base64,")


# ============================================================
# TC-PIPE-10: Extraction Config Defaults
# ============================================================

def test_extraction_config_defaults():
    """ExtractionConfig should have sensible defaults"""
    from core.image_embedding import ExtractionConfig, ImageFormat

    config = ExtractionConfig()

    assert config.min_width == 50
    assert config.min_height == 50
    assert config.output_format == ImageFormat.PNG
    assert config.skip_duplicates is True


# ============================================================
# TC-PIPE-11: Pipeline Config Defaults
# ============================================================

def test_pipeline_config_defaults():
    """PipelineConfig should have sensible defaults"""
    from core.image_embedding import PipelineConfig

    config = PipelineConfig()

    assert config.min_image_size == 50
    assert config.max_width_ratio == 0.8
    assert config.with_captions is True
