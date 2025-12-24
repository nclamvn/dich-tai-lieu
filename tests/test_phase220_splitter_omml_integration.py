"""
Phase 2.2.0 - LaTeX Splitter + OMML Integration Tests

Tests the end-to-end integration of the LaTeX equation splitter
with OMML rendering in the Academic DOCX Builder.

Test Coverage:
1. Splitter activation (only when latex_source exists but not latex_equation_primary)
2. Confident splitting → successful OMML rendering
3. Non-confident splitting → fallback to Phase 2.1.2/2.1.0 path
4. Compound LaTeX blocks → clean equation extraction
5. Phase 2.1.2 enrichment still works (no regression)
6. Edge cases (empty segments, conversion failures, etc.)
7. Various equation types (display math, environments, etc.)
8. Fallback chain integrity (2.2.0 → 2.1.2 → 2.1.0 → text)
"""

import pytest
import tempfile
from pathlib import Path
from docx import Document
from unittest.mock import patch, MagicMock

from core.export.docx_academic_builder import (
    build_academic_docx,
    AcademicLayoutConfig,
)
from core.structure.semantic_model import DocNode, DocNodeType
from core.rendering import omml_converter
from core.latex.eq_splitter import SplitEquationResult


# ============================================================================
# TEST 1: Splitter Activation Conditions
# ============================================================================

def test_splitter_activates_when_latex_source_without_primary():
    """Test that Phase 2.2.0 splitter is called when latex_source exists but latex_equation_primary doesn't."""
    if not omml_converter.is_pandoc_available():
        pytest.skip("Pandoc not available")

    config = AcademicLayoutConfig(equation_rendering_mode="omml")

    # Create node with latex_source but NO latex_equation_primary
    # This is the "Track A - LaTeX source available but not yet enriched" scenario
    compound_latex = r"$$ \sup_{n,d \in \mathbb{N}} \left\| \sum_{j=1}^n f(jd) \right\|_H < \infty $$"

    node = DocNode(
        node_type=DocNodeType.EQUATION_BLOCK,
        text="[PDF extracted text - may be garbled]",
        metadata={
            'latex_source': compound_latex,  # Phase 2.1.0 extracted this
            # NO 'latex_equation_primary' - Phase 2.1.2 hasn't run or failed
        }
    )

    nodes = [node]

    with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as f:
        output_path = f.name

    try:
        # Mock the splitter to verify it was called
        with patch('core.latex.eq_splitter.split_latex_equations') as mock_splitter:
            # Configure mock to return confident result
            mock_splitter.return_value = SplitEquationResult(
                original=compound_latex,
                text_segments=[],
                equation_segments=[r"\sup_{n,d \in \mathbb{N}} \left\| \sum_{j=1}^n f(jd) \right\|_H < \infty"],
                is_confident=True,
                reason=None
            )

            build_academic_docx(nodes, output_path, config)

            # Verify splitter was called with latex_source
            mock_splitter.assert_called_once_with(compound_latex)

    finally:
        Path(output_path).unlink(missing_ok=True)


def test_splitter_not_activated_when_primary_exists():
    """Test that Phase 2.2.0 splitter is NOT called when latex_equation_primary exists."""
    if not omml_converter.is_pandoc_available():
        pytest.skip("Pandoc not available")

    config = AcademicLayoutConfig(equation_rendering_mode="omml")

    # Node with BOTH latex_source and latex_equation_primary
    # This means Phase 2.1.2 enrichment already ran successfully
    node = DocNode(
        node_type=DocNodeType.EQUATION_BLOCK,
        text="[PDF text]",
        metadata={
            'latex_source': r"compound block with text $x^2$ and stuff",
            'latex_equation_primary': r"x^2",  # Phase 2.1.2 already enriched
        }
    )

    nodes = [node]

    with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as f:
        output_path = f.name

    try:
        with patch('core.latex.eq_splitter.split_latex_equations') as mock_splitter:
            build_academic_docx(nodes, output_path, config)

            # Splitter should NOT have been called
            mock_splitter.assert_not_called()

    finally:
        Path(output_path).unlink(missing_ok=True)


# ============================================================================
# TEST 2: Confident Splitting → OMML Success
# ============================================================================

def test_confident_splitting_produces_omml():
    """Test that confident splitter results lead to successful OMML rendering."""
    if not omml_converter.is_pandoc_available():
        pytest.skip("Pandoc not available")

    config = AcademicLayoutConfig(equation_rendering_mode="omml")

    # Compound LaTeX block with display math (should split confidently)
    compound_latex = r"$$ E = mc^2 $$"

    node = DocNode(
        node_type=DocNodeType.EQUATION_BLOCK,
        text="[PDF garbled text]",
        metadata={'latex_source': compound_latex}
    )

    nodes = [node]

    with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as f:
        output_path = f.name

    try:
        build_academic_docx(nodes, output_path, config)

        # Verify DOCX was created
        assert Path(output_path).exists()

        # Open and check for OMML elements
        doc = Document(output_path)
        para_elem = doc.paragraphs[0]._element
        omml_ns = '{http://schemas.openxmlformats.org/officeDocument/2006/math}'
        omath_elems = para_elem.findall(f'.//{omml_ns}oMath')

        # Should have OMML rendering (Phase 2.2.0 split successfully)
        assert len(omath_elems) > 0, "Expected OMML rendering from confident split"

    finally:
        Path(output_path).unlink(missing_ok=True)


def test_environment_block_splitting_produces_omml():
    """Test that environment blocks are split confidently and render as OMML."""
    if not omml_converter.is_pandoc_available():
        pytest.skip("Pandoc not available")

    config = AcademicLayoutConfig(equation_rendering_mode="omml")

    # LaTeX environment block (should split confidently)
    env_latex = r"\begin{equation} \int_0^\infty e^{-x^2} dx = \frac{\sqrt{\pi}}{2} \end{equation}"

    node = DocNode(
        node_type=DocNodeType.EQUATION_BLOCK,
        text="[PDF text]",
        metadata={'latex_source': env_latex}
    )

    nodes = [node]

    with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as f:
        output_path = f.name

    try:
        build_academic_docx(nodes, output_path, config)

        doc = Document(output_path)
        para_elem = doc.paragraphs[0]._element
        omml_ns = '{http://schemas.openxmlformats.org/officeDocument/2006/math}'
        omath_elems = para_elem.findall(f'.//{omml_ns}oMath')

        # Should have OMML from environment block
        assert len(omath_elems) > 0

    finally:
        Path(output_path).unlink(missing_ok=True)


# ============================================================================
# TEST 3: Non-Confident Splitting → Fallback
# ============================================================================

def test_non_confident_splitting_falls_back():
    """Test that non-confident splitter results trigger fallback to regular path."""
    if not omml_converter.is_pandoc_available():
        pytest.skip("Pandoc not available")

    config = AcademicLayoutConfig(equation_rendering_mode="omml")

    # Text with inline math (should NOT be confident)
    mixed_content = r"Given a function $f: \mathbb{N} \to H$ in space $H$."

    node = DocNode(
        node_type=DocNodeType.EQUATION_BLOCK,
        text=mixed_content,
        metadata={'latex_source': mixed_content}
    )

    nodes = [node]

    with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as f:
        output_path = f.name

    try:
        build_academic_docx(nodes, output_path, config)

        # Verify DOCX created (should not crash)
        assert Path(output_path).exists()

        # Document should exist and contain fallback content
        # The exact fallback behavior depends on whether latex_source can be converted directly
        # Main point: should not crash due to non-confident split
        doc = Document(output_path)
        assert len(doc.paragraphs) > 0

    finally:
        Path(output_path).unlink(missing_ok=True)


# ============================================================================
# TEST 4: Compound LaTeX Block Extraction
# ============================================================================

def test_compound_latex_block_extraction():
    """Test that compound LaTeX blocks are extracted and rendered correctly."""
    if not omml_converter.is_pandoc_available():
        pytest.skip("Pandoc not available")

    config = AcademicLayoutConfig(equation_rendering_mode="omml")

    # Real-world compound block from arXiv (with delimiters and whitespace)
    compound_latex = r"$$   \sum_{j=1}^n f(jd) \leq C \cdot n   $$"

    node = DocNode(
        node_type=DocNodeType.EQUATION_BLOCK,
        text="garbled PDF text here",
        metadata={'latex_source': compound_latex}
    )

    nodes = [node]

    with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as f:
        output_path = f.name

    try:
        build_academic_docx(nodes, output_path, config)

        doc = Document(output_path)
        para_elem = doc.paragraphs[0]._element
        omml_ns = '{http://schemas.openxmlformats.org/officeDocument/2006/math}'
        omath_elems = para_elem.findall(f'.//{omml_ns}oMath')

        # Should successfully extract and render OMML
        assert len(omath_elems) > 0, "Compound block should be split and rendered"

    finally:
        Path(output_path).unlink(missing_ok=True)


# ============================================================================
# TEST 5: Phase 2.1.2 Enrichment Still Works (No Regression)
# ============================================================================

def test_phase212_enrichment_not_broken():
    """Test that Phase 2.1.2 enriched equations still work (no regression)."""
    if not omml_converter.is_pandoc_available():
        pytest.skip("Pandoc not available")

    config = AcademicLayoutConfig(equation_rendering_mode="omml")

    # Node with Phase 2.1.2 enrichment (latex_equation_primary exists)
    node = DocNode(
        node_type=DocNodeType.EQUATION_BLOCK,
        text="garbled",
        metadata={
            'latex_source': r"Some compound block $x^2 + y^2 = z^2$ with text",
            'latex_equation_primary': r"x^2 + y^2 = z^2",  # Phase 2.1.2 extracted clean equation
        }
    )

    nodes = [node]

    with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as f:
        output_path = f.name

    try:
        build_academic_docx(nodes, output_path, config)

        doc = Document(output_path)
        para_elem = doc.paragraphs[0]._element
        omml_ns = '{http://schemas.openxmlformats.org/officeDocument/2006/math}'
        omath_elems = para_elem.findall(f'.//{omml_ns}oMath')

        # Should render OMML using Phase 2.1.2's clean equation
        assert len(omath_elems) > 0, "Phase 2.1.2 enriched equations should still render"

    finally:
        Path(output_path).unlink(missing_ok=True)


# ============================================================================
# TEST 6: Edge Cases
# ============================================================================

def test_splitter_empty_result_falls_back():
    """Test that empty splitter results fall back gracefully."""
    config = AcademicLayoutConfig(equation_rendering_mode="omml")

    # Mock splitter to return confident but empty segments
    with patch('core.latex.eq_splitter.split_latex_equations') as mock_splitter:
        mock_splitter.return_value = SplitEquationResult(
            original="something",
            text_segments=[],
            equation_segments=[],  # Empty!
            is_confident=True,
            reason=None
        )

        # Phase 2.1.1: Also mock inject_omml_as_display to force fallback
        with patch('core.rendering.omml_converter.inject_omml_as_display', return_value=False):
            node = DocNode(
                node_type=DocNodeType.EQUATION_BLOCK,
                text="fallback text",
                metadata={'latex_source': "something"}
            )

            nodes = [node]

            with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as f:
                output_path = f.name

            try:
                # Should not crash
                build_academic_docx(nodes, output_path, config)

                # Should have fallback text
                doc = Document(output_path)
                para_text = doc.paragraphs[0].text
                assert "fallback text" in para_text

            finally:
                Path(output_path).unlink(missing_ok=True)


def test_splitter_omml_conversion_failure_falls_back():
    """Test that splitter success + OMML failure still falls back to regular path."""
    if not omml_converter.is_pandoc_available():
        pytest.skip("Pandoc not available")

    config = AcademicLayoutConfig(equation_rendering_mode="omml")

    # Create a valid split but mock OMML conversion to fail
    node = DocNode(
        node_type=DocNodeType.EQUATION_BLOCK,
        text="$x^2$",
        metadata={'latex_source': r"$$ x^2 $$"}
    )

    nodes = [node]

    with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as f:
        output_path = f.name

    try:
        # Mock latex_to_omml to return None (conversion failed)
        with patch('core.rendering.omml_converter.latex_to_omml', return_value=None):
            build_academic_docx(nodes, output_path, config)

            # Should not crash, should have fallback text
            doc = Document(output_path)
            assert len(doc.paragraphs) > 0

            # Check that it fell back to regular path (should try latex_source)
            # The exact text depends on fallback chain

    finally:
        Path(output_path).unlink(missing_ok=True)


def test_splitter_exception_handling():
    """Test that exceptions in splitter don't crash document generation."""
    config = AcademicLayoutConfig(equation_rendering_mode="omml")

    with patch('core.latex.eq_splitter.split_latex_equations', side_effect=RuntimeError("Splitter error")):
        node = DocNode(
            node_type=DocNodeType.EQUATION_BLOCK,
            text="fallback text",
            metadata={'latex_source': "something"}
        )

        nodes = [node]

        with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as f:
            output_path = f.name

        try:
            # Should not crash
            build_academic_docx(nodes, output_path, config)

            # Should still create document
            assert Path(output_path).exists()
            doc = Document(output_path)
            assert len(doc.paragraphs) > 0

        finally:
            Path(output_path).unlink(missing_ok=True)


# ============================================================================
# TEST 7: Fallback Chain Integrity
# ============================================================================

def test_fallback_chain_phase220_to_212_to_210():
    """Test the complete fallback chain: 2.2.0 → 2.1.2 → 2.1.0 → text."""
    if not omml_converter.is_pandoc_available():
        pytest.skip("Pandoc not available")

    config = AcademicLayoutConfig(equation_rendering_mode="omml")

    # Test Case 1: Phase 2.2.0 path (latex_source only)
    node1 = DocNode(
        node_type=DocNodeType.EQUATION_BLOCK,
        text="text fallback",
        metadata={'latex_source': r"$$ x = 1 $$"}
    )

    # Test Case 2: Phase 2.1.2 path (latex_equation_primary exists)
    node2 = DocNode(
        node_type=DocNodeType.EQUATION_BLOCK,
        text="text fallback",
        metadata={
            'latex_source': r"compound $x = 2$ block",
            'latex_equation_primary': r"x = 2"
        }
    )

    # Test Case 3: Phase 2.1.0 path (only latex_source, splitter fails)
    node3 = DocNode(
        node_type=DocNodeType.EQUATION_BLOCK,
        text="text fallback",
        metadata={'latex_source': r"x = 3"}
    )

    # Test Case 4: Text fallback only (no metadata)
    node4 = DocNode(
        node_type=DocNodeType.EQUATION_BLOCK,
        text="x = 4"
    )

    nodes = [node1, node2, node3, node4]

    with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as f:
        output_path = f.name

    try:
        # Should handle all cases without crashing
        build_academic_docx(nodes, output_path, config)

        doc = Document(output_path)
        assert len(doc.paragraphs) >= 4

        # All equations should have been processed somehow
        # (either OMML or fallback text)

    finally:
        Path(output_path).unlink(missing_ok=True)


# ============================================================================
# TEST 8: Multiple Equations with Mixed Metadata
# ============================================================================

def test_multiple_equations_mixed_metadata():
    """Test document with multiple equations having different metadata states."""
    if not omml_converter.is_pandoc_available():
        pytest.skip("Pandoc not available")

    config = AcademicLayoutConfig(equation_rendering_mode="omml")

    nodes = [
        # Equation 1: Phase 2.2.0 path (splitter candidate)
        DocNode(
            node_type=DocNodeType.EQUATION_BLOCK,
            text="eq1",
            metadata={'latex_source': r"$$ a = b $$"}
        ),
        # Equation 2: Phase 2.1.2 path (already enriched)
        DocNode(
            node_type=DocNodeType.EQUATION_BLOCK,
            text="eq2",
            metadata={
                'latex_source': r"compound",
                'latex_equation_primary': r"c = d"
            }
        ),
        # Equation 3: No metadata (text only)
        DocNode(
            node_type=DocNodeType.EQUATION_BLOCK,
            text="e = f"
        )
    ]

    with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as f:
        output_path = f.name

    try:
        build_academic_docx(nodes, output_path, config)

        # All equations should be processed
        doc = Document(output_path)
        assert len(doc.paragraphs) >= 3

    finally:
        Path(output_path).unlink(missing_ok=True)


# ============================================================================
# TEST 9: Latex Text Mode Not Affected by Splitter
# ============================================================================

def test_latex_text_mode_not_affected():
    """Test that latex_text rendering mode doesn't trigger splitter."""
    config = AcademicLayoutConfig(equation_rendering_mode="latex_text")

    node = DocNode(
        node_type=DocNodeType.EQUATION_BLOCK,
        text="x^2",
        metadata={'latex_source': r"$$ x^2 $$"}
    )

    nodes = [node]

    with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as f:
        output_path = f.name

    try:
        with patch('core.latex.eq_splitter.split_latex_equations') as mock_splitter:
            build_academic_docx(nodes, output_path, config)

            # Splitter should NOT be called in latex_text mode
            mock_splitter.assert_not_called()

            # Should have text rendering
            doc = Document(output_path)
            para_text = doc.paragraphs[0].text
            assert "x^2" in para_text

    finally:
        Path(output_path).unlink(missing_ok=True)


# ============================================================================
# TEST 10: Real-World arXiv Pattern
# ============================================================================

def test_real_world_arxiv_pattern():
    """Test with real-world arXiv LaTeX pattern."""
    if not omml_converter.is_pandoc_available():
        pytest.skip("Pandoc not available")

    config = AcademicLayoutConfig(equation_rendering_mode="omml")

    # Real pattern from arXiv papers
    arxiv_latex = r"""
    \begin{equation}
    \sup_{n,d \in \mathbb{N}} \left\| \sum_{j=1}^n f(jd) \right\|_H < \infty
    \end{equation}
    """

    node = DocNode(
        node_type=DocNodeType.EQUATION_BLOCK,
        text="[OCR garbled equation]",
        metadata={'latex_source': arxiv_latex.strip()}
    )

    nodes = [node]

    with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as f:
        output_path = f.name

    try:
        build_academic_docx(nodes, output_path, config)

        doc = Document(output_path)
        para_elem = doc.paragraphs[0]._element
        omml_ns = '{http://schemas.openxmlformats.org/officeDocument/2006/math}'
        omath_elems = para_elem.findall(f'.//{omml_ns}oMath')

        # Should render real arXiv equation as OMML
        # (environment block should be detected and split confidently)
        assert len(omath_elems) > 0, "Real arXiv equation should render as OMML"

    finally:
        Path(output_path).unlink(missing_ok=True)


# ============================================================================
# Pytest Configuration
# ============================================================================

if __name__ == "__main__":
    pytest.main([__file__, "-v"])
