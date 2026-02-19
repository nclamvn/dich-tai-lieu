"""
APS V2 API Router

FastAPI endpoints for Claude-Native Universal Publishing.
"""

import logging
from pathlib import Path
from typing import List

from fastapi import APIRouter, UploadFile, File, Form, HTTPException, Body, Request
from fastapi.responses import FileResponse

from .aps_v2_models import (
    PublishTextRequest,
    JobResponseV2,
    ProfileListResponse,
    HealthResponseV2,
    ErrorResponseV2,
    JobStatusV2,
)
from .aps_v2_service import get_v2_service
from .rate_limiter import limiter, rate_limit_config

# Import unified client for provider validation
from ai_providers.unified_client import (
    validate_providers_before_job,
    get_unified_client,
    AllProvidersUnavailableError,
)

logger = logging.getLogger(__name__)

router = APIRouter(prefix="/api/v2", tags=["APS V2 - Universal Publishing"])


# ==================== PUBLISH ENDPOINTS ====================

@router.post(
    "/publish",
    response_model=JobResponseV2,
    responses={400: {"model": ErrorResponseV2}},
    summary="Start publishing job with file upload",
)
@limiter.limit(rate_limit_config.get_limit("vision"))
async def publish_file(
    request: Request,
    file: UploadFile = File(..., description="Document to publish"),
    source_language: str = Form(default="en"),
    target_language: str = Form(default="vi"),
    profile_id: str = Form(default="novel"),
    output_formats: str = Form(default="docx"),  # Comma-separated
    use_vision: bool = Form(default=True, description="Use Claude Vision for PDF reading (recommended)"),
    api_key: str = Form(default="", description="User API key (optional, overrides server config)"),
    docx_template: str = Form(default="auto", description="DOCX template: 'ebook', 'academic', 'business', or 'auto'"),
    pdf_template: str = Form(default="auto", description="PDF template: 'ebook', 'academic', 'business', or 'auto'"),
    provider: str = Form(default="auto", description="AI provider: 'auto', 'openai', 'anthropic'"),
    model: str = Form(default="", description="Model name (e.g., 'gpt-4o', 'claude-sonnet-4-20250514')"),
):
    """
    Upload a document and start the publishing pipeline.

    **Vision Mode (default=True):**
    For PDFs, Claude Vision reads the document like a human would,
    perfectly preserving mathematical formulas and layout.

    Supported input formats: .txt, .md, .docx, .pdf

    **Publishing Profiles:**
    - `novel` - Fiction/literary works
    - `poetry` - Poetry collections
    - `essay` - Essays and articles
    - `business_report` - Business reports
    - `white_paper` - White papers
    - `academic_paper` - Academic papers
    - `arxiv_paper` - arXiv preprints (LaTeX)
    - `thesis` - Thesis/dissertation
    - `textbook` - Educational textbooks
    - `technical_doc` - Technical documentation
    - `api_doc` - API documentation
    - `user_manual` - User manuals

    **DOCX Templates (docx_template):**
    - `auto` - Auto-select based on document type (default)
    - `ebook` - Professional book layout (Georgia font, trade paperback)
    - `academic` - Academic paper format (Times New Roman, A4)
    - `business` - Corporate report style (Calibri, clean layout)

    **PDF Templates (pdf_template):**
    - `auto` - Auto-select based on document type (default)
    - `ebook` - Trade paperback size, serif fonts, elegant chapters
    - `academic` - A4, formal styling, 1.5 line spacing, running headers
    - `business` - A4 narrow margins, sans-serif, blue accents
    """
    service = get_v2_service()

    try:
        # ===== PRE-VALIDATION: Check AI providers before starting =====
        # If user provided api_key, skip validation (trust user key)
        if not api_key:
            success, message, status = await validate_providers_before_job()
            if not success:
                logger.error(f"Provider validation failed: {message}")
                raise HTTPException(
                    status_code=503,
                    detail={
                        "error": "No AI providers available",
                        "message": message,
                        "providers": status.get("statuses", {}),
                        "action": "Please check your API keys and billing status"
                    }
                )
            logger.info(f"Provider validated: {message}")
        else:
            logger.info("Using user-provided API key, skipping server validation")

        # Save uploaded file
        content_bytes = await file.read()
        file_path = await service.save_upload(file.filename, content_bytes)

        # Read content (Vision mode for PDFs, with source language for OCR routing)
        content = await service.read_upload(
            file_path,
            use_vision=use_vision,
            source_lang=source_language
        )

        if not content.strip():
            raise HTTPException(status_code=400, detail="Empty document")

        # Parse output formats
        formats = [f.strip() for f in output_formats.split(",") if f.strip()]
        if not formats:
            formats = ["docx"]

        # Create job with Vision mode and provider selection
        job = await service.create_job(
            source_file=file.filename,
            content=content,
            source_language=source_language,
            target_language=target_language,
            profile_id=profile_id,
            output_formats=formats,
            use_vision=use_vision,
            api_key=api_key if api_key else None,
            docx_template=docx_template,
            pdf_template=pdf_template,
            provider=provider if provider != "auto" else None,
            model=model if model else None,
        )

        return service.get_job_response(job)

    except HTTPException:
        raise
    except AllProvidersUnavailableError as e:
        logger.error(f"All providers unavailable: {e}")
        raise HTTPException(
            status_code=503,
            detail={
                "error": "All AI providers unavailable",
                "message": str(e),
                "action": "Please check your API keys and billing status"
            }
        )
    except Exception as e:
        logger.error(f"Publish failed: {e}")
        raise HTTPException(status_code=500, detail=str(e))


@router.post(
    "/publish/text",
    response_model=JobResponseV2,
    responses={400: {"model": ErrorResponseV2}},
    summary="Start publishing job with text content",
)
async def publish_text(request: PublishTextRequest):
    """
    Submit text content directly (no file upload).

    Useful for:
    - API integrations
    - Text from other sources
    - Testing
    """
    service = get_v2_service()

    try:
        if not request.content.strip():
            raise HTTPException(status_code=400, detail="Empty content")

        # Create job
        job = await service.create_job(
            source_file=request.filename,
            content=request.content,
            source_language=request.source_language,
            target_language=request.target_language,
            profile_id=request.profile_id,
            output_formats=request.output_formats,
            docx_template=request.docx_template,
            pdf_template=request.pdf_template,
        )

        return service.get_job_response(job)

    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Publish text failed: {e}")
        raise HTTPException(status_code=500, detail=str(e))


# ==================== JOB ENDPOINTS ====================

@router.get(
    "/jobs",
    response_model=List[JobResponseV2],
    summary="List all V2 jobs",
)
async def list_jobs(limit: int = 50):
    """List all V2 publishing jobs, most recent first."""
    service = get_v2_service()
    all_jobs = service._repo.get_all_jobs(limit=limit)
    results = []
    for job_data in all_jobs:
        # Ensure job is in memory cache for get_job_response
        job_id = job_data.get("job_id")
        if job_id not in service._jobs:
            service._jobs[job_id] = job_data
        results.append(service.get_job_response(job_data))
    return results


@router.delete(
    "/jobs",
    summary="Bulk delete V2 jobs",
)
async def bulk_delete_jobs(job_ids: List[str] = Body(...)):
    """Delete multiple V2 jobs at once."""
    service = get_v2_service()
    deleted = 0
    for job_id in job_ids:
        try:
            service._repo.delete(job_id)
            service._jobs.pop(job_id, None)
            deleted += 1
        except Exception:
            pass
    return {"deleted": deleted}


@router.get(
    "/jobs/current",
    response_model=JobResponseV2,
    responses={404: {"model": ErrorResponseV2}},
    summary="Get current running job",
)
async def get_current_job():
    """
    Get the most recent running/pending job.

    Used by frontend to reconnect after page refresh.
    Returns 404 if no running job exists.
    """
    service = get_v2_service()

    # Find running/pending job
    for job in service._jobs.values():
        status = job.get("status")
        if status:
            status_val = status.value if hasattr(status, 'value') else str(status)
            if status_val in ["pending", "running", "vision_reading", "translating"]:
                return service.get_job_response(job)

    raise HTTPException(status_code=404, detail="No running job")


@router.get(
    "/jobs/{job_id}",
    response_model=JobResponseV2,
    responses={404: {"model": ErrorResponseV2}},
    summary="Get job status",
)
async def get_job(job_id: str):
    """
    Get the current status of a publishing job.

    Poll this endpoint to track progress.
    """
    service = get_v2_service()

    job = service.get_job(job_id)
    if not job:
        raise HTTPException(status_code=404, detail=f"Job not found: {job_id}")

    return service.get_job_response(job)


@router.get(
    "/jobs/{job_id}/reader-content",
    summary="Get reader-friendly content for in-app ebook reader",
)
async def get_reader_content(job_id: str):
    """
    Get LayoutDNA-structured content for the in-app ebook reader.

    Returns chapters split at H1 headings with typed regions.
    Falls back to translated chunks or plain text if LayoutDNA unavailable.
    """
    service = get_v2_service()

    job = service.get_job(job_id)
    if not job:
        raise HTTPException(status_code=404, detail=f"Job not found: {job_id}")

    status = job.get("status")
    status_val = status.value if hasattr(status, "value") else str(status)
    if status_val not in ("complete", "completed"):
        raise HTTPException(status_code=400, detail=f"Job not complete. Status: {status_val}")

    chapters = []
    has_layout_dna = False
    content_source = "none"

    # Check if LayoutDNA exists (for metadata only — it contains SOURCE text, not translated)
    layout_dna_data = job.get("layout_dna")
    if layout_dna_data and isinstance(layout_dna_data, dict) and layout_dna_data.get("regions"):
        has_layout_dna = True

    # Strategy 1: translated_chunks (in-memory translated content)
    translated_chunks = job.get("translated_chunks", [])
    if translated_chunks:
        chapters = _chunks_to_chapters(translated_chunks)
        content_source = "translated_chunks"
        logger.info(f"[{job_id}] Reader: using translated_chunks ({len(translated_chunks)} chunks)")

    # Strategy 2: read output files (these contain TRANSLATED text)
    if not chapters:
        for fmt in ("md", "txt", "docx"):
            path = service.get_output_path(job_id, fmt)
            if path and path.exists():
                try:
                    if fmt == "docx":
                        from docx import Document as DocxDocument
                        doc = DocxDocument(str(path))
                        text = "\n\n".join(p.text for p in doc.paragraphs if p.text.strip())
                    else:
                        text = path.read_text(encoding="utf-8", errors="ignore")
                    if text.strip():
                        chapters = _text_to_chapters(text, job.get("source_file", "Document"))
                        content_source = f"output_file:{fmt}"
                        logger.info(f"[{job_id}] Reader: using output file ({fmt})")
                        break
                except Exception as e:
                    logger.warning(f"[{job_id}] Failed to read {fmt} output: {e}")

    # Strategy 3 (LAST RESORT): LayoutDNA regions — WARNING: this is SOURCE text
    if not chapters and has_layout_dna:
        try:
            from api.services.layout_dna import LayoutDNA
            dna = LayoutDNA.from_dict(layout_dna_data)
            if dna.regions:
                chapters = _split_dna_to_chapters(dna)
                content_source = "source_layout_dna"
                logger.warning(f"[{job_id}] Reader: falling back to LayoutDNA SOURCE text — translated content not available")
        except Exception as e:
            logger.warning(f"[{job_id}] LayoutDNA parse failed: {e}")

    # Metadata
    total_words = sum(
        len(r.get("content", "").split())
        for ch in chapters
        for r in ch.get("regions", [])
        if r.get("type") in ("text", "heading")
    )
    metadata = {
        "total_chapters": len(chapters),
        "total_words": total_words,
        "total_regions": sum(len(ch.get("regions", [])) for ch in chapters),
        "tables": sum(1 for ch in chapters for r in ch.get("regions", []) if r.get("type") == "table"),
        "formulas": sum(1 for ch in chapters for r in ch.get("regions", []) if r.get("type") == "formula"),
        "has_layout_dna": has_layout_dna,
        "content_source": content_source,
    }

    # Quality summary
    quality = {}
    eqs = job.get("eqs")
    if eqs:
        quality["eqs_grade"] = eqs.get("eqs_grade", eqs.get("grade", ""))
        quality["eqs_score"] = eqs.get("eqs_score", eqs.get("score", 0))
    consistency = job.get("consistency")
    if consistency:
        quality["consistency_score"] = consistency.get("score", 0)
        quality["consistency_passed"] = consistency.get("passed", False)
    qapr = job.get("qapr")
    if qapr:
        quality["provider"] = qapr.get("provider", qapr.get("selected_provider", ""))
        quality["routing_mode"] = qapr.get("mode", "")

    return {
        "job_id": job_id,
        "title": job.get("source_file", "Document"),
        "source_language": job.get("source_language", ""),
        "target_language": job.get("target_language", ""),
        "chapters": chapters,
        "metadata": metadata,
        "quality": quality,
    }


def _split_dna_to_chapters(dna) -> list:
    """Split LayoutDNA regions into chapters at H1 boundaries."""
    from api.services.layout_dna import RegionType

    chapters = []
    current_regions = []
    current_title = "Document"
    chapter_index = 0

    for region in dna.regions:
        rd = _region_to_dict(region)

        if (region.type == RegionType.HEADING
                and region.level == 1
                and current_regions):
            chapter_index += 1
            chapters.append({
                "id": f"ch-{chapter_index}",
                "title": current_title,
                "regions": current_regions,
            })
            current_regions = []
            current_title = region.content.strip() or f"Chapter {chapter_index + 1}"

        current_regions.append(rd)

        if region.type == RegionType.HEADING and region.level == 1:
            current_title = region.content.strip() or current_title

    if current_regions:
        chapter_index += 1
        chapters.append({
            "id": f"ch-{chapter_index}",
            "title": current_title,
            "regions": current_regions,
        })

    return chapters


def _region_to_dict(region) -> dict:
    """Convert a LayoutDNA Region to reader-friendly dict."""
    from api.services.layout_dna import RegionType

    base = {"type": region.type.value, "content": region.content}

    if region.type == RegionType.HEADING:
        base["level"] = region.level or 2
    elif region.type == RegionType.TABLE:
        base["html"] = _markdown_table_to_html(region.content)
        base["caption"] = region.metadata.get("caption", "")
    elif region.type == RegionType.FORMULA:
        base["latex"] = region.content
        base["display_text"] = region.content
        base["inline"] = region.metadata.get("inline", False)
    elif region.type == RegionType.LIST:
        items = [line.lstrip("•-*0123456789.) ").strip()
                 for line in region.content.split("\n") if line.strip()]
        base["items"] = items
        ordered = any(region.content.lstrip().startswith(f"{i}")
                       for i in "0123456789")
        base["ordered"] = ordered
    elif region.type == RegionType.IMAGE:
        base["alt_text"] = region.metadata.get("alt_text", region.content.strip())
    elif region.type == RegionType.CODE:
        pass  # content is already the code
    else:
        base["style"] = "body"

    return base


def _markdown_table_to_html(md_table: str) -> str:
    """Convert markdown table to simple HTML."""
    lines = [l.strip() for l in md_table.strip().split("\n") if l.strip()]
    if not lines:
        return ""

    html_parts = ["<table>"]
    for i, line in enumerate(lines):
        # Skip separator lines like |---|---|
        if set(line.replace("|", "").replace("-", "").replace(":", "").strip()) == set():
            continue
        cells = [c.strip() for c in line.strip("|").split("|")]
        tag = "th" if i == 0 else "td"
        row = "".join(f"<{tag}>{c}</{tag}>" for c in cells)
        html_parts.append(f"<tr>{row}</tr>")
    html_parts.append("</table>")
    return "\n".join(html_parts)


def _chunks_to_chapters(chunks: list) -> list:
    """Convert translated chunks to chapters."""
    chapters = []
    for i, chunk in enumerate(chunks):
        text = chunk if isinstance(chunk, str) else chunk.get("text", str(chunk))
        chapters.append({
            "id": f"ch-{i + 1}",
            "title": f"Section {i + 1}",
            "regions": [{"type": "text", "content": text, "style": "body"}],
        })
    return chapters


def _text_to_chapters(text: str, title: str = "Document") -> list:
    """Convert plain text to chapters, splitting on heading patterns."""
    import re
    # Try splitting on markdown H1 headers
    parts = re.split(r'\n(?=# [^\n])', text)
    if len(parts) > 1:
        chapters = []
        for i, part in enumerate(parts):
            lines = part.strip().split("\n", 1)
            ch_title = lines[0].lstrip("# ").strip() if lines[0].startswith("#") else f"Section {i + 1}"
            body = lines[1].strip() if len(lines) > 1 else lines[0]
            chapters.append({
                "id": f"ch-{i + 1}",
                "title": ch_title,
                "regions": [{"type": "text", "content": body, "style": "body"}],
            })
        return chapters

    # Single chapter fallback
    return [{
        "id": "ch-1",
        "title": title.rsplit(".", 1)[0] if "." in title else title,
        "regions": [{"type": "text", "content": text, "style": "body"}],
    }]


@router.get(
    "/jobs/{job_id}/download/{format_type}",
    summary="Download output file",
)
async def download_output(job_id: str, format_type: str):
    """
    Download the generated output file.

    **Format types:** docx, pdf, epub, html, md, latex

    If the requested format wasn't generated, will attempt on-demand conversion.
    """
    service = get_v2_service()

    job = service.get_job(job_id)
    if not job:
        raise HTTPException(status_code=404, detail=f"Job not found: {job_id}")

    # Check job is complete
    if job["status"] != JobStatusV2.COMPLETE:
        raise HTTPException(
            status_code=400,
            detail=f"Job not complete. Status: {job['status'].value}"
        )

    # Get output path
    output_path = service.get_output_path(job_id, format_type)

    # If format not found, try on-demand conversion
    if not output_path:
        output_path = await _convert_on_demand(job_id, format_type, service)
        if not output_path:
            raise HTTPException(
                status_code=404,
                detail=f"Output not found for format: {format_type}. Available formats: {list(job.get('output_paths', {}).keys())}"
            )

    # Determine media type
    media_types = {
        "docx": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        "pdf": "application/pdf",
        "epub": "application/epub+zip",
        "html": "text/html",
        "md": "text/markdown",
        "latex": "application/x-latex",
    }

    media_type = media_types.get(format_type, "application/octet-stream")

    return FileResponse(
        path=str(output_path),
        media_type=media_type,
        filename=output_path.name,
    )


async def _convert_on_demand(job_id: str, target_format: str, service) -> Path:
    """
    Convert to requested format on-demand from existing DOCX.

    Returns the output path if successful, None otherwise.
    """
    # Get existing docx path
    docx_path = service.get_output_path(job_id, "docx")
    if not docx_path or not docx_path.exists():
        logger.warning(f"[{job_id}] No DOCX found for on-demand conversion")
        return None

    output_dir = docx_path.parent
    base_name = f"{job_id}_translated"
    target_path = output_dir / f"{base_name}.{target_format}"

    # Check if already exists
    if target_path.exists():
        return target_path

    try:
        if target_format == "md":
            # Convert DOCX to Markdown
            from docx import Document
            doc = Document(str(docx_path))

            md_lines = []
            for para in doc.paragraphs:
                text = para.text.strip()
                if not text:
                    md_lines.append("")
                    continue

                # Detect headings by style
                style_name = para.style.name.lower() if para.style else ""
                if "heading 1" in style_name or "title" in style_name:
                    md_lines.append(f"# {text}")
                elif "heading 2" in style_name:
                    md_lines.append(f"## {text}")
                elif "heading 3" in style_name:
                    md_lines.append(f"### {text}")
                else:
                    md_lines.append(text)

                md_lines.append("")  # Empty line after each paragraph

            target_path.write_text("\n".join(md_lines), encoding="utf-8")
            logger.info(f"[{job_id}] On-demand MD conversion successful: {target_path}")
            return target_path

        elif target_format in ["pdf", "html", "epub"]:
            # Use pandoc for conversion
            import subprocess
            import shutil

            if not shutil.which("pandoc"):
                logger.error("pandoc not available for conversion")
                return None

            # For PDF, try LibreOffice first (better formatting)
            if target_format == "pdf" and shutil.which("soffice"):
                result = subprocess.run(
                    ["soffice", "--headless", "--convert-to", "pdf", "--outdir", str(output_dir), str(docx_path)],
                    capture_output=True,
                    timeout=120
                )
                # LibreOffice outputs with original filename
                lo_output = output_dir / f"{docx_path.stem}.pdf"
                if result.returncode == 0 and lo_output.exists():
                    if lo_output != target_path:
                        lo_output.rename(target_path)
                    logger.info(f"[{job_id}] On-demand PDF conversion (LibreOffice) successful")
                    return target_path

            # Fallback to pandoc
            cmd = ["pandoc", str(docx_path), "-o", str(target_path)]
            if target_format == "pdf":
                # Use xelatex for Unicode support
                if shutil.which("xelatex"):
                    cmd.extend(["--pdf-engine=xelatex"])
                elif shutil.which("pdflatex"):
                    cmd.extend(["--pdf-engine=pdflatex"])
            elif target_format == "html":
                cmd.extend(["--standalone"])

            result = subprocess.run(cmd, capture_output=True, timeout=120)
            if result.returncode == 0 and target_path.exists():
                logger.info(f"[{job_id}] On-demand {target_format.upper()} conversion successful")
                return target_path
            else:
                logger.error(f"pandoc conversion failed: {result.stderr.decode()[:500]}")
                return None

        else:
            logger.warning(f"Unsupported format for on-demand conversion: {target_format}")
            return None

    except Exception as e:
        logger.error(f"[{job_id}] On-demand conversion error: {e}")
        return None


@router.post(
    "/jobs/{job_id}/cancel",
    response_model=JobResponseV2,
    responses={404: {"model": ErrorResponseV2}},
    summary="Cancel a running job",
)
async def cancel_job(job_id: str):
    """
    Cancel a running publishing job.
    """
    service = get_v2_service()

    job = service.get_job(job_id)
    if not job:
        raise HTTPException(status_code=404, detail=f"Job not found: {job_id}")

    success = await service.cancel_job(job_id)

    if not success:
        raise HTTPException(
            status_code=400,
            detail="Job cannot be cancelled (already complete or failed)"
        )

    return service.get_job_response(job)


# ==================== PROFILE ENDPOINTS ====================

@router.get(
    "/profiles",
    response_model=ProfileListResponse,
    summary="List publishing profiles",
)
async def list_publishing_profiles():
    """
    Get all available publishing profiles.

    Each profile defines formatting rules for a specific genre/document type.
    """
    service = get_v2_service()

    profiles = service.get_profiles()

    return ProfileListResponse(
        profiles=profiles,
        total=len(profiles),
    )


@router.get(
    "/profiles/{profile_id}",
    summary="Get profile details",
)
async def get_profile_detail(profile_id: str):
    """
    Get detailed information about a specific publishing profile.
    """
    service = get_v2_service()

    profile = service.get_profile_by_id(profile_id)
    if not profile:
        raise HTTPException(status_code=404, detail=f"Profile not found: {profile_id}")

    return profile


# ==================== CACHE MANAGEMENT ====================

@router.delete(
    "/cache",
    summary="Clear all cache",
)
async def clear_all_cache():
    """
    Clear all cached jobs and output files.

    Use this to force fresh translations.
    """
    service = get_v2_service()
    count = service.clear_cache()

    return {
        "success": True,
        "cleared_jobs": count,
        "message": f"Cleared {count} jobs from cache"
    }


@router.delete(
    "/cache/{job_id}",
    summary="Clear specific job cache",
)
async def clear_job_cache(job_id: str):
    """
    Clear a specific job from cache.
    """
    service = get_v2_service()
    success = service.clear_job(job_id)

    if success:
        return {"success": True, "message": f"Job {job_id} cleared"}
    else:
        raise HTTPException(status_code=404, detail=f"Job {job_id} not found")


@router.get(
    "/cache/stats",
    summary="Get cache statistics",
)
async def get_cache_stats():
    """
    Get current cache statistics.
    """
    service = get_v2_service()
    return service.get_cache_stats()


# ==================== HEALTH ====================

@router.get(
    "/health",
    response_model=HealthResponseV2,
    summary="Health check",
)
async def health_check():
    """
    Check API health and dependencies.
    """
    service = get_v2_service()
    return service.health_check()


@router.get(
    "/providers/status",
    summary="Check LLM provider status",
)
async def provider_status():
    """
    Check the status of all LLM providers.

    Shows which providers are available, which have billing issues,
    and which one is currently active.

    Returns detailed status for each provider:
    - available: Whether the provider is working
    - no_credit: Provider has billing/credit issues
    - invalid_key: API key is invalid
    - not_configured: API key not set
    """
    client = get_unified_client()
    return await client.get_status_summary()
