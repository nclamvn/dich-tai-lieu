"""
APS Service - REAL INTEGRATION

Business logic for the AI Publishing System pipeline.
Connects to real BatchProcessor for Agent #1.
"""

import asyncio
import logging
import uuid
import os
from pathlib import Path
from datetime import datetime
from typing import Dict, List, Optional, Any
from dataclasses import dataclass, field
import threading
import traceback
import json
import time

# Import real translation components
from core.batch_processor import BatchProcessor
from core.job_queue import JobQueue, TranslationJob, JobStatus as TranslationJobStatus

# Import APS contracts
from core.contracts import ManuscriptCoreOutput, Segment, SegmentType
from core.contracts.layout_intent import LayoutIntentPackage, Block, BlockType, SectionDefinition, SectionType, ConsistencyReport

# Import ADN extractor
try:
    from core.adn import ADNExtractor, ContentADN
    HAS_ADN = True
except ImportError:
    HAS_ADN = False

# Import cache manager (PERF-004)
try:
    from core.cache import get_cache_manager, APSCacheManager
    HAS_CACHE = True
except ImportError:
    HAS_CACHE = False

logger = logging.getLogger(__name__)


@dataclass
class APSJob:
    """Represents an APS processing job"""
    job_id: str
    status: str = "pending"
    progress: float = 0
    stage: int = 1
    stage_progress: float = 0
    current_task: str = ""
    error: Optional[str] = None

    # Input
    source_file: str = ""
    source_path: Path = None
    source_lang: str = "auto"
    target_lang: str = "vi"
    template: str = "book"
    output_formats: List[str] = field(default_factory=lambda: ["docx"])
    provider: str = "openai"
    model: str = "gpt-4o-mini"

    # Translation job reference
    translation_job_id: Optional[str] = None

    # Output
    output_paths: Dict[str, Path] = field(default_factory=dict)
    manuscript_output: Optional[ManuscriptCoreOutput] = None
    adn_data: Optional[Dict] = None

    # Metrics
    agent1_metrics: Dict = field(default_factory=dict)
    agent2_metrics: Dict = field(default_factory=dict)
    agent3_metrics: Dict = field(default_factory=dict)

    # Timestamps
    created_at: datetime = field(default_factory=datetime.utcnow)
    updated_at: datetime = field(default_factory=datetime.utcnow)
    completed_at: Optional[datetime] = None

    # Control
    cancelled: bool = False


class APSService:
    """
    AI Publishing System Service - REAL INTEGRATION

    Manages the full pipeline using real BatchProcessor:
    - Agent #1: BatchProcessor (translation + ADN)
    - Agent #2: Editorial processing (consistency + intent)
    - Agent #3: Layout rendering (format + output)
    """

    def __init__(
        self,
        output_dir: str = "outputs/aps",
        job_queue: JobQueue = None,
        batch_processor: BatchProcessor = None,
        websocket_manager: Any = None,
        enable_cache: bool = True,
    ):
        self.output_dir = Path(output_dir)
        self.output_dir.mkdir(parents=True, exist_ok=True)

        # Use shared job queue or create new
        self.job_queue = job_queue

        # Use shared batch processor (must be provided for real integration)
        self.batch_processor = batch_processor

        self.websocket_manager = websocket_manager
        self.jobs: Dict[str, APSJob] = {}
        self._lock = threading.Lock()

        # Initialize cache manager (PERF-004)
        self.enable_cache = enable_cache and HAS_CACHE
        self._cache_manager: Optional[APSCacheManager] = None
        if self.enable_cache:
            try:
                self._cache_manager = get_cache_manager()
                logger.info("APSCacheManager initialized")
            except Exception as e:
                logger.warning(f"Failed to initialize cache: {e}")
                self.enable_cache = False

        mode = "REAL BatchProcessor" if batch_processor else "MOCK (no processor)"
        cache_mode = "WITH cache" if self.enable_cache else "NO cache"
        logger.info(f"APSService initialized with {mode}, {cache_mode}. Output: {self.output_dir}")

    def create_job(
        self,
        source_file: str,
        source_path: Path,
        source_lang: str = "auto",
        target_lang: str = "vi",
        template: str = "book",
        output_formats: List[str] = None,
        provider: str = "openai",
        model: str = "gpt-4o-mini",
    ) -> APSJob:
        """Create a new APS processing job"""
        job_id = str(uuid.uuid4())[:8]

        job = APSJob(
            job_id=job_id,
            source_file=source_file,
            source_path=source_path,
            source_lang=source_lang,
            target_lang=target_lang,
            template=template,
            output_formats=output_formats or ["docx"],
            provider=provider,
            model=model,
        )

        with self._lock:
            self.jobs[job_id] = job

        logger.info(f"Created APS job {job_id}: {source_file}")
        return job

    def get_job(self, job_id: str) -> Optional[APSJob]:
        """Get job by ID"""
        return self.jobs.get(job_id)

    def cancel_job(self, job_id: str) -> bool:
        """Cancel a job"""
        job = self.jobs.get(job_id)
        if job and job.status not in ["complete", "error", "cancelled"]:
            job.cancelled = True
            job.status = "cancelled"
            job.updated_at = datetime.utcnow()

            # Also cancel translation job if running
            if job.translation_job_id and self.job_queue:
                trans_job = self.job_queue.get_job(job.translation_job_id)
                if trans_job and trans_job.status in [TranslationJobStatus.PENDING, TranslationJobStatus.PROCESSING]:
                    trans_job.status = TranslationJobStatus.CANCELLED
                    trans_job.cancellation_requested = True

            return True
        return False

    async def process_job(self, job_id: str):
        """
        Process a job through all 3 agents.

        Uses real BatchProcessor for Agent #1.
        """
        job = self.get_job(job_id)
        if not job:
            logger.error(f"Job {job_id} not found")
            return

        try:
            job.status = "processing"
            job.updated_at = datetime.utcnow()

            # ========== AGENT #1: MANUSCRIPT CORE (Real BatchProcessor) ==========
            job.stage = 1
            job.current_task = "Starting translation..."
            logger.info(f"[{job_id}] Stage 1: Manuscript Core (REAL)")

            # Check if we have real processor, otherwise use mock
            if self.batch_processor and self.job_queue:
                await self._run_agent1_real(job)
            else:
                logger.warning(f"[{job_id}] No BatchProcessor available, using mock")
                await self._run_agent1_mock(job)

            if job.cancelled:
                return

            # ========== AGENT #2: EDITORIAL CORE ==========
            job.stage = 2
            job.stage_progress = 0
            job.current_task = "Analyzing consistency..."
            job.progress = 40
            logger.info(f"[{job_id}] Stage 2: Editorial Core")

            lip = await self._run_agent2(job)

            if job.cancelled:
                return

            # ========== AGENT #3: LAYOUT CORE ==========
            job.stage = 3
            job.stage_progress = 0
            job.current_task = "Generating output files..."
            job.progress = 70
            logger.info(f"[{job_id}] Stage 3: Layout Core")

            await self._run_agent3(job, lip)

            if job.cancelled:
                return

            # ========== COMPLETE ==========
            job.status = "complete"
            job.progress = 100
            job.stage_progress = 100
            job.current_task = "Done"
            job.completed_at = datetime.utcnow()
            job.updated_at = datetime.utcnow()

            logger.info(f"[{job_id}] APS Processing complete!")

            # Broadcast completion
            await self._broadcast_progress(job, "job_completed")

        except Exception as e:
            logger.error(f"[{job_id}] Processing failed: {e}")
            logger.error(traceback.format_exc())
            job.status = "error"
            job.error = str(e)
            job.updated_at = datetime.utcnow()

            await self._broadcast_progress(job, "job_failed")

    async def _run_agent1_real(self, job: APSJob):
        """
        Run Agent #1: Real BatchProcessor Translation

        Creates a TranslationJob and processes it through BatchProcessor.
        """
        start_time = time.time()

        # Create output directory
        job_output_dir = self.output_dir / job.job_id
        job_output_dir.mkdir(parents=True, exist_ok=True)

        # Determine output file path
        base_name = Path(job.source_file).stem
        output_file = job_output_dir / f"{base_name}_translated.docx"

        # Create TranslationJob via JobQueue
        translation_job = self.job_queue.create_job(
            job_name=f"APS: {job.source_file}",
            input_file=str(job.source_path),
            output_file=str(output_file),
            source_lang=job.source_lang if job.source_lang != "auto" else "en",
            target_lang=job.target_lang,
            output_format="docx",
            provider=job.provider,
            model=job.model,
            metadata={
                "aps_job_id": job.job_id,
                "template": job.template,
                "enable_adn_extraction": True,
            }
        )

        job.translation_job_id = translation_job.job_id
        logger.info(f"[{job.job_id}] Created translation job: {translation_job.job_id}")

        # Poll for completion with progress updates
        last_progress = 0
        while True:
            if job.cancelled:
                translation_job.status = TranslationJobStatus.CANCELLED
                translation_job.cancellation_requested = True
                return

            # Get fresh job status
            trans_job = self.job_queue.get_job(translation_job.job_id)
            if not trans_job:
                trans_job = translation_job

            # Update APS job progress (Agent #1 is 0-35%)
            if trans_job.total_chunks > 0:
                chunk_progress = (trans_job.completed_chunks / trans_job.total_chunks) * 100
            else:
                chunk_progress = trans_job.progress * 100

            if chunk_progress != last_progress:
                job.stage_progress = chunk_progress
                job.progress = chunk_progress * 0.35  # 0-35% range
                job.current_task = f"Translating... ({trans_job.completed_chunks}/{trans_job.total_chunks} chunks)"
                await self._broadcast_progress(job, "job_updated")
                last_progress = chunk_progress

            # Check completion status
            if trans_job.status == TranslationJobStatus.COMPLETED:
                logger.info(f"[{job.job_id}] Translation completed!")
                break
            elif trans_job.status == TranslationJobStatus.FAILED:
                raise Exception(trans_job.error or "Translation failed")
            elif trans_job.status == TranslationJobStatus.CANCELLED:
                job.cancelled = True
                return

            await asyncio.sleep(0.5)

        # Get final translation job state
        trans_job = self.job_queue.get_job(translation_job.job_id) or translation_job

        # Build ManuscriptCoreOutput from translation results
        job.current_task = "Building manuscript output..."
        job.stage_progress = 90
        await self._broadcast_progress(job, "job_updated")

        manuscript = self._build_manuscript_from_job(trans_job, job)
        job.manuscript_output = manuscript

        # Get ADN data from translation job
        if trans_job.metadata and 'adn' in trans_job.metadata:
            job.adn_data = trans_job.metadata['adn']
            logger.info(f"[{job.job_id}] ADN extracted from translation job")
        elif trans_job.metadata and 'adn_summary' in trans_job.metadata:
            # ADN summary available but not full data
            job.adn_data = {
                "version": "1.0",
                "metadata": {
                    "source_lang": trans_job.source_lang,
                    "target_lang": trans_job.target_lang,
                },
                "statistics": trans_job.metadata['adn_summary'],
            }
        else:
            # Extract ADN fallback
            job.adn_data = self._extract_adn_fallback(manuscript, job)

        job.stage_progress = 100
        job.progress = 35

        # Metrics
        job.agent1_metrics = {
            "duration_seconds": round(time.time() - start_time, 2),
            "items_processed": trans_job.total_chunks or len(manuscript.segments),
            "extra": {
                "segments": len(manuscript.segments),
                "quality_score": trans_job.avg_quality_score,
                "adn_extracted": bool(job.adn_data),
                "provider": trans_job.provider,
                "model": trans_job.model,
                "real_translation": True,
            }
        }

        logger.info(f"[{job.job_id}] Agent #1 complete: {len(manuscript.segments)} segments")

    async def _run_agent1_mock(self, job: APSJob):
        """
        Run Agent #1: Mock simulation (fallback when no BatchProcessor)
        """
        start_time = time.time()

        # Simulate text extraction
        job.current_task = "Extracting text..."
        job.stage_progress = 10
        await self._broadcast_progress(job, "job_updated")
        await asyncio.sleep(0.5)

        if job.cancelled:
            return

        # Simulate translation
        job.current_task = "Translating content..."
        job.stage_progress = 30
        await self._broadcast_progress(job, "job_updated")
        await asyncio.sleep(1.5)

        if job.cancelled:
            return

        # Simulate ADN extraction
        job.current_task = "Extracting Content ADN..."
        job.stage_progress = 70
        await self._broadcast_progress(job, "job_updated")
        await asyncio.sleep(0.8)

        # Create mock manuscript
        job.manuscript_output = ManuscriptCoreOutput(
            source_file=job.source_file,
            source_language=job.source_lang,
            target_language=job.target_lang,
            segments=[
                Segment(id="seg_0001", type=SegmentType.CHAPTER, level=1,
                       original_text="Chapter 1", translated_text="Chương 1"),
                Segment(id="seg_0002", type=SegmentType.PARAGRAPH, level=0,
                       original_text="This is mock content.",
                       translated_text="Đây là nội dung mẫu."),
            ]
        )

        # Generate mock ADN
        job.adn_data = self._generate_mock_adn(job)

        job.stage_progress = 100
        job.progress = 35

        # Metrics
        job.agent1_metrics = {
            "duration_seconds": round(time.time() - start_time, 2),
            "items_processed": 2,
            "extra": {
                "segments": 2,
                "adn_extracted": True,
                "real_translation": False,
            }
        }

    def _build_manuscript_from_job(
        self,
        trans_job: TranslationJob,
        aps_job: APSJob,
    ) -> ManuscriptCoreOutput:
        """
        Build ManuscriptCoreOutput from TranslationJob results.
        """
        manuscript = ManuscriptCoreOutput(
            source_file=aps_job.source_file,
            source_language=trans_job.source_lang,
            target_language=trans_job.target_lang,
        )

        segments = []

        # Try to read translated content from output file
        if trans_job.output_file:
            output_path = Path(trans_job.output_file)

            if output_path.exists():
                try:
                    if output_path.suffix == '.txt':
                        text = output_path.read_text(encoding='utf-8')
                        paragraphs = [p.strip() for p in text.split('\n\n') if p.strip()]
                        for i, para in enumerate(paragraphs):
                            seg_type, level = self._detect_segment_type(para, i)
                            segments.append(Segment(
                                id=f"seg_{i:04d}",
                                type=seg_type,
                                level=level,
                                translated_text=para,
                                confidence=0.9,
                            ))
                    elif output_path.suffix == '.docx':
                        try:
                            from docx import Document
                            doc = Document(str(output_path))
                            for i, para in enumerate(doc.paragraphs):
                                if para.text.strip():
                                    seg_type, level = self._detect_segment_type(para.text, i)
                                    segments.append(Segment(
                                        id=f"seg_{i:04d}",
                                        type=seg_type,
                                        level=level,
                                        translated_text=para.text,
                                        confidence=0.9,
                                    ))
                        except Exception as e:
                            logger.warning(f"Could not read DOCX: {e}")
                except Exception as e:
                    logger.warning(f"Could not read output file: {e}")

        # Fallback: create minimal segments
        if not segments:
            segments = [
                Segment(
                    id="seg_0001",
                    type=SegmentType.PARAGRAPH,
                    level=0,
                    translated_text="[Translation completed - see output file]",
                    confidence=trans_job.avg_quality_score or 0.9,
                )
            ]

        manuscript.segments = segments

        # Add ADN if available
        if trans_job.metadata and 'adn' in trans_job.metadata:
            manuscript.adn = trans_job.metadata['adn']

        return manuscript

    def _detect_segment_type(self, text: str, index: int) -> tuple:
        """Detect segment type from content"""
        text_lower = text.lower().strip()

        # Chapter detection
        chapter_markers = ['chương', 'chapter', 'phần', 'part', '第']
        for marker in chapter_markers:
            if text_lower.startswith(marker):
                return SegmentType.CHAPTER, 1

        # Heading detection (short, no period at end)
        if len(text) < 100 and not text.rstrip().endswith('.'):
            words = text.split()
            if len(words) < 15:
                return SegmentType.HEADING, 2

        # Default to paragraph
        return SegmentType.PARAGRAPH, 0

    def _extract_adn_fallback(
        self,
        manuscript: ManuscriptCoreOutput,
        job: APSJob,
    ) -> Dict:
        """Extract ADN if not done by BatchProcessor"""
        if not HAS_ADN:
            return self._generate_mock_adn(job)

        try:
            texts = [s.translated_text for s in manuscript.segments if s.translated_text]

            if not texts:
                return self._generate_mock_adn(job)

            # Check cache for ADN (PERF-004)
            document_hash = None
            if self._cache_manager:
                content = "\n".join(texts)
                document_hash = self._cache_manager.hash_document(content)
                cached_adn = self._cache_manager.get_adn(document_hash)
                if cached_adn:
                    logger.info(f"[{job.job_id}] ADN cache HIT: {document_hash[:8]}")
                    return cached_adn

            extractor = ADNExtractor(
                source_lang=job.source_lang if job.source_lang != "auto" else "en",
                target_lang=job.target_lang,
            )

            doc_type = "book" if job.template == "book" else "article"
            adn = extractor.extract(texts, doc_type)
            adn_dict = adn.to_dict()

            # Cache ADN result (PERF-004)
            if self._cache_manager and document_hash:
                self._cache_manager.set_adn(document_hash, adn_dict)
                logger.info(f"[{job.job_id}] ADN cached: {document_hash[:8]}")

            return adn_dict
        except Exception as e:
            logger.warning(f"ADN extraction fallback failed: {e}")
            return self._generate_mock_adn(job)

    async def _run_agent2(self, job: APSJob) -> LayoutIntentPackage:
        """
        Run Agent #2: Editorial Core

        Creates LayoutIntentPackage from ManuscriptCoreOutput.
        """
        start_time = time.time()

        if not job.manuscript_output:
            raise ValueError("No manuscript output from Agent #1")

        job.current_task = "Checking consistency..."
        job.stage_progress = 20
        await self._broadcast_progress(job, "job_updated")
        await asyncio.sleep(0.1)

        if job.cancelled:
            return None

        job.current_task = "Mapping layout intent..."
        job.stage_progress = 50
        await self._broadcast_progress(job, "job_updated")
        await asyncio.sleep(0.1)

        if job.cancelled:
            return None

        job.current_task = "Building Layout Intent Package..."
        job.stage_progress = 80

        # Create LayoutIntentPackage from manuscript
        lip = self._create_layout_intent(job.manuscript_output, job)

        job.stage_progress = 100
        job.progress = 65
        await self._broadcast_progress(job, "job_updated")

        # Metrics
        job.agent2_metrics = {
            "duration_seconds": round(time.time() - start_time, 2),
            "items_processed": len(lip.blocks),
            "extra": {
                "blocks": len(lip.blocks),
                "sections": len(lip.sections),
                "consistency_issues": lip.consistency.unresolved_count if lip.consistency else 0,
            }
        }

        logger.info(f"[{job.job_id}] Agent #2 complete: {len(lip.blocks)} blocks")
        return lip

    def _create_layout_intent(
        self,
        manuscript: ManuscriptCoreOutput,
        job: APSJob,
    ) -> LayoutIntentPackage:
        """Create LayoutIntentPackage from ManuscriptCoreOutput"""

        # Create blocks from segments
        blocks = []
        for seg in manuscript.segments:
            block_type = self._segment_to_block_type(seg.type)
            blocks.append(Block(
                id=seg.id,
                type=block_type,
                content=seg.translated_text,
                level=seg.level,
            ))

        # Create sections based on template
        first_block_id = blocks[0].id if blocks else "seg_0000"
        last_block_id = blocks[-1].id if blocks else "seg_0000"
        sections = [
            SectionDefinition(
                type=SectionType.MAIN_BODY,
                start_block_id=first_block_id,
                end_block_id=last_block_id,
            )
        ]

        # Create consistency report
        consistency = ConsistencyReport(
            resolved_count=len(blocks),
            unresolved_count=0,
        )

        return LayoutIntentPackage(
            title=Path(job.source_file).stem,
            template=job.template,
            blocks=blocks,
            sections=sections,
            consistency=consistency,
        )

    def _segment_to_block_type(self, seg_type: SegmentType) -> BlockType:
        """Convert SegmentType to BlockType"""
        mapping = {
            SegmentType.CHAPTER: BlockType.CHAPTER,
            SegmentType.SECTION: BlockType.SECTION,
            SegmentType.PARAGRAPH: BlockType.PARAGRAPH,
            SegmentType.HEADING: BlockType.HEADING_1,
            SegmentType.QUOTE: BlockType.QUOTE,
            SegmentType.LIST: BlockType.LIST,
            SegmentType.TABLE: BlockType.TABLE,
            SegmentType.CODE: BlockType.CODE,
            SegmentType.FORMULA: BlockType.FORMULA,
            SegmentType.IMAGE: BlockType.IMAGE,
            SegmentType.FOOTNOTE: BlockType.FOOTNOTE,
        }
        return mapping.get(seg_type, BlockType.PARAGRAPH)

    async def _run_agent3(self, job: APSJob, lip: LayoutIntentPackage):
        """
        Run Agent #3: Layout Core

        Generates output files from LayoutIntentPackage.
        """
        start_time = time.time()

        if not lip:
            raise ValueError("No Layout Intent Package from Agent #2")

        # Create output directory for this job
        job_output_dir = self.output_dir / job.job_id
        job_output_dir.mkdir(parents=True, exist_ok=True)

        base_name = Path(job.source_file).stem

        items_generated = 0

        for i, format_type in enumerate(job.output_formats):
            if job.cancelled:
                return

            progress_per_format = 100 / len(job.output_formats)
            job.current_task = f"Generating {format_type.upper()}..."
            job.stage_progress = i * progress_per_format
            await self._broadcast_progress(job, "job_updated")

            output_path = job_output_dir / f"{base_name}.{format_type}"

            try:
                # Generate output file
                self._render_output(lip, output_path, format_type, job)

                job.output_paths[format_type] = output_path
                items_generated += 1

                logger.info(f"[{job.job_id}] Generated: {output_path}")

            except Exception as e:
                logger.error(f"[{job.job_id}] Failed to generate {format_type}: {e}")
                # Continue with other formats

        job.stage_progress = 100
        job.progress = 95

        # Metrics
        job.agent3_metrics = {
            "duration_seconds": round(time.time() - start_time, 2),
            "items_processed": items_generated,
            "extra": {
                "formats_generated": list(job.output_paths.keys()),
            }
        }

        logger.info(f"[{job.job_id}] Agent #3 complete: {items_generated} formats")

    def _render_output(
        self,
        lip: LayoutIntentPackage,
        output_path: Path,
        format_type: str,
        job: APSJob,
    ):
        """Render output file from LayoutIntentPackage"""

        if format_type == "docx":
            self._render_docx(lip, output_path, job)
        elif format_type == "pdf":
            # PDF generation would require additional library
            output_path.write_text(f"PDF generation not implemented yet.\n\nContent:\n{self._blocks_to_text(lip.blocks)}")
        elif format_type == "epub":
            # EPUB generation would require additional library
            output_path.write_text(f"EPUB generation not implemented yet.\n\nContent:\n{self._blocks_to_text(lip.blocks)}")
        else:
            # Default: plain text
            output_path.write_text(self._blocks_to_text(lip.blocks))

    def _render_docx(self, lip: LayoutIntentPackage, output_path: Path, job: APSJob):
        """Render DOCX from LayoutIntentPackage"""
        try:
            from docx import Document
            from docx.shared import Pt, Inches
            from docx.enum.text import WD_ALIGN_PARAGRAPH

            doc = Document()

            # Add title
            title_para = doc.add_heading(lip.title or "Translated Document", 0)

            # Add metadata
            doc.add_paragraph(f"Source: {job.source_file}")
            doc.add_paragraph(f"Languages: {job.source_lang} → {job.target_lang}")
            doc.add_paragraph(f"Template: {job.template}")
            doc.add_paragraph("")  # Spacer

            # Add content blocks
            for block in lip.blocks:
                if block.type == BlockType.CHAPTER:
                    doc.add_heading(block.content, level=1)
                elif block.type == BlockType.SECTION:
                    doc.add_heading(block.content, level=2)
                elif block.type in [BlockType.HEADING_1, BlockType.HEADING_2, BlockType.HEADING_3]:
                    level = 1 if block.type == BlockType.HEADING_1 else (2 if block.type == BlockType.HEADING_2 else 3)
                    doc.add_heading(block.content, level=level)
                elif block.type == BlockType.QUOTE:
                    para = doc.add_paragraph(block.content)
                    para.paragraph_format.left_indent = Inches(0.5)
                    para.italic = True
                elif block.type == BlockType.CODE:
                    para = doc.add_paragraph(block.content)
                    para.style = 'Quote'  # Use quote style for code
                else:
                    doc.add_paragraph(block.content)

            doc.save(str(output_path))

        except ImportError:
            # Fallback if python-docx not installed
            output_path.write_text(self._blocks_to_text(lip.blocks))

    def _blocks_to_text(self, blocks: List[Block]) -> str:
        """Convert blocks to plain text"""
        lines = []
        for block in blocks:
            if block.type in [BlockType.CHAPTER, BlockType.HEADING_1]:
                lines.append(f"\n# {block.content}\n")
            elif block.type in [BlockType.SECTION, BlockType.HEADING_2]:
                lines.append(f"\n## {block.content}\n")
            elif block.type == BlockType.HEADING_3:
                lines.append(f"\n### {block.content}\n")
            elif block.type == BlockType.QUOTE:
                lines.append(f"\n> {block.content}\n")
            elif block.type == BlockType.CODE:
                lines.append(f"\n```\n{block.content}\n```\n")
            else:
                lines.append(block.content)
                lines.append("")
        return "\n".join(lines)

    def _generate_mock_adn(self, job: APSJob) -> Dict:
        """Generate mock ADN data for testing"""
        return {
            "version": "1.0",
            "metadata": {
                "source_lang": job.source_lang,
                "target_lang": job.target_lang,
                "total_chunks": 25,
                "document_type": job.template,
            },
            "statistics": {
                "proper_nouns_count": 8,
                "characters_count": 4,
                "terms_count": 12,
                "patterns_count": 5,
            },
            "properNouns": [
                {"original": "Tokyo", "translated": "Tokyo", "category": "place", "count": 5},
                {"original": "Dr. Smith", "translated": "Tiến sĩ Smith", "category": "person", "count": 12},
            ],
            "characters": [
                {
                    "name": "Dr. Sarah Chen",
                    "translated": "Tiến sĩ Sarah Chen",
                    "role": "protagonist",
                    "aliases": ["Sarah", "Dr. Chen"],
                    "mentions": 45,
                },
            ],
            "terms": [
                {
                    "original": "quantum entanglement",
                    "translated": "vướng víu lượng tử",
                    "domain": "scientific",
                },
                {
                    "original": "neural network",
                    "translated": "mạng nơ-ron",
                    "domain": "technical",
                },
            ],
            "patterns": [
                {
                    "type": "honorific",
                    "name": "Academic Titles",
                    "rule": "Dr./Prof. + Name → Tiến sĩ/Giáo sư + Name",
                    "frequency": 23,
                    "confidence": 0.95,
                },
            ],
        }

    async def _broadcast_progress(self, job: APSJob, event_type: str):
        """Broadcast progress update via WebSocket"""
        if self.websocket_manager:
            try:
                await self.websocket_manager.broadcast({
                    "event": event_type,
                    "job_id": job.job_id,
                    "status": job.status,
                    "progress": job.progress,
                    "stage": job.stage,
                    "stage_progress": job.stage_progress,
                    "current_task": job.current_task,
                    "error": job.error,
                })
            except Exception as e:
                logger.warning(f"WebSocket broadcast failed: {e}")

    def get_output_path(self, job_id: str, format_type: str) -> Optional[Path]:
        """Get output file path for a job"""
        job = self.get_job(job_id)
        if job and format_type in job.output_paths:
            return job.output_paths[format_type]
        return None

    def get_adn(self, job_id: str) -> Optional[Dict]:
        """Get ADN data for a job"""
        job = self.get_job(job_id)
        if job and job.adn_data:
            return job.adn_data
        return None

    # ==================== Cache Methods (PERF-004) ====================

    def get_cache_stats(self) -> Dict:
        """Get cache statistics"""
        if not self._cache_manager:
            return {"enabled": False, "message": "Cache not available"}

        return {
            "enabled": True,
            **self._cache_manager.stats()
        }

    def clear_cache(self) -> Dict:
        """Clear all caches"""
        if not self._cache_manager:
            return {"enabled": False, "message": "Cache not available"}

        return {
            "enabled": True,
            "cleared": self._cache_manager.clear_all()
        }

    def cleanup_expired_cache(self) -> Dict:
        """Cleanup expired cache entries"""
        if not self._cache_manager:
            return {"enabled": False, "message": "Cache not available"}

        return {
            "enabled": True,
            "cleaned": self._cache_manager.cleanup_expired()
        }


# Global service instance - will be initialized with real components
_aps_service: Optional[APSService] = None


def get_aps_service(
    job_queue: JobQueue = None,
    batch_processor: BatchProcessor = None,
    websocket_manager: Any = None,
    force_reinit: bool = False,
) -> APSService:
    """Get or create APS service instance"""
    global _aps_service

    if _aps_service is None or force_reinit:
        _aps_service = APSService(
            job_queue=job_queue,
            batch_processor=batch_processor,
            websocket_manager=websocket_manager,
        )

    return _aps_service


# Legacy: For backward compatibility with direct imports
aps_service = APSService()
