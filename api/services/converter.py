"""
Document format conversion service.

Extracted from api/routes/job_outputs.py — pure conversion logic,
no FastAPI or HTTP concerns.
"""
import subprocess
from pathlib import Path

from config.logging_config import get_logger

logger = get_logger(__name__)


MEDIA_TYPES = {
    "docx": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    "pdf": "application/pdf",
    "txt": "text/plain",
    "html": "text/html",
    "md": "text/markdown",
    "srt": "text/plain",
}


def get_media_type(fmt: str) -> str:
    """Return the MIME type for a given format string."""
    return MEDIA_TYPES.get(fmt, "application/octet-stream")


def convert_to_markdown(source_path: Path, target_path: Path) -> Path:
    """Convert a DOCX file to Markdown.

    Args:
        source_path: Path to source .docx file.
        target_path: Path to write .md output.

    Returns:
        target_path on success.

    Raises:
        Exception: If conversion fails.
    """
    from docx import Document

    doc = Document(str(source_path))

    md_lines = []
    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            md_lines.append("")
            continue

        style_name = para.style.name.lower() if para.style else ""
        if "heading 1" in style_name or "title" in style_name:
            md_lines.append(f"# {text}")
        elif "heading 2" in style_name:
            md_lines.append(f"## {text}")
        elif "heading 3" in style_name:
            md_lines.append(f"### {text}")
        else:
            md_lines.append(text)

        md_lines.append("")  # blank line after each paragraph

    with open(target_path, "w", encoding="utf-8") as f:
        f.write("\n".join(md_lines))

    return target_path


def convert_to_txt(source_path: Path, target_path: Path) -> Path:
    """Convert a DOCX file to plain text.

    Args:
        source_path: Path to source .docx file.
        target_path: Path to write .txt output.

    Returns:
        target_path on success.

    Raises:
        Exception: If conversion fails.
    """
    from docx import Document

    doc = Document(str(source_path))
    text_lines = [para.text for para in doc.paragraphs]

    with open(target_path, "w", encoding="utf-8") as f:
        f.write("\n\n".join(text_lines))

    return target_path


def convert_to_pdf(source_path: Path, target_path: Path, output_dir: Path) -> Path:
    """Convert a DOCX file to PDF.

    Tries LibreOffice first, falls back to reportlab.

    Args:
        source_path: Path to source .docx file.
        target_path: Path to write .pdf output.
        output_dir: Directory for LibreOffice --outdir.

    Returns:
        target_path on success.

    Raises:
        Exception: If all conversion methods fail.
    """
    # Method 1: LibreOffice
    try:
        result = subprocess.run(
            ["soffice", "--headless", "--convert-to", "pdf",
             "--outdir", str(output_dir), str(source_path)],
            capture_output=True,
            timeout=60,
        )
        if result.returncode == 0 and target_path.exists():
            return target_path
    except Exception as e:
        logger.warning("LibreOffice PDF conversion failed: %s", e)

    # Method 2: reportlab
    from docx import Document
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
    from reportlab.pdfbase import pdfmetrics
    from reportlab.pdfbase.ttfonts import TTFont

    try:
        pdfmetrics.registerFont(
            TTFont("DejaVu", "/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf")
        )
        font_name = "DejaVu"
    except Exception as e:
        logger.debug("DejaVu font not available, using Helvetica: %s", e)
        font_name = "Helvetica"

    doc = Document(str(source_path))
    pdf = SimpleDocTemplate(
        str(target_path), pagesize=A4,
        rightMargin=72, leftMargin=72,
        topMargin=72, bottomMargin=72,
    )

    styles = getSampleStyleSheet()
    styles.add(ParagraphStyle(name="Vietnamese", fontName=font_name,
                              fontSize=11, leading=14))

    story = []
    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            text = text.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")
            story.append(Paragraph(text, styles["Vietnamese"]))
            story.append(Spacer(1, 6))

    pdf.build(story)
    return target_path


async def convert_document_format(
    source_path: Path, target_format: str,
    output_dir: Path, base_name: str,
) -> Path:
    """Convert a document to the requested format.

    This is the main entry point used by the route layer.

    Args:
        source_path: Path to original output file.
        target_format: One of 'md', 'txt', 'pdf'.
        output_dir: Directory for output file.
        base_name: Stem name for target file.

    Returns:
        Path to the converted file.

    Raises:
        ValueError: If target_format is unsupported.
        Exception: If conversion itself fails.
    """
    target_path = output_dir / f"{base_name}.{target_format}"

    if target_format == "md":
        return convert_to_markdown(source_path, target_path)
    elif target_format == "txt":
        return convert_to_txt(source_path, target_path)
    elif target_format == "pdf":
        return convert_to_pdf(source_path, target_path, output_dir)
    else:
        raise ValueError(f"Unsupported target format: {target_format}")
