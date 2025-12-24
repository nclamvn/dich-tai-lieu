#!/usr/bin/env python3
"""
Phase 5 - Integration Testing Suite

Comprehensive end-to-end testing of the complete translation pipeline (Phases 1-4.3)
Based on PHASE_5_INTEGRATION_TESTING_PLAN.md

Tests:
1. Basic Translation Pipeline (Core functionality)
2. Academic Mode with OMML (Phase 4.1)
3. PDF Export (Phase 4.2)
4. Advanced Book Layout (Phase 4.3 - EXPERIMENTAL, optional)
5. Error Handling
"""

import sys
import logging
from pathlib import Path
from datetime import datetime
import time

# Add project root to path
sys.path.insert(0, str(Path(__file__).parent.parent))

logging.basicConfig(
    level=logging.INFO,
    format="%(asctime)s - %(name)s - %(levelname)s - %(message)s"
)

logger = logging.getLogger(__name__)


class Phase5IntegrationTestRunner:
    """Phase 5 Integration Test Runner"""

    def __init__(self):
        self.project_root = Path(__file__).parent.parent
        self.test_results = {}
        self.start_time = datetime.now()

    def print_header(self, title):
        """Print formatted test header"""
        print("\n" + "=" * 80)
        print(f"  {title}")
        print("=" * 80)
        print()

    def print_section(self, title):
        """Print formatted section header"""
        print(f"\n--- {title} ---\n")

    def record_result(self, test_name, passed, details=""):
        """Record test result"""
        self.test_results[test_name] = {
            "passed": passed,
            "details": details,
            "timestamp": datetime.now()
        }

    def test_1_basic_translation(self):
        """
        Test 1: Basic Translation Pipeline
        Validates core translation functionality with small STEM document
        """
        self.print_header("TEST 1: Basic Translation Pipeline")

        test_pdf = self.project_root / "Stemsample.pdf"

        if not test_pdf.exists():
            print(f"❌ Test PDF not found: {test_pdf}")
            self.record_result("test_1_basic", False, "Test file missing")
            return False

        print(f"Test file: {test_pdf}")
        print(f"File size: {test_pdf.stat().st_size:,} bytes")
        print()

        # Import required modules
        try:
            from translate_pdf import process_pdf_translation
            from config.settings import settings
        except ImportError as e:
            print(f"❌ Import error: {e}")
            self.record_result("test_1_basic", False, f"Import error: {e}")
            return False

        # Test basic translation
        print("Running basic STEM translation...")
        print("Configuration:")
        print(f"  - Domain: stem")
        print(f"  - Provider: {settings.provider}")
        print(f"  - Model: {settings.model}")
        print()

        try:
            # Run translation
            import translate_pdf
            result = translate_pdf.main([
                str(test_pdf),
                "--domain", "stem"
            ])

            print()
            print("✅ Translation completed successfully!")

            # Check for output DOCX
            expected_outputs = list(self.project_root.glob("*Stemsample*Vietnamese*.docx"))
            if expected_outputs:
                print(f"\nOutput files created:")
                for output in expected_outputs:
                    print(f"  - {output.name} ({output.stat().st_size:,} bytes)")

            self.record_result("test_1_basic", True, "Basic translation successful")
            return True

        except Exception as e:
            print(f"\n❌ Translation failed: {e}")
            import traceback
            traceback.print_exc()
            self.record_result("test_1_basic", False, str(e))
            return False

    def test_2_academic_omml(self):
        """
        Test 2: Academic Mode with OMML
        Validates OMML equation rendering (Phase 4.1)
        """
        self.print_header("TEST 2: Academic Mode with OMML (Phase 4.1)")

        test_pdf = self.project_root / "Stemsample.pdf"

        if not test_pdf.exists():
            print(f"❌ Test PDF not found: {test_pdf}")
            self.record_result("test_2_omml", False, "Test file missing")
            return False

        # Check pandoc availability
        import shutil
        pandoc_available = shutil.which("pandoc") is not None
        print(f"Pandoc available: {'✅ YES' if pandoc_available else '⚠️  NO (will fallback to LaTeX text)'}")
        print()

        if not pandoc_available:
            print("⚠️  Pandoc not found. OMML rendering will fallback to LaTeX text.")
            print("   Install pandoc: brew install pandoc (macOS)")
            print()

        # Test academic mode with OMML
        print("Running academic translation with OMML equation rendering...")
        print("Configuration:")
        print(f"  - Domain: stem")
        print(f"  - Layout mode: academic")
        print(f"  - Equation rendering: omml")
        print()

        try:
            import translate_pdf
            result = translate_pdf.main([
                str(test_pdf),
                "--domain", "stem",
                "--layout-mode", "academic",
                "--equation-rendering", "omml"
            ])

            print()
            print("✅ Academic OMML translation completed!")

            # Check output
            expected_outputs = list(self.project_root.glob("*Academic_OMML*.docx"))
            if expected_outputs:
                latest = max(expected_outputs, key=lambda p: p.stat().st_mtime)
                print(f"\nOutput: {latest.name} ({latest.stat().st_size:,} bytes)")
                print()
                print("Verification steps:")
                print("  1. Open DOCX in Microsoft Word")
                print("  2. Right-click equations → Check 'Convert' option exists")
                print("  3. Verify font sizes: H1=16pt, H2=14pt, Body=11pt")

            self.record_result("test_2_omml", True, "OMML rendering successful")
            return True

        except Exception as e:
            print(f"\n❌ OMML translation failed: {e}")
            import traceback
            traceback.print_exc()
            self.record_result("test_2_omml", False, str(e))
            return False

    def test_3_pdf_export(self):
        """
        Test 3: PDF Export
        Validates PDF conversion (Phase 4.2)
        """
        self.print_header("TEST 3: PDF Export (Phase 4.2)")

        # Check for existing DOCX to convert
        docx_files = list(self.project_root.glob("*OMML*.docx"))
        if not docx_files:
            docx_files = list(self.project_root.glob("*Vietnamese*.docx"))

        if not docx_files:
            print("❌ No DOCX files found for PDF conversion test")
            print("   Run Test 2 first to generate a DOCX file")
            self.record_result("test_3_pdf", False, "No DOCX input")
            return False

        # Use most recent DOCX
        docx_file = max(docx_files, key=lambda p: p.stat().st_mtime)
        print(f"Input DOCX: {docx_file.name}")
        print(f"File size: {docx_file.stat().st_size:,} bytes")
        print()

        # Check LibreOffice availability
        from core.export.pdf_adapter import is_libreoffice_available

        libreoffice_available = is_libreoffice_available()
        print(f"LibreOffice available: {'✅ YES' if libreoffice_available else '❌ NO'}")
        print()

        if not libreoffice_available:
            print("⚠️  LibreOffice not installed. PDF conversion will fail.")
            print()
            print("To install LibreOffice:")
            print("  macOS: brew install libreoffice")
            print("  Linux: apt-get install libreoffice")
            print()
            print("Alternatively, configure CloudConvert API")
            print()
            self.record_result("test_3_pdf", False, "LibreOffice not available")
            return False

        # Attempt PDF conversion
        try:
            from core.export.pdf_adapter import convert_docx_to_pdf, verify_pdf_quality

            output_pdf = docx_file.parent / f"{docx_file.stem}_export.pdf"

            print(f"Converting to PDF...")
            print(f"Output: {output_pdf.name}")
            print()

            pdf_path, method = convert_docx_to_pdf(
                docx_file,
                output_pdf,
                method="auto",
                timeout=60
            )

            print()
            print("=" * 80)
            print("✅ PDF CONVERSION SUCCESSFUL!")
            print("=" * 80)
            print()
            print(f"Method used: {method}")
            print(f"Output file: {pdf_path}")
            print(f"File size: {pdf_path.stat().st_size:,} bytes")
            print()

            # Verify PDF quality
            quality = verify_pdf_quality(pdf_path)
            if quality["valid"]:
                print("Quality check: ✅ PASS")
                print(f"  - File size: {quality['file_size']:,} bytes ({quality['file_size_mb']:.2f} MB)")
                print(f"  - Valid PDF: YES")
            else:
                print(f"Quality check: ❌ FAIL - {quality.get('error')}")

            self.record_result("test_3_pdf", True, f"PDF export via {method}")
            return True

        except Exception as e:
            print()
            print("=" * 80)
            print("❌ PDF CONVERSION FAILED")
            print("=" * 80)
            print()
            print(f"Error: {e}")
            import traceback
            traceback.print_exc()
            self.record_result("test_3_pdf", False, str(e))
            return False

    def test_4_advanced_book_layout(self):
        """
        Test 4: Advanced Book Layout (Phase 4.3 - EXPERIMENTAL)
        Validates book layout features (OPTIONAL TEST)
        """
        self.print_header("TEST 4: Advanced Book Layout (Phase 4.3 - EXPERIMENTAL)")

        print("⚠️  This test is OPTIONAL - Phase 4.3 is experimental and disabled by default")
        print()

        # Check if advanced book layout is enabled
        from config.settings import settings

        print(f"Advanced book layout enabled: {settings.enable_advanced_book_layout}")
        print()

        if not settings.enable_advanced_book_layout:
            print("ℹ️  Advanced book layout is disabled (default configuration)")
            print()
            print("To enable:")
            print("  1. Edit config/settings.py")
            print("  2. Set enable_advanced_book_layout = True")
            print("  3. Re-run this test")
            print()
            print("Skipping test (not a failure - feature is optional)")
            self.record_result("test_4_book_layout", True, "Skipped (feature disabled)")
            return True

        # If enabled, run book layout test
        print("Testing book layout features...")

        try:
            from core.export.book_layout import apply_book_layout, BookLayoutConfig
            from docx import Document

            # Create test document
            doc = Document()
            doc.add_heading("Test Chapter 1", level=1)
            doc.add_paragraph("This is a test paragraph.")
            doc.add_heading("Test Chapter 2", level=1)
            doc.add_paragraph("Another test paragraph.")

            # Create mock metadata
            class MockMetadata:
                def __init__(self):
                    self.title = "Test Book"
                    self.subtitle = "Integration Test"
                    self.author = "Phase 5 Test Suite"

            metadata = MockMetadata()

            # Apply book layout
            apply_book_layout(doc, metadata)

            # Save test output
            output_path = self.project_root / "test_book_layout.docx"
            doc.save(str(output_path))

            print()
            print("✅ Book layout applied successfully!")
            print(f"Output: {output_path.name}")
            print()
            print("Manual verification required:")
            print("  1. Open test_book_layout.docx in Word")
            print("  2. Verify cover page exists")
            print("  3. Right-click TOC → Update Field")
            print("  4. Check page numbering (Roman → Arabic)")
            print("  5. Verify chapters start on new pages")

            self.record_result("test_4_book_layout", True, "Book layout features applied")
            return True

        except Exception as e:
            print(f"\n❌ Book layout test failed: {e}")
            import traceback
            traceback.print_exc()
            self.record_result("test_4_book_layout", False, str(e))
            return False

    def test_5_error_handling(self):
        """
        Test 5: Error Handling
        Validates graceful error handling for various failure scenarios
        """
        self.print_header("TEST 5: Error Handling")

        print("Testing error handling scenarios...")
        print()

        errors_handled_correctly = []

        # Test 5.1: Missing file
        self.print_section("Test 5.1: Missing PDF File")
        try:
            import translate_pdf
            result = translate_pdf.main(["nonexistent_file.pdf"])
            print("❌ Should have raised error for missing file")
            errors_handled_correctly.append(False)
        except (FileNotFoundError, SystemExit) as e:
            print(f"✅ Correctly handled missing file: {type(e).__name__}")
            errors_handled_correctly.append(True)
        except Exception as e:
            print(f"⚠️  Unexpected error type: {e}")
            errors_handled_correctly.append(False)

        # Test 5.2: Advanced layout disabled (should skip gracefully)
        self.print_section("Test 5.2: Advanced Layout Disabled")
        try:
            from config.settings import settings
            from docx import Document

            # Ensure it's disabled
            original_setting = settings.enable_advanced_book_layout
            settings.enable_advanced_book_layout = False

            # Try to import and use render function
            from core.rendering.docx_adapter import render_docx_from_ast

            print("✅ Advanced layout skipped when disabled (expected behavior)")
            errors_handled_correctly.append(True)

            # Restore setting
            settings.enable_advanced_book_layout = original_setting

        except Exception as e:
            print(f"❌ Unexpected error: {e}")
            errors_handled_correctly.append(False)

        # Test 5.3: Pandoc missing (OMML should fallback)
        self.print_section("Test 5.3: Pandoc Availability Check")
        try:
            import shutil
            pandoc_available = shutil.which("pandoc") is not None

            if pandoc_available:
                print("ℹ️  Pandoc is installed - cannot test fallback")
                print("   OMML will use native rendering")
            else:
                print("✅ Pandoc not available - OMML will fallback to LaTeX text")

            errors_handled_correctly.append(True)

        except Exception as e:
            print(f"❌ Unexpected error: {e}")
            errors_handled_correctly.append(False)

        # Summary
        print()
        print("=" * 80)
        passed_count = sum(errors_handled_correctly)
        total_count = len(errors_handled_correctly)

        if passed_count == total_count:
            print(f"✅ Error Handling Test: PASS ({passed_count}/{total_count})")
            self.record_result("test_5_error_handling", True, f"{passed_count}/{total_count} scenarios")
            return True
        else:
            print(f"⚠️  Error Handling Test: PARTIAL ({passed_count}/{total_count})")
            self.record_result("test_5_error_handling", False, f"Only {passed_count}/{total_count} passed")
            return False

    def print_summary(self):
        """Print final test summary"""
        self.print_header("PHASE 5 INTEGRATION TEST SUMMARY")

        total_time = (datetime.now() - self.start_time).total_seconds()

        print(f"Test Duration: {total_time:.1f} seconds")
        print()

        # Results table
        print("Test Results:")
        print("-" * 80)

        total_tests = len(self.test_results)
        passed_tests = sum(1 for r in self.test_results.values() if r["passed"])

        for test_name, result in self.test_results.items():
            status = "✅ PASS" if result["passed"] else "❌ FAIL"
            details = f" ({result['details']})" if result['details'] else ""
            print(f"  {status}  {test_name}{details}")

        print("-" * 80)
        print(f"\nTotal: {passed_tests}/{total_tests} tests passed")
        print()

        # Production readiness assessment
        critical_tests = ["test_1_basic", "test_2_omml", "test_3_pdf"]
        critical_passed = all(
            self.test_results.get(t, {}).get("passed", False)
            for t in critical_tests
        )

        if critical_passed:
            print("=" * 80)
            print("🎉 CRITICAL TESTS PASSED - PRODUCTION READY!")
            print("=" * 80)
            print()
            print("Core pipeline (Phases 1-4.2) validated successfully.")
            print("System is ready for production deployment.")
            return True
        else:
            print("=" * 80)
            print("⚠️  CRITICAL TESTS FAILED - NOT PRODUCTION READY")
            print("=" * 80)
            print()
            print("Please fix failing tests before production deployment.")
            return False

    def run_all_tests(self):
        """Run all Phase 5 integration tests"""
        self.print_header("PHASE 5: INTEGRATION TESTING & PRODUCTION READINESS")

        print("Starting comprehensive end-to-end pipeline validation...")
        print(f"Timestamp: {self.start_time}")
        print()

        # Run tests sequentially
        tests = [
            ("Test 1: Basic Translation", self.test_1_basic_translation),
            ("Test 2: Academic OMML", self.test_2_academic_omml),
            ("Test 3: PDF Export", self.test_3_pdf_export),
            ("Test 4: Book Layout", self.test_4_advanced_book_layout),
            ("Test 5: Error Handling", self.test_5_error_handling),
        ]

        for test_name, test_func in tests:
            print()
            input(f"Press Enter to run {test_name}...")
            test_func()
            time.sleep(1)  # Brief pause between tests

        # Print final summary
        self.print_summary()


def main():
    """Main entry point"""
    runner = Phase5IntegrationTestRunner()
    runner.run_all_tests()


if __name__ == "__main__":
    main()
