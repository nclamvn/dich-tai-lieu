#!/usr/bin/env python3
"""Render-coverage soak: the live renderer must not lose source content.

Successor of ``soak_ast_vs_engine.py``. That harness compared the legacy
engines against the AST stack to de-risk the stage-4 default flip; the AST
side won on every sample/format and stage 5 retired the engines, so the
comparison leg is gone. What remains — and what this guards — is the absolute
invariant: rendering a document must keep (essentially) every content word of
the source.

Coverage floors are pinned from the recorded parity results
(docs/SOAK_RENDER_COVERAGE.md): DOCX rendered at 1.00 on every sample; PDF at
0.97–1.00 (extraction noise). A drop below the floor exits non-zero.

Run:
    python scripts/soak_render_coverage.py --report soak_report.md
    python scripts/soak_render_coverage.py --keep ./soak_out

CI runs the same check as ``tests/eval/test_render_coverage.py``.
"""

from __future__ import annotations

import argparse
import asyncio
import re
import sys
import tempfile
import zipfile
from pathlib import Path

ROOT = Path(__file__).resolve().parents[1]
if str(ROOT) not in sys.path:
    sys.path.insert(0, str(ROOT))

# Coverage floors per format, pinned from the recorded results.
FLOORS = {"docx": 0.99, "pdf": 0.95}

SAMPLES: dict[str, str] = {
    "prose_inline": (
        "# Chương một\n\n"
        "Đây là đoạn văn có **chữ đậm**, *chữ nghiêng* và `mã nguồn` xen kẽ. "
        "Nội dung tiếng Việt đủ dấu để kiểm tra phông chữ.\n\n"
        "## Mục 1.1\n\n"
        "Đoạn thứ hai bình thường, không định dạng đặc biệt nào cả.\n"
    ),
    "lists": (
        "# Danh sách\n\n"
        "Một vài mục cần theo dõi:\n\n"
        "- Mục **quan trọng** đầu tiên\n"
        "- Mục *nhấn mạnh* thứ hai\n"
        "- Mục có `tham số` kỹ thuật\n\n"
        "Các bước có thứ tự:\n\n"
        "1. Chuẩn bị dữ liệu\n"
        "2. Chạy xử lý\n"
        "3. Kiểm tra kết quả\n"
    ),
    "table": (
        "# Bảng số liệu\n\n"
        "| Hạng mục | Giá trị | Ghi chú |\n"
        "| --- | --- | --- |\n"
        "| Doanh thu | 1200 | tăng trưởng |\n"
        "| Chi phí | 800 | ổn định |\n"
        "| Lợi nhuận | 400 | khả quan |\n"
    ),
    "mixed_book": (
        "# Lời nói đầu\n\n"
        "Cuốn sách này trình bày phương pháp làm việc với tài liệu đa định dạng.\n\n"
        "> Một câu trích dẫn có **ý nghĩa** quan trọng.\n\n"
        "# Chương 1: Khởi đầu\n\n"
        "Nội dung chương một với nhiều đoạn văn khác nhau để kiểm tra bố cục.\n\n"
        "## 1.1 Bối cảnh\n\n"
        "Phần bối cảnh mô tả vấn đề cần giải quyet và cách tiếp cận tổng thể.\n"
    ),
}

_WORD_RE = re.compile(r"\w+", re.UNICODE)
_MARKER_RE = re.compile(r"[*`_#>|>-]+")


def _tokens(text: str) -> list[str]:
    """Lowercased word tokens (Unicode-aware), for content comparison."""
    return [w.lower() for w in _WORD_RE.findall(text)]


def _source_tokens(markdown: str) -> set[str]:
    """Content tokens in the source, with Markdown markers stripped.

    Leading list markers (ordered ``1.`` / ``2)`` and unordered ``-``/``*``/``+``)
    are removed per line so bare list ordinals are NOT counted as content: the
    DOCX renderer emits native Word auto-numbered lists (the ordinal is Word's
    list numbering, not paragraph text). Mid-line digits (real data, e.g. table
    values) are preserved.
    """
    cleaned = []
    for line in markdown.splitlines():
        line = re.sub(r"^\s*(?:\d+[.)]|[-*+])\s+", "", line)  # leading list marker
        cleaned.append(_MARKER_RE.sub(" ", line))
    return set(_tokens("\n".join(cleaned)))


def _docx_text(path: Path) -> str:
    import docx as _docx

    doc = _docx.Document(str(path))
    parts = [p.text for p in doc.paragraphs]
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                parts.append(cell.text)
    return "\n".join(parts)


def _pdf_text(path: Path) -> str:
    import pypdf

    return "\n".join((pg.extract_text() or "") for pg in pypdf.PdfReader(str(path)).pages)


def _epub_text(path: Path) -> str:
    with zipfile.ZipFile(path) as z:
        raw = "".join(
            z.read(n).decode("utf-8", "ignore")
            for n in z.namelist()
            if n.endswith((".xhtml", ".html"))
        )
    return re.sub(r"<[^>]+>", " ", raw)


async def _render(converter, fmt: str, markdown: str, out: Path) -> None:
    if fmt == "docx":
        await converter.convert_markdown_to_docx_professional(markdown, out, title="Soak", author="Harness")
    elif fmt == "pdf":
        await converter.convert_markdown_to_pdf_professional(markdown, out, title="Soak", author="Harness")
    else:
        raise ValueError(fmt)


def _coverage(source: set[str], produced: set[str]) -> float:
    return 1.0 if not source else len(source & produced) / len(source)


async def _run(report_path: Path | None, keep_dir: Path | None) -> int:
    from core_v2.output_converter import OutputConverter

    out_root = keep_dir or Path(tempfile.mkdtemp(prefix="soak_render_"))
    out_root.mkdir(parents=True, exist_ok=True)
    converter = OutputConverter(temp_dir=out_root / "_tmp")

    extractors = {"docx": _docx_text, "pdf": _pdf_text}
    rows: list[dict] = []
    regressions: list[str] = []

    for name, md in SAMPLES.items():
        src = _source_tokens(md)
        for fmt, extract in extractors.items():
            out = out_root / f"{name}.{fmt}"
            await _render(converter, fmt, md, out)
            cov = _coverage(src, set(_tokens(extract(out))))
            rows.append({"sample": name, "fmt": fmt, "coverage": cov})
            if cov + 1e-9 < FLOORS[fmt]:
                regressions.append(
                    f"{name}/{fmt}: coverage {cov:.3f} < floor {FLOORS[fmt]:.2f}"
                )

    report = _format_report(rows, regressions)
    print(report)
    if report_path:
        report_path.write_text(report, encoding="utf-8")
    if not keep_dir:
        import shutil

        shutil.rmtree(out_root, ignore_errors=True)
    return 1 if regressions else 0


def _format_report(rows: list[dict], regressions: list[str]) -> str:
    lines = ["# Render-coverage soak", "", "| sample | fmt | coverage | floor |", "| --- | --- | --- | --- |"]
    for r in rows:
        lines.append(f"| {r['sample']} | {r['fmt']} | {r['coverage']:.3f} | {FLOORS[r['fmt']]:.2f} |")
    lines.append("")
    lines.append("REGRESSIONS:" if regressions else "No coverage regressions.")
    lines.extend(f"- {r}" for r in regressions)
    return "\n".join(lines)


def main() -> int:
    parser = argparse.ArgumentParser(description="Render-coverage soak (AST stack)")
    parser.add_argument("--report", type=Path, default=None, help="write the markdown report here")
    parser.add_argument("--keep", type=Path, default=None, help="keep rendered artifacts in this dir")
    args = parser.parse_args()
    return asyncio.run(_run(args.report, args.keep))


if __name__ == "__main__":
    raise SystemExit(main())
