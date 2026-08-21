#!/usr/bin/env python3
"""Soak harness: legacy engine vs AST output pipeline (no API key).

Renders sample Markdown documents through BOTH live professional converters —
``OUTPUT_PIPELINE=engine`` (legacy ``docx_engine`` / ``pdf_engine``) and
``OUTPUT_PIPELINE=ast`` (``DocumentAST`` + ``core/rendering``) — for DOCX and PDF,
then compares the *content* each output preserves against the source.

The question this answers before the Option-A stage-4 default flip: **does the AST
path drop any source content the engine keeps?** It reports, per sample/format, the
source-token coverage of each pipeline and the token deltas between them. The AST
path is expected to cover the source at least as well as the engine (it also adds
front-matter — cover page / TOC — which shows up as extra AST-only tokens, not a
loss).

Usage:
    python scripts/soak_ast_vs_engine.py [--report OUT.md] [--keep DIR]

Exit code is 0 unless a pipeline fails to render or the AST path covers the source
*worse* than the engine on any sample (a real regression signal).
"""

from __future__ import annotations

import argparse
import asyncio
import os
import re
import tempfile
import zipfile
from pathlib import Path

# --------------------------------------------------------------------------- #
# Sample corpus — varied structure, all offline (no translation / no key).
# --------------------------------------------------------------------------- #
SAMPLES: dict[str, str] = {
    "prose_inline": (
        "# Chương một\n\n"
        "Đây là đoạn văn có **chữ đậm**, *chữ nghiêng* và `mã nguồn` xen kẽ. "
        "Nội dung tiếng Việt đủ dấu để kiểm tra phông chữ.\n\n"
        "## Mục 1.1\n\n"
        "Đoạn thứ hai bình thường, không định dạng đặc biệt nào cả.\n"
    ),
    "lists": (
        "# Danh sách\n\n"
        "Một vài mục cần theo dõi:\n\n"
        "- Mục **quan trọng** đầu tiên\n"
        "- Mục *nhấn mạnh* thứ hai\n"
        "- Mục có `tham số` kỹ thuật\n\n"
        "Các bước có thứ tự:\n\n"
        "1. Chuẩn bị dữ liệu\n"
        "2. Chạy xử lý\n"
        "3. Kiểm tra kết quả\n"
    ),
    "table": (
        "# Bảng số liệu\n\n"
        "| Hạng mục | Giá trị | Ghi chú |\n"
        "| --- | --- | --- |\n"
        "| Doanh thu | 1200 | tăng trưởng |\n"
        "| Chi phí | 800 | ổn định |\n"
        "| Lợi nhuận | 400 | khả quan |\n"
    ),
    "mixed_book": (
        "# Lời nói đầu\n\n"
        "Cuốn sách này trình bày phương pháp làm việc với tài liệu đa định dạng.\n\n"
        "> Một câu trích dẫn có **ý nghĩa** quan trọng.\n\n"
        "# Chương 1: Khởi đầu\n\n"
        "Nội dung chương một với nhiều đoạn văn khác nhau để kiểm tra bố cục.\n\n"
        "## 1.1 Bối cảnh\n\n"
        "Phần bối cảnh mô tả vấn đề cần giải quyet và cách tiếp cận tổng thể.\n"
    ),
}

_WORD_RE = re.compile(r"\w+", re.UNICODE)
_MARKER_RE = re.compile(r"[*`_#>|>-]+")


def _tokens(text: str) -> list[str]:
    """Lowercased word tokens (Unicode-aware), for content comparison."""
    return [w.lower() for w in _WORD_RE.findall(text)]


def _source_tokens(markdown: str) -> set[str]:
    """Content tokens in the source, with Markdown markers stripped.

    Leading list markers (ordered ``1.`` / ``2)`` and unordered ``-``/``*``/``+``)
    are removed per line so bare list ordinals are NOT counted as content: the AST
    DOCX renderer emits native Word auto-numbered lists (the ordinal is Word's list
    numbering, not paragraph text), whereas the legacy engine embeds the digit
    literally. Mid-line digits (real data, e.g. table values) are preserved.
    """
    cleaned = []
    for line in markdown.splitlines():
        line = re.sub(r"^\s*(?:\d+[.)]|[-*+])\s+", "", line)  # leading list marker
        cleaned.append(_MARKER_RE.sub(" ", line))
    return set(_tokens("\n".join(cleaned)))


def _docx_text(path: Path) -> str:
    import docx as _docx

    doc = _docx.Document(str(path))
    parts = [p.text for p in doc.paragraphs]
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                parts.append(cell.text)
    return "\n".join(parts)


def _pdf_text(path: Path) -> str:
    import pypdf

    return "\n".join((pg.extract_text() or "") for pg in pypdf.PdfReader(str(path)).pages)


def _epub_text(path: Path) -> str:
    with zipfile.ZipFile(path) as z:
        raw = "".join(
            z.read(n).decode("utf-8", "ignore")
            for n in z.namelist()
            if n.endswith((".xhtml", ".html"))
        )
    return re.sub(r"<[^>]+>", " ", raw)


async def _render(converter, fmt: str, markdown: str, out: Path, pipeline: str) -> None:
    os.environ["OUTPUT_PIPELINE"] = pipeline
    if fmt == "docx":
        await converter.convert_markdown_to_docx_professional(markdown, out, title="Soak", author="Harness")
    elif fmt == "pdf":
        await converter.convert_markdown_to_pdf_professional(markdown, out, title="Soak", author="Harness")
    else:
        raise ValueError(fmt)


def _coverage(source: set[str], produced: set[str]) -> float:
    if not source:
        return 1.0
    return len(source & produced) / len(source)


async def _run(report_path: Path | None, keep_dir: Path | None) -> int:
    from core_v2.output_converter import OutputConverter

    workdir = keep_dir or Path(tempfile.mkdtemp(prefix="soak_"))
    workdir.mkdir(parents=True, exist_ok=True)
    converter = OutputConverter(temp_dir=workdir / "_tmp")

    rows: list[dict] = []
    regressions: list[str] = []

    for name, md in SAMPLES.items():
        src = _source_tokens(md)
        for fmt, extract in (("docx", _docx_text), ("pdf", _pdf_text)):
            row = {"sample": name, "format": fmt}
            try:
                eng = workdir / f"{name}.engine.{fmt}"
                ast = workdir / f"{name}.ast.{fmt}"
                await _render(converter, fmt, md, eng, "engine")
                await _render(converter, fmt, md, ast, "ast")
                eng_tok = set(_tokens(extract(eng)))
                ast_tok = set(_tokens(extract(ast)))
                row["engine_cov"] = _coverage(src, eng_tok)
                row["ast_cov"] = _coverage(src, ast_tok)
                # source content the engine kept but AST dropped == a real regression
                dropped = sorted((src & eng_tok) - ast_tok)
                row["ast_dropped_vs_engine"] = dropped
                row["ast_only_extra"] = len(ast_tok - eng_tok)
                if row["ast_cov"] + 1e-9 < row["engine_cov"]:
                    regressions.append(
                        f"{name}/{fmt}: AST coverage {row['ast_cov']:.2f} < engine {row['engine_cov']:.2f}"
                    )
                if dropped:
                    regressions.append(f"{name}/{fmt}: AST dropped source tokens kept by engine: {dropped}")
            except Exception as e:  # a pipeline failing to render is itself a finding
                row["error"] = f"{type(e).__name__}: {e}"
                regressions.append(f"{name}/{fmt}: render error — {row['error']}")
            rows.append(row)

    report = _format_report(rows, regressions)
    print(report)
    if report_path:
        report_path.write_text(report, encoding="utf-8")
        print(f"\n[soak] report written to {report_path}")
    if keep_dir:
        print(f"[soak] rendered artifacts kept in {workdir}")
    return 1 if regressions else 0


def _format_report(rows: list[dict], regressions: list[str]) -> str:
    lines = [
        "# Soak report — legacy engine vs AST pipeline",
        "",
        "Source-token **coverage** = fraction of the source's content words that "
        "appear in the rendered output (1.00 = nothing dropped). `ast_dropped` lists "
        "any source word the engine kept but the AST path lost — this must be empty. "
        "`ast_extra` counts AST-only tokens (mostly cover-page / TOC front matter).",
        "",
        "| sample | fmt | engine cov | ast cov | ast_dropped | ast_extra |",
        "| --- | --- | --- | --- | --- | --- |",
    ]
    for r in rows:
        if "error" in r:
            lines.append(f"| {r['sample']} | {r['format']} | — | — | ERROR | {r['error']} |")
            continue
        dropped = ", ".join(r["ast_dropped_vs_engine"]) or "—"
        lines.append(
            f"| {r['sample']} | {r['format']} | {r['engine_cov']:.2f} | {r['ast_cov']:.2f} "
            f"| {dropped} | {r['ast_only_extra']} |"
        )
    lines += ["", "## Verdict", ""]
    if regressions:
        lines.append("**REGRESSIONS FOUND** — the AST path is not yet at parity:")
        lines += [f"- {r}" for r in regressions]
    else:
        lines.append(
            "**No content regressions.** On every sample and format the AST pipeline "
            "covers the source content at least as well as the legacy engine, and drops "
            "no source token the engine keeps. (AST-only extras are front matter.)"
        )
    return "\n".join(lines)


def main() -> int:
    ap = argparse.ArgumentParser(description=__doc__)
    ap.add_argument("--report", type=Path, default=None, help="write the markdown report here")
    ap.add_argument("--keep", type=Path, default=None, help="keep rendered artifacts in this dir")
    args = ap.parse_args()
    return asyncio.run(_run(args.report, args.keep))


if __name__ == "__main__":
    raise SystemExit(main())
