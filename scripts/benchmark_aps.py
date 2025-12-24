#!/usr/bin/env python3
"""
APS Performance Benchmark

Tests the APS pipeline with various document sizes
and measures timing for each stage.

Components tested:
- ADN Extraction (Agent #1 subset)
- ManuscriptCoreOutput creation
- LayoutIntentPackage creation (Agent #2)
- DOCX rendering (Agent #3)
"""

import asyncio
import time
import json
import sys
import os
from pathlib import Path
from typing import Dict, List, Any
from dataclasses import dataclass, field

# Add project root to path
sys.path.insert(0, str(Path(__file__).parent.parent))

from core.contracts import ManuscriptCoreOutput, Segment, SegmentType
from core.contracts.layout_intent import (
    LayoutIntentPackage, Block, BlockType,
    SectionDefinition, SectionType, ConsistencyReport
)
from core.adn import ADNExtractor


@dataclass
class BenchmarkResult:
    """Results from a single benchmark run"""
    name: str
    input_size: int  # characters
    segment_count: int

    # Timing (seconds)
    adn_time: float = 0
    manuscript_time: float = 0
    lip_time: float = 0
    docx_time: float = 0
    total_time: float = 0

    # Memory (MB)
    peak_memory: float = 0

    # Output
    output_size: int = 0  # bytes
    adn_items: int = 0

    def to_dict(self) -> Dict:
        return {
            "name": self.name,
            "input_size": self.input_size,
            "segment_count": self.segment_count,
            "adn_time": round(self.adn_time, 4),
            "manuscript_time": round(self.manuscript_time, 4),
            "lip_time": round(self.lip_time, 4),
            "docx_time": round(self.docx_time, 4),
            "total_time": round(self.total_time, 4),
            "peak_memory_mb": round(self.peak_memory, 2),
            "output_size_bytes": self.output_size,
            "adn_items": self.adn_items,
            "throughput_chars_per_sec": round(self.input_size / self.total_time, 1) if self.total_time > 0 else 0,
        }


def generate_test_document(num_segments: int, chars_per_segment: int = 500) -> List[str]:
    """Generate test document with specified size"""
    segments = []

    # Sample content for realistic text
    sample_paragraphs = [
        "Dr. Smith and Professor Williams conducted extensive research at Harvard University on quantum mechanics and neural networks.",
        "The machine learning algorithms developed by Microsoft and Google showed promising results in natural language processing applications.",
        "According to the latest findings published in Nature, the breakthrough in artificial intelligence could revolutionize medical diagnosis.",
        "The quantum entanglement theory proposed by Einstein has been experimentally verified by teams at MIT and Stanford University.",
        "Our methodology involves analyzing large datasets using advanced statistical techniques and deep learning architectures.",
    ]

    for i in range(num_segments):
        if i % 15 == 0:
            # Chapter heading
            text = f"Chapter {i // 15 + 1}: Important Research Findings and Methodology"
        elif i % 8 == 0:
            # Section heading
            text = f"Section {i}.{i % 5 + 1}: Detailed Analysis and Results"
        else:
            # Regular paragraph - use sample content
            base_idx = i % len(sample_paragraphs)
            base_text = sample_paragraphs[base_idx]

            # Repeat to reach desired length
            repeats = max(1, chars_per_segment // len(base_text))
            text = " ".join([base_text] * repeats)[:chars_per_segment]

        segments.append(text)

    return segments


def get_memory_usage() -> float:
    """Get current memory usage in MB"""
    try:
        import psutil
        process = psutil.Process(os.getpid())
        return process.memory_info().rss / 1024 / 1024
    except ImportError:
        return 0


def benchmark_adn_extraction(segments: List[str]) -> tuple:
    """Benchmark ADN extraction"""
    start_time = time.time()
    start_memory = get_memory_usage()

    extractor = ADNExtractor(source_lang="en", target_lang="vi")
    adn = extractor.extract(segments, "book")

    end_time = time.time()
    peak_memory = get_memory_usage() - start_memory

    # Count ADN items
    adn_dict = adn.to_dict()
    adn_items = (
        len(adn_dict.get('proper_nouns', [])) +
        len(adn_dict.get('characters', [])) +
        len(adn_dict.get('terms', [])) +
        len(adn_dict.get('patterns', []))
    )

    return adn, end_time - start_time, max(0, peak_memory), adn_items


def benchmark_manuscript_creation(segments: List[str], adn) -> tuple:
    """Benchmark ManuscriptCoreOutput creation"""
    start_time = time.time()
    start_memory = get_memory_usage()

    manuscript = ManuscriptCoreOutput(
        source_file="benchmark_test.txt",
        source_language="en",
        target_language="vi",
    )

    segs = []
    for i, text in enumerate(segments):
        if text.startswith("Chapter"):
            seg_type = SegmentType.CHAPTER
            level = 1
        elif text.startswith("Section"):
            seg_type = SegmentType.HEADING
            level = 2
        else:
            seg_type = SegmentType.PARAGRAPH
            level = 0

        segs.append(Segment(
            id=f"seg_{i:04d}",
            type=seg_type,
            level=level,
            original_text=text,
            translated_text=text,  # Simulating translated content
            confidence=0.95,
        ))

    manuscript.segments = segs
    if adn:
        manuscript.adn = adn.to_dict()

    end_time = time.time()
    peak_memory = get_memory_usage() - start_memory

    return manuscript, end_time - start_time, max(0, peak_memory)


def benchmark_lip_creation(manuscript: ManuscriptCoreOutput) -> tuple:
    """Benchmark LayoutIntentPackage creation"""
    start_time = time.time()
    start_memory = get_memory_usage()

    # Create blocks from segments
    blocks = []
    for seg in manuscript.segments:
        # Map segment type to block type
        type_mapping = {
            SegmentType.CHAPTER: BlockType.CHAPTER,
            SegmentType.HEADING: BlockType.HEADING_1,
            SegmentType.PARAGRAPH: BlockType.PARAGRAPH,
        }
        block_type = type_mapping.get(seg.type, BlockType.PARAGRAPH)

        blocks.append(Block(
            id=seg.id,
            type=block_type,
            content=seg.translated_text,
            level=seg.level,
        ))

    # Create sections
    first_block_id = blocks[0].id if blocks else "seg_0000"
    last_block_id = blocks[-1].id if blocks else "seg_0000"
    sections = [
        SectionDefinition(
            type=SectionType.MAIN_BODY,
            start_block_id=first_block_id,
            end_block_id=last_block_id,
        )
    ]

    # Create consistency report
    consistency = ConsistencyReport(
        resolved_count=len(blocks),
        unresolved_count=0,
    )

    lip = LayoutIntentPackage(
        title="Benchmark Document",
        template="book",
        blocks=blocks,
        sections=sections,
        consistency=consistency,
    )

    end_time = time.time()
    peak_memory = get_memory_usage() - start_memory

    return lip, end_time - start_time, max(0, peak_memory)


def benchmark_docx_render(lip: LayoutIntentPackage, output_path: str) -> tuple:
    """Benchmark DOCX rendering"""
    start_time = time.time()
    start_memory = get_memory_usage()

    try:
        from docx import Document
        from docx.shared import Pt, Inches

        doc = Document()

        # Add title
        doc.add_heading(lip.title or "Document", 0)

        # Add content blocks
        for block in lip.blocks:
            if block.type == BlockType.CHAPTER:
                doc.add_heading(block.content, level=1)
            elif block.type in [BlockType.HEADING_1, BlockType.HEADING_2]:
                level = 1 if block.type == BlockType.HEADING_1 else 2
                doc.add_heading(block.content, level=level)
            elif block.type == BlockType.QUOTE:
                para = doc.add_paragraph(block.content)
                para.paragraph_format.left_indent = Inches(0.5)
            else:
                doc.add_paragraph(block.content)

        doc.save(str(output_path))

    except ImportError as e:
        # Fallback to plain text if docx not available
        with open(output_path, 'w') as f:
            for block in lip.blocks:
                f.write(block.content + "\n\n")

    end_time = time.time()
    peak_memory = get_memory_usage() - start_memory

    output_size = Path(output_path).stat().st_size if Path(output_path).exists() else 0

    return end_time - start_time, max(0, peak_memory), output_size


def run_benchmark(name: str, num_segments: int, chars_per_segment: int = 500) -> BenchmarkResult:
    """Run full benchmark"""
    print(f"\n{'='*60}")
    print(f"Benchmark: {name}")
    print(f"Segments: {num_segments}, Chars/segment: {chars_per_segment}")
    print(f"{'='*60}")

    # Generate test data
    segments = generate_test_document(num_segments, chars_per_segment)
    total_chars = sum(len(s) for s in segments)
    print(f"Total characters: {total_chars:,}")

    result = BenchmarkResult(
        name=name,
        input_size=total_chars,
        segment_count=num_segments,
    )

    # Benchmark ADN extraction
    print("\n[1/4] ADN Extraction...")
    adn, time_adn, mem_adn, adn_items = benchmark_adn_extraction(segments)
    result.adn_time = time_adn
    result.adn_items = adn_items
    result.peak_memory = max(result.peak_memory, mem_adn)
    print(f"  Time: {time_adn:.4f}s, Memory: {mem_adn:.1f}MB, Items: {adn_items}")

    # Benchmark Manuscript creation
    print("\n[2/4] Manuscript Creation...")
    manuscript, time_ms, mem_ms = benchmark_manuscript_creation(segments, adn)
    result.manuscript_time = time_ms
    result.peak_memory = max(result.peak_memory, mem_ms)
    print(f"  Time: {time_ms:.4f}s, Memory: {mem_ms:.1f}MB, Segments: {len(manuscript.segments)}")

    # Benchmark LIP creation
    print("\n[3/4] Layout Intent Package...")
    lip, time_lip, mem_lip = benchmark_lip_creation(manuscript)
    result.lip_time = time_lip
    result.peak_memory = max(result.peak_memory, mem_lip)
    print(f"  Time: {time_lip:.4f}s, Memory: {mem_lip:.1f}MB, Blocks: {len(lip.blocks)}")

    # Benchmark DOCX render
    output_path = f"/tmp/benchmark_{name.replace(' ', '_')}.docx"
    print("\n[4/4] DOCX Rendering...")
    time_docx, mem_docx, output_size = benchmark_docx_render(lip, output_path)
    result.docx_time = time_docx
    result.output_size = output_size
    result.peak_memory = max(result.peak_memory, mem_docx)
    print(f"  Time: {time_docx:.4f}s, Memory: {mem_docx:.1f}MB, Output: {output_size:,} bytes")

    # Total
    result.total_time = result.adn_time + result.manuscript_time + result.lip_time + result.docx_time

    print(f"\n{'─'*60}")
    print(f"TOTAL: {result.total_time:.4f}s")
    print(f"Throughput: {total_chars / result.total_time:,.0f} chars/sec")

    # Cleanup
    Path(output_path).unlink(missing_ok=True)

    return result


def run_all_benchmarks() -> List[BenchmarkResult]:
    """Run all benchmark scenarios"""
    print("\n" + "="*70)
    print("              APS PERFORMANCE BENCHMARK")
    print("="*70)

    scenarios = [
        ("Tiny (1 page)", 10, 300),
        ("Small (5 pages)", 50, 400),
        ("Medium (20 pages)", 200, 500),
        ("Large (50 pages)", 500, 500),
        ("XLarge (100 pages)", 1000, 500),
        ("Book (200 pages)", 2000, 500),
    ]

    results = []

    for name, segments, chars in scenarios:
        try:
            result = run_benchmark(name, segments, chars)
            results.append(result)
        except Exception as e:
            print(f"\n[ERROR] Benchmark '{name}' failed: {e}")
            import traceback
            traceback.print_exc()

    return results


def print_summary(results: List[BenchmarkResult]):
    """Print benchmark summary"""
    print("\n" + "="*100)
    print("                                    BENCHMARK SUMMARY")
    print("="*100)

    print(f"\n{'Scenario':<20} {'Segs':>8} {'Chars':>12} {'ADN':>10} {'Manuscript':>12} {'LIP':>10} {'DOCX':>10} {'Total':>10} {'Throughput':>14}")
    print("-"*100)

    for r in results:
        print(f"{r.name:<20} {r.segment_count:>8} {r.input_size:>12,} {r.adn_time:>9.3f}s {r.manuscript_time:>11.4f}s {r.lip_time:>9.4f}s {r.docx_time:>9.3f}s {r.total_time:>9.3f}s {r.input_size/r.total_time:>12,.0f}/s")

    print("-"*100)

    # Analysis
    print("\n" + "="*70)
    print("                        ANALYSIS")
    print("="*70)

    # Find bottleneck
    total_adn = sum(r.adn_time for r in results)
    total_ms = sum(r.manuscript_time for r in results)
    total_lip = sum(r.lip_time for r in results)
    total_docx = sum(r.docx_time for r in results)
    total_all = total_adn + total_ms + total_lip + total_docx

    print(f"\n  Time Distribution (across all benchmarks):")
    print(f"    ADN Extraction:       {total_adn:>8.3f}s ({total_adn/total_all*100:5.1f}%)")
    print(f"    Manuscript Creation:  {total_ms:>8.3f}s ({total_ms/total_all*100:5.1f}%)")
    print(f"    LIP Creation:         {total_lip:>8.3f}s ({total_lip/total_all*100:5.1f}%)")
    print(f"    DOCX Rendering:       {total_docx:>8.3f}s ({total_docx/total_all*100:5.1f}%)")

    # Identify bottleneck
    times = {
        "ADN Extraction": total_adn,
        "Manuscript Creation": total_ms,
        "LIP Creation": total_lip,
        "DOCX Rendering": total_docx,
    }
    bottleneck = max(times, key=times.get)

    print(f"\n  BOTTLENECK: {bottleneck} ({times[bottleneck]/total_all*100:.1f}% of total time)")

    # Scaling analysis
    if len(results) >= 2:
        small = results[1]  # Small
        large = results[-1]  # Largest

        scaling_factor = large.segment_count / small.segment_count
        time_factor = large.total_time / small.total_time

        print(f"\n  Scaling Analysis (Small → Book):")
        print(f"    Input scaling:  {scaling_factor:.1f}x segments")
        print(f"    Time scaling:   {time_factor:.1f}x")

        if time_factor > scaling_factor * 1.5:
            print(f"    [WARNING] Super-linear scaling detected! O(n^2) or worse behavior.")
            print(f"              Optimization needed for large documents.")
        elif time_factor < scaling_factor * 0.8:
            print(f"    [EXCELLENT] Sub-linear scaling - very efficient!")
        else:
            print(f"    [OK] Linear scaling - expected O(n) behavior.")

    # Memory analysis
    max_memory = max(r.peak_memory for r in results)
    if max_memory > 0:
        print(f"\n  Memory Usage:")
        print(f"    Peak memory: {max_memory:.1f} MB")
        if max_memory > 500:
            print(f"    [WARNING] High memory usage. Consider streaming for large files.")

    # Recommendations
    print(f"\n  RECOMMENDATIONS:")
    if bottleneck == "ADN Extraction":
        print("    - Consider caching ADN results")
        print("    - Optimize regex patterns in extractor")
        print("    - Add incremental ADN extraction")
    elif bottleneck == "DOCX Rendering":
        print("    - Use streaming document generation")
        print("    - Pre-compile document templates")
        print("    - Consider parallel block rendering")
    elif bottleneck == "Manuscript Creation":
        print("    - Use lazy segment loading")
        print("    - Optimize segment type detection")

    print("\n" + "="*70)


def save_results(results: List[BenchmarkResult], output_file: str = "benchmark_results.json"):
    """Save results to JSON file"""
    data = {
        "timestamp": time.strftime("%Y-%m-%d %H:%M:%S"),
        "python_version": sys.version,
        "results": [r.to_dict() for r in results],
        "summary": {
            "total_benchmarks": len(results),
            "total_time": sum(r.total_time for r in results),
            "avg_throughput": sum(r.input_size/r.total_time for r in results) / len(results) if results else 0,
        }
    }

    with open(output_file, 'w') as f:
        json.dump(data, f, indent=2)

    print(f"\n  Results saved to: {output_file}")


if __name__ == "__main__":
    print("\nStarting APS Performance Benchmark...")
    print("This will test the pipeline with various document sizes.\n")

    results = run_all_benchmarks()
    print_summary(results)
    save_results(results, "/tmp/aps_benchmark_results.json")

    print("\nBenchmark complete!")
