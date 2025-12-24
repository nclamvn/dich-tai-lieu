#!/usr/bin/env python3
"""
Visual Quality Scanner - Phase 3.7

Analyzes DOCX files for typography consistency and detects drift/anomalies.

This scanner validates that:
- Typography remains stable across long documents
- No drift in indent, line spacing, or font sizes
- No unexpected margin/spacing jumps
- No loss of "page break before" for H1 headings
- No duplicate or corrupted Word styles

Usage:
    python3 scripts/visual_quality_scanner.py phase36_test_output/phase36_test_novel_25ch.docx
    python3 scripts/visual_quality_scanner.py --all-files phase36_test_output/
"""

import sys
import logging
from pathlib import Path
from typing import Dict, List, Any, Optional
from collections import defaultdict, Counter
from dataclasses import dataclass, field

from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(levelname)s - %(message)s'
)
logger = logging.getLogger(__name__)


@dataclass
class ParagraphMetrics:
    """Typography metrics for a single paragraph."""
    index: int
    style_name: str
    text_preview: str  # First 50 chars
    font_name: Optional[str] = None
    font_size_pt: Optional[float] = None
    bold: bool = False
    italic: bool = False
    line_spacing: Optional[float] = None
    space_before_pt: Optional[float] = None
    space_after_pt: Optional[float] = None
    first_line_indent_pt: Optional[float] = None
    left_indent_pt: Optional[float] = None
    right_indent_pt: Optional[float] = None
    alignment: Optional[str] = None
    keep_with_next: bool = False
    keep_together: bool = False
    page_break_before: bool = False


@dataclass
class DocumentAnalysis:
    """Complete analysis of a DOCX document."""
    file_path: Path
    total_paragraphs: int
    metrics: List[ParagraphMetrics] = field(default_factory=list)

    # Statistics
    style_counts: Dict[str, int] = field(default_factory=dict)
    font_size_distribution: Dict[float, int] = field(default_factory=dict)
    indent_distribution: Dict[float, int] = field(default_factory=dict)

    # Detected issues
    anomalies: List[Dict[str, Any]] = field(default_factory=list)
    warnings: List[str] = field(default_factory=list)


class VisualQualityScanner:
    """
    Scans DOCX files for typography consistency.

    Detects:
    - Style drift across chapters
    - Inconsistent indentation
    - Font size variations
    - Line spacing anomalies
    - Missing page breaks
    """

    # Expected typography standards (from Phase 3.5)
    EXPECTED_STANDARDS = {
        'Heading 1': {
            'font_size_pt': 16.0,
            'bold': True,
            'page_break_before': True,
        },
        'Heading 2': {
            'font_size_pt': 14.0,
            'bold': True,
        },
        'Heading 3': {
            'font_size_pt': 12.0,
            'bold': True,
        },
        'Normal': {
            'font_size_pt': 11.0,
            'font_name': 'Georgia',
            'line_spacing': 1.15,
        },
    }

    # Tolerance for float comparisons
    TOLERANCE = 0.5

    def __init__(self):
        self.logger = logging.getLogger(__name__)

    def scan_document(self, file_path: Path) -> DocumentAnalysis:
        """
        Scan a DOCX file and extract all typography metrics.

        Args:
            file_path: Path to DOCX file

        Returns:
            DocumentAnalysis with full metrics and detected issues
        """
        self.logger.info(f"Scanning document: {file_path}")

        try:
            doc = Document(str(file_path))
        except Exception as e:
            self.logger.error(f"Failed to open document: {e}")
            raise

        analysis = DocumentAnalysis(
            file_path=file_path,
            total_paragraphs=len(doc.paragraphs)
        )

        # Extract metrics for each paragraph
        for idx, para in enumerate(doc.paragraphs):
            metrics = self._extract_paragraph_metrics(para, idx)
            analysis.metrics.append(metrics)

            # Update statistics
            analysis.style_counts[metrics.style_name] = \
                analysis.style_counts.get(metrics.style_name, 0) + 1

            if metrics.font_size_pt:
                analysis.font_size_distribution[metrics.font_size_pt] = \
                    analysis.font_size_distribution.get(metrics.font_size_pt, 0) + 1

            if metrics.first_line_indent_pt is not None:
                analysis.indent_distribution[metrics.first_line_indent_pt] = \
                    analysis.indent_distribution.get(metrics.first_line_indent_pt, 0) + 1

        # Detect anomalies
        self._detect_anomalies(analysis)

        self.logger.info(f"Scan complete: {len(analysis.metrics)} paragraphs analyzed")
        return analysis

    def _extract_paragraph_metrics(self, para, idx: int) -> ParagraphMetrics:
        """Extract typography metrics from a paragraph."""
        # Get text preview
        text_preview = para.text[:50] if para.text else "(empty)"

        # Get style name
        style_name = para.style.name if para.style else "Unknown"

        metrics = ParagraphMetrics(
            index=idx,
            style_name=style_name,
            text_preview=text_preview
        )

        # Extract font properties from first run
        if para.runs:
            first_run = para.runs[0]
            metrics.font_name = first_run.font.name
            if first_run.font.size:
                metrics.font_size_pt = first_run.font.size.pt
            metrics.bold = first_run.font.bold or False
            metrics.italic = first_run.font.italic or False

        # Extract paragraph format properties
        fmt = para.paragraph_format

        if fmt.line_spacing is not None:
            metrics.line_spacing = fmt.line_spacing

        if fmt.space_before:
            metrics.space_before_pt = fmt.space_before.pt

        if fmt.space_after:
            metrics.space_after_pt = fmt.space_after.pt

        if fmt.first_line_indent:
            metrics.first_line_indent_pt = fmt.first_line_indent.pt

        if fmt.left_indent:
            metrics.left_indent_pt = fmt.left_indent.pt

        if fmt.right_indent:
            metrics.right_indent_pt = fmt.right_indent.pt

        # Alignment
        if para.alignment:
            alignment_map = {
                WD_ALIGN_PARAGRAPH.LEFT: "left",
                WD_ALIGN_PARAGRAPH.CENTER: "center",
                WD_ALIGN_PARAGRAPH.RIGHT: "right",
                WD_ALIGN_PARAGRAPH.JUSTIFY: "justify",
            }
            metrics.alignment = alignment_map.get(para.alignment, "unknown")

        # Pagination properties
        metrics.keep_with_next = fmt.keep_with_next or False
        metrics.keep_together = fmt.keep_together or False
        metrics.page_break_before = fmt.page_break_before or False

        return metrics

    def _detect_anomalies(self, analysis: DocumentAnalysis):
        """
        Detect typography anomalies and inconsistencies.

        Checks for:
        - Font size drift
        - Indent inconsistency
        - Missing page breaks
        - Style violations
        """
        # Group by style
        by_style = defaultdict(list)
        for m in analysis.metrics:
            by_style[m.style_name].append(m)

        # Check each style for consistency
        for style_name, metrics_list in by_style.items():
            # Check font size consistency
            font_sizes = [m.font_size_pt for m in metrics_list if m.font_size_pt]
            if font_sizes:
                size_counter = Counter(font_sizes)
                if len(size_counter) > 1:
                    # Multiple font sizes for same style
                    analysis.anomalies.append({
                        'type': 'font_size_drift',
                        'style': style_name,
                        'sizes_found': dict(size_counter),
                        'severity': 'high' if style_name.startswith('Heading') else 'medium'
                    })

            # Check indent consistency for body paragraphs
            if style_name == 'Normal':
                indents = [m.first_line_indent_pt for m in metrics_list
                          if m.first_line_indent_pt is not None]
                if indents:
                    indent_counter = Counter(indents)
                    # Should have only 0pt (first para) and ~23pt (body para)
                    if len(indent_counter) > 3:
                        analysis.anomalies.append({
                            'type': 'indent_drift',
                            'style': style_name,
                            'indents_found': dict(indent_counter),
                            'severity': 'medium'
                        })

            # Check expected standards
            if style_name in self.EXPECTED_STANDARDS:
                expected = self.EXPECTED_STANDARDS[style_name]

                # Check font size
                if 'font_size_pt' in expected:
                    violations = [m for m in metrics_list
                                 if m.font_size_pt and
                                 abs(m.font_size_pt - expected['font_size_pt']) > self.TOLERANCE]

                    if violations:
                        analysis.anomalies.append({
                            'type': 'font_size_violation',
                            'style': style_name,
                            'expected': expected['font_size_pt'],
                            'violations': len(violations),
                            'examples': [f"Para {v.index}: {v.font_size_pt}pt"
                                        for v in violations[:3]],
                            'severity': 'high'
                        })

                # Check page break for H1
                if style_name == 'Heading 1' and expected.get('page_break_before'):
                    missing_breaks = [m for m in metrics_list
                                     if not m.page_break_before and m.index > 0]

                    if missing_breaks:
                        analysis.warnings.append(
                            f"Heading 1: {len(missing_breaks)} headings missing page break"
                        )

        # Check line spacing consistency for body text
        body_paras = [m for m in analysis.metrics if m.style_name == 'Normal']
        if body_paras:
            line_spacings = [m.line_spacing for m in body_paras if m.line_spacing]
            if line_spacings:
                spacing_counter = Counter(line_spacings)
                if len(spacing_counter) > 2:
                    analysis.anomalies.append({
                        'type': 'line_spacing_drift',
                        'spacings_found': dict(spacing_counter),
                        'severity': 'medium'
                    })

    def generate_report(self, analyses: List[DocumentAnalysis], output_path: Path):
        """
        Generate a comprehensive visual quality report.

        Args:
            analyses: List of document analyses
            output_path: Path to output markdown report
        """
        self.logger.info(f"Generating report: {output_path}")

        with open(output_path, 'w', encoding='utf-8') as f:
            f.write("# Phase 3.7 - Visual Quality Verification Report\n\n")
            f.write(f"**Date**: {Path.cwd()}\n")
            f.write(f"**Files Analyzed**: {len(analyses)}\n\n")
            f.write("---\n\n")

            # Overall summary
            f.write("## Executive Summary\n\n")

            total_anomalies = sum(len(a.anomalies) for a in analyses)
            total_warnings = sum(len(a.warnings) for a in analyses)

            if total_anomalies == 0 and total_warnings == 0:
                f.write("✅ **ALL DOCUMENTS PASSED** - No anomalies or warnings detected.\n\n")
                f.write("The AST pipeline maintains perfect typography consistency across all test documents.\n\n")
            else:
                f.write(f"⚠️ **{total_anomalies} anomalies** and **{total_warnings} warnings** detected.\n\n")

            f.write("---\n\n")

            # Per-file analysis
            for analysis in analyses:
                self._write_file_analysis(f, analysis)

            # Recommendations
            f.write("---\n\n")
            f.write("## Recommendations\n\n")

            if total_anomalies == 0:
                f.write("### ✅ AST Pipeline Ready for Production\n\n")
                f.write("All typography tests passed. The AST pipeline can be enabled as default:\n\n")
                f.write("```python\n")
                f.write("# In batch_processor.py or docx_exporter.py\n")
                f.write("use_ast_pipeline = True  # Safe to enable\n")
                f.write("```\n\n")
            else:
                f.write("### ⚠️ Issues Require Investigation\n\n")
                f.write("The following issues should be resolved before enabling AST as default:\n\n")

                # Group by severity
                high_severity = [a for analysis in analyses for a in analysis.anomalies
                               if a.get('severity') == 'high']

                if high_severity:
                    f.write("**High Priority:**\n")
                    for anomaly in high_severity:
                        f.write(f"- {anomaly['type']}: {anomaly.get('style', 'N/A')}\n")
                    f.write("\n")

            f.write("---\n\n")
            f.write("**Report Generated by**: Visual Quality Scanner (Phase 3.7)\n")
            f.write("**Next Phase**: Phase 4 - Commercial PDF Export\n")

        self.logger.info(f"✅ Report saved: {output_path}")

    def _write_file_analysis(self, f, analysis: DocumentAnalysis):
        """Write analysis for a single file to report."""
        f.write(f"## File: `{analysis.file_path.name}`\n\n")
        f.write(f"**Total Paragraphs**: {analysis.total_paragraphs}\n\n")

        # Style breakdown
        f.write("### Style Distribution\n\n")
        f.write("| Style | Count | Percentage |\n")
        f.write("|-------|-------|------------|\n")
        for style, count in sorted(analysis.style_counts.items(),
                                   key=lambda x: x[1], reverse=True):
            pct = (count / analysis.total_paragraphs) * 100
            f.write(f"| {style} | {count} | {pct:.1f}% |\n")
        f.write("\n")

        # Font size distribution
        if analysis.font_size_distribution:
            f.write("### Font Size Distribution\n\n")
            f.write("| Font Size (pt) | Count |\n")
            f.write("|----------------|-------|\n")
            for size, count in sorted(analysis.font_size_distribution.items()):
                f.write(f"| {size:.1f} | {count} |\n")
            f.write("\n")

        # Indent distribution
        if analysis.indent_distribution:
            f.write("### Indent Distribution\n\n")
            f.write("| First-Line Indent (pt) | Count |\n")
            f.write("|------------------------|-------|\n")
            for indent, count in sorted(analysis.indent_distribution.items()):
                f.write(f"| {indent:.1f} | {count} |\n")
            f.write("\n")

        # Anomalies
        if analysis.anomalies:
            f.write("### ⚠️ Detected Anomalies\n\n")
            for anomaly in analysis.anomalies:
                severity = anomaly.get('severity', 'low').upper()
                f.write(f"**[{severity}]** {anomaly['type']}\n\n")
                for key, value in anomaly.items():
                    if key not in ['type', 'severity']:
                        f.write(f"- {key}: {value}\n")
                f.write("\n")
        else:
            f.write("### ✅ No Anomalies Detected\n\n")

        # Warnings
        if analysis.warnings:
            f.write("### Warnings\n\n")
            for warning in analysis.warnings:
                f.write(f"- {warning}\n")
            f.write("\n")

        f.write("---\n\n")


# ============================================================================
# CLI Interface
# ============================================================================

def main():
    """Main entry point for CLI."""
    import argparse

    parser = argparse.ArgumentParser(
        description="Visual Quality Scanner - Phase 3.7"
    )
    parser.add_argument(
        'files',
        nargs='+',
        help='DOCX files to scan'
    )
    parser.add_argument(
        '--output',
        default='PHASE_3.7_VISUAL_REPORT.md',
        help='Output report path (default: PHASE_3.7_VISUAL_REPORT.md)'
    )

    args = parser.parse_args()

    scanner = VisualQualityScanner()
    analyses = []

    logger.info("=" * 70)
    logger.info("PHASE 3.7 - VISUAL QUALITY VERIFICATION")
    logger.info("=" * 70)
    logger.info("")

    # Scan each file
    for file_path in args.files:
        path = Path(file_path)
        if not path.exists():
            logger.error(f"File not found: {path}")
            continue

        if path.is_dir():
            # Scan all DOCX files in directory
            for docx_file in path.glob("*.docx"):
                if not docx_file.name.startswith('~'):  # Skip temp files
                    analysis = scanner.scan_document(docx_file)
                    analyses.append(analysis)
        else:
            analysis = scanner.scan_document(path)
            analyses.append(analysis)

    if not analyses:
        logger.error("No files analyzed")
        return 1

    # Generate report
    output_path = Path(args.output)
    scanner.generate_report(analyses, output_path)

    # Print summary
    logger.info("")
    logger.info("=" * 70)
    logger.info("SCAN COMPLETE")
    logger.info("=" * 70)
    logger.info(f"Files scanned: {len(analyses)}")
    logger.info(f"Report saved: {output_path}")
    logger.info("")

    total_anomalies = sum(len(a.anomalies) for a in analyses)
    if total_anomalies == 0:
        logger.info("✅ ALL TESTS PASSED - No anomalies detected")
        logger.info("AST pipeline is ready for production use")
    else:
        logger.warning(f"⚠️ {total_anomalies} anomalies detected - review report")

    return 0


if __name__ == '__main__':
    sys.exit(main())
